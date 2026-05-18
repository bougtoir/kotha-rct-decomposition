#!/usr/bin/env python3
"""Create cover letter for Lancet Regional Health – Western Pacific submission."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
from datetime import date

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
LRH_DIR = os.path.join(OUTPUT_DIR, 'lancet_rh_wp')
os.makedirs(LRH_DIR, exist_ok=True)

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

unadj_d = reg["model1_unadjusted"]["cohens_d"]
attenuation = (1 - reg["adjusted_cpsp_test"]["cohens_d"] / unadj_d) * 100

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

today = date.today().strftime('%B %d, %Y')

# Date
p = doc.add_paragraph(today)
doc.add_paragraph()

# Addressee
doc.add_paragraph('Professor Chi Chiu Leung')
doc.add_paragraph('Editor-in-Chief')
doc.add_paragraph('The Lancet Regional Health \u2013 Western Pacific')
doc.add_paragraph()

# Salutation
doc.add_paragraph('Dear Professor Leung,')
doc.add_paragraph()

# Body
body_paras = [
    (
        'I am pleased to submit the enclosed manuscript entitled '
        '\u201cDo Cultural Labels Predict Analgesic Need? '
        'A 47-Prefecture Ecological Study of Pain Prescribing '
        'Variation in Japan\u201d '
        'for consideration as an Original Article in '
        'The Lancet Regional Health \u2013 Western Pacific.'
    ),
    (
        'Clinicians frequently rely on cultural labels\u2014such as \u201cJapanese patients '
        'are stoic\u201d\u2014when making analgesic prescribing decisions. '
        'Yet all prior comparisons have treated Japan as a culturally homogeneous unit. '
        'Using population-complete insurance claims data from Japan\u2019s National Database '
        '(NDB Open Data, 10th edition, covering approximately 125 million insured individuals), '
        'we mapped perioperative and chronic pain-related prescribing across all 47 prefectures '
        'for the first time to test whether cultural labels predict analgesic use.'
    ),
    (
        'Our principal findings are:'
    ),
]

for text in body_paras:
    doc.add_paragraph(text)

findings = [
    f'Acute perioperative analgesic prescribing varied 1\u00b797-fold across prefectures '
    f'(Kruskal\u2013Wallis P<0\u00b7001), with significant regional clustering.',

    f'Tohoku\u2014traditionally considered Japan\u2019s most stoic region\u2014prescribed '
    f'more, not fewer, analgesics (Cohen\u2019s d=0\u00b787), contradicting the stereotype '
    f'of universal pain endurance.',

    f'The apparent Tohoku excess in neuropathic pain prescribing was largely explained by '
    f'confounding disease prevalence (especially diabetes; r=0\u00b787), with '
    f'{attenuation:.0f}% attenuation after adjustment.',

    f'These findings demonstrate that cultural labels do not predict analgesic need '
    f'at the population level, with implications for any clinician\u2014surgeon, internist, '
    f'or pain specialist\u2014who prescribes analgesics across the Western Pacific region.',
]

for finding in findings:
    p = doc.add_paragraph(finding, style='List Bullet')

doc.add_paragraph()

more_paras = [
    (
        'We believe this manuscript is particularly suited to The Lancet Regional Health \u2013 '
        'Western Pacific for several reasons. First, it directly addresses within-region '
        'health variation using population-complete data from a major Western Pacific nation, '
        'aligning with the journal\u2019s mission to reduce regional health inequities. '
        'Second, the ecological design using NDB Open Data parallels the approach of '
        'Taira et al (Lancet Reg Health West Pac 2021; 12: 100170), who demonstrated '
        'regional inequality in dental care utilisation using the same data source. '
        'Third, the clinical message\u2014that cultural stereotyping should not replace '
        'individualised pain assessment\u2014has immediate relevance for clinicians '
        'treating Japanese and other East Asian patients throughout the Western Pacific region.'
    ),
    (
        'The study is reported following the STROBE statement and RECORD extension for '
        'observational studies using routinely-collected health data. '
        'All data are publicly available from the Japanese Ministry of Health, Labour and Welfare. '
        'Analysis code is available on GitHub.'
    ),
    (
        'The manuscript has not been published previously, is not under consideration elsewhere, '
        'and the author has no competing interests to declare. '
        'The author approves the manuscript and agrees to its submission.'
    ),
    (
        'Thank you for considering this manuscript. I look forward to hearing from you.'
    ),
]

for text in more_paras:
    doc.add_paragraph(text)

doc.add_paragraph()
doc.add_paragraph('Yours sincerely,')
doc.add_paragraph()

sig = doc.add_paragraph()
run = sig.add_run('Tatsuki Onishi, MD')
run.bold = True

doc.add_paragraph('Department of Anesthesiology')
doc.add_paragraph('[Institution]')
doc.add_paragraph('[Address], [City], [Postal code], Japan')
doc.add_paragraph('E-mail: [email]')

outpath = os.path.join(LRH_DIR, 'LRH_cover_letter.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
