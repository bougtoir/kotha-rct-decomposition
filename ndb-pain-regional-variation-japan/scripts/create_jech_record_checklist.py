#!/usr/bin/env python3
"""Create STROBE/RECORD checklist for JECH submission.

Adapted from Lancet RH version with JECH-specific section references.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JECH_DIR = os.path.join(OUTPUT_DIR, 'jech')
os.makedirs(JECH_DIR, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(8)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('STROBE Statement + RECORD Extension Checklist')
run.bold = True
run.font.size = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    'Do Cultural Labels Predict Analgesic Need? '
    'A 47-Prefecture Ecological Study of Pain Prescribing Variation in Japan'
)
run2.font.size = Pt(10)

doc.add_paragraph()

items = [
    ('Title and abstract', '1', 'STROBE 1(a): Indicate the study design with a commonly used term in the title or abstract', 'Title includes "ecological study"', 'Title page'),
    ('Title and abstract', '1', 'STROBE 1(b): Provide an informative and balanced summary of what was done and found', 'Structured abstract (Background, Methods, Results, Conclusions)', 'Abstract'),
    ('Title and abstract', '1', 'RECORD 1.1: The type of data used should be specified in the title or abstract', '"NDB Open Data" specified in abstract', 'Abstract'),
    ('Title and abstract', '1', 'RECORD 1.2: If applicable, the geographic region and timeframe should be specified', 'Japan, 47 prefectures, April 2023\u2013March 2024', 'Abstract'),
    ('Title and abstract', '1', 'RECORD 1.3: If linkage between databases was conducted, this should be specified', 'N/A \u2013 single data source', 'N/A'),
    ('Introduction', '2', 'STROBE 2: Explain the scientific background and rationale for the investigation', 'Cultural pain stoicism, within-country variation, NDB data', 'Introduction'),
    ('Introduction', '3', 'STROBE 3: State specific objectives, including any prespecified hypotheses', 'Three objectives: map acute variation, adjust chronic proxy, integrate phases', 'Introduction, last paragraph'),
    ('Methods', '4', 'STROBE 4: Present key elements of study design early in the paper', 'Ecological study, STROBE + RECORD reporting', 'Methods: Study design'),
    ('Methods', '4', 'RECORD 4.1: Describe sources of data and methods of assessment', 'NDB Open Data 10th edition described', 'Methods: Data source'),
    ('Methods', '5', 'STROBE 5: Describe the setting, locations, and relevant dates', '47 prefectures, 9 regions, April 2023\u2013March 2024', 'Methods: Data source'),
    ('Methods', '6', 'STROBE 6(a): Cross-sectional \u2013 describe eligibility criteria and sources of data', 'All insurance claims captured; aggregate data', 'Methods: Data source'),
    ('Methods', '7', 'STROBE 7: Clearly define all outcomes, exposures, predictors, confounders', 'Drug classes, formulation counts, confounder proxies defined', 'Methods: Phase 1, Phase 2, Confounders'),
    ('Methods', '8', 'STROBE 8: For each variable, give sources of data and assessment', 'NDB drug classification codes, formulation counts', 'Methods'),
    ('Methods', '9', 'STROBE 9: Describe efforts to address potential sources of bias', 'Per-surgery standardisation, confounder adjustment, perioperative design', 'Methods, Discussion'),
    ('Methods', '10', 'STROBE 10: Explain how the study size was arrived at', 'All 47 prefectures included (complete enumeration)', 'Methods: Data source'),
    ('Methods', '11', 'STROBE 11: Explain how quantitative variables were handled', 'Per-surgery ratios, per-capita rates, regression models', 'Methods: Statistical analysis'),
    ('Methods', '12', 'STROBE 12(a): Describe all statistical methods', 'Kruskal\u2013Wallis, Mann\u2013Whitney U, multiple regression, Cohen\u2019s d', 'Methods: Statistical analysis'),
    ('Methods', '12', 'STROBE 12(d): Cross-sectional \u2013 describe analytical methods for subgroups', 'Models 1\u20135 with progressive confounder adjustment', 'Methods: Statistical analysis'),
    ('Methods', '12', 'RECORD 12.1: Authors should describe data cleaning and linkage', 'Suppressed cells (<10 events); single-source aggregate data', 'Methods: Data source'),
    ('Results', '13', 'STROBE 13(a): Report numbers of individuals at each stage of study', 'N=47 prefectures, 7.9M surgeries, 274M analgesic units', 'Results'),
    ('Results', '14', 'STROBE 14: Descriptive data on demographics and confounders', 'Regional summary statistics in Table 1', 'Results, Table 1'),
    ('Results', '15', 'STROBE 15: Report outcome data', 'Phase 1 and Phase 2 indices reported by region', 'Results'),
    ('Results', '16', 'STROBE 16(a): Give unadjusted estimates and confounding-adjusted estimates', 'Models 1 (unadjusted) through 5 (fully adjusted) in Table 2', 'Results, Table 2'),
    ('Results', '17', 'STROBE 17: Report other analyses done', 'SCR analysis, drug class sub-analyses, acute\u2013chronic correlation', 'Results'),
    ('Discussion', '18', 'STROBE 18: Summarise key results with reference to study objectives', 'Principal findings summarised', 'Discussion, first paragraph'),
    ('Discussion', '19', 'STROBE 19: Discuss limitations, potential biases, and ecological fallacy', 'Ecological fallacy, no diagnosis codes, unmeasured confounders', 'Discussion: Strengths and limitations'),
    ('Discussion', '20', 'STROBE 20: Give a cautious overall interpretation of results', 'Individualised assessment recommended over stereotypes', 'Discussion: Conclusion'),
    ('Discussion', '21', 'STROBE 21: Discuss the generalisability of the study results', 'Applicable to multicultural settings; framework replicable', 'Discussion'),
    ('Other', '22', 'STROBE 22: Give the source of funding and role of funders', 'No funding received', 'Funding'),
    ('Other', '', 'RECORD 22.1: Data governance and ethics', 'Public aggregate data; no ethical approval required', 'Ethics approval'),
    ('Other', '', 'Patient and public involvement', 'No patients or public involved in study design', 'Patient and public involvement'),
    ('Other', '', 'Data accessibility', 'NDB URL provided; analysis code on GitHub', 'Data availability statement'),
]


def set_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


table = doc.add_table(rows=1 + len(items), cols=5, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Section', 'Item #', 'Checklist Item', 'Reported (Location)', 'Page/Section']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(8)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for idx, (section, item_num, description, location, page) in enumerate(items):
    row = table.rows[idx + 1].cells
    row[0].text = section
    row[1].text = item_num
    row[2].text = description
    row[3].text = location
    row[4].text = page
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(7.5)

set_borders(table)

widths = [Inches(1.0), Inches(0.5), Inches(4.5), Inches(3.0), Inches(1.5)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

outpath = os.path.join(JECH_DIR, 'JECH_STROBE_RECORD_checklist.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
