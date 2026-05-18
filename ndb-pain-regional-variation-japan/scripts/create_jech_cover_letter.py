#!/usr/bin/env python3
"""Create cover letter for JECH submission."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
from datetime import date

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JECH_DIR = os.path.join(OUTPUT_DIR, 'jech')
os.makedirs(JECH_DIR, exist_ok=True)

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

unadj_d = reg["model1_unadjusted"]["cohens_d"]
attenuation = (1 - reg["adjusted_cpsp_test"]["cohens_d"] / unadj_d) * 100

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

today = date.today().strftime('%B %d, %Y')

doc.add_paragraph(today)
doc.add_paragraph()

doc.add_paragraph('Professor Vittal Katikireddi and Dr Anna Pearce')
doc.add_paragraph('Editors-in-Chief')
doc.add_paragraph('Journal of Epidemiology and Community Health')
doc.add_paragraph()

doc.add_paragraph('Dear Professors Katikireddi and Pearce,')
doc.add_paragraph()

body_paras = [
    (
        'I am pleased to submit the enclosed manuscript entitled '
        '\u201cDo Cultural Labels Predict Analgesic Need? '
        'A 47-Prefecture Ecological Study of Pain Prescribing '
        'Variation in Japan\u201d '
        'for consideration as an Original Research article in the '
        'Journal of Epidemiology and Community Health.'
    ),
    (
        'Clinicians frequently rely on cultural labels\u2014such as \u201cJapanese patients '
        'are stoic\u201d\u2014when making analgesic prescribing decisions. '
        'Yet all prior comparisons have treated Japan as a culturally homogeneous unit. '
        'Using population-complete insurance claims data from Japan\u2019s National Database '
        '(NDB Open Data, covering approximately 125 million insured individuals), '
        'we mapped perioperative and chronic pain-related prescribing across all 47 prefectures '
        'for the first time.'
    ),
    'Our principal findings are:',
]

for text in body_paras:
    doc.add_paragraph(text)

findings = [
    'Acute perioperative analgesic prescribing varied 1.97-fold across prefectures '
    '(Kruskal\u2013Wallis P<0.001), with significant regional clustering.',

    'Tohoku\u2014traditionally considered Japan\u2019s most stoic region\u2014prescribed '
    'more, not fewer, analgesics (Cohen\u2019s d=0.87), contradicting the stereotype.',

    f'The apparent Tohoku excess in neuropathic pain prescribing was largely explained by '
    f'confounding disease prevalence (especially diabetes; r=0.87), with '
    f'{attenuation:.0f}% attenuation after adjustment.',

    'These findings demonstrate that within-country prescribing heterogeneity '
    'challenges the validity of national-level cultural stereotypes in clinical practice.',
]

for finding in findings:
    doc.add_paragraph(finding, style='List Bullet')

doc.add_paragraph()

more_paras = [
    (
        'We believe this manuscript aligns well with JECH\u2019s scope and readership. '
        'The study addresses socioeconomic and cultural determinants of health at the '
        'population level, using a novel ecological framework applied to population-complete '
        'national claims data. The finding that cultural stereotypes fail to predict '
        'prescribing patterns has implications for health equity and clinical practice '
        'globally\u2014not only in Japan but wherever national-level cultural generalisations '
        'inform clinical decisions. The within-database confounder-adjustment methodology '
        'demonstrated here is directly replicable in other countries with national claims '
        'databases.'
    ),
    (
        'The study is reported following the STROBE statement and RECORD extension for '
        'observational studies using routinely-collected health data. '
        'All data are publicly available from the Japanese Ministry of Health, Labour and Welfare. '
        'Analysis code is available on GitHub.'
    ),
    (
        'The manuscript has not been published previously, is not under consideration elsewhere, '
        'and all authors have approved the submitted version. '
        'There are no conflicts of interest to declare. No funding was received for this study.'
    ),
    (
        'Thank you for considering this manuscript. I look forward to your response.'
    ),
]

for text in more_paras:
    doc.add_paragraph(text)

doc.add_paragraph()
doc.add_paragraph('Sincerely,')
doc.add_paragraph()
doc.add_paragraph('Tatsuki Onishi, MD')
doc.add_paragraph('Department of Anesthesiology')
doc.add_paragraph('[Institution]')
doc.add_paragraph('E-mail: [email]')

outpath = os.path.join(JECH_DIR, 'JECH_cover_letter.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
