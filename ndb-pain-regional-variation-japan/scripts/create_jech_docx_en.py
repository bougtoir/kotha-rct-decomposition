#!/usr/bin/env python3
"""Create English manuscript for JECH (Journal of Epidemiology and Community Health).

Key JECH / BMJ requirements (Original Research):
- Abstract: structured (Background, Methods, Results, Conclusions) ≤250 words
- Body: ~3,000 words (excl. abstract/refs/legends)
- "What is already known on this topic" / "What this study adds" box
- References: Vancouver style — numbered [n] in order of first appearance
- Figures: uploaded as separate TIFF/EPS files (not embedded)
- Tables: Word format, placed in main text where first cited
- STROBE + RECORD checklist for observational studies using routinely-collected data
- Patient and public involvement statement
- Double-spaced, continuous line numbering
- Data sharing statement
- Contributor statement
- Ethics statement
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re
import csv
import numpy as np
from collections import defaultdict

# ============================================================
# Paths
# ============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JECH_DIR = os.path.join(OUTPUT_DIR, 'jech')
os.makedirs(JECH_DIR, exist_ok=True)

FIG_DIR = OUTPUT_DIR

# ============================================================
# Load data
# ============================================================
with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'scr_summary.json'), 'r') as f:
    scr = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_EN = {
    '北海道': 'Hokkaido', '東北': 'Tohoku', '関東': 'Kanto',
    '北陸・甲信越': 'Hokuriku-Koshinetsu', '東海': 'Tokai', '近畿': 'Kinki',
    '中国': 'Chugoku', '四国': 'Shikoku', '九州・沖縄': 'Kyushu-Okinawa',
}
REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

# ============================================================
# REFERENCES — Vancouver style, numbered in order of first appearance
# ============================================================
ref_list = [
    # 1 Callister
    'Callister LC. Cultural influences on pain perceptions and behaviors. '
    'Home Health Care Manag Pract 2003;15:207\u201311. '
    'doi:10.1177/1084822303015004003',
    # 2 Rogger
    'Rogger R, Bello C, Romero CS, et al. '
    'Cultural framing and the impact on acute pain and pain services. '
    'Curr Pain Headache Rep 2023;27:429\u201336. '
    'doi:10.1007/s11916-023-01134-3',
    # 3 Zborowski
    'Zborowski M. People in Pain. San Francisco: Jossey-Bass, 1969.',
    # 4 Okolo
    'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: '
    'a review of cross-cultural studies. Int J Sci Res Arch 2024;11:2550\u20136. '
    'doi:10.30574/ijsra.2024.11.1.0298',
    # 5 Hobara
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
    'between Japanese and Euro-Americans. Eur J Pain 2005;9:389\u201393. '
    'doi:10.1016/j.ejpain.2004.09.006',
    # 6 Feng
    'Feng Y, Herdman M, van Nooten F, et al. An exploration of differences between Japan '
    'and two European countries in the self-reporting and valuation of pain and discomfort '
    'on the EQ-5D. Qual Life Res 2017;26:2067\u201378. '
    'doi:10.1007/s11136-017-1541-7',
    # 7 Cohen
    'Cohen D, Nisbett RE, Bowdle BF, et al. Insult, aggression, and the southern '
    'culture of honor: an "experimental ethnography." J Pers Soc Psychol 1996;70:945\u201360. '
    'doi:10.1037/0022-3514.70.5.945',
    # 8 Kumagai
    'Kumagai S. Media representations reproducing images of Tohoku: the Tohoku '
    'reconstruction corner in "Secret Kenmin SHOW." Kotoba 2020;41:21\u201338. [in Japanese]',
    # 9 Takeda
    'Takeda K, Yarimizu K. Regional differences in the pain expression uzuku. '
    'NINJAL Research Papers 2016;10:85\u2013107. [in Japanese]',
    # 10 Pfizer
    'Pfizer Japan Inc. 47-prefecture survey on chronic pain. '
    'https://www.pfizer.co.jp/pfizer/company/press/2017 (accessed 1 Feb 2025).',
    # 11 MHLW
    'Ministry of Health, Labour and Welfare. NDB Open Data, 10th edition. '
    'https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html '
    '(accessed 15 Jan 2025).',
    # 12 Wakaizumi
    'Wakaizumi K, Tanaka C, Shinohara Y, et al. Geographical variation in high-impact '
    'chronic pain and psychological associations at the regional level: a multilevel analysis '
    'of a large-scale internet-based cross-sectional survey. Front Public Health 2024;12:1482177. '
    'doi:10.3389/fpubh.2024.1482177',
    # 13 Matsuoka
    'Matsuoka Y, Morishima T, Sato A, et al. Population-based claims study '
    'of regional and hospital function differences in opioid prescribing for cancer patients '
    'who died in hospital in Japan. Jpn J Clin Oncol 2025;55:hyaf149.',
    # 14 Taira
    'Taira K, Mori T, Ishimaru M, et al. Regional inequality in dental care utilisation '
    'in Japan: an ecological study using the National Database of Health Insurance Claims. '
    'Lancet Reg Health West Pac 2021;12:100170. '
    'doi:10.1016/j.lanwpc.2021.100170',
    # 15 von Elm (STROBE)
    'von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational '
    'Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. '
    'Lancet 2007;370:1453\u20137. '
    'doi:10.1016/S0140-6736(07)61602-X',
    # 16 Benchimol (RECORD)
    'Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using '
    'Observational Routinely-collected health Data (RECORD) statement. '
    'PLoS Med 2015;12:e1001885. '
    'doi:10.1371/journal.pmed.1001885',
    # 17 Anderson
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain: '
    'causes and consequences of unequal care. J Pain 2009;10:1187\u2013204. '
    'doi:10.1016/j.jpain.2009.06.010',
    # 18 Campbell
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
    'Pain Manag 2012;2:219\u201330. '
    'doi:10.2217/pmt.12.7',
    # 19 Befu
    'Befu H. Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. '
    'Melbourne: Trans Pacific Press, 2001.',
    # 20 Burgess
    'Burgess C. The "illusion" of homogeneous Japan and national character: '
    'discourse as a tool to transcend the "myth" vs. "reality" binary. '
    'Asia Pac J 2010;8(9):1\u201322.',
    # 21 Raja (IASP)
    'Raja SN, Carr DB, Cohen M, et al. The revised International Association for the '
    'Study of Pain definition of pain: concepts, challenges, and compromises. '
    'Pain 2020;161:1976\u201382. '
    'doi:10.1097/j.pain.0000000000001939',
    # 22 Onishi
    'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor of respiration recovery '
    'and quantification of nociception anti-nociception balance compared to opioid effect site '
    'concentration: a prospective, observational study. F1000Research 2024;13:233. '
    'doi:10.12688/f1000research.147085.2',
    # 23 Kehlet
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
    'Lancet 2006;367:1618\u201325. '
    'doi:10.1016/S0140-6736(06)68700-X',
]


def cite(*nums):
    """Return Vancouver superscript citation marker string."""
    return '{' + ','.join(str(n) for n in nums) + '}'


def add_ref_runs(p, text):
    """Parse text with {n} or {n,m} markers and create runs with font-based superscript."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(10)
        else:
            p.add_run(part)


# ============================================================
# Document creation
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

# Page numbers
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def wc(text):
    return len(re.sub(r'\{[^}]+\}', '', text).split())


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def add_inline_figure(caption_text, fig_num):
    """Insert figure placeholder and caption (figures submitted as independent TIFF files)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    r0 = p.add_run(f'[Insert Figure {fig_num} here]')
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f'Figure {fig_num}. ')
    r.bold = True
    r.font.size = Pt(10)
    r2 = cap.add_run(caption_text)
    r2.font.size = Pt(10)


# ============================================================
# Computed values
# ============================================================
unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100

scr_analgesic_range = scr['analgesic_inpatient']['scr_range']
scr_analgesic_ratio = scr['analgesic_inpatient']['variation_ratio']
scr_neuro_range = scr['neuropathic_outpatient']['scr_range']
scr_neuro_ratio = scr['neuropathic_outpatient']['variation_ratio']
scr_neuro_tohoku = scr['neuropathic_outpatient']['scr_tohoku_mean']
scr_neuro_non_tohoku = scr['neuropathic_outpatient']['scr_non_tohoku_mean']

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    'Do Cultural Labels Predict Analgesic Need? '
    'A 47-Prefecture Ecological Study of Pain Prescribing '
    'Variation in Japan'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Tatsuki Onishi, MD')
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Department of Anesthesiology, [Institution], [Address], [City], [Postal code], Japan')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Corresponding author: ')
r.bold = True
p.add_run(
    'Tatsuki Onishi, MD, Department of Anesthesiology, [Institution], '
    '[Address], [City], [Postal code], Japan. E-mail: [email]'
)

doc.add_paragraph()

# Word count line
p = doc.add_paragraph()
r = p.add_run('Word count: ')
r.bold = True
p.add_run('[body word count] (excluding abstract, references, tables, and figure legends)')

doc.add_page_break()

# ============================================================
# ABSTRACT (Background, Methods, Results, Conclusions) ≤250 words
# ============================================================
add_heading_text('Abstract', level=1)

abstract_bg = (
    'Clinicians frequently rely on cultural labels when estimating analgesic requirements, '
    'yet whether pain-related prescribing varies within Japan has not been examined '
    'at the population level.'
)

abstract_methods = (
    'This ecological study used Japan\u2019s National Database of Health Insurance Claims '
    '(NDB Open Data, 10th edition; population-complete, April 2023\u2013March 2024) to map '
    'pain-related prescribing across all 47 prefectures. '
    'Phase 1 examined acute perioperative analgesic prescribing per surgery. '
    'Phase 2 examined outpatient neuropathic pain drug prescribing as a chronic postsurgical '
    'pain proxy, adjusted for confounding diseases using multiple regression. '
    'Standardised claim ratios (SCR) confirmed findings after age-sex standardisation.'
)

abstract_results = (
    f'Acute analgesic prescribing varied 1\u00b797-fold across prefectures '
    f'(Kruskal\u2013Wallis P<0.001). '
    f'Tohoku\u2014traditionally considered Japan\u2019s most stoic region\u2014prescribed more, '
    f'not fewer, analgesics (Cohen\u2019s d=0.87). '
    f'Unadjusted neuropathic pain prescribing showed a large Tohoku excess '
    f'(d={unadj_d:.2f}), but diabetes drug prescribing was a strong confounder (r=0.87). '
    f'After adjustment, the Tohoku effect was attenuated by {attenuation:.0f}% '
    f'(P={reg["adjusted_cpsp_test"]["p_value"]:.2f}). '
    f'SCR confirmed {scr_analgesic_ratio:.2f}-fold variation after age-sex standardisation.'
)

abstract_conclusions = (
    'Nearly twofold within-country variation demonstrates that cultural labels '
    'do not predict analgesic need. Confounding diseases substantially modified the '
    'apparent regional pattern. Clinicians should base analgesic decisions on '
    'individual assessment rather than cultural assumptions.'
)

for label, text in [
    ('Background', abstract_bg),
    ('Methods', abstract_methods),
    ('Results', abstract_results),
    ('Conclusion', abstract_conclusions),
]:
    p = doc.add_paragraph()
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

abstract_total = sum(wc(t) for t in [abstract_bg, abstract_methods, abstract_results,
                                      abstract_conclusions])
print(f'Abstract word count: {abstract_total} (JECH limit: 250)')

doc.add_page_break()

# ============================================================
# WHAT IS ALREADY KNOWN / WHAT THIS STUDY ADDS
# ============================================================
add_heading_text('What is already known on this topic', level=2)

known_items = [
    'Cross-cultural studies characterise Japanese people as stoic toward pain, '
    'yet all prior comparisons treat Japan as a homogeneous unit.',
    'Regional variation in chronic pain prevalence and cancer opioid prescribing '
    'has been documented in Japan, but no study has examined perioperative '
    'analgesic prescribing at the population level.',
    'Cultural stereotypes can influence clinician pain assessment and prescribing behaviour, '
    'leading to systematic under- or over-treatment.',
]
for item in known_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_text('What this study adds', level=2)

adds_items = [
    'Acute perioperative analgesic prescribing varies 1.97-fold across Japan\u2019s '
    '47 prefectures, refuting the notion of uniform stoicism.',
    'Tohoku\u2014Japan\u2019s culturally "most stoic" region\u2014prescribes more, not fewer, '
    'analgesics, and this apparent excess is largely explained by confounding disease prevalence.',
    'Within-country prescribing heterogeneity challenges the validity of any national-level '
    'cultural stereotype in clinical decision-making.',
]
for item in adds_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_text('How this study might affect research, practice or policy', level=2)

policy_items = [
    'Clinicians should base analgesic decisions on individual assessment rather than '
    'cultural assumptions about pain tolerance\u2014the nearly twofold within-Japan variation '
    'demonstrates that no national-level stereotype reliably predicts prescribing need.',
    'Ecological studies using neuropathic pain drug prescribing as a chronic pain proxy '
    'must adjust for confounding diseases (especially diabetes); without adjustment, '
    'regional differences in disease prevalence can be misinterpreted as pain behaviour differences.',
    'The within-database confounder-adjustment framework demonstrated here is replicable '
    'in any country with a national claims database and could inform future population-level '
    'pain research.',
]
for item in policy_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('Introduction', level=1)

intro_parts = []

intro1 = (
    f'When prescribing analgesics, clinicians routinely make assumptions about how much pain '
    f'a patient can tolerate. These assumptions are often anchored to cultural '
    f'labels.{cite(1,2)} Since Zborowski\u2019s landmark observation that ethnic groups '
    f'differ in pain behaviour,{cite(3)} a large body of literature has established that '
    f'cultural norms influence pain reporting, treatment-seeking, and analgesic use.{cite(4)} '
    f'Japan is a case in point: Hobara found that Japanese respondents rated pain behaviours '
    f'as less appropriate than Euro-Americans,{cite(5)} and Feng et al showed that Japanese '
    f'participants were far less willing to trade time to avoid pain on the EQ-5D than '
    f'respondents in the UK and the Netherlands.{cite(6)} '
    f'The resulting cultural label\u2014\u201cJapanese patients are stoic\u201d\u2014carries '
    f'clinical risk whenever a clinician\u2014whether surgeon, internist, or '
    f'anaesthesiologist\u2014assumes that a stoic presentation indicates lower analgesic need.'
)
p = doc.add_paragraph()
add_ref_runs(p, intro1)
intro_parts.append(intro1)

intro2 = (
    f'Most cross-cultural pain studies compare behaviour between nations, treating each country '
    f'as a homogeneous unit. However, substantial cultural heterogeneity exists within countries. '
    f'Cohen and Nisbett demonstrated that the \u201cculture of honour\u201d in the southern '
    f'United States produces measurably different behavioural responses from the North, including '
    f'physiological stress responses.{cite(7)} '
    f'In Japan, regional cultural identities remain strong despite the discourse of national '
    f'homogeneity. The Tohoku region (northeastern Honshu) is traditionally perceived as '
    f'embodying stoic endurance,{cite(8)} and even the Japanese word for throbbing pain '
    f'(uzuku) shows distinct regional usage patterns that map onto historical dialect '
    f'boundaries.{cite(9)} A 2017 Pfizer Japan survey found that the proportion of chronic '
    f'pain patients \u201cenduring pain without seeking treatment\u201d ranged from 48\u00b77% '
    f'to 81\u00b76% across prefectures.{cite(10)} '
    f'Yet whether such differences translate into measurable differences in healthcare '
    f'utilisation at the population level remains unknown.'
)
p = doc.add_paragraph()
add_ref_runs(p, intro2)
intro_parts.append(intro2)

intro3 = (
    f'Japan\u2019s universal health insurance system, standardised drug pricing, and the '
    f'National Database of Health Insurance Claims (NDB)\u2014capturing virtually all '
    f'reimbursed healthcare utilisation for approximately 125 million insured '
    f'individuals{cite(11)}\u2014provide a unique setting for ecological analysis of '
    f'regional prescribing variation. Recent work has documented up to 1\u00b76-fold regional '
    f'variation in chronic pain prevalence using a large internet survey{cite(12)} and '
    f'4-fold variation in cancer opioid prescribing using hospital claims data.{cite(13)} '
    f'Taira et al demonstrated substantial regional inequality in dental care utilisation '
    f'using NDB-derived standardised claim ratios across all 47 prefectures.{cite(14)} '
    f'However, no study has applied this population-complete ecological framework '
    f'to pain-related prescribing.'
)
p = doc.add_paragraph()
add_ref_runs(p, intro3)
intro_parts.append(intro3)

intro4 = (
    'The perioperative setting offers a further methodological advantage: because all '
    'patients in Phase 1 are hospitalised for surgery, healthcare access\u2014a major confounder '
    'in community-based pain studies\u2014is neutralised by design. '
    'This study had three objectives: (1) map regional variation in acute perioperative '
    'analgesic prescribing across 47 prefectures; (2) examine outpatient neuropathic pain '
    'prescribing as a chronic postsurgical pain (CPSP) proxy after adjustment for '
    'confounding diseases; and (3) integrate acute and chronic pain findings at the '
    'population level.'
)
doc.add_paragraph(intro4)
intro_parts.append(intro4)

intro_total = sum(wc(t) for t in intro_parts)
print(f'Introduction word count: {intro_total}')

# ============================================================
# METHODS
# ============================================================
add_heading_text('Methods', level=1)

methods_parts = []

add_heading_text('Study design and reporting', level=2)
m1 = (
    f'This ecological study analysed prefecture-level aggregate data from the NDB Open Data. '
    f'It is reported following the Strengthening the Reporting of Observational Studies in '
    f'Epidemiology (STROBE) statement{cite(15)} and the REporting of studies Conducted using '
    f'Observational Routinely-collected health Data (RECORD) extension.{cite(16)} '
    f'As only publicly available aggregate data were used, ethical approval was not required '
    f'under Japan\u2019s Ethical Guidelines for Medical and Biological Research Involving '
    f'Human Subjects.'
)
p = doc.add_paragraph()
add_ref_runs(p, m1)
methods_parts.append(m1)

add_heading_text('Data source', level=2)
m2 = (
    f'The 10th edition of the NDB Open Data (April 2023\u2013March 2024) was used.{cite(11)} '
    f'The NDB captures claims from all insurers within Japan\u2019s universal coverage system, '
    f'encompassing approximately 125 million insured individuals. '
    f'Aggregate prescription and procedure data are published at the prefecture level (n=47) '
    f'with suppression of cells containing fewer than ten events. '
    f'Prefecture-level population estimates (October 2023, by five-year age group and sex) '
    f'from the Statistics Bureau of Japan were used to compute per-capita rates and '
    f'standardised claim ratios (SCR).'
)
p = doc.add_paragraph()
add_ref_runs(p, m2)
methods_parts.append(m2)

add_heading_text('Regional classification', level=2)
m3 = (
    'Prefectures were grouped into nine standard regional blocks following '
    'the classification used by the Statistics Bureau of Japan: '
    'Hokkaido (1 prefecture), Tohoku (6: Aomori, Iwate, Miyagi, Akita, Yamagata, '
    'Fukushima), Kanto (7), Hokuriku-Koshinetsu (6), Tokai (4), '
    'Kinki (6), Chugoku (5), Shikoku (4), and Kyushu-Okinawa (8). '
    'Tohoku was designated as the primary region of interest a priori, based on its '
    'traditional cultural characterisation as embodying patient endurance (gaman).'
)
doc.add_paragraph(m3)
methods_parts.append(m3)

add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
m4 = (
    'Inpatient prescription data were extracted for three analgesic drug classes commonly '
    'used in perioperative pain management: '
    'Class 114 (antipyretic analgesics including NSAIDs and acetaminophen), '
    'Class 811 (opium alkaloid narcotics including morphine and codeine), '
    'and Class 821 (synthetic narcotics including fentanyl and pethidine). '
    'Inpatient surgical procedure counts were extracted from the K Surgery section '
    'of the claims data. '
    'The analgesic-per-surgery index was calculated for each prefecture as '
    'total inpatient analgesic prescription units divided by total inpatient '
    'surgical procedure count. This ratio provides a standardised measure of analgesic '
    'intensity that accounts for differences in surgical volume between prefectures. '
    'Separate sub-analyses were conducted for each drug class to assess consistency.'
)
doc.add_paragraph(m4)
methods_parts.append(m4)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing', level=2)
m5 = (
    'Five classes of outpatient oral neuropathic pain medications were extracted: '
    'pregabalin (78 formulations), mirogabalin (8 formulations), '
    'duloxetine (33 formulations), tramadol (3 formulations), and neurotropin (1 formulation). '
    'These agents are first-line or commonly used medications for neuropathic pain in Japan '
    'and serve as a population-level proxy for chronic neuropathic pain burden. '
    'The neuropathic pain prescribing-per-surgery index was calculated as total outpatient '
    'neuropathic pain drug quantity divided by total inpatient surgical procedure count, '
    'using surgery count as a denominator that normalises for healthcare system capacity. '
    'Per-capita neuropathic pain prescribing rates (units per thousand population) were '
    'additionally computed using prefecture population data.'
)
doc.add_paragraph(m5)
methods_parts.append(m5)

add_heading_text('Confounder disease proxies', level=2)
m6 = (
    'Four confounder disease proxies were extracted from outpatient data: '
    'oral hypoglycaemic agents (261 formulations; proxy for diabetic neuropathy, the most '
    'common cause of neuropathic pain in Japan), '
    'herpes zoster antivirals (47 formulations; proxy for postherpetic neuralgia), '
    'antidepressants excluding duloxetine (128 formulations; proxy for depression, '
    'which shares gabapentinoid prescribing), '
    'and anxiolytics (112 formulations; proxy for anxiety disorders). '
    'Each proxy was expressed per surgery to maintain consistency with the primary outcome. '
    'Outpatient nerve block procedure counts (73 codes) served as an additional independent '
    'CPSP-related proxy.'
)
doc.add_paragraph(m6)
methods_parts.append(m6)

add_heading_text('Statistical analysis', level=2)
m7a = (
    'Regional differences in Phase 1 were assessed using the Kruskal\u2013Wallis test '
    'across nine regional blocks, followed by post-hoc Mann\u2013Whitney U tests with Bonferroni '
    'correction for pairwise comparisons. Effect sizes were quantified using Cohen\u2019s d. '
    'For Phase 2, five regression models were fitted to examine the Tohoku regional effect '
    'with progressive confounder adjustment: '
    'Model 1 (unadjusted Tohoku vs non-Tohoku comparison), '
    'Model 2 (neuropathic pain ~ diabetes + herpes + antidepressants + anxiolytics + Tohoku indicator), '
    'Model 3 (core neuropathic drugs only ~ same confounders), '
    'Model 4 (nerve blocks ~ same confounders), and '
    'Model 5 (neuropathic pain ~ acute analgesic index + confounders).'
)
doc.add_paragraph(m7a)
methods_parts.append(m7a)

m7b = (
    'The adjusted CPSP index was derived as residuals from regressing neuropathic pain '
    'prescribing on the four confounder proxies. '
    'Standardised claim ratios (SCR) were computed by indirect age-sex standardisation '
    f'following Taira et al.{cite(14)} '
    'National age-sex-specific prescription rates (18 five-year age groups \u00d7 2 sexes) '
    'from the NDB were applied to each prefecture\u2019s population structure. '
    'All analyses used Python 3.11 (NumPy 1.24, SciPy 1.11).'
)
p = doc.add_paragraph()
add_ref_runs(p, m7b)
methods_parts.append(m7b)

add_heading_text('Patient and public involvement', level=2)
m_ppi = (
    'Patients or members of the public were not involved in the design, conduct, '
    'reporting, or dissemination plans of this research.'
)
doc.add_paragraph(m_ppi)
methods_parts.append(m_ppi)

methods_total = sum(wc(t) for t in methods_parts)
print(f'Methods word count: {methods_total}')

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

results_parts = []

add_heading_text('Phase 1: Regional variation in acute perioperative analgesic prescribing', level=2)
r1 = (
    'During April 2023\u2013March 2024, the NDB recorded 7,903,515 inpatient surgical procedures '
    'and 274,579,851 analgesic prescription units across 47 prefectures. '
    'The national mean analgesic-per-surgery index was 35.78 (SD 5.56), '
    'ranging from 25.20 (Gifu) to 49.75 (Kagoshima)\u2014a 1.97-fold difference '
    '(Kruskal\u2013Wallis P<0.001 across nine regions; table 1).'
)
doc.add_paragraph(r1)
results_parts.append(r1)

# === TABLE 1 INLINE ===
p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(14)
r_cap = p_cap.add_run('Table 1. ')
r_cap.bold = True
r_cap.font.size = Pt(10)
p_cap.add_run(
    'Regional summary of inpatient analgesic prescribing per surgery across nine regional blocks. '
    'Values are analgesic-per-surgery index (mean \u00b1 SD). '
    'Kruskal\u2013Wallis P<0.001.'
).font.size = Pt(10)

t1 = doc.add_table(rows=1 + len(REGION_ORDER), cols=4, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t1.rows[0].cells
for i, h in enumerate(['Region', 'n (prefectures)', 'Mean \u00b1 SD', 'Range']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = t1.rows[idx + 1].cells
    row[0].text = REGION_EN[reg_name]
    row[1].text = str(len(vals))
    row[2].text = f'{np.mean(vals):.2f} \u00b1 {np.std(vals, ddof=1):.2f}' if len(vals) > 1 else f'{np.mean(vals):.2f}'
    row[3].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t1)
doc.add_paragraph()

r2 = (
    'Substantial regional clustering was observed. Tokai and Kinki (western Japan) had the lowest '
    'indices, while Kyushu-Okinawa and Hokkaido had the highest. '
    'Tohoku, culturally perceived as Japan\u2019s most stoic region, '
    'ranked seventh of nine with a mean index of 39.97 (SD 3.53), '
    'significantly above the non-Tohoku mean of 35.17 '
    '(Mann\u2013Whitney U=190, P=0.031; Cohen\u2019s d=0.87). '
    'All six Tohoku prefectures ranked in the upper half nationally. '
    'This pattern was consistent across all three drug classes: '
    'NSAIDs (P=0.044), opioid alkaloids (P=0.003), and synthetic opioids (P=0.001).'
)
doc.add_paragraph(r2)
results_parts.append(r2)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing (unadjusted)', level=2)
r3 = (
    f'Nationally, outpatient neuropathic pain drug prescriptions totalled 2,289,549,163 units, '
    f'comprising pregabalin (40.2%), neurotropin (20.1%), mirogabalin (19.6%), '
    f'duloxetine (15.3%), and tramadol (4.9%). '
    f'Tohoku had a markedly higher neuropathic pain prescribing-per-surgery index '
    f'({reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; P<0.001; '
    f'd={reg["model1_unadjusted"]["cohens_d"]:.2f}), '
    f'with Iwate (566.7), Aomori (519.3), and Akita (461.1) '
    f'occupying the top three nationally (figure 1).'
)
doc.add_paragraph(r3)
results_parts.append(r3)

# === FIGURE 1 placeholder ===
add_inline_figure(
    'Outpatient neuropathic pain drug prescribing per surgery by prefecture (unadjusted). '
    'Tohoku prefectures (red) cluster at the high end. Dashed line = national mean.',
    1
)

add_heading_text('Confounder analysis and adjustment', level=2)
r4a = (
    f'Neuropathic pain prescribing showed strong correlations with confounder disease proxies. '
    f'Diabetes drug prescribing was the strongest correlate (r=0.87, P<0.001), '
    f'followed by anxiolytics (r=0.75), antidepressants (r=0.46), '
    f'and herpes antivirals (r=0.19). '
    f'These four confounders collectively explained '
    f'{reg["model2_adjusted"]["R2"]*100:.1f}% of between-prefecture variance in neuropathic pain '
    f'prescribing (R\u00b2={reg["model2_adjusted"]["R2"]:.3f} in Model 2; figure 2).'
)
doc.add_paragraph(r4a)
results_parts.append(r4a)

# === FIGURE 2 placeholder ===
add_inline_figure(
    'Correlation between neuropathic pain prescribing and confounder disease proxies. '
    'Each dot represents one prefecture. Tohoku prefectures are marked with red borders. '
    'Diabetes drugs show the strongest correlation (r=0.87).',
    2
)

r4b = (
    f'After adjustment for all four confounders, the Tohoku effect was attenuated '
    f'and became nonsignificant in Model 2 '
    f'(\u03b2={reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
    f'P={reg["model2_adjusted"]["tohoku_p"]:.2f}). '
    f'This was consistent across specifications: '
    f'Model 3 (core neuropathic drugs only; '
    f'\u03b2={reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P={reg["model3_core_neuropathic"]["tohoku_p"]:.2f}), '
    f'Model 4 (nerve blocks; P={reg["model4_nerve_blocks"]["tohoku_p"]:.2f}), and '
    f'Model 5 (integrated; \u03b2={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P={reg["model5_integrated"]["tohoku_p"]:.2f}; table 2).'
)
doc.add_paragraph(r4b)
results_parts.append(r4b)

# === TABLE 2 INLINE ===
p_cap2 = doc.add_paragraph()
p_cap2.paragraph_format.space_before = Pt(14)
r_cap2 = p_cap2.add_run('Table 2. ')
r_cap2.bold = True
r_cap2.font.size = Pt(10)
p_cap2.add_run(
    'Regression models for outpatient neuropathic pain prescribing with Tohoku '
    'indicator and confounder adjustment.'
).font.size = Pt(10)

models = [
    ('Model 1 (unadjusted)', '\u2014', '\u2014',
     f'd={reg["model1_unadjusted"]["cohens_d"]:.2f}', 'P<0.001'),
    ('Model 2 (all confounders)',
     f'{reg["model2_adjusted"]["tohoku_coef"]:.1f}',
     f'{reg["model2_adjusted"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}', ''),
    ('Model 3 (core neuro)',
     f'{reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}',
     f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model3_core_neuropathic"]["R2"]:.3f}', ''),
    ('Model 4 (nerve blocks)',
     f'{reg["model4_nerve_blocks"]["tohoku_coef"]:.1f}',
     f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model4_nerve_blocks"]["R2"]:.3f}', ''),
    ('Model 5 (integrated)',
     f'{reg["model5_integrated"]["tohoku_coef"]:.1f}',
     f'{reg["model5_integrated"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model5_integrated"]["R2"]:.3f}', ''),
]

t2 = doc.add_table(rows=1 + len(models), cols=5, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'Tohoku \u03b2', 'Tohoku P', 'Fit', 'Note']):
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, (name, beta, pval, fit, note) in enumerate(models):
    row = t2.rows[idx + 1].cells
    row[0].text = name
    row[1].text = beta
    row[2].text = pval
    row[3].text = fit
    row[4].text = note
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t2)
doc.add_paragraph()

r5a = (
    f'The adjusted CPSP index showed a dramatically different geographic pattern '
    f'from the unadjusted data. '
    f'The Tohoku mean shifted from markedly positive to a modest, nonsignificant excess '
    f'({reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs '
    f'{reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f'P={reg["adjusted_cpsp_test"]["p_value"]:.2f}; '
    f'd={reg["adjusted_cpsp_test"]["cohens_d"]:.2f}; figure 3). '
    f'Chugoku emerged as having the highest adjusted index, while Tokai had the lowest.'
)
doc.add_paragraph(r5a)
results_parts.append(r5a)

# === FIGURE 3 placeholder ===
add_inline_figure(
    'Regional comparison of neuropathic pain prescribing: (a) unadjusted and '
    '(b) after confounder adjustment. Tohoku (red) shifts from the highest region to '
    'mid-range after adjustment. Error bars = SD.',
    3
)

add_heading_text('Phase 1\u2013Phase 2 integration', level=2)
r6 = (
    f'Acute perioperative prescribing correlated positively with unadjusted neuropathic '
    f'pain prescribing (r=0.38, P=0.008). After confounder adjustment, this '
    f'correlation was attenuated (r=0.29, P=0.052). '
    f'In Model 5, the acute pain index remained a significant predictor of chronic neuropathic '
    f'pain prescribing (\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}), '
    f'while the Tohoku effect remained nonsignificant '
    f'(\u03b2={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P={reg["model5_integrated"]["tohoku_p"]:.2f}). '
    f'After adjustment, the Tohoku effect was attenuated by {attenuation:.0f}%.'
)
doc.add_paragraph(r6)
results_parts.append(r6)

add_heading_text('Age-sex standardised claim ratios', level=2)
r7 = (
    f'After indirect age-sex standardisation, inpatient analgesic SCR ranged from '
    f'{scr_analgesic_range[0]:.1f} to {scr_analgesic_range[1]:.1f} '
    f'({scr_analgesic_ratio:.2f}-fold variation), confirming that the prescribing heterogeneity '
    f'was not attributable to differences in prefectural age-sex composition. '
    f'Outpatient neuropathic pain drug SCR ranged from {scr_neuro_range[0]:.1f} to '
    f'{scr_neuro_range[1]:.1f} ({scr_neuro_ratio:.2f}-fold). '
    f'Tohoku retained elevated neuropathic pain SCR (mean {scr_neuro_tohoku:.1f}) '
    f'relative to non-Tohoku prefectures ({scr_neuro_non_tohoku:.1f}), consistent with '
    f'the crude analysis.'
)
doc.add_paragraph(r7)
results_parts.append(r7)

results_total = sum(wc(t) for t in results_parts)
print(f'Results word count: {results_total}')

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('Discussion', level=1)

disc_parts = []

d1 = (
    'This study is the first to map perioperative and chronic pain-related prescribing across '
    'all 47 prefectures of Japan, leveraging freely available NDB Open Data that capture '
    'virtually the entire insured population. The answer to the question posed in the title '
    'is no: cultural labels do not predict analgesic need. Three principal findings support '
    'this conclusion and carry implications for any clinician who prescribes analgesics.'
)
doc.add_paragraph(d1)
disc_parts.append(d1)

add_heading_text('Within-country heterogeneity challenges cultural stereotypes', level=2)
d2 = (
    f'Despite Japan\u2019s well-documented cultural stoicism,{cite(5,6)} '
    f'we found 1.97-fold variation in acute perioperative analgesic prescribing. '
    f'This parallels Cohen and Nisbett\u2019s finding that the US \u201cculture of '
    f'honour\u201d produces regional behavioural differences within a single '
    f'nation.{cite(7)} Japan\u2019s pain culture is not monolithic; regional '
    f'demographics, healthcare infrastructure, and clinical practices generate '
    f'heterogeneity beneath the surface of a nationally shared cultural norm.'
)
p = doc.add_paragraph()
add_ref_runs(p, d2)
disc_parts.append(d2)

add_heading_text('Implications for clinical practice', level=2)
d3 = (
    f'The 1.97-fold within-Japan variation has direct clinical relevance for any physician '
    f'who writes an analgesic prescription\u2014not only pain specialists but also surgeons, '
    f'internists, and general practitioners. A large body of evidence demonstrates that '
    f'cultural stereotypes influence clinician pain assessment and prescribing behaviour. '
    f'Anderson et al showed that racial and ethnic minorities in the United States consistently '
    f'receive less adequate pain management across acute, chronic, cancer, and palliative '
    f'settings.{cite(17)} Campbell and Edwards identified that clinician expectations about '
    f'a patient\u2019s cultural pain behaviour can lead to systematic under- or '
    f'over-treatment.{cite(18)} Rogger et al emphasised that cultural framing affects not '
    f'only patient reporting but also how clinicians interpret and respond to pain cues.{cite(2)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3)
disc_parts.append(d3)

d3b = (
    f'The broader lesson is that within-country heterogeneity in pain prescribing '
    f'challenges the use of national cultural stereotypes in clinical practice\u2014'
    f'not only in Japan but in any country where cultural generalisations guide prescribing '
    f'decisions. The nihonjinron discourse (theories of Japanese '
    f'uniqueness) has long promoted the notion that Japanese people constitute a uniform '
    f'population sharing a single set of behavioural norms.{cite(19)} Yet this '
    f'\u201chegemony of homogeneity,\u201d as Befu termed it, is an ideological construct '
    f'rather than an empirical fact. Burgess showed that the \u201cillusion\u201d of '
    f'homogeneous Japan has tangible consequences for social policy and public '
    f'perception.{cite(20)} In the clinical context, the combination of two '
    f'stereotypes\u2014\u201cJapanese are stoic\u201d and \u201cJapanese are '
    f'homogeneous\u201d\u2014creates a doubly misleading assumption: that all Japanese '
    f'patients will tolerate pain equally and require less analgesia. '
    f'Our finding of 1.97-fold within-Japan variation directly refutes this assumption. '
    f'If this degree of heterogeneity exists within a society widely regarded as culturally '
    f'uniform, the same fragility of cultural labels likely applies to every other '
    f'national-level pain stereotype.'
)
p = doc.add_paragraph()
add_ref_runs(p, d3b)
disc_parts.append(d3b)

d3c = (
    f'Put simply, there is no such entity as \u201cthe Japanese patient\u201d whose pain '
    f'behaviour can be predicted from nationality alone\u2014there are only individual patients '
    f'from 47 diverse prefectures, each with distinct demographic profiles, clinical environments, '
    f'and pain-related prescribing cultures. '
    f'The revised IASP definition describes pain as inherently '
    f'subjective,{cite(21)} and no cultural label can substitute for direct measurement '
    f'of a patient\u2019s nociceptive state. Objective nociception monitoring, such as '
    f'normalised pulse volume, may help standardise perioperative assessment regardless of '
    f'cultural background.{cite(22)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3c)
disc_parts.append(d3c)

add_heading_text('Confounders explain Tohoku\u2019s apparent excess', level=2)
d4 = (
    f'The most methodologically important finding is that the dramatic regional variation '
    f'in neuropathic pain prescribing (unadjusted d={unadj_d:.2f} for Tohoku vs rest) '
    f'was largely explained by confounding disease proxies. '
    f'Diabetes drug prescribing alone correlated at r=0.87 with neuropathic pain '
    f'prescribing, reflecting the known high prevalence of diabetic neuropathy requiring '
    f'gabapentinoids. After adjustment, the Tohoku effect was attenuated by {attenuation:.0f}% '
    f'and became nonsignificant. '
    f'This has important implications for ecological pain research: studies using neuropathic '
    f'pain drug prescribing as a population-level CPSP proxy must account for confounding '
    f'diseases. Without such adjustment, regional differences in diabetes prevalence could be '
    f'misinterpreted as differences in CPSP. The within-database confounder adjustment '
    f'demonstrated here\u2014using disease-specific drug proxies from the same data '
    f'source\u2014provides a replicable framework for other countries with national claims '
    f'databases.'
)
doc.add_paragraph(d4)
disc_parts.append(d4)

add_heading_text('A population-level acute\u2013chronic pain continuum', level=2)
d5 = (
    f'The positive correlation between Phase 1 (acute) and Phase 2 (chronic, adjusted) '
    f'indices (r=0.29, P=0.052) suggests a modest link between regional acute pain '
    f'management intensity and subsequent chronic pain-related prescribing. '
    f'While ecological correlations cannot establish causation, this finding is consistent '
    f'with individual-level evidence that the intensity of acute postoperative pain is a '
    f'risk factor for CPSP.{cite(23)} '
    f'The significant contribution of the acute pain index in Model 5 '
    f'(\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}), even after confounder adjustment, '
    f'supports the hypothesis that regional patterns of acute pain management have downstream '
    f'consequences for chronic pain burden.'
)
p = doc.add_paragraph()
add_ref_runs(p, d5)
disc_parts.append(d5)

add_heading_text('Regional stigma in an era of mobility', level=2)
d_mob = (
    'The premise of region-based stereotyping is increasingly untenable. '
    'A patient presenting in a Singapore hospital may have been born in Japan, '
    'raised in Australia, and working across Southeast Asia\u2014rendering any assumption '
    'based on nationality meaningless. Medicine is an encounter between a clinician '
    'and an individual patient, not a cultural category. Individualised assessment '
    'provides the only reliable basis for analgesic decision-making in a mobile, '
    'multicultural world.'
)
doc.add_paragraph(d_mob)
disc_parts.append(d_mob)

add_heading_text('Strengths and limitations', level=2)
d6a = (
    'Strengths of this study include the use of population-complete data covering all '
    'insurance-reimbursed healthcare in Japan, the novel integration of acute and chronic '
    'pain proxies within a single analytical framework, the transparent '
    'confounder-adjustment methodology, and the use of a perioperative design that '
    'neutralises healthcare access as a confounder (all Phase 1 patients are inpatients '
    'by definition).'
)
doc.add_paragraph(d6a)
disc_parts.append(d6a)

d6b = (
    'The main limitations are inherent to the ecological design. '
    'The unit of analysis is the prefecture, not the individual patient; ecological '
    'correlations may not reflect individual-level associations (ecological fallacy). '
    'NDB Open Data lack diagnosis codes, so the neuropathic pain drug proxy '
    'captures all indications, not CPSP specifically. '
    'Unmeasured confounders such as surgical case mix, physician density, and '
    'regional prescribing culture may contribute to residual variation.'
)
doc.add_paragraph(d6b)
disc_parts.append(d6b)

add_heading_text('Conclusion', level=2)
d7 = (
    'Cultural labels do not predict analgesic need. Despite Japan\u2019s culturally ingrained '
    'norm of pain endurance, perioperative and chronic pain-related prescribing varies up to '
    f'1.97-fold across prefectures. Any clinician who adjusts analgesic dosing based on the '
    f'label \u201cJapanese patient\u201d risks both over- and under-treatment. '
    f'Individualised pain assessment should replace culturally stereotyped assumptions.'
)
doc.add_paragraph(d7)
disc_parts.append(d7)

disc_total = sum(wc(t) for t in disc_parts)
print(f'Discussion word count: {disc_total}')

# ============================================================
# CONTRIBUTORS
# ============================================================
doc.add_paragraph()
add_heading_text('Contributors', level=1)
doc.add_paragraph(
    'TO conceived and designed the study, obtained and verified the underlying data, '
    'performed all analyses, created visualisations, and wrote the manuscript. '
    'TO had full access to all data and is the guarantor.')

# ============================================================
# FUNDING
# ============================================================
doc.add_paragraph()
add_heading_text('Funding', level=1)
doc.add_paragraph('None.')

# ============================================================
# COMPETING INTERESTS
# ============================================================
doc.add_paragraph()
add_heading_text('Competing interests', level=1)
doc.add_paragraph('None declared.')

# ============================================================
# ETHICS APPROVAL
# ============================================================
doc.add_paragraph()
add_heading_text('Ethics approval', level=1)
doc.add_paragraph(
    'This study used only publicly available aggregate data from the NDB Open Data. '
    'Ethical approval was not required under Japan\u2019s Ethical Guidelines for '
    'Medical and Biological Research Involving Human Subjects.')

# ============================================================
# DATA AVAILABILITY STATEMENT
# ============================================================
doc.add_paragraph()
add_heading_text('Data availability statement', level=1)
doc.add_paragraph(
    'The NDB Open Data used in this study are publicly available from the Ministry of Health, '
    'Labour and Welfare website '
    '(https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html). '
    'Analysis code is available at '
    'https://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japan.')

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
doc.add_paragraph()
add_heading_text('Acknowledgements', level=1)
doc.add_paragraph(
    'The author thanks the Ministry of Health, Labour and Welfare for making '
    'the NDB Open Data publicly available. '
    'Parts of data processing and manuscript preparation were assisted by generative AI '
    '(Claude, Anthropic). The author takes full responsibility for the accuracy and content '
    'of the manuscript.')

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_text('References', level=1)

for i, ref_text in enumerate(ref_list, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i} ')
    run_num.bold = True
    p.add_run(ref_text)

# ============================================================
# FIGURE LEGENDS (at end per BMJ guidelines)
# ============================================================
doc.add_page_break()
add_heading_text('Figure legends', level=1)

legends = [
    ('Figure 1.', 'Outpatient neuropathic pain drug prescribing per surgery by prefecture '
     '(unadjusted). Tohoku prefectures (red) cluster at the high end. '
     'Dashed line = national mean.'),
    ('Figure 2.', 'Correlation between neuropathic pain prescribing and confounder disease '
     'proxies. Each dot represents one prefecture. Tohoku prefectures are marked with red '
     'borders. Diabetes drugs show the strongest correlation (r=0.87).'),
    ('Figure 3.', 'Regional comparison of neuropathic pain prescribing: (a) unadjusted and '
     '(b) after confounder adjustment. Tohoku (red) shifts from the highest region to '
     'mid-range after adjustment. Error bars = SD.'),
]

for label, text in legends:
    p = doc.add_paragraph()
    r = p.add_run(label + ' ')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

# ============================================================
# SAVE
# ============================================================
outpath = os.path.join(JECH_DIR, 'JECH_manuscript_EN.docx')
doc.save(outpath)
print(f'\nSaved: {outpath}')

body_total = intro_total + methods_total + results_total + disc_total
print(f'\nBody word count: {body_total} (JECH target: ~3,000)')
print(f'Abstract word count: {abstract_total} (limit: 250)')
print(f'References: {len(ref_list)} (limit: 40)')
print(f'Display items: 3 figures + 2 tables = 5')

# Verification
print('\n--- Vancouver reference order verification ---')
for i, ref_text in enumerate(ref_list, 1):
    print(f'  [{i}] {ref_text[:70]}...')
