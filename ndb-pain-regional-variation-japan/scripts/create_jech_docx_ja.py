#!/usr/bin/env python3
"""Create Japanese manuscript for JECH submission.

Same content as EN version but with all text in Japanese.
JECH format: structured abstract, "What is already known" / "What this study adds" boxes.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re
import csv
import numpy as np
from collections import defaultdict

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JECH_DIR = os.path.join(OUTPUT_DIR, 'jech')
os.makedirs(JECH_DIR, exist_ok=True)

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'scr_summary.json'), 'r') as f:
    scr = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

ref_list = [
    'Callister LC. Cultural influences on pain perceptions and behaviors. Home Health Care Manag Pract 2003;15:207\u201311.',
    'Rogger R, Bello C, Romero CS, et al. Cultural framing and the impact on acute pain and pain services. Curr Pain Headache Rep 2023;27:429\u201336.',
    'Zborowski M. People in Pain. San Francisco: Jossey-Bass, 1969.',
    'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: a review of cross-cultural studies. Int J Sci Res Arch 2024;11:2550\u20136.',
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences between Japanese and Euro-Americans. Eur J Pain 2005;9:389\u201393.',
    'Feng Y, Herdman M, van Nooten F, et al. An exploration of differences between Japan and two European countries in the self-reporting and valuation of pain and discomfort on the EQ-5D. Qual Life Res 2017;26:2067\u201378.',
    'Cohen D, Nisbett RE, Bowdle BF, et al. Insult, aggression, and the southern culture of honor. J Pers Soc Psychol 1996;70:945\u201360.',
    'Kumagai S. メディアが再生産する東北のイメージ. ことば 2020;41:21\u201338.',
    'Takeda K, Yarimizu K. 痛み表現「うずく」の地域差. 国語研プロジェクトレビュー 2016;10:85\u2013107.',
    'Pfizer Japan Inc. 47都道府県 慢性疼痛調査. https://www.pfizer.co.jp/pfizer/company/press/2017 (2025年2月1日アクセス).',
    '厚生労働省. NDBオープンデータ 第10回. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html (2025年1月15日アクセス).',
    'Wakaizumi K, Tanaka C, Shinohara Y, et al. Geographical variation in high-impact chronic pain. Front Public Health 2024;12:1482177.',
    'Matsuoka Y, Morishima T, Sato A, et al. Population-based claims study of regional differences in opioid prescribing. Jpn J Clin Oncol 2025;55:hyaf149.',
    'Taira K, Mori T, Ishimaru M, et al. Regional inequality in dental care utilization in Japan. Lancet Reg Health West Pac 2021;12:100170.',
    'von Elm E, Altman DG, Egger M, et al. The STROBE statement. Lancet 2007;370:1453\u20137.',
    'Benchimol EI, Smeeth L, Guttmann A, et al. The RECORD statement. PLoS Med 2015;12:e1001885.',
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain. J Pain 2009;10:1187\u2013204.',
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. Pain Manag 2012;2:219\u201330.',
    'Befu H. Hegemony of Homogeneity. Melbourne: Trans Pacific Press, 2001.',
    'Burgess C. The "illusion" of homogeneous Japan. Asia Pac J 2010;8(9):1\u201322.',
    'Raja SN, Carr DB, Cohen M, et al. The revised IASP definition of pain. Pain 2020;161:1976\u201382.',
    'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor. F1000Research 2024;13:233.',
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain. Lancet 2006;367:1618\u201325.',
]


def cite(*nums):
    return '{' + ','.join(str(n) for n in nums) + '}'


def add_ref_runs(p, text):
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(10)
        else:
            p.add_run(part)


doc = Document()
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Yu Mincho'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 2.0

for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fc1)
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = ' PAGE '
    run._r.append(it)
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc2)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def add_inline_figure(caption_text, fig_num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    r0 = p.add_run(f'[図{fig_num}をここに挿入]')
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f'図{fig_num}. ')
    r.bold = True
    r.font.size = Pt(10)
    cap.add_run(caption_text).font.size = Pt(10)


unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100

scr_analgesic_range = scr['analgesic_inpatient']['scr_range']
scr_analgesic_ratio = scr['analgesic_inpatient']['variation_ratio']
scr_neuro_range = scr['neuropathic_outpatient']['scr_range']
scr_neuro_ratio = scr['neuropathic_outpatient']['variation_ratio']
scr_neuro_tohoku = scr['neuropathic_outpatient']['scr_tohoku_mean']
scr_neuro_non_tohoku = scr['neuropathic_outpatient']['scr_non_tohoku_mean']

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

# ============================================================
# タイトルページ
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    '文化的ラベルは鎮痛薬の必要性を予測するか？\n'
    '日本47都道府県における疼痛関連処方の地域差に関する生態学的研究'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run('大西 竜希').font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('麻酔科学教室, [所属機関], [住所], [市], [郵便番号], 日本')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('責任著者: ')
r.bold = True
p.add_run('大西 竜希, 麻酔科学教室, [所属機関], [住所]. E-mail: [email]')

doc.add_page_break()

# ============================================================
# 抄録（背景、方法、結果、結論）
# ============================================================
add_heading_text('抄録', level=1)

abs_sections = [
    ('背景:', (
        '臨床医は鎮痛薬の必要量を推定する際、文化的ラベルに頻繁に依拠する。'
        'しかし疼痛関連処方が日本国内で地域差を示すかは集団レベルで検証されていない。'
    )),
    ('方法:', (
        f'本生態学的研究は、NDBオープンデータ第10回（2023年4月〜2024年3月、全保険請求）を用い、'
        f'47都道府県の疼痛関連処方を分析した。Phase 1は入院周術期鎮痛薬処方を手術件数で標準化、'
        f'Phase 2は外来神経障害性疼痛薬処方を慢性術後痛プロキシとし交絡疾患で調整した。'
        f'標準化請求比（SCR）で年齢性別標準化後の所見を確認した。'
    )),
    ('結果:', (
        f'急性鎮痛薬処方は都道府県間で1.97倍変動した（Kruskal–Wallis P<0.001）。'
        f'伝統的に最も忍耐強いとされる東北は鎮痛薬処方が多かった（Cohen\'s d=0.87）。'
        f'未調整の神経障害性疼痛処方に大きな東北超過（d={unadj_d:.2f}）があったが、'
        f'糖尿病薬処方が強い交絡因子であり（r=0.87）、調整後は{attenuation:.0f}%減弱した'
        f'（P={reg["adjusted_cpsp_test"]["p_value"]:.2f}）。'
        f'SCRは年齢性別標準化後も{scr_analgesic_ratio:.2f}倍の変動を確認した。'
    )),
    ('結論:', (
        '国内約2倍の変動は文化的ラベルが鎮痛薬必要性を予測しないことを示す。'
        '交絡疾患が見かけの地域パターンを大きく修飾した。'
        '臨床医は文化的仮定ではなく個別評価に基づく処方を行うべきである。'
    )),
]

for label, text in abs_sections:
    p = doc.add_paragraph()
    r = p.add_run(f'{label} ')
    r.bold = True
    p.add_run(text)

doc.add_page_break()

# ============================================================
# このトピックについてすでに知られていること / 本研究が加えること
# ============================================================
add_heading_text('このトピックについてすでに知られていること', level=2)
known_items = [
    '異文化間研究は日本人を痛みに忍耐強いと特徴づけてきたが、'
    'すべての先行比較は日本を均質な単位として扱っている。',
    '日本では慢性疼痛有病率やがん性オピオイド処方の地域変動が報告されているが、'
    '周術期鎮痛薬処方の集団レベルでの検討はない。',
    '文化的ステレオタイプは臨床医の疼痛評価と処方行動に影響し、'
    '系統的な治療不足・過剰につながりうる。',
]
for item in known_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_text('本研究が加えること', level=2)
adds_items = [
    '急性周術期鎮痛薬処方は47都道府県間で1.97倍変動し、'
    '画一的な忍耐文化という概念を否定する。',
    '文化的に「最も忍耐強い」とされる東北は鎮痛薬処方が多く、'
    'この見かけの超過は交絡疾患有病率で大部分説明される。',
    '国内の処方異質性は、臨床上の意思決定における国レベルの'
    '文化的ステレオタイプの妥当性を問う。',
]
for item in adds_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_text('本研究が研究・実践・政策にどのような影響を与えうるか', level=2)
policy_items = [
    '臨床医は鎮痛薬の意思決定を痛み耐性に関する文化的仮定ではなく個別評価に基づくべきである'
    '――国内約2倍の変動は、国レベルのステレオタイプが処方必要性を予測しないことを示す。',
    '神経障害性疼痛薬処方を慢性疼痛の代理指標とする生態学的研究では、'
    '交絡疾患（特に糖尿病）の調整が不可欠である。調整なしでは疾患有病率の地域差が'
    '疼痛行動の差異と誤解される可能性がある。',
    '本研究で実証したデータベース内交絡調整フレームワークは、'
    '全国レセプトデータベースを有する他国でも再現可能であり、'
    '今後の集団レベルの疼痛研究に資する。',
]
for item in policy_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 緒言
# ============================================================
add_heading_text('緒言', level=1)

intro1 = (
    f'臨床医は鎮痛薬処方において、文化的ラベルに基づく仮定を日常的に行う。{cite(1,2)}'
    f'Zborowski以来、{cite(3)}文化的規範が痛みの報告と鎮痛薬使用に影響することが'
    f'確立されている。{cite(4)}Hobaraは日本人がEuro-Americansより痛み行動を不適切と'
    f'評価することを示し、{cite(5)}Feng et alはEQ-5Dにおいて日本人が痛みを避ける'
    f'ための時間トレードオフに消極的であることを示した。{cite(6)}'
    f'「日本人患者は忍耐強い」というラベルは、忍耐的態度を低い鎮痛需要と解釈する'
    f'リスクを伴う。'
)
p = doc.add_paragraph()
add_ref_runs(p, intro1)

intro2 = (
    f'ほとんどの研究は各国を均質な単位として扱う。しかし国内にも異質性が存在する。'
    f'CohenとNisbettは米国南部の「名誉の文化」が北部と異なる行動反応を生むことを'
    f'実証した。{cite(7)}日本でも東北は忍耐の体現と認識され、{cite(8)}'
    f'痛みの方言「うずく」にも地域差がある。{cite(9)}'
    f'ファイザー調査では慢性疼痛患者の「痛みを我慢」割合が48.7%〜81.6%であった。{cite(10)}'
)
p = doc.add_paragraph()
add_ref_runs(p, intro2)

intro3 = (
    f'日本の国民皆保険制度とNDB{cite(11)}は地域処方変動の生態学的分析に理想的である。'
    f'慢性疼痛有病率に1.6倍、{cite(12)}がん性オピオイド処方に4倍の変動{cite(13)}が'
    f'報告されている。Taira et alはNDB由来SCRで歯科受療格差を実証した。{cite(14)}'
    f'しかし疼痛関連処方にこの枠組みを適用した研究はない。'
)
p = doc.add_paragraph()
add_ref_runs(p, intro3)

doc.add_paragraph(
    '本研究の目的は、(1)47都道府県における急性周術期鎮痛薬処方の地域変動をマッピングし、'
    '(2)外来神経障害性疼痛薬処方を交絡疾患調整後のCPSPプロキシとして検討し、'
    '(3)急性・慢性疼痛所見を集団レベルで統合することである。'
)

# ============================================================
# 方法
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインと報告', level=2)
m1 = (
    f'本生態学的研究はNDBオープンデータの都道府県レベル集計データを分析した。'
    f'STROBE声明{cite(15)}およびRECORD拡張{cite(16)}に従い報告する。'
    f'公開集計データのみを使用したため、「人を対象とする生命科学・医学系研究に関する'
    f'倫理指針」に基づき倫理審査は不要であった。'
)
p = doc.add_paragraph()
add_ref_runs(p, m1)

add_heading_text('データソース', level=2)
m2 = (
    f'NDBオープンデータ第10回（2023年4月〜2024年3月）を使用した。{cite(11)}'
    f'NDBは全保険者の請求データを収集し約1億2500万人を網羅する。'
    f'SCR算出に総務省統計局の都道府県別人口推計（2023年10月）を用いた。'
)
p = doc.add_paragraph()
add_ref_runs(p, m2)

add_heading_text('地域分類', level=2)
doc.add_paragraph(
    '都道府県は9地方ブロックに分類した：北海道（1）、東北（6）、関東（7）、'
    '北陸・甲信越（6）、東海（4）、近畿（6）、中国（5）、四国（4）、九州・沖縄（8）。'
    '東北は我慢の文化的特徴づけに基づき先験的に主要関心地域とした。'
)

add_heading_text('Phase 1: 急性周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '入院処方データから3種の鎮痛薬分類を抽出した：'
    '114類（解熱鎮痛消炎薬）、811類（アヘンアルカロイド系麻薬）、821類（合成麻薬）。'
    '鎮痛薬／手術指数を各都道府県で算出した。'
)

add_heading_text('Phase 2: 外来神経障害性疼痛薬処方', level=2)
doc.add_paragraph(
    '外来内服の神経障害性疼痛薬5種を抽出した：プレガバリン、ミロガバリン、'
    'デュロキセチン、トラマドール、ノイロトロピン。処方／手術指数と人口当たり率を算出した。'
)

add_heading_text('交絡疾患プロキシ', level=2)
doc.add_paragraph(
    '4種の交絡プロキシを抽出した：経口血糖降下薬（糖尿病性神経障害）、'
    '帯状疱疹抗ウイルス薬（帯状疱疹後神経痛）、デュロキセチンを除く抗うつ薬（うつ病）、'
    '抗不安薬（不安障害）。神経ブロック処置件数を追加プロキシとした。'
)

add_heading_text('統計解析', level=2)
m7 = (
    f'Phase 1はKruskal–Wallis検定（Bonferroni補正付きMann–Whitney U検定）で評価した。'
    f'Phase 2では5つの回帰モデルで東北効果を検討した。'
    f'調整後CPSP指数は4交絡プロキシへの回帰残差として算出した。'
    f'Taira et al{cite(14)}の方法でSCRを算出した。Python 3.11で全解析を実施した。'
)
p = doc.add_paragraph()
add_ref_runs(p, m7)

add_heading_text('患者・市民参画', level=2)
doc.add_paragraph('患者または市民は本研究の設計、実施、報告、普及計画に関与していない。')

# ============================================================
# 結果
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('Phase 1: 急性周術期鎮痛薬処方の地域変動', level=2)
doc.add_paragraph(
    '2023年4月〜2024年3月に、NDBは47都道府県で7,903,515件の入院手術と'
    '274,579,851単位の鎮痛薬処方を記録した。全国平均の鎮痛薬／手術指数は35.78（SD 5.56）で、'
    '岐阜25.20〜鹿児島49.75の1.97倍の差を示した（9地方間Kruskal–Wallis P<0.001；表1）。'
)

# === 表1 ===
p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(14)
r_c = p_cap.add_run('表1. ')
r_c.bold = True
r_c.font.size = Pt(10)
p_cap.add_run(
    '9地方ブロック別の入院鎮痛薬処方／手術指数。Kruskal–Wallis P<0.001。'
).font.size = Pt(10)

t1 = doc.add_table(rows=1 + len(REGION_ORDER), cols=4, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['地方', '都道府県数', '平均 \u00b1 SD', '範囲']):
    t1.rows[0].cells[i].text = h
    for run in t1.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = t1.rows[idx + 1].cells
    row[0].text = reg_name
    row[1].text = str(len(vals))
    row[2].text = (
        f'{np.mean(vals):.2f} \u00b1 {np.std(vals, ddof=1):.2f}'
        if len(vals) > 1
        else f'{np.mean(vals):.2f}'
    )
    row[3].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)
set_table_borders(t1)
doc.add_paragraph()

doc.add_paragraph(
    '東北は9地方中7位で平均指数39.97（SD 3.53）と非東北平均35.17を有意に上回った'
    '（Mann–Whitney U=190, P=0.031; Cohen\'s d=0.87）。'
    'NSAIDs（P=0.044）、オピオイドアルカロイド（P=0.003）、合成麻薬（P=0.001）で一貫した。'
)

add_heading_text('Phase 2: 外来神経障害性疼痛薬処方', level=2)
doc.add_paragraph(
    f'東北の処方指数は顕著に高く'
    f'（{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; P<0.001; '
    f'd={unadj_d:.2f}；図1）。糖尿病薬処方が最強の交絡因子であった（r=0.87；図2）。'
)

add_inline_figure(
    '都道府県別の外来神経障害性疼痛薬処方／手術指数（未調整）。東北（赤）が高値に集中。', 1
)

doc.add_paragraph(
    f'4交絡因子調整後、東北効果は減弱し非有意となった'
    f'（\u03b2={reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
    f'P={reg["model2_adjusted"]["tohoku_p"]:.2f}; '
    f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}; 表2）。'
)

add_inline_figure(
    '神経障害性疼痛処方と交絡疾患プロキシの相関。糖尿病薬が最強の相関（r=0.87）。', 2
)

# === 表2 ===
p_cap2 = doc.add_paragraph()
p_cap2.paragraph_format.space_before = Pt(14)
r_c2 = p_cap2.add_run('表2. ')
r_c2.bold = True
r_c2.font.size = Pt(10)
p_cap2.add_run('外来神経障害性疼痛処方の回帰モデル。').font.size = Pt(10)

models = [
    ('モデル1（未調整）', '\u2014', '\u2014', f'd={unadj_d:.2f}'),
    ('モデル2（全交絡）', f'{reg["model2_adjusted"]["tohoku_coef"]:.1f}',
     f'{reg["model2_adjusted"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}'),
    ('モデル3（コア神経障害薬）', f'{reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}',
     f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model3_core_neuropathic"]["R2"]:.3f}'),
    ('モデル4（神経ブロック）', f'{reg["model4_nerve_blocks"]["tohoku_coef"]:.1f}',
     f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model4_nerve_blocks"]["R2"]:.3f}'),
    ('モデル5（統合）', f'{reg["model5_integrated"]["tohoku_coef"]:.1f}',
     f'{reg["model5_integrated"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model5_integrated"]["R2"]:.3f}'),
]
t2 = doc.add_table(rows=1 + len(models), cols=4, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['モデル', '東北 \u03b2', '東北 P値', '適合度']):
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
for idx, (name, beta, pval, fit) in enumerate(models):
    row = t2.rows[idx + 1].cells
    row[0].text = name
    row[1].text = beta
    row[2].text = pval
    row[3].text = fit
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)
set_table_borders(t2)
doc.add_paragraph()

doc.add_paragraph(
    f'調整後CPSP指数は未調整と異なるパターンを示し、東北の超過は非有意となった'
    f'（P={reg["adjusted_cpsp_test"]["p_value"]:.2f}; d={adj_d:.2f}; 図3）。'
    f'年齢性別標準化後、鎮痛薬SCRは{scr_analgesic_range[0]:.1f}〜{scr_analgesic_range[1]:.1f}'
    f'（{scr_analgesic_ratio:.2f}倍）であった。'
)

add_inline_figure(
    '地方別神経障害性疼痛処方：(a)未調整、(b)交絡調整後。東北（赤）は調整後に中位に移動。', 3
)

# ============================================================
# 考察
# ============================================================
add_heading_text('考察', level=1)

doc.add_paragraph(
    '本研究はNDBオープンデータを用いて日本全47都道府県の周術期・慢性疼痛関連処方を'
    '初めてマッピングした。文化的ラベルは鎮痛薬の必要性を予測しない。'
)

add_heading_text('国内異質性が文化的ステレオタイプを問う', level=2)
d2 = (
    f'日本の文化的忍耐にもかかわらず、{cite(5,6)}'
    f'急性鎮痛薬処方に1.97倍の変動を見出した。'
    f'CohenとNisbettが米国内の地域差を示したように、{cite(7)}'
    f'日本の疼痛文化も画一的ではない。'
)
p = doc.add_paragraph()
add_ref_runs(p, d2)

add_heading_text('臨床実践への示唆', level=2)
d3 = (
    f'Anderson et alは民族的少数者の疼痛管理不足を示し、{cite(17)}'
    f'Campbell and Edwardsは文化的期待による系統的誤管理を同定した。{cite(18)}'
    f'日本人論による均質性の「覇権」{cite(19)}とBurgessの均質日本の「幻想」{cite(20)}は、'
    f'「日本人は忍耐強く均質」という二重のステレオタイプを生む。'
    f'1.97倍の変動はこれを直接否定する。'
    f'IASPの痛みの定義は本質的に主観的であり、{cite(21)}'
    f'客観的侵害受容モニタリングが評価を標準化しうる。{cite(22)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3)

add_heading_text('交絡因子が東北の見かけの超過を説明する', level=2)
doc.add_paragraph(
    f'神経障害性疼痛処方の東北超過（未調整d={unadj_d:.2f}）は交絡疾患で大部分説明された。'
    f'糖尿病薬でr=0.87の相関を示し、調整後に東北効果は{attenuation:.0f}%減弱し非有意となった。'
)

add_heading_text('モビリティの時代における地域的スティグマ', level=2)
doc.add_paragraph(
    '地域に基づくステレオタイピングの前提は現代においてますます維持できない。'
    '医療は臨床医と個人の患者との出会いであり、文化的カテゴリーとの出会いではない。'
    '個別化された評価こそが鎮痛薬処方の唯一の信頼できる基盤である。'
)

add_heading_text('強みと限界', level=2)
doc.add_paragraph(
    '強みは全保険請求をカバーする全数データ、急性・慢性疼痛プロキシの統合、'
    '透明な交絡調整手法、周術期デザインによる医療アクセス中立化である。'
    '限界は生態学的デザイン固有のもの（分析単位は都道府県、NDBに診断コードなし、'
    '手術ケースミックス等の未測定交絡因子）である。'
)

add_heading_text('結論', level=2)
doc.add_paragraph(
    '文化的ラベルは鎮痛薬の必要性を予測しない。周術期・慢性疼痛関連処方は都道府県間で'
    '1.97倍変動する。「日本人患者」というラベルで処方を調整することは過剰・過少投与の'
    'リスクを伴う。個別化された疼痛評価が文化的ステレオタイプに基づく仮定に取って代わるべきである。'
)

# ============================================================
# 末尾セクション
# ============================================================
doc.add_paragraph()
add_heading_text('著者貢献', level=1)
doc.add_paragraph(
    'TO: 研究の着想・設計、データ取得・検証、全解析の実施、可視化、原稿執筆。'
    'TOは全データへのフルアクセスを有し、保証人である。'
)

doc.add_paragraph()
add_heading_text('資金', level=1)
doc.add_paragraph('該当なし。')

doc.add_paragraph()
add_heading_text('利益相反', level=1)
doc.add_paragraph('申告すべき利益相反はない。')

doc.add_paragraph()
add_heading_text('倫理審査', level=1)
doc.add_paragraph(
    '本研究はNDBオープンデータ（公開集計データ）のみを使用した。'
    '「人を対象とする生命科学・医学系研究に関する倫理指針」に基づき倫理審査は不要であった。'
)

doc.add_paragraph()
add_heading_text('データ利用可能性', level=1)
doc.add_paragraph(
    'NDBオープンデータは厚生労働省ウェブサイト'
    '（https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html）から公開。'
    '解析コード: https://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japan'
)

doc.add_paragraph()
add_heading_text('謝辞', level=1)
doc.add_paragraph(
    'NDBオープンデータを公開している厚生労働省に感謝する。'
    'データ処理と原稿作成の一部に生成AI（Claude, Anthropic）を使用した。'
    '著者は原稿の正確性と内容に全責任を負う。'
)

doc.add_page_break()
add_heading_text('参考文献', level=1)
for i, ref_text in enumerate(ref_list, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i} ')
    run_num.bold = True
    p.add_run(ref_text)

outpath = os.path.join(JECH_DIR, 'JECH_manuscript_JA.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
