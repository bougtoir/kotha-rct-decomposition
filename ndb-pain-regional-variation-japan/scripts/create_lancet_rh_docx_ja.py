#!/usr/bin/env python3
"""Create Japanese manuscript for Lancet Regional Health – Western Pacific.

Same structure as EN version but with all text in Japanese.
Figures/tables embedded INLINE (using JA figure versions).
Vancouver-style superscript numbered references.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re
import csv
import numpy as np
from collections import defaultdict

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
LRH_DIR = os.path.join(OUTPUT_DIR, 'lancet_rh_wp')
os.makedirs(LRH_DIR, exist_ok=True)
FIG_DIR = OUTPUT_DIR

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'scr_summary.json'), 'r') as f:
    scr = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

ref_list = [
    'Callister LC. Cultural influences on pain perceptions and behaviors. Home Health Care Manag Pract 2003; 15: 207\u201311.',
    'Rogger R, Bello C, Romero CS, et al. Cultural framing and the impact on acute pain and pain services. Curr Pain Headache Rep 2023; 27: 429\u201336.',
    'Zborowski M. People in Pain. San Francisco: Jossey-Bass, 1969.',
    'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: a review of cross-cultural studies. Int J Sci Res Arch 2024; 11: 2550\u20136.',
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences between Japanese and Euro-Americans. Eur J Pain 2005; 9: 389\u201393.',
    'Feng Y, Herdman M, van Nooten F, et al. An exploration of differences between Japan and two European countries in the self-reporting and valuation of pain and discomfort on the EQ-5D. Qual Life Res 2017; 26: 2067\u201378.',
    'Cohen D, Nisbett RE, Bowdle BF, Schwarz N. Insult, aggression, and the southern culture of honor. J Pers Soc Psychol 1996; 70: 945\u201360.',
    'Kumagai S. メディアが再生産する東北のイメージ. ことば 2020; 41: 21\u201338.',
    'Takeda K, Yarimizu K. 痛み表現「うずく」の地域差. 国語研プロジェクトレビュー 2016; 10: 85\u2013107.',
    'Pfizer Japan Inc. 47都道府県 慢性疼痛調査. https://www.pfizer.co.jp/pfizer/company/press/2017 (2025年2月1日アクセス).',
    '厚生労働省. NDBオープンデータ 第10回. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html (2025年1月15日アクセス).',
    'Wakaizumi K, Tanaka C, Shinohara Y, et al. Geographical variation in high-impact chronic pain. Front Public Health 2024; 12: 1482177.',
    'Matsuoka Y, Morishima T, Sato A, et al. Population-based claims study of regional differences in opioid prescribing. Jpn J Clin Oncol 2025; 55: hyaf149.',
    'Taira K, Mori T, Ishimaru M, et al. Regional inequality in dental care utilization in Japan. Lancet Reg Health West Pac 2021; 12: 100170.',
    'von Elm E, Altman DG, Egger M, et al. The STROBE statement. Lancet 2007; 370: 1453\u20137.',
    'Benchimol EI, Smeeth L, Guttmann A, et al. The RECORD statement. PLoS Med 2015; 12: e1001885.',
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain. J Pain 2009; 10: 1187\u2013204.',
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. Pain Manag 2012; 2: 219\u201330.',
    'Befu H. Hegemony of Homogeneity. Melbourne: Trans Pacific Press, 2001.',
    'Burgess C. The "illusion" of homogeneous Japan. Asia Pac J 2010; 8(9): 1\u201322.',
    'Raja SN, Carr DB, Cohen M, et al. The revised IASP definition of pain. Pain 2020; 161: 1976\u201382.',
    'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor. F1000Research 2024; 13: 233.',
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain. Lancet 2006; 367: 1618\u201325.',
]

def cite(*nums):
    return '{' + ','.join(str(n) for n in nums) + '}'

def add_ref_runs(p, text):
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(10)
        else:
            p.add_run(part)

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Yu Mincho'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 2.0

for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); run._r.append(fc1)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; run._r.append(it)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end'); run._r.append(fc2)

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)

def add_inline_figure(fig_path, caption_text, fig_num):
    """図のプレースホルダーとキャプションを挿入（図は独立TIFFファイルとして投稿）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    r0 = p.add_run(f'[図{fig_num}をここに挿入]')
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f'図{fig_num}. ')
    r.bold = True
    r.font.size = Pt(10)
    cap.add_run(caption_text).font.size = Pt(10)

unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100

scr_analgesic_range = scr['analgesic_inpatient']['scr_range']
scr_analgesic_ratio = scr['analgesic_inpatient']['variation_ratio']
scr_neuro_range = scr['neuropathic_outpatient']['scr_range']
scr_neuro_ratio = scr['neuropathic_outpatient']['variation_ratio']
scr_neuro_tohoku = scr['neuropathic_outpatient']['scr_tohoku_mean']
scr_neuro_non_tohoku = scr['neuropathic_outpatient']['scr_non_tohoku_mean']

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

# ============================================================
# タイトルページ
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    '文化的ラベルは鎮痛薬の必要性を予測するか？\n'
    '日本47都道府県における疼痛関連処方の地域差に関する生態学的研究'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run('大西 竜希').font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('麻酔科学教室, [所属機関], [住所], [市], [郵便番号], 日本')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('責任著者: ')
r.bold = True
p.add_run('大西 竜希, 麻酔科学教室, [所属機関], [住所]. E-mail: [email]')

doc.add_page_break()

# ============================================================
# 抄録
# ============================================================
add_heading_text('抄録', level=1)

abs_sections = [
    ('背景', (
        '臨床医は鎮痛薬の必要量を推定する際、「日本人患者は我慢強い」などの文化的ラベルに頻繁に依拠する。'
        'しかし疼痛関連処方が日本国内で実際に地域差を示すかは集団レベルで検証されていない。'
        'NDB（レセプト情報・特定健診等情報データベース）オープンデータは、'
        '文化的ラベルが鎮痛薬使用を予測するかを全47都道府県で直接検証することを可能にする。'
    )),
    ('方法', (
        f'本生態学的研究は、NDBオープンデータ第10回（2023年4月〜2024年3月、全保険請求データ）を用い、'
        f'47都道府県・9地方ブロックにおける疼痛関連処方を分析した。'
        f'Phase 1では入院周術期鎮痛薬処方（NSAIDs、麻薬性鎮痛薬）を手術件数で標準化した。'
        f'Phase 2では外来神経障害性疼痛薬処方（プレガバリン、ミロガバリン、デュロキセチン、'
        f'トラマドール、ノイロトロピン）を慢性術後痛（CPSP）のプロキシとし、'
        f'交絡疾患（糖尿病性神経障害、帯状疱疹後神経痛、うつ病、不安障害）で調整した。'
        f'間接年齢性別標準化による標準化請求比（SCR）で人口構成の影響を確認した。'
    )),
    ('結果', (
        f'Phase 1では、鎮痛薬／手術指数は1.97倍の都道府県間変動を示した'
        f'（岐阜25.20〜鹿児島49.75; Kruskal–Wallis P<0.001）。'
        f'伝統的に最も忍耐強い地域とされる東北は、鎮痛薬処方が少ないどころか多かった'
        f'（平均39.97 vs 非東北35.17; Cohen\'s d=0.87）。'
        f'Phase 2では、未調整の神経障害性疼痛処方に大きな東北の超過が見られたが'
        f'（d={unadj_d:.2f}）、交絡疾患調整後、東北効果は{attenuation:.0f}%減弱し'
        f'非有意となった（P={reg["adjusted_cpsp_test"]["p_value"]:.2f}）。'
        f'SCRは鎮痛薬処方に年齢性別標準化後も{scr_analgesic_ratio:.2f}倍の変動を確認した。'
    )),
    ('解釈', (
        '国内で約2倍の疼痛関連処方変動は、文化的ラベルが集団レベルで鎮痛薬の必要性を予測しないことを示す。'
        '外科医、内科医、痛み専門医を問わず、文化的背景ではなく'
        '個別の評価に基づく処方が不可欠である。'
    )),
    ('資金', '該当なし。'),
]

for label, text in abs_sections:
    p = doc.add_paragraph()
    r = p.add_run(f'{label} ')
    r.bold = True
    p.add_run(text)

doc.add_page_break()

# ============================================================
# Research in Context（日本語）
# ============================================================
add_heading_text('Research in Context', level=1)

p = doc.add_paragraph()
r = p.add_run('本研究以前のエビデンス ')
r.bold = True; r.italic = True
p.add_run(
    '2025年1月までにPubMedおよびGoogle Scholarを「pain」AND「regional variation」AND「Japan」'
    'で検索した（言語制限なし）。異文化間研究は日本人を一貫して痛みに忍耐強いと特徴づけてきたが、'
    '国内での地域間比較は行われていなかった。Wakaizumi et al（2024）はインターネット調査で'
    '都道府県間1.6倍の慢性疼痛有病率変動を報告し、Matsuoka et al（2025）はがん性オピオイド'
    '処方に4倍の変動を見出した。全47都道府県を対象とした全数レセプトデータによる'
    '周術期鎮痛薬・神経障害性疼痛薬処方のマッピングは行われていなかった。'
)

p = doc.add_paragraph()
r = p.add_run('本研究の付加価値 ')
r.bold = True; r.italic = True
p.add_run(
    'NDBオープンデータ（約1億2500万人の全被保険者をカバー）を用い、急性周術期鎮痛薬処方に'
    '1.97倍の変動と有意な地域クラスタリングを見出した。伝統的に最も忍耐強いとされる東北は、'
    '鎮痛薬処方が多かった。神経障害性疼痛処方の東北超過は交絡疾患（特に糖尿病）で'
    '大部分が説明された。年齢性別標準化請求比（SCR）でも処方の異質性は同様に確認された。'
)

p = doc.add_paragraph()
r = p.add_run('全エビデンスからの示唆 ')
r.bold = True; r.italic = True
p.add_run(
    '日本国内の疼痛関連処方は均一ではなく、国の文化的ラベルが個人の疼痛行動を'
    '予測するという仮定を否定する。西太平洋地域およびそれ以外の臨床医は、'
    '日本人および他の東アジア患者の疼痛評価において民族的ステレオタイピングを避けるべきである。'
)

doc.add_page_break()

# ============================================================
# 緒言
# ============================================================
add_heading_text('緒言', level=1)

intro1 = (
    f'痛みは普遍的な経験であるが、その表現と管理は文化により大きく形作られる。{cite(1,2)}'
    f'Zborowski以来、{cite(3)}文化的規範が痛みの報告、受療行動、鎮痛薬使用に影響することが'
    f'広く確立されている。{cite(4)}日本人は一貫して痛みに忍耐強いと特徴づけられてきた。'
    f'Hobaraは日本人がEuro-Americansより痛み行動を不適切と評価することを示し、{cite(5)}'
    f'Feng et alはEQ-5Dにおいて日本人が英蘭の回答者より痛みを避けるための時間トレードオフに'
    f'はるかに消極的であることを示した。{cite(6)}この文化的忍耐は我慢（がまん）の概念に'
    f'集約されるが、臨床医が忍耐的な態度を低い鎮痛需要と解釈するリスクを伴う。'
)
p = doc.add_paragraph()
add_ref_runs(p, intro1)

intro2 = (
    f'ほとんどの研究は国家間の痛み行動を比較し、各国を均質な単位として扱う。'
    f'しかし、国内にも実質的な異質性が存在する。CohenとNisbettは米国南部の「名誉の文化」が'
    f'北部と異なる行動反応を生むことを実証した。{cite(7)}'
    f'日本でも地域的アイデンティティは強く、東北は伝統的に忍耐の体現と認識されており、{cite(8)}'
    f'痛みを表す方言「うずく」にも地域差がある。{cite(9)}'
    f'ファイザー日本の調査では、慢性疼痛患者の「痛みを我慢している」割合は'
    f'都道府県間で48.7%〜81.6%の範囲であった。{cite(10)}'
)
p = doc.add_paragraph()
add_ref_runs(p, intro2)

intro3 = (
    f'日本の国民皆保険制度、標準化された薬価、NDB{cite(11)}は地域処方変動の生態学的分析に'
    f'理想的な環境を提供する。近年、慢性疼痛有病率に1.6倍の地域差{cite(12)}や、'
    f'がん性オピオイド処方に4倍の変動{cite(13)}が報告されている。'
    f'Taira et alはNDB由来の標準化請求比を用いて歯科受療の地域格差を実証した。{cite(14)}'
    f'しかし、疼痛関連処方にこの全数生態学的枠組みを適用した研究はない。'
)
p = doc.add_paragraph()
add_ref_runs(p, intro3)

intro4 = (
    f'本研究の目的は、(1)47都道府県における急性周術期鎮痛薬処方の地域変動をマッピングし、'
    f'(2)外来神経障害性疼痛薬処方を交絡疾患調整後のCPSPプロキシとして検討し、'
    f'(3)急性・慢性疼痛所見を集団レベルで統合することである。'
)
p = doc.add_paragraph()
add_ref_runs(p, intro4)

# ============================================================
# 方法
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインと報告', level=2)
m1 = f'本生態学的研究はNDBオープンデータの都道府県レベル集計データを分析した。STROBE声明{cite(15)}およびRECORD拡張{cite(16)}に従い報告する。公開集計データのみを使用したため、倫理審査は不要であった。'
p = doc.add_paragraph()
add_ref_runs(p, m1)

add_heading_text('データソース', level=2)
m2 = f'NDBオープンデータ第10回（2023年4月〜2024年3月）を使用した。{cite(11)}NDBは日本の国民皆保険制度下の全保険者の請求データを収集し、約1億2500万人の被保険者を網羅する。集計データは都道府県レベルで公表され、10件未満のセルは秘匿される。人口当たりの割合および標準化請求比（SCR）の算出には、総務省統計局の都道府県別5歳階級・男女別人口推計（2023年10月）を用いた。'
p = doc.add_paragraph()
add_ref_runs(p, m2)

add_heading_text('地域分類', level=2)
doc.add_paragraph(
    '都道府県は総務省統計局の分類に従い9地方ブロックに分類した：'
    '北海道（1）、東北（6: 青森、岩手、宮城、秋田、山形、福島）、'
    '関東（7）、北陸・甲信越（6）、東海（4）、近畿（6）、中国（5）、四国（4）、'
    '九州・沖縄（8）。東北は我慢の文化的特徴づけに基づき、先験的に主要関心地域とした。'
)

add_heading_text('Phase 1: 急性周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '入院処方データから3種の鎮痛薬分類を抽出した：'
    '114類（解熱鎮痛消炎薬）、811類（アヘンアルカロイド系麻薬）、821類（合成麻薬）。'
    '入院手術件数はK手術区分から抽出した。鎮痛薬／手術指数を各都道府県で算出した。'
    '薬効分類別のサブ解析も実施した。'
)

add_heading_text('Phase 2: 外来神経障害性疼痛薬処方', level=2)
doc.add_paragraph(
    '外来内服の神経障害性疼痛薬5種を抽出した：プレガバリン（78製剤）、'
    'ミロガバリン（8製剤）、デュロキセチン（33製剤）、トラマドール（3製剤）、'
    'ノイロトロピン（1製剤）。神経障害性疼痛処方／手術指数と人口当たり処方率を算出した。'
)

add_heading_text('交絡疾患プロキシ', level=2)
doc.add_paragraph(
    '外来データから4種の交絡疾患プロキシを抽出した：経口血糖降下薬（261製剤；糖尿病性神経障害）、'
    '帯状疱疹抗ウイルス薬（47製剤；帯状疱疹後神経痛）、デュロキセチンを除く抗うつ薬（128製剤；うつ病）、'
    '抗不安薬（112製剤；不安障害）。神経ブロック処置件数（73コード）を独立したCPSPプロキシとした。'
)

add_heading_text('統計解析', level=2)
doc.add_paragraph(
    'Phase 1の地域差はKruskal–Wallis検定（Bonferroni補正付きpost-hoc Mann–Whitney U検定）で評価した。'
    'Phase 2では5つの回帰モデルを適合させた：モデル1（未調整）、'
    'モデル2（全交絡調整）、モデル3（コア神経障害性疼痛薬のみ）、'
    'モデル4（神経ブロック）、モデル5（急性指数＋交絡の統合モデル）。'
    '調整後CPSP指数は4交絡プロキシへの回帰残差として算出した。'
    'Python 3.11（NumPy 1.24, SciPy 1.11）で全解析を実施した。'
)

add_heading_text('年齢性別標準化', level=2)
m_scr = (
    f'Taira et al{cite(14)}の方法に従い、間接年齢性別標準化により標準化請求比（SCR）を算出した。'
    f'NDBの性別・年齢階級別テーブルから全国の年齢性別別処方率（18の5歳階級×2性別）を計算し、'
    f'各都道府県の人口構成に適用した。SCR = (観測値 / 期待値) × 100。'
)
p = doc.add_paragraph()
add_ref_runs(p, m_scr)

# ============================================================
# 結果
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('Phase 1: 急性周術期鎮痛薬処方の地域変動', level=2)
doc.add_paragraph(
    '2023年4月〜2024年3月に、NDBは47都道府県で7,903,515件の入院手術と'
    '274,579,851単位の鎮痛薬処方を記録した。全国平均の鎮痛薬／手術指数は35.78（SD 5.56）で、'
    '岐阜25.20〜鹿児島49.75の1.97倍の差を示した（9地方間Kruskal–Wallis P<0.001；表1）。'
)

# === 表1 インライン ===
p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(14)
r_c = p_cap.add_run('表1. ')
r_c.bold = True; r_c.font.size = Pt(10)
p_cap.add_run('9地方ブロック別の入院鎮痛薬処方／手術指数の地域要約。Kruskal–Wallis P<0.001。').font.size = Pt(10)

t1 = doc.add_table(rows=1 + len(REGION_ORDER), cols=4, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['地方', '都道府県数', '平均 \u00b1 SD', '範囲']):
    t1.rows[0].cells[i].text = h
    for run in t1.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = t1.rows[idx + 1].cells
    row[0].text = reg_name
    row[1].text = str(len(vals))
    row[2].text = f'{np.mean(vals):.2f} \u00b1 {np.std(vals, ddof=1):.2f}' if len(vals) > 1 else f'{np.mean(vals):.2f}'
    row[3].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)
set_table_borders(t1)
doc.add_paragraph()

doc.add_paragraph(
    '東北は、文化的に日本で最も忍耐強い地域と認識されているが、9地方中7位で'
    '平均指数39.97（SD 3.53）と非東北平均35.17を有意に上回った'
    '（Mann–Whitney U=190, P=0.031; Cohen\'s d=0.87）。'
    '東北6県すべてが全国上位半分にランクした。NSAIDs（P=0.044）、'
    'オピオイドアルカロイド（P=0.003）、合成麻薬（P=0.001）で一貫した。'
)

add_heading_text('Phase 2: 外来神経障害性疼痛薬処方（未調整）', level=2)
doc.add_paragraph(
    f'全国で外来神経障害性疼痛薬は2,289,549,163単位処方された'
    f'（プレガバリン40.2%、ノイロトロピン20.1%、ミロガバリン19.6%、'
    f'デュロキセチン15.3%、トラマドール4.9%）。東北の処方指数は顕著に高く'
    f'（{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; P<0.001; '
    f'd={unadj_d:.2f}）、岩手（566.7）、青森（519.3）、秋田（461.1）が全国上位3を占めた（図1）。'
)

add_inline_figure(
    os.path.join(FIG_DIR, 'fig1_neuropathic_unadjusted.png'),
    '都道府県別の外来神経障害性疼痛薬処方／手術指数（未調整）。東北（赤）が高値に集中。破線＝全国平均。',
    1
)

add_heading_text('交絡分析と調整', level=2)
doc.add_paragraph(
    f'神経障害性疼痛処方は糖尿病薬処方と最も強く相関した（r=0.87, P<0.001）。'
    f'4交絡因子で分散の{reg["model2_adjusted"]["R2"]*100:.1f}%を説明した'
    f'（R\u00b2={reg["model2_adjusted"]["R2"]:.3f}; 図2）。'
    f'調整後、東北効果はモデル2（\u03b2={reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
    f'P={reg["model2_adjusted"]["tohoku_p"]:.2f}）で減弱し非有意となった（表2）。'
)

add_inline_figure(
    os.path.join(FIG_DIR, 'fig2_confounder_correlations.png'),
    '神経障害性疼痛処方と交絡疾患プロキシの相関。各点は1都道府県。糖尿病薬が最強の相関（r=0.87）。',
    2
)

# === 表2 インライン ===
p_cap2 = doc.add_paragraph()
p_cap2.paragraph_format.space_before = Pt(14)
r_c2 = p_cap2.add_run('表2. ')
r_c2.bold = True; r_c2.font.size = Pt(10)
p_cap2.add_run('外来神経障害性疼痛処方の回帰モデル。').font.size = Pt(10)

models = [
    ('モデル1（未調整）', '\u2014', '\u2014', f'd={unadj_d:.2f}'),
    ('モデル2（全交絡）', f'{reg["model2_adjusted"]["tohoku_coef"]:.1f}', f'{reg["model2_adjusted"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}'),
    ('モデル3（コア神経障害薬）', f'{reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}', f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model3_core_neuropathic"]["R2"]:.3f}'),
    ('モデル4（神経ブロック）', f'{reg["model4_nerve_blocks"]["tohoku_coef"]:.1f}', f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model4_nerve_blocks"]["R2"]:.3f}'),
    ('モデル5（統合）', f'{reg["model5_integrated"]["tohoku_coef"]:.1f}', f'{reg["model5_integrated"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model5_integrated"]["R2"]:.3f}'),
]
t2 = doc.add_table(rows=1 + len(models), cols=4, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['モデル', '東北 \u03b2', '東北 P値', '適合度']):
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(9)
for idx, (name, beta, pval, fit) in enumerate(models):
    row = t2.rows[idx + 1].cells
    row[0].text = name; row[1].text = beta; row[2].text = pval; row[3].text = fit
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)
set_table_borders(t2)
doc.add_paragraph()

doc.add_paragraph(
    f'調整後CPSP指数は未調整と大きく異なるパターンを示した。'
    f'東北の超過は非有意となり（{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs '
    f'{reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f'P={reg["adjusted_cpsp_test"]["p_value"]:.2f}; d={adj_d:.2f}; 図3）、'
    f'中国が調整後指数で最高、東海が最低となった。'
)

add_inline_figure(
    os.path.join(FIG_DIR, 'fig4_region_unadj_vs_adj.png'),
    '地方別神経障害性疼痛処方の比較：(a)未調整、(b)交絡調整後。東北（赤）は調整後に中位に移動。',
    3
)

add_heading_text('Phase 1–Phase 2の統合', level=2)
doc.add_paragraph(
    f'急性鎮痛薬指数は未調整神経障害性疼痛処方と正の相関を示した（r=0.38, P=0.008）。'
    f'交絡調整後は減弱した（r=0.29, P=0.052）。モデル5では急性疼痛指数が有意な予測因子であり'
    f'（\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}）、'
    f'東北効果は非有意であった。調整後、東北効果は{attenuation:.0f}%減弱した。'
)

add_heading_text('年齢性別標準化請求比', level=2)
doc.add_paragraph(
    f'間接年齢性別標準化後、入院鎮痛薬SCRは{scr_analgesic_range[0]:.1f}〜{scr_analgesic_range[1]:.1f}'
    f'（{scr_analgesic_ratio:.2f}倍の変動）であり、処方の異質性が都道府県の人口構成の違いでは'
    f'説明できないことを確認した。外来神経障害性疼痛薬SCRは{scr_neuro_range[0]:.1f}〜'
    f'{scr_neuro_range[1]:.1f}（{scr_neuro_ratio:.2f}倍）。東北は神経障害性疼痛SCRが高く'
    f'（平均{scr_neuro_tohoku:.1f} vs 非東北{scr_neuro_non_tohoku:.1f}）、粗解析と一致した。'
)

# ============================================================
# 考察
# ============================================================
add_heading_text('考察', level=1)

doc.add_paragraph(
    f'本研究は、NDBオープンデータを用いて日本全47都道府県の周術期および慢性疼痛関連処方を'
    f'初めてマッピングした。タイトルの問いに対する答えは「いいえ」である：'
    f'文化的ラベルは鎮痛薬の必要性を予測しない。'
    f'3つの主要な知見が、鎮痛薬を処方するすべての臨床医に示唆を持つ。'
)

add_heading_text('文化的ラベルの失敗：一文化内で1.97倍の変動', level=2)
d2 = (
    f'日本の痛みに対する文化的忍耐が十分に文書化されているにもかかわらず、{cite(5,6)}'
    f'急性周術期鎮痛薬処方に1.97倍の都道府県間変動を見出した。'
    f'Cohen and Nisbettが米国内の地域文化差が行動の違いを生むことを示したように、{cite(7)}'
    f'日本の疼痛文化も画一的ではない。'
)
p = doc.add_paragraph()
add_ref_runs(p, d2)

add_heading_text('すべての処方医にとっての意義', level=2)
d3 = (
    f'1.97倍の国内変動は、痛み専門医のみならず外科医・内科医・一般診療医など'
    f'鎮痛薬を処方するすべての医師に直接的な臨床的意義を持つ。'
    f'Anderson et alは民族的少数者が一貫して不十分な疼痛管理を受けることを示した。{cite(17)}'
    f'Campbell and Edwardsは文化的期待による系統的な治療不足・過剰を同定した。{cite(18)}'
    f'より広い教訓は、国内の疼痛処方の異質性が、臨床実践における国レベルの文化的ステレオタイプの使用を問うということである。'
    f'日本人論による均質性の「覇権」{cite(19)}とBurgessが指摘した均質な日本の「幻想」{cite(20)}は、'
    f'「日本人は忍耐強い」と「日本人は均質」という二重のステレオタイプを作り出す。'
    f'文化的に均一と広く見なされる社会内でこれほどの異質性が存在するなら、'
    f'他のあらゆる国レベルの疼痛ステレオタイプにも同様の脆弱性があると推定される。'
    f'IASPの改訂された痛みの定義は本質的に主観的であり、{cite(21)}'
    f'客観的侵害受容モニタリングが文化的背景に依存しない周術期評価を標準化しうる。{cite(22)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3)

add_heading_text('交絡因子が東北の見かけの超過を説明する', level=2)
doc.add_paragraph(
    f'方法論的に最も重要な知見は、神経障害性疼痛処方の劇的な地域変動'
    f'（未調整d={unadj_d:.2f}）が交絡疾患プロキシで大部分説明されたことである。'
    f'糖尿病薬処方のみでr=0.87の相関を示した。調整後、東北効果は{attenuation:.0f}%減弱し非有意となった。'
    f'これは生態学的疼痛研究への重要な示唆を持つ。'
)

add_heading_text('集団レベルの急性–慢性疼痛連続体', level=2)
d5 = (
    f'Phase 1（急性）とPhase 2（慢性、調整後）の正の相関（r=0.29, P=0.052）は、'
    f'急性術後痛がCPSPのリスク因子であるという個人レベルのエビデンスと整合する。{cite(23)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d5)

add_heading_text('モビリティの時代における地域的スティグマ', level=2)
doc.add_paragraph(
    '統計的エビデンスを超えて、地域に基づくステレオタイピングの前提そのものが'
    '現代の日本および西太平洋地域においてますます維持できなくなっている。'
    '国内移動、都市化、高齢化する労働力は各都道府県の人口構成を根本的に変容させた。'
    'シンガポールの病院に来院する患者は、日本で生まれ、オーストラリアで育ち、'
    '東南アジア各国で働いているかもしれない——国籍や民族に基づく仮定は無意味である。'
    '国際的には、アジア太平洋地域における医療ツーリズムと駐在員医療の成長により、'
    '臨床医は多様な文化的背景を持つ患者にますます頻繁に遭遇する。'
    'この文脈で、「東北の患者は忍耐強い」や「日本人は痛みをよく耐える」という地域的・'
    '国家的ステレオタイプを適用することは、不正確であるだけでなく潜在的に有害である。'
    '医療は、臨床医と個人の患者との出会いであり続けなければならない'
    '——臨床医と文化的カテゴリーとの出会いではない。'
    '構造化された疼痛スケール、機能的アウトカム指標、そして利用可能な場合は'
    '客観的侵害受容モニタリングによる個別化された評価こそが、'
    'モバイルで多文化的な世界における鎮痛薬処方の唯一の信頼できる基盤である。'
)

add_heading_text('強みと限界', level=2)
d6b = (
    f'主な限界は生態学的デザインに固有のものである。分析単位は個人ではなく都道府県（生態学的誤謬）。'
    f'NDBオープンデータは診断コードを含まず、CPSPを直接同定できない。'
    f'SCRによる間接年齢性別標準化は人口構成の影響を除外したが、'
    f'個人レベルの交絡因子は集計NDBデータでは測定不能である。'
)
p = doc.add_paragraph()
add_ref_runs(p, d6b)

add_heading_text('結論', level=2)
doc.add_paragraph(
    '文化的ラベルは鎮痛薬の必要性を予測するか？本研究の全数データは「いいえ」と示す。'
    '日本の文化的な痛みの忍耐（我慢）にもかかわらず、周術期および慢性疼痛関連処方は'
    '都道府県間で最大1.97倍変動する。'
    '外科医、内科医、麻酔科医を問わず、「日本人患者」というラベルで処方を調整することは、'
    '過剰投与と過少投与の両方のリスクを伴う。'
    'この原則はあらゆる文化的ラベルに適用される。個別化された疼痛評価が文化的ステレオタイプに'
    '基づく仮定に取って代わるべきである。'
)

# ============================================================
# 末尾セクション
# ============================================================
doc.add_paragraph()
add_heading_text('著者貢献', level=1)
doc.add_paragraph('TO: 研究の着想・設計、データ取得・検証、全解析の実施、可視化、原稿執筆。TOは全データへのフルアクセスを有し、投稿の決定に最終責任を負った。')

doc.add_paragraph()
add_heading_text('利益相反', level=1)
doc.add_paragraph('著者は利益相反がないことを宣言する。')

doc.add_paragraph()
add_heading_text('データ共有', level=1)
doc.add_paragraph('本研究で使用したNDBオープンデータは厚生労働省ウェブサイト（https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html）から公開されている。解析コードはhttps://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japanで利用可能。')

doc.add_paragraph()
add_heading_text('謝辞', level=1)
doc.add_paragraph('NDBオープンデータを公開している厚生労働省に感謝する。データ処理と原稿作成の一部に生成AI（Claude, Anthropic）を使用した。著者は原稿の正確性と内容に全責任を負う。')

doc.add_page_break()
add_heading_text('参考文献', level=1)
for i, ref_text in enumerate(ref_list, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i} ')
    run_num.bold = True
    p.add_run(ref_text)

outpath = os.path.join(LRH_DIR, 'LRH_manuscript_JA.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
