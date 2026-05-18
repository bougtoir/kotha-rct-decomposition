#!/usr/bin/env python3
"""Create separate editable tables (docx) for Lancet Regional Health – WP (English)."""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import csv
import numpy as np
from collections import defaultdict

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
LRH_DIR = os.path.join(OUTPUT_DIR, 'lancet_rh_wp')
os.makedirs(LRH_DIR, exist_ok=True)

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_EN = {
    '北海道': 'Hokkaido', '東北': 'Tohoku', '関東': 'Kanto',
    '北陸・甲信越': 'Hokuriku-Koshinetsu', '東海': 'Tokai', '近畿': 'Kinki',
    '中国': 'Chugoku', '四国': 'Shikoku', '九州・沖縄': 'Kyushu-Okinawa',
}
REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

unadj_d = reg["model1_unadjusted"]["cohens_d"]

def set_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)

# ============================================================
# TABLE 1
# ============================================================
doc1 = Document()
style = doc1.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

p = doc1.add_paragraph()
r = p.add_run('Table 1. ')
r.bold = True
p.add_run(
    'Regional summary of inpatient analgesic prescribing per surgery across nine regional blocks '
    '(NDB Open Data, April 2023\u2013March 2024). '
    'Values are analgesic-per-surgery index (mean \u00b1 SD). '
    'Kruskal\u2013Wallis P<0\u00b7001.'
)

t1 = doc1.add_table(rows=1 + len(REGION_ORDER) + 1, cols=4, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Region', 'n (prefectures)', 'Mean \u00b1 SD', 'Range']):
    t1.rows[0].cells[i].text = h
    for run in t1.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = t1.rows[idx + 1].cells
    row[0].text = REGION_EN[reg_name]
    row[1].text = str(len(vals))
    row[2].text = f'{np.mean(vals):.2f} \u00b1 {np.std(vals, ddof=1):.2f}' if len(vals) > 1 else f'{np.mean(vals):.2f}'
    row[3].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'

# National total row
all_vals = [r['acute_analgesic_per_surgery'] for r in rows]
total_row = t1.rows[-1].cells
total_row[0].text = 'National'
for run in total_row[0].paragraphs[0].runs:
    run.bold = True
total_row[1].text = '47'
total_row[2].text = f'{np.mean(all_vals):.2f} \u00b1 {np.std(all_vals, ddof=1):.2f}'
total_row[3].text = f'{min(all_vals):.2f}\u2013{max(all_vals):.2f}'

set_borders(t1)

outpath1 = os.path.join(LRH_DIR, 'LRH_Table1_EN.docx')
doc1.save(outpath1)
print(f'Saved: {outpath1}')

# ============================================================
# TABLE 2
# ============================================================
doc2 = Document()
style = doc2.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

p = doc2.add_paragraph()
r = p.add_run('Table 2. ')
r.bold = True
p.add_run(
    'Regression models for outpatient neuropathic pain prescribing with Tohoku indicator '
    'and progressive confounder adjustment.'
)

models = [
    ('Model 1 (unadjusted)', '\u2014', 'P<0\u00b7001', f'd={unadj_d:.2f}', 'Tohoku vs non-Tohoku t-test'),
    ('Model 2 (all confounders)', f'{reg["model2_adjusted"]["tohoku_coef"]:.1f}', f'{reg["model2_adjusted"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}', 'Diabetes + herpes + antidep + anxiolytics + Tohoku'),
    ('Model 3 (core neuro)', f'{reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}', f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model3_core_neuropathic"]["R2"]:.3f}', 'Pregabalin + mirogabalin only'),
    ('Model 4 (nerve blocks)', f'{reg["model4_nerve_blocks"]["tohoku_coef"]:.1f}', f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model4_nerve_blocks"]["R2"]:.3f}', 'Outpatient nerve block procedures'),
    ('Model 5 (integrated)', f'{reg["model5_integrated"]["tohoku_coef"]:.1f}', f'{reg["model5_integrated"]["tohoku_p"]:.3f}', f'R\u00b2={reg["model5_integrated"]["R2"]:.3f}', 'Acute pain index + confounders'),
]

t2 = doc2.add_table(rows=1 + len(models), cols=5, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'Tohoku \u03b2', 'Tohoku P', 'Fit', 'Variables']):
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

for idx, (name, beta, pval, fit, variables) in enumerate(models):
    row = t2.rows[idx + 1].cells
    row[0].text = name
    row[1].text = beta
    row[2].text = pval
    row[3].text = fit
    row[4].text = variables

set_borders(t2)

outpath2 = os.path.join(LRH_DIR, 'LRH_Table2_EN.docx')
doc2.save(outpath2)
print(f'Saved: {outpath2}')
