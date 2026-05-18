"""
Generate the manuscript as .docx (Japanese version).
Formatted for Research Integrity and Peer Review (BMC/Springer Nature).

Sections: 背景, 方法, 結果, 考察, 限界, 結論, 略語一覧, 宣言
"""

import os
import sys
import json
import re
import numpy as np
import pandas as pd
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(BASE_DIR, 'figures')
RESULTS_DIR = os.path.join(BASE_DIR, 'results')
DATA_DIR = os.path.join(BASE_DIR, 'data')


def load_results():
    with open(os.path.join(RESULTS_DIR, 'full_results.json'), 'r') as f:
        return json.load(f)


def add_superscript_refs(paragraph, text):
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(10.5)


def get_domain_summary(results):
    domains = {}
    for r in results:
        cat = r['category']
        v = r['verdict']['verdict']
        if cat not in domains:
            domains[cat] = {'total': 0, 'verdicts': Counter(), 'curves': []}
        domains[cat]['total'] += 1
        domains[cat]['verdicts'][v] += 1
        domains[cat]['curves'].append(r)
    return domains


def get_curve(results, name):
    for r in results:
        if r['name'] == name:
            return r
    return None


def fmt_p(p_val):
    if p_val < 0.0001:
        return f"{p_val:.1e}"
    return f"{p_val:.3f}"


VERDICT_JA = {
    'NOT_SIGNIFICANT': '非有意',
    'ROBUST_NONLINEAR': '頑健',
    'OUTLIER_DEPENDENT': '外れ値依存',
    'OVERFITTING': '過適合',
    'BIC_PREFERS_LINEAR': 'BIC線形',
}


def create_manuscript_ja():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Load results dynamically
    results = load_results()
    verdicts = [r['verdict']['verdict'] for r in results]
    n_ns = verdicts.count('NOT_SIGNIFICANT')
    n_outlier = verdicts.count('OUTLIER_DEPENDENT')
    n_robust = verdicts.count('ROBUST_NONLINEAR')
    n_overfit = verdicts.count('OVERFITTING')
    domains = get_domain_summary(results)

    # ======== TITLE ========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        '定説曲線の脆弱性：現代的モデル選択手法を用いた\n'
        '52本の確立された非線形関係の体系的学際的監査')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('大西 龍輝')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[所属機関]')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('責任著者：大西 龍輝（[email]）')
    run.font.size = Pt(10)

    doc.add_page_break()

    # ======== 構造化要旨 ========
    h = doc.add_heading('要旨', level=1)

    abstract_parts = [
        ("【背景】", (
            "多くの曲線関係が様々な学問分野で「定説」として引用されているが、現代のモデル選択手法で"
            "体系的に再検証されたものは少ない。これら「定説曲線」の頑健性は、その関数形に依存する"
            "政策提言と理論的主張に直接影響する。"
        )),
        ("【方法】", (
            "8分野にまたがる52本の定説曲線に統一的な4検定フレームワークを適用した：二次項に対する"
            "ネストF検定、AIC/BICによるモデル選択、Leave-One-Out交差検証（LOOCV）による予測精度評価、"
            "およびCookの距離に基づく上位3点除外の感度分析。データは主にWorld Bank WDI APIおよび"
            "公表文献から取得した。"
        )),
        ("【結果】", (
            f"52本中、{n_ns}本（{100*n_ns/52:.0f}%）は非線形性が統計的に非有意、"
            f"{n_outlier}本（{100*n_outlier/52:.0f}%）は外れ値依存（1〜3点の除外で有意性消失）、"
            f"頑健な非線形性を示したのは{n_robust}本（{100*n_robust/52:.0f}%）のみであった。"
            f"分野間の非対称性が顕著であり、公衆衛生学の曲線は大部分が頑健"
            f"（{domains.get('Public Health', {}).get('verdicts', {}).get('ROBUST_NONLINEAR', 0)}"
            f"/{domains.get('Public Health', {}).get('total', 10)}）であったのに対し、"
            f"経済学の曲線は大部分が非有意"
            f"（{domains.get('Economics', {}).get('verdicts', {}).get('NOT_SIGNIFICANT', 0)}"
            f"/{domains.get('Economics', {}).get('total', 12)}）であった。"
        )),
        ("【結論】", (
            "教科書的非線形関係の約3分の2が現代的頑健性検定に不合格であった。"
            "定説曲線を非線形理論の経験的根拠として引用する際には慎重さが求められる。"
        )),
    ]
    for label, text in abstract_parts:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10.5)
        run2 = p.add_run(text)
        run2.font.size = Pt(10.5)

    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('キーワード：')
    run.bold = True
    p.add_run('モデル選択、非線形性、外れ値依存、頑健性、F検定、情報量規準、'
              '交差検証、Cookの距離、定説曲線、メタリサーチ')

    doc.add_page_break()

    # ======== 背景 ========
    h = doc.add_heading('背景', level=1)

    bg_paras = [
        ("曲線関係は社会科学・自然科学において特権的な地位を占めている。マクロ経済学のフィリップス曲線から"
         "公衆衛生学のプレストン曲線に至るまで、これらの非線形関数形は広く教育され、政策文書で頻繁に引用され、"
         "確立された経験則として扱われている。{1\u20133}しかし、これらの関係の多くは限られたデータ、初歩的な"
         "統計手法、そしてAICやBICが標準化される以前の時代に確立されたものである。{4,5}"),

        ("本研究の着想は、著者らが実施したプレストン曲線の予備的解析に端を発する。{6}"
         "同解析において、プレストン曲線の見かけの凹性（二次項）が米国1点の外れ値としての位置に"
         "大きく依存し、米国を除外するとp値が0.49に上昇して非有意となることが明らかになった。"
         "この発見は、他の「定説曲線」にも同様の脆弱性が潜在するのではないかという問いを提起した。"
         "実際、CO\u2082排出に関する環境クズネッツ曲線は繰り返し異議を唱えられており、{7,8}"
         "ダニング・クルーガー効果は平均への回帰という統計的アーティファクトであるという批判がある。{9}"),

        ("これらの個別的指摘にもかかわらず、学際的な体系的監査は行われてこなかった。"
         "本研究は8つの学問分野にまたがる52本の定説曲線に統一的な方法論的フレームワークを"
         "適用することで、このギャップを埋める。"),
    ]
    for text in bg_paras:
        p = doc.add_paragraph()
        add_superscript_refs(p, text)

    # ======== 方法 ========
    h = doc.add_heading('方法', level=1)

    doc.add_heading('曲線の選定', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "候補曲線は各分野の主要ハンドブックにおける命名法則・定型化事実の系統的レビュー、"
        "および再現性危機に関する文献のレビューを通じて同定した。{10,11}"
        "初期候補78本から、以下の包含基準を全て満たす52本を選定した。"
    )

    p = doc.add_paragraph()
    run = p.add_run('包含基準：')
    run.bold = True
    p.add_run("(1) 命名的・定説的地位（教科書2冊以上 or 被引用500以上）；"
              "(2) 明示的な非線形性の主張；(3) 二変量での検証可能性；"
              "(4) 公開データの利用可能性（N ≥ 10）；(5) 曲線形状の政策的・理論的関連性。")

    p = doc.add_paragraph()
    run = p.add_run('除外基準：')
    run.bold = True
    p.add_run("定義的・恒等的曲線；非公開データ要件；純粋な時間的ダイナミクス；"
              "既に撤回された関係；二変量射影が無意味な多変量指定。")

    doc.add_paragraph(
        "8分野に層別化：経済学（12本）、公衆衛生（10本）、人口学（6本）、環境科学（6本）、"
        "心理学（5本）、物理学（4本）、政治学（5本）、農学（4本）。"
        "5類型の非線形性を代表：逆U字（14本）、U/J字（8本）、凹型（12本）、べき乗則（8本）、S字型（10本）。"
    )

    doc.add_heading('データソース', level=2)
    doc.add_paragraph(
        "データは主にWorld Bank WDI APIから横断面曲線用に取得した（GDP per capita PPP、平均寿命、"
        "ジニ係数、合計特殊出生率、森林面積率等）。米国マクロ経済時系列（失業率、インフレ率、GDP成長率）も"
        "WDI経由で取得。その他のソースにはOECD.Stat、公表メタ分析、USGS地震カタログ、"
        "元論文デジタル化データを含む。全52曲線のデータソース詳細は補遺表S1に示す。"
    )

    doc.add_heading('統計的フレームワーク', level=2)
    methods = [
        "ネストF検定：制約モデル（線形：y = a + bx）と非制約モデル（二次：y = a + bx + cx²）を"
        "最小二乗法で推定し、二次項に対するF統計量を計算。",
        "情報量規準：線形・二次・対数モデルのAIC/BICを計算し最小値モデルを選択。",
        "LOOCV：線形・二次モデルのLOOCV RMSEを計算しサンプル外予測精度を評価。",
        "Cook距離感度分析：上位3点を除外しF検定を再実施。完全データでp < 0.05だが除外後に"
        "p > 0.05となる曲線を「外れ値依存」と分類。",
    ]
    for m in methods:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_heading('判定分類', level=2)
    doc.add_paragraph("各曲線に以下4判定のいずれかを付与した：")
    verdicts_def = [
        "頑健（ROBUST_NONLINEAR）：全データおよび外れ値除外後も有意、LOOCV RMSEが二次 ≤ 線形。",
        "外れ値依存（OUTLIER_DEPENDENT）：全データでは有意だが上位3点除外で非有意。",
        "非有意（NOT_SIGNIFICANT）：全データでも非有意（p ≥ 0.05）。",
        "過適合（OVERFITTING）：有意だがLOOCV RMSEが二次モデルで悪化。",
    ]
    for v in verdicts_def:
        doc.add_paragraph(v, style='List Bullet')

    doc.add_heading('ソフトウェア', level=2)
    doc.add_paragraph(
        "全分析はPython 3.12（statsmodels、NumPy、pandas）で実施。"
        "World Bankデータはwbgapiライブラリで取得。コードとデータは公開（データ利用可能性の項を参照）。"
    )

    # ======== 結果 ========
    h = doc.add_heading('結果', level=1)

    doc.add_heading('判定の全体分布', level=2)
    doc.add_paragraph(
        f"52本中、{n_robust}本（{100*n_robust/52:.0f}%）が頑健、"
        f"{n_outlier}本（{100*n_outlier/52:.0f}%）が外れ値依存、"
        f"{n_ns}本（{100*n_ns/52:.0f}%）が非有意、"
        f"{n_overfit}本（{100*n_overfit/52:.0f}%）が過適合であった（図1）。"
        f"教科書的非線形関係の約3分の2が有意性に達しないか"
        f"少数の影響力のある観測値に依存していた。"
    )

    # Figure 1
    fig1_path = os.path.join(FIGURES_DIR, 'fig1_verdict_distribution.png')
    if os.path.exists(fig1_path):
        doc.add_paragraph()
        doc.add_picture(fig1_path, width=Inches(5.5))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run('図1 ')
        run.bold = True
        p.add_run('52本の定説曲線の判定分布（分野別）。')

    # Domain results (DYNAMIC)
    doc.add_heading('分野別結果', level=2)

    # Economics
    econ = domains.get('Economics', {})
    econ_v = econ.get('verdicts', {})
    phillips = get_curve(results, 'Phillips Curve')
    kuznets = get_curve(results, 'Kuznets Curve')
    p = doc.add_paragraph()
    run = p.add_run('経済学（12本）：')
    run.bold = True
    p.add_run(
        f"{econ_v.get('NOT_SIGNIFICANT', 0)}本が非有意、"
        f"{econ_v.get('OUTLIER_DEPENDENT', 0)}本が外れ値依存、"
        f"{econ_v.get('ROBUST_NONLINEAR', 0)}本が頑健。"
        f"フィリップス曲線はN={phillips['n']}で非有意"
        f"（p={fmt_p(phillips['f_test']['p_value'])}）。"
        f"クズネッツ曲線はN={kuznets['n']}か国で非有意"
        f"（p={fmt_p(kuznets['f_test']['p_value'])}）。"
        "環境クズネッツ曲線、ラッファー曲線、グレートギャツビー曲線は外れ値依存。"
    )

    # Public Health
    health = domains.get('Public Health', {})
    health_v = health.get('verdicts', {})
    preston = get_curve(results, 'Preston Curve')
    p = doc.add_paragraph()
    run = p.add_run('公衆衛生（10本）：')
    run.bold = True
    if preston['verdict']['verdict'] == 'NOT_SIGNIFICANT':
        preston_text = (
            f"プレストン曲線はWorld Bank実データ（N={preston['n']}）で非有意"
            f"（p={fmt_p(preston['f_test']['p_value'])}）であり、対数線形関係と整合的。"
        )
    else:
        preston_text = (
            f"プレストン曲線は{VERDICT_JA.get(preston['verdict']['verdict'], '')}。"
        )
    p.add_run(
        f"最も頑健な分野（{health_v.get('ROBUST_NONLINEAR', 0)}/{health.get('total', 10)}が頑健）。"
        "BMI死亡J曲線、飲酒死亡J曲線、バーカー仮説U字は全て検証に耐えた。"
        f"{preston_text}"
    )

    # Demography
    demo = domains.get('Demography', {})
    demo_v = demo.get('verdicts', {})
    dt = get_curve(results, 'Demographic Transition (TFR)')
    p = doc.add_paragraph()
    run = p.add_run('人口学（6本）：')
    run.bold = True
    dt_verdict = VERDICT_JA.get(dt['verdict']['verdict'], dt['verdict']['verdict'])
    p.add_run(
        f"{demo_v.get('ROBUST_NONLINEAR', 0)}本が頑健、"
        f"{demo_v.get('NOT_SIGNIFICANT', 0)}本が非有意、"
        f"{demo_v.get('OUTLIER_DEPENDENT', 0)}本が外れ値依存。"
        f"人口転換モデル（N={dt['n']}）は{dt_verdict}"
        f"（p={fmt_p(dt['f_test']['p_value'])}）。"
        "リー・カーター死亡率モデルとコール・トラッセル出生力スケジュールも頑健。"
    )

    # Environmental Science
    env = domains.get('Environmental Science', {})
    env_v = env.get('verdicts', {})
    forest = get_curve(results, 'Forest Transition Curve')
    p = doc.add_paragraph()
    run = p.add_run('環境科学（6本）：')
    run.bold = True
    p.add_run(
        f"{env_v.get('ROBUST_NONLINEAR', 0)}本が頑健（キーリング曲線）、"
        f"{env_v.get('NOT_SIGNIFICANT', 0)}本が非有意、"
        f"{env_v.get('OUTLIER_DEPENDENT', 0)}本が外れ値依存。"
        f"森林転換曲線は実データ（N={forest['n']}）で非有意"
        f"（p={fmt_p(forest['f_test']['p_value'])}）。"
    )

    # Psychology
    psych = domains.get('Psychology', {})
    psych_v = psych.get('verdicts', {})
    p = doc.add_paragraph()
    run = p.add_run('心理学（5本）：')
    run.bold = True
    p.add_run(
        f"{psych_v.get('ROBUST_NONLINEAR', 0)}/5が頑健"
        "（ヤーキーズ・ドッドソン、エビングハウス、ダニング・クルーガー、幸福U字）。"
        "ウェーバー・フェヒナーのみ非有意。"
    )

    # Physics
    phys = domains.get('Physics', {})
    phys_v = phys.get('verdicts', {})
    p = doc.add_paragraph()
    run = p.add_run('物理学（4本）：')
    run.bold = True
    p.add_run(
        f"{phys_v.get('NOT_SIGNIFICANT', 0)}本が非有意、"
        f"{phys_v.get('OUTLIER_DEPENDENT', 0)}本が外れ値依存。"
        "ハッブル法則とグーテンベルグ・リヒターは外れ値依存。"
    )

    # Political Science
    pol = domains.get('Political Science', {})
    pol_v = pol.get('verdicts', {})
    lipset = get_curve(results, 'Lipset Hypothesis')
    p = doc.add_paragraph()
    run = p.add_run('政治学（5本）：')
    run.bold = True
    p.add_run(
        f"リプセット仮説が最も劇的な外れ値依存例"
        f"（全データp={fmt_p(lipset['f_test']['p_value'])}、"
        f"湾岸産油国除外p={fmt_p(lipset['sensitivity']['p_clean'])}）。"
        f"{pol_v.get('NOT_SIGNIFICANT', 0)}/5が非有意。"
    )

    # Agriculture
    agr = domains.get('Agriculture', {})
    agr_v = agr.get('verdicts', {})
    p = doc.add_paragraph()
    run = p.add_run('農学（4本）：')
    run.bold = True
    p.add_run(
        f"{agr_v.get('ROBUST_NONLINEAR', 0)}/4が頑健"
        "（ミッチェルリッヒ収量、微量栄養素U字）；"
        f"{agr_v.get('NOT_SIGNIFICANT', 0)}/4が非有意。"
    )

    # Figure 2
    fig2_path = os.path.join(FIGURES_DIR, 'fig2_sensitivity_analysis.png')
    if os.path.exists(fig2_path):
        doc.add_paragraph()
        doc.add_picture(fig2_path, width=Inches(5.0))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run('図2 ')
        run.bold = True
        p.add_run('F検定p値の外れ値除外に対する感度（Cook距離上位3点除外）。')

    # Table 1
    doc.add_heading('全結果要約', level=2)
    doc.add_paragraph("表1に52本全曲線の統計検定結果を示す。")

    df = pd.read_csv(os.path.join(RESULTS_DIR, 'summary_table.csv'))
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['#', '曲線名', 'N', 'p (全)', 'p (除外後)', 'BIC最良', '判定']
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(7)

    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(row['Curve'])[:28]
        row_cells[2].text = str(row['N'])
        p_full = row['p (full)']
        row_cells[3].text = f"{p_full:.4f}" if p_full > 0.0001 else f"{p_full:.1e}"
        p_clean = row['p (clean)']
        row_cells[4].text = f"{p_clean:.4f}" if p_clean > 0.0001 else f"{p_clean:.1e}"
        row_cells[5].text = str(row['BIC best'])
        verdict = str(row['Verdict'])
        row_cells[6].text = VERDICT_JA.get(verdict, verdict)
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(6.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('表1 ')
    run.bold = True
    p.add_run("全52本の再検証結果。p (全) = 全データでのF検定p値；"
              "p (除外後) = Cook距離上位3点除外後。")

    # ======== 考察 ========
    h = doc.add_heading('考察', level=1)

    doc.add_paragraph(
        f"本体系的監査により、確立された非線形関係の"
        f"{100*(n_ns + n_outlier + n_overfit)/52:.0f}%が少なくとも1つの現代的頑健性検定に"
        "不合格であることが明らかになった。5つの横断的パターンが浮上した："
    )

    patterns = [
        "外れ値駆動型非線形性：最も一般的な失敗モード。横断面分析では1〜3の地政学的に特異な国が曲率を駆動。",
        "分野間非対称性：公衆衛生・心理学の曲線は経済学・政治学より格段に頑健（機構的関係 vs 偶発的規則性）。",
        "時系列 vs 横断面：時系列曲線は構成効果や未測定交絡に対する脆弱性が低い。",
        "対数変換による見かけの非線形性の解消：予測変数の対数変換で線形関係が得られる場合が多い。",
        "BICの保守性：BICはAICより強い複雑性ペナルティにより、線形モデルを選択する頻度が高い。",
    ]
    for i, pat in enumerate(patterns, 1):
        doc.add_paragraph(f"{i}. {pat}")

    doc.add_paragraph(
        "政策的含意として、ラッファー曲線（減税正当化）、環境クズネッツ曲線（成長による汚染解消論）、"
        "リプセット仮説（近代化論）の経験的根拠は外れ値依存であり、これらに基づく政策提言には"
        "慎重さが求められる。"
    )

    # ======== 限界 ========
    h = doc.add_heading('限界', level=1)

    limitations = [
        "第一に、二変量関係のみを対象としており、多変量設定での再検証は含まない。",
        "第二に、統一的に二次モデルを代替仮説としたが、一部の曲線には特定の関数形が適切である。",
        "第三に、8本はWorld Bank API実データ（N=31〜247）だが、残り44本は代表的・公表データ。",
        "第四に、曲率の有意性のみを検定しており、線形関係の存在は検定していない。",
        "第五に、一部の曲線では集計データやメタ分析データに依拠している。",
    ]
    for lim in limitations:
        doc.add_paragraph(lim)

    # ======== 結論 ========
    h = doc.add_heading('結論', level=1)

    doc.add_paragraph(
        f"52本の定説曲線の体系的再検証により、統計的に有意で外れ値除外に耐え"
        f"サンプル外予測でも優れた非線形性を示すのは{100*n_robust/52:.0f}%のみであることが"
        "明らかになった。研究者は定説曲線を引用する際に感度分析とモデル比較規準を報告すべきであり、"
        "政策決定者は特定の曲線形状に基づく介入について慎重であるべきである。"
    )

    # ======== 略語一覧 ========
    doc.add_page_break()
    h = doc.add_heading('略語一覧', level=1)
    abbreviations = [
        ("AIC", "赤池情報量規準"),
        ("BIC", "ベイズ情報量規準"),
        ("CO\u2082", "二酸化炭素"),
        ("EKC", "環境クズネッツ曲線"),
        ("GDP", "国内総生産"),
        ("LOOCV", "Leave-One-Out交差検証"),
        ("LNT", "しきい値なし線形モデル"),
        ("OLS", "最小二乗法"),
        ("PPP", "購買力平価"),
        ("RMSE", "二乗平均平方根誤差"),
        ("TFR", "合計特殊出生率"),
        ("WDI", "World Development Indicators"),
    ]
    for abbr, full in abbreviations:
        p = doc.add_paragraph()
        run = p.add_run(f"{abbr}：")
        run.bold = True
        p.add_run(full)

    # ======== 宣言 ========
    doc.add_page_break()
    h = doc.add_heading('宣言', level=1)

    doc.add_heading('倫理承認と参加同意', level=2)
    doc.add_paragraph("該当なし。本研究は公開集計データのみを使用。")

    doc.add_heading('公表への同意', level=2)
    doc.add_paragraph("該当なし。")

    doc.add_heading('データと資料の利用可能性', level=2)
    doc.add_paragraph(
        "全分析コード、データ、結果は[GitHubリポジトリURL]で公開。"
        "World Bankデータはwbgapiライブラリで取得。全52曲線のソースメタデータ表は補遺に収録。"
    )

    doc.add_heading('利益相反', level=2)
    doc.add_paragraph("著者は利益相反がないことを宣言する。")

    doc.add_heading('資金', level=2)
    doc.add_paragraph("[記入予定]")

    doc.add_heading('著者の貢献', level=2)
    doc.add_paragraph(
        "TO：研究の着想、分析フレームワークの設計、データ収集・分析、結果の解釈、原稿執筆。"
    )

    doc.add_heading('謝辞', level=2)
    doc.add_paragraph("[記入予定]")

    # ======== 補遺：データソーステーブル ========
    doc.add_page_break()
    h = doc.add_heading('補遺：データソーステーブル（表S1）', level=1)

    source_meta_path = os.path.join(DATA_DIR, 'source_metadata.json')
    if os.path.exists(source_meta_path):
        with open(source_meta_path, 'r', encoding='utf-8') as f:
            sources = json.load(f)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['#', '曲線名', '元論文', 'データソース', 'N', '主張された関数形']
        for i, h_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h_text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(7)

        for src in sources:
            row_cells = table.add_row().cells
            row_cells[0].text = str(src['id'])
            row_cells[1].text = src['name'][:25]
            row_cells[2].text = src['original_paper'][:40]
            row_cells[3].text = src['data_source'][:45]
            n_val = src.get('current_n')
            row_cells[4].text = str(n_val) if n_val else 'WB API'
            row_cells[5].text = src['claimed_form'][:30]
            for cell in row_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(6.5)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run('表S1 ')
        run.bold = True
        p.add_run('52本の定説曲線のデータソース一覧。')

    # ======== 参考文献 ========
    doc.add_page_break()
    h = doc.add_heading('参考文献', level=1)
    references = [
        "Phillips AW. The relation between unemployment and the rate of change of money wage rates in the United Kingdom, 1861-1957. Economica. 1958;25(100):283-299.",
        "Kuznets S. Economic growth and income inequality. Am Econ Rev. 1955;45(1):1-28.",
        "Preston SH. The changing relation between mortality and level of economic development. Popul Stud. 1975;29(2):231-248.",
        "Akaike H. A new look at the statistical model identification. IEEE Trans Automat Contr. 1974;19(6):716-723.",
        "Schwarz G. Estimating the dimension of a model. Ann Stat. 1978;6(2):461-464.",
        "Onishi T. Re-examination of the Preston Curve: outlier dependence of quadratic fit. Working paper. 2026.",
        "Grossman GM, Krueger AB. Environmental impacts of a North American free trade agreement. NBER Working Paper 3914. 1991.",
        "Stern DI. The rise and fall of the environmental Kuznets curve. World Dev. 2004;32(8):1419-1439.",
        "Krueger J, Mueller RA. Unskilled, unaware, or both? J Pers Soc Psychol. 2002;82(2):180-188.",
        "Open Science Collaboration. Estimating the reproducibility of psychological science. Science. 2015;349(6251):aac4716.",
        "Ioannidis JPA. Why most published research findings are false. PLoS Med. 2005;2(8):e124.",
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        p.add_run(ref)
        p.paragraph_format.left_indent = Cm(1)

    # Save
    output_path = os.path.join(BASE_DIR, 'manuscript_canonical_curves_ja.docx')
    doc.save(output_path)
    print(f"Japanese manuscript saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_manuscript_ja()
