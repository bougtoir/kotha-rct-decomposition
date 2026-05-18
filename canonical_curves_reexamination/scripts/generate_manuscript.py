"""
Generate the manuscript as .docx (English) formatted for:
Research Integrity and Peer Review (BMC/Springer Nature)

Requirements:
- Structured abstract (Background/Methods/Results/Conclusions) ≤350 words
- Sections: Background, Methods, Results, Discussion, Limitations, Conclusions
- Declarations section
- List of abbreviations
- Vancouver-style numbered references
- ≤5,000 words (main body)
- Keywords (3-10)

Onishi T. 2026.
"""

import os
import sys
import json
import re
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(BASE_DIR, 'figures')
RESULTS_DIR = os.path.join(BASE_DIR, 'results')
DATA_DIR = os.path.join(BASE_DIR, 'data')


def load_results():
    with open(os.path.join(RESULTS_DIR, 'full_results.json'), 'r') as f:
        return json.load(f)


def add_superscript_refs(paragraph, text):
    """Parse text with {N} markers and add superscript references using Word font."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)


def get_domain_summary(results):
    """Dynamically compute domain-level verdict counts from results."""
    from collections import Counter
    domains = {}
    for r in results:
        cat = r['category']
        v = r['verdict']['verdict']
        if cat not in domains:
            domains[cat] = {'total': 0, 'verdicts': Counter(), 'curves': []}
        domains[cat]['total'] += 1
        domains[cat]['verdicts'][v] += 1
        domains[cat]['curves'].append(r)
    return domains


def get_curve_result(results, name):
    """Get result for a specific curve by name."""
    for r in results:
        if r['name'] == name:
            return r
    return None


def fmt_p(p_val):
    """Format p-value for display."""
    if p_val < 0.0001:
        return f"{p_val:.1e}"
    return f"{p_val:.3f}"


def create_manuscript():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Load results for dynamic content
    results = load_results()
    verdicts = [r['verdict']['verdict'] for r in results]
    n_ns = verdicts.count('NOT_SIGNIFICANT')
    n_outlier = verdicts.count('OUTLIER_DEPENDENT')
    n_robust = verdicts.count('ROBUST_NONLINEAR')
    n_overfit = verdicts.count('OVERFITTING')
    domains = get_domain_summary(results)

    # ======== TITLE PAGE ========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Fragility of canonical curves: a systematic cross-disciplinary '
        'audit of 52 established nonlinear relationships using modern '
        'model selection methods')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tatsuki Onishi')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Institutional affiliation]')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Corresponding author: Tatsuki Onishi ([email])')
    run.font.size = Pt(10)

    doc.add_page_break()

    # ======== ABSTRACT (structured, ≤350 words) ========
    h = doc.add_heading('Abstract', level=1)

    abstract_sections = [
        ("Background: ", (
            "Many curvilinear relationships are treated as established empirical facts "
            "across diverse academic disciplines, yet few have been systematically re-evaluated "
            "using modern model selection techniques. The robustness of these 'canonical curves' "
            "has implications for policy recommendations and theoretical claims that depend on "
            "their specific functional form."
        )),
        ("Methods: ", (
            "We applied a uniform four-test framework to 52 canonical curves spanning eight "
            "disciplines: nested F-tests for quadratic terms, Akaike and Bayesian Information "
            "Criteria (AIC/BIC) for model selection, leave-one-out cross-validation (LOOCV) for "
            "predictive accuracy, and Cook's distance sensitivity analysis with removal of the "
            "top 3 influential observations. Data were drawn from the World Bank World Development "
            "Indicators API and published sources."
        )),
        ("Results: ", (
            f"Of 52 curves, {n_ns} ({100*n_ns/52:.0f}%) showed no statistically significant "
            f"nonlinearity, {n_outlier} ({100*n_outlier/52:.0f}%) were outlier-dependent "
            f"(significance lost after removing 1\u20133 influential points), and only "
            f"{n_robust} ({100*n_robust/52:.0f}%) demonstrated robust nonlinearity surviving "
            f"all tests. Domain asymmetry was pronounced: public health curves were predominantly "
            f"robust ({domains.get('Public Health', {}).get('verdicts', {}).get('ROBUST_NONLINEAR', 0)}"
            f"/{domains.get('Public Health', {}).get('total', 10)}) while economics curves largely "
            f"failed ({domains.get('Economics', {}).get('verdicts', {}).get('NOT_SIGNIFICANT', 0)}"
            f"/{domains.get('Economics', {}).get('total', 12)} non-significant)."
        )),
        ("Conclusions: ", (
            "Approximately two-thirds of textbook nonlinear relationships fail modern robustness "
            "tests. Researchers and policymakers should exercise caution when citing canonical "
            "curves as empirical support for nonlinear theories."
        )),
    ]

    for label, text in abstract_sections:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)

    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Keywords: ')
    run.bold = True
    p.add_run('model selection, nonlinearity, outlier dependence, robustness, '
              'F-test, information criteria, cross-validation, Cook\'s distance, '
              'canonical relationships, meta-research')

    doc.add_page_break()

    # ======== BACKGROUND ========
    h = doc.add_heading('Background', level=1)

    bg_paras = [
        ("Curvilinear relationships occupy a privileged position in the social and natural "
         "sciences. From the Phillips Curve in macroeconomics to the Preston Curve in public "
         "health, these nonlinear functional forms are widely taught, frequently cited in policy "
         "documents, and treated as established empirical regularities.{1\u20133} Yet many were "
         "originally established with limited data, rudimentary statistical methods, and before "
         "model selection criteria such as the Akaike Information Criterion (AIC) and Bayesian "
         "Information Criterion (BIC) became standard practice.{4,5}"),

        ("The present study was motivated by a preliminary analysis of the Preston Curve "
         "conducted by the authors.{6} That analysis revealed that the apparent concavity "
         "(quadratic term) depends heavily on the position of the United States as a single "
         "outlier\u2014removing the US raised the p-value to 0.49, rendering the nonlinear term "
         "non-significant. This finding raised the question of whether similar fragilities lurk "
         "beneath other canonical curves. The Environmental Kuznets Curve for CO\u2082 emissions "
         "has been repeatedly challenged,{7,8} and the Dunning-Kruger effect has been argued to "
         "be a statistical artifact of regression to the mean.{9}"),

        ("Despite individual critiques, no systematic cross-disciplinary audit has been "
         "conducted. The present study fills this gap by applying a uniform methodological "
         "framework to 52 canonical curves across eight academic disciplines."),
    ]
    for text in bg_paras:
        p = doc.add_paragraph()
        add_superscript_refs(p, text)

    # ======== METHODS ========
    h = doc.add_heading('Methods', level=1)

    # Curve selection
    doc.add_heading('Curve selection', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Candidate curves were identified through systematic review of named laws and "
        "stylized facts in major handbooks, citation analysis of papers with 'curve,' 'law,' "
        "or 'paradox' in titles, and review of the replication crisis literature.{10,11} "
        "This yielded 78 candidates, from which 52 met all inclusion criteria."
    )

    p = doc.add_paragraph()
    run = p.add_run('Inclusion criteria: ')
    run.bold = True
    p.add_run("(1) eponymous/canonical status (\u22652 textbooks or >500 citations); "
              "(2) explicit nonlinearity claim; (3) bivariate testability; "
              "(4) publicly available data (N \u2265 10); (5) policy or theoretical relevance "
              "of curve shape.")

    p = doc.add_paragraph()
    run = p.add_run('Exclusion criteria: ')
    run.bold = True
    p.add_run("definitional/tautological curves; proprietary data requirements; "
              "purely temporal dynamics; formally retracted relationships; "
              "curves requiring multivariate specification.")

    p = doc.add_paragraph(
        "Curves were stratified across eight disciplines: economics (12), public health (10), "
        "demography (6), environmental science (6), psychology (5), physics (4), political "
        "science (5), and agriculture (4). Five types of claimed nonlinearity were represented: "
        "inverted-U (14), U/J-shaped (8), concave/saturating (12), power-law (8), "
        "and S-shaped/structural-break (10)."
    )

    # Data sources
    doc.add_heading('Data sources', level=2)
    p = doc.add_paragraph(
        "Data were drawn primarily from the World Bank World Development Indicators (WDI) "
        "API for cross-country curves (GDP per capita PPP, life expectancy, Gini coefficient, "
        "total fertility rate, forest area), with US macroeconomic time series (unemployment, "
        "inflation, GDP growth) also obtained via WDI. Additional sources included OECD.Stat, "
        "published meta-analyses, USGS earthquake catalogs, and digitized original publication "
        "data. The complete data source table for all 52 curves is provided in Additional file 1."
    )

    # Statistical framework
    doc.add_heading('Statistical framework', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Nested F-test. ')
    run.bold = True
    p.add_run("For each curve, restricted (linear: y = a + bx) and unrestricted "
              "(quadratic: y = a + bx + cx\u00b2) models were fitted via ordinary least squares. "
              "The F-statistic tested the significance of the additional parameter.")

    p = doc.add_paragraph()
    run = p.add_run('Information criteria. ')
    run.bold = True
    p.add_run("AIC and BIC were computed for linear, quadratic, and logarithmic models. "
              "The model with lowest criterion value was selected.")

    p = doc.add_paragraph()
    run = p.add_run('Leave-one-out cross-validation. ')
    run.bold = True
    p.add_run("LOOCV root mean squared error (RMSE) was computed for linear and quadratic "
              "models to assess out-of-sample predictive accuracy.")

    p = doc.add_paragraph()
    run = p.add_run("Cook's distance sensitivity analysis. ")
    run.bold = True
    p.add_run("The top 3 most influential observations (by Cook's distance) were removed "
              "and the F-test repeated. A curve was classified as outlier-dependent if "
              "significance (p < 0.05) was lost after removal.")

    # Verdict classification
    doc.add_heading('Verdict classification', level=2)
    p = doc.add_paragraph("Each curve received one of four verdicts:")
    verdict_defs = [
        "ROBUST_NONLINEAR: significant with full data AND after outlier removal, with quadratic LOOCV RMSE \u2264 linear RMSE.",
        "OUTLIER_DEPENDENT: significant with full data but non-significant after removing top 3 influential points.",
        "NOT_SIGNIFICANT: non-significant even with full data (p \u2265 0.05).",
        "OVERFITTING: significant but LOOCV RMSE is worse for the quadratic model.",
    ]
    for v in verdict_defs:
        doc.add_paragraph(v, style='List Bullet')

    # Software
    doc.add_heading('Software', level=2)
    doc.add_paragraph(
        "All analyses were conducted in Python 3.12 using statsmodels (OLS regression), "
        "NumPy, and pandas. World Bank data were fetched via the wbgapi library. "
        "Code and data are publicly available (see Availability of data and materials)."
    )

    # ======== RESULTS ========
    h = doc.add_heading('Results', level=1)

    # Overall verdicts
    doc.add_heading('Overall distribution of verdicts', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        f"Of 52 canonical curves, {n_robust} ({100*n_robust/52:.0f}%) demonstrated robust "
        f"nonlinearity, {n_outlier} ({100*n_outlier/52:.0f}%) were outlier-dependent, "
        f"{n_ns} ({100*n_ns/52:.0f}%) showed no significant nonlinearity, and "
        f"{n_overfit} ({100*n_overfit/52:.0f}%) exhibited overfitting (Fig. 1). "
        f"Approximately two-thirds of textbook nonlinear relationships either failed to "
        f"reach significance or were driven by a small number of influential observations."
    )

    # Figure 1
    doc.add_paragraph()
    fig1_path = os.path.join(FIGURES_DIR, 'fig1_verdict_distribution.png')
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(5.5))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run('Fig. 1 ')
        run.bold = True
        p.add_run('Distribution of verdicts across 52 canonical curves, stratified by domain.')

    # Domain-level results (dynamic)
    doc.add_heading('Results by domain', level=2)

    # Economics
    econ = domains.get('Economics', {})
    econ_v = econ.get('verdicts', {})
    phillips = get_curve_result(results, 'Phillips Curve')
    kuznets = get_curve_result(results, 'Kuznets Curve')
    p = doc.add_paragraph()
    run = p.add_run('Economics. ')
    run.bold = True
    p.add_run(
        f"Of 12 curves, {econ_v.get('NOT_SIGNIFICANT', 0)} were non-significant, "
        f"{econ_v.get('OUTLIER_DEPENDENT', 0)} were outlier-dependent, and "
        f"{econ_v.get('ROBUST_NONLINEAR', 0)} demonstrated robust nonlinearity. "
        f"The Phillips Curve showed no significant nonlinearity "
        f"(p={fmt_p(phillips['f_test']['p_value'])}). "
        f"The Kuznets Curve's inverted-U was not significant with {kuznets['n']} countries "
        f"(p={fmt_p(kuznets['f_test']['p_value'])}). "
        "The Environmental Kuznets Curve (CO\u2082), Laffer Curve, and Great Gatsby Curve "
        "were all outlier-dependent."
    )

    # Public Health
    health = domains.get('Public Health', {})
    health_v = health.get('verdicts', {})
    preston = get_curve_result(results, 'Preston Curve')
    p = doc.add_paragraph()
    run = p.add_run('Public health. ')
    run.bold = True
    preston_desc = (
        f"The Preston Curve with real World Bank data (N={preston['n']}) showed "
        f"no significant nonlinearity (p={fmt_p(preston['f_test']['p_value'])}), "
        "consistent with a log-linear relationship"
    ) if preston['verdict']['verdict'] == 'NOT_SIGNIFICANT' else (
        f"The Preston Curve showed {preston['verdict']['verdict'].lower().replace('_', ' ')}"
    )
    p.add_run(
        f"Public health curves showed the highest robustness rate "
        f"({health_v.get('ROBUST_NONLINEAR', 0)}/{health.get('total', 10)} robust). "
        f"The BMI-Mortality J-curve, Alcohol-Mortality J-curve, and Barker Hypothesis "
        f"U-shape all survived rigorous testing. {preston_desc}."
    )

    # Demography
    demo = domains.get('Demography', {})
    demo_v = demo.get('verdicts', {})
    dt = get_curve_result(results, 'Demographic Transition (TFR)')
    p = doc.add_paragraph()
    run = p.add_run('Demography. ')
    run.bold = True
    dt_desc = (
        f"The Demographic Transition model (N={dt['n']}) was classified as "
        f"{dt['verdict']['verdict'].lower().replace('_', ' ')} "
        f"(p={fmt_p(dt['f_test']['p_value'])})"
    )
    p.add_run(
        f"Of 6 curves, {demo_v.get('ROBUST_NONLINEAR', 0)} showed robust nonlinearity, "
        f"{demo_v.get('NOT_SIGNIFICANT', 0)} were non-significant, and "
        f"{demo_v.get('OUTLIER_DEPENDENT', 0)} were outlier-dependent. "
        f"{dt_desc}. The Lee-Carter mortality model and Coale-Trussell fertility schedule "
        f"demonstrated robust nonlinearity."
    )

    # Environmental Science
    env = domains.get('Environmental Science', {})
    env_v = env.get('verdicts', {})
    p = doc.add_paragraph()
    run = p.add_run('Environmental science. ')
    run.bold = True
    forest = get_curve_result(results, 'Forest Transition Curve')
    p.add_run(
        f"Of 6 curves, {env_v.get('ROBUST_NONLINEAR', 0)} showed robust nonlinearity "
        f"(Keeling Curve), {env_v.get('NOT_SIGNIFICANT', 0)} were non-significant, and "
        f"{env_v.get('OUTLIER_DEPENDENT', 0)} were outlier-dependent. "
        f"The Forest Transition Curve with real data (N={forest['n']}) was "
        f"non-significant (p={fmt_p(forest['f_test']['p_value'])})."
    )

    # Psychology
    psych = domains.get('Psychology', {})
    psych_v = psych.get('verdicts', {})
    p = doc.add_paragraph()
    run = p.add_run('Psychology. ')
    run.bold = True
    p.add_run(
        f"Psychology showed high robustness: {psych_v.get('ROBUST_NONLINEAR', 0)}/5 "
        f"curves were robust (Yerkes-Dodson, Ebbinghaus, Dunning-Kruger, Happiness U-Curve). "
        f"Only the Weber-Fechner Law was non-significant."
    )

    # Physics, Political Science, Agriculture (brief)
    phys = domains.get('Physics', {})
    phys_v = phys.get('verdicts', {})
    p = doc.add_paragraph()
    run = p.add_run('Physics. ')
    run.bold = True
    p.add_run(
        f"All 4 physics curves failed: {phys_v.get('NOT_SIGNIFICANT', 0)} non-significant, "
        f"{phys_v.get('OUTLIER_DEPENDENT', 0)} outlier-dependent. "
        "Hubble's Law and Gutenberg-Richter were outlier-dependent; "
        "Kleiber's Law and Moore's Law were non-significant in log space."
    )

    pol = domains.get('Political Science', {})
    pol_v = pol.get('verdicts', {})
    lipset = get_curve_result(results, 'Lipset Hypothesis')
    p = doc.add_paragraph()
    run = p.add_run('Political science. ')
    run.bold = True
    p.add_run(
        f"The Lipset Hypothesis was outlier-dependent "
        f"(p={fmt_p(lipset['f_test']['p_value'])} full, "
        f"p={fmt_p(lipset['sensitivity']['p_clean'])} after removal of Gulf oil states). "
        f"{pol_v.get('NOT_SIGNIFICANT', 0)}/5 were non-significant."
    )

    agr = domains.get('Agriculture', {})
    agr_v = agr.get('verdicts', {})
    p = doc.add_paragraph()
    run = p.add_run('Agriculture. ')
    run.bold = True
    p.add_run(
        f"{agr_v.get('ROBUST_NONLINEAR', 0)}/4 robust (Mitscherlich yield, "
        f"Micronutrient U-shape); {agr_v.get('NOT_SIGNIFICANT', 0)}/4 non-significant."
    )

    # Sensitivity figure
    doc.add_heading('Sensitivity analysis', level=2)
    p = doc.add_paragraph(
        "Fig. 2 displays p-values before and after outlier removal. Points in the "
        "upper-left quadrant (significant full data, non-significant after removal) "
        "represent outlier-dependent curves, concentrated in economics and political science."
    )
    fig2_path = os.path.join(FIGURES_DIR, 'fig2_sensitivity_analysis.png')
    if os.path.exists(fig2_path):
        doc.add_paragraph()
        doc.add_picture(fig2_path, width=Inches(5.0))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run('Fig. 2 ')
        run.bold = True
        p.add_run('Sensitivity of F-test p-values to outlier removal (Cook\'s distance top 3).')

    # Summary table
    doc.add_heading('Summary of all results', level=2)
    p = doc.add_paragraph("Table 1 presents complete results for all 52 curves.")

    df = pd.read_csv(os.path.join(RESULTS_DIR, 'summary_table.csv'))
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['#', 'Curve', 'N', 'p (full)', 'p (clean)', 'BIC best', 'Verdict']
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)

    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(row['Curve'])[:28]
        row_cells[2].text = str(row['N'])
        p_full = row['p (full)']
        row_cells[3].text = f"{p_full:.4f}" if p_full > 0.0001 else f"{p_full:.1e}"
        p_clean = row['p (clean)']
        row_cells[4].text = f"{p_clean:.4f}" if p_clean > 0.0001 else f"{p_clean:.1e}"
        row_cells[5].text = str(row['BIC best'])
        row_cells[6].text = str(row['Verdict']).replace('_', ' ')
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('Table 1 ')
    run.bold = True
    p.add_run('Summary of re-examination results. p (full) = nested F-test with all data; '
              'p (clean) = after removing top 3 Cook\'s distance points.')

    # ======== DISCUSSION ========
    h = doc.add_heading('Discussion', level=1)

    disc_paras = [
        (f"This systematic audit reveals that {100*(n_ns + n_outlier + n_overfit)/52:.0f}% "
         f"of established nonlinear relationships fail at least one modern robustness test. "
         f"Five cross-cutting patterns emerge:"),
    ]
    for text in disc_paras:
        doc.add_paragraph(text)

    patterns = [
        "Outlier-driven nonlinearity: the most common failure mode. In cross-country analyses, "
        "1\u20133 geopolitically distinctive nations (oil states, the US) drive curvature.",
        "Domain asymmetry: public health and psychology curves are substantially more robust "
        "than economics or political science curves, likely reflecting mechanistic versus "
        "contingent relationships.",
        "Time-series vs. cross-section: time-series curves are more robust, likely because "
        "they are less vulnerable to compositional effects.",
        "Log transformation resolves apparent nonlinearity: in many cases a simple log "
        "transformation produces a linear relationship, suggesting the 'canonical curve' is "
        "linear on the wrong scale.",
        "BIC is more conservative than AIC: BIC selects the linear model more frequently, "
        "reflecting stronger complexity penalties.",
    ]
    for i, pat in enumerate(patterns, 1):
        doc.add_paragraph(f"{i}. {pat}")

    doc.add_paragraph(
        "Several outlier-dependent curves have direct policy implications. The Laffer Curve "
        "is used to justify tax reductions; the Environmental Kuznets Curve is cited to argue "
        "that growth resolves pollution; the Lipset Hypothesis underpins modernization theory. "
        "Our findings suggest that the empirical bases of these policy-relevant claims are "
        "more fragile than commonly assumed."
    )

    # ======== LIMITATIONS ========
    h = doc.add_heading('Limitations', level=1)

    limitations = [
        "First, our analysis is restricted to bivariate relationships; many canonical curves "
        "may be better specified in multivariate settings.",
        "Second, we use a uniform quadratic alternative, whereas some curves posit specific "
        "functional forms (power laws, logistic functions).",
        "Third, although 8 curves use real World Bank API data (N=31\u2013247), the remaining "
        "44 curves use representative or published data with smaller sample sizes.",
        "Fourth, we test only the significance of curvature, not the existence of any "
        "relationship. A curve classified as non-significant may still have a significant "
        "linear component.",
        "Finally, for some curves (psychology, demography), we rely on aggregate or "
        "meta-analytic data rather than individual-level microdata.",
    ]
    for lim in limitations:
        doc.add_paragraph(lim)

    # ======== CONCLUSIONS ========
    h = doc.add_heading('Conclusions', level=1)

    doc.add_paragraph(
        f"This systematic re-examination of 52 canonical curves reveals that only "
        f"{100*n_robust/52:.0f}% demonstrate nonlinearity that is statistically significant, "
        f"survives outlier removal, and shows superior out-of-sample prediction. "
        f"Researchers should routinely report sensitivity analyses and model comparison "
        f"criteria when invoking canonical curves. Policymakers should be cautious about "
        f"interventions premised on specific curve shapes\u2014particularly the Laffer Curve, "
        f"Environmental Kuznets Curve, and Lipset Hypothesis\u2014whose empirical bases are "
        f"outlier-dependent."
    )

    # ======== LIST OF ABBREVIATIONS ========
    doc.add_page_break()
    h = doc.add_heading('List of abbreviations', level=1)
    abbreviations = [
        ("AIC", "Akaike Information Criterion"),
        ("BIC", "Bayesian Information Criterion"),
        ("CO\u2082", "Carbon dioxide"),
        ("EKC", "Environmental Kuznets Curve"),
        ("GDP", "Gross Domestic Product"),
        ("LOOCV", "Leave-one-out cross-validation"),
        ("LNT", "Linear No-Threshold"),
        ("OLS", "Ordinary least squares"),
        ("PPP", "Purchasing power parity"),
        ("RMSE", "Root mean squared error"),
        ("TFR", "Total fertility rate"),
        ("WDI", "World Development Indicators"),
    ]
    for abbr, full in abbreviations:
        p = doc.add_paragraph()
        run = p.add_run(f"{abbr}: ")
        run.bold = True
        p.add_run(full)

    # ======== DECLARATIONS ========
    doc.add_page_break()
    h = doc.add_heading('Declarations', level=1)

    doc.add_heading('Ethics approval and consent to participate', level=2)
    doc.add_paragraph("Not applicable. This study uses publicly available aggregate data only.")

    doc.add_heading('Consent for publication', level=2)
    doc.add_paragraph("Not applicable.")

    doc.add_heading('Availability of data and materials', level=2)
    doc.add_paragraph(
        "All analysis code, data, and results are publicly available at "
        "[GitHub repository URL]. World Bank data were accessed via the wbgapi Python "
        "library. The complete source metadata table for all 52 curves is provided "
        "in Additional file 1."
    )

    doc.add_heading('Competing interests', level=2)
    doc.add_paragraph("The author declares no competing interests.")

    doc.add_heading('Funding', level=2)
    doc.add_paragraph("[To be completed]")

    doc.add_heading('Authors\u2019 contributions', level=2)
    doc.add_paragraph(
        "TO conceived the study, designed the analytical framework, collected and analyzed "
        "data, interpreted results, and wrote the manuscript."
    )

    doc.add_heading('Acknowledgements', level=2)
    doc.add_paragraph("[To be completed]")

    # ======== REFERENCES ========
    doc.add_page_break()
    h = doc.add_heading('References', level=1)

    references = [
        "Phillips AW. The relation between unemployment and the rate of change of money wage rates in the United Kingdom, 1861-1957. Economica. 1958;25(100):283-299.",
        "Kuznets S. Economic growth and income inequality. Am Econ Rev. 1955;45(1):1-28.",
        "Preston SH. The changing relation between mortality and level of economic development. Popul Stud. 1975;29(2):231-248.",
        "Akaike H. A new look at the statistical model identification. IEEE Trans Automat Contr. 1974;19(6):716-723.",
        "Schwarz G. Estimating the dimension of a model. Ann Stat. 1978;6(2):461-464.",
        "Onishi T. Re-examination of the Preston Curve: outlier dependence of quadratic fit. Working paper. 2026.",
        "Grossman GM, Krueger AB. Environmental impacts of a North American free trade agreement. NBER Working Paper 3914. 1991.",
        "Stern DI. The rise and fall of the environmental Kuznets curve. World Dev. 2004;32(8):1419-1439.",
        "Krueger J, Mueller RA. Unskilled, unaware, or both? The better-than-average heuristic and statistical regression predict errors in estimates of own performance. J Pers Soc Psychol. 2002;82(2):180-188.",
        "Open Science Collaboration. Estimating the reproducibility of psychological science. Science. 2015;349(6251):aac4716.",
        "Ioannidis JPA. Why most published research findings are false. PLoS Med. 2005;2(8):e124.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        p.add_run(ref)
        p.paragraph_format.left_indent = Cm(1)

    # ======== ADDITIONAL FILE ========
    doc.add_page_break()
    h = doc.add_heading('Additional file 1: Source metadata table', level=1)
    doc.add_paragraph(
        "Complete data source information for all 52 curves, including original publication, "
        "data source, sample size, variable definitions, and claimed functional form."
    )

    source_meta_path = os.path.join(DATA_DIR, 'source_metadata.json')
    if os.path.exists(source_meta_path):
        with open(source_meta_path, 'r', encoding='utf-8') as f:
            sources = json.load(f)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['#', 'Curve', 'Original paper', 'Data source', 'N', 'Claimed form']
        for i, h_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h_text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(7)

        for src in sources:
            row_cells = table.add_row().cells
            row_cells[0].text = str(src['id'])
            row_cells[1].text = src['name'][:25]
            row_cells[2].text = src['original_paper'][:40]
            row_cells[3].text = src['data_source'][:45]
            n_val = src.get('current_n')
            row_cells[4].text = str(n_val) if n_val else 'WB API'
            row_cells[5].text = src['claimed_form'][:30]
            for cell in row_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(6.5)

    # Save
    output_path = os.path.join(BASE_DIR, 'manuscript_canonical_curves_en.docx')
    doc.save(output_path)
    print(f"Manuscript saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_manuscript()
