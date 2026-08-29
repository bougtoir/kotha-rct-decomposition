#!/usr/bin/env python3
"""Build an editable docx with JBS supplementary tables S1-S5."""
import argparse
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from lxml import etree


BASE = os.path.dirname(os.path.abspath(__file__))

CAPTIONS = [
    'Supplementary Table S1: Study-level data for magnesium in AMI',
    'Supplementary Table S2: Study-level data for statins in heart failure',
    'Supplementary Table S3: Existing approaches to mitigate event dilution in RCTs',
    'Supplementary Table S4: Module H assessment checklist mapped to GRADE domains',
    'Supplementary Table S5: Module H assessment --- Standard GRADE vs. KOTHA-enhanced',
]

# Table indices in the source CCTC documents
# KOTHA_Framework_CCTC.docx table order: 0=Abbreviations, 1=Table 1 (existing approaches),
#   2=Table 2 (Module H checklist), 3=Table 3 (Bayesian integration), 4=Table 4 (Module H assessment),
#   5=Table 5 (operating characteristics)
# KOTHA_Framework_CCTC_supplementary_tables.docx table order: 0=S1 (magnesium), 1=S2 (statins)
TABLE_SOURCES = [
    ('KOTHA_Framework_CCTC_supplementary_tables.docx', 0),
    ('KOTHA_Framework_CCTC_supplementary_tables.docx', 1),
    ('KOTHA_Framework_CCTC.docx', 1),
    ('KOTHA_Framework_CCTC.docx', 2),
    ('KOTHA_Framework_CCTC.docx', 4),
]


def _copy_table(src_tbl):
    """Deep-copy a docx table element."""
    return etree.fromstring(etree.tostring(src_tbl._tbl))


def _set_table_font(table):
    """Set Times New Roman 11 pt for all cells."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)


def build(src_docx=None, supp_docx=None, out_docx='KOTHA_Framework_JBS_supplementary_tables.docx'):
    src_docx = src_docx or os.path.join(BASE, 'KOTHA_Framework_CCTC.docx')
    supp_docx = supp_docx or os.path.join(BASE, 'KOTHA_Framework_CCTC_supplementary_tables.docx')

    docs = {
        src_docx: Document(src_docx),
        supp_docx: Document(supp_docx),
        os.path.basename(src_docx): Document(src_docx),
        os.path.basename(supp_docx): Document(supp_docx),
    }

    dst = Document()
    style = dst.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    title = dst.add_paragraph()
    run = title.add_run('Supplementary Tables')
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dst.add_paragraph()

    for caption, (source_path, idx) in zip(CAPTIONS, TABLE_SOURCES):
        src = docs[source_path]
        if idx >= len(src.tables):
            raise IndexError(f'Source {source_path} has {len(src.tables)} tables; requested index {idx}')
        src_tbl = src.tables[idx]

        p = dst.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        placeholder = dst.add_table(rows=1, cols=1)
        placeholder._tbl.getparent().replace(placeholder._tbl, _copy_table(src_tbl))
        _set_table_font(dst.tables[-1])
        dst.add_paragraph()

    dst.save(out_docx)
    print(f'Saved {out_docx}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--src', default=None)
    parser.add_argument('--supp', default=None)
    parser.add_argument('--out', default='KOTHA_Framework_ClinicalTrials_supplementary_tables.docx')
    args = parser.parse_args()
    build(args.src, args.supp, args.out)
