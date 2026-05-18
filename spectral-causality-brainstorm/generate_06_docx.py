"""
Generate docx for 06_spectral_causality_academic.md
Uses python-docx with OML (Office Math Markup Language) for equations.
"""

import os
import re
import lxml.etree as ET
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(BASE_DIR, 'figures')

# ============================================================
# Math XML helper: wrap LaTeX-like text as Word equation via OMML
# ============================================================

def add_math_paragraph(doc, math_text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Add a paragraph containing math as styled monospace text.
    For proper OMML we'd need a full LaTeX-to-OMML converter;
    here we use a clean monospace rendering that preserves readability."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(math_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    return p


def add_display_equation(doc, equation_text, label=None):
    """Add a display equation (centered, with optional label)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(equation_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    if label:
        tab_run = p.add_run(f'    ({label})')
        tab_run.font.size = Pt(10)
    return p


def add_figure(doc, image_path, caption, width=Inches(5.5)):
    """Add a figure with caption to the document, centered."""
    if not os.path.exists(image_path):
        # Fallback: add placeholder text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[図: {caption} — 画像ファイルが見つかりません: {image_path}]')
        run.italic = True
        run.font.size = Pt(9)
        return
    # Image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run()
    run.add_picture(image_path, width=width)
    # Caption paragraph
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(4)
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.italic = True


def generate_docx():
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ============================================================
    # Title
    # ============================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('スペクトル因果性の数理的基礎')
    title_run.bold = True
    title_run.font.size = Pt(18)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        '— 有向グラフのスペクトル理論に基づく因果推論の新しいアプローチ —'
    )
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(80, 80, 80)

    reader_p = doc.add_paragraph()
    reader_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = reader_p.add_run(
        '想定読者: 線形代数（固有値分解）と基礎的な確率論を既習の学部上級生〜大学院生'
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    # ============================================================
    # 1. 導入
    # ============================================================
    doc.add_heading('1. 導入', level=1)

    doc.add_heading('1.1 問題設定', level=2)
    doc.add_paragraph(
        '因果推論（causal inference）の中心的な問い — 「X は Y の原因か？」— に対して、'
        '様々なアプローチが提案されてきた。代表的なものとして：'
    )
    items = [
        '構造方程式モデル（SEM）と do-calculus (Pearl, 2009): 介入に基づく反事実的定義',
        '潜在結果モデル (Rubin, 1974): 処置群と対照群の潜在結果の差',
        'LiNGAM (Shimizu et al., 2006): データの非ガウス性を利用した因果方向の同定',
        'Granger因果 (Granger, 1969): 時系列における予測改善に基づく因果性',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '本稿では、これらとは異なる原理 — グラフのスペクトル構造（固有値・固有ベクトル）から'
        '因果的方向性を読み取る — に基づく手法を定式化する。この手法をスペクトル因果性'
        '（spectral causality）と呼ぶ。'
    )

    doc.add_heading('1.2 基本的着想', level=2)
    doc.add_paragraph(
        'n 個の変数 {X\u2081, ..., X\u2099} の間に因果関係があるとする。'
        'これらの関係を有向グラフ G = (V, E) で表現したとき、グラフのラプラシアン行列のスペクトル'
        '（固有値と固有ベクトル）には、因果的方向性に関する情報が含まれうる。'
    )
    doc.add_paragraph(
        '特に、磁気ラプラシアン（magnetic Laplacian）と呼ばれるエルミート行列を用いると、'
        'エッジの方向性が固有ベクトルの複素位相（complex phase）として符号化され、'
        '因果方向の推定が可能になる。'
    )

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig6_causal_dag.png'),
        '図1: DirectLiNGAMによる推定因果DAG（UCI心疾患データ, n=297）。'
        '辺の太さは効果量 |Bᵢⱼ| に比例。Age が因果的最上流に位置する。')

    doc.add_heading('1.3 本稿の構成', level=2)
    doc.add_paragraph(
        '§2でグラフラプラシアンの基礎を復習し、§3で磁気ラプラシアンを導入する。'
        '§4でスペクトル因果性を厳密に定式化し（§4.1.1 でデータ駆動の非対称統計量DPIを導入する）、'
        '§5でHodge分解との関係を示す。'
        '§6で既存手法との相補性と増強拡張性を論じ、§7で関連研究・先行文献を概観し、'
        '§8で実データ（UCI心疾患データ）への適用例を示す。§9以降で詳細解析と展望を議論する。'
    )

    doc.add_paragraph('本手法の主要な貢献は以下の5点に集約される：')
    contributions = [
        'DPI（非対称統計量）の導入: α = 0（ドメイン知識なし）でも有向辺を検出し因果方向を推定可能'
        '（UCI心疾患データで9本の有向辺, 67% LiNGAM方向一致）',
        'DAG仮定不要: Hodge分解によりDAG成分（勾配）とフィードバック成分（カール）を自然に分離',
        'ドメイン知識による精度向上: r_gradient: 0.581（α = 0）→ 0.859（α = 0.6）と滑らかに改善',
        'LiNGAM連携: ドメイン知識がない場合、LiNGAMの推定DAGから高確信辺を C に設定可能',
        'ECDアンサンブル: Hill の9基準の網羅性向上（H6/H7/H9をカバー）',
    ]
    for c in contributions:
        doc.add_paragraph(c, style='List Number')

    # ============================================================
    # 2. 準備
    # ============================================================
    doc.add_heading('2. 準備：グラフラプラシアンの基礎', level=1)

    doc.add_heading('2.1 無向グラフのラプラシアン', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定義 2.1')
    run.bold = True
    p.add_run(
        '（グラフラプラシアン）重み付き無向グラフ G = (V, E, w)（|V| = n, w: E → ℝ₊）に対して、'
        '重み付き隣接行列 W ∈ ℝⁿˣⁿ, 次数行列 D = diag(d₁, ..., dₙ)（dᵢ = Σⱼ Wᵢⱼ）を用いて、以下を定義する：'
    )

    add_display_equation(doc, 'L = D − W    （非正規化ラプラシアン）')
    add_display_equation(doc, '\u2112 = I − D⁻¹ᐟ² W D⁻¹ᐟ²    （正規化ラプラシアン）')

    p = doc.add_paragraph()
    run = p.add_run('命題 2.1')
    run.bold = True
    p.add_run('（基本性質）L および \u2112 について以下が成り立つ：')

    props = [
        '(i) L は対称半正定値行列であり、固有値は 0 = λ₁ ≤ λ₂ ≤ ... ≤ λₙ を満たす。',
        '(ii) λ₁ = 0 に対応する固有ベクトルは 1 = (1, ..., 1)ᵀ（定数ベクトル）。',
        '(iii) λ₂ > 0 であることは、G が連結であることと同値（Fiedler値）。',
        '(iv) 任意のベクトル f ∈ ℝⁿ に対して、fᵀLf = Σ₍ᵢ,ⱼ₎∈E wᵢⱼ(fᵢ − fⱼ)² ≥ 0。',
    ]
    for prop in props:
        doc.add_paragraph(prop, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('証明のスケッチ')
    run.italic = True
    p.add_run(
        '：(iv) は L の二次形式を展開すれば直接示せる。(i) は (iv) から従う。'
        '(ii) は L1 = 0 の直接計算による。(iii) は代数的連結度の定理。 □'
    )

    doc.add_paragraph(
        '性質 (iv) は重要である：fᵀLf が小さいほど、f は隣接ノードで類似した値をとる — '
        'つまり、ラプラシアンの低固有値固有ベクトルはグラフ上で滑らかな信号を表す。'
    )

    doc.add_heading('2.2 スペクトル分解の幾何学的意味', level=2)
    doc.add_paragraph(
        '\u2112 のスペクトル分解 \u2112 = UΛUᵀ（U = [u₁, ..., uₙ], Λ = diag(λ₁, ..., λₙ)）において：'
    )
    items = [
        'uₖ の各成分 uₖ(i) = ノード i が第 k 固有モードにどれだけ「荷重（load）」するかを表す',
        'λₖ = 第 k モードの「周波数」（大きいほど高周波 = 局所変動）',
        'u₂（第2固有ベクトル, Fiedler vector）はグラフの最適2分割を与える',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'この枠組みは、信号処理におけるフーリエ変換のグラフ上への一般化'
        '（Graph Signal Processing; GSP）の基礎となっている (Shuman et al., 2013)。'
    )

    doc.add_heading('2.3 問題：無向ラプラシアンは方向性を失う', level=2)
    doc.add_paragraph(
        'L = D − W は対称行列であるため、エッジの方向性 i → j と j → i を区別できない。'
        '因果推論では「X が Y の原因」という方向性が本質的であり、無向ラプラシアンでは情報が不足する。'
    )
    doc.add_paragraph(
        '有向グラフのラプラシアン Ld = Dout − W を直接用いる手もあるが、Ld は一般に非対称であり、'
        '固有値が複素数になりうる。これは理論的に扱いにくい。'
    )

    # ============================================================
    # 3. 磁気ラプラシアン
    # ============================================================
    doc.add_heading('3. 磁気ラプラシアン：方向性の複素位相符号化', level=1)

    doc.add_heading('3.1 物理的背景', level=2)
    doc.add_paragraph(
        '磁気ラプラシアンの名前は量子力学に由来する。磁場 B 中の荷電粒子のハミルトニアンは '
        'H = (p − eA)²/2m（A はベクトルポテンシャル）であり、粒子が閉じた経路を一周すると '
        'Aharonov-Bohm 位相 exp(i∮A·dr) を獲得する。この位相の向き依存性が、'
        'グラフ上のエッジ方向性の符号化に利用できる。'
    )

    doc.add_heading('3.2 定義', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定義 3.1')
    run.bold = True
    p.add_run(
        '（磁気ラプラシアン; de Resende & da Costa, 2020; Zhang et al., 2021）'
        '重み付き有向グラフ G = (V, E, w) と電荷パラメータ q ∈ [0, 0.5] に対して、'
        'エルミート隣接行列 H⁽ᑫ⁾ ∈ ℂⁿˣⁿ を以下で定義する：'
    )

    add_display_equation(doc, 'H⁽ᑫ⁾ᵢⱼ = wᵢⱼ · exp(i · 2πq · σᵢⱼ)')

    doc.add_paragraph(
        'ここで σᵢⱼ ∈ {−1, 0, +1} はエッジの方向性符号であり：'
    )
    items = [
        'σᵢⱼ = +1  if i → j',
        'σᵢⱼ = −1  if j → i',
        'σᵢⱼ =  0  if エッジなし',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '重み wᵢⱼ は対称化して用いる（wᵢⱼ = wⱼᵢ = (w_orig_ij + w_orig_ji)/2）。'
        '正規化磁気ラプラシアンを以下で定義する：'
    )

    add_display_equation(doc, '\u2112⁽ᑫ⁾ = I − D⁻¹ᐟ² H⁽ᑫ⁾ D⁻¹ᐟ²')

    doc.add_paragraph('ここで D = diag(d₁, ..., dₙ), dᵢ = Σⱼ |H⁽ᑫ⁾ᵢⱼ|。')

    # Proposition 3.1
    p = doc.add_paragraph()
    run = p.add_run('命題 3.1')
    run.bold = True
    p.add_run('（磁気ラプラシアンの基本性質）')

    props = [
        '(i) H⁽ᑫ⁾ はエルミート行列である：H⁽ᑫ⁾ⱼᵢ = H̄⁽ᑫ⁾ᵢⱼ。',
        '(ii) \u2112⁽ᑫ⁾ はエルミート半正定値であり、固有値は実数かつ非負。',
        '(iii) 固有ベクトルは一般に複素数値をとる。',
        '(iv) q = 0 のとき、\u2112⁽⁰⁾ は通常の正規化ラプラシアン \u2112 に退化する。',
    ]
    for prop in props:
        doc.add_paragraph(prop, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('命題 3.1 (i) の証明')
    run.bold = True
    run.italic = True
    p.add_run('：')

    add_display_equation(doc, 'H⁽ᑫ⁾ⱼᵢ = wⱼᵢ · exp(i · 2πq · σⱼᵢ)')

    doc.add_paragraph(
        'wⱼᵢ = wᵢⱼ（対称化済み）かつ σⱼᵢ = −σᵢⱼ より：'
    )

    add_display_equation(doc,
        'H⁽ᑫ⁾ⱼᵢ = wᵢⱼ · exp(−i · 2πq · σᵢⱼ) = w̄ᵢⱼ · exp(i · 2πq · σᵢⱼ) = H̄⁽ᑫ⁾ᵢⱼ   □')

    doc.add_heading('3.3 電荷パラメータ q の意味', level=2)
    doc.add_paragraph(
        'q は方向性に対する感度を制御するパラメータである：'
    )

    # Table for q values
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['q', '位相 2πq', '効果']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    rows_data = [
        ['0', '0', '方向性を完全無視。exp(i·0) = 1 より実行列に退化'],
        ['0.25', 'π/2', '最大方向性感度。exp(iπ/2) = i, exp(−iπ/2) = −i'],
        ['0.5', 'π', '方向を反転。exp(iπ) = −1'],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('注意 3.1')
    run.bold = True
    p.add_run(
        ' q = 0.25 のとき、H⁽ᑫ⁾ᵢⱼ = i·wᵢⱼ（i → j のエッジ）かつ '
        'H⁽ᑫ⁾ⱼᵢ = −i·wᵢⱼ となり、方向性が虚数単位 i によって最も鋭く分離される。'
    )

    doc.add_heading('3.4 固有ベクトルの複素位相と方向性', level=2)
    doc.add_paragraph(
        '\u2112⁽ᑫ⁾ の固有ベクトル uₖ ∈ ℂⁿ の各成分は極形式で表すことができる：'
    )

    add_display_equation(doc, 'uₖ(j) = |uₖ(j)| · exp(i · θₖ(j))')

    doc.add_paragraph(
        'ここで |uₖ(j)| は振幅（ノード j がモード k にどれだけ荷重するか）、'
        'θₖ(j) = arg(uₖ(j)) は位相角である。'
    )
    p = doc.add_paragraph()
    run = p.add_run('核心的な主張')
    run.bold = True
    p.add_run(
        '：q > 0 のとき、位相角 θₖ(j) の順序が因果的フローの方向を反映する。'
        '因果の上流（原因側）のノードと下流（結果側）のノードは、'
        '固有ベクトル上で異なる位相角を持つ。'
    )

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig2_magnetic_laplacian_q.png'),
        '図3: 磁気ラプラシアン固有ベクトルの複素平面上の分布。'
        '各ノードの位相角 θₖ(j) が因果的上流（左）から下流（右）へと回転する。')

    # ============================================================
    # 4. スペクトル因果性の定式化
    # ============================================================
    doc.add_heading('4. スペクトル因果性の定式化', level=1)

    doc.add_heading('4.1 ユーティリティ有向グラフ', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定義 4.1')
    run.bold = True
    p.add_run(
        '（ユーティリティ有向グラフ）n 個の変数 {X₁, ..., Xₙ} に対して、'
        'ユーティリティ関数 U: {1, ..., n}² → ℝ≥₀ を：'
    )

    add_display_equation(doc, 'U(i, j) = 「変数 Xᵢ の情報が変数 Xⱼ に関する問いにどれだけ有用か」')

    doc.add_paragraph('と定義する。ユーティリティ有向グラフ G_U = (V, E, w, σ) は：')
    items = [
        'V = {1, ..., n}',
        'w(i, j) = (U(i, j) + U(j, i)) / 2（対称化された重み）',
        'σ(i, j) = sign(U(i, j) − U(j, i))（方向性符号）',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # -- 4.1.1 DPI --
    doc.add_heading('4.1.1 データ駆動成分の構成：方向性予測指標（DPI）', level=3)

    doc.add_paragraph(
        'ユーティリティ関数のデータ駆動成分として、従来は |ρ̂ᵢⱼ|（相関係数の絶対値）が'
        '用いられてきた。しかし |ρ̂ᵢⱼ| = |ρ̂ⱼᵢ| であるため、データ駆動成分が完全に対称となり、'
        'α = 0（ドメイン知識なし）では方向性信号がゼロになる。これは統計的因果推論としての'
        '要件を満たさない。この理論的限界を克服するため、方向性予測指標（Directional '
        'Predictability Index, DPI）を提案する。'
    )

    p = doc.add_paragraph()
    run = p.add_run('定義 4.1a')
    run.bold = True
    p.add_run(
        '（方向性予測指標; DPI）n 個の変数 {X₁, ..., Xₙ} の観測データ X ∈ ℝ^(N×n) に対して、'
        'DPI行列 D_DPI ∈ ℝ^(n×n) を：'
    )

    add_display_equation(doc, 'D_DPI(i → j) = |ρ̂ᵢⱼ| · (1 + γ · Ā(i, j))')

    doc.add_paragraph(
        'と定義する。ここで γ > 0 は方向性強度パラメータ（本稿では γ = 1）であり、'
        'Ā(i,j) は 3 つの正規化非対称統計量の平均である：'
    )

    add_display_equation(doc, 'Ā(i,j) = (1/3) [Â_reg(i,j) + Â_ANM(i,j) + Â_ent(i,j)]')

    doc.add_paragraph('各成分は以下のように定義される：')

    doc.add_paragraph(
        '(i) 回帰係数非対称性 Â_reg：非標準化データにおける単回帰係数 '
        'β_j|i = Cov(Xᵢ, Xⱼ)/Var(Xᵢ) は、Var(Xᵢ) ≠ Var(Xⱼ) のとき '
        '|β_j|i| ≠ |β_i|j| となり非対称である。この非対称性を [−1, 1] に正規化する。'
    )
    doc.add_paragraph(
        '(ii) ANM残差独立性 Â_ANM：加法的ノイズモデル（Additive Noise Model）の原理に基づき、'
        '各ペア (i, j) に対して Xⱼ = βXᵢ + ε の残差 ε̂ と Xᵢ の独立性を HSIC'
        '（Hilbert-Schmidt Independence Criterion; カーネル帯域幅はメディアンヒューリスティック）'
        'で評価する。HSIC値が小さいほど独立性が高く、Xᵢ → Xⱼ の方向がもっともらしい。'
    )
    doc.add_paragraph(
        '(iii) 条件付きエントロピー縮減 Â_ent：Xᵢ を知ることによる Xⱼ のエントロピー縮減量 '
        'H(Xⱼ) − H(Xⱼ|Xᵢ) を kNN推定量で計算する。'
        'H(Xⱼ) − H(Xⱼ|Xᵢ) ≠ H(Xᵢ) − H(Xᵢ|Xⱼ) のとき方向性情報を持つ。'
    )

    p = doc.add_paragraph()
    run = p.add_run('命題 4.0a')
    run.bold = True
    p.add_run(' D_DPI は一般に非対称であり、D_DPI(i → j) ≠ D_DPI(j → i) である。')

    p = doc.add_paragraph()
    run = p.add_run('証明')
    run.italic = True
    p.add_run(
        '：Ā(i,j) = −Ā(j,i)（各成分の正規化非対称性は反対称）であるから、'
        'D_DPI(i → j) = |ρ̂ᵢⱼ|(1 + γĀ(i,j)) ≠ |ρ̂ᵢⱼ|(1 − γĀ(i,j)) = D_DPI(j → i)'
        '（Ā(i,j) ≠ 0 のとき）。□'
    )

    p = doc.add_paragraph()
    run = p.add_run('ユーティリティ関数のハイブリッド構成')
    run.bold = True
    p.add_run('：上記を用いて')

    add_display_equation(doc, 'U(i, j) = α · C_domain(i, j) + (1 − α) · D_DPI(i → j)')

    doc.add_paragraph(
        'とする。α = 0（ドメイン知識なし）でも D_DPI の非対称性により方向性信号が保たれ、'
        'α > 0 でドメイン知識の注入により精度が向上する。'
    )

    doc.add_heading('4.2 スペクトル因果結合度と因果方向', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定義 4.2')
    run.bold = True
    p.add_run(
        '（スペクトル因果結合度; Spectral Causal Coupling, SCC）'
        '磁気ラプラシアン \u2112⁽ᑫ⁾ の固有値分解に対して、ノード i, j のスペクトル因果結合度を：'
    )

    add_display_equation(doc, 'SCC(i, j) = Σₖ f(λₖ) · |uₖ(i)| · |uₖ(j)| · cos(θₖ(i) − θₖ(j))')

    doc.add_paragraph(
        'と定義する。ここで f: ℝ≥₀ → ℝ≥₀ は固有値重み関数（典型的には f(λ) = λ）、'
        'θₖ(i) = arg(uₖ(i))。'
    )

    p = doc.add_paragraph()
    run = p.add_run('命題 4.1')
    run.bold = True
    p.add_run(' SCC は対称である：SCC(i, j) = SCC(j, i)。')

    p = doc.add_paragraph()
    run = p.add_run('証明')
    run.italic = True
    p.add_run('：cos(α − β) = cos(β − α) より直ちに従う。 □')

    doc.add_paragraph(
        'SCCは因果的結合の強さを測るが、方向は測れない。方向の定量化には以下を用いる。'
    )

    p = doc.add_paragraph()
    run = p.add_run('定義 4.3')
    run.bold = True
    p.add_run('（スペクトル因果方向; Spectral Causal Direction, SCD）')

    add_display_equation(doc, 'SCD(i, j) = Σₖ f(λₖ) · |uₖ(i)| · |uₖ(j)| · sin(θₖ(i) − θₖ(j))')

    p = doc.add_paragraph()
    run = p.add_run('命題 4.2')
    run.bold = True
    p.add_run(' SCD は反対称である：SCD(i, j) = −SCD(j, i)。')

    p = doc.add_paragraph()
    run = p.add_run('証明')
    run.italic = True
    p.add_run('：sin(α − β) = −sin(β − α) より直ちに従う。 □')

    p = doc.add_paragraph()
    run = p.add_run('系 4.1')
    run.bold = True
    p.add_run('（自己因果方向はゼロ）SCD(i, i) = 0。')

    doc.add_paragraph(
        'SCD(i, j) > 0 は「i から j への因果的方向」を、SCD(i, j) < 0 は逆方向を示唆する。'
    )

    doc.add_heading('4.3 SCC と SCD の統一的理解', level=2)

    doc.add_paragraph(
        'SCC と SCD は、複素内積の実部と虚部として統一的に理解できる。'
    )

    p = doc.add_paragraph()
    run = p.add_run('命題 4.3')
    run.bold = True
    p.add_run('（複素因果指標）以下の Complex Causal Index (CCI) を定義すると：')

    add_display_equation(doc, 'CCI(i, j) = Σₖ f(λₖ) · |uₖ(i)| · |uₖ(j)| · exp(i(θₖ(i) − θₖ(j)))')

    doc.add_paragraph('SCC と SCD は CCI の実部と虚部に対応する：')

    add_display_equation(doc, 'SCC(i, j) = Re[CCI(i, j)],    SCD(i, j) = Im[CCI(i, j)]')

    p = doc.add_paragraph()
    run = p.add_run('証明')
    run.italic = True
    p.add_run('：exp(iα) = cosα + i·sinα（Euler公式）を適用すればよい。 □')

    p = doc.add_paragraph()
    run = p.add_run('幾何学的解釈')
    run.bold = True
    p.add_run(
        '：CCI を複素平面上のベクトルとみなすと、偏角 arg(CCI(i,j)) が因果の方向を、'
        '絶対値 |CCI(i,j)| が因果的結合の強さを表す。'
    )

    doc.add_heading('4.4 SCD行列の性質', level=2)

    p = doc.add_paragraph()
    run = p.add_run('命題 4.4')
    run.bold = True
    p.add_run('')

    props = [
        '(i) S は歪対称（skew-symmetric）：S = −Sᵀ。',
        '(ii) tr(S) = 0（対角成分はすべて0）。',
        '(iii) q = 0 のとき S = O（ゼロ行列）。すなわち、方向性情報がなければ因果方向は推定できない。',
    ]
    for prop in props:
        doc.add_paragraph(prop, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('証明')
    run.italic = True
    p.add_run(
        '：(i) は命題4.2の行列版。(ii) は系4.1から。'
        '(iii) は q = 0 のとき θₖ(i) = 0 または π（実固有ベクトル）なので '
        'sin(θₖ(i) − θₖ(j)) = 0。 □'
    )

    doc.add_paragraph(
        '性質 (iii) は重要である：スペクトル因果性は、方向性情報（q > 0）がなければ機能しない。'
        'これはLiNGAMが非ガウス性なしには機能しないのと対照的である。'
    )

    doc.add_heading('4.5 因果順序の推定', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定義 4.4')
    run.bold = True
    p.add_run('（スペクトル因果スコア）各ノード i のスペクトル因果スコアを：')

    add_display_equation(doc, 's(i) = Σⱼ≠ᵢ SCD(i, j)')

    doc.add_paragraph(
        'と定義する。s(i) が大きいノードほど「原因側（上流）」、小さいノードほど「結果側（下流）」。'
    )

    p = doc.add_paragraph()
    run = p.add_run('注意 4.2')
    run.bold = True
    p.add_run(' S の歪対称性より Σᵢ s(i) = 0 であり、スコアは零和（zero-sum）である。')

    # ============================================================
    # 5. Hodge分解
    # ============================================================
    doc.add_heading('5. Hodge分解：因果フローの直交分解', level=1)

    doc.add_heading('5.1 グラフ上の微分形式', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定義 5.1')
    run.bold = True
    p.add_run('（鎖複体）グラフ G = (V, E) に対して、以下の線形写像を定義する：')

    items = [
        '0-コチェイン C⁰ = ℝ|V|（ノード上の関数）',
        '1-コチェイン C¹ = ℝ|E|（エッジ上の関数 = フロー）',
        'コバウンダリ作用素 δ₀: C⁰ → C¹：(δ₀f)(i→j) = f(j) − f(i)（勾配）',
        'コバウンダリ作用素 δ₁: C¹ → C²：三角形上のカール',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2 Hodge分解定理', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定理 5.1')
    run.bold = True
    p.add_run('（グラフ上の Hodge 分解; Jiang et al., 2011）'
              '任意の 1-コチェイン（エッジフロー）ω ∈ C¹ は、以下のように直交分解される：')

    add_display_equation(doc, 'ω = δ₀φ  +  δ₁*ψ  +  h')
    doc.add_paragraph('       勾配成分    カール成分    調和成分')

    # Table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['成分', '数学的意味', '因果的解釈']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    rows_data = [
        ['δ₀φ（勾配）', 'ポテンシャル差に駆動されるフロー', '因果的フロー（DAG的な一方向の流れ）'],
        ['δ₁*ψ（カール）', '局所的な循環フロー', 'フィードバックループ（局所的な相互作用）'],
        ['h（調和）', '大域的な循環フロー', '恒常性維持（全身性の調節メカニズム）'],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig3_hodge_decomposition.png'),
        '図4: Hodge分解による情報フローの直交分解。'
        '勾配成分（DAG的因果フロー）85.9%、カール成分（フィードバック）14.1%。')

    doc.add_heading('5.3 因果ポテンシャル', level=2)

    p = doc.add_paragraph()
    run = p.add_run('定義 5.2')
    run.bold = True
    p.add_run(
        '（因果ポテンシャル）勾配成分 δ₀φ におけるポテンシャル関数 φ: V → ℝ を因果ポテンシャルと呼ぶ。'
        'φ は以下の最小二乗問題の解として求まる：'
    )

    add_display_equation(doc, 'φ = argmin_φ̃ Σ₍ᵢ,ⱼ₎∈E (ω(i,j) − (φ̃(j) − φ̃(i)))²')

    doc.add_paragraph('これはグラフラプラシアンに関するポアソン方程式 Lφ = δ₀*ω に帰着する。')

    p = doc.add_paragraph()
    run = p.add_run('注意 5.1')
    run.bold = True
    p.add_run(' 勾配成分のエネルギー比：')

    add_display_equation(doc, 'r_gradient = ‖δ₀φ‖² / ‖ω‖²')

    doc.add_paragraph(
        'は、データがDAG的構造にどの程度適合するかの指標となる。'
        'r_gradient ≈ 1 ならばDAG仮定が妥当、r_gradient ≪ 1 ならばフィードバックが支配的である。'
    )

    # ============================================================
    # 6. 既存手法との相補性と増強拡張性
    # ============================================================
    doc.add_heading('6. 既存手法との相補性と増強拡張性', level=1)

    doc.add_paragraph(
        '本節では、スペクトル因果性を既存手法と競合するものとしてではなく、'
        'それらと相補的に機能し、相互に増強する枠組みとして位置づける。'
    )

    doc.add_heading('6.1 LiNGAMとの相補性', level=2)

    doc.add_paragraph(
        'LiNGAM (Shimizu et al., 2006) は x = Bx + e（e ~ 非ガウス独立）を仮定し、'
        '非ガウス性を利用して因果効果行列 B を同定する。二手法の能力は排他的ではなく、'
        '相補的に分布する。'
    )

    # 相補性テーブル（4列）
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['能力', 'LiNGAMが提供', 'スペクトル因果性が提供', '相補的結合']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    rows_data = [
        ['因果方向の同定', '◎ 識別可能性保証', '○ DPIによる方向推定', 'LiNGAM高確信辺でスペクトルを初期化'],
        ['効果量の定量', '◎ Bijを直接推定', '△ SCD値で相対強度', 'LiNGAM効果量 × SCD方向性の統合'],
        ['フィードバック検出', '✗（DAG仮定で不可）', '◎ Hodgeカール成分', 'LiNGAM DAGの「残差」をスペクトルが回収'],
        ['グローバル因果構造', '△ 辺ごとの独立推定', '◎ スペクトル分解で全体構造', '局所 + 大域の統合'],
        ['Hill H6/H7/H9', '✗', '◎ ユーティリティ経由', 'ECDアンサンブルで9基準を網羅'],
        ['識別可能性', '◎ 理論保証あり', '△ 仮説段階', 'LiNGAMの理論保証が「アンカー」に'],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph(
        '核心的洞察: LiNGAMは局所的・統計的に因果方向を同定するのに対し、'
        'スペクトル因果性は大域的・構造的に因果フローを捕捉する。'
        '両者は原理的に異なる情報源を用いるため、一方の弱点が他方の強みで補完される。'
    )

    doc.add_heading('6.1.1 双方向の増強', level=3)
    doc.add_paragraph('LiNGAM → スペクトル因果性の増強:')
    for item in [
        'LiNGAMの推定DAGから高確信辺を抽出し、Cdomainとして注入する「二段ロケット」戦略',
        'LiNGAMの識別可能性保証がスペクトル因果性の方向推定の「外部検証基準」として機能',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('スペクトル因果性 → LiNGAMの増強:')
    for item in [
        'LiNGAMが禁止するフィードバックループをHodgeカール成分で定量化 → DAG仮定の妥当性を事後検証',
        '因果ポテンシャル φ による「介入可能性」の定量化は、LiNGAMの因果順序に臨床的解釈を付与',
        'Hill H6/H7/H9の計算的評価は、LiNGAMの統計的因果推定に疫学的妥当性を追加',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.2 Granger因果との相補性', level=2)

    doc.add_paragraph(
        'Granger因果は時系列データ（縦断）を対象とし時間的先行性を方向性の源泉とする。'
        'スペクトル因果性は横断スナップショットを対象とし構造的非対称性（DPI + ドメイン知識）を用いる。'
        '両者は因果推論の時間軸に関して相補的であり、時間ラグ付きユーティリティグラフの構築により'
        '統合可能である。'
    )

    doc.add_heading('6.3 因果の梯子における位置づけ: Level間の架橋', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['レベル', '問い', '代表手法']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    rows_data = [
        ['3: 反事実', '「もし X=x だったら Y は？」', '潜在結果モデル, do-calculus'],
        ['2: 介入', '「X を操作したら Y は変わるか？」', 'RCT, IV, MR'],
        ['1.5: 情報的因果 ★', '「X を知ると Y について何が分かるか？」', 'スペクトル因果性, Utility Causality'],
        ['1: 関連', '「X と Y は共変動するか？」', '相関, 回帰'],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph(
        '重要なのは、スペクトル因果性が Level 1.5 に閉じ込められるのではなく、'
        '他のレベルの手法と結合することでLevel間を架橋する点にある。'
        'DPIが相関（Level 1）から方向性情報を抽出し Level 1.5へ引き上げ、'
        '因果ポテンシャル φ が介入可能性を示唆して Level 2の優先順位付けに貢献する。'
    )

    doc.add_heading('6.4 Hillの9基準: 単独手法の限界とアンサンブルによる網羅', level=2)

    doc.add_paragraph(
        'いかなる単独の計算的手法もHillの9基準を網羅できない。'
        'これは単独手法の限界であると同時に、相補的アンサンブルの必要性を動機づける。'
    )

    # Hill criteria table (6 columns)
    table = doc.add_table(rows=10, cols=6)
    table.style = 'Table Grid'
    headers = ['Hill基準', 'LiNGAM', 'Granger', 'RCT', 'スペクトル因果性', 'ECD(アンサンブル)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
    rows_data = [
        ['H1: 強さ', '◎', '○', '◎', '○', '◎'],
        ['H2: 一貫性', '△', '△', '△', '○', '◎'],
        ['H3: 特異性', '◎', '◎', '◎', '△', '◎'],
        ['H4: 時間性', '—', '◎', '◎', '○', '◎'],
        ['H5: 量反応', '○', '○', '◎', '△', '○'],
        ['H6: 妥当性', '—', '—', '—', '◎', '◎'],
        ['H7: 整合性', '—', '—', '—', '◎', '◎'],
        ['H8: 実験', '—', '—', '◎', '—', '○'],
        ['H9: 類似性', '—', '—', '—', '◎', '◎'],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph(
        'ECDアンサンブルとして統合することで、単独手法では到達不可能な9基準の広範なカバレッジを実現する。'
    )

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig5_hill_radar.png'),
        '図2: Hill の9基準に対する各手法のカバレッジ。'
        '単独手法ではカバーできない基準がECDアンサンブルにより網羅される。')

    doc.add_heading('6.5 増強拡張性のフレームワーク', level=2)

    doc.add_paragraph(
        'スペクトル因果性の設計は、複数のレイヤーで外部手法・知識源の接続（プラグイン）を'
        '許容するモジュラー構造を持つ。'
    )

    doc.add_heading('6.5.1 DPIのモジュラー拡張', level=3)
    doc.add_paragraph(
        'DPIは3つの非対称統計量の平均として定義されるが、この構成は開かれた設計である。'
        '追加の非対称統計量（転送エントロピー差、LiNGAM Bijの符号、非線形Granger、'
        'Knockoff統計量、LLM因果スコア等）を [-1,1] に正規化すれば即座に統合可能であり、'
        '新たな非対称統計量の発見がそのまま手法全体の精度向上に直結する。'
    )

    doc.add_heading('6.5.2 ユーティリティ関数のインターフェース設計', level=3)
    doc.add_paragraph(
        'U(i,j) = α·Cdomain(i,j) + (1−α)·DDPI(i→j) において、Cdomainは'
        '任意のドメイン知識源（臨床専門家、LiNGAM推定DAG、先行RCT結果、LLMメタ知識、'
        '文献マイニング）を受け入れる汎用インターフェースとして機能する。'
        '知識の種類や品質に応じて α と Cdomain を柔軟に構成でき、'
        '新たなドメイン知識源が利用可能になるたびにシステムの精度を向上させることが可能である。'
    )

    doc.add_heading('6.5.3 Hodge分解の後処理としての汎用性', level=3)
    doc.add_paragraph(
        'Hodge分解はスペクトル因果性に固有のステップではなく、任意の因果推定手法'
        '（LiNGAM, NOTEARS, PC, GES等）の出力を後処理する汎用ツールとしても機能する。'
        '推定されたDAGが実際にどの程度DAG的であるかを r_gradient で定量化し、'
        'カール成分の大きな辺を「DAG仮定の違反が疑われるフィードバック候補」として報告できる。'
    )

    doc.add_heading('6.5.4 段階的精度向上の設計原理', level=3)
    stages = [
        'Stage 0: データのみ（α=0, DPI基本3成分）→ r_gradient=0.581, 67% LiNGAM一致',
        'Stage 1: DPI拡張（追加非対称統計量）→ 成分数 K の増加で方向推定が安定化',
        'Stage 2: 外部DAG注入（LiNGAM → Cdomain, 低α）→ r_gradient=0.859',
        'Stage 3: ドメイン知識注入（専門家/RCT/LLM → Cdomain）→ 知識品質に応じてさらに改善',
        'Stage 4: ECDアンサンブル → Hill 9基準の網羅的評価',
    ]
    for s in stages:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph(
        '各段階は前段階の結果を包含し、新たな情報源の追加が破壊的変更なく精度を向上させる。'
        'この単調な精度向上の保証が、実運用における段階的導入を可能にする。'
    )

    # ============================================================
    # 7. 実データ
    # ============================================================
    doc.add_heading('7. 実データによる例示', level=1)

    doc.add_heading('7.1 データと変数', level=2)
    doc.add_paragraph(
        'UCI Heart Disease Dataset (Cleveland subset; Detrano et al., 1989) の連続変数5つを用いた：'
        'X = (Age, RestingBP, Cholesterol, MaxHR, STDepression)。標本数 n = 297。'
    )

    doc.add_heading('7.2 LiNGAM による因果順序', level=2)
    doc.add_paragraph(
        '推定因果順序: Age → MaxHR → STDep → RestBP → Chol'
    )
    doc.add_paragraph('主要な因果効果:')
    items = [
        'B₄₂ = −0.395: Age → MaxHR（加齢による最大心拍数低下）',
        'B₂₁ = +0.309: Age → RestingBP（加齢による血圧上昇）',
        'B₅₄ = −0.348: MaxHR → STDepression（運動耐容能低下による心筋虚血）',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.3 Hodge分解の結果', level=2)

    add_display_equation(doc, '‖δ₀φ‖² / ‖ω‖² = 85.9%    （勾配成分 = DAG的因果フロー）')
    add_display_equation(doc, '‖δ₁*ψ‖² / ‖ω‖² = 14.1%    （カール成分 = フィードバック）')

    doc.add_paragraph('因果ポテンシャル φ（降順 = 因果的上流から）:')

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['順位', '変数', 'φ']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    rows_data = [
        ['1', 'Age', '0.000'],
        ['2', 'Cholesterol', '−0.168'],
        ['3', 'Resting BP', '−0.204'],
        ['4', 'Max Heart Rate', '−0.204'],
        ['5', 'ST Depression', '−0.324'],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph(
        'Age が最上流、ST Depression が最下流という結果は臨床的に妥当である。'
        'LiNGAM の因果順序との Kendall 順位相関は τ = 0.50。'
    )

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig4_direction_comparison.png'),
        '図5: 3手法（LiNGAM, スペクトル因果性, Granger）による因果方向の比較。'
        'スペクトル因果性はLiNGAMと67%の方向一致を示す。')

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig7_lingam_vs_spectral.png'),
        '図6: LiNGAM DAG vs スペクトル因果性 DCG の3条件比較。'
        '左: LiNGAM（DAG仮定）、中: スペクトル因果性（α=0.6）、右: スペクトル因果性（α=0, DPIのみ）。')

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig9_ecd_pruning_analysis.png'),
        '図7: ECDアンサンブルとプルーニング解析。'
        '(A) 因果ポテンシャル比較、(B) フィードバック率別プルーニング、(C) 因果上流性 vs 介入可能性。')

    # ============================================================
    # 8. DAG転移点の解析：ドメイン知識はどこまで必要か
    # ============================================================
    doc.add_heading('8. DAG転移点の解析：ドメイン知識はどこまで必要か', level=1)

    doc.add_paragraph(
        '本節は、スペクトル因果性の実用上最も重要な問いに答える：'
        '「DAG的構造が出現するために、ドメイン知識はどの程度必要か？」。'
        '結果は直観に反する発見の連続であり、方法の設計原理に深く関わる。'
    )

    # --- 8.1 α掃引：旧モデルの不連続相転移 ---
    doc.add_heading('8.1 α掃引実験：二つのモデルの対比', level=2)

    p = doc.add_paragraph()
    run = p.add_run('8.1.1 旧モデル（|ρ| ベース）：不連続相転移の発見')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        '旧モデル（|ρ̂ᵢⱼ| をデータ駆動成分に使用）で α を 0 → 1 まで掃引すると、'
        '予想外の不連続相転移が観測される：'
    )

    # α sweep table (old model)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['α', 'r_gradient', '辺数']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    old_data = [
        ['0', '未定義（0/0）', '0'],
        ['10⁻⁶', '0.859', '9'],
        ['10⁻⁴', '0.859', '9'],
        ['0.5', '0.859', '9'],
        ['1.0', '0.859', '9'],
    ]
    for ri, row_data in enumerate(old_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('命題 8.1')
    run.bold = True
    p.add_run(
        '（r_gradient のスカラー倍不変性）'
        ' 対称なデータ行列 |ρ̂| に非対称行列 C をスカラー α > 0 で混合するとき、'
        'ユーティリティ行列の非対称成分は A(α) = α·(C − C⊤) となる。'
        'Hodge分解の勾配比率 r_gradient はフローの構造にのみ依存し、'
        'フローの大きさには依存しない。'
    )

    add_display_equation(doc,
        'A_old(α) = α·(C − C⊤) + (1−α)·(|ρ| − |ρ|⊤) = α·(C − C⊤)')

    doc.add_paragraph(
        '証明のスケッチ：r_gradient = ‖δ₀φ‖² / ‖ω‖² において、ω → cω（c > 0）とスケールすると、'
        'φ → cφ であり、分子・分母ともに c² でスケールされるため比率は不変。□'
    )

    p = doc.add_paragraph()
    run = p.add_run('実用的含意: ')
    run.bold = True
    p.add_run(
        'αの値そのものは実質的に無意味。α = 0.01 と α = 0.9 で結果は同一。'
        'α はDAG構造の「スイッチ」であり、「ボリューム」ではない。'
        '重要なのは α > 0 か否か、すなわち非対称な方向性信号が存在するか否かのみである。'
    )

    p = doc.add_paragraph()
    run = p.add_run('8.1.2 DPIモデル：滑らかな相転移への移行')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        '§4.1.1 で導入した DPI をデータ駆動成分に用いると、状況は質的に変化する：'
    )

    # α sweep table (DPI model)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(['α', 'r_gradient', '辺数', '非対称ノルム']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    dpi_data = [
        ['0.0', '0.581', '9', '0.815'],
        ['0.1', '0.598', '9', '0.755'],
        ['0.3', '0.688', '9', '0.719'],
        ['0.5', '0.792', '10', '0.803'],
        ['0.6', '0.824', '10', '0.882'],
        ['1.0', '0.859', '9', '1.327'],
    ]
    for ri, row_data in enumerate(dpi_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph(
        '旧モデルからDPIモデルへの移行は、一次相転移（不連続）から二次相転移（連続）への質的変化に相当する。'
        'DPIのデータ駆動的方向性信号は「残留磁化」の役割を果たし、'
        '外部磁場 h = 0（α = 0）でも部分的な秩序状態を維持する。'
    )

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig8_alpha_sweep.png'),
        '図8: α掃引によるDAG転移解析。'
        '(A) r_gradient は α=0 で 0.581 から始まり滑らかに 0.859 へ到達。'
        '(B) 辺数とLiNGAM一致率。(C) 非対称ノルム。(D) 位相図。')

    # --- 8.2 知識の質 p_flip ---
    doc.add_heading('8.2 真の閾値：知識の「量」ではなく「質」', level=2)

    doc.add_paragraph(
        '§8.1 の結果は、DAG構造の出現を支配するのが α（知識の量）ではないことを示した。'
        '真の閾値は知識の質（方向性の正確さ）にある。'
    )

    doc.add_paragraph(
        '正しいドメイン知識の辺方向を p_flip の割合でランダムに反転させた実験（200試行, α = 0.6）：'
    )

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['p_flip', 'r_gradient（平均 ± SD）', '解釈']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    pflip_data = [
        ['0.0', '0.859 ± 0.000', '完全に正しい → 高DAG'],
        ['0.1', '0.576 ± 0.242', '10%誤り → 急落'],
        ['0.2', '0.443 ± 0.226', 'ほぼランダム水準'],
        ['0.3', '0.371 ± 0.214', '最低点（最大循環）'],
        ['0.4', '0.434 ± 0.214', '回復開始'],
        ['0.5', '0.516 ± 0.232', '半分反転'],
        ['0.7', '0.733 ± 0.164', '大部分反転 → DAG回復'],
        ['1.0', '0.859 ± 0.000', '完全反転 → 逆DAGで回復'],
    ]
    for ri, row_data in enumerate(pflip_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('U字型カーブの驚き: ')
    run.bold = True
    p.add_run(
        'p_flip = 0（全て正しい）と p_flip = 1（全て逆転）で同じDAG度。'
        '最低点は p_flip ≈ 0.3。'
    )

    items = [
        'p = 0: 全辺が整合的 → 強いDAG',
        'p = 1: 全辺が逆転しているが、互いに整合的 → 逆方向の強いDAG',
        'p ≈ 0.3: 一部が正方向、一部が逆方向 → 矛盾する方向指示がカール（循環）を最大化',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('「部分的な誤情報は完全な無知より悪い」')
    run.bold = True
    run.italic = True
    p.add_run(
        ' — これはドメイン知識に基づく因果推論一般に対する警告である。'
    )

    doc.add_paragraph(
        '臨界閾値：p*_flip ≈ 0.15（辺方向の85%以上が正しければDAG構造が維持される）。'
    )

    # --- 8.3 LOEO ---
    doc.add_heading('8.3 Leave-One-Edge-Out：根ノードの方向性が骨格', level=2)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['除去した辺', 'Δr_gradient', '重要度']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    loeo_data = [
        ['Age ↔ STDep', '−0.267', '★★★（最重要）'],
        ['Age ↔ MaxHR', '−0.098', '★★'],
        ['Age ↔ Chol', '−0.069', '★★'],
        ['Age ↔ RestBP', '−0.040', '★'],
        ['Chol ↔ STDep', '−0.054', '★'],
        ['RestBP ↔ MaxHR', '+0.015', '除去で改善'],
    ]
    for ri, row_data in enumerate(loeo_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph(
        'Age（根ノード = 外生変数）に関連する辺がDAG構造の骨格を形成している。'
        '「この変数は他に影響されない（外生的だ）」という最小限の知識が最大のレバレッジを持つ。'
        '全ての辺の方向を知る必要はなく、根ノード1つの同定が全体構造を大きく改善する。'
    )

    # --- 8.4 ランダム知識 ---
    doc.add_heading('8.4 ランダム知識との比較', level=2)

    doc.add_paragraph(
        'ドメイン知識をランダム行列に置換した場合（50試行平均）：r_gradient ≈ 0.4（構造なし）。'
        '正しい臨床知識の0.859と大きく異なる。'
        'αを大きくしてもDAG度が改善しない点が注目に値する — '
        'これは命題8.1と整合する：r_gradient は非対称成分の構造に依存し、大きさには依存しない。'
    )

    # --- 8.5 Ising類推 ---
    doc.add_heading('8.5 相転移の物理学的類推：Ising模型との対応', level=2)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(['物理系（Ising模型）', '因果推定系（スペクトル因果性）']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    ising_data = [
        ['温度 T', '知識品質の逆数 p_flip'],
        ['秩序パラメータ（磁化率）', 'r_gradient（DAG度）'],
        ['外部磁場 h', '知識量 α / DPIの方向性信号'],
        ['相転移温度 Tc', 'p*_flip ≈ 0.15'],
        ['強磁性相（低温）', 'DAG的因果構造'],
        ['常磁性相（高温）', '循環的（DCG）構造'],
    ]
    for ri, row_data in enumerate(ising_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph(
        'U字型カーブの物理的解釈：p_flip の増加は「スピンの一部を逆転させる」ことに対応する。'
        'p ≈ 0.3 では正方向群と逆方向群の境界が最も複雑になり、'
        'フラストレーション（矛盾）が最大化する — これがカール成分の最大化に対応する。'
    )

    # --- 8.6 三つの閾値と実用指針 ---
    doc.add_heading('8.6 三つの閾値のまとめと実用指針', level=2)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(['閾値', '値', '意味', '実用的含意']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    thresh_data = [
        ['α*（知識量）', 'DPIにより連続化', 'α=0でもr_gradient=0.581', 'αの厳密な設定は不要'],
        ['p*_flip（知識品質）', '≈ 0.15', '85%以上正しければDAG', '少数でも確実な知識が最良'],
        ['Δr*（骨格辺）', '根ノード（Age等）', '外生変数の方向性がDAG維持に必須', '最小限の知識が最大効果'],
    ]
    for ri, row_data in enumerate(thresh_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph('')  # spacer

    p = doc.add_paragraph()
    run = p.add_run('α設定ガイドライン：')
    run.bold = True

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['状況', '推奨α', '理由']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    guide_data = [
        ['ドメイン知識に高い確信あり', '0.01–0.1', '構造は同一。データの相関重みも活かせる'],
        ['ドメイン知識に不確実性あり', 'LiNGAMのDAGをCに', 'データ駆動でDAG構造を獲得'],
        ['ドメイン知識がない', 'α=0（DPIのみ）', 'DPIが部分的DAGを提供'],
        ['循環を含む分析がしたい', 'α=0でHodge分解', 'DAG仮定を置かない分析'],
    ]
    for ri, row_data in enumerate(guide_data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('最重要メッセージ: ')
    run.bold = True
    p.add_run(
        '重要なのは α の値ではなく、(1) C_domain の方向性が正しいかどうか、'
        '(2) 根ノード（外生変数）の同定、の二点である。'
    )

    # ============================================================
    # 9. 理論的課題と展望
    # ============================================================
    doc.add_heading('9. 理論的課題と展望', level=1)

    # --- 9.1 識別可能性 ---
    doc.add_heading('9.1 識別可能性', level=2)

    p = doc.add_paragraph()
    run = p.add_run('9.1.1 初期モデルの識別可能性問題')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'LiNGAM には明確な識別可能性条件（非ガウス＋線形＋DAG＋共通原因なし → 因果方向が一意に同定）がある。'
        '初期のスペクトル因果性（|ρ̂ᵢⱼ| ベース）には識別可能性の理論がなかった。'
    )

    p = doc.add_paragraph()
    run = p.add_run('予想 9.1')
    run.bold = True
    p.add_run(' 以下の条件下で、SCD は因果方向と一致する：')

    items = [
        '1. ユーティリティ非対称性 U(i,j) − U(j,i) が真の因果方向と同符号',
        '2. ユーティリティ重み w(i,j) が因果効果の強さの単調関数',
        '3. グラフがDAG的構造を持つ（r_gradient ≈ 1）',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Number')

    p = doc.add_paragraph()
    run = p.add_run('条件1の循環性問題: ')
    run.bold = True
    p.add_run(
        '条件1は「正しい因果方向をユーティリティに入力すれば、SCDが正しい因果方向を出力する」と述べており、'
        '循環論法である。|ρ̂ᵢⱼ| が対称であったため、ユーティリティの非対称性は外部から注入する以外に満たす方法がなかった。'
    )

    p = doc.add_paragraph()
    run = p.add_run('9.1.2 DPI導入による循環性の解消')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        '§4.1.1 で導入した DPI は、条件1の循環性を部分的に解消する。'
        'DPIの3成分はそれぞれ、データの統計的性質から直接的に非対称な方向性信号を抽出するため、'
        '「真の因果方向を事前に知っている」必要がない：'
    )

    # DPI identifiability table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['DPI成分', '識別可能性の理論的根拠', '成立条件']):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    dpi_ident = [
        ['回帰係数非対称性', 'Var(Xi) ≠ Var(Xj) の下で |β(j|i)| ≠ |β(i|j)|。LiNGAM と同一の情報源', '分散比が非自明'],
        ['ANM残差独立性', 'Hoyer et al. (2009)の定理：ANMの下で正しい因果方向の残差のみが入力と独立。理論的保証が既に存在', 'ANM仮定'],
        ['条件付きエントロピー縮減', '独立メカニズム仮定の下で、原因→結果の方向でエントロピー縮減が大きい', 'ICM仮定'],
    ]
    for ri, row_data in enumerate(dpi_ident):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run('命題 9.1a')
    run.bold = True
    p.add_run(
        '（DPIの部分的識別可能性）'
        ' データ生成過程が加法ノイズモデル Xⱼ = f(Xᵢ) + ε（ε ⊥ Xᵢ）に従うとき、'
        'DPIのANM成分 Â_ANM(i,j) > 0 が n → ∞ で保証される。'
        'これにより、α = 0 においてもユーティリティ行列の非対称性は外部知識に依存しない。'
    )

    p = doc.add_paragraph()
    run = p.add_run('9.1.3 識別可能性の統合ロードマップ')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'DPI導入後の識別可能性は、3つのフェーズで段階的に到達可能である：'
    )

    p = doc.add_paragraph()
    run = p.add_run('Phase 1: DPI成分レベルの識別可能性（到達済み）')
    run.bold = True
    doc.add_paragraph(
        'DPIの各成分は、それぞれの仮定の下で既に理論的識別可能性を持つ。'
        '特にANM成分は Hoyer et al. (2009) の定理に基づき、加法ノイズモデルの下で方向同定が保証される。'
        '回帰係数非対称性は、線形非ガウスモデルの下で LiNGAM の識別可能性理論がそのまま適用可能である。'
    )

    p = doc.add_paragraph()
    run = p.add_run('Phase 2: スペクトル伝播の整合性（到達可能）')
    run.bold = True
    doc.add_paragraph(
        'DPIが根ノードの方向を正しく同定したとき、Hodge分解がその方向情報をグラフ全体に'
        '伝播する機構の証明が求められる。'
        '(i) ポアソン方程式の一意性：因果ポテンシャル φ は Lφ = δ₀*ω の解として一意に定まる。'
        '(ii) 根ノード方向の伝播：§8.3 の LOEO 解析は、根ノード（Age）の方向情報のみで '
        'r_gradient が 0.314 → 0.581 に改善することを実験的に示した。'
        '特殊ケース（ツリーDAG + 線形SEM）での証明が最初の到達目標となる。'
    )

    p = doc.add_paragraph()
    run = p.add_run('Phase 3: 完全な識別可能性（到達困難 — だが実用上は不要）')
    run.bold = True
    doc.add_paragraph(
        '「全辺の方向を同時に正しく推定する」完全な識別可能性は、スペクトル等価性問題'
        '（異なる因果グラフが同一のスペクトル構造を持ちうる）のため証明が困難である。'
        'しかし、Phase 3 の到達は実用上は必要ない：'
    )
    items = [
        'LiNGAM の識別可能性も有限標本では確率的保証に帰着し、完全な一致は保証されない',
        '§8.2 の p_flip 解析は、方向推定の85%以上が正しければDAG構造が維持されることを示した',
        'ECDアンサンブルにより、LiNGAMの識別可能性保証をコア辺で「借用」し、'
        '残りの構造をスペクトル因果性が補完する戦略が可能',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('9.1.4 実用的な識別可能性の到達点')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'スペクトル因果性の識別可能性は、以下の3つの実用的保証として定式化されるべきである：'
    )
    items = [
        '整合性（Consistency）：n → ∞ で DPI の各成分は真の非対称性に収束する',
        '部分的識別可能性：ANM仮定または非ガウス仮定の下で、少なくとも根ノード群の方向は同定可能であり、'
        'Phase 2 の伝播機構によりグラフ全体の因果順序が定まる',
        '頑健性限界：p*_flip ≈ 0.15 により、DPIの方向精度が85%以上であればDAG構造が実用的に維持される',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Number')

    p = doc.add_paragraph()
    run = p.add_run('ECDアンサンブルの役割: ')
    run.bold = True
    p.add_run(
        'ドメイン知識も DPI 単独の精度も不十分な場合、LiNGAM の識別可能性保証をコア辺で借用し、'
        'スペクトル因果性がフィードバック定量化・Hill基準カバレッジ・グローバル構造推定で補完する。'
        'この「識別可能性の分業」が、完全な識別可能性に代わる現実的な統合パスである。'
    )

    # --- 9.2 ユーティリティ関数の構成とDPIの役割 ---
    doc.add_heading('9.2 ユーティリティ関数の構成とDPIの役割', level=2)

    doc.add_paragraph(
        '初期モデルではデータ駆動成分に |ρ̂ᵢⱼ|（相関係数の絶対値）を用いていた。'
        'しかし |ρ̂ᵢⱼ| は対称であり、α = 0 では方向性信号が消失する。'
    )

    p = doc.add_paragraph()
    run = p.add_run('ベース手法の理論的限界: ')
    run.bold = True
    p.add_run(
        '対称な統計量を用いるかぎり、ドメイン知識なし（α = 0）での因果方向推定は原理的に不可能である。'
        '本稿では §4.1.1 で定義した DPI（方向性予測指標）を導入することにより、この理論的限界を克服した。'
    )

    doc.add_paragraph(
        'DPIの導入により、スペクトル因果性の運用は以下の段階的フレームワークとなる：'
    )
    items = [
        '(a) ドメイン知識なし（α = 0）: U(i,j) = D_DPI(i → j)。DPIの非対称性のみで因果方向を推定。'
        'UCI心疾患データで r_gradient = 0.581, 9本の有向辺を検出, LiNGAM方向一致率67%。',
        '(b) ドメイン知識あり（α > 0）: U(i,j) = α·C_domain + (1−α)·D_DPI。'
        'r_gradient は 0.581 → 0.859 へ滑らかに増加。',
        '(c) LiNGAMとのアンサンブル（ECD）: LiNGAMの推定DAGから高確信辺のみを C_domain に設定可能。'
        'Hill の9基準の網羅性も向上（H6/H7/H9をカバー）。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # --- 9.3 今後の方向性 ---
    doc.add_heading('9.3 今後の方向性', level=2)
    items = [
        '識別可能性 Phase 2 の証明：特殊ケース（ツリーDAG + 線形SEM）でのスペクトル伝播の整合性証明',
        'ECDパイプラインの検証：MIMIC-IV、日本健診コホート（n > 10⁵）での再現性評価',
        '経時データへの拡張：時間的ユーティリティグラフの構築と Eigentrajectories の抽出',
        'p_flip の自動推定：LiNGAMとの因果方向一致率からドメイン知識品質を推定',
        'プルーニング閾値の自動化：フィードバック率のブートストラップ信頼区間に基づく統計的閾値',
        'ドメイン知識のエンコード改善：「情報的影響度」ではなく「介入的因果強度」を',
        'データ適応的 α：非対称行列とデータの相関構造の整合性ブートストラップテスト',
        '相転移の理論的証明：命題8.1の一般化（Hodge理論の帰結として）',
        '多変量スケーリング：50変数以上（> 1225辺）への拡張、Leave-One-Node-Out解析',
        '他のデータセットでの再現性検証：U字型カーブと p*_flip ≈ 0.15 の再現性',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Number')

    # ============================================================
    # 記号一覧
    # ============================================================
    doc.add_heading('記号一覧', level=1)

    symbols = [
        ('G = (V, E, w)', '重み付き（有向）グラフ'),
        ('W, D', '隣接行列, 次数行列'),
        ('L = D − W', '非正規化グラフラプラシアン'),
        ('\u2112 = I − D⁻¹ᐟ²WD⁻¹ᐟ²', '正規化グラフラプラシアン'),
        ('H⁽ᑫ⁾', 'エルミート隣接行列（磁気ラプラシアン用）'),
        ('\u2112⁽ᑫ⁾ = I − D⁻¹ᐟ²H⁽ᑫ⁾D⁻¹ᐟ²', '正規化磁気ラプラシアン'),
        ('q', '電荷パラメータ（方向性感度, [0, 0.5]）'),
        ('σᵢⱼ', 'エッジ方向性符号（{−1, 0, +1}）'),
        ('uₖ, λₖ', '第 k 固有ベクトル, 固有値'),
        ('θₖ(i) = arg(uₖ(i))', 'ノード i の第 k モードにおける位相角'),
        ('U(i,j)', 'ユーティリティ関数'),
        ('SCC(i,j)', 'スペクトル因果結合度（対称）'),
        ('SCD(i,j)', 'スペクトル因果方向（反対称）'),
        ('CCI(i,j)', '複素因果指標（SCC + i·SCD）'),
        ('φ(i)', '因果ポテンシャル（Hodge分解より）'),
        ('r_gradient', '勾配エネルギー比（DAG適合度）'),
        ('D_DPI(i → j)', '方向性予測指標（非対称データ駆動成分）'),
        ('Ā(i,j)', '正規化非対称統計量の平均（DPIの方向性成分）'),
        ('γ', 'DPIの方向性強度パラメータ'),
        ('α', 'ドメイン知識の混合比率'),
    ]

    table = doc.add_table(rows=len(symbols) + 1, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '記号'
    table.cell(0, 1).text = '意味'
    for paragraph in table.cell(0, 0).paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for paragraph in table.cell(0, 1).paragraphs:
        for run in paragraph.runs:
            run.bold = True

    for ri, (sym, meaning) in enumerate(symbols):
        table.cell(ri + 1, 0).text = sym
        table.cell(ri + 1, 1).text = meaning
        for ci in range(2):
            for paragraph in table.cell(ri + 1, ci).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # ============================================================
    # 参考文献
    # ============================================================
    doc.add_heading('参考文献', level=1)

    refs = [
        'Pearl, J. (2009). Causality: Models, Reasoning, and Inference (2nd ed.). Cambridge University Press.',
        'Shimizu, S., Hoyer, P.O., Hyvärinen, A. & Kerminen, A. (2006). A linear non-Gaussian acyclic model for causal discovery. JMLR, 7, 2003–2030.',
        'Shimizu, S. et al. (2011). DirectLiNGAM: A direct method for learning a linear non-Gaussian structural equation model. JMLR, 12, 1225–1248.',
        'Hill, A.B. (1965). The environment and disease: Association or causation? Proc. R. Soc. Med., 58, 295–300.',
        'Granger, C.W.J. (1969). Investigating causal relations by econometric models and cross-spectral methods. Econometrica, 37(3), 424–438.',
        'Shuman, D.I. et al. (2013). The emerging field of signal processing on graphs. IEEE Signal Processing Magazine, 30(3), 83–98.',
        'Zhang, X. et al. (2022). MagNet: A neural network for directed graphs. NeurIPS 2021.',
        'de Resende, B.M.F. & da Costa, L.F. (2020). Characterization of large directed networks through the spectra of the magnetic Laplacian. Chaos, 30(7), 073141.',
        'Seifert, B., Wendler, C. & Püschel, M. (2023). Causal Fourier analysis on directed acyclic graphs and posets. IEEE Trans. Signal Processing, 71, 3516–3530.',
        'Jiang, X. et al. (2011). Statistical ranking and combinatorial Hodge theory. Mathematical Programming, 127, 203–244.',
        'Maehara, K. & Ohkawa, Y. (2019). Modeling latent flows on single-cell data using the Hodge decomposition. bioRxiv.',
        'Kotoku, J. et al. (2020). Causal relations of health indices inferred statistically using DirectLiNGAM. PLOS ONE, 15(12), e0243229.',
        'Okuda, S. et al. (2025). Operationalizing longitudinal causal discovery under real-world workflow constraints. arXiv:2602.23800.',
        'Detrano, R. et al. (1989). International application of a new probability algorithm for the diagnosis of coronary artery disease. Am. J. Cardiol., 64, 304–310.',
        'Rubin, D.B. (1974). Estimating causal effects of treatments in randomized and nonrandomized studies. J. Educ. Psychol., 66(5), 688–701.',
        'Hoyer, P.O. et al. (2009). Nonlinear causal discovery with additive noise models. NeurIPS 2008, 21, 689–696.',
        'Janzing, D. & Schölkopf, B. (2010). Causal inference using the algorithmic Markov condition. IEEE Trans. IT, 56(10), 5168–5194.',
    ]

    for i, ref in enumerate(refs):
        p = doc.add_paragraph()
        run_num = p.add_run(f'[{i+1}] ')
        run_num.bold = True
        run_num.font.size = Pt(10)
        run_text = p.add_run(ref)
        run_text.font.size = Pt(10)

    # Save
    output_path = os.path.join(BASE_DIR, '06_spectral_causality_academic.docx')
    doc.save(output_path)
    print(f'Saved: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_docx()
