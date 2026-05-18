"""
Generate docx files for:
  05_spectral_causality_explainer.docx (general audience)
  06_spectral_causality_academic.docx (university students)
  07_lingam_vs_spectral_comparison.docx (comparison report)

All figures are embedded inline after first mention.
LaTeX math ($...$, $$...$$) is converted to human-readable formatted text
using Unicode symbols and Word-native superscript/subscript.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

FIGURES_DIR = Path('/home/ubuntu/repos/wip/spectral-causality-brainstorm/figures')
OUTPUT_DIR = Path('/home/ubuntu/repos/wip/spectral-causality-brainstorm')

# Map markdown figure references to actual files
FIGURE_MAP = {
    'figures/fig1_three_approaches.png': FIGURES_DIR / 'fig1_three_approaches.png',
    'figures/fig2_magnetic_laplacian_q.png': FIGURES_DIR / 'fig2_magnetic_laplacian_q.png',
    'figures/fig3_hodge_decomposition.png': FIGURES_DIR / 'fig3_hodge_decomposition.png',
    'figures/fig4_direction_comparison.png': FIGURES_DIR / 'fig4_direction_comparison.png',
    'figures/fig5_hill_radar.png': FIGURES_DIR / 'fig5_hill_radar.png',
    'figures/fig6_causal_dag.png': FIGURES_DIR / 'fig6_causal_dag.png',
    'figures/fig7_lingam_vs_spectral.png': FIGURES_DIR / 'fig7_lingam_vs_spectral.png',
    'figures/fig8_alpha_sweep.png': FIGURES_DIR / 'fig8_alpha_sweep.png',
    'figures/fig9_ecd_pruning_analysis.png': FIGURES_DIR / 'fig9_ecd_pruning_analysis.png',
}

# ── LaTeX → Unicode conversion tables ────────────────────────────────────────

GREEK_MAP = {
    '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
    '\\epsilon': 'ε', '\\varepsilon': 'ε', '\\zeta': 'ζ', '\\eta': 'η',
    '\\theta': 'θ', '\\iota': 'ι', '\\kappa': 'κ', '\\lambda': 'λ',
    '\\mu': 'μ', '\\nu': 'ν', '\\xi': 'ξ', '\\pi': 'π',
    '\\rho': 'ρ', '\\sigma': 'σ', '\\tau': 'τ', '\\upsilon': 'υ',
    '\\phi': 'φ', '\\varphi': 'φ', '\\chi': 'χ', '\\psi': 'ψ',
    '\\omega': 'ω',
    '\\Gamma': 'Γ', '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ',
    '\\Xi': 'Ξ', '\\Pi': 'Π', '\\Sigma': 'Σ', '\\Phi': 'Φ',
    '\\Psi': 'Ψ', '\\Omega': 'Ω',
}

OPERATOR_MAP = {
    '\\approx': '≈', '\\neq': '≠', '\\leq': '≤', '\\geq': '≥',
    '\\times': '×', '\\cdot': '·', '\\pm': '±', '\\infty': '∞',
    '\\to': '→', '\\rightarrow': '→', '\\leftarrow': '←',
    '\\Rightarrow': '⇒', '\\Leftarrow': '⇐',
    '\\nabla': '∇', '\\partial': '∂',
    '\\sum': 'Σ', '\\prod': '∏', '\\int': '∫',
    '\\in': '∈', '\\notin': '∉', '\\subset': '⊂',
    '\\cup': '∪', '\\cap': '∩', '\\emptyset': '∅',
    '\\forall': '∀', '\\exists': '∃',
    '\\equiv': '≡', '\\sim': '∼', '\\propto': '∝',
    '\\ldots': '…', '\\cdots': '⋯', '\\dots': '…',
    '\\langle': '⟨', '\\rangle': '⟩',
    '\\ll': '≪', '\\gg': '≫',
    '\\top': 'ᵀ',
    '\\quad': '  ', '\\qquad': '    ',
}


def preprocess_latex(s):
    """Replace LaTeX commands with Unicode equivalents."""
    # \left, \right, \bigl, \bigr → remove
    s = re.sub(r'\\(?:left|right|bigl|bigr|Bigl|Bigr)(?![a-zA-Z])', '', s)

    # \text{...}, \mathrm{...}, \mathbf{...}, \textbf{...}, \operatorname{...}
    s = re.sub(r'\\text\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\textbf\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\operatorname\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\boldsymbol\{([^}]*)\}', r'\1', s)

    # \mathcal{X} → calligraphic Unicode
    _cal = {'L': 'ℒ', 'H': 'ℋ', 'O': '𝒪', 'N': '𝒩', 'F': 'ℱ'}
    s = re.sub(r'\\mathcal\{(\w)\}', lambda m: _cal.get(m.group(1), m.group(1)), s)

    # \mathbb{R} etc.
    _bb = {'R': 'ℝ', 'C': 'ℂ', 'Z': 'ℤ', 'N': 'ℕ', 'Q': 'ℚ'}
    s = re.sub(r'\\mathbb\{(\w)\}', lambda m: _bb.get(m.group(1), m.group(1)), s)

    # \hat{x} → x̂, \bar{x} → x̄, \tilde{x} → x̃
    s = re.sub(r'\\hat\{([^}]*)\}', lambda m: m.group(1) + '\u0302', s)
    s = re.sub(r'\\bar\{([^}]*)\}', lambda m: m.group(1) + '\u0304', s)
    s = re.sub(r'\\tilde\{([^}]*)\}', lambda m: m.group(1) + '\u0303', s)
    s = re.sub(r'\\overline\{([^}]*)\}', lambda m: m.group(1) + '\u0304', s)

    # \frac{a}{b} → a/b
    s = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'\1/\2', s)

    # \sqrt{x} → √(x)
    s = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', s)

    # \underbrace{content}_{label} → content
    s = re.sub(r'\\underbrace\{([^}]*)\}_\{[^}]*\}', r'\1', s)
    s = re.sub(r'\\underbrace\{([^}]*)\}', r'\1', s)

    # \begin{cases}...\end{cases} → simplify
    s = re.sub(r'\\begin\{cases\}', '{ ', s)
    s = re.sub(r'\\end\{cases\}', ' }', s)
    s = s.replace('\\\\', ' ; ')  # line breaks in cases

    # \| → ‖ (norm)
    s = s.replace('\\|', '‖')

    # Greek letters (sort by length desc to avoid partial matches)
    for cmd, char in sorted(GREEK_MAP.items(), key=lambda x: -len(x[0])):
        s = s.replace(cmd, char)

    # Operators (sort by length desc)
    for cmd, char in sorted(OPERATOR_MAP.items(), key=lambda x: -len(x[0])):
        s = s.replace(cmd, char)

    # Common function names: \exp → exp, \sin → sin, etc.
    for func in ['exp', 'sin', 'cos', 'tan', 'log', 'ln', 'max', 'min',
                 'arg', 'det', 'tr', 'diag', 'rank', 'dim', 'ker', 'lim']:
        s = s.replace('\\' + func, func)

    # \, \; \: \! spacing → thin space or nothing
    s = re.sub(r'\\[,;:!]', '', s)

    # Remove stray backslashes before known safe chars
    s = re.sub(r'\\([{}])', r'\1', s)

    return s


def latex_to_segments(latex_str):
    """
    Parse LaTeX into segments for Word rendering.
    Returns list of (text, 'normal'|'sub'|'super').
    """
    s = preprocess_latex(latex_str)
    segments = []
    i = 0
    while i < len(s):
        if s[i] in '_^' and i + 1 < len(s):
            fmt = 'sub' if s[i] == '_' else 'super'
            i += 1
            if i < len(s) and s[i] == '{':
                try:
                    j = s.index('}', i + 1)
                    segments.append((s[i + 1:j], fmt))
                    i = j + 1
                except ValueError:
                    segments.append((s[i + 1:], fmt))
                    break
            elif i < len(s):
                segments.append((s[i], fmt))
                i += 1
        elif s[i] in '{}':
            i += 1
        else:
            j = i
            while j < len(s) and s[j] not in '_^{}':
                j += 1
            if j > i:
                segments.append((s[i:j], 'normal'))
            i = j
    return segments


def add_math_runs(paragraph, latex_str, base_size=None):
    """Add LaTeX math as formatted Word runs (Cambria Math + sub/super).
    Returns list of created runs so callers can apply bold/italic."""
    segments = latex_to_segments(latex_str)
    runs = []
    for text, fmt in segments:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.font.name = 'Cambria Math'
        if base_size:
            run.font.size = base_size
        if fmt == 'sub':
            run.font.subscript = True
        elif fmt == 'super':
            run.font.superscript = True
        runs.append(run)
    return runs


# ── Markdown → docx helpers ──────────────────────────────────────────────────

def set_cell_text(cell, text, bold=False, size=9):
    """Set table cell text, handling $math$ and **bold** if present."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if '$' in text or '**' in text:
        add_formatted_text(p, text, base_size=Pt(size))
        if bold:
            for run in p.runs:
                run.bold = True
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold


def add_table_from_md(doc, header_row, data_rows):
    ncols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header_row):
        set_cell_text(table.rows[0].cells[i], h.strip(), bold=True, size=9)
    for r, row_data in enumerate(data_rows):
        for c, val in enumerate(row_data):
            if c < ncols:
                set_cell_text(table.rows[r + 1].cells[c], val.strip(), size=9)
    return table


def parse_md_table(lines):
    """Parse markdown table lines into header and data rows."""
    if len(lines) < 2:
        return None, None
    header = [c.strip() for c in lines[0].strip('|').split('|')]
    # Skip separator line
    data = []
    for line in lines[2:]:
        if line.strip().startswith('|'):
            row = [c.strip() for c in line.strip('|').split('|')]
            data.append(row)
    return header, data


def add_formatted_text(paragraph, text, base_size=None):
    """Add text with markdown bold/italic/code/math formatting.

    Three-level parser to handle nested patterns like **$math$**:
      Level 1: bold (**...**)
      Level 2: math ($...$)
      Level 3: italic (*...*), code (`...`), superscript refs ({...})
    """
    # Level 1: split on bold (non-greedy to handle **a** x **b**)
    BOLD_RE = re.compile(r'(\*\*.+?\*\*)')
    parts = BOLD_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            _render_with_math(paragraph, part[2:-2], base_size, bold=True)
        else:
            _render_with_math(paragraph, part, base_size, bold=False)


def _render_with_math(paragraph, text, base_size=None, bold=False):
    """Level 2: split on $math$, then delegate non-math to _render_styled."""
    MATH_RE = re.compile(r'(\$[^$]+\$)')
    parts = MATH_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('$') and part.endswith('$'):
            runs = add_math_runs(paragraph, part[1:-1], base_size)
            if bold:
                for r in runs:
                    r.bold = True
        else:
            _render_styled(paragraph, part, base_size, bold=bold)


def _render_styled(paragraph, text, base_size=None, bold=False):
    """Level 3: handle *italic*, `code`, {ref}, and plain text."""
    TOKEN_RE = re.compile(r'(\*[^*]+\*|`[^`]+`|\{[^}]+\})')
    parts = TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.bold = bold
            if base_size:
                run.font.size = base_size
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = base_size or Pt(9)
            run.bold = bold
        elif part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            if base_size:
                run.font.size = base_size
            run.bold = bold
        else:
            run = paragraph.add_run(part)
            if base_size:
                run.font.size = base_size
            run.bold = bold


# ── Main converter ───────────────────────────────────────────────────────────

def md_to_docx(md_path, docx_path, title):
    """Convert markdown to docx with inline figures and rendered math."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Title
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False
    in_reference_section = False

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ──
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1)
                code_lines = []
                in_code_block = False
            else:
                if in_table and table_lines:
                    header, data = parse_md_table(table_lines)
                    if header and data:
                        add_table_from_md(doc, header, data)
                        doc.add_paragraph()
                    table_lines = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table detection ──
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                header, data = parse_md_table(table_lines)
                if header and data:
                    add_table_from_md(doc, header, data)
                    doc.add_paragraph()
                table_lines = []
                in_table = False

        # ── Skip TOC links ──
        if line.strip().startswith('[') and '](#' in line:
            i += 1
            continue

        # ── Display math: $$...$$ ──
        stripped = line.strip()
        if stripped.startswith('$$'):
            if stripped.endswith('$$') and len(stripped) > 4:
                # Single-line display math
                latex = stripped[2:-2].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_math_runs(p, latex)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            else:
                # Multi-line display math
                math_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip() == '$$':
                    math_lines.append(lines[i].strip())
                    i += 1
                latex = ' '.join(math_lines)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_math_runs(p, latex)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Headings ──
        if line.startswith('# ') and not line.startswith('## '):
            i += 1
            continue
        elif line.startswith('## '):
            heading_text = line.lstrip('#').strip()
            p = doc.add_heading('', level=1)
            add_formatted_text(p, heading_text)
            # Track if we're in the reference section
            if '参考文献' in heading_text:
                in_reference_section = True
            else:
                in_reference_section = False
            i += 1
            continue
        elif line.startswith('### '):
            heading_text = line.lstrip('#').strip()
            p = doc.add_heading('', level=2)
            add_formatted_text(p, heading_text)
            i += 1
            continue
        elif line.startswith('#### '):
            heading_text = line.lstrip('#').strip()
            p = doc.add_heading('', level=3)
            add_formatted_text(p, heading_text)
            i += 1
            continue

        # ── Figures: ![caption](path) ──
        fig_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if fig_match:
            fig_path = fig_match.group(2)
            if fig_path in FIGURE_MAP and FIGURE_MAP[fig_path].exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(FIGURE_MAP[fig_path]), width=Inches(5.5))
            i += 1
            continue

        # ── Figure caption (italic line *...*) ──
        if line.strip().startswith('*') and line.strip().endswith('*'):
            caption_text = line.strip().strip('*')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(p, caption_text, base_size=Pt(9))
            # Make non-math runs italic
            for run in p.runs:
                if run.font.name != 'Cambria Math':
                    run.italic = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # ── Horizontal rule ──
        if line.strip() == '---':
            i += 1
            continue

        # ── Blockquote ──
        if line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            add_formatted_text(p, text)
            i += 1
            continue

        # ── Bullet points ──
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip().lstrip('-*').strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # ── Numbered list ──
        num_match = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if num_match:
            num_str = num_match.group(1)
            text = num_match.group(2)
            if in_reference_section:
                # Use explicit numbering to prevent Word auto-numbering issues
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Cm(-1.0)
                run = p.add_run(f'{num_str}. ')
                run.font.size = Pt(10)
                add_formatted_text(p, text)
            else:
                p = doc.add_paragraph(style='List Number')
                add_formatted_text(p, text)
            i += 1
            continue

        # ── Regular paragraph ──
        if line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())

        i += 1

    # Flush remaining table
    if in_table and table_lines:
        header, data = parse_md_table(table_lines)
        if header and data:
            add_table_from_md(doc, header, data)

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    # Generate 05 (general audience explainer)
    md_to_docx(
        OUTPUT_DIR / '05_spectral_causality_explainer.md',
        OUTPUT_DIR / '05_spectral_causality_explainer.docx',
        'スペクトル因果性 — 「音の科学」で因果関係を見つける'
    )

    # Generate 06 (academic)
    md_to_docx(
        OUTPUT_DIR / '06_spectral_causality_academic.md',
        OUTPUT_DIR / '06_spectral_causality_academic.docx',
        'スペクトル因果性の数理的基礎\n— 有向グラフのスペクトル理論に基づく因果推論の新しいアプローチ —'
    )

    # Generate 07 (comparison report)
    md_to_docx(
        OUTPUT_DIR / '07_lingam_vs_spectral_comparison.md',
        OUTPUT_DIR / '07_lingam_vs_spectral_comparison.docx',
        'LiNGAM vs スペクトル因果性：UCI心疾患データによる構造比較'
    )
