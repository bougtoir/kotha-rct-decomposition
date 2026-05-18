"""
Generate cover letter and supplementary submission documents for JCI.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def generate_cover_letter():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run('[Date]')

    doc.add_paragraph('')

    # Addressee
    p = doc.add_paragraph()
    p.add_run('Editors-in-Chief').bold = True
    doc.add_paragraph('Professor Elias Bareinboim, Professor Iv\u00e1n D\u00edaz, and Professor Jin Tian')
    doc.add_paragraph('Journal of Causal Inference')
    doc.add_paragraph('De Gruyter')

    doc.add_paragraph('')

    # Subject
    p = doc.add_paragraph()
    run = p.add_run('Re: Submission of manuscript entitled ')
    run2 = p.add_run('"Mathematical Foundations of Spectral Causality: '
        'A Novel Approach to Causal Inference Based on Spectral Theory of Directed Graphs"')
    run2.italic = True

    doc.add_paragraph('')

    # Salutation
    doc.add_paragraph('Dear Editors,')

    doc.add_paragraph('')

    # Body paragraph 1: Introduction and suitability
    doc.add_paragraph(
        'We are pleased to submit the above-titled manuscript for consideration as a '
        'Research Article in the Journal of Causal Inference. This manuscript presents '
        'spectral causality, a novel framework for causal inference based on the spectral '
        'theory of directed graphs, which we believe aligns closely with the journal\u2019s '
        'scope and mission to advance causal research across disciplinary boundaries.'
    )

    # Body paragraph 2: Why JCI is the right venue
    doc.add_paragraph(
        'We consider JCI the ideal venue for this work for three reasons. First, our framework '
        'introduces a fundamentally new mathematical approach to causal inference\u2014using '
        'the magnetic Laplacian and Hodge decomposition from spectral graph theory\u2014that '
        'complements the existing do-calculus and potential outcomes paradigms rather than '
        'competing with them. JCI\u2019s stated mission to develop \u201ca shared language\u201d '
        'across causal inference disciplines makes it the natural home for such cross-paradigm '
        'work. Second, our Directional Predictability Index (DPI) ensemble directly addresses '
        'core JCI topics including identifiability, statistical estimation, and sensitivity '
        'analysis. Third, the framework\u2019s explicit connection to Hill\u2019s nine '
        'epidemiological criteria bridges the gap between computational causal methods and '
        'applied biostatistics\u2014a gap that existing discipline-specific journals tend '
        'to leave unaddressed.'
    )

    # Body paragraph 3: Novelty and inspiration
    doc.add_paragraph(
        'We believe this manuscript will inspire discussion in the causal inference community '
        'for several reasons:')

    items = [
        'It demonstrates that spectral graph theory can extract causal directional '
        'information from observational data without requiring the acyclicity assumption '
        'that underpins LiNGAM and related methods, thereby accommodating feedback loops '
        'that are ubiquitous in clinical and biological systems.',

        'The DPI enables purely data-driven causal direction estimation (\u03b1 = 0) by '
        'combining regression coefficient asymmetry, additive noise model residual '
        'independence (HSIC), and conditional entropy reduction\u2014achieving 67% agreement '
        'with LiNGAM on the UCI Heart Disease dataset without any domain knowledge.',

        'The Hodge decomposition provides the first principled quantification of the '
        'gradient (DAG-like) and curl (feedback) components of causal flow, offering a '
        'new tool for assessing when the DAG assumption is appropriate.',

        'The Ensemble Causal Discovery (ECD) pipeline, combining LiNGAM\u2019s '
        'identifiability guarantees with spectral causality\u2019s feedback quantification, '
        'achieves broader coverage of Hill\u2019s nine criteria than any single method alone.',

        'The DAG transition point analysis reveals that knowledge quality (p*_flip \u2248 0.15) '
        'rather than quantity (\u03b1) is the critical threshold\u2014a finding with broad '
        'implications for domain knowledge integration in causal discovery.'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Body paragraph 4: Driving academic discussion
    doc.add_paragraph(
        'This work drives academic discussion forward by challenging the implicit assumption '
        'in much of the causal inference literature that methods must choose between the '
        'structural (Pearl) and statistical (Rubin/LiNGAM) traditions. Spectral causality '
        'occupies a distinct position\u2014\u201cLevel 1.5\u201d on the ladder of '
        'causation\u2014that bridges observational association (Level 1) and interventional '
        'causation (Level 2), opening new theoretical questions about the relationship '
        'between spectral structure and causal identifiability.'
    )

    # Body paragraph 5: Declarations
    doc.add_paragraph(
        'We confirm that this manuscript has not been previously published in any language '
        'and is not under simultaneous consideration by another journal. All authors have '
        'approved the manuscript and agree with its submission to Journal of Causal Inference. '
        'The analysis code and generated figures are available in a supplementary repository, '
        'and the UCI Heart Disease dataset used in the empirical illustration is publicly '
        'available. The authors declare no conflicts of interest.'
    )

    doc.add_paragraph('')

    # Suggested reviewers (optional, but helpful)
    p = doc.add_paragraph()
    run = p.add_run('Suggested reviewers:')
    run.bold = True

    reviewers = [
        'Prof. Markus P\u00fcschel (ETH Z\u00fcrich, Switzerland) \u2014 Pioneer of '
        'causal Fourier analysis on DAGs; expert on graph signal processing.',
        'Prof. Michael Schaub (RWTH Aachen University, Germany) \u2014 Expert on Hodge '
        'decomposition and spectral methods for network analysis.',
        'Prof. Kun Zhang (Carnegie Mellon University, USA) \u2014 Expert on causal '
        'discovery, non-linear causal models, and identifiability theory.',
    ]
    for r in reviewers:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_paragraph('')

    # Closing
    doc.add_paragraph(
        'We look forward to your consideration and welcome any feedback from the reviewers.')

    doc.add_paragraph('')
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.add_run('[Author Name]').bold = True
    doc.add_paragraph('[Affiliation]')
    doc.add_paragraph('[Email]')
    doc.add_paragraph('[ORCID: xxxx-xxxx-xxxx-xxxx]')

    out_path = os.path.join(BASE_DIR, 'cover_letter_jci.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


def generate_submission_checklist():
    """Generate a submission checklist document for JCI via ScholarOne."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('JCI Submission Checklist')
    run.bold = True
    run.font.size = Pt(16)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run('ScholarOne Manuscripts: https://mc.manuscriptcentral.com/dgjci')

    doc.add_paragraph('')

    doc.add_heading('Required Documents', level=2)

    items = [
        ('\u2610', 'Manuscript (docx or PDF)', 
         '06_spectral_causality_academic_en.docx \u2014 Full manuscript with inline figures, '
         'Vancouver references [1]\u2013[30], data availability statement, COI disclosure. '
         'For initial submission, a single PDF with text + figures is recommended.'),

        ('\u2610', 'Cover letter',
         'cover_letter_jci.docx \u2014 Explains suitability for JCI, novelty, '
         'not previously published, no simultaneous submission, suggested reviewers.'),

        ('\u2610', 'License to Publish (after acceptance)',
         'Download from: https://www.degruyterbrill.com/publication/journal_key/JCI/'
         'downloadAsset/JCI_License%20to%20Publish.pdf \u2014 '
         'Print, sign, scan, and submit. Required after acceptance, not at initial submission.'),

        ('\u2610', 'Figure files (if separate)',
         'spectral_causality_figures_en.pptx contains all 8 figures. '
         'For initial review: figures are already embedded in the manuscript PDF/docx. '
         'For revision: provide individual figure files (EPS, JPG, TIFF, or PDF). '
         'Note: PowerPoint files are NOT accepted as final figure format by JCI.'),
    ]

    for check, title, detail in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{check} {title}')
        run.bold = True
        p2 = doc.add_paragraph(detail)
        p2.paragraph_format.left_indent = Inches(0.5)
        for r in p2.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_heading('ScholarOne Submission Steps', level=2)

    steps = [
        'Go to https://mc.manuscriptcentral.com/dgjci',
        'Create an Author account (if first time)',
        'Select "Author Center" \u2192 "Click here to submit a new manuscript"',
        'Article type: "Research Article"',
        'Enter title, abstract, keywords, and author information',
        'Upload manuscript file (PDF recommended for initial submission)',
        'Upload cover letter as a supplementary file',
        'Suggest 2\u20133 reviewers (names, affiliations, emails)',
        'Review submission summary and submit',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Author Information to Fill In', level=2)

    placeholders = [
        ('Author Name', '[Full name]'),
        ('Affiliation', '[Department, University/Institution, City, Country]'),
        ('Email', '[Corresponding author email]'),
        ('ORCID', '[xxxx-xxxx-xxxx-xxxx]'),
        ('Phone', '[Optional]'),
    ]

    table = doc.add_table(rows=len(placeholders) + 1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Field'
    table.rows[0].cells[1].text = 'Value'
    for r in table.rows[0].cells[0].paragraphs[0].runs:
        r.bold = True
    for r in table.rows[0].cells[1].paragraphs[0].runs:
        r.bold = True
    for i, (field, value) in enumerate(placeholders, 1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = value

    doc.add_heading('Accepted Figure Formats (for revision)', level=2)
    doc.add_paragraph(
        'JCI accepts: EPS, BMP, JPG, TIFF, GIF, PDF. PowerPoint is NOT accepted. '
        'For the initial submission, figures embedded in the manuscript PDF are sufficient. '
        'If revision is requested, export individual figures from the pptx as high-resolution '
        'PNG/TIFF (300 DPI minimum) or PDF.')

    doc.add_heading('Post-Acceptance Requirements', level=2)
    doc.add_paragraph(
        '\u2022 License to Publish form (signed, scanned PDF)\n'
        '\u2022 Editable text file (LaTeX strongly preferred for final version; docx acceptable)\n'
        '\u2022 Individual figure files in accepted formats (EPS, TIFF, PDF preferred)\n'
        '\u2022 LaTeX template available from JCI website')

    doc.add_heading('Notes', level=2)
    doc.add_paragraph(
        '\u2022 APC: \u20ac1,000 (Open Access, CC-BY)\n'
        '\u2022 Review model: Single-blind\n'
        '\u2022 Minimum 2 reviewers\n'
        '\u2022 No length restriction\n'
        '\u2022 Data sharing: Data availability statement required (already included in manuscript)')

    out_path = os.path.join(BASE_DIR, 'jci_submission_checklist.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


def generate_title_page():
    """Generate a separate title page document."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run('Mathematical Foundations of Spectral Causality:\n'
        'A Novel Approach to Causal Inference Based on\n'
        'Spectral Theory of Directed Graphs')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('')

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Author Name]')
    run.font.size = Pt(14)

    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Department, University/Institution]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[City, Country]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph('')

    # Corresponding author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Corresponding Author:')
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('[Author Name]').font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('[Email address]').font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('[ORCID: xxxx-xxxx-xxxx-xxxx]').font.size = Pt(11)

    doc.add_paragraph('')

    # Word count
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Word count: ~11,300 words | Figures: 8 | Tables: 15 | References: 30').font.size = Pt(10)

    doc.add_paragraph('')

    # Running title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Running title: ')
    run.bold = True
    run.font.size = Pt(10)
    p.add_run('Spectral Causality via Magnetic Laplacian').font.size = Pt(10)

    doc.add_paragraph('')

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(
        'spectral causality, magnetic Laplacian, Hodge decomposition, causal inference, '
        'directed acyclic graph, feedback loop, Hill\u2019s criteria, LiNGAM, ensemble '
        'causal discovery'
    ).font.size = Pt(10)

    out_path = os.path.join(BASE_DIR, 'title_page_jci.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


if __name__ == '__main__':
    generate_cover_letter()
    generate_submission_checklist()
    generate_title_page()
