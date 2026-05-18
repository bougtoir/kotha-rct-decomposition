"""
Generate English docx for 06_spectral_causality_academic_en.md
For JCI (Journal of Causal Inference) submission.
Parses the markdown and produces a well-formatted docx with inline figures.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(BASE_DIR, 'figures')
MD_FILE = os.path.join(BASE_DIR, '06_spectral_causality_academic_en.md')

# Figure mapping: figure number -> filename
FIGURE_MAP = {
    1: 'fig6_causal_dag.png',
    2: 'fig5_hill_radar.png',
    3: 'fig2_magnetic_laplacian_q.png',
    4: 'fig3_hodge_decomposition.png',
    5: 'fig4_direction_comparison.png',
    6: 'fig7_lingam_vs_spectral.png',
    7: 'fig9_ecd_pruning_analysis.png',
    8: 'fig8_alpha_sweep.png',
}


def add_figure(doc, image_path, caption, width=Inches(5.5)):
    """Add a figure with caption to the document, centered."""
    if not os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[Figure: {caption} — image file not found: {image_path}]')
        run.italic = True
        run.font.size = Pt(9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run()
    run.add_picture(image_path, width=width)
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(4)
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.italic = True


def convert_inline_math(text):
    """Convert inline LaTeX ($...$) to Unicode approximations for Word."""
    # Greek letters
    greek = {
        r'\alpha': '\u03b1', r'\beta': '\u03b2', r'\gamma': '\u03b3',
        r'\delta': '\u03b4', r'\varepsilon': '\u03b5', r'\epsilon': '\u03b5',
        r'\zeta': '\u03b6', r'\eta': '\u03b7', r'\theta': '\u03b8',
        r'\iota': '\u03b9', r'\kappa': '\u03ba', r'\lambda': '\u03bb',
        r'\mu': '\u03bc', r'\nu': '\u03bd', r'\xi': '\u03be',
        r'\pi': '\u03c0', r'\rho': '\u03c1', r'\sigma': '\u03c3',
        r'\tau': '\u03c4', r'\phi': '\u03c6', r'\chi': '\u03c7',
        r'\psi': '\u03c8', r'\omega': '\u03c9',
        r'\Gamma': '\u0393', r'\Delta': '\u0394', r'\Lambda': '\u039b',
        r'\Sigma': '\u03a3', r'\Phi': '\u03a6', r'\Psi': '\u03a8',
        r'\Omega': '\u03a9',
    }
    # Math symbols
    symbols = {
        r'\in': '\u2208', r'\to': '\u2192', r'\rightarrow': '\u2192',
        r'\leftarrow': '\u2190', r'\leftrightarrow': '\u2194',
        r'\Rightarrow': '\u21d2', r'\Leftrightarrow': '\u21d4',
        r'\leq': '\u2264', r'\geq': '\u2265', r'\neq': '\u2260',
        r'\approx': '\u2248', r'\sim': '\u223c', r'\infty': '\u221e',
        r'\partial': '\u2202', r'\nabla': '\u2207', r'\sum': '\u2211',
        r'\prod': '\u220f', r'\int': '\u222b', r'\times': '\u00d7',
        r'\cdot': '\u00b7', r'\circ': '\u2218', r'\perp': '\u22a5',
        r'\top': '\u22a4', r'\forall': '\u2200', r'\exists': '\u2203',
        r'\emptyset': '\u2205', r'\subset': '\u2282', r'\supset': '\u2283',
        r'\cup': '\u222a', r'\cap': '\u2229', r'\wedge': '\u2227',
        r'\vee': '\u2228', r'\neg': '\u00ac', r'\ll': '\u226a',
        r'\gg': '\u226b', r'\star': '\u22c6', r'\ast': '*',
        r'\pm': '\u00b1', r'\mp': '\u2213',
        r'\quad': '  ', r'\qquad': '    ',
        r'\dots': '\u2026', r'\cdots': '\u22ef', r'\ldots': '\u2026',
        r'\square': '\u25a1',
    }
    # Math operators
    operators = {
        r'\operatorname{tr}': 'tr', r'\operatorname{diag}': 'diag',
        r'\operatorname{sign}': 'sign', r'\operatorname{Re}': 'Re',
        r'\operatorname{Im}': 'Im', r'\arg': 'arg', r'\min': 'min',
        r'\max': 'max', r'\exp': 'exp', r'\log': 'log', r'\cos': 'cos',
        r'\sin': 'sin', r'\text{tr}': 'tr', r'\text{diag}': 'diag',
    }

    def process_math(m):
        s = m.group(1)
        # Apply replacements
        for k, v in operators.items():
            s = s.replace(k, v)
        for k, v in greek.items():
            s = s.replace(k, v)
        for k, v in symbols.items():
            s = s.replace(k, v)
        # \mathbb{R} etc
        s = re.sub(r'\\mathbb\{([A-Z])\}', lambda m2: {'R': '\u211d', 'C': '\u2102', 'Z': '\u2124', 'N': '\u2115', 'Q': '\u211a'}.get(m2.group(1), m2.group(1)), s)
        # \mathcal{L}
        s = re.sub(r'\\mathcal\{L\}', '\u2112', s)
        s = re.sub(r'\\mathcal\{([A-Z])\}', lambda m2: m2.group(1), s)
        # \mathbf{x}
        s = re.sub(r'\\mathbf\{([^}]+)\}', r'\1', s)
        # \text{...}
        s = re.sub(r'\\text\{([^}]+)\}', r'\1', s)
        # \mathrm{...}
        s = re.sub(r'\\mathrm\{([^}]+)\}', r'\1', s)
        # \hat{x}
        s = re.sub(r'\\hat\{([^}]+)\}', lambda m2: m2.group(1) + '\u0302', s)
        s = re.sub(r'\\hat([a-zA-Z])', lambda m2: m2.group(1) + '\u0302', s)
        # \bar{x}
        s = re.sub(r'\\bar\{([^}]+)\}', lambda m2: m2.group(1) + '\u0304', s)
        s = re.sub(r'\\bar([a-zA-Z])', lambda m2: m2.group(1) + '\u0304', s)
        # \tilde{x}
        s = re.sub(r'\\tilde\{([^}]+)\}', lambda m2: m2.group(1) + '\u0303', s)
        # \overline{x}
        s = re.sub(r'\\overline\{([^}]+)\}', lambda m2: m2.group(1) + '\u0304', s)
        # \bigl, \bigr, \left, \right
        s = re.sub(r'\\(?:bigl?|bigr?|left|right)([()[\]|.])', r'\1', s)
        s = re.sub(r'\\(?:bigl?|bigr?|left|right)', '', s)
        # Fractions \frac{a}{b} -> a/b
        s = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', s)
        # Subscript _{...}
        sub_map = {'0': '\u2080', '1': '\u2081', '2': '\u2082', '3': '\u2083',
                    '4': '\u2084', '5': '\u2085', '6': '\u2086', '7': '\u2087',
                    '8': '\u2088', '9': '\u2089', '+': '\u208a', '-': '\u208b',
                    '=': '\u208c', '(': '\u208d', ')': '\u208e',
                    'i': '\u1d62', 'j': '\u2c7c', 'k': '\u2096', 'n': '\u2099',
                    'a': '\u2090', 'e': '\u2091', 'o': '\u2092', 'r': '\u1d63',
                    's': '\u209b', 't': '\u209c', 'x': '\u2093', 'u': '\u1d64',
                    'v': '\u1d65', 'p': '\u209a'}
        sup_map = {'0': '\u2070', '1': '\u00b9', '2': '\u00b2', '3': '\u00b3',
                    '4': '\u2074', '5': '\u2075', '6': '\u2076', '7': '\u2077',
                    '8': '\u2078', '9': '\u2079', '+': '\u207a', '-': '\u207b',
                    '=': '\u207c', '(': '\u207d', ')': '\u207e',
                    'n': '\u207f', 'i': '\u2071', '*': '\u002a', 'T': '\u1d40'}

        def sub_repl(m2):
            content = m2.group(1)
            return ''.join(sub_map.get(c, c) for c in content)

        def sup_repl(m2):
            content = m2.group(1)
            return ''.join(sup_map.get(c, c) for c in content)

        s = re.sub(r'_\{([^}]+)\}', sub_repl, s)
        s = re.sub(r'_([a-z0-9*])', lambda m2: sub_map.get(m2.group(1), '_' + m2.group(1)), s)
        s = re.sub(r'\^\{([^}]+)\}', sup_repl, s)
        s = re.sub(r'\^([a-z0-9*TnN])', lambda m2: sup_map.get(m2.group(1), '^' + m2.group(1)), s)
        # Remove remaining backslashes for simple commands
        s = re.sub(r'\\[a-zA-Z]+', '', s)
        return s

    # Process inline math $...$  (not $$...$$)
    text = re.sub(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', process_math, text)
    return text


def add_rich_paragraph(doc, text, bold_prefix=None, italic=False, style=None, alignment=None):
    """Add a paragraph with inline math converted and optional bold prefix."""
    text = convert_inline_math(text)
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        text_after = text
    else:
        text_after = text
    # Handle **bold** and *italic* in text
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text_after)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            if italic:
                run.italic = True
    return p


def add_display_equation(doc, eq_text, label=None):
    """Add a centered equation."""
    eq_text = convert_inline_math('$' + eq_text + '$')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(eq_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    if label:
        tab_run = p.add_run(f'    ({label})')
        tab_run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(convert_inline_math(h))
        run.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(convert_inline_math(str(val)))
            run.font.size = Pt(9)
    return table


def generate_docx():
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ============================================================
    # Title Page
    # ============================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(72)
    title_run = title_p.add_run('Mathematical Foundations of Spectral Causality')
    title_run.bold = True
    title_run.font.size = Pt(18)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        'A Novel Approach to Causal Inference Based on '
        'Spectral Theory of Directed Graphs'
    )
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(80, 80, 80)

    # Author placeholder
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_before = Pt(24)
    ar = author_p.add_run('[Author Name]')
    ar.font.size = Pt(12)

    aff_p = doc.add_paragraph()
    aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    af = aff_p.add_run('[Affiliation]')
    af.font.size = Pt(10)
    af.font.color.rgb = RGBColor(100, 100, 100)

    corr_p = doc.add_paragraph()
    corr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = corr_p.add_run('Corresponding author: [email]')
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ============================================================
    # Abstract
    # ============================================================
    doc.add_heading('Abstract', level=1)
    add_rich_paragraph(doc,
        'We propose **spectral causality**, a novel framework for causal inference based on the '
        'spectral theory of directed graphs. The key mathematical tools are the **magnetic '
        'Laplacian**\u2014a Hermitian matrix that encodes edge directionality as complex '
        'phase\u2014and the **Hodge decomposition**, which orthogonally separates edge flows '
        'into a gradient (DAG-like) component and a curl (feedback) component. Unlike LiNGAM, '
        'our framework does not require the acyclicity assumption, thereby accommodating feedback '
        'loops that are ubiquitous in biological and clinical systems.')

    add_rich_paragraph(doc,
        'To enable purely data-driven causal direction estimation, we introduce the **Directional '
        'Predictability Index (DPI)**, an ensemble of three asymmetric statistics (regression '
        'coefficient asymmetry, additive noise model residual independence via HSIC, and '
        'conditional entropy reduction). With DPI, the framework estimates causal directions even '
        'in the absence of domain knowledge (\u03b1 = 0): on the UCI Heart Disease dataset '
        '(n = 297, 5 variables), the method detects 9 directed edges with 67% agreement with '
        'LiNGAM, achieving a gradient energy ratio r_gradient = 0.581 that smoothly improves to '
        '0.859 as domain knowledge increases.')

    add_rich_paragraph(doc,
        'We further develop an **Ensemble Causal Discovery (ECD)** pipeline that combines '
        "LiNGAM\u2019s identifiability guarantees with spectral causality\u2019s feedback "
        "quantification, achieving broader coverage of Hill\u2019s nine criteria for "
        'epidemiological causality. A detailed analysis of **DAG transition points** reveals that '
        'the critical threshold is not the amount of domain knowledge (\u03b1) but its quality '
        '(p*_flip \u2248 0.15: directions must be at least 85% correct to maintain DAG structure).')

    kw_p = doc.add_paragraph()
    kr = kw_p.add_run('Keywords: ')
    kr.bold = True
    kw_p.add_run(
        'spectral causality, magnetic Laplacian, Hodge decomposition, causal inference, '
        'directed acyclic graph, feedback loop, Hill\u2019s criteria, LiNGAM, ensemble causal discovery'
    )

    # ============================================================
    # 1. Introduction
    # ============================================================
    doc.add_heading('1. Introduction', level=1)

    doc.add_heading('1.1 Problem Setting', level=2)
    add_rich_paragraph(doc,
        'The central question of causal inference\u2014\u201cIs X a cause of Y?\u201d\u2014has '
        'been addressed through diverse approaches. Notable frameworks include:')

    items_1 = [
        'Structural equation models (SEMs) and do-calculus [1]: Counterfactual definitions based on interventions',
        'The potential outcomes framework [2]: Differences in potential outcomes between treated and control groups',
        'LiNGAM [3]: Identification of causal direction by exploiting non-Gaussianity of data',
        'Granger causality [4]: Causality defined by predictive improvement in time series',
    ]
    for item in items_1:
        doc.add_paragraph(item, style='List Bullet')

    add_rich_paragraph(doc,
        'In this paper, we formulate a method based on a fundamentally different principle: '
        '**reading causal directionality from the spectral structure (eigenvalues and '
        'eigenvectors) of a graph**. We term this approach **spectral causality**.')

    doc.add_heading('1.2 Core Idea', level=2)
    add_rich_paragraph(doc,
        'Suppose n variables {X\u2081, \u2026, X\u2099} are causally related. When these '
        'relationships are represented as a directed graph G = (V, E), the spectrum of the '
        "graph\u2019s Laplacian matrix (its eigenvalues and eigenvectors) can encode information "
        'about causal directionality.')

    add_rich_paragraph(doc,
        '**Remark 1.1** (Graph types and causal models). Graphs used in causal inference are '
        'not necessarily restricted to directed acyclic graphs (DAGs). While LiNGAM imposes a '
        'DAG assumption, real biological systems universally contain feedback loops (e.g., '
        'inflammation \u2192 organ damage \u2192 inflammation). Spectral causality accommodates '
        'directed cyclic graphs (DCGs), as the Hodge decomposition (\u00a75) quantifies cyclic '
        'flows via the curl component. Figure 1 shows an example of a causal DAG estimated by '
        'LiNGAM under the acyclicity assumption.')

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig6_causal_dag.png'),
        'Figure 1: Causal DAG estimated by DirectLiNGAM [5] on the UCI Heart Disease dataset '
        '(Cleveland subset, n = 297) with five clinical variables. Blue edges indicate positive '
        'causal effects; red edges indicate negative effects.')

    add_rich_paragraph(doc,
        'In particular, by using a magnetic Laplacian\u2014a Hermitian matrix\u2014edge '
        'directionality is encoded as complex phase in the eigenvectors, enabling estimation '
        'of causal direction.')

    doc.add_heading('1.3 Structure of This Paper and Main Contributions', level=2)
    add_rich_paragraph(doc,
        '\u00a72 reviews the fundamentals of graph Laplacians, \u00a73 introduces the magnetic '
        'Laplacian, \u00a74 formally defines spectral causality (with DPI introduced in '
        '\u00a74.1.1), and \u00a75 establishes the connection to Hodge decomposition. \u00a76 '
        'discusses complementarity and augmentative extensibility with existing methods, and '
        '\u00a77 surveys related work. \u00a78 presents an empirical illustration. '
        '\u00a7\u00a79\u201312 detail the structural comparison with LiNGAM, ECD ensemble, '
        'DAG transition point analysis, and cyclic pruning. \u00a713 discusses theoretical '
        'challenges and future directions.')

    add_rich_paragraph(doc, 'The main contributions of this work are:')
    contributions = [
        'Introduction of DPI (asymmetric statistics): Enables detection of directed edges '
        'and estimation of causal direction even at \u03b1 = 0 (without domain knowledge). '
        'On the UCI Heart Disease dataset, 9 directed edges are detected with 67% agreement '
        'with LiNGAM.',
        'No DAG assumption required: Hodge decomposition naturally separates DAG (gradient) '
        'and feedback (curl) components. This accommodates clinically correct feedback structures.',
        'Smooth improvement with domain knowledge: r_gradient improves smoothly from 0.581 '
        '(\u03b1 = 0) to 0.859 (\u03b1 = 0.6).',
        'LiNGAM integration: When domain knowledge is absent, high-confidence edges from '
        "LiNGAM\u2019s estimated DAG can be injected as C (a \u201ctwo-stage rocket\u201d strategy).",
        "ECD ensemble: Improves coverage of Hill\u2019s nine criteria (covers H6/H7/H9).",
    ]
    for c in contributions:
        doc.add_paragraph(c, style='List Number')

    # ============================================================
    # 2. Preliminaries
    # ============================================================
    doc.add_heading('2. Preliminaries: Fundamentals of Graph Laplacians', level=1)

    doc.add_heading('2.1 The Laplacian of an Undirected Graph', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Definition 2.1 ')
    run.bold = True
    p.add_run(
        '(Graph Laplacian). For a weighted undirected graph G = (V, E, w) '
        '(|V| = n, w: E \u2192 \u211d\u208b), define the weighted adjacency matrix '
        'W \u2208 \u211d\u207f\u02e3\u207f and the degree matrix D = diag(d\u2081, \u2026, d\u2099) '
        '(d\u1d62 = \u03a3\u2c7c W\u1d62\u2c7c). The Laplacians are:'
    )

    add_display_equation(doc, 'L = D \u2212 W \\quad (unnormalized Laplacian)')
    add_display_equation(doc, '\u2112 = I \u2212 D\u207b\u00b9\u02f2 W D\u207b\u00b9\u02f2 \\quad (normalized Laplacian)')

    p = doc.add_paragraph()
    run = p.add_run('Proposition 2.1 ')
    run.bold = True
    p.add_run('(Basic Properties).')

    props = [
        '(i) L is a symmetric positive semi-definite matrix with eigenvalues 0 = \u03bb\u2081 \u2264 \u03bb\u2082 \u2264 \u2026 \u2264 \u03bb\u2099.',
        '(ii) The eigenvector corresponding to \u03bb\u2081 = 0 is 1 = (1, \u2026, 1)\u1d40 (constant vector).',
        '(iii) \u03bb\u2082 > 0 if and only if G is connected (Fiedler value).',
        '(iv) For any vector f \u2208 \u211d\u207f, f\u1d40Lf = \u03a3\u208d\u1d62,\u2c7c\u208e\u2208E w\u1d62\u2c7c(f\u1d62 \u2212 f\u2c7c)\u00b2 \u2265 0.',
    ]
    for prop in props:
        doc.add_paragraph(prop, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Proof sketch')
    run.italic = True
    p.add_run(
        ': (iv) follows from direct expansion of the quadratic form of L. (i) follows from (iv). '
        '(ii) is obtained by direct computation of L\u00b71 = 0. (iii) is the algebraic '
        'connectivity theorem. \u25a1'
    )

    add_rich_paragraph(doc,
        'Property (iv) is key: a smaller f\u1d40Lf means that f takes similar values on '
        'adjacent nodes\u2014i.e., the low-eigenvalue eigenvectors of the Laplacian represent '
        '**smooth signals on the graph**.')

    doc.add_heading('2.2 Geometric Meaning of the Spectral Decomposition', level=2)
    add_rich_paragraph(doc,
        'In the spectral decomposition \u2112 = U\u039bU\u1d40 (U = [u\u2081, \u2026, u\u2099], '
        '\u039b = diag(\u03bb\u2081, \u2026, \u03bb\u2099)):')

    items_2 = [
        'u\u2096(i) = the "loading" of node i on the k-th eigenmode',
        '\u03bb\u2096 = the "frequency" of the k-th mode (larger values = higher frequency = more local variation)',
        'u\u2082 (the Fiedler vector) yields the optimal bipartition of the graph',
    ]
    for item in items_2:
        doc.add_paragraph(item, style='List Bullet')

    add_rich_paragraph(doc,
        'This framework is the foundation of **Graph Signal Processing (GSP)** [6]\u2014a '
        'generalization of the Fourier transform to graphs.')

    doc.add_heading('2.3 Limitation: The Undirected Laplacian Loses Directionality', level=2)
    add_rich_paragraph(doc,
        'Since L = D \u2212 W is a symmetric matrix, it cannot distinguish edge direction '
        'i \u2192 j from j \u2192 i. In causal inference, the directionality \u201cX causes '
        'Y\u201d is essential, and the undirected Laplacian provides insufficient information.')

    # ============================================================
    # 3. Magnetic Laplacian
    # ============================================================
    doc.add_heading('3. The Magnetic Laplacian: Encoding Directionality via Complex Phase', level=1)

    doc.add_heading('3.1 Physical Background', level=2)
    add_rich_paragraph(doc,
        'The name \u201cmagnetic Laplacian\u201d derives from quantum mechanics. The Hamiltonian '
        'of a charged particle in a magnetic field B is H = (p \u2212 eA)\u00b2/2m (where A is '
        'the vector potential), and a particle traversing a closed loop acquires an Aharonov\u2013Bohm '
        'phase exp(i \u222b A \u00b7 dr). The direction dependence of this phase is exploited to '
        'encode edge directionality on graphs.')

    doc.add_heading('3.2 Definition', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Definition 3.1 ')
    run.bold = True
    p.add_run(
        '(Magnetic Laplacian [7, 8]). For a weighted directed graph G = (V, E, w) and a charge '
        'parameter q \u2208 [0, 0.5], define the Hermitian adjacency matrix '
        'H\u207e\u207d\u1d60\u207e\u207d \u2208 \u2102\u207f\u02e3\u207f as:'
    )

    add_display_equation(doc,
        'H(q)_ij = w_ij \u00b7 exp(i \u00b7 2\u03c0q \u00b7 \u03c3_ij)')

    add_rich_paragraph(doc,
        'where \u03c3\u1d62\u2c7c \u2208 {\u22121, 0, +1} is the edge directionality sign: '
        '+1 if i \u2192 j, \u22121 if j \u2192 i, 0 if no edge. The weights w\u1d62\u2c7c are '
        'symmetrized. The normalized magnetic Laplacian is:')

    add_display_equation(doc,
        '\u2112(q) = I \u2212 D\u207b\u00b9\u02f2 H(q) D\u207b\u00b9\u02f2')

    p = doc.add_paragraph()
    run = p.add_run('Proposition 3.1 ')
    run.bold = True
    p.add_run('(Basic Properties of the Magnetic Laplacian).')

    props_3 = [
        '(i) H(q) is Hermitian: H(q)_ji = H\u0304(q)_ij.',
        '(ii) \u2112(q) is Hermitian positive semi-definite with real, non-negative eigenvalues.',
        '(iii) The eigenvectors are generally complex-valued.',
        '(iv) When q = 0, \u2112(0) degenerates to the standard normalized Laplacian \u2112 (no directional information).',
    ]
    for prop in props_3:
        doc.add_paragraph(prop, style='List Bullet')

    doc.add_heading('3.3 Meaning of the Charge Parameter q', level=2)

    add_table(doc,
        ['q', 'Phase 2\u03c0q', 'Effect'],
        [
            ['0', '0', 'Directionality completely ignored. exp(i\u00b70) = 1, reducing to real matrix'],
            ['0.25', '\u03c0/2', 'Maximum directional sensitivity. e^(i\u03c0/2) = i, e^(\u2212i\u03c0/2) = \u2212i'],
            ['0.5', '\u03c0', 'Direction reversal. e^(i\u03c0) = \u22121'],
        ])

    add_rich_paragraph(doc,
        '**Remark 3.1** At q = 0.25, H(q)_ij = i\u00b7w_ij (for i \u2192 j edges), providing '
        'the sharpest separation of directionality via the imaginary unit i.')

    doc.add_heading('3.4 Complex Phase of Eigenvectors and Directionality', level=2)
    add_rich_paragraph(doc,
        'Each component of the eigenvector u\u2096 \u2208 \u2102\u207f of \u2112(q) can be '
        'written in polar form: u\u2096(j) = |u\u2096(j)| \u00b7 exp(i\u00b7\u03b8\u2096(j)), '
        'where |u\u2096(j)| is the amplitude and \u03b8\u2096(j) = arg(u\u2096(j)) is the phase angle.')

    add_rich_paragraph(doc,
        '**Core claim**: When q > 0, the ordering of phase angles \u03b8\u2096(j) reflects '
        'the direction of causal flow. Upstream (cause-side) and downstream (effect-side) nodes '
        'acquire distinct phase angles in the eigenvectors.')

    # ============================================================
    # 4. Formulation of Spectral Causality
    # ============================================================
    doc.add_heading('4. Formulation of Spectral Causality', level=1)

    doc.add_heading('4.1 The Utility Directed Graph', level=2)
    add_rich_paragraph(doc,
        'We encode causal directionality between variables using the asymmetry of utility '
        '(clinical usefulness).')

    p = doc.add_paragraph()
    run = p.add_run('Definition 4.1 ')
    run.bold = True
    p.add_run(
        '(Utility Directed Graph). For n variables {X\u2081, \u2026, X\u2099}, define the '
        'utility function U: {1, \u2026, n}\u00b2 \u2192 \u211d\u2265\u2080 as '
        'U(i, j) = "How useful is information about X\u1d62 for questions about X\u2c7c". '
        'The utility directed graph G_U = (V, E, w, \u03c3) is defined by: '
        'w(i,j) = (U(i,j) + U(j,i))/2 (symmetrized weight) and '
        '\u03c3(i,j) = sign(U(i,j) \u2212 U(j,i)) (directionality sign).'
    )

    doc.add_heading('4.1.1 Data-Driven Component: The Directional Predictability Index (DPI)', level=3)

    add_rich_paragraph(doc,
        'Previously, |\u03c1\u0302\u1d62\u2c7c| (absolute correlation) was used as the '
        'data-driven component. However, since |\u03c1\u0302\u1d62\u2c7c| = |\u03c1\u0302\u2c7c\u1d62|, '
        'the data-driven component is perfectly symmetric, yielding zero directional signal '
        'when \u03b1 = 0 (no domain knowledge). This fails to satisfy the requirements of '
        'statistical causal inference.')

    p = doc.add_paragraph()
    run = p.add_run('Definition 4.1a ')
    run.bold = True
    p.add_run(
        '(Directional Predictability Index; DPI). For observed data X \u2208 \u211d^(N\u00d7n) '
        'from n variables, the DPI matrix D_DPI \u2208 \u211d^(n\u00d7n) is:'
    )

    add_display_equation(doc,
        'D_DPI(i\u2192j) = |\u03c1\u0302\u1d62\u2c7c| \u00b7 (1 + \u03b3 \u00b7 A\u0304(i,j))')

    add_rich_paragraph(doc,
        'where \u03b3 > 0 is the directional strength parameter (we use \u03b3 = 1 throughout), '
        'and A\u0304(i,j) is the mean of three normalized asymmetric statistics:')

    add_display_equation(doc,
        'A\u0304(i,j) = (1/3)[A\u0302_reg(i,j) + A\u0302_ANM(i,j) + A\u0302_ent(i,j)]')

    add_rich_paragraph(doc,
        '**(i) Regression coefficient asymmetry** A\u0302_reg: On unstandardized data, the simple '
        'regression coefficient \u03b2_j|i = Cov(X\u1d62,X\u2c7c)/Var(X\u1d62) is asymmetric when '
        'Var(X\u1d62) \u2260 Var(X\u2c7c).')

    add_rich_paragraph(doc,
        '**(ii) ANM residual independence** A\u0302_ANM: Based on the Additive Noise Model (ANM) '
        'principle, for each pair (i, j), we fit X\u2c7c = \u03b2X\u1d62 + \u03b5 and evaluate '
        'the independence of the residual \u03b5\u0302 from X\u1d62 using HSIC (Hilbert\u2013Schmidt '
        'Independence Criterion; kernel bandwidth via median heuristic).')

    add_rich_paragraph(doc,
        '**(iii) Conditional entropy reduction** A\u0302_ent: The entropy reduction '
        'H(X\u2c7c) \u2212 H(X\u2c7c|X\u1d62) is estimated via k-NN estimators.')

    p = doc.add_paragraph()
    run = p.add_run('Proposition 4.0a ')
    run.bold = True
    p.add_run('D_DPI is generally asymmetric: D_DPI(i\u2192j) \u2260 D_DPI(j\u2192i).')

    p = doc.add_paragraph()
    run = p.add_run('Proof')
    run.italic = True
    p.add_run(
        ': Since A\u0304(i,j) = \u2212A\u0304(j,i) (each component\u2019s normalized asymmetry '
        'is antisymmetric), D_DPI(i\u2192j) = |\u03c1\u0302|(1 + \u03b3A\u0304(i,j)) \u2260 '
        '|\u03c1\u0302|(1 \u2212 \u03b3A\u0304(i,j)) = D_DPI(j\u2192i) when A\u0304(i,j) \u2260 0. \u25a1'
    )

    add_rich_paragraph(doc,
        '**Hybrid utility function**: U(i,j) = \u03b1 \u00b7 C_domain(i,j) + (1\u2212\u03b1) '
        '\u00b7 D_DPI(i\u2192j). At \u03b1 = 0, the asymmetry of D_DPI preserves the directional '
        'signal; at \u03b1 > 0, domain knowledge injection improves accuracy.')

    doc.add_heading('4.2 Spectral Causal Coupling and Causal Direction', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Definition 4.2 ')
    run.bold = True
    p.add_run(
        '(Spectral Causal Coupling; SCC). For the eigendecomposition \u2112(q) = U\u039bU* '
        '(U = [u\u2081, \u2026, u\u2099]), the spectral causal coupling between nodes i, j is:'
    )

    add_display_equation(doc,
        'SCC(i,j) = \u03a3\u2096 f(\u03bb\u2096) \u00b7 |u\u2096(i)| \u00b7 |u\u2096(j)| '
        '\u00b7 cos(\u03b8\u2096(i) \u2212 \u03b8\u2096(j))')

    add_rich_paragraph(doc,
        'where f: \u211d\u2265\u2080 \u2192 \u211d\u2265\u2080 is the eigenvalue weighting function '
        '(typically f(\u03bb) = \u03bb). SCC is **symmetric**: SCC(i,j) = SCC(j,i).')

    p = doc.add_paragraph()
    run = p.add_run('Definition 4.3 ')
    run.bold = True
    p.add_run('(Spectral Causal Direction; SCD).')

    add_display_equation(doc,
        'SCD(i,j) = \u03a3\u2096 f(\u03bb\u2096) \u00b7 |u\u2096(i)| \u00b7 |u\u2096(j)| '
        '\u00b7 sin(\u03b8\u2096(i) \u2212 \u03b8\u2096(j))')

    add_rich_paragraph(doc,
        'SCD is **antisymmetric**: SCD(i,j) = \u2212SCD(j,i). SCD(i,j) > 0 suggests causal '
        'direction from i to j; SCD(i,j) < 0 suggests the reverse.')

    doc.add_heading('4.3 Unified Understanding of SCC and SCD', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Proposition 4.3 ')
    run.bold = True
    p.add_run(
        '(Complex Causal Index). Define the Complex Causal Index (CCI) as: '
        'CCI(i,j) = \u03a3\u2096 f(\u03bb\u2096)|u\u2096(i)||u\u2096(j)| '
        'exp(i(\u03b8\u2096(i) \u2212 \u03b8\u2096(j))). Then SCC = Re[CCI] and SCD = Im[CCI]. '
        'Geometrically, arg(CCI(i,j)) encodes the direction and |CCI(i,j)| encodes the strength '
        'of causal coupling.'
    )

    doc.add_heading('4.4 Properties of the SCD Matrix', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Proposition 4.4 ')
    run.bold = True
    p.add_run('The SCD matrix S \u2208 \u211d^(n\u00d7n) satisfies:')

    props_4 = [
        '(i) S is skew-symmetric: S = \u2212S\u1d40.',
        '(ii) tr(S) = 0.',
        '(iii) When q = 0, S = O (the zero matrix). Without directional information, causal direction cannot be estimated.',
    ]
    for prop in props_4:
        doc.add_paragraph(prop, style='List Bullet')

    add_rich_paragraph(doc,
        'Property (iii) is crucial: spectral causality **cannot function without directional '
        "information** (q > 0). This parallels LiNGAM\u2019s inability to function without "
        'non-Gaussianity.')

    doc.add_heading('4.5 Estimation of Causal Order', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Definition 4.4 ')
    run.bold = True
    p.add_run(
        '(Spectral Causal Score). The spectral causal score of each node i is: '
        's(i) = \u03a3_j\u2260i SCD(i,j). Nodes with larger s(i) are "upstream (cause-side)"; '
        'nodes with smaller s(i) are "downstream (effect-side)".'
    )

    # ============================================================
    # 5. Hodge Decomposition
    # ============================================================
    doc.add_heading('5. Hodge Decomposition: Orthogonal Decomposition of Causal Flows', level=1)

    doc.add_heading('5.1 Differential Forms on Graphs', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Definition 5.1 ')
    run.bold = True
    p.add_run(
        '(Chain Complex). For a graph G = (V, E), define: '
        '0-cochain C\u2070 = \u211d^|V| (functions on nodes), '
        '1-cochain C\u00b9 = \u211d^|E| (functions on edges = flows), '
        'Coboundary operator \u03b4\u2080: C\u2070 \u2192 C\u00b9: (\u03b4\u2080f)(i\u2192j) = f(j) \u2212 f(i) (gradient), '
        'Coboundary operator \u03b4\u2081: C\u00b9 \u2192 C\u00b2: curl on triangles.'
    )

    doc.add_heading('5.2 The Hodge Decomposition Theorem', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Theorem 5.1 ')
    run.bold = True
    p.add_run('(Hodge Decomposition on Graphs; Jiang et al. [9]). '
        'Any 1-cochain (edge flow) \u03c9 \u2208 C\u00b9 admits the orthogonal decomposition:')

    add_display_equation(doc,
        '\u03c9 = \u03b4\u2080\u03c6 (gradient) + \u03b4\u2081*\u03c8 (curl) + h (harmonic)')

    add_table(doc,
        ['Component', 'Mathematical meaning', 'Causal interpretation'],
        [
            ['\u03b4\u2080\u03c6 (gradient)', 'Flow driven by potential differences', 'Causal flow (unidirectional, DAG-like)'],
            ['\u03b4\u2081*\u03c8 (curl)', 'Local cyclic flow', 'Feedback loops (local interactions)'],
            ['h (harmonic)', 'Global cyclic flow', 'Homeostatic regulation (system-wide)'],
        ])

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig3_hodge_decomposition.png'),
        'Figure 4: Hodge decomposition of causal flows. Edge flows are orthogonally '
        'decomposed into gradient (DAG-like), curl (feedback), and harmonic (global cycle) components.')

    doc.add_heading('5.3 Causal Potential', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Definition 5.2 ')
    run.bold = True
    p.add_run(
        '(Causal Potential). The potential function \u03c6: V \u2192 \u211d in the gradient '
        'component \u03b4\u2080\u03c6 is the solution to the Poisson equation on the graph '
        'Laplacian: L\u03c6 = \u03b4\u2080*\u03c9. Since L is positive semi-definite, '
        '\u03c6 is unique up to an additive constant.'
    )

    add_rich_paragraph(doc,
        'The gradient energy ratio r_gradient = ||\u03b4\u2080\u03c6||\u00b2/||\u03c9||\u00b2 '
        'serves as an indicator of how well the data conforms to a DAG structure. '
        'r_gradient \u2248 1 implies the DAG assumption is adequate; r_gradient \u226a 1 implies '
        'feedback dominates.')

    # ============================================================
    # 6. Complementarity and Augmentative Extensibility
    # ============================================================
    doc.add_heading('6. Complementarity and Augmentative Extensibility with Existing Methods', level=1)

    add_rich_paragraph(doc,
        'This section positions spectral causality not as a **competitor** to existing methods '
        'but as a framework that **complements and mutually augments** them.')

    doc.add_heading('6.1 Complementarity with LiNGAM', level=2)

    add_rich_paragraph(doc,
        'LiNGAM (Linear Non-Gaussian Acyclic Model; Shimizu et al., 2006) assumes: '
        'x = Bx + e, e ~ non-Gaussian, independent. Identifiability hinges on (I \u2212 B) '
        'becoming lower triangular under a causal-order permutation.')

    add_table(doc,
        ['Capability', 'LiNGAM provides', 'Spectral causality provides', 'Complementary combination'],
        [
            ['Causal direction ID', 'Identifiability guarantee [3]', 'Direction estimation via DPI', 'Initialize spectral with LiNGAM edges'],
            ['Effect size', 'Direct estimation of B_ij', 'Relative strength via SCD', 'Integrate effect size with SCD'],
            ['Feedback detection', 'Impossible (DAG assumption)', 'Hodge curl component', "Spectral recovers LiNGAM's DAG residual"],
            ['Global structure', 'Edge-by-edge estimation', 'Spectral decomposition of whole graph', 'Local (LiNGAM) + global (spectral)'],
            ['Hill H6/H7/H9', 'Not covered', 'Covered via utility', 'ECD covers all 9 criteria'],
        ])

    add_rich_paragraph(doc,
        '**Bidirectional augmentation**: LiNGAM \u2192 Spectral: inject high-confidence LiNGAM '
        'edges as C_domain (\u201ctwo-stage rocket\u201d). Spectral \u2192 LiNGAM: quantify '
        'feedback loops via Hodge curl, provide interventionability via \u03c6, evaluate Hill '
        'H6/H7/H9 via utility.')

    doc.add_heading('6.2 Complementarity with Granger Causality', level=2)

    add_table(doc,
        ['Property', 'Granger Causality', 'Spectral Causality'],
        [
            ['Data type', 'Time series (longitudinal)', 'Cross-sectional snapshot'],
            ['Source of directionality', 'Temporal precedence', 'Structural asymmetry (DPI + domain knowledge)'],
            ['Unit of analysis', 'Sequential pairwise tests', 'Spectral structure of the whole graph'],
        ])

    doc.add_heading("6.3 Position on the Ladder of Causation", level=2)

    add_table(doc,
        ['Level', 'Question', 'Representative methods'],
        [
            ['3: Counterfactual', 'What would have happened if X = x?', 'Potential outcomes, do-calculus'],
            ['2: Intervention', 'Would Y change if we manipulate X?', 'RCT, IV, Mendelian randomization'],
            ['1.5: Informational causality \u22c6', 'What can we learn about Y from X?', 'Spectral causality, utility causality'],
            ['1: Association', 'Do X and Y co-vary?', 'Correlation, regression'],
        ])

    add_rich_paragraph(doc,
        'Spectral causality **bridges between levels**: DPI extracts directional information '
        'from correlations (Level 1 \u2192 1.5), the causal potential \u03c6 suggests '
        'interventionability (Level 1.5 \u2192 2), and RCT results can be reflected as domain '
        'knowledge (Level 2 \u2192 1.5).')

    doc.add_heading("6.4 Hill\u2019s Nine Criteria", level=2)

    add_rich_paragraph(doc,
        "Against Hill\u2019s nine criteria [10] for epidemiological causal judgment, "
        '**no single computational method can cover all nine criteria** (Figure 2).')

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig5_hill_radar.png'),
        "Figure 2: Radar chart showing each method\u2019s coverage of Hill\u2019s nine criteria. "
        'LiNGAM excels at H1 and H3 but lacks H6/H7/H9. The ECD ensemble integrates both, '
        'covering nearly all criteria.')

    doc.add_heading('6.5 Framework for Augmentative Extensibility', level=2)

    add_rich_paragraph(doc,
        'Spectral causality has a modular design permitting plug-in connection of external methods '
        'at multiple layers.')

    add_rich_paragraph(doc,
        '**DPI modular extension**: DPI is defined as the mean of K asymmetric statistics; '
        'additional components (transfer entropy, LiNGAM B_ij sign, nonlinear Granger, knockoff '
        'statistics, LLM causal scores) can augment DPI once normalized to [\u22121, 1].')

    add_rich_paragraph(doc,
        '**Staged accuracy improvement path**:')

    add_table(doc,
        ['Stage', 'Input', 'Method', 'Expected r_gradient'],
        [
            ['Stage 0', 'Data only (\u03b1 = 0)', 'DPI alone', '~0.58'],
            ['Stage 1', 'Data + LiNGAM (\u03b1 = 0.1\u20130.3)', 'ECD (two-stage rocket)', '~0.55\u20130.70'],
            ['Stage 2', 'Data + domain knowledge (\u03b1 = 0.3\u20130.6)', 'Spectral + expert C', '~0.70\u20130.86'],
            ['Stage 3', 'Data + domain + RCT (\u03b1 = 0.5\u20130.8)', 'Spectral + intervention validation', '~0.86+'],
            ['Stage 4', 'Full ECD ensemble', "All methods + Hill\u2019s 9 criteria", 'Comprehensive'],
        ])

    # ============================================================
    # 7. Related Work
    # ============================================================
    doc.add_heading('7. Related Work, Prior Literature, and Survey of Adjacent Fields', level=1)

    doc.add_heading('7.1 Prior Literature: Mathematical Foundations', level=2)

    add_rich_paragraph(doc,
        '**7.1.1 Magnetic Laplacian on Directed Graphs.** Fanuel & Suykens [11, 12] introduced '
        'the deformed Laplacian for spectral ranking and community detection in directed networks. '
        'de Resende & da Costa [7] studied spectra of magnetic Laplacians for large directed '
        'networks. Zhang et al. [8] developed MagNet, a GNN built on the magnetic Laplacian.')

    add_rich_paragraph(doc,
        '**7.1.2 Hodge Decomposition for Ranking and Flow Analysis.** Jiang et al. [9] established '
        'the theoretical foundation for Hodge decomposition on graphs. Maehara & Ohkawa [13, 14] '
        'extended Hodge decomposition to single-cell RNA sequencing data (ddHodge; Nature '
        'Communications, 2025).')

    add_rich_paragraph(doc,
        '**7.1.3 DAG-Based Graph Signal Processing.** Seifert, Wendler & P\u00fcschel [15] '
        'developed Causal Fourier Analysis on DAGs and posets. Misiakos, Mihal & P\u00fcschel [16] '
        '(ICASSP 2024) extended this to learning graph structures from time-series data. '
        'Stankovic et al. [17] proposed zero-padding techniques for Fourier analysis on DAGs.')

    doc.add_heading('7.2 Related Methods', level=2)

    add_rich_paragraph(doc,
        '**7.2.1 Continuous DAG Learning.** NOTEARS [18] and GOLEM [19] formulate DAG structure '
        'learning as continuous optimization. M\'Charrak et al. [20] proposed DAG learning for '
        'nonlinear models and possibly cyclic graphical models.')

    add_rich_paragraph(doc,
        '**7.2.2 Information-Theoretic Causality.** Transfer Entropy (TE) [21] extends Granger '
        'causality to the information-theoretic setting. Convergent Cross Mapping (CCM) [22] '
        'detects causality in deterministic dynamical systems.')

    add_rich_paragraph(doc,
        '**7.2.3 LiNGAM Extensions and Medical Applications.** DirectLiNGAM [5] serves as '
        'our baseline. Kotoku et al. [23] applied DirectLiNGAM to Osaka health checkup data. '
        'Okuda et al. [24] proposed workflow-constrained Longitudinal LiNGAM for Japanese '
        'health checkup cohorts.')

    doc.add_heading('7.3 Survey of Adjacent Fields', level=2)

    add_rich_paragraph(doc,
        '**7.3.1 LLMs and Causal Inference.** Le, Xia & Chen [25] proposed MAC (Multi-Agent '
        'Causal discovery) using LLM agents. Sheth, Fatemi & Fritz [26] evaluated LLMs\u2019 '
        'causal reasoning in CausalGraph2LLM.')

    add_rich_paragraph(doc,
        '**7.3.2 Biological Network Analysis.** Wein et al. [27] proposed GNN-based brain '
        'network causal inference. Bernal-Gonzalez et al. [28] proposed logical digraphs for '
        'biological control networks.')

    add_rich_paragraph(doc,
        '**7.3.3 Systematic Review.** Liu et al. [29] conducted a scoping review of causal '
        'discovery in observational medical research.')

    # ============================================================
    # 8. Empirical Illustration
    # ============================================================
    doc.add_heading('8. Empirical Illustration with Real Data', level=1)

    doc.add_heading('8.1 Data and Variables', level=2)
    add_rich_paragraph(doc,
        'We use five continuous variables from the UCI Heart Disease Dataset (Cleveland subset; '
        'Detrano et al. [30]):')

    add_rich_paragraph(doc,
        'X = (Age, RestingBP, Cholesterol, MaxHR, STDepression), sample size n = 297. '
        'All variables are standardized (mean 0, variance 1).')

    doc.add_heading('8.2 LiNGAM Causal Order (Baseline)', level=2)
    add_rich_paragraph(doc,
        'Applying DirectLiNGAM [5], the estimated causal order is: '
        'Age \u2192 MaxHR \u2192 STDep \u2192 RestBP \u2192 Chol. Major causal effects: '
        'B\u2084\u2082 = \u22120.395 (Age \u2192 MaxHR), '
        'B\u2082\u2081 = +0.309 (Age \u2192 RestingBP), '
        'B\u2085\u2084 = \u22120.348 (MaxHR \u2192 STDepression).')

    doc.add_heading('8.3 Eigenvectors of the Magnetic Laplacian', level=2)
    add_rich_paragraph(doc,
        'We construct the utility directed graph (mixing clinical knowledge and correlation at '
        '\u03b1 = 0.6) and compute the magnetic Laplacian at q \u2208 {0, 0.1, 0.25}.')

    add_table(doc,
        ['Variable', '|u\u2082| (amplitude)', '\u03b8\u2082 (phase angle, degrees)'],
        [
            ['Age', '0.53', '0.0\u00b0'],
            ['Resting BP', '0.35', '164.6\u00b0'],
            ['Cholesterol', '0.42', '\u221284.3\u00b0'],
            ['Max HR', '0.47', '34.7\u00b0'],
            ['ST Depression', '0.44', '\u221240.6\u00b0'],
        ])

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig2_magnetic_laplacian_q.png'),
        'Figure 3: Second eigenvector of the magnetic Laplacian plotted on the complex plane. '
        'At q = 0, all points lie on the real axis. At q = 0.25, variables spread across the '
        'complex plane, with phase angle ordering encoding causal flow direction.')

    doc.add_heading('8.4 Hodge Decomposition Results', level=2)

    add_rich_paragraph(doc,
        'Hodge decomposition yields: gradient (DAG-like causal flow) = 85.9%, curl (feedback) = 14.1%.')

    add_table(doc,
        ['Variable', '\u03c6 (Hodge potential)', 'Interpretation'],
        [
            ['Age', '0.000', 'Most upstream (exogenous)'],
            ['Cholesterol', '\u22120.127', ''],
            ['Resting BP', '\u22120.170', ''],
            ['Max HR', '\u22120.093', ''],
            ['ST Depression', '\u22120.255', 'Most downstream'],
        ])

    doc.add_heading('8.5 DPI-Based Analysis (\u03b1 = 0)', level=2)
    add_rich_paragraph(doc,
        'With DPI as the data-driven component at \u03b1 = 0 (no domain knowledge): '
        '**9 directed edges detected** (vs. 0 edges with the old |\u03c1|-based model), '
        '**r_gradient = 0.581**, **67% agreement with LiNGAM direction**. '
        'This demonstrates that spectral causality can perform causal direction estimation '
        'from data alone.')

    # ============================================================
    # 9. Structural Comparison
    # ============================================================
    doc.add_heading('9. LiNGAM vs. Spectral Causality: Structural Comparison', level=1)

    doc.add_heading('9.1 Edge-by-Edge Direction Comparison', level=2)
    add_rich_paragraph(doc,
        'Comparing causal directions estimated by three methods (LiNGAM, spectral causality with '
        '\u03b1 = 0.6, and DPI with \u03b1 = 0) (Figure 5):')

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig4_direction_comparison.png'),
        'Figure 5: Comparison of causal directions estimated by three methods\u2014LiNGAM, '
        'spectral causality (\u03b1 = 0.6), and DPI (\u03b1 = 0)\u2014on the UCI Heart Disease '
        'dataset (5 variables, 10 possible edges).')

    doc.add_heading('9.2 "Informational Direction" vs. "Interventional Causation"', level=2)
    add_rich_paragraph(doc,
        'When spectral causality suggests a direction opposite to LiNGAM, this is not necessarily '
        'an error. LiNGAM identifies **interventional causal direction**, while spectral causality '
        'captures **informational direction**. This distinction is precisely the Level 1.5 vs. '
        'Level 2 distinction on the Ladder of Causation (\u00a76.3).')

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig7_lingam_vs_spectral.png'),
        'Figure 6: Comparison of DAG (LiNGAM) and DCG (spectral causality) structures. LiNGAM '
        'produces a strict DAG under the acyclicity assumption, while spectral causality permits '
        'cycles (feedback loops), with edge feedback rates quantified by Hodge decomposition.')

    # ============================================================
    # 10. ECD Ensemble
    # ============================================================
    doc.add_heading('10. ECD Ensemble and Causal Upstream Potential', level=1)

    doc.add_heading('10.1 ECD Pipeline', level=2)
    add_rich_paragraph(doc,
        "Using LiNGAM\u2019s estimated result as domain knowledge: "
        'U_ECD(i\u2192j) = \u03b1 \u00b7 C_LiNGAM(i,j) + (1\u2212\u03b1) \u00b7 |corr(X\u1d62,X\u2c7c)|, '
        'where C_LiNGAM(i,j) = |B_ji|.')

    add_table(doc,
        ['Metric', 'Clinical knowledge (\u03b1 = 0.6)', 'ECD/LiNGAM (\u03b1 = 0.3)'],
        [
            ['r_gradient', '0.859', '0.555'],
            ['Edge count', '9', '6'],
            ['Hodge \u03c6 order', 'Age > Chol > BP \u2248 MaxHR > ST', 'Age > MaxHR > STDep > Chol > RestBP'],
        ])

    add_rich_paragraph(doc,
        "**Key finding**: The Hodge causal potential order under ECD is nearly identical to "
        "LiNGAM\u2019s causal order (only the bottom two variables are swapped).")

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig9_ecd_pruning_analysis.png'),
        'Figure 7: (A) Hodge decomposition of the ECD structure (blue = gradient, red = curl). '
        '(B) Correspondence between causal potential \u03c6 and interventionability \u03b9. '
        '(C) Edge-level feedback rates. (D) U-shaped relationship between domain knowledge '
        'quality (p_flip) and DAG degree.')

    doc.add_heading('10.2 Correspondence Between Causal Upstream Potential and Interventionability', level=2)

    add_table(doc,
        ['Variable', 'Hodge \u03c6', 'Interventionability \u03b9', 'Clinical rationale'],
        [
            ['Age', '0.000', 'Impossible (\u03b9 = 0)', 'Irreversible biological process'],
            ['MaxHR', '\u22120.204', 'Difficult (\u03b9 \u2248 0.3)', 'Depends on aging and constitution'],
            ['STDep', '\u22120.324', 'Indirect (\u03b9 \u2248 0.5)', 'Ischemia improved by PCI/CABG'],
            ['Chol', '\u22120.168', 'Easy (\u03b9 \u2248 0.9)', 'Statins'],
            ['RestBP', '\u22120.204', 'Easy (\u03b9 \u2248 0.8)', 'Antihypertensives'],
        ])

    add_rich_paragraph(doc,
        'Non-interventionable variables are exogenous and sit at the root of the DAG. '
        'A purely mathematical quantity (graph spectral structure) acquires practical clinical '
        'meaning as \u201cactionability.\u201d')

    # ============================================================
    # 11. DAG Transition Point Analysis
    # ============================================================
    doc.add_heading('11. DAG Transition Point Analysis', level=1)

    add_rich_paragraph(doc,
        'This section addresses the most practically important question: **"How much domain '
        'knowledge is required for DAG-like structure to emerge?"**')

    doc.add_heading('11.1 \u03b1-Sweep: Contrasting Two Models', level=2)

    add_rich_paragraph(doc,
        '**11.1.1 Old Model (|\u03c1|-based): Discontinuous Phase Transition.** '
        'r_gradient is identical at \u03b1 = 10\u207b\u2076 and \u03b1 = 1 (0.859). No smooth '
        'threshold exists.')

    p = doc.add_paragraph()
    run = p.add_run('Proposition 11.1 ')
    run.bold = True
    p.add_run(
        '(Scale-Invariance). When mixing a symmetric data matrix |\u03c1\u0302| with an asymmetric '
        'matrix C at scalar \u03b1 > 0, the asymmetric component is '
        'A_old(\u03b1) = \u03b1 \u00b7 (C \u2212 C\u1d40). Since r_gradient depends only on '
        'the structure (not magnitude) of the flow, r_gradient(\u03b1) = r_gradient(1) for all '
        '\u03b1 > 0. \u03b1 is a \u201cswitch,\u201d not a \u201cvolume knob.\u201d'
    )

    add_rich_paragraph(doc,
        '**11.1.2 DPI Model: Smooth Phase Transition.** With DPI, the transition becomes smooth:')

    add_table(doc,
        ['\u03b1', 'r_gradient', 'Edge count', 'Asymmetric norm'],
        [
            ['0.0', '0.581', '9', '0.815'],
            ['0.1', '0.598', '9', '0.755'],
            ['0.3', '0.688', '9', '0.719'],
            ['0.5', '0.792', '10', '0.803'],
            ['0.6', '0.824', '10', '0.882'],
            ['1.0', '0.859', '9', '1.327'],
        ])

    add_figure(doc,
        os.path.join(FIG_DIR, 'fig8_alpha_sweep.png'),
        'Figure 8: \u03b1-sweep with the DPI-based model. (A) r_gradient starts at 0.581 for '
        '\u03b1 = 0 and smoothly reaches 0.859 with increasing domain knowledge. (B) Edge count '
        'and LiNGAM agreement rate. (C) Asymmetric norm. (D) Phase diagram.')

    doc.add_heading('11.2 The True Threshold: Quality, Not Quantity', level=2)

    add_rich_paragraph(doc,
        'We randomly flip a fraction p_flip of the edge directions in the correct domain '
        'knowledge C_true (200 trials, \u03b1 = 0.6):')

    add_table(doc,
        ['p_flip', 'r_gradient (mean \u00b1 SD)', 'Interpretation'],
        [
            ['0.0', '0.859 \u00b1 0.000', 'Fully correct \u2192 high DAG'],
            ['0.1', '0.576 \u00b1 0.242', '10% error \u2192 sharp drop'],
            ['0.2', '0.443 \u00b1 0.226', 'Near random level'],
            ['0.3', '0.371 \u00b1 0.214', 'Minimum (maximum cyclicity)'],
            ['0.5', '0.516 \u00b1 0.232', 'Half flipped'],
            ['0.7', '0.733 \u00b1 0.164', 'Mostly flipped'],
            ['1.0', '0.859 \u00b1 0.000', 'Fully flipped \u2192 reversed DAG'],
        ])

    add_rich_paragraph(doc,
        '**The U-shaped curve**: p_flip = 0 and p_flip = 1 yield the same DAG degree; the '
        'minimum is at p_flip \u2248 0.3. Partial misinformation is worse than complete '
        'ignorance. The critical threshold is **p*_flip \u2248 0.15** (85%+ correct directions '
        'maintains DAG structure).')

    doc.add_heading('11.3 Leave-One-Edge-Out: Root Node Is the Backbone', level=2)

    add_table(doc,
        ['Removed edge', '\u0394r_gradient', 'Importance'],
        [
            ['Age \u2194 STDep', '\u22120.267', 'Highest'],
            ['Age \u2194 MaxHR', '\u22120.098', 'High'],
            ['Age \u2194 Chol', '\u22120.069', 'High'],
            ['Age \u2194 RestBP', '\u22120.040', 'Moderate'],
            ['Chol \u2194 STDep', '\u22120.054', 'Moderate'],
            ['RestBP \u2194 MaxHR', '+0.015', 'Removal improves'],
        ])

    add_rich_paragraph(doc,
        '**Finding**: Edges involving Age (root node = exogenous variable) form the backbone of '
        'the DAG structure. The minimal knowledge \u201cthis variable is not influenced by '
        'others (exogenous)\u201d provides maximum leverage.')

    doc.add_heading('11.4 Physical Analogy: Ising Model Correspondence', level=2)

    add_table(doc,
        ['Physical system (Ising model)', 'Causal estimation system (spectral causality)'],
        [
            ['Temperature T', 'Inverse of knowledge quality p_flip'],
            ['Order parameter (magnetization)', 'r_gradient (DAG degree)'],
            ['External field h', 'Knowledge quantity \u03b1 / DPI directional signal'],
            ['Critical temperature T_c', 'p*_flip \u2248 0.15'],
            ['Ferromagnetic phase (low T)', 'DAG causal structure (r_gradient > 0.5)'],
            ['Paramagnetic phase (high T)', 'Cyclic (DCG) structure (r_gradient < 0.5)'],
            ['Residual magnetization', "DPI\u2019s data-driven directional signal"],
        ])

    doc.add_heading('11.5 Summary of Three Thresholds', level=2)

    add_table(doc,
        ['Threshold', 'Value', 'Meaning', 'Practical implication'],
        [
            ['\u03b1* (knowledge quantity)', 'Smoothed by DPI', 'r_gradient = 0.581 at \u03b1 = 0', 'Precise \u03b1 setting unnecessary'],
            ['p*_flip (knowledge quality)', '\u2248 0.15', '85%+ correct maintains DAG', 'Few certain edges > many uncertain'],
            ['\u0394r* (backbone edges)', 'Root node (e.g., Age)', 'Exogenous variable essential', 'Minimal knowledge, maximum effect'],
        ])

    # ============================================================
    # 12. Cyclic Pruning
    # ============================================================
    doc.add_heading('12. Cyclic Structure Pruning and Practical Deployment', level=1)

    doc.add_heading('12.1 Feedback Is Clinically Correct', level=2)
    add_rich_paragraph(doc,
        'DAGs are mathematically convenient, but clinically, cyclic models are often more accurate: '
        'Exercise tolerance \u2194 Ischemia (MaxHR \u2192 STDep \u2192 \u2026 \u2192 MaxHR decline); '
        'Hypertension \u2194 Ischemia (RestBP \u2192 myocardial hypertrophy \u2192 worsened ischemia \u2192 '
        'sympathetic activation \u2192 RestBP increase).')

    doc.add_heading('12.2 Edge-Level Feedback Analysis', level=2)

    add_table(doc,
        ['Edge', 'Gradient direction', 'Feedback rate', 'Clinical interpretation'],
        [
            ['Age \u2192 RestBP', 'Age \u2192 RestBP', '0%', 'Pure unidirectional causation'],
            ['Age \u2192 Chol', 'Age \u2192 Chol', '1%', 'Pure unidirectional causation'],
            ['RestBP \u2194 STDep', 'STDep \u2192 RestBP', '24%', 'Weak hypertension-ischemia cycle'],
            ['Age \u2194 MaxHR', 'Age \u2192 MaxHR', '34%', 'Aging-fitness decline cycle'],
            ['MaxHR \u2194 STDep', 'MaxHR \u2192 STDep', '73%', 'Strong exercise-ischemia loop'],
        ])

    add_rich_paragraph(doc,
        "The **73% feedback rate for MaxHR \u2194 STDep** indicates that LiNGAM\u2019s DAG "
        'assumption (unidirectional MaxHR \u2192 STDep) misses a clinically real feedback loop.')

    doc.add_heading('12.3 Practical Deployment Pipeline', level=2)
    add_rich_paragraph(doc, 'Recommended operational workflow:')

    steps = [
        'Step 1: Estimate DAG via LiNGAM (no domain knowledge required)',
        'Step 2: Prune via bootstrap stability (retain edges appearing in >80% of trials)',
        'Step 3: Set retained edges as C_LiNGAM with low \u03b1 (0.01\u20130.1)',
        'Step 4: Apply spectral causality (Hodge decomposition): confirm DAG flow, quantify feedback, identify cyclic edges',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # ============================================================
    # 13. Theoretical Challenges and Outlook
    # ============================================================
    doc.add_heading('13. Theoretical Challenges and Outlook', level=1)

    doc.add_heading('13.1 Identifiability', level=2)

    add_rich_paragraph(doc,
        '**13.1.1 The Identifiability Problem of the Initial Model.** The initial spectral '
        'causality model (|\u03c1\u0302|-based) lacked an identifiability theory. Conjecture 13.1 '
        'requires that U(i,j) \u2212 U(j,i) has the same sign as the true causal direction\u2014'
        'a **circular argument** since |\u03c1\u0302| is symmetric.')

    add_rich_paragraph(doc,
        '**13.1.2 Resolution via DPI.** DPI partially resolves the circularity. Each DPI component '
        'directly extracts asymmetric directional signals from data:')

    add_table(doc,
        ['DPI component', 'Identifiability basis', 'Required conditions'],
        [
            ['Regression \u03b2 asymmetry', 'Var(X\u1d62) \u2260 Var(X\u2c7c) \u2192 |\u03b2_j|i| \u2260 |\u03b2_i|j|', 'Non-trivial variance ratio'],
            ['ANM residual independence', 'Hoyer et al. (2009) theorem: ANM identifiable', 'ANM assumption'],
            ['Conditional entropy', 'ICM assumption (Janzing & Sch\u00f6lkopf, 2010)', 'ICM assumption'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Proposition 13.1a ')
    run.bold = True
    p.add_run(
        '(Partial Identifiability of DPI). Under the additive noise model, the ANM component of '
        'DPI correctly identifies the direction as n \u2192 \u221e. This guarantees '
        'D_DPI(i\u2192j) \u2260 D_DPI(j\u2192i) at \u03b1 = 0, without external knowledge.'
    )

    add_rich_paragraph(doc,
        '**13.1.3 Integrated Identifiability Roadmap.**')

    add_rich_paragraph(doc,
        '**Phase 1 (Achieved)**: DPI component-level identifiability. The ANM component is '
        'guaranteed by Hoyer et al. (2009); regression \u03b2 asymmetry exploits the same '
        "information source as LiNGAM, so LiNGAM\u2019s identifiability theory applies.")

    add_rich_paragraph(doc,
        '**Phase 2 (Achievable)**: Spectral propagation consistency. When DPI correctly identifies '
        'root node directions, the Poisson equation L\u03c6 = \u03b4\u2080*\u03c9 uniquely '
        'determines \u03c6 (up to a constant). The LOEO experiment (\u00a711.3) provides '
        'empirical support.')

    add_rich_paragraph(doc,
        '**Phase 3 (Difficult but practically unnecessary)**: Complete identifiability is hard due '
        'to the spectral equivalence problem. However, ECD ensemble allows borrowing '
        "LiNGAM\u2019s identifiability for core edges while spectral causality complements "
        'with feedback quantification and Hill criteria coverage.')

    add_rich_paragraph(doc,
        '**13.1.4 Practical Identifiability Milestones:** (1) Consistency: as n \u2192 \u221e, '
        'each DPI component converges to the true asymmetry. (2) Partial identifiability: under '
        'ANM or non-Gaussian assumptions, root node directions are identifiable. (3) Robustness '
        'bound: p*_flip \u2248 0.15 guarantees DAG structure when DPI accuracy exceeds 85%.')

    doc.add_heading('13.2 Utility Function Construction and the Role of DPI', level=2)

    add_rich_paragraph(doc,
        'With DPI, spectral causality operates in a staged framework: '
        '(a) No domain knowledge (\u03b1 = 0): DPI asymmetry alone (r_gradient = 0.581). '
        '(b) With domain knowledge (\u03b1 > 0): smooth improvement to 0.859. '
        '(c) LiNGAM ensemble (ECD): borrowing identifiability for core edges.')

    doc.add_heading('13.3 Future Directions', level=2)

    directions = [
        'Phase 2 identifiability proof: Spectral propagation consistency for tree DAG + linear SEM',
        'ECD pipeline validation: Reproducibility on MIMIC-IV and Japanese health checkup cohorts (n > 10\u2075; Okuda et al. [24])',
        'Extension to longitudinal data: Temporal utility graphs and eigentrajectory extraction',
        'Automatic p_flip estimation: Meta-method from LiNGAM directional agreement rates',
        'Automated pruning thresholds: Bootstrap confidence intervals of feedback rates',
        'Improved domain knowledge encoding: From "informational influence" to "interventional causal strength"',
        'Data-adaptive \u03b1: Bootstrap consistency test between asymmetric matrix and data correlation structure',
        'Phase transition proof: Rigorous proof that r_gradient depends only on the sign structure of the asymmetric component',
        'Multi-variable scaling: From 5 variables (10 edges) to 50+ variables (>1225 edges)',
        'Reproducibility: U-shaped curve and p*_flip \u2248 0.15 validation on MIMIC-IV and other datasets',
    ]
    for i, d in enumerate(directions, 1):
        doc.add_paragraph(f'{i}. {d}', style='List Number')

    # ============================================================
    # Data Availability Statement
    # ============================================================
    doc.add_heading('Data Availability Statement', level=1)
    add_rich_paragraph(doc,
        'The UCI Heart Disease Dataset (Cleveland subset) used in this study is publicly '
        'available at the UCI Machine Learning Repository '
        '(https://archive.ics.uci.edu/ml/datasets/Heart+Disease). The analysis code and '
        'all generated figures are available in the supplementary repository.')

    # ============================================================
    # Conflict of Interest
    # ============================================================
    doc.add_heading('Conflict of Interest', level=1)
    doc.add_paragraph('The authors declare no conflict of interest.')

    # ============================================================
    # References
    # ============================================================
    doc.add_heading('References', level=1)

    refs = [
        '[1] Pearl, J. (2009). Causality: Models, Reasoning, and Inference (2nd ed.). Cambridge University Press.',
        '[2] Rubin, D.B. (1974). Estimating causal effects of treatments in randomized and nonrandomized studies. Journal of Educational Psychology, 66(5), 688\u2013701.',
        '[3] Shimizu, S., Hoyer, P.O., Hyvarinen, A. & Kerminen, A. (2006). A linear non-Gaussian acyclic model for causal discovery. Journal of Machine Learning Research, 7, 2003\u20132030.',
        '[4] Granger, C.W.J. (1969). Investigating causal relations by econometric models and cross-spectral methods. Econometrica, 37(3), 424\u2013438.',
        '[5] Shimizu, S., Inazumi, T., Sogawa, Y., Hyvarinen, A., Kawahara, Y., Washio, T., Hoyer, P.O. & Bollen, K. (2011). DirectLiNGAM: A direct method for learning a linear non-Gaussian structural equation model. Journal of Machine Learning Research, 12, 1225\u20131248.',
        '[6] Shuman, D.I., Narang, S.K., Frossard, P., Ortega, A. & Vandergheynst, P. (2013). The emerging field of signal processing on graphs. IEEE Signal Processing Magazine, 30(3), 83\u201398.',
        '[7] de Resende, B.M.F. & da Costa, L.F. (2020). Characterization and comparison of large directed networks through the spectra of the magnetic Laplacian. Chaos, 30(7), 073141.',
        '[8] Zhang, X., He, Y., Bruger, N., Hooi, B. & Zhu, L. (2022). MagNet: A neural network for directed graphs. In NeurIPS 2021.',
        '[9] Jiang, X., Lim, L.H., Yao, Y. & Ye, Y. (2011). Statistical ranking and combinatorial Hodge theory. Mathematical Programming, 127, 203\u2013244.',
        "[10] Hill, A.B. (1965). The environment and disease: Association or causation? Proceedings of the Royal Society of Medicine, 58, 295\u2013300.",
        '[11] Fanuel, M. & Suykens, J.A.K. (2017). Deformed Laplacians and spectral ranking in directed networks. arXiv:1511.00492.',
        '[12] Fanuel, M., Alaiz, C.M. & Suykens, J.A.K. (2017). Magnetic eigenmaps for community detection in directed networks. Physical Review E, 95, 022302.',
        '[13] Maehara, K. & Ohkawa, Y. (2019). Modeling latent flows on single-cell data using the Hodge decomposition. bioRxiv.',
        '[14] Maehara, K. & Ohkawa, Y. (2025). Geometry-preserving vector field reconstruction of high-dimensional cell-state dynamics using ddHodge. Nature Communications, 16, 11342.',
        '[15] Seifert, B., Wendler, C. & P\u00fcschel, M. (2023). Causal Fourier analysis on directed acyclic graphs and posets. IEEE Transactions on Signal Processing, 71, 3516\u20133530.',
        '[16] Misiakos, P., Mihal, V. & P\u00fcschel, M. (2024). Learning signals and graphs from time-series graph data with few causes. In IEEE ICASSP 2024.',
        '[17] Stankovic, L. et al. (2024). Fourier analysis of signals on directed acyclic graphs (DAG) using graph zero-padding. arXiv:2311.01073.',
        '[18] Zheng, X., Aragam, B., Ravikumar, P. & Xing, E.P. (2018). DAGs with NO TEARS: Continuous optimization for structure learning. In NeurIPS 2018.',
        '[19] Ng, I., Ghassami, A. & Zhang, K. (2020). On the role of sparsity and DAG constraints for learning linear DAG models. In NeurIPS 2020.',
        "[20] M'Charrak, I., Luengo-Sanchez, S. & Ruiz, C. (2025). Causal structure learning in directed, possibly cyclic, graphical models. Journal of Causal Inference, 13(1), 20240078.",
        '[21] Schreiber, T. (2000). Measuring information transfer. Physical Review Letters, 85(2), 461\u2013464.',
        '[22] Sugihara, G. et al. (2012). Detecting causality in complex ecosystems. Science, 338(6106), 496\u2013500.',
        '[23] Kotoku, J. et al. (2020). Causal relations of health indices inferred statistically using the DirectLiNGAM algorithm from big data of Osaka prefecture health checkups. PLoS ONE, 15(12), e0243229.',
        '[24] Okuda, S. et al. (2025). Workflow-constrained longitudinal LiNGAM for causal discovery from Japanese health checkup cohort. Journal of Biomedical Informatics, 152, 104612.',
        '[25] Le, T.D., Xia, F. & Chen, J. (2024). MAC: Multi-agent causal discovery with large language models. arXiv:2402.09812.',
        '[26] Sheth, I., Fatemi, B. & Fritz, M. (2024). CausalGraph2LLM: Evaluating LLMs for causal queries. arXiv:2405.00527.',
        '[27] Wein, S. et al. (2024). Graph neural networks for brain network causal inference. NeuroImage, 289, 120551.',
        '[28] Bernal-Gonzalez, J.P. et al. (2023). Logical digraphs for biological control networks. Frontiers in Physics, 11, 1272674.',
        '[29] Liu, Y. et al. (2024). Causal discovery from observational medical data: A scoping review. BMC Medical Research Methodology, 24(1), 89.',
        '[30] Detrano, R. et al. (1989). International application of a new probability algorithm for the diagnosis of coronary artery disease. American Journal of Cardiology, 64(5), 304\u2013310.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9)

    # ============================================================
    # Save
    # ============================================================
    out_path = os.path.join(BASE_DIR, '06_spectral_causality_academic_en.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


if __name__ == '__main__':
    generate_docx()
