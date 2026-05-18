#!/usr/bin/env python3
"""Generate Electoral Studies submission documents (cover letter, title page, highlights, declarations)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import date

OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'output', 'submission')
os.makedirs(OUT_DIR, exist_ok=True)


# ══════════════════════════════════════════════
# 1. Cover Letter
# ══════════════════════════════════════════════
def create_cover_letter():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15

    today = date.today().strftime('%B %d, %Y')

    # Date
    p = doc.add_paragraph(today)
    p.paragraph_format.space_after = Pt(12)

    # Addressee
    lines = [
        'Professor Rosie Campbell, Professor Oliver Heath, and Professor Nick Vivyan',
        'Co-Editors in Chief',
        'Electoral Studies',
    ]
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph('')

    # Salutation
    p = doc.add_paragraph('Dear Professors Campbell, Heath, and Vivyan,')
    p.paragraph_format.space_after = Pt(12)

    # Body
    paragraphs = [
        'We are pleased to submit our manuscript entitled "Trust-Adjusted Transparent Scoring '
        'with Unified Knowledge Integration (TATSUKI): An Agent-Based Model of Accountability-Driven '
        'Electoral Reform with Empirical Calibration" for consideration for publication in Electoral Studies.',

        'This paper introduces TATSUKI, a novel electoral mechanism that links candidate trust '
        'coefficients to measured policy fulfillment through influence functions, creating continuous '
        'accountability incentives within representative democracy. Unlike binary electoral sanctions '
        '(re-election vs. defeat), TATSUKI modulates candidates\u2019 effective vote share based on '
        'their demonstrated pledge fulfillment, thereby incentivizing sincere policy delivery.',

        'We believe this work is particularly well-suited for Electoral Studies for three reasons. '
        'First, it directly addresses the core question of how electoral institutions can enhance '
        'democratic accountability\u2014a central concern of your readership. Second, the paper bridges '
        'theoretical mechanism design with empirical evidence: we calibrate the model using '
        'the Polimeter pledge-tracking project (1,050 coded promises from the Trudeau government '
        'across three parliamentary terms, 2015\u20132025) and cross-national benchmarks from '
        'Thomson et al. (2017, AJPS; 20,000+ pledges, 12 countries, 57 elections). Third, the '
        'counterfactual analysis\u2014retroactively applying TATSUKI to real-world Canadian data\u2014'
        'demonstrates the mechanism\u2019s empirical plausibility and provides a template for '
        'future applied research on electoral accountability.',

        'Key contributions of the manuscript include: (i) a formal specification of the TATSUKI '
        'mechanism with three families of influence functions (concave, linear, sigmoid); '
        '(ii) an ODD-protocol compliant agent-based model demonstrating TATSUKI\u2019s effects on '
        'accountability scores, candidate-type evolution, and voter welfare; (iii) adversarial '
        'robustness analysis using genetic algorithm search; and (iv) empirical calibration using '
        'Polimeter and Thomson et al. data, with counterfactual trust trajectory analysis for '
        'a real-world minority government.',

        'The manuscript has not been published elsewhere and is not under consideration by any '
        'other journal. All authors have approved the manuscript and agree with its submission '
        'to Electoral Studies. The authors declare no competing interests.',

        'We confirm that this submission complies with the journal\u2019s double-anonymized review '
        'policy: the manuscript contains no author-identifying information. Author details are '
        'provided in the separate title page.',

        'Thank you for considering our submission. We look forward to receiving your decision.',
    ]

    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('')

    # Closing
    closing_lines = [
        'Sincerely,',
        '',
        '[Corresponding Author Name]',
        '[Affiliation]',
        '[Email Address]',
        '[ORCID]',
    ]
    for line in closing_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    out = os.path.join(OUT_DIR, 'Cover_Letter_Electoral_Studies.docx')
    doc.save(out)
    print(f'Cover letter saved to {out}')


# ══════════════════════════════════════════════
# 2. Title Page (separate for double-anonymized review)
# ══════════════════════════════════════════════
def create_title_page():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Title
    p = doc.add_paragraph()
    run = p.add_run(
        'Trust-Adjusted Transparent Scoring with Unified Knowledge Integration (TATSUKI):\n'
        'An Agent-Based Model of Accountability-Driven Electoral Reform\n'
        'with Empirical Calibration'
    )
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Author 1 Name]')
    run.font.size = Pt(12)
    run = p.add_run('a')
    run.font.superscript = True
    run = p.add_run('*')
    run.font.superscript = True
    p.paragraph_format.space_after = Pt(12)

    # Affiliations
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('a')
    run.font.superscript = True
    run = p.add_run(' [Affiliation, City, Country]')
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(24)

    # Corresponding author
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('* Corresponding author:')
    run.bold = True
    p.paragraph_format.space_after = Pt(3)

    info = [
        'Email: [email@example.com]',
        'Address: [Full postal address]',
        'ORCID: [0000-0000-0000-0000]',
    ]
    for line in info:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(1.0)

    # Word count
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Word count: ')
    run.bold = True
    p.add_run('approximately 9,500 words (excluding references and figure captions)')
    p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    run = p.add_run('Number of figures: ')
    run.bold = True
    p.add_run('7')
    p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    run = p.add_run('Number of tables: ')
    run.bold = True
    p.add_run('0')
    p.paragraph_format.space_after = Pt(12)

    # Acknowledgments
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Acknowledgments')
    run.bold = True
    p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        '[Add acknowledgments here. Note: acknowledgments are placed on the title page '
        'to maintain double-anonymized review of the main manuscript.]'
    )

    # Funding
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Funding')
    run.bold = True
    p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        '[Specify funding sources, if any. If none, state: '
        '"This research did not receive any specific grant from funding agencies '
        'in the public, commercial, or not-for-profit sectors."]'
    )

    out = os.path.join(OUT_DIR, 'Title_Page.docx')
    doc.save(out)
    print(f'Title page saved to {out}')


# ══════════════════════════════════════════════
# 3. Highlights
# ══════════════════════════════════════════════
def create_highlights():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph()
    run = p.add_run('Highlights')
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(18)

    # Elsevier highlights: 3-5 bullet points, max 85 characters each
    highlights = [
        'TATSUKI links candidate trust coefficients to measured pledge fulfillment.',
        'Agent-based model shows TATSUKI raises accountability and resists gaming.',
        'Concave and sigmoid influence functions optimally balance incentives.',
        'Model calibrated with 1,050 real pledges from Polimeter and Thomson data.',
        'Counterfactual analysis confirms stable trust trajectories for Trudeau.',
    ]

    for h in highlights:
        p = doc.add_paragraph(h, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
        # Verify length
        if len(h) > 85:
            print(f'  WARNING: Highlight exceeds 85 chars ({len(h)}): {h[:50]}...')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.italic = True
    p.add_run('Each highlight must be no more than 85 characters including spaces. '
              'These are entered separately during the Editorial Manager submission process.').italic = True

    out = os.path.join(OUT_DIR, 'Highlights.docx')
    doc.save(out)
    print(f'Highlights saved to {out}')


# ══════════════════════════════════════════════
# 4. Declarations
# ══════════════════════════════════════════════
def create_declarations():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Declaration of Competing Interest
    p = doc.add_paragraph()
    run = p.add_run('Declaration of Competing Interest')
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        'The authors declare that they have no known competing financial interests or personal '
        'relationships that could have appeared to influence the work reported in this paper.'
    )

    doc.add_paragraph('')

    # CRediT Author Statement
    p = doc.add_paragraph()
    run = p.add_run('CRediT Author Statement')
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        '[Author 1]: Conceptualization, Methodology, Software, Formal analysis, '
        'Investigation, Data curation, Writing \u2013 Original Draft, Writing \u2013 Review & Editing, '
        'Visualization.'
    )

    doc.add_paragraph('')

    # Data Availability Statement
    p = doc.add_paragraph()
    run = p.add_run('Data Availability Statement')
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        'The Polimeter data used for empirical calibration are publicly available at '
        'https://polimeter.org. The Thomson et al. (2017) cross-national pledge fulfillment '
        'data are available in the published article (American Journal of Political Science, '
        '61(3), 527\u2013542). The simulation code and calibration scripts are available '
        '[in the supplementary materials / at the repository URL upon acceptance].'
    )

    doc.add_paragraph('')

    # Ethics Statement
    p = doc.add_paragraph()
    run = p.add_run('Ethics Statement')
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        'This study uses only publicly available aggregate data (Polimeter pledge tracking data '
        'and published cross-national statistics). No human subjects were involved in data '
        'collection. Ethical approval was not required for this research.'
    )

    out = os.path.join(OUT_DIR, 'Declarations.docx')
    doc.save(out)
    print(f'Declarations saved to {out}')


# ══════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════
if __name__ == '__main__':
    create_cover_letter()
    create_title_page()
    create_highlights()
    create_declarations()
    print('\nAll submission documents created successfully!')
    print(f'Output directory: {OUT_DIR}')
