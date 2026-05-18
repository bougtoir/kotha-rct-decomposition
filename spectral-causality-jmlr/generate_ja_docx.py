#!/usr/bin/env python3
"""
Generate Japanese DOCX manuscript from JMLR spectral causality paper.
All figures and tables are placed inline immediately after first mention.
Uses font-based superscript for citation numbers (not Unicode).
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FIGURES_DIR = "figures"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, style="Normal", bold=False, italic=False, font_size=None, alignment=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    return p


def add_math_para(doc, text, font_size=10):
    """Add a paragraph with math-like text (italicized)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(font_size)
    return p


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Add a figure with caption immediately inline."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=width)
    else:
        p.add_run(f"[図: {img_path}]")

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    return cap


def add_table(doc, headers, rows, caption=None):
    """Add a table with optional caption."""
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.space_before = Pt(12)
        run = cap_p.add_run(caption)
        run.bold = True
        run.font.size = Pt(9)

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing
    return table


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Yu Mincho"
    font.size = Pt(10.5)

    # ============================================================
    # Title Page
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run("スペクトル因果性:\n磁気ラプラシアンとHodge分解による因果方向推定")
    run.bold = True
    run.font.size = Pt(16)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("大西 達貴")
    run.font.size = Pt(12)

    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = affil.add_run("麻酔科学教室\n大学病院, 日本")
    run.font.size = Pt(10)
    affil.paragraph_format.space_after = Pt(24)

    # ============================================================
    # Abstract
    # ============================================================
    add_heading(doc, "要旨", level=1)
    add_para(doc, (
        "本論文ではスペクトル因果性を提案する。これは磁気ラプラシアン—辺の方向性を複素位相として"
        "符号化するエルミート行列—のスペクトル構造を活用して、観測変数間の因果方向を推定する枠組みである。"
        "古典的なラプラシアン手法が方向情報を破棄するのに対し、磁気ラプラシアンの固有ベクトルは"
        "方向依存の位相を持ち、因果的上流ノードと下流ノードを分離する。"
        "対称的なSpectral Causal Coupling (SCC) と反対称的なSpectral Causal Direction (SCD) の"
        "2つの相補的指標を定義し、これらが単一の複素因果指標の実部・虚部として導出されることを示す。"
    ))
    add_para(doc, (
        "中心的な理論的貢献はDirectional Predictability Index (DPI) である。"
        "回帰係数の非対称性、加法ノイズモデル (ANM) 残差独立性、条件付きエントロピー削減の"
        "3つの非対称統計量からなるデータ駆動型の指標であり、"
        "ANM仮定下でDPIが外部知識なしに因果方向の整合的推定量となることを証明する。"
    ))
    add_para(doc, (
        "本枠組みをグラフ上のHodge分解と接続し、因果辺フローの勾配-回転-調和分解が"
        "DAG構造下でトポロジカルソートを回復する因果ポテンシャルを生成することを証明する。"
        "勾配エネルギー比 r_gradient のスケール不変性を確立し、"
        "因果構造の出現が知識の量ではなく質に支配される相転移現象を特徴付ける。"
    ))
    add_para(doc, (
        "UCI心疾患データセット (n=297, 臨床変数5個) での実験により、"
        "DPIを用いたスペクトル因果性はα=0 (ドメイン知識なし) で9本の有向辺を検出し、"
        "Direct Linear Non-Gaussian Acyclic Model (DirectLiNGAM; Shimizu et al., 2011) と67%の方向一致を達成する。"
        "Hodge分解はDAGベース手法では不可視な臨床的に意味のあるフィードバックループを明らかにする。"
        "LiNGAMの識別可能性保証とスペクトル因果性のフィードバック定量化を統合する"
        "Ensemble Causal Direction (ECD) パイプラインを提案し、"
        "Hillの因果推論9基準の包括的カバレッジを達成する。"
    ))

    kw = doc.add_paragraph()
    run = kw.add_run("キーワード: ")
    run.bold = True
    kw.add_run("因果探索, 磁気ラプラシアン, Hodge分解, スペクトルグラフ理論, 有向グラフ, 因果推論")

    # ============================================================
    # Section 1: Introduction
    # ============================================================
    add_heading(doc, "1. はじめに", level=1)

    add_para(doc, (
        "因果推論—XがYの原因であるか否かの判定—は科学・医学の中心的問題である。"
        "主要なアプローチとして、do計算を用いた構造方程式モデル (Pearl, 2009)、"
        "潜在結果 (Rubin, 1974)、線形非ガウス非巡回モデル (LiNGAM) (Shimizu et al., 2006)、"
        "時系列データに対するGranger因果性 (Granger, 1969) がある。"
    ))
    add_para(doc, (
        "本論文では根本的に異なる原理を提案する: グラフラプラシアンのスペクトル構造 "
        "(固有値・固有ベクトル) から因果方向を読み取る。このアプローチをスペクトル因果性と呼ぶ。"
    ))

    add_para(doc, "基本的アイデア", bold=True)
    add_para(doc, (
        "n個の変数 {X_1,...,X_n} に推定される因果関係を重み付き有向グラフ G=(V,E,w) として表現し、"
        "磁気ラプラシアン L^(q) はエルミート行列であり、その固有ベクトルは複素数値を取る。"
        "固有ベクトルの位相角が辺の方向性を符号化し、スペクトル構造のみから因果方向推定を可能にする。"
    ))

    add_para(doc, "DAGベース手法との主要な差異", bold=True)
    add_para(doc, (
        "LiNGAMおよび関連手法はDAG (有向非巡回グラフ) を仮定する。"
        "生物医学系ではフィードバックループが偏在している "
        "(例: 炎症→臓器障害→炎症)。"
        "スペクトル因果性はHodge分解を通じて有向巡回グラフ (DCG) に対応し、"
        "因果フローの勾配 (DAG的) 成分と回転 (フィードバック) 成分を分離する。"
    ))

    add_para(doc, "貢献", bold=True)
    add_para(doc, (
        "(1) Directional Predictability Index (DPI): α=0 (ドメイン知識なし) での因果方向推定を可能にする"
        "データ駆動型非対称統計量（第4節）。\n"
        "(2) DAGフリー因果推論: Hodge分解により勾配 (DAG互換フロー) とカール (フィードバック) を分離し、"
        "勾配エネルギー比 r_gradient をDAG適合度の定量的指標として提供（第5節）。\n"
        "(3) スケール不変性定理: r_gradient がユーティリティ行列の非対称成分の構造 (符号パターン) のみに"
        "依存し、大きさには依存しないことを証明（第8節）。\n"
        "(4) Ensemble Causal Direction (ECD): LiNGAMの局所的識別可能性とスペクトル因果性の"
        "大域的構造解析を統合するパイプライン（第9節）。\n"
        "(5) DPIの部分的識別可能性: ANM下でDPIのANM成分が n→∞ で因果方向を整合的に識別することを証明（第10節）。"
    ))

    add_para(doc, (
        "図1に本研究で統合する3つの相補的アプローチの概念概要を示す。"
    ))

    # Figure 1: Three approaches (conceptual overview)
    add_figure(doc,
        f"{FIGURES_DIR}/fig1_three_approaches.png",
        "図1: 因果推論への3つの相補的アプローチ: (1) 構造方程式モデルとDAGベース手法 (例: LiNGAM)、"
        "(2) 磁気ラプラシアンによるスペクトルグラフ手法、(3) Hodge理論的フロー分解。"
        "スペクトル因果性は(2) と(3) を統合する。",
        width=Inches(5.5)
    )

    # ============================================================
    # Section 2: Preliminaries
    # ============================================================
    add_heading(doc, "2. 準備: グラフラプラシアン", level=1)

    add_para(doc, "定義 2.1 (グラフラプラシアン)", bold=True)
    add_para(doc, (
        "重み付き無向グラフ G=(V,E,w), |V|=n に対し、重み付き隣接行列 W ∈ R^{n×n}、"
        "次数行列 D = diag(d_1,...,d_n) (d_i = Σ_j W_{ij})、グラフラプラシアンを定義する:\n"
        "L = D - W（非正規化）\n"
        "L_norm = I - D^{-1/2} W D^{-1/2}（正規化）"
    ))

    add_para(doc, "命題 2.1 (基本性質)", bold=True)
    add_para(doc, (
        "(i) L は対称半正定値であり、固有値は 0 = λ_1 ≤ λ_2 ≤ ... ≤ λ_n。\n"
        "(ii) λ_1 = 0 の固有ベクトルは 1 = (1,...,1)^T。\n"
        "(iii) λ_2 > 0 ⟺ G は連結（代数的連結性/Fiedler値）。\n"
        "(iv) 任意の f ∈ R^n に対し f^T L f = Σ_{(i,j)∈E} w_{ij}(f_i - f_j)^2 ≥ 0。"
    ))

    add_para(doc, (
        "性質 (iv) は根本的である: f^T L f は隣接ノードに類似の値を割り当てる f に対して小さくなる。"
        "低固有値の固有ベクトルはグラフ上の「滑らかな信号」を表し、"
        "グラフ信号処理の基礎をなす (Shuman et al., 2013)。"
    ))

    add_para(doc, "注意 2.1 (無向ラプラシアンの限界)", bold=True, italic=True)
    add_para(doc, (
        "L = D - W は対称であるため、方向 i→j と j→i を区別できない。"
        "有向ラプラシアン L_d = D_out - W は一般に非対称で複素固有値を持ち、"
        "スペクトル解析が困難である。磁気ラプラシアン（第3節）は方向を複素位相として"
        "符号化しつつ、エルミート性（したがって実固有値）を保存することでこの問題を解決する。"
    ))

    # ============================================================
    # Section 3: Magnetic Laplacian
    # ============================================================
    add_heading(doc, "3. 磁気ラプラシアン", level=1)

    add_heading(doc, "3.1 物理的動機", level=2)
    add_para(doc, (
        "磁気ラプラシアンは量子力学に起源を持つ。磁場 B 中のベクトルポテンシャル A の下で、"
        "荷電粒子のハミルトニアンは H = (p - eA)^2 / 2m である。"
        "閉じたループを横断する粒子はAharonov-Bohm位相 exp(i ∮ A·dr) を獲得し、"
        "その方向依存性をグラフ辺の方向符号化に転用できる。"
    ))

    add_heading(doc, "3.2 定義", level=2)
    add_para(doc, "定義 3.1 (磁気ラプラシアン)", bold=True)
    add_para(doc, (
        "重み付き有向グラフ G=(V,E,w) とチャージパラメータ q ∈ [0, 0.5] に対し、"
        "エルミート隣接行列 H^(q) ∈ C^{n×n} を次のように定義する:\n"
        "H^(q)_{ij} = w_{ij} · exp(i · 2πq · σ_{ij})\n"
        "ここで w_{ij} は対称化重み、σ_{ij} ∈ {-1, 0, +1} は方向符号である。"
        "正規化磁気ラプラシアンは:\n"
        "L^(q) = I - D^{-1/2} H^(q) D^{-1/2}"
    ))

    add_para(doc, "命題 3.1 (磁気ラプラシアンの性質)", bold=True)
    add_para(doc, (
        "(i) H^(q) はエルミート: H^(q)_{ji} = conj(H^(q)_{ij})。\n"
        "(ii) L^(q) はエルミート半正定値で、実非負固有値を持つ。\n"
        "(iii) L^(q) の固有ベクトルは一般に複素数値。\n"
        "(iv) q=0 では L^(0) は標準正規化ラプラシアンに帰着する（方向情報なし）。"
    ))

    add_heading(doc, "3.3 チャージパラメータ q", level=2)

    add_table(doc,
        ["q", "位相 2πq", "効果"],
        [
            ["0", "0", "方向なし。exp(0)=1; 実行列。"],
            ["0.25", "π/2", "最大感度。e^{iπ/2}=i, e^{-iπ/2}=-i。"],
            ["0.5", "π", "方向反転。e^{iπ}=-1。"],
        ],
        caption="表 3.1: チャージパラメータの効果"
    )

    add_heading(doc, "3.4 位相角と方向性", level=2)
    add_para(doc, (
        "L^(q) の各固有ベクトル u_k ∈ C^n は極座標分解を許容する:\n"
        "u_k(j) = |u_k(j)| · exp(i · θ_k(j))\n"
        "ここで |u_k(j)| は振幅、θ_k(j) = arg(u_k(j)) はノード j における位相角である。"
    ))

    add_para(doc, "定理 3.1 (位相-方向対応)", bold=True)
    add_para(doc, (
        "q > 0 かつ整合的因果順序を持つユーティリティ有向グラフに対し、"
        "低周波固有ベクトルの位相角 θ_k(j) は因果的上流ノード (ソース) と"
        "下流ノード (シンク) を複素平面上で分離する。具体的には、Fiedler固有ベクトル u_2 において、"
        "ノード i がノード j より因果的に上流 (φ(i) > φ(j)) ならば、"
        "θ_2(i) と θ_2(j) は異なる角度セクターを占める傾向がある。"
    ))
    add_para(doc, (
        "証明: 有向辺 i→j での位相蓄積が鍵となる。q=0.25 では "
        "H^(q)_{ij} = i·w_{ij}, H^(q)_{ji} = -i·w_{ij} であり、"
        "固有方程式の非対称位相因子が各有向辺に沿って厳密に正の位相増分 Δθ > 0 を強制する。"
        "第1段階 (補題A.1): パスグラフ v_1→...→v_n では、正規化磁気ラプラシアンの"
        "固有方程式が再帰式 (1-λ)u(v_k) = (-iw/d_{v_k})u(v_{k-1}) + (iw/d_{v_k})u(v_{k+1}) "
        "を与える。u(v_k) = r_k exp(iθ_k) と極座標表示し実部を取ると、"
        "1-λ_2 > 0 かつ r_k > 0 から θ_{k+1}-θ_k > 0 が要求される。"
        "第2段階 (定理A.1): tree DAGに対し木の深さに関する構造的帰納法を用いる。"
        "基底（星グラフ）で θ_2(c_j) = θ_2(r) + π/2。"
        "帰納段階では部分木仮説と推移律を組み合わせる。"
        "非一様重みへの摂動論法で結果を拡張。完全な証明は付録Aに記載。□"
    ), italic=True)

    # ============================================================
    # Section 4: Formulation
    # ============================================================
    add_heading(doc, "4. スペクトル因果性: 定式化", level=1)

    add_heading(doc, "4.1 ユーティリティ有向グラフ", level=2)
    add_para(doc, "定義 4.1 (ユーティリティ有向グラフ)", bold=True)
    add_para(doc, (
        "n個の変数 {X_1,...,X_n} に対し、ユーティリティ関数 "
        "U: {1,...,n}^2 → R_≥0 を定義する。U(i,j) は「X_i に関する情報が X_j についての"
        "推論にどれほど有用か」を定量化する。ユーティリティ有向グラフ G_U は:\n"
        "・対称化重み: w(i,j) = (U(i,j) + U(j,i)) / 2\n"
        "・方向符号: σ(i,j) = sign(U(i,j) - U(j,i))"
    ))

    add_heading(doc, "4.2 Directional Predictability Index (DPI)", level=2)
    add_para(doc, (
        "データ駆動成分として |ρ̂_{ij}|（絶対相関）を使用する場合の致命的問題は、"
        "|ρ̂_{ij}| = |ρ̂_{ji}| であるためα=0でユーティリティ行列が対称となり、"
        "外部知識なしに方向推定が不可能となる循環性である。"
    ))

    add_para(doc, "定義 4.2 (Directional Predictability Index)", bold=True)
    add_para(doc, (
        "観測データ X ∈ R^{N×n} に対し、DPI行列 D_DPI ∈ R^{n×n} を定義する:\n"
        "D_DPI(i→j) = |ρ̂_{ij}| · (1 + γ · Ā(i,j))\n"
        "ここで γ > 0 は方向強度パラメータ（γ=1を使用）、Ā(i,j) は3つの正規化非対称統計量の平均:\n"
        "Ā(i,j) = (1/3)[Â_reg(i,j) + Â_ANM(i,j) + Â_ent(i,j)]"
    ))
    add_para(doc, (
        "3成分:\n"
        "(i) 回帰係数非対称性 Â_reg: |β_{j|i}| - |β_{i|j}| の正規化差。\n"
        "(ii) ANM残差独立性 Â_ANM: X_j = βX_i + ε を回帰し、HSIC (Hilbert-Schmidt独立性基準) で"
        "残差の独立性を評価。低HSICは i→j を支持。\n"
        "(iii) 条件付きエントロピー削減 Â_ent: H(X_j|X_i) vs H(X_i|X_j) の正規化差。"
        "低い条件付きエントロピーは因果方向を支持。"
    ))

    add_para(doc, "命題 4.1 (DPIの非対称性)", bold=True)
    add_para(doc, (
        "ANM仮定下で Â_ANM(i,j) ≠ Â_ANM(j,i)。"
        "したがって D_DPI(i→j) ≠ D_DPI(j→i) であり、α=0でも方向推定が可能となる。"
    ))

    add_heading(doc, "4.3 ハイブリッドユーティリティ関数", level=2)
    add_para(doc, (
        "ドメイン知識とデータ駆動成分を混合するハイブリッドユーティリティ関数:\n"
        "U(i→j) = α · C_domain(i,j) + (1-α) · D_DPI(i→j)\n"
        "ここで α ∈ [0,1] はドメイン知識の混合比率、"
        "C_domain(i,j) ∈ [0,1] はドメイン知識行列であり、"
        "変数 X_i が X_j に因果的影響を及ぼすと信じられる度合いを符号化する。"
        "情報源としては専門家評価、文献ベースの評定、または構造化オントロジーが利用可能である。"
        "二値隣接行列と異なり C_domain は段階的信頼度を許容し、"
        "非対称性 C_domain(i,j) ≠ C_domain(j,i) が方向的事前知識を符号化する。"
    ))

    # Figure 2: DPI Architecture
    add_figure(doc,
        f"{FIGURES_DIR}/fig_dpi_architecture.png",
        "図2: Directional Predictability Index (DPI) のアーキテクチャ。"
        "3つの独立な非対称統計量（回帰係数非対称性 Â_reg、ANM残差独立性 Â_ANM、"
        "条件付きエントロピー削減 Â_ent）を正規化・平均し、対称相関 |ρ_{ij}| を"
        "非対称DPIスコアに変換する。モジュラー設計により各成分を独立に拡張可能。",
        width=Inches(5.5)
    )

    add_heading(doc, "4.4 SCC, SCD, 複素因果指標", level=2)
    add_para(doc, (
        "定義 4.3 (Spectral Causal Coupling, SCC): "
        "SCC(i,j) = Σ_k f(λ_k) |u_k(i)| |u_k(j)| cos(θ_k(i) - θ_k(j))\n"
        "定義 4.4 (Spectral Causal Direction, SCD): "
        "SCD(i,j) = Σ_k f(λ_k) |u_k(i)| |u_k(j)| sin(θ_k(i) - θ_k(j))\n"
        "SCC(i,j) = SCC(j,i)（対称）、SCD(i,j) = -SCD(j,i)（反対称）。"
    ))

    add_para(doc, "定理 4.1 (複素因果指標)", bold=True)
    add_para(doc, (
        "CCI(i,j) = Σ_k f(λ_k) |u_k(i)| |u_k(j)| exp(i(θ_k(i) - θ_k(j))) と定義すると、"
        "SCC(i,j) = Re[CCI(i,j)]、SCD(i,j) = Im[CCI(i,j)] が成り立つ。"
    ))

    # Figure 3: CCI Complex Plane
    add_figure(doc,
        f"{FIGURES_DIR}/fig_cci_complex_plane.png",
        "図3: 複素因果指標 CCI(i,j) の複素平面プロット。"
        "実軸 (SCC) は結合強度、虚軸 (SCD) は因果方向を表す。"
        "赤点は順方向 (i→j)、青点は逆方向を示す。"
        "大多数のペアが下半平面に集中し、Ageの上流ルートノードとしての優位性を反映。",
        width=Inches(5.5)
    )

    add_para(doc, "命題 4.2 (SCD行列の性質)", bold=True)
    add_para(doc, (
        "S_{ij} = SCD(i,j) で定義される行列 S は:\n"
        "(i) S = -S^T（反対称）。(ii) tr(S) = 0。(iii) q=0 ならば S = O。"
    ))
    add_para(doc, (
        "性質 (iii) は重要: スペクトル因果性は方向情報 (q > 0) を必要とする。"
        "これはLiNGAMの非ガウス性要件に対応する。"
    ))

    # ============================================================
    # Section 5: Hodge
    # ============================================================
    add_heading(doc, "5. Hodge分解と因果フロー", level=1)

    add_heading(doc, "5.1 グラフ上の微分形式", level=2)
    add_para(doc, "定義 5.1 (コチェイン複体)", bold=True)
    add_para(doc, (
        "有向グラフ G=(V,E) に対し:\n"
        "・C^0 = R^{|V|}（ノード関数）、C^1 = R^{|E|}（辺フロー）\n"
        "・コバウンダリ δ_0: C^0 → C^1: (δ_0 f)(i→j) = f(j) - f(i)（離散勾配）\n"
        "・コバウンダリ δ_1: C^1 → C^2: 三角形上の回転"
    ))

    add_heading(doc, "5.2 Hodge分解定理", level=2)
    add_para(doc, "定理 5.1 (グラフ上のHodge分解; Jiang et al. 2011, Lim 2020)", bold=True)
    add_para(doc, (
        "任意の辺フロー ω ∈ C^1 は以下のように一意に分解される:\n"
        "ω = δ_0 φ（勾配）+ δ_1* ψ（回転）+ h（調和）\n"
        "3成分は互いに直交する。"
    ))

    add_table(doc,
        ["成分", "数学的意味", "因果的解釈"],
        [
            ["δ_0 φ（勾配）", "ポテンシャル駆動フロー", "因果フロー（DAG的、一方向）"],
            ["δ_1* ψ（回転）", "局所的循環", "フィードバックループ（局所的相互作用）"],
            ["h（調和）", "大域的循環", "恒常性調節（全身的）"],
        ],
        caption="表 5.1: Hodge分解の因果的解釈"
    )

    add_heading(doc, "5.3 因果ポテンシャル", level=2)
    add_para(doc, "定義 5.2 (因果ポテンシャル)", bold=True)
    add_para(doc, (
        "因果ポテンシャル φ: V → R は勾配成分 δ_0 φ のポテンシャル関数であり、"
        "グラフ上のPoisson方程式 Lφ = δ_0* ω を解くことで得られる。"
        "L が半正定値であるため、φ は加法定数を除いて一意である。"
    ))

    add_para(doc, "命題 5.1 (因果ポテンシャルはトポロジカルソートを回復する)", bold=True)
    add_para(doc, (
        "ω が純粋なDAGフロー (回転・調和成分がゼロ) を表す場合、"
        "φ によるノードの順序付けはDAGのトポロジカルソートに一致する。\n"
        "証明: ω = δ_0 φ のとき、すべての有向辺 i→j に対し "
        "ω(i→j) = φ(j) - φ(i) > 0 であるから φ(j) > φ(i)。"
        "これはトポロジカル順序の条件そのものである。□"
    ))

    add_para(doc, "定義 5.3 (勾配エネルギー比)", bold=True)
    add_para(doc, (
        "r_gradient = ||δ_0 φ||^2 / ||ω||^2\n"
        "DAG互換成分が捕捉するフローエネルギーの割合を測定する。"
        "r_gradient ≈ 1 はDAG適合を示し、r_gradient << 1 は優勢なフィードバック構造を示す。"
    ))

    # ============================================================
    # Section 6: Related Work
    # ============================================================
    add_heading(doc, "6. 関連研究", level=1)
    add_para(doc, (
        "有向グラフ上の磁気ラプラシアン: Fanuel et al. (2018) はコミュニティ検出に磁気ラプラシアンを使用。"
        "Singer & Wu (2012) は物理的な磁気ラプラシアンとベクトル束の接続を確立。"
        "Bandeira et al. (2014) は角度同期化を一般化した。"
        "これらは因果推論には適用しておらず、本研究はこの接続を確立する最初の研究である。"
    ))
    add_para(doc, (
        "Hodge分解: Jiang et al. (2011) はペアワイズ比較データにHodge分解を適用。"
        "Lim (2020) はグラフ上のHodge理論の包括的レビューを提供。"
        "因果フローのHodge分解を因果推論の文脈に位置づける本研究は新規である。"
    ))
    add_para(doc, (
        "DAG学習: Zheng et al. (2018) のNOTEARSは連続的DAG制約を提案。"
        "LiNGAM (Shimizu et al., 2006) は非ガウス性を利用。"
        "スペクトル因果性はDAG仮定を必要とせず、フィードバック構造を積極的に定量化する。"
    ))
    add_para(doc, (
        "LiNGAMの医療応用: DirectLiNGAM (Shimizu et al., 2011) は逐次非ガウス性検定で因果順序を同定。"
        "Kotoku (2020) は大阪健康診断データ (n≈10^5) に適用。"
        "本ECDパイプラインはLiNGAMの出力をスペクトル因果性の初期化として使用する。"
    ))

    # ============================================================
    # Section 7: Experiments
    # ============================================================
    add_heading(doc, "7. 実験", level=1)

    add_heading(doc, "7.1 データと変数", level=2)
    add_para(doc, (
        "UCI心疾患データセット (Cleveland部分集合; Detrano et al., 1989) から"
        "5つの連続臨床変数を使用する:\n"
        "X = (X_1, X_2, X_3, X_4, X_5) = (年齢, 安静時血圧, コレステロール, 最大心拍数, ST低下)\n"
        "標本サイズ N=297。全変数は平均0、分散1に標準化。"
    ))

    add_para(doc, "ドメイン知識行列", bold=True)
    add_para(doc, (
        "確立された心血管生理学に基づき C_domain を構築する (表7.0)。"
        "各要素 C_domain(i,j) ∈ [0,1] は変数 X_i が X_j に及ぼすと信じられる"
        "因果影響の強度を表し、医学文献 (Detrano et al., 1989) から導出した。"
        "例えば C_domain(年齢, 安静時血圧) = 0.6 は加齢に伴う動脈硬化と"
        "安静時血圧上昇の確立された効果を反映する。"
    ))

    add_table(doc,
        ["C(i,j)", "年齢", "安静時BP", "コレステロール", "最大HR", "ST低下"],
        [
            ["年齢",         "---", "0.6", "0.4", "0.5", "0.3"],
            ["安静時BP",     "0.1", "---", "0.2", "0.3", "0.4"],
            ["コレステロール", "0.1", "0.3", "---", "0.1", "0.3"],
            ["最大HR",       "0.1", "0.2", "0.1", "---", "0.5"],
            ["ST低下",       "0.0", "0.1", "0.0", "0.2", "---"],
        ],
        caption="表 7.0: UCI心疾患変数のドメイン知識行列 C_domain。要素 (i,j) は変数 i の変数 j への因果影響の強度を表す。非対称性 C(i,j) ≠ C(j,i) が方向的事前知識を符号化する。"
    )

    add_heading(doc, "7.2 LiNGAMベースライン", level=2)
    add_para(doc, (
        "DirectLiNGAM (Shimizu et al., 2011) を適用し、推定因果順序:\n"
        "X_1 ≺ X_4 ≺ X_5 ≺ X_2 ≺ X_3 (年齢→最大心拍数→ST低下→安静時血圧→コレステロール)\n"
        "主要因果効果: B_{42} = -0.395 (年齢→最大心拍数)、B_{21} = +0.309 (年齢→安静時血圧)、"
        "B_{54} = -0.348 (最大心拍数→ST低下)。"
    ))

    add_heading(doc, "7.3 磁気ラプラシアンの固有ベクトル", level=2)
    add_para(doc, (
        "式 (4) により α=0.6 で C_domain (表7.0) を60%の重み、データ駆動DPIを40%の重みで"
        "結合してユーティリティ有向グラフを構成し、"
        "q ∈ {0, 0.1, 0.25} で L^(q) を計算する。\n"
        "q=0 では全固有ベクトルは実数値で方向情報は利用不可。"
        "q=0.25 では固有ベクトルが複素数値となり、情報を含む位相角を持つ (表7.1)。"
    ))

    # Table 7.1
    add_table(doc,
        ["変数", "|u_2| (振幅)", "θ_2 (位相角, 度)"],
        [
            ["年齢", "0.53", "0.0°"],
            ["安静時血圧", "0.35", "164.6°"],
            ["コレステロール", "0.42", "-84.3°"],
            ["最大心拍数", "0.47", "34.7°"],
            ["ST低下", "0.44", "-40.6°"],
        ],
        caption="表 7.1: q=0.25 でのFiedler固有ベクトル (u_2) の位相角。因果上流ノード (年齢, 最大心拍数) は正の位相にクラスタ、下流 (コレステロール, ST低下) は負の位相にクラスタ。"
    )

    add_para(doc, (
        "図4は固有ベクトル成分の複素平面分布を示し、q の増加により上流ノードと下流ノードが分離される。"
    ))

    # Figure 2: Magnetic Laplacian
    add_figure(doc,
        f"{FIGURES_DIR}/fig2_magnetic_laplacian_q.png",
        "図4: 磁気ラプラシアンのFiedler固有ベクトルの複素平面プロット (q=0, 0.1, 0.25)。"
        "q=0 では全点が実軸上。q の増加に伴い変数が複素平面に展開し、位相角順序が因果フロー方向を反映。",
        width=Inches(5.0)
    )

    add_heading(doc, "7.4 Hodge分解結果", level=2)
    add_para(doc, (
        "辺フロー ω(i,j) = w(i,j) · σ(i,j) にHodge分解を適用:\n"
        "||δ_0 φ||^2 / ||ω||^2 = 85.9% (勾配 = DAG的因果フロー)\n"
        "||δ_1* ψ||^2 / ||ω||^2 = 14.1% (回転 = フィードバックループ)\n"
        "r_gradient = 85.9% は主にDAG的構造を示す (図5)。"
    ))

    # Figure 3: Hodge decomposition
    add_figure(doc,
        f"{FIGURES_DIR}/fig3_hodge_decomposition.png",
        "図5: Hodge分解結果。(A) フローエネルギーの85.9%が勾配 (DAG) 成分、14.1%が回転 (フィードバック) 成分。"
        "(B) 因果ポテンシャル φ: 年齢が最上流、ST低下が最下流。",
        width=Inches(5.5)
    )

    add_para(doc, (
        "因果ポテンシャル φ は 年齢 > コレステロール > 安静時血圧 ≈ 最大心拍数 > ST低下 "
        "の順に変数を並べ (表7.2)、年齢を最上流、ST低下を最下流とする。臨床的に合理的。"
    ))

    # Table 7.2
    add_table(doc,
        ["順位", "変数", "φ"],
        [
            ["1", "年齢", "0.000"],
            ["2", "コレステロール", "-0.168"],
            ["3", "安静時血圧", "-0.204"],
            ["4", "最大心拍数", "-0.204"],
            ["5", "ST低下", "-0.324"],
        ],
        caption="表 7.2: Hodge分解による因果ポテンシャル φ (α=0.6)"
    )

    add_heading(doc, "7.5 手法比較", level=2)
    add_para(doc, (
        "全 C(5,2)=10 変数ペアについて3条件—DirectLiNGAM (A)、ドメイン知識付きスペクトル因果性 "
        "(B, α=0.6)、DPI単独スペクトル因果性 (C, α=0)—の因果方向を比較すると、"
        "一致と情報を含む不一致の両方が明らかになる。"
    ))

    add_para(doc, (
        "手法間の不一致は情報的である: スペクトル因果性 (BまたはC) がDirectLiNGAM (A) "
        "と逆方向を示す場合、その変数ペアは「診断マーカー」関係を持つことが多い。"
        "これは「介入的因果性」(レベル2) と「情報的因果性」(レベル1.5) の区別を反映する。"
    ))

    add_heading(doc, "7.6 三条件構造比較", level=2)

    add_table(doc,
        ["条件", "手法", "グラフ型", "ドメイン知識"],
        [
            ["(A)", "DirectLiNGAM", "DAG", "なし"],
            ["(B)", "スペクトル (α=0.6)", "DCG", "臨床60% + データ40%"],
            ["(C)", "スペクトル (α=0, DPI)", "DCG", "なし (純データ駆動)"],
        ],
        caption="表 7.3: 三条件比較"
    )

    add_para(doc, (
        "条件 (C) はDPIによりα=0で9本の有向辺を検出し、r_gradient = 0.581、"
        "LiNGAMとの方向一致率67%。α=0からα=0.6への遷移で r_gradient は0.581から0.824に滑らかに改善。"
    ))

    # ============================================================
    # Section 8: Phase Transition
    # ============================================================
    add_heading(doc, "8. 因果構造の相転移", level=1)

    add_heading(doc, "8.1 勾配エネルギー比のスケール不変性", level=2)

    add_para(doc, "定理 8.1 (スケール不変性)", bold=True)
    add_para(doc, (
        "U_α = α·C + (1-α)·S_data とし、S_data は対称行列 (例: |ρ̂_{ij}|)、"
        "C は非対称ドメイン知識行列とする。U_α の反対称成分は "
        "A(α) = U_α - U_α^T = α·(C - C^T) であり、S_data - S_data^T = 0 であるため:\n"
        "r_gradient(α) = r_gradient(1)  (すべての α > 0 に対して)"
    ))
    add_para(doc, (
        "証明: Hodge分解は線形: ω_α = α·ω_1 ならば φ_α = α·φ_1 であり、"
        "r_gradient(α) = ||δ_0(αφ_1)||^2 / ||αω_1||^2 = α^2||δ_0 φ_1||^2 / α^2||ω_1||^2 = r_gradient(1)。□"
    ), italic=True)

    add_para(doc, "系 8.1", bold=True)
    add_para(doc, (
        "対称データ駆動成分では、α はスイッチ (オン/オフ) として機能し、ダイヤルではない: "
        "α = 10^{-6} と α = 1 は同一の r_gradient を与える。"
    ))
    add_para(doc, (
        "証明: データ駆動成分が D_data = D_data^T（対称）を満たす場合、"
        "ユーティリティ関数の反対称部分は A(α) = α·(C_domain - C_domain^T)/2 "
        "（任意の α>0）。定理 8.1 により r_gradient は A の符号構造のみに依存し、"
        "これは全ての α>0 で同一である。□"
    ), italic=True)

    add_para(doc, "注意 8.1", bold=True, italic=True)
    add_para(doc, (
        "DPIをデータ駆動成分とする場合、α=0 での反対称成分が非ゼロとなり "
        "r_gradient(0) = 0.581 > 0。α=0 から α=1 への遷移は滑らかで、"
        "旧モデルの一次相転移に替わる二次相転移となる (図6)。"
    ))

    # Figure 6: Alpha sweep
    add_figure(doc,
        f"{FIGURES_DIR}/fig8_alpha_sweep.png",
        "図6: DPIによるα掃引解析。(A) r_gradient は0.581 (α=0) から0.859 (α=1) に滑らかに増加。"
        "(B) 検出辺数とLiNGAM一致率。(C) 非対称ノルム。(D) 相図。",
        width=Inches(5.5)
    )

    add_heading(doc, "8.2 知識品質 vs 知識量", level=2)

    add_para(doc, "定理 8.2 (U字型品質カーブ)", bold=True)
    add_para(doc, (
        "DAG G=(V,E) の辺方向を C_true に符号化し、各辺方向を独立に確率 p_flip で反転させた "
        "知識行列 C_p を定義する。このとき:"
        "(i) E[r_gradient(0)] = E[r_gradient(1)] = r_gradient* (真のDAGの勾配比)。"
        "(ii) r_gradient(p) は p_min ∈ (0, 0.5) で最小値を取る。"
        "(iii) p → E[r_gradient(p)] は p_flip = 0.5 に関して対称。"
    ))
    add_para(doc, (
        "証明: (i) p_flip=1 では全辺が反転: σ_{ij} → -σ_{ij}。辺フローは ω_1 = -ω_0 となる。"
        "Hodge分解の線形性から δ_0 φ_1 = -δ_0 φ_0 であり "
        "r_gradient(1) = ||δ_0 φ_0||^2 / ||ω_0||^2 = r_gradient(0)。"
        "(iii) 確率 p で C_true を反転した分布は、確率 1-p で C_reversed を反転した分布と一致。"
        "r_gradient はωの全体符号反転に不変なので E[r_gradient(p)] = E[r_gradient(1-p)]。"
        "(ii) ω_p = Σ_e s_e · ω_e (s_e ∈ {±1}, Pr(s_e=-1)=p) と書くと、"
        "E[||δ_0 φ_p||^2] = (1-2p)^2 ||δ_0 φ_0||^2。一方 E[||ω_p||^2] は p に依存しない (s_e^2=1)。"
        "よって E[r_gradient(p)] = (1-2p)^2 · r_gradient*。"
        "これは p=0.5 で最小値0を取る下向き放物線。実験上の最小値が p_min ≈ 0.3 < 0.5 "
        "となるのは、α<1 でのDPIデータ駆動成分が非反転非対称寄与を加えるため。□"
    ), italic=True)

    add_para(doc, (
        "DAG維持の臨界閾値 (r_gradient > 0.5) は p_flip* ≈ 0.15: "
        "辺方向の少なくとも85%が正しい必要がある (図7)。"
    ))

    # Figure 7: p_flip U-curve
    add_figure(doc,
        f"{FIGURES_DIR}/fig_pflip_ucurve.png",
        "図7: 知識品質の相転移（定理8.2）。"
        "辺方向反転率 p_flip の増加に伴い、勾配エネルギー比 r_gradient はU字型カーブを描く。"
        "理論予測 (1-2p)^2 r*（赤破線）は完全対称な反転を仮定。"
        "実測値（青、200試行、α=0.6）はDPIデータ駆動成分（40%重み）が不変のため乖離し、"
        "p=0.5でr_gradientはゼロに到達しない。"
        "臨界閾値 p*_flip ≈ 0.15 以下でDAG構造が維持される。",
        width=Inches(5.5)
    )

    add_para(doc, "注意 8.2 (生兵法は大怪我のもと)", bold=True, italic=True)
    add_para(doc, (
        "知識ベース因果推論に対する一般的警告: 少数の誤った方向性主張が"
        "推定因果構造を壊滅的に歪曲しうる。"
    ))

    add_heading(doc, "8.3 スケルトン辺: ルートノード優位性", level=2)

    add_table(doc,
        ["除去辺", "Δr_gradient", "重要度"],
        [
            ["年齢 ↔ ST低下", "-0.267", "致命的"],
            ["年齢 ↔ 最大心拍数", "-0.098", "高"],
            ["年齢 ↔ コレステロール", "-0.069", "高"],
            ["安静時血圧 ↔ 最大心拍数", "+0.015", "無視可能"],
        ],
        caption="表 8.1: 辺除去重要度分析"
    )

    add_para(doc, (
        "最大レバレッジのための最小知識は、1つの外生 (ルート) 変数の同定である。"
    ))

    # ============================================================
    # Section 9: ECD
    # ============================================================
    add_heading(doc, "9. Ensemble Causal Direction (ECD)", level=1)

    add_heading(doc, "9.1 ECDパイプライン", level=2)
    add_para(doc, (
        "ECDパイプラインはLiNGAMの出力をスペクトル因果性のドメイン知識として使用する:\n"
        "U_ECD(i→j) = α · C_LiNGAM(i,j) + (1-α) · |ρ̂_{ij}|\n"
        "ここで C_LiNGAM(i,j) = |B_{ji}|。α=0.3 でHodgeポテンシャル順序は "
        "年齢 > 最大心拍数 > ST低下 > コレステロール > 安静時血圧 となり、"
        "LiNGAMの因果順序に近似する (表9.1)。"
    ))

    add_table(doc,
        ["", "臨床知識 (α=0.6)", "ECD (α=0.3)", "LiNGAM"],
        [
            ["r_gradient", "0.859", "0.555", "---"],
            ["辺数", "9", "6", "6"],
            ["順序", "年齢 > コレステロール > 血圧 ≈ 心拍数 > ST", "年齢 > 心拍数 > ST > コレステロール > 血圧", "年齢 > 心拍数 > ST > 血圧 > コレステロール"],
        ],
        caption="表 9.1: 臨床知識、ECD (LiNGAM初期化)、LiNGAM単独の因果順序比較"
    )

    add_heading(doc, "9.2 因果ポテンシャルと介入可能性", level=2)
    add_para(doc, (
        "因果ポテンシャル φ は臨床的介入可能性 ι との著しい対応を示す:"
    ))

    add_table(doc,
        ["変数", "φ", "ι", "臨床的理由"],
        [
            ["年齢", "0.000", "0 (不可能)", "不可逆的生物学的過程"],
            ["最大心拍数", "-0.204", "≈0.3 (困難)", "加齢・体質依存"],
            ["ST低下", "-0.324", "≈0.5 (間接的)", "PCI/CABGで改善可能"],
            ["コレステロール", "-0.168", "≈0.9 (容易)", "スタチン"],
            ["安静時血圧", "-0.204", "≈0.8 (容易)", "降圧薬"],
        ],
        caption="表 9.2: 因果ポテンシャル φ と介入可能性 ι の対応"
    )

    add_para(doc, "命題 9.1 (ポテンシャル-介入可能性対応)", bold=True)
    add_para(doc, (
        "DAG G=(V,E) と構造方程式モデル X_j = f_j(pa(j), ε_j) に対し、"
        "介入可能性を ι(j) = 1 - Var(ε_j)/Var(X_j) ∈ [0,1] と定義する。"
        "任意のルートノード r (pa(r) = ∅) と非ルートノード j に対し:"
        "(i) φ(r) ≥ φ(j)、(ii) ι(r) = 0 ≤ ι(j)。"
        "したがって φ と ι は極値で逆順序を持つ。"
    ))
    add_para(doc, (
        "証明: (i) 命題 5.1 より φ はDAG上のトポロジカル順序を誘導し、"
        "ルートノードは最大位置を占める。"
        "(ii) ルートノード r では X_r = f_r(ε_r) であり、ι(r) = 0。"
        "非ルートノード j では pa(j) ≠ ∅ であり、faithfulness条件下で "
        "Var(X_j) > Var(ε_j) となるため ι(j) > 0。"
        "DAG内で深いノードほど低い φ と高い ι を持つ。□"
    ), italic=True)

    add_heading(doc, "9.3 Hillの9基準カバレッジ", level=2)
    add_para(doc, (
        "単一の計算手法ではHillの因果判定9基準すべてをカバーできない (Hill, 1965)。"
        "ECDアンサンブルは大幅に広いカバレッジを達成する。"
        "具体的にはLiNGAMはH1 (強さ) とH3 (特異性) をカバーするがH6/H7/H9は欠如。"
        "スペクトル因果性はユーティリティ関数を通じてH6 (妥当性), H7 (整合性), H9 (類似性) をカバーし、"
        "ECDアンサンブルは両方を統合しほぼ完全なカバレッジを達成する。"
    ))

    add_heading(doc, "9.4 フィードバック解析とプルーニング", level=2)

    add_table(doc,
        ["辺", "勾配方向", "フィードバック率", "解釈"],
        [
            ["年齢→安静時血圧", "年齢→安静時血圧", "0%", "純一方向因果"],
            ["年齢→コレステロール", "年齢→コレステロール", "1%", "純一方向因果"],
            ["最大心拍数↔ST低下", "最大心拍数→ST低下", "73%", "強い運動-虚血ループ"],
        ],
        caption="表 9.3: 辺ごとのフィードバック率"
    )

    add_para(doc, (
        "最大心拍数↔ST低下の73%フィードバックは、LiNGAMのDAG仮定 "
        "(最大心拍数→ST低下のみ) が臨床的に重要なフィードバックループ"
        "—運動不耐症→虚血→心筋酸素需要増大→さらなる運動不耐症—を見落としていることを示す。"
    ))

    # ============================================================
    # Section 10: Identifiability
    # ============================================================
    add_heading(doc, "10. 識別可能性", level=1)

    add_heading(doc, "10.1 循環性の問題", level=2)
    add_para(doc, (
        "|ρ̂_{ij}| を用いた元の定式化は循環性を抱えていた: "
        "SCDが真の因果方向に一致する条件が、ユーティリティの非対称性に"
        "既に正しい方向が符号化されていることを要求した。"
        "|ρ̂_{ij}| が対称であるため、外部知識 (α > 0) でのみ達成可能であった。"
    ))

    add_heading(doc, "10.2 DPIによる循環性の解消", level=2)
    add_para(doc, "定理 10.1 (DPIの部分的識別可能性)", bold=True)
    add_para(doc, (
        "データ生成過程がANMに従うとする: X_j = f(X_i) + ε, ε ⊥ X_i。"
        "このとき、DPIのANM成分は Â_ANM(i,j) > 0 (正しく i→j を識別) を "
        "N→∞ で確率1に収束して満たす。"
    ))
    add_para(doc, (
        "証明: ANM下で、順方向残差 ε̂_{j|i} = X_j - f̂(X_i) は真のノイズ ε に収束し、"
        "仮定より X_i と独立。したがって HSIC(ε̂_{j|i}, X_i) → 0。"
        "逆方向では ε̂_{i|j} は X_j と独立な量に収束しない (Hoyer et al., 2009 の識別可能性結果)。"
        "よって HSIC(ε̂_{i|j}, X_j) > c > 0。正規化差は極限で正。□"
    ), italic=True)

    add_para(doc, "系 10.1", bold=True)
    add_para(doc, (
        "ANM仮定下、α=0 (ドメイン知識なし) で D_DPI(i→j) ≠ D_DPI(j→i) が"
        "正しい方向の符号で成立し、非循環的データ駆動基盤を提供する。"
    ))
    add_para(doc, (
        "証明: α=0 ではユーティリティ関数は U(i,j) = D_DPI(i→j) に帰着。"
        "定理 10.1 により、各DPI成分はANM仮定下で識別可能であり (命題 4.1)、"
        "Ā(i,j) ≠ 0 が真の因果辺に対し正しい符号で成立する。"
        "したがって D_DPI(i→j) ≠ D_DPI(j→i) であり、"
        "ドメイン知識なしに非対称ユーティリティグラフが定義可能。□"
    ), italic=True)

    add_heading(doc, "10.3 識別可能性ロードマップ", level=2)
    add_para(doc, (
        "Phase 1 (成分レベル, 達成済): 各DPI成分はそれぞれの仮定下で確立された識別可能性保証を持つ。\n"
        "Phase 2 (スペクトル伝播, 達成可能): DPIがルートノード方向を正しく同定した場合、"
        "Hodge分解がPoisson方程式を通じてこの情報を大域的に伝播する。\n"
        "Phase 3 (完全識別可能性, 困難だが実用的に不要): 全辺の同時正確方向推定の証明は"
        "スペクトル同値性とHodge回転成分の巡回許容性により困難。"
        "しかし実用的に不要: ECDパイプラインは核心辺にLiNGAMの識別可能性を「借用」し、"
        "フィードバック定量化をスペクトル因果性に委譲する。"
    ))

    # ============================================================
    # Section 11: Discussion
    # ============================================================
    add_heading(doc, "11. 議論と将来の方向性", level=1)

    add_heading(doc, "11.1 情報的因果性 vs 介入的因果性", level=2)
    add_para(doc, (
        "スペクトル因果性はPearlの因果の梯子における「レベル1.5」で動作する"
        "—連関 (レベル1) より深いが介入的因果性 (レベル2) ではない。"
        "LiNGAMとの不一致はこの区別を反映: スペクトル因果性は「情報的方向」"
        "(「Xを知ることでYについて情報を得る」) を捕捉し、"
        "これは診断マーカーペアでは「介入的方向」(「Xを操作するとYが変化する」) と逆転しうる。"
    ))

    add_heading(doc, "11.2 理論的含意", level=2)
    add_para(doc, (
        "スケール不変性定理 (定理8.1) は実用的に重要な帰結を持つ: 非対称成分の符号構造が"
        "固定されれば、ドメイン知識行列 C_domain の正確な数値は r_gradient に無関係である。"
        "すなわち大まかな順序判断（「年齢は安静時血圧に先行する」）が正確な定量的推定と"
        "同等の構造情報を持つ。実務者にとって、方向が正しい限り粗い専門家知識でも"
        "有意味な因果構造回復が可能であることを意味する。"
    ))
    add_para(doc, (
        "U字型品質曲線 (定理8.2) は重要だが見落とされがちなリスクを定式化する: "
        "部分的に誤ったドメイン知識の導入はドメイン知識を全く使わないよりも悪い結果を"
        "もたらしうる。臨界閾値 p*_flip ≈ 0.15 は、主張された辺方向の約15%超が誤りの場合、"
        "純粋にデータ駆動のベースライン (α=0) 以下に因果構造が劣化することを意味する。"
        "知識ベース因果発見パイプラインへの直接的含意: 専門家知識導入前に方向性主張の"
        "信頼性を推定し、信頼度が低い場合は α=0 (DPIのみ) の使用を検討すべきである。"
    ))
    add_para(doc, (
        "位相-方向対応 (定理3.1) はスペクトルグラフ理論と因果順序付けの深い接続を確立する: "
        "磁気ラプラシアンの複素固有ベクトルが位相角を通じてトポロジカルソートを符号化する。"
        "この幾何学的視点はLiNGAMの代数的識別可能性結果を補完し、非ガウス性仮定の検証が"
        "困難な設定でスペクトル手法が代替的識別可能性条件を提供しうることを示唆する。"
    ))

    add_heading(doc, "11.3 実験結果の解釈", level=2)
    add_para(doc, (
        "DirectLiNGAMとスペクトル因果性の実験的比較は不一致に体系的パターンを明らかにする。"
        "10変数ペアのうち不一致は2カテゴリにクラスタする。"
    ))
    add_para(doc, (
        "診断マーカー逆転: (コレステロール, ST低下) や (安静時血圧, 最大心拍数) 等のペアで"
        "スペクトル因果性はLiNGAM方向を逆転させる。これらは下流変数（結果）が臨床で"
        "上流変数（原因）の診断指標として日常的に使用されるペアに正確に対応する。"
        "スペクトル枠組みの「情報的方向」は臨床推論の予測的構造を捕捉し、LiNGAMが回復する"
        "介入的構造と相補的（矛盾ではない）であることを示唆する。"
    ))
    add_para(doc, (
        "フィードバック検出: MaxHR↔STDep辺の73%カール成分は臨床的に有意である: "
        "運動不耐性→虚血→心筋酸素需要増加→さらなる運動制限という病的フィードバックループを"
        "構成する。DAGベース手法は単方向辺 MaxHR→STDep を強制し、この双方向構造を失う。"
        "Hodge分解の辺ごとフィードバック比率定量化能力は、DAG仮定が妥当な辺 (<10%カール) と"
        "フィードバックモデリングが不可欠な辺を判定する原理的基準を提供する。"
    ))
    add_para(doc, (
        "α掃引の滑らかな遷移 (r_gradient: α=0で0.581→α=1で0.859) は、ドメイン知識と"
        "データ駆動推論が冗長ではなく相補的な情報を提供することを実証する。"
        "DPI単独 (α=0) で既にLiNGAMとの67%方向一致を達成する事実はDPIの有意性を検証し、"
        "ドメイン知識からの追加20+ポイントは信頼できる専門家入力の価値を強調する。"
    ))

    add_heading(doc, "11.4 仮定の検討", level=2)
    add_para(doc, (
        "加法ノイズモデル (ANM): 部分的識別可能性定理 (定理10.1) はANMデータ生成過程を仮定する。"
        "これはLiNGAMの線形性＋非ガウス性の同時要件より弱いが、乗法ノイズ、"
        "分散不均一モデル、事後非線形モデルを除外する。DPIのアンサンブル設計"
        "（3つの非対称成分の平均）はANM設定を超える経験的頑健性を提供するが、"
        "より広いモデルクラスへの形式的保証はオープン問題である。"
    ))
    add_para(doc, (
        "tree DAG構造: 位相-方向対応の証明 (付録A) は正の辺重みを持つtree DAGを仮定する。"
        "一般DAGへの拡張は複数の親を持つノードの処理を要し、固有ベクトル位相は単調関数でなく"
        "親位相の加重平均となる。予備的数値実験では有界入次数の疎DAGで近似的に成立するが、"
        "形式的証明は未達成。"
    ))
    add_para(doc, (
        "固定チャージパラメータ q: スペクトル解析はチャージパラメータ q に依存し、"
        "対称結合情報 (q=0) と方向位相情報 (q>0) のトレードオフを制御する。"
        "原理的な q 選択手続き—カーネル法のバンド幅選択に類似—は枠組みの実用性を強化する。"
        "外部基準 (LiNGAM一致率やブートストラップ安定性) に対する交差検証が有望な方向である。"
    ))
    add_para(doc, (
        "ドメイン知識品質: 本枠組みは非負で因果影響強度に大まかに校正された"
        "ドメイン知識行列 C_domain へのアクセスを仮定する。定理8.1により符号構造のみが"
        "r_gradient に重要であるが、SCCとSCDの値は大きさに依存する。"
        "C_domain の摂動に対するスペクトル指標の感度分析は信頼区間を提供し、自然な拡張である。"
    ))

    add_heading(doc, "11.5 計算量", level=2)
    add_para(doc, (
        "磁気ラプラシアン固有分解は O(n^3)。大規模 n ではランダム化SVDにより O(nk^2)。"
        "Hodge分解は疎グラフで O(|E|)。DPI計算は O(n^2·N)。"
    ))

    add_heading(doc, "11.6 限界", level=2)
    add_para(doc, (
        "(1) UCI心疾患実験は変数5個のみ; MIMIC-IV等のより大規模データセットでの検証が必要。\n"
        "(2) チャージパラメータ q が結果に大きく影響; 原則的な選択基準の開発が必要。\n"
        "(3) 完全識別可能性 (Phase 3) はオープン問題。"
    ))

    add_heading(doc, "11.7 将来の方向性", level=2)
    add_para(doc, (
        "(1) tree DAG + 線形SEMに対するPhase 2識別可能性の証明。\n"
        "(2) MIMIC-IVおよび大規模コホートデータでのECD検証。\n"
        "(3) 時間拡張: 時間遅延ユーティリティグラフと固有値軌跡抽出。\n"
        "(4) LiNGAM一致率からの p_flip の自動推定。\n"
        "(5) ブートストラップ検定によるデータ適応型 α。"
    ))

    # ============================================================
    # References (abbreviated)
    # ============================================================
    add_heading(doc, "参考文献", level=1)
    refs = [
        "Bandeira, A. S., Singer, A., & Spielman, D. A. (2014). A Cheeger inequality for the graph connection Laplacian. SIAM J. Matrix Analysis, 35, 2009-2053.",
        "Detrano, R. et al. (1989). International application of a new probability algorithm for the diagnosis of coronary artery disease. Am. J. Cardiol., 64, 304-310.",
        "Fanuel, M. et al. (2018). Magnetic eigenmaps for the visualization of directed networks. Applied and Computational Harmonic Analysis, 44, 189-199.",
        "Granger, C. W. J. (1969). Investigating causal relations by econometric models and cross-spectral methods. Econometrica, 37, 424-438.",
        "Hill, A. B. (1965). The environment and disease: Association or causation? Proc. Royal Soc. Med., 58, 295-300.",
        "Hoyer, P. et al. (2009). Nonlinear causal discovery with additive noise models. NeurIPS.",
        "Janzing, D. et al. (2010). Causal inference using the algorithmic Markov condition. IEEE Trans. IT, 56, 5168-5194.",
        "Jiang, X. et al. (2011). Statistical ranking and combinatorial Hodge theory. Mathematical Programming, 127, 203-244.",
        "Kotoku, J. (2020). Causal relations of health indices inferred by LiNGAM. Osaka health checkup data.",
        "Le, T. et al. (2024). Multi-agent causal discovery (MAC). AAAI.",
        "Lim, L.-H. (2020). Hodge Laplacians on graphs. SIAM Review, 62, 685-715.",
        "Pearl, J. (2009). Causality (2nd ed.). Cambridge University Press.",
        "Peters, J. et al. (2014). Causal discovery with continuous additive noise models. JMLR, 15, 2009-2053.",
        "Rubin, D. B. (1974). Estimating causal effects of treatments in randomized and nonrandomized studies. J. Educ. Psych., 66, 688-701.",
        "Shimizu, S. et al. (2006). A linear non-Gaussian acyclic model for causal discovery. JMLR, 7, 2003-2030.",
        "Shimizu, S. et al. (2011). DirectLiNGAM: A direct method for learning a linear non-Gaussian structural equation model. JMLR, 12, 1225-1248.",
        "Shuman, D. I. et al. (2013). The emerging field of signal processing on graphs. IEEE Signal Processing Magazine, 30, 83-98.",
        "Singer, A. & Wu, H.-T. (2012). Vector diffusion maps and the connection Laplacian. Comm. Pure Appl. Math., 65, 1067-1144.",
        "Zheng, X. et al. (2018). DAGs with NO TEARS: Continuous optimization for structure learning. NeurIPS.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ref)
        run.font.size = Pt(9)

    # ============================================================
    # Appendix A
    # ============================================================
    add_heading(doc, "付録A: 位相-方向対応の証明 (定理 3.1)", level=1)

    add_para(doc, (
        "定理 3.1 を2段階で証明する: まずパスグラフに対して (補題 A.1)、"
        "次にtree DAGに対して構造的帰納法により拡張する (定理 A.1)。"
    ))

    add_para(doc, "仮定 A.1", bold=True)
    add_para(doc, (
        "G=(V,E) は n ノードの根付きtree DAGで、構造方程式モデルは "
        "X_j = Σ_{i∈pa(j)} β_{ij} X_i + ε_j (β_{ij} > 0)。"
        "ユーティリティグラフは辺重み w_{ij} > 0 と方向符号 σ_{ij} = +1 を割り当てる。"
    ))

    add_heading(doc, "A.1 パスグラフ", level=2)
    add_para(doc, "補題 A.1 (パスグラフ上の位相単調性)", bold=True)
    add_para(doc, (
        "有向パスグラフ P = v_1→v_2→...→v_n (一様重み w > 0, q=0.25) に対し、"
        "Fiedler固有ベクトル u_2 の位相角は θ_2(v_1) < θ_2(v_2) < ... < θ_2(v_n) を満たす。"
    ))
    add_para(doc, (
        "証明: q=0.25 でエルミート隣接行列は H^(q)_{v_k,v_{k+1}} = iw, "
        "H^(q)_{v_{k+1},v_k} = -iw を満たす。内部ノード v_k の次数は d_{v_k} = 2w。"
        "正規化磁気ラプラシアンの固有方程式から:"
    ))
    add_para(doc, (
        "(1-λ)u(v_k) = (-iw/d_{v_k})·u(v_{k-1}) + (iw/d_{v_k})·u(v_{k+1})    ...(A.1)"
    ), italic=True)
    add_para(doc, (
        "u(v_k) = r_k exp(iθ_k) と極座標表示し (A.1) に代入すると、実部から:"
    ))
    add_para(doc, (
        "(1-λ_2)r_k = (w/d_{v_k})[r_{k+1} sin(θ_{k+1}-θ_k) + r_{k-1} sin(θ_k-θ_{k-1})]"
    ), italic=True)
    add_para(doc, (
        "1-λ_2 > 0 かつ r_k > 0 であるため右辺は正でなければならず、"
        "θ_{k+1} - θ_k > 0 (各ステップで正の位相増分) が要求される。"
        "一様重みでは対称性により一定の位相増分が得られる:"
    ))
    add_para(doc, (
        "Δθ = θ_{k+1} - θ_k = arctan(2πq·w̃/(1+w̃²)) > 0"
    ), italic=True)
    add_para(doc, (
        "ここで w̃ = w/d_{v_k} は正規化重み。各 Δθ > 0 より "
        "θ_2(v_1) < θ_2(v_2) < ... < θ_2(v_n)。□"
    ))

    add_heading(doc, "A.2 Tree DAGへの拡張", level=2)
    add_para(doc, "定理 A.1 (tree DAG上の位相-方向対応)", bold=True)
    add_para(doc, (
        "仮定 A.1 の下、q=0.25 で L^(q) のFiedler固有ベクトル u_2 に対し、"
        "ノード i がノード j の祖先ならば θ_2(i) < θ_2(j)。"
    ))
    add_para(doc, (
        "証明: 木の深さに関する構造的帰納法による。"
    ))
    add_para(doc, (
        "基底: 深さ1の木は星グラフ r→{c_1,...,c_m}。"
        "葉 c_j (d_{c_j}=w) の固有方程式から "
        "u(c_j) = (r_0/(√m(1-λ))) · exp(i(θ_0+π/2)) を得る。"
        "よって θ_2(c_j) = θ_0 + π/2 > θ_0 = θ_2(r) であり全子に対して成立。"
    ))
    add_para(doc, (
        "帰納段階: 深さ ≤ d の全部分木で結果が成立すると仮定。"
        "深さ d+1 の根 r について、基底と同じ計算で θ_2(r) < θ_2(c_j) (全子 c_j)。"
        "各部分木 T_j 内では帰納仮定により θ_2(c_j) < θ_2(v) (c_j の全子孫 v)。"
        "推移律より θ_2(r) < θ_2(c_j) < θ_2(v)。"
    ))
    add_para(doc, (
        "非一様重みの摂動論法: 辺重み w_{ij} が異なる場合も、"
        "正規化重み比 w̃_{ij} = w_{ij}/√(d_i d_j) ∈ (0,1] であるため "
        "Δθ_{ij} ∈ (0, π/2] が保証され、根から葉への全経路で位相の厳密な単調性が保存される。□"
    ))

    add_para(doc, "注意 A.1", bold=True, italic=True)
    add_para(doc, (
        "tree DAGへの制限は本質的である: V構造を持つ一般DAGでは、"
        "異なる経路からの位相寄与が干渉し単調性が弱まりうる。"
        "経験的には密なグラフでも対応は強いが (第7節)、"
        "木を超えた形式的証明の拡張はオープン問題として残る。"
    ))

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = "spectral_causality_ja.docx"
    doc.save(output_path)
    print(f"Saved {output_path}")
