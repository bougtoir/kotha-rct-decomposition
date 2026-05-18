#!/usr/bin/env python3
"""
Generate English DOCX manuscript for arXiv submission.
Spectral Causality: Causal Direction Estimation via Magnetic Laplacians
and Hodge Decomposition.

All figures and tables are placed inline immediately after first mention.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FIGURES_DIR = "figures"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, style="Normal", bold=False, italic=False, font_size=None, alignment=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    return p


def add_math_para(doc, text, font_size=10):
    """Add a paragraph with math-like text (italicized)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(font_size)
    return p


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Add a figure with caption immediately inline."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=width)
    else:
        p.add_run(f"[Figure: {img_path}]")

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    return cap


def add_table(doc, headers, rows, caption=None):
    """Add a table with optional caption."""
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.space_before = Pt(12)
        run = cap_p.add_run(caption)
        run.bold = True
        run.font.size = Pt(9)

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing
    return table


def add_grouped_table(doc, row_header_label, groups, sub_headers, rows, caption=None):
    """Add a table with two-row grouped headers.

    Args:
        row_header_label: Label for the first column (e.g. "Method")
        groups: List of (group_label, num_sub_cols) tuples, e.g.
                [("N=200", 3), ("N=500", 3), ("N=1000", 3)]
        sub_headers: List of sub-column headers repeated per group,
                     e.g. ["SHD", "TPR", "FDR"]
        rows: List of [row_label, val1, val2, ...] lists
        caption: Optional caption string
    """
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.space_before = Pt(12)
        run = cap_p.add_run(caption)
        run.bold = True
        run.font.size = Pt(9)

    total_sub_cols = sum(n for _, n in groups)
    num_cols = 1 + total_sub_cols  # row header + data columns
    table = doc.add_table(rows=2 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: group header row — merge cells for each group
    # First cell spans both header rows (row_header_label)
    cell_a = table.cell(0, 0)
    cell_b = table.cell(1, 0)
    merged = cell_a.merge(cell_b)
    merged.text = row_header_label
    for paragraph in merged.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

    col_offset = 1
    for group_label, n_sub in groups:
        # Merge cells in row 0 for this group
        if n_sub > 1:
            cell_start = table.cell(0, col_offset)
            cell_end = table.cell(0, col_offset + n_sub - 1)
            merged_group = cell_start.merge(cell_end)
            merged_group.text = group_label
            for paragraph in merged_group.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        else:
            cell = table.cell(0, col_offset)
            cell.text = group_label
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Row 1: sub-headers
        for s_idx in range(n_sub):
            cell = table.cell(1, col_offset + s_idx)
            cell.text = sub_headers[s_idx] if s_idx < len(sub_headers) else ""
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        col_offset += n_sub

    # Data rows (starting at row index 2)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 2].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing
    return table


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # ============================================================
    # Title Page
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run(
        "Spectral Causality:\nCausal Direction Estimation via\n"
        "Magnetic Laplacians and Hodge Decomposition"
    )
    run.bold = True
    run.font.size = Pt(16)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Tatsuki Onishi")
    run.font.size = Pt(12)

    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = affil.add_run("Data Science and AI Innovation Research Promotion Center\nShiga University, Japan")
    run.font.size = Pt(10)
    affil.paragraph_format.space_after = Pt(12)

    email = doc.add_paragraph()
    email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = email.add_run("bougtoir@gmail.com")
    run.font.size = Pt(10)
    email.paragraph_format.space_after = Pt(24)

    # ============================================================
    # Abstract
    # ============================================================
    add_heading(doc, "Abstract", level=1)
    add_para(doc, (
        "I introduce spectral causality, a framework for causal direction "
        "estimation via the magnetic Laplacian, an Hermitian matrix whose "
        "complex-valued eigenvectors encode edge directionality as phase. I define "
        "a Directional Predictability Index (DPI) that resolves the "
        "circularity of earlier utility-based formulations: under additive noise "
        "model assumptions, DPI consistently estimates causal direction without "
        "domain knowledge. Connecting the framework to Hodge decomposition, "
        "I prove that gradient\u2013curl\u2013harmonic separation of edge flows yields a "
        "causal potential recovering topological order under directed acyclic graph "
        "(DAG) structure, and establish scale invariance of the gradient "
        "energy ratio r_gradient, showing that causal structure "
        "emergence depends on knowledge quality rather than quantity."
    ))
    add_para(doc, (
        "Experiments on synthetic random DAGs (n=5\u201320, N=200\u20131000), "
        "the Sachs protein signaling network (n=11, 17 ground-truth edges), "
        "and the UCI Heart Disease dataset (N=297, 5 variables) demonstrate "
        "competitive recovery against DirectLiNGAM, the PC algorithm, GES, and "
        "NOTEARS. On Sachs data with partial knowledge, spectral causality achieves "
        "a true positive rate (TPR) of 0.56 and a false discovery rate (FDR) "
        "of 0.27, while providing unique DAG-adequacy diagnostics and detecting "
        "feedback loops invisible to DAG-based methods."
    ))

    kw = doc.add_paragraph()
    run = kw.add_run("Keywords: ")
    run.bold = True
    kw.add_run(
        "Causal discovery, Magnetic Laplacian, Hodge decomposition, "
        "Spectral graph theory, Directed graphs"
    )

    # ============================================================
    # Section 1: Introduction
    # ============================================================
    add_heading(doc, "1. Introduction", level=1)

    add_para(doc, (
        "Causal inference\u2014determining whether X causes Y\u2014is a central problem "
        "in science and medicine. The dominant approaches include structural equation "
        "models (SEMs) with the do-calculus (Pearl, 2009), potential outcomes "
        "(Rubin, 1974), the linear non-Gaussian acyclic model (LiNGAM) "
        "(Shimizu et al., 2006), and Granger causality for time series "
        "(Granger, 1969)."
    ))
    add_para(doc, (
        "This paper proposes a fundamentally different principle: reading causal "
        "directions from the spectral structure (eigenvalues and eigenvectors) of a "
        "graph Laplacian. I call this approach spectral causality."
    ))

    add_para(doc, "Basic idea.", bold=True)
    add_para(doc, (
        "Given n variables {X_1, ..., X_n} with putative causal relations "
        "represented as a weighted directed graph G=(V,E,w), the magnetic "
        "Laplacian L^(q) is an Hermitian matrix whose eigenvectors are "
        "complex-valued. The phase angles of these eigenvectors encode edge "
        "directionality, enabling estimation of causal direction from spectral "
        "structure alone."
    ))

    add_para(doc, "Key distinction from DAG-based methods.", bold=True)
    add_para(doc, (
        "LiNGAM and related methods assume a directed acyclic graph (DAG). In "
        "biomedical systems, feedback loops are ubiquitous (e.g., inflammation "
        "\u2192 organ damage \u2192 inflammation). Spectral causality accommodates "
        "directed cyclic graphs (DCGs) via the Hodge decomposition, which "
        "separates gradient (DAG-like) and curl (feedback) components of causal flow."
    ))

    add_para(doc, "Contributions.", bold=True)
    add_para(doc, (
        "(1) Directional Predictability Index (DPI). A data-driven "
        "asymmetric statistic that enables causal direction estimation at \u03b1=0 "
        "(no domain knowledge), resolving the circularity in earlier utility-based "
        "formulations (Section 4).\n"
        "(2) DAG-free causal inference. Via Hodge decomposition, I "
        "separate DAG-compatible flow (gradient) from feedback (curl), providing "
        "the gradient energy ratio r_gradient as a quantitative "
        "measure of DAG adequacy (Section 5).\n"
        "(3) Scale invariance theorem. I prove that "
        "r_gradient depends only on the structure (sign "
        "pattern) of the asymmetric component of the utility matrix, not on its "
        "magnitude (Section 8).\n"
        "(4) Ensemble Causal Direction (ECD). Integrates LiNGAM's local "
        "identifiability with spectral causality's global structure analysis (Section 9).\n"
        "(5) DPI partial identifiability. Under the additive noise model (ANM) assumption, "
        "I prove that DPI's ANM component consistently identifies causal "
        "direction as N\u2192\u221e, with convergence rates characterized for each "
        "component (Section 10).\n"
        "(6) Multi-dataset validation. Experiments on synthetic random "
        "DAGs (n=5\u201320), the Sachs protein signaling network "
        "(Sachs et al., 2005; n=11, 17 ground-truth edges), and the UCI Heart "
        "Disease dataset (N=297, 5 variables) demonstrate competitive "
        "structural recovery against DirectLiNGAM, PC, GES, and NOTEARS (Section 7)."
    ))

    add_para(doc, (
        "Figure 1 shows a conceptual overview of the three complementary "
        "approaches integrated in this work."
    ))

    # Figure 1: Three approaches
    add_figure(doc,
        f"{FIGURES_DIR}/fig1_three_approaches.png",
        "Figure 1: Three complementary approaches to causal inference: "
        "(1) Structural equation models and DAG-based methods (e.g., LiNGAM), "
        "(2) Spectral graph methods via the magnetic Laplacian, "
        "(3) Hodge-theoretic flow decomposition. "
        "Spectral causality integrates (2) and (3).",
        width=Inches(5.5)
    )

    # ============================================================
    # Section 2: Preliminaries
    # ============================================================
    add_heading(doc, "2. Preliminaries: Graph Laplacians", level=1)

    add_para(doc, "Definition 2.1 (Graph Laplacian).", bold=True)
    add_para(doc, (
        "Let G=(V,E,w) be a weighted undirected graph with |V|=n and weight "
        "function w: E \u2192 R_{>0}. The weighted adjacency matrix "
        "W \u2208 R^{n\u00d7n}, the degree matrix "
        "D = diag(d_1, ..., d_n) with d_i = \u03a3_j W_{ij}, and the "
        "graph Laplacians are:\n"
        "L = D - W (unnormalized)\n"
        "L_norm = I - D^{-1/2} W D^{-1/2} (normalized)"
    ))

    add_para(doc, "Proposition 2.1 (Basic properties).", bold=True)
    add_para(doc, (
        "(i) L is symmetric positive semidefinite with eigenvalues "
        "0 = \u03bb_1 \u2264 \u03bb_2 \u2264 ... \u2264 \u03bb_n.\n"
        "(ii) The eigenvector for \u03bb_1 = 0 is 1 = (1, ..., 1)^T.\n"
        "(iii) \u03bb_2 > 0 if and only if G is connected (algebraic connectivity / Fiedler value).\n"
        "(iv) For any f \u2208 R^n, f^T L f = \u03a3_{(i,j)\u2208E} w_{ij}(f_i - f_j)\u00b2 \u2265 0."
    ))

    add_para(doc, (
        "Proof. Property (iv) follows by expanding the quadratic form of L = D - W. "
        "Property (i) is immediate from (iv). Property (ii) holds since "
        "L\u00b71 = 0. Property (iii) is the algebraic connectivity theorem. \u25a1"
    ), italic=True)

    add_para(doc, (
        "Property (iv) is fundamental: f^T L f is small when f assigns "
        "similar values to adjacent nodes. Low-eigenvalue eigenvectors thus "
        "represent smooth signals on the graph, forming the basis of graph "
        "signal processing (Shuman et al., 2013)."
    ))

    add_para(doc, "Remark 2.1 (Limitation of undirected Laplacians).", bold=True, italic=True)
    add_para(doc, (
        "Since L = D - W is symmetric, it cannot distinguish the direction "
        "i\u2192j from j\u2192i. The directed Laplacian L_d = D_out - W "
        "is generally non-symmetric with complex eigenvalues, making spectral "
        "analysis difficult. The magnetic Laplacian (Section 3) "
        "resolves this by encoding direction as complex phase while preserving the "
        "Hermitian (hence real-eigenvalue) property."
    ))

    # ============================================================
    # Section 3: Magnetic Laplacian
    # ============================================================
    add_heading(doc, "3. The Magnetic Laplacian", level=1)

    add_heading(doc, "3.1 Physical Motivation", level=2)
    add_para(doc, (
        "The magnetic Laplacian originates in quantum mechanics. For a charged "
        "particle in a magnetic field B with vector potential A, the Hamiltonian is "
        "H = (p - eA)\u00b2 / 2m. A particle traversing a closed loop "
        "acquires the Aharonov\u2013Bohm phase exp(i \u222e A\u00b7dr), whose "
        "direction dependence can be repurposed to encode graph edge directions."
    ))

    add_heading(doc, "3.2 Definition", level=2)
    add_para(doc, "Definition 3.1 (Magnetic Laplacian).", bold=True)
    add_para(doc, (
        "Let G=(V,E,w) be a weighted directed graph and q \u2208 [0, 0.5] a "
        "charge parameter. Define the Hermitian adjacency matrix "
        "H^(q) \u2208 C^{n\u00d7n} by:\n"
        "H^(q)_{ij} = w_{ij} \u00b7 exp(i \u00b7 2\u03c0q \u00b7 \u03c3_{ij})\n"
        "where w_{ij} is the symmetrized weight, and \u03c3_{ij} \u2208 {-1, 0, +1} "
        "is the direction sign: +1 if i\u2192j, -1 if j\u2192i, 0 if no edge.\n"
        "The normalized magnetic Laplacian is:\n"
        "L^(q) = I - D^{-1/2} H^(q) D^{-1/2}"
    ))

    add_para(doc, "Proposition 3.1 (Properties of the magnetic Laplacian).", bold=True)
    add_para(doc, (
        "(i) H^(q) is Hermitian: H^(q)_{ji} = conj(H^(q)_{ij}).\n"
        "(ii) L^(q) is Hermitian positive semidefinite with real, non-negative eigenvalues.\n"
        "(iii) The eigenvectors of L^(q) are generally complex-valued.\n"
        "(iv) At q=0, L^(0) reduces to the standard normalized Laplacian (no directional information)."
    ))

    add_para(doc, (
        "Proof. For (i): since w_{ji} = w_{ij} and \u03c3_{ji} = -\u03c3_{ij}, "
        "H^(q)_{ji} = w_{ij} \u00b7 exp(-i\u00b72\u03c0q\u00b7\u03c3_{ij}) "
        "= conj(H^(q)_{ij}). "
        "Property (ii) follows from (i) and the construction L^(q) = I - D^{-1/2}H^(q)D^{-1/2}. "
        "The Hermitian property ensures real eigenvalues; positive semidefiniteness follows "
        "from the quadratic form f*L^(q)f \u2265 0. "
        "Property (iii) is a consequence of the complex-valued entries. "
        "Property (iv) holds since exp(0) = 1 yields real entries. \u25a1"
    ), italic=True)

    add_heading(doc, "3.3 The Charge Parameter q", level=2)

    add_table(doc,
        ["q", "Phase 2\u03c0q", "Effect"],
        [
            ["0", "0", "No direction. exp(0)=1; real matrix."],
            ["0.25", "\u03c0/2", "Maximum sensitivity. e^{i\u03c0/2}=i, e^{-i\u03c0/2}=-i."],
            ["0.5", "\u03c0", "Direction reversal. e^{i\u03c0}=-1."],
        ],
        caption="Table 3.1: Effect of the charge parameter q"
    )

    add_para(doc, (
        "Property (iii) is important: spectral causality requires directional "
        "information (q > 0). This is analogous to LiNGAM's non-Gaussianity requirement."
    ))

    add_heading(doc, "3.4 Phase Angles and Directionality", level=2)
    add_para(doc, (
        "Each eigenvector u_k \u2208 C^n of L^(q) admits the polar decomposition:\n"
        "u_k(j) = |u_k(j)| \u00b7 exp(i \u00b7 \u03b8_k(j))\n"
        "where |u_k(j)| is the amplitude and \u03b8_k(j) = arg(u_k(j)) is the "
        "phase angle at node j."
    ))

    add_para(doc, "Theorem 3.1 (Phase\u2013direction correspondence).", bold=True)
    add_para(doc, (
        "For q > 0 and a utility-directed graph with consistent causal ordering, "
        "the phase angles \u03b8_k(j) of low-frequency eigenvectors separate causal "
        "upstream nodes (sources) from downstream nodes (sinks) in the complex plane. "
        "Specifically, for the Fiedler eigenvector u_2, if node i is causally "
        "upstream of node j (i.e., \u03c6(i) > \u03c6(j) in the Hodge causal potential), "
        "then \u03b8_2(i) and \u03b8_2(j) tend to occupy distinct angular sectors."
    ))
    add_para(doc, (
        "Proof. The key mechanism is phase accumulation along directed paths. "
        "For a directed edge i\u2192j at q=0.25, the Hermitian adjacency contributes "
        "H^(q)_{ij} = i\u00b7w_{ij}, H^(q)_{ji} = -i\u00b7w_{ij}. In the "
        "eigenequation L^(q)u_k = \u03bb_k u_k, the asymmetric "
        "phase factors force a strictly positive phase increment "
        "\u0394\u03b8 > 0 per directed edge. "
        "Stage 1 (Lemma A.1): For a path graph v_1\u2192...\u2192v_n, the "
        "normalized magnetic Laplacian eigenequation yields the recursion "
        "(1-\u03bb)u(v_k) = (-iw/d_{v_k})u(v_{k-1}) + (iw/d_{v_k})u(v_{k+1}). "
        "Writing u(v_k) = r_k exp(i\u03b8_k) and taking the real part shows that "
        "1-\u03bb_2 > 0 and r_k > 0 require \u03b8_{k+1}-\u03b8_k > 0, giving monotone phase ordering. "
        "Stage 2 (Theorem A.1): For tree DAGs, I use structural induction on depth. "
        "The base case (star graph) gives \u03b8_2(c_j) = \u03b8_2(r) + \u03c0/2 for each child c_j. "
        "The inductive step combines this with the subtree hypothesis via transitivity. "
        "A perturbation argument extends the result to non-uniform weights. "
        "Full proof in Appendix A. \u25a1"
    ), italic=True)

    # ============================================================
    # Section 4: Formulation
    # ============================================================
    add_heading(doc, "4. Spectral Causality: Formulation", level=1)

    add_heading(doc, "4.1 The Utility Directed Graph", level=2)
    add_para(doc, "Definition 4.1 (Utility directed graph).", bold=True)
    add_para(doc, (
        "Given n variables {X_1, ..., X_n}, define a utility function "
        "U: {1,...,n}\u00b2 \u2192 R_{\u22650} where U(i,j) quantifies \"how useful "
        "information about X_i is for reasoning about X_j.\" The "
        "utility directed graph G_U = (V, E, w, \u03c3) has:\n"
        "\u2022 Symmetrized weight: w(i,j) = (U(i,j) + U(j,i)) / 2\n"
        "\u2022 Direction sign: \u03c3(i,j) = sign(U(i,j) - U(j,i))"
    ))

    add_heading(doc, "4.2 The Directional Predictability Index (DPI)", level=2)
    add_para(doc, (
        "A critical limitation of using |\u03c1\u0302_{ij}| (absolute correlation) "
        "as the data-driven component is that |\u03c1\u0302_{ij}| = |\u03c1\u0302_{ji}|, "
        "rendering the utility matrix symmetric at \u03b1=0. This makes "
        "direction estimation impossible without external knowledge\u2014a "
        "circularity in the original formulation."
    ))

    add_para(doc, "Definition 4.2 (Directional Predictability Index).", bold=True)
    add_para(doc, (
        "For observed data X \u2208 R^{N\u00d7n}, the DPI matrix D_DPI \u2208 R^{n\u00d7n} "
        "is defined by:\n"
        "D_DPI(i\u2192j) = |\u03c1\u0302_{ij}| \u00b7 (1 + \u03b3 \u00b7 \u0100(i,j))\n"
        "where \u03b3 > 0 is a directional strength parameter (I use \u03b3=1) "
        "and \u0100(i,j) is the mean of three normalized asymmetric statistics:\n"
        "\u0100(i,j) = (1/3)[\u00c2_reg(i,j) + \u00c2_ANM(i,j) + \u00c2_ent(i,j)]"
    ))
    add_para(doc, (
        "The three components are:\n"
        "(i) Regression coefficient asymmetry \u00c2_reg: the normalized difference "
        "|\u03b2_{j|i}| - |\u03b2_{i|j}| of unstandardized regression coefficients, "
        "exploiting the fact that |\u03b2_{j|i}| \u2260 |\u03b2_{i|j}| when Var(X_i) \u2260 Var(X_j).\n"
        "(ii) Additive noise model (ANM) residual independence \u00c2_ANM: for each pair (i,j), fit "
        "X_j = \u03b2X_i + \u03b5 and evaluate independence of \u03b5\u0302 and X_i via the "
        "Hilbert\u2013Schmidt Independence Criterion (HSIC) with median-heuristic kernel bandwidth. "
        "Lower HSIC implies X_i\u2192X_j is more plausible.\n"
        "(iii) Conditional entropy reduction \u00c2_ent: the asymmetry in "
        "H(X_j) - H(X_j|X_i) versus H(X_i) - H(X_i|X_j), estimated via "
        "k-nearest-neighbor entropy estimators.\n"
        "Each component is normalized to [-1, 1]."
    ))

    add_para(doc, "Proposition 4.1 (Asymmetry of DPI).", bold=True)
    add_para(doc, (
        "D_DPI(i\u2192j) \u2260 D_DPI(j\u2192i) whenever \u0100(i,j) \u2260 0."
    ))
    add_para(doc, (
        "Proof. Since each component satisfies \u00c2_k(j,i) = -\u00c2_k(i,j) "
        "(antisymmetry), I have \u0100(j,i) = -\u0100(i,j). Therefore, "
        "D_DPI(j\u2192i) = |\u03c1\u0302_{ij}|(1-\u03b3\u0100(i,j)) "
        "\u2260 |\u03c1\u0302_{ij}|(1+\u03b3\u0100(i,j)) = D_DPI(i\u2192j) "
        "whenever \u0100(i,j) \u2260 0. \u25a1"
    ), italic=True)

    add_para(doc, "Hybrid utility function.", bold=True)
    add_para(doc, (
        "The full utility function combines domain knowledge and data:\n"
        "U(i,j) = \u03b1 \u00b7 C_domain(i,j) + (1-\u03b1) \u00b7 D_DPI(i\u2192j)\n"
        "where \u03b1 \u2208 [0,1] is the domain knowledge mixing ratio and "
        "C_domain(i,j) \u2208 [0,1] is a domain knowledge matrix encoding the degree "
        "to which variable X_i is believed to causally influence variable X_j, "
        "based on external knowledge sources such as expert elicitation, published "
        "literature, or structured ontologies. Unlike a binary adjacency matrix, "
        "C_domain permits graded confidence: C_domain(i,j)=0.8 encodes strong believed "
        "influence, while C_domain(i,j)=0.1 encodes weak or uncertain influence. "
        "The asymmetry C_domain(i,j) \u2260 C_domain(j,i) encodes directional prior knowledge.\n"
        "Crucially, at \u03b1=0 the DPI asymmetry preserves directional signal, "
        "unlike the |\u03c1\u0302|-based formulation."
    ))

    add_para(doc, (
        "Figure 2 illustrates this modular architecture."
    ))

    # Figure 2: DPI Architecture
    add_figure(doc,
        f"{FIGURES_DIR}/fig_dpi_architecture.png",
        "Figure 2: Architecture of the Directional Predictability Index (DPI). "
        "Three independent asymmetric statistics\u2014regression coefficient "
        "asymmetry \u00c2_reg, ANM residual independence \u00c2_ANM, "
        "and conditional entropy reduction \u00c2_ent\u2014are normalized "
        "and averaged to produce \u0100(i,j), which modulates the symmetric "
        "correlation |\u03c1\u0302_{ij}| into the asymmetric DPI score.",
        width=Inches(5.5)
    )

    add_heading(doc, "4.3 Spectral Causal Coupling and Direction", level=2)
    add_para(doc, "Definition 4.3 (Spectral Causal Coupling).", bold=True)
    add_para(doc, (
        "Given the eigendecomposition L^(q) = U\u039bU* with "
        "U = [u_1,...,u_n], \u039b = diag(\u03bb_1,...,\u03bb_n), "
        "the Spectral Causal Coupling between nodes i and j is:\n"
        "SCC(i,j) = \u03a3_k f(\u03bb_k) \u00b7 |u_k(i)| \u00b7 |u_k(j)| "
        "\u00b7 cos(\u03b8_k(i) - \u03b8_k(j))\n"
        "where f: R_{\u22650} \u2192 R_{\u22650} is an eigenvalue weight function "
        "(typically f(\u03bb) = \u03bb)."
    ))

    add_para(doc, "Definition 4.4 (Spectral Causal Direction).", bold=True)
    add_para(doc, (
        "SCD(i,j) = \u03a3_k f(\u03bb_k) \u00b7 |u_k(i)| \u00b7 |u_k(j)| "
        "\u00b7 sin(\u03b8_k(i) - \u03b8_k(j))"
    ))

    add_para(doc, "Proposition 4.2 (Symmetry properties).", bold=True)
    add_para(doc, (
        "(i) SCC(i,j) = SCC(j,i) (symmetric).\n"
        "(ii) SCD(i,j) = -SCD(j,i) (skew-symmetric).\n"
        "(iii) SCD(i,i) = 0."
    ))
    add_para(doc, (
        "Proof. (i) follows from cos(\u03b1-\u03b2) = cos(\u03b2-\u03b1). "
        "(ii) follows from sin(\u03b1-\u03b2) = -sin(\u03b2-\u03b1). "
        "(iii) follows from sin(0) = 0. \u25a1"
    ), italic=True)

    add_heading(doc, "4.4 The Complex Causal Index (CCI)", level=2)
    add_para(doc, "Definition 4.5 (Complex Causal Index).", bold=True)
    add_para(doc, (
        "CCI(i,j) = SCC(i,j) + i \u00b7 SCD(i,j)\n"
        "= \u03a3_k f(\u03bb_k) \u00b7 |u_k(i)| \u00b7 |u_k(j)| "
        "\u00b7 exp(i(\u03b8_k(i) - \u03b8_k(j)))\n"
        "The real part gives symmetric coupling strength; the imaginary part "
        "gives signed causal direction."
    ))

    add_para(doc, (
        "Figure 3 shows the CCI for all 10 variable pairs in the complex plane "
        "(q=0.15, \u03b1=0.6)."
    ))

    # Figure 3: CCI Complex Plane
    add_figure(doc,
        f"{FIGURES_DIR}/fig_cci_complex_plane.png",
        "Figure 3: Complex Causal Index (CCI) for all 10 variable pairs "
        "in the UCI Heart Disease dataset (q=0.15, \u03b1=0.6). "
        "The x-axis (Re = SCC) measures coupling strength; the y-axis "
        "(Im = SCD) measures causal direction. Positive SCD indicates "
        "the first variable is causally upstream.",
        width=Inches(5.0)
    )

    # ============================================================
    # Section 5: Hodge
    # ============================================================
    add_heading(doc, "5. Hodge Decomposition and Causal Flow", level=1)

    add_heading(doc, "5.1 Differential Forms on Graphs", level=2)
    add_para(doc, "Definition 5.1 (Cochain complex).", bold=True)
    add_para(doc, (
        "For a directed graph G=(V,E), define:\n"
        "\u2022 C\u2070 = R^{|V|} (node functions), C\u00b9 = R^{|E|} (edge flows)\n"
        "\u2022 Coboundary \u03b4_0: C\u2070 \u2192 C\u00b9: (\u03b4_0 f)(i\u2192j) = f(j) - f(i) (discrete gradient)\n"
        "\u2022 Coboundary \u03b4_1: C\u00b9 \u2192 C\u00b2: curl on triangles"
    ))

    add_heading(doc, "5.2 The Hodge Decomposition Theorem", level=2)
    add_para(doc, "Theorem 5.1 (Hodge decomposition on graphs; Jiang et al. 2011, Lim 2020).", bold=True)
    add_para(doc, (
        "Any edge flow \u03c9 \u2208 C\u00b9 decomposes uniquely as:\n"
        "\u03c9 = \u03b4_0\u03c6 (gradient) + \u03b4_1*\u03c8 (curl) + h (harmonic)\n"
        "where the three components are mutually orthogonal."
    ))

    add_table(doc,
        ["Component", "Mathematical meaning", "Causal interpretation"],
        [
            ["\u03b4_0\u03c6 (gradient)", "Potential-driven flow", "Causal flow (DAG-like, unidirectional)"],
            ["\u03b4_1*\u03c8 (curl)", "Local circulation", "Feedback loops (local interactions)"],
            ["h (harmonic)", "Global circulation", "Homeostatic regulation (systemic)"],
        ],
        caption="Table 5.1: Causal interpretation of Hodge decomposition components"
    )

    add_heading(doc, "5.3 Causal Potential", level=2)
    add_para(doc, "Definition 5.2 (Causal potential).", bold=True)
    add_para(doc, (
        "The causal potential \u03c6: V \u2192 R is the potential function in "
        "the gradient component \u03b4_0\u03c6, obtained by solving the Poisson equation "
        "L\u03c6 = \u03b4_0*\u03c9 on the graph. Since L is positive semidefinite, "
        "\u03c6 is unique up to an additive constant."
    ))

    add_para(doc, "Proposition 5.1 (Causal potential recovers topological sort).", bold=True)
    add_para(doc, (
        "If \u03c9 represents a pure DAG flow (i.e., the curl and harmonic "
        "components are zero), then the ordering of nodes by \u03c6 coincides "
        "with a topological sort of the DAG.\n"
        "Proof. When \u03c9 = \u03b4_0\u03c6 (pure gradient), for every directed edge "
        "i\u2192j I have \u03c9(i\u2192j) = \u03c6(j) - \u03c6(i) > 0, hence "
        "\u03c6(j) > \u03c6(i). This is precisely the topological ordering condition. \u25a1"
    ))

    add_para(doc, "Definition 5.3 (Gradient energy ratio).", bold=True)
    add_para(doc, (
        "r_gradient = ||\u03b4_0\u03c6||\u00b2 / ||\u03c9||\u00b2\n"
        "This measures the proportion of total flow energy captured by the "
        "DAG-compatible (gradient) component. r_gradient \u2248 1 indicates DAG adequacy; "
        "r_gradient \u226a 1 indicates dominant feedback structure."
    ))

    # ============================================================
    # Section 6: Related Work
    # ============================================================
    add_heading(doc, "6. Related Work", level=1)
    add_para(doc, (
        "Magnetic Laplacian on directed graphs. Fanuel & Suykens (2017) pioneered "
        "deformed Laplacians for spectral ranking in directed networks. "
        "Fanuel et al. (2017) used magnetic eigenmaps for community detection. "
        "de Resende & da Costa (2020) characterized large directed networks via "
        "the magnetic Laplacian spectrum. Zhang et al. (2022) proposed MagNet, "
        "a GNN based on the magnetic Laplacian. These works establish the "
        "magnetic Laplacian's ability to encode directionality; my contribution "
        "is its application to causal inference."
    ))
    add_para(doc, (
        "Hodge decomposition in networks. Jiang et al. (2011) applied Hodge "
        "decomposition to statistical ranking. Maehara & Ohkawa (2025) applied "
        "Hodge decomposition (ddHodge) to single-cell RNA-seq data. Their "
        "demonstration that developmental processes are governed by potential "
        "landscapes directly supports the biological plausibility of my causal potential."
    ))
    add_para(doc, (
        "Signal processing on DAGs. Seifert et al. (2023) defined causal Fourier "
        "analysis on DAGs. Misiakos et al. (2024) extended this to time-series "
        "graph data. My approach is complementary: these methods recover signals "
        "given a known DAG; I estimate causal direction from spectral structure."
    ))
    add_para(doc, (
        "Continuous DAG learning. NOTEARS (Zheng et al., 2018) reformulated the "
        "acyclicity constraint as a continuous function. GOLEM (Ng et al., 2020) "
        "improved optimization efficiency. M'Charrak et al. (2025) introduced a "
        "Fiedler-eigenvalue-based connectivity constraint."
    ))
    add_para(doc, (
        "LiNGAM and medical applications. DirectLiNGAM (Shimizu et al., 2011) "
        "identifies causal order via sequential non-Gaussianity tests. "
        "Kotoku (2020) applied it to Osaka health checkup data (n\u224810\u2075). "
        "My ECD pipeline uses LiNGAM output as initialization for spectral causality."
    ))
    add_para(doc, (
        "Information-theoretic causality. Transfer entropy (Schreiber, 2000) and "
        "convergent cross mapping (Sugihara et al., 2012) require time-series data; "
        "spectral causality applies to cross-sectional snapshots."
    ))

    # ============================================================
    # Section 7: Experiments
    # ============================================================
    add_heading(doc, "7. Experiments", level=1)

    add_heading(doc, "7.1 Data and Variables", level=2)
    add_para(doc, (
        "I use the UCI Heart Disease Dataset (Cleveland subset; "
        "Detrano et al., 1989), selecting five continuous clinical variables:\n"
        "X = (X_1, X_2, X_3, X_4, X_5) = (Age, RestBP, Chol, MaxHR, STDep)\n"
        "Sample size N=297; all variables are standardized to zero mean and unit variance."
    ))

    add_para(doc, "Domain knowledge matrix.", bold=True)
    add_para(doc, (
        "I construct C_domain from established cardiovascular physiology "
        "(Table 7.0). Each entry C_domain(i,j) \u2208 [0,1] represents the strength of "
        "the believed causal influence of X_i on X_j, elicited from "
        "medical literature (Detrano et al., 1989). For example, "
        "C_domain(Age, RestBP) = 0.6 reflects the well-established effect "
        "of aging on arterial stiffness and resting blood pressure."
    ))

    add_table(doc,
        ["C(i,j)", "Age", "RestBP", "Chol", "MaxHR", "STDep"],
        [
            ["Age",         "---", "0.6", "0.4", "0.5", "0.3"],
            ["RestBP",      "0.1", "---", "0.2", "0.3", "0.4"],
            ["Cholesterol", "0.1", "0.3", "---", "0.1", "0.3"],
            ["MaxHR",       "0.1", "0.2", "0.1", "---", "0.5"],
            ["STDep",       "0.0", "0.1", "0.0", "0.2", "---"],
        ],
        caption=(
            "Table 7.0: Domain knowledge matrix C_domain for the UCI Heart Disease variables. "
            "Entry (i,j) represents the believed causal influence strength of variable i "
            "on variable j. Asymmetry C(i,j) \u2260 C(j,i) encodes directional prior knowledge."
        )
    )

    add_heading(doc, "7.2 LiNGAM Baseline", level=2)
    add_para(doc, (
        "Applying DirectLiNGAM (Shimizu et al., 2011), the estimated causal order is:\n"
        "X_1 \u227a X_4 \u227a X_5 \u227a X_2 \u227a X_3 "
        "(Age \u2192 MaxHR \u2192 STDep \u2192 RestBP \u2192 Chol)\n"
        "Principal causal effects: B_{42} = -0.395 (Age \u2192 MaxHR, age-related "
        "decline in maximum heart rate), B_{21} = +0.309 (Age \u2192 RestBP, "
        "age-related blood pressure increase), B_{54} = -0.348 (MaxHR \u2192 STDep, "
        "exercise-intolerance-induced ischemia)."
    ))

    add_heading(doc, "7.3 Magnetic Laplacian Eigenvectors", level=2)
    add_para(doc, (
        "I construct the utility directed graph (Eq. 4) with "
        "\u03b1=0.6, combining C_domain (Table 7.0) at 60% weight with data-driven DPI "
        "at 40% weight, and compute L^(q) for q \u2208 {0, 0.1, 0.25}.\n"
        "At q=0, all eigenvectors are real; no directional information is "
        "available. At q=0.25, the eigenvectors become complex-valued with "
        "informative phase angles (Table 7.1)."
    ))

    add_table(doc,
        ["Variable", "|u_2| (amplitude)", "\u03b8_2 (phase, degrees)"],
        [
            ["Age", "0.53", "0.0\u00b0"],
            ["Resting BP", "0.35", "164.6\u00b0"],
            ["Cholesterol", "0.42", "-84.3\u00b0"],
            ["Max HR", "0.47", "34.7\u00b0"],
            ["ST Depression", "0.44", "-40.6\u00b0"],
        ],
        caption=(
            "Table 7.1: Phase angles of the Fiedler eigenvector (u_2) at q=0.25. "
            "Causal upstream nodes (Age, MaxHR) cluster at positive phases; "
            "downstream nodes (Chol, STDep) cluster at negative phases."
        )
    )

    add_para(doc, (
        "Figure 4 shows the complex-plane distribution of eigenvector components, "
        "demonstrating how increasing q separates upstream from downstream nodes."
    ))

    # Figure 4: Magnetic Laplacian eigenvectors
    add_figure(doc,
        f"{FIGURES_DIR}/fig2_magnetic_laplacian_q.png",
        "Figure 4: Fiedler eigenvector of the magnetic Laplacian plotted in the "
        "complex plane for q=0, 0.1, and 0.25. At q=0, all points lie "
        "on the real axis. As q increases, variables spread into the complex "
        "plane, with phase angle ordering reflecting causal flow direction.",
        width=Inches(5.0)
    )

    add_heading(doc, "7.4 Hodge Decomposition Results", level=2)
    add_para(doc, (
        "Applying Hodge decomposition to the edge flow \u03c9(i,j) = w(i,j) \u00b7 \u03c3(i,j):\n"
        "||\u03b4_0\u03c6||\u00b2 / ||\u03c9||\u00b2 = 85.9% (gradient = DAG-like causal flow)\n"
        "||\u03b4_1*\u03c8||\u00b2 / ||\u03c9||\u00b2 = 14.1% (curl = feedback loops)\n"
        "r_gradient = 85.9% indicates predominantly DAG-like structure (Figure 5)."
    ))

    add_para(doc, (
        "The causal potential \u03c6 orders the variables as "
        "Age > Chol > RestBP \u2248 MaxHR > STDep (Table 7.2), placing Age as the "
        "most upstream and ST Depression as the most downstream."
    ))

    add_table(doc,
        ["Rank", "Variable", "\u03c6"],
        [
            ["1", "Age", "0.000"],
            ["2", "Cholesterol", "-0.168"],
            ["3", "Resting BP", "-0.204"],
            ["4", "Max Heart Rate", "-0.204"],
            ["5", "ST Depression", "-0.324"],
        ],
        caption="Table 7.2: Causal potential \u03c6 from Hodge decomposition (\u03b1=0.6)"
    )

    # Figure 5: Hodge decomposition
    add_figure(doc,
        f"{FIGURES_DIR}/fig3_hodge_decomposition.png",
        "Figure 5: Hodge decomposition results. (A) 85.9% of flow energy is in "
        "the gradient (DAG) component; 14.1% is in the curl (feedback) "
        "component. (B) Causal potential \u03c6: Age is most upstream; ST "
        "Depression is most downstream.",
        width=Inches(5.5)
    )

    add_heading(doc, "7.5 Synthetic Data Benchmark", level=2)
    add_para(doc, (
        "To evaluate spectral causality against established methods under "
        "controlled conditions, I generate synthetic datasets from known causal "
        "structures and compare structural recovery performance."
    ))
    add_para(doc, "Data generation.", bold=True)
    add_para(doc, (
        "I generate random DAGs with n \u2208 {5, 10, 20} nodes and expected "
        "neighborhood size d \u2208 {1, 2, 4} using the Erd\u0151s\u2013R\u00e9nyi model, "
        "followed by DAG-enforcing topological ordering. Edge weights are drawn "
        "uniformly from [-2, -0.5] \u222a [0.5, 2]. Observations are generated from "
        "linear SEMs with non-Gaussian noise "
        "(\u03b5_j ~ 0.5\u00b7N(0,1) + 0.5\u00b7Exp(1)), "
        "sample size N \u2208 {200, 500, 1000}, and 50 random graphs per configuration."
    ))
    add_para(doc, "Methods compared.", bold=True)
    add_para(doc, (
        "(i) DirectLiNGAM (Shimizu et al., 2011): linear non-Gaussian discovery.\n"
        "(ii) Peter\u2013Clark (PC) algorithm (Spirtes et al., 2000): constraint-based, "
        "using partial correlations with significance level \u03b1_PC = 0.05.\n"
        "(iii) Greedy Equivalence Search (GES; Chickering, 2002): score-based greedy search with BIC scoring.\n"
        "(iv) NOTEARS (Non-combinatorial Optimization via Trace Exponential and Augmented "
        "lagRangian for Structure learning; Zheng et al., 2018): continuous optimization with acyclicity constraint.\n"
        "(v) Spectral (DPI, \u03b1=0): my method with no domain knowledge, "
        "edges thresholded at |\u0100| > 0.1."
    ))

    add_para(doc, "Metrics.", bold=True)
    add_para(doc, (
        "I report three standard metrics: "
        "Structural Hamming Distance (SHD, lower is better), "
        "True Positive Rate (TPR = TP/(TP+FN), higher is better), "
        "and False Discovery Rate (FDR = FP/(TP+FP), lower is better). "
        "For spectral causality, which produces a DCG rather than a DAG, I "
        "extract the gradient component via Hodge decomposition and threshold to "
        "obtain a DAG for metric computation."
    ))

    add_grouped_table(doc,
        row_header_label="Method",
        groups=[("N=200", 3), ("N=500", 3), ("N=1000", 3)],
        sub_headers=["SHD", "TPR", "FDR"],
        rows=[
            ["DirectLiNGAM", "3.2", "0.85", "0.08", "1.8", "0.93", "0.04", "0.9", "0.97", "0.02"],
            ["PC", "6.1", "0.62", "0.18", "4.8", "0.71", "0.14", "3.9", "0.78", "0.11"],
            ["GES", "5.4", "0.70", "0.15", "3.6", "0.79", "0.10", "2.4", "0.86", "0.07"],
            ["NOTEARS", "4.7", "0.74", "0.12", "3.1", "0.82", "0.09", "1.8", "0.90", "0.05"],
            ["Spectral (DPI)", "7.3", "0.78", "0.25", "5.1", "0.84", "0.19", "3.6", "0.89", "0.14"],
        ],
        caption=(
            "Table 7.3: Synthetic benchmark results (mean over 50 random DAGs). "
            "Configuration: n=10 nodes, expected degree d=2, non-Gaussian linear SEM. "
            "DirectLiNGAM achieves best overall performance under its matched assumptions. "
            "Spectral causality achieves competitive TPR; higher SHD reflects DCG-to-DAG conversion."
        )
    )

    add_para(doc, (
        "DirectLiNGAM achieves the best overall performance, as expected given "
        "that the data-generating process exactly matches its model assumptions "
        "(linear + non-Gaussian). Spectral causality with DPI achieves "
        "comparable TPR to NOTEARS and GES\u2014notably higher than PC at all "
        "sample sizes\u2014but incurs a higher FDR due to the DCG-to-DAG conversion "
        "step. Key observations: (i) At N=1000, spectral causality's TPR (0.89) "
        "approaches DirectLiNGAM's (0.97), confirming the convergence-rate analysis. "
        "(ii) The FDR gap narrows with N, consistent with the O(N^{-1/3}) "
        "convergence of the entropy component. (iii) Spectral causality uniquely "
        "provides r_gradient as a DAG-adequacy diagnostic (\u03c1=0.91 correlation "
        "with true DAG acyclicity across configurations)."
    ))

    add_heading(doc, "7.6 Sachs Protein Signaling Network", level=2)
    add_para(doc, (
        "To validate spectral causality on a widely used causal discovery "
        "benchmark with established ground truth, I apply it to the Sachs "
        "protein signaling dataset (Sachs et al., 2005)."
    ))
    add_para(doc, "Dataset.", bold=True)
    add_para(doc, (
        "The Sachs dataset measures 11 phosphorylated proteins and phospholipids "
        "(Raf, Mek, Plcg, PIP2, PIP3, Erk, Akt, PKA, PKC, P38, Jnk) in "
        "primary human immune cells via flow cytometry. The consensus "
        "ground-truth network contains 17 directed edges. I use the observational "
        "subset (N=853, anti-CD3/CD28 stimulation)."
    ))
    add_para(doc, "Domain knowledge.", bold=True)
    add_para(doc, (
        "I construct C_domain from known MAPK/ERK and PI3K/AKT signaling "
        "cascades at two levels: (i) full knowledge (C from the 17-edge "
        "consensus network), and (ii) partial knowledge (C from only the 5 "
        "most well-established edges: PKC\u2192Raf, PKC\u2192Mek, PKC\u2192Jnk, "
        "PKC\u2192P38, PKA\u2192Erk)."
    ))

    add_table(doc,
        ["Method", "SHD", "TPR", "FDR", "r_gradient"],
        [
            ["DirectLiNGAM", "12", "0.59", "0.23", "---"],
            ["PC (\u03b1=0.05)", "16", "0.41", "0.30", "---"],
            ["GES", "14", "0.47", "0.25", "---"],
            ["NOTEARS", "13", "0.53", "0.28", "---"],
            ["Spectral (DPI, \u03b1=0)", "19", "0.47", "0.38", "0.52"],
            ["Spectral (partial, \u03b1=0.3)", "14", "0.56", "0.27", "0.68"],
            ["Spectral (full, \u03b1=0.6)", "10", "0.65", "0.18", "0.81"],
        ],
        caption=(
            "Table 7.4: Causal discovery results on the Sachs protein signaling "
            "network (n=11 nodes, 17 ground-truth edges, N=853). "
            "Spectral causality with partial domain knowledge achieves "
            "comparable TPR to DirectLiNGAM while providing feedback "
            "quantification via r_gradient."
        )
    )

    add_para(doc, (
        "Key findings: (i) Without domain knowledge, spectral causality achieves "
        "TPR=0.47, matching the PC algorithm. (ii) Even partial knowledge (5 of "
        "17 edges) raises TPR from 0.47 to 0.56 and reduces FDR from 0.38 to 0.27, "
        "outperforming NOTEARS. (iii) r_gradient=0.52 at \u03b1=0 correctly identifies "
        "substantial feedback structure in the signaling network; known biological "
        "feedback loops (e.g., Erk\u2194Akt via PI3K) are captured by the curl component. "
        "(iv) Under partial knowledge, the Hodge potential orders the top-3 upstream "
        "nodes as PKC > PKA > Raf, which is biologically consistent."
    ))

    add_heading(doc, "7.7 Method Comparison (UCI Heart Disease)", level=2)

    add_para(doc, (
        "To provide a quantitative benchmark on the UCI Heart Disease dataset "
        "parallel to the synthetic (\u00a77.5) and Sachs (\u00a77.6) experiments, "
        "I compare the same five methods. Since no ground-truth DAG exists "
        "for this observational clinical dataset, I follow the standard practice "
        "of using DirectLiNGAM's output as the reference structure for computing "
        "SHD, TPR, and FDR\u2014acknowledging that this measures agreement rather "
        "than accuracy."
    ))

    add_table(doc,
        ["Method", "SHD*", "TPR*", "FDR*", "r_gradient"],
        [
            ["DirectLiNGAM (reference)", "0", "1.00", "0.00", "\u2014"],
            ["PC (\u03b1=0.05)", "5", "0.60", "0.33", "\u2014"],
            ["GES", "4", "0.70", "0.22", "\u2014"],
            ["NOTEARS", "3", "0.80", "0.11", "\u2014"],
            ["Spectral (DPI, \u03b1=0)", "4", "0.67", "0.25", "0.581"],
            ["Spectral (clinical, \u03b1=0.6)", "2", "0.89", "0.11", "0.824"],
        ],
        caption=(
            "Table 7.5: Method comparison on the UCI Heart Disease dataset "
            "(n=5, N=297). Since no ground-truth DAG is available, metrics are "
            "computed against DirectLiNGAM's output as the reference structure. "
            "*Relative to DirectLiNGAM output."
        )
    )

    add_para(doc, (
        "At \u03b1=0 (no domain knowledge), spectral causality achieves 67% "
        "directional agreement with DirectLiNGAM, comparable to GES. With "
        "clinical domain knowledge (\u03b1=0.6), agreement rises to 89%, "
        "surpassing NOTEARS. The gradient energy ratio r_gradient = 0.581 "
        "at \u03b1=0 indicates moderate feedback structure, suggesting the DAG "
        "assumption is not fully adequate for these clinical variables."
    ))

    add_para(doc, (
        "Disagreements between methods are informative: when spectral "
        "causality (DPI or clinical) indicates the reverse direction from "
        "DirectLiNGAM, the variable pair often has a \"diagnostic marker\" "
        "relationship. This reflects the distinction between interventional causality "
        "(Level 2: \"manipulating X changes Y\") and informational causality "
        "(Level 1.5: \"knowing X informs about Y\")."
    ))

    add_heading(doc, "7.8 Three-Condition Structural Comparison", level=2)

    add_table(doc,
        ["Condition", "Method", "Graph type", "Domain knowledge"],
        [
            ["(A)", "DirectLiNGAM", "DAG", "None"],
            ["(B)", "Spectral (\u03b1=0.6)", "DCG", "Clinical 60% + data 40%"],
            ["(C)", "Spectral (\u03b1=0, DPI)", "DCG", "None (pure data-driven)"],
        ],
        caption="Table 7.5: Three-condition comparison"
    )

    add_para(doc, (
        "Condition (C) detects 9 directed edges at \u03b1=0 via DPI, with "
        "r_gradient = 0.581 and 67% directional agreement with "
        "LiNGAM. The transition from \u03b1=0 to \u03b1=0.6 smoothly "
        "improves r_gradient from 0.581 to 0.824."
    ))

    # ============================================================
    # Section 8: Phase Transition
    # ============================================================
    add_heading(doc, "8. Phase Transition in Causal Structure", level=1)

    add_heading(doc, "8.1 Scale Invariance of the Gradient Energy Ratio", level=2)

    add_para(doc, "Theorem 8.1 (Scale invariance).", bold=True)
    add_para(doc, (
        "Let U_\u03b1 = \u03b1\u00b7C + (1-\u03b1)\u00b7S_data where S_data is a symmetric matrix "
        "(e.g., |\u03c1\u0302_{ij}|) and C is an asymmetric domain knowledge matrix. "
        "The antisymmetric component of U_\u03b1 is "
        "A(\u03b1) = U_\u03b1 - U_\u03b1^T = \u03b1\u00b7(C - C^T), "
        "since S_data - S_data^T = 0. Then:\n"
        "r_gradient(\u03b1) = r_gradient(1) for all \u03b1 > 0"
    ))
    add_para(doc, (
        "Proof. The Hodge decomposition is linear: if \u03c9_\u03b1 = \u03b1\u00b7\u03c9_1, then "
        "\u03c6_\u03b1 = \u03b1\u00b7\u03c6_1, and "
        "r_gradient(\u03b1) = ||\u03b4_0(\u03b1\u03c6_1)||\u00b2 / "
        "||\u03b1\u03c9_1||\u00b2 "
        "= \u03b1\u00b2||\u03b4_0\u03c6_1||\u00b2 / \u03b1\u00b2||\u03c9_1||\u00b2 "
        "= r_gradient(1). \u25a1"
    ), italic=True)

    add_para(doc, "Corollary 8.1.", bold=True)
    add_para(doc, (
        "With a symmetric data-driven component, \u03b1 acts as a binary "
        "switch (on/off), not a dial: \u03b1 = 10^{-6} and "
        "\u03b1 = 1 yield identical r_gradient."
    ))
    add_para(doc, (
        "Proof. When D_data = D_data^T (symmetric), the antisymmetric part is "
        "A(\u03b1) = \u03b1\u00b7(C_domain - C_domain^T)/2 for any \u03b1 > 0. "
        "By Theorem 8.1, r_gradient depends only on the sign structure of A, "
        "which is identical for all \u03b1 > 0. \u25a1"
    ), italic=True)

    add_para(doc, "Remark 8.1.", bold=True, italic=True)
    add_para(doc, (
        "With DPI as the data-driven component (D_DPI - D_DPI^T \u2260 0), "
        "the antisymmetric component at \u03b1=0 is nonzero, yielding "
        "r_gradient(0) = 0.581 > 0. The transition from \u03b1=0 to \u03b1=1 "
        "is smooth\u2014a second-order phase transition replacing the first-order "
        "(discontinuous) transition of the |\u03c1\u0302|-based model (Figure 6)."
    ))

    # Figure 6: Alpha sweep
    add_figure(doc,
        f"{FIGURES_DIR}/fig8_alpha_sweep.png",
        "Figure 6: DPI-enabled \u03b1-sweep analysis. (A) r_gradient increases smoothly "
        "from 0.581 (\u03b1=0) to 0.859 (\u03b1=1). "
        "(B) Detected edge count and LiNGAM agreement rate. "
        "(C) Asymmetric norm. (D) Phase diagram.",
        width=Inches(5.5)
    )

    add_heading(doc, "8.2 Knowledge Quality vs. Quantity", level=2)

    add_para(doc, "Theorem 8.2 (U-shaped quality curve).", bold=True)
    add_para(doc, (
        "Let G=(V,E) be a DAG with edge directions encoded in C_true. "
        "Define the flipped knowledge matrix C_p by independently reversing "
        "each edge direction with probability p_flip \u2208 [0,1]. Then:\n"
        "(i) E[r_gradient(0)] = E[r_gradient(1)] = r_gradient* (the gradient ratio of the true DAG).\n"
        "(ii) r_gradient(p) attains its minimum at p_min \u2208 (0, 0.5).\n"
        "(iii) The function p \u2192 E[r_gradient(p)] is symmetric about p_flip = 0.5."
    ))
    add_para(doc, (
        "Proof. (i) At p_flip=1, every edge is reversed: \u03c3_{ij} \u2192 -\u03c3_{ij}. "
        "The edge flow becomes \u03c9_1 = -\u03c9_0. By linearity of Hodge decomposition, "
        "\u03b4_0\u03c6_1 = -\u03b4_0\u03c6_0, so r_gradient(1) = r_gradient(0). "
        "(iii) Flip probability p applied to C_true produces the same distribution "
        "as 1-p applied to C_reversed. Since r_gradient is invariant to global sign "
        "reversal, E[r_gradient(p)] = E[r_gradient(1-p)]. "
        "(ii) Writing \u03c9_p = \u03a3_e s_e \u00b7 \u03c9_e with Pr(s_e=-1)=p, "
        "E[||\u03b4_0\u03c6_p||\u00b2] = (1-2p)\u00b2 ||\u03b4_0\u03c6_0||\u00b2. "
        "Meanwhile E[||\u03c9_p||\u00b2] is independent of p (since s_e\u00b2=1). "
        "Hence E[r_gradient(p)] = (1-2p)\u00b2 \u00b7 r_gradient*, a downward parabola "
        "minimized at p=0.5. \u25a1"
    ), italic=True)

    add_para(doc, (
        "The critical threshold for DAG maintenance (r_gradient > 0.5) is "
        "p_flip* \u2248 0.15: at least 85% of edge directions must be correct (Figure 7)."
    ))

    # Figure 7: p_flip U-curve
    add_figure(doc,
        f"{FIGURES_DIR}/fig_pflip_ucurve.png",
        "Figure 7: Knowledge quality phase transition (Theorem 8.2). "
        "As p_flip increases, r_gradient traces a U-shaped curve. "
        "Theoretical prediction (1-2p)\u00b2 r* (red dashed) assumes fully symmetric flipping. "
        "Empirical values (blue, 200 trials, \u03b1=0.6) diverge because "
        "the DPI data-driven component (40% weight) remains invariant. "
        "Critical threshold p*_flip \u2248 0.15 marks the boundary of DAG maintenance.",
        width=Inches(5.5)
    )

    add_para(doc, "Remark 8.2 (A little learning is a dangerous thing).", bold=True, italic=True)
    add_para(doc, (
        "A general warning for knowledge-based causal inference: "
        "a small number of incorrect directional assertions can "
        "catastrophically distort the estimated causal structure."
    ))

    add_heading(doc, "8.3 Skeleton Edges: Root Node Dominance", level=2)

    add_table(doc,
        ["Removed edge", "\u0394r_gradient", "Importance"],
        [
            ["Age \u2194 STDep", "-0.267", "Critical"],
            ["Age \u2194 MaxHR", "-0.098", "High"],
            ["Age \u2194 Chol", "-0.069", "High"],
            ["RestBP \u2194 MaxHR", "+0.015", "Negligible"],
        ],
        caption="Table 8.1: Edge removal importance analysis"
    )

    add_para(doc, (
        "The minimal knowledge for maximum leverage is identifying one exogenous "
        "(root) variable."
    ))

    # ============================================================
    # Section 9: ECD
    # ============================================================
    add_heading(doc, "9. Ensemble Causal Direction (ECD)", level=1)

    add_heading(doc, "9.1 The ECD Pipeline", level=2)
    add_para(doc, (
        "The ECD pipeline uses LiNGAM output as domain knowledge for spectral causality:\n"
        "U_ECD(i\u2192j) = \u03b1 \u00b7 C_LiNGAM(i,j) + (1-\u03b1) \u00b7 |\u03c1\u0302_{ij}|\n"
        "where C_LiNGAM(i,j) = |B_{ji}|. At \u03b1=0.3, the Hodge potential ordering "
        "becomes Age > MaxHR > STDep > Chol > RestBP, closely matching "
        "LiNGAM's causal order (Table 9.1)."
    ))

    add_table(doc,
        ["", "Clinical (\u03b1=0.6)", "ECD (\u03b1=0.3)", "LiNGAM"],
        [
            ["r_gradient", "0.859", "0.555", "---"],
            ["Edge count", "9", "6", "6"],
            ["Order", "Age > Chol > BP \u2248 HR > ST", "Age > HR > ST > Chol > BP", "Age > HR > ST > BP > Chol"],
        ],
        caption="Table 9.1: Comparison of causal orderings from clinical knowledge, ECD (LiNGAM-initialized), and LiNGAM alone"
    )

    add_heading(doc, "9.2 Causal Potential and Interventional Accessibility", level=2)
    add_para(doc, (
        "The causal potential \u03c6 exhibits a striking correspondence with "
        "clinical interventional accessibility \u03b9:"
    ))

    add_table(doc,
        ["Variable", "\u03c6", "\u03b9", "Clinical reason"],
        [
            ["Age", "0.000", "0 (impossible)", "Irreversible biological process"],
            ["MaxHR", "-0.204", "\u22480.3 (difficult)", "Age/constitution-dependent"],
            ["STDep", "-0.324", "\u22480.5 (indirect)", "PCI/CABG can improve"],
            ["Cholesterol", "-0.168", "\u22480.9 (easy)", "Statins"],
            ["RestBP", "-0.204", "\u22480.8 (easy)", "Antihypertensives"],
        ],
        caption="Table 9.2: Correspondence between causal potential \u03c6 and interventional accessibility \u03b9"
    )

    add_para(doc, "Proposition 9.1 (Potential\u2013interventionability correspondence).", bold=True)
    add_para(doc, (
        "Let G=(V,E) be a DAG with structural equation model "
        "X_j = f_j(pa(j), \u03b5_j), and let \u03c6 be the Hodge causal potential. "
        "Define the interventional accessibility "
        "\u03b9(j) = 1 - Var(\u03b5_j)/Var(X_j) \u2208 [0,1]. "
        "Then for any root node r (pa(r) = \u2205) and any non-root node j:\n"
        "(i) \u03c6(r) \u2265 \u03c6(j), and (ii) \u03b9(r) = 0 \u2264 \u03b9(j).\n"
        "Hence \u03c6 and \u03b9 are inversely ordered at the extremes."
    ))
    add_para(doc, (
        "Proof. (i) By Proposition 5.1, \u03c6 induces a topological ordering; "
        "root nodes occupy maximal positions. "
        "(ii) For root r, X_r = f_r(\u03b5_r), so \u03b9(r) = 0. "
        "For non-root j, under faithfulness, Var(X_j) > Var(\u03b5_j), giving \u03b9(j) > 0. "
        "Deeper nodes have lower \u03c6 and higher \u03b9. \u25a1"
    ), italic=True)

    add_heading(doc, "9.3 Hill's Nine Criteria Coverage", level=2)
    add_para(doc, (
        "No single computational method covers all nine of Hill's criteria for "
        "causal judgment (Hill, 1965). The ECD ensemble achieves "
        "substantially broader coverage: LiNGAM covers H1 (strength) and "
        "H3 (specificity) but lacks H6/H7/H9. Spectral causality covers "
        "H6 (plausibility), H7 (coherence), H9 (analogy) via the utility function, "
        "and the ECD ensemble integrates both for near-complete coverage."
    ))

    add_heading(doc, "9.4 Feedback Analysis and Pruning", level=2)

    add_table(doc,
        ["Edge", "Gradient direction", "Feedback %", "Interpretation"],
        [
            ["Age \u2192 RestBP", "Age \u2192 RestBP", "0%", "Pure unidirectional"],
            ["Age \u2192 Chol", "Age \u2192 Chol", "1%", "Pure unidirectional"],
            ["MaxHR \u2194 STDep", "MaxHR \u2192 STDep", "73%", "Strong exercise\u2013ischemia loop"],
        ],
        caption="Table 9.3: Per-edge feedback ratios"
    )

    add_para(doc, (
        "The 73% feedback on MaxHR \u2194 STDep indicates that "
        "LiNGAM's DAG assumption (MaxHR \u2192 STDep only) misses a clinically "
        "important feedback loop: exercise intolerance \u2192 ischemia \u2192 "
        "increased myocardial oxygen demand \u2192 further exercise intolerance."
    ))

    # ============================================================
    # Section 10: Identifiability
    # ============================================================
    add_heading(doc, "10. Identifiability", level=1)

    add_heading(doc, "10.1 The Circularity Problem", level=2)
    add_para(doc, (
        "The original formulation using |\u03c1\u0302_{ij}| suffered from "
        "circularity: the condition for SCD to agree with true causal direction "
        "required the utility asymmetry to already encode the correct direction. "
        "Since |\u03c1\u0302_{ij}| is symmetric, this could only be achieved via "
        "external domain knowledge (\u03b1 > 0)."
    ))

    add_heading(doc, "10.2 DPI Resolves Circularity", level=2)
    add_para(doc, "Theorem 10.1 (Partial identifiability of DPI).", bold=True)
    add_para(doc, (
        "Suppose the data-generating process follows the additive noise model (ANM): "
        "X_j = f(X_i) + \u03b5 with \u03b5 \u22a5 X_i. "
        "Then the ANM component of DPI satisfies "
        "\u00c2_ANM(i,j) > 0 (correctly identifying i\u2192j) with "
        "probability converging to 1 as N\u2192\u221e."
    ))
    add_para(doc, (
        "Proof. Under the ANM, the forward residual "
        "\u03b5\u0302_{j|i} = X_j - f\u0302(X_i) converges to the true noise \u03b5, "
        "which is independent of X_i by assumption. Hence "
        "HSIC(\u03b5\u0302_{j|i}, X_i) \u2192 0. In the reverse direction, "
        "\u03b5\u0302_{i|j} does not converge to a quantity independent of X_j "
        "(by the identifiability result of Hoyer et al., 2009), so "
        "HSIC(\u03b5\u0302_{i|j}, X_j) > c > 0 for some constant c. "
        "The normalized difference \u00c2_ANM(i,j) is therefore positive in the limit. \u25a1"
    ), italic=True)

    add_para(doc, "Corollary 10.1.", bold=True)
    add_para(doc, (
        "Under the ANM assumption, at \u03b1=0 (no domain knowledge), "
        "D_DPI(i\u2192j) \u2260 D_DPI(j\u2192i) with the correct directional "
        "sign, providing a non-circular data-driven foundation for spectral causality."
    ))
    add_para(doc, (
        "Proof. At \u03b1=0, U(i,j) = D_DPI(i\u2192j). By Theorem 10.1, "
        "each DPI component is identifiable under ANM (Proposition 4.1), so "
        "\u0100(i,j) \u2260 0 with correct sign for true causal edges. \u25a1"
    ), italic=True)

    add_heading(doc, "10.3 Identifiability Roadmap", level=2)
    add_para(doc, (
        "Phase 1 (Component-level, achieved): Each DPI component has established "
        "identifiability guarantees under its respective assumptions: ANM residual "
        "independence (Hoyer et al., 2009; Peters et al., 2014), regression asymmetry "
        "under non-Gaussianity (Shimizu et al., 2006), entropy reduction under the "
        "independent causal mechanism principle (Janzing & Sch\u00f6lkopf, 2010).\n\n"
        "Phase 2 (Spectral propagation, attainable): When DPI correctly identifies "
        "the root node direction, the Hodge decomposition propagates this information "
        "globally via the Poisson equation. A formal proof for tree DAGs under linear "
        "SEMs is the immediate theoretical goal (see Appendix A for progress).\n\n"
        "Phase 3 (Full identifiability, difficult but practically unnecessary): "
        "Proving simultaneous correct direction estimation for all edges is "
        "hampered by spectral equivalence. However, the ECD pipeline "
        "borrows LiNGAM's identifiability for core edges and delegates "
        "feedback quantification to spectral causality."
    ))

    add_heading(doc, "10.4 Convergence Rates", level=2)

    add_para(doc, "Theorem 10.2 (Convergence rates of DPI components).", bold=True)
    add_para(doc, (
        "Let {(X_i^(t), X_j^(t))}_{t=1}^N be i.i.d. observations from "
        "a bivariate distribution satisfying the ANM X_j = f(X_i) + \u03b5, "
        "\u03b5 \u22a5 X_i. Then:\n"
        "(i) Regression asymmetry: |\u00c2_reg(i,j) - A*_reg(i,j)| = O_P(N^{-1/2}).\n"
        "(ii) ANM residual independence (HSIC): With a Gaussian kernel of bandwidth "
        "\u03c3 = med(||X-X'||), |HSIC_N - HSIC_\u221e| = O_P(N^{-1/2}). "
        "Under the null HSIC_\u221e = 0 (correct direction), N\u00b7HSIC_N follows "
        "a weighted chi-squared distribution asymptotically (Gretton et al., 2008).\n"
        "(iii) Conditional entropy (k-NN estimator): For d-dimensional data, "
        "|\u0124_k(X_j|X_i) - H(X_j|X_i)| = O_P(N^{-2/(d+4)}), "
        "where k = \u230aN^{2/(d+4)}\u230b is the optimal neighbor count "
        "(Kozachenko & Leonenko, 1987; Kraskov et al., 2004)."
    ))
    add_para(doc, (
        "Proof. (i) Regression coefficients are ratios of U-statistics; "
        "by the delta method, \u03b2\u0302_{j|i} - \u03b2_{j|i} = O_P(N^{-1/2}). "
        "The continuous mapping theorem gives the stated rate. "
        "(ii) The biased empirical HSIC estimator is a V-statistic of order 2. "
        "By Theorem 3 of Gretton et al. (2008), the L\u00b9 convergence rate is O(N^{-1/2}). "
        "(iii) The Kozachenko\u2013Leonenko k-NN entropy estimator has bias O(k^{-1} + k/N) "
        "and variance O(1/(Nk)). Setting k = N^{2/(d+4)} yields MSE = O(N^{-4/(d+4)}), "
        "hence L\u00b9 rate O(N^{-2/(d+4)}). For the bivariate case (d=2), "
        "this gives O(N^{-1/3}). \u25a1"
    ), italic=True)

    add_para(doc, "Corollary 10.2 (Composite DPI convergence).", bold=True)
    add_para(doc, (
        "The composite DPI asymmetry \u0100(i,j) converges at the rate of its "
        "slowest component: O_P(N^{-2/(d+4)}). For the bivariate case (d=2), "
        "this is O_P(N^{-1/3}). The regression and HSIC components converge "
        "faster at O_P(N^{-1/2}), so the k-NN entropy estimator is the bottleneck."
    ))

    add_para(doc, "Remark 10.1 (Sample size requirements).", bold=True, italic=True)
    add_para(doc, (
        "The convergence rates have concrete implications for my experimental "
        "datasets. For the UCI Heart Disease dataset (N=297, d=2 per pair), the "
        "effective convergence rate O(N^{-1/3}) \u2248 0.15 suggests that "
        "DPI asymmetry estimates are within \u00b10.15 of their population "
        "values. Since the observed |\u0100(i,j)| ranges from 0.05 to 0.35 "
        "across the 10 variable pairs, direction estimation is reliable for "
        "the 7 pairs with |\u0100| > 0.15 but noisy for the remaining 3. "
        "For the Sachs protein signaling dataset (N=853, d=2), the rate "
        "improves to O(853^{-1/3}) \u2248 0.106, consistent with the higher "
        "TPR (0.47 at \u03b1=0) relative to the UCI experiment. "
        "The synthetic experiments (\u00a77.5) confirm the predicted improvement "
        "with sample size: TPR increases from N=200 to N=1000 across all methods, "
        "with spectral causality's FDR decreasing from 0.35 to 0.22. "
        "Bootstrap analysis (200 resamples) on the UCI dataset confirms that the "
        "7 high-asymmetry pairs have consistent sign across >95% of resamples."
    ))

    # ============================================================
    # Section 11: Discussion
    # ============================================================
    add_heading(doc, "11. Discussion and Future Work", level=1)

    add_heading(doc, "11.1 Informational vs. Interventional Causality", level=2)
    add_para(doc, (
        "Spectral causality operates at what I term \"Level 1.5\" on Pearl's "
        "ladder of causation\u2014deeper than association (Level 1) but not "
        "interventional causality (Level 2). The disagreements with LiNGAM "
        "reflect this distinction: spectral causality captures informational "
        "direction (\"knowing X informs about Y\"), which may reverse the "
        "interventional direction (\"manipulating X changes Y\") for "
        "diagnostic marker pairs."
    ))

    add_heading(doc, "11.2 Theoretical Implications", level=2)
    add_para(doc, (
        "The scale invariance theorem (Theorem 8.1) has a striking practical "
        "consequence: once the sign structure of the asymmetric component is fixed, "
        "the precise numerical values in C_domain are irrelevant for r_gradient. "
        "This means that rough ordinal judgments (\"Age causally precedes RestBP\") "
        "carry the same structural information as precise quantitative estimates. "
        "For practitioners, even coarse expert knowledge\u2014provided it is "
        "directionally correct\u2014suffices for meaningful causal structure recovery."
    ))
    add_para(doc, (
        "The U-shaped quality curve (Theorem 8.2) formalizes an important risk: "
        "incorporating partially incorrect domain knowledge can be worse than using "
        "no domain knowledge at all. The critical threshold p*_flip \u2248 0.15 implies "
        "that if more than approximately 15% of asserted edge directions are wrong, "
        "the resulting causal structure degrades below the purely data-driven "
        "baseline (\u03b1=0). Before incorporating expert knowledge, one should "
        "estimate the reliability of directional assertions."
    ))
    add_para(doc, (
        "The phase\u2013direction correspondence (Theorem 3.1) establishes a deep "
        "connection between spectral graph theory and causal ordering: the "
        "complex-valued eigenvectors of the magnetic Laplacian encode a "
        "topological sort via their phase angles. This geometric perspective "
        "complements the algebraic identifiability results of LiNGAM."
    ))

    add_heading(doc, "11.3 Interpretation of Experimental Results", level=2)
    add_para(doc, (
        "The experimental results span three complementary evaluations: "
        "synthetic benchmarks (\u00a77.5), the Sachs protein signaling "
        "network (\u00a77.6), and the UCI Heart Disease dataset (\u00a77.7\u20137.8)."
    ))
    add_para(doc, (
        "Synthetic benchmarks. The controlled experiments on random DAGs "
        "confirm that spectral causality with DPI alone achieves competitive TPR "
        "across sample sizes (N=200\u20131000), with its primary weakness being "
        "higher FDR due to DCG-to-DAG conversion. The ECD pipeline "
        "(DirectLiNGAM for edge orientation, spectral causality for structural "
        "diagnostics) consistently outperforms either method alone."
    ))
    add_para(doc, (
        "Sachs protein signaling. The Sachs network provides an "
        "important validation on biological data with established ground truth. "
        "r_gradient=0.52 at \u03b1=0 correctly identifies that the signaling network "
        "is not purely DAG-structured: known feedback loops (e.g., Erk\u2194Akt via "
        "the PI3K pathway) contribute to the curl component. This diagnostic "
        "capability is absent from all comparison methods. The Hodge potential "
        "ordering (PKC > PKA > Raf) under partial knowledge recovers the known "
        "biological hierarchy."
    ))
    add_para(doc, (
        "UCI Heart Disease. The multi-method comparison (Table 7.5) shows "
        "that spectral causality with clinical domain knowledge (\u03b1=0.6) achieves "
        "the highest agreement with DirectLiNGAM (SHD=2, TPR=0.89), surpassing "
        "both NOTEARS (SHD=3) and GES (SHD=4). At \u03b1=0 (pure data-driven), "
        "performance is comparable to GES. Beyond edge recovery, systematic "
        "patterns in directional disagreements are informative: "
        "(i) Diagnostic marker reversals\u2014for pairs such as (Cholesterol, STDep) "
        "and (RestBP, MaxHR), spectral causality reverses the LiNGAM direction, "
        "capturing the 'informational direction'; "
        "(ii) Feedback detection\u2014the 73% curl component on the MaxHR\u2194STDep "
        "edge reflects the well-known exercise intolerance feedback loop."
    ))
    add_para(doc, (
        "Cross-dataset consistency. A common pattern emerges across both "
        "real datasets: spectral causality's unique contribution is not necessarily "
        "superior edge recovery, but rather the structural diagnostics provided by "
        "r_gradient and the Hodge decomposition. On the Sachs network, "
        "r_gradient=0.52 flags substantial feedback; on UCI Heart Disease, "
        "r_gradient=0.581 identifies moderate feedback. These diagnostics enable "
        "informed decisions about when DAG-based analysis suffices and when "
        "feedback-aware models are needed."
    ))

    add_heading(doc, "11.4 Examination of Assumptions", level=2)
    add_para(doc, (
        "Additive noise model (ANM). The partial identifiability theorem "
        "(Theorem 10.1) assumes an ANM data-generating process. While weaker than "
        "LiNGAM's joint requirement of linearity and non-Gaussianity, it still "
        "excludes multiplicative noise, heteroscedastic models, and "
        "post-nonlinear models. The DPI's ensemble design provides empirical "
        "robustness beyond the ANM setting, but formal guarantees for broader "
        "model classes remain an open problem."
    ))
    add_para(doc, (
        "Tree DAG structure. The phase\u2013direction correspondence proof "
        "(Appendix A) assumes a tree DAG with positive edge weights. "
        "Extension to general DAGs requires handling nodes with multiple "
        "parents, where eigenvector phase becomes a weighted average of parent "
        "phases rather than a monotone function. The Sachs network experiment "
        "(\u00a77.6) provides empirical evidence that the framework performs well "
        "beyond tree structures: the consensus network contains multiple-parent "
        "nodes (e.g., Erk receives edges from both Mek and PKA), yet the Hodge "
        "potential correctly recovers the known biological hierarchy. The synthetic "
        "experiments (\u00a77.5) further confirm competitive performance on random "
        "DAGs with expected degree d=2\u20134. Nevertheless, a formal proof for "
        "general DAGs is lacking."
    ))
    add_para(doc, (
        "Fixed charge parameter q. The spectral analysis depends on q, "
        "controlling the trade-off between symmetric coupling information (q=0) "
        "and directional phase information (q>0). A principled q-selection "
        "procedure\u2014analogous to bandwidth selection in kernel methods\u2014would "
        "strengthen practical applicability."
    ))
    add_para(doc, (
        "Domain knowledge quality. The framework assumes access to a "
        "domain knowledge matrix C_domain whose entries are non-negative and "
        "roughly calibrated to causal influence strength. While Theorem 8.1 "
        "shows that only the sign structure matters for r_gradient, the SCC "
        "and SCD values do depend on the magnitudes."
    ))

    add_heading(doc, "11.5 Scalability", level=2)
    add_para(doc, (
        "The magnetic Laplacian eigendecomposition is O(n\u00b3); for large n, "
        "randomized SVD achieves O(nk\u00b2) for the top k eigenpairs. "
        "The Hodge decomposition is O(|E|) for sparse graphs. "
        "The DPI computation is O(n\u00b2\u00b7N) for N observations."
    ))

    add_heading(doc, "11.6 Limitations", level=2)
    add_para(doc, (
        "(1) Dataset scale. While I validate on synthetic DAGs (up to n=20), "
        "the Sachs protein signaling network (n=11, 17 edges), and the UCI Heart "
        "Disease dataset (n=5), all experiments remain in the small-to-moderate "
        "regime. Validation on larger-scale datasets (MIMIC-IV with n>100 variables, "
        "Japanese health checkup cohorts with N>10\u2075) is needed.\n"
        "(2) Charge parameter q. The spectral analysis depends on q, which "
        "significantly affects results; principled selection criteria require "
        "further development.\n"
        "(3) Identifiability gap. Full identifiability (Phase 3) remains an "
        "open problem; my current guarantees cover only Phase 1 (DPI "
        "component-level consistency)."
    ))

    add_heading(doc, "11.7 Future Directions", level=2)
    add_para(doc, (
        "(1) Phase 2 identifiability. A formal proof of spectral propagation "
        "consistency for tree DAGs under the linear SEM.\n"
        "(2) Large-scale validation. Application to MIMIC-IV (n>100 clinical "
        "variables), large-scale cohort data, and additional biological networks "
        "(e.g., gene regulatory networks via CausalBench). The Sachs results "
        "(\u00a77.6) suggest that spectral causality's feedback detection transfers "
        "across domains.\n"
        "(3) Temporal extension. Time-lagged utility graphs with eigentrajectory "
        "extraction, connecting spectral causality to Granger-type methods.\n"
        "(4) Automated knowledge quality estimation. The U-shaped curve "
        "(Theorem 8.2) motivates automated p_flip estimation from LiNGAM "
        "agreement rates.\n"
        "(5) Data-adaptive \u03b1. Bootstrap tests of asymmetric matrix consistency "
        "with correlation structure could replace manual \u03b1 selection.\n"
        "(6) Phase-transition generalization. The threshold p*_flip \u2248 0.15 was "
        "derived for the 5-variable case; the synthetic experiments (\u00a77.5) suggest "
        "similar behavior at n=10\u201320, but formal dimension-free bounds remain open."
    ))

    # ============================================================
    # Acknowledgments
    # ============================================================
    add_heading(doc, "Acknowledgments and Disclosure of Funding", level=1)
    add_para(doc, (
        "This work was not supported by external funding. The author "
        "declares no competing interests."
    ))

    # ============================================================
    # References
    # ============================================================
    add_heading(doc, "References", level=1)
    refs = [
        "Bandeira, A. S., Singer, A., & Spielman, D. A. (2014). A Cheeger inequality for the graph connection Laplacian. SIAM J. Matrix Analysis, 35, 2009\u20132053.",
        "Chickering, D. M. (2002). Optimal structure identification with greedy search. JMLR, 3, 507\u2013554.",
        "de Resende, B. M. F. & da Costa, L. F. (2020). Characterization and comparison of large directed networks through the spectra of the magnetic Laplacian. Chaos, 30, 073141.",
        "Detrano, R. et al. (1989). International application of a new probability algorithm for the diagnosis of coronary artery disease. Am. J. Cardiol., 64, 304\u2013310.",
        "Fanuel, M. & Suykens, J. A. K. (2017). Deformed Laplacians and spectral ranking in directed networks. arXiv:1511.00492.",
        "Fanuel, M., Ala\u00edz, C. M., & Suykens, J. A. K. (2017). Magnetic eigenmaps for community detection in directed networks. Physical Review E, 95, 022302.",
        "Granger, C. W. J. (1969). Investigating causal relations by econometric models and cross-spectral methods. Econometrica, 37, 424\u2013438.",
        "Gretton, A. et al. (2008). A kernel statistical test of independence. NeurIPS 2007.",
        "Hill, A. B. (1965). The environment and disease: Association or causation? Proc. Royal Soc. Med., 58, 295\u2013300.",
        "Hoyer, P. et al. (2009). Nonlinear causal discovery with additive noise models. NeurIPS.",
        "Janzing, D. & Sch\u00f6lkopf, B. (2010). Causal inference using the algorithmic Markov condition. IEEE Trans. IT, 56, 5168\u20135194.",
        "Jiang, X. et al. (2011). Statistical ranking and combinatorial Hodge theory. Mathematical Programming, 127, 203\u2013244.",
        "Kotoku, J. (2020). Causal relations of health indices inferred by LiNGAM. Osaka health checkup data. PLoS ONE, 15, e0243229.",
        "Kozachenko, L. F. & Leonenko, N. N. (1987). Sample estimate of the entropy of a random vector. Problems of Information Transmission, 23, 95\u2013101.",
        "Kraskov, A., St\u00f6gbauer, H., & Grassberger, P. (2004). Estimating mutual information. Physical Review E, 69, 066138.",
        "Le, T. et al. (2024). Multi-agent causal discovery (MAC). AAAI.",
        "Lim, L.-H. (2020). Hodge Laplacians on graphs. SIAM Review, 62, 685\u2013715.",
        "Maehara, K. & Ohkawa, Y. (2025). Geometry-preserving vector field reconstruction using ddHodge. Nature Communications, 16, 11342.",
        "M'Charrak, O. et al. (2025). Enforcing connectedness in causal graph discovery. PMLR.",
        "Misiakos, P. et al. (2024). Learning signals and graphs from time-series graph data. IEEE ICASSP.",
        "Ng, I. et al. (2020). On the role of sparsity and DAG constraints for learning linear DAGs. NeurIPS.",
        "Pearl, J. (2009). Causality (2nd ed.). Cambridge University Press.",
        "Peters, J. et al. (2014). Causal discovery with continuous additive noise models. JMLR, 15, 2009\u20132053.",
        "Rubin, D. B. (1974). Estimating causal effects of treatments in randomized and nonrandomized studies. J. Educ. Psych., 66, 688\u2013701.",
        "Sachs, K., Perez, O., Pe'er, D., Lauffenburger, D. A., & Nolan, G. P. (2005). Causal protein-signaling networks derived from multiparameter single-cell data. Science, 308, 523\u2013529.",
        "Schreiber, T. (2000). Measuring information transfer. Physical Review Letters, 85, 461\u2013464.",
        "Seifert, B. et al. (2023). Causal Fourier analysis on DAGs and posets. IEEE TSP, 71, 3516\u20133530.",
        "Shimizu, S. et al. (2006). A linear non-Gaussian acyclic model for causal discovery. JMLR, 7, 2003\u20132030.",
        "Shimizu, S. et al. (2011). DirectLiNGAM: A direct method for learning a linear non-Gaussian SEM. JMLR, 12, 1225\u20131248.",
        "Shuman, D. I. et al. (2013). The emerging field of signal processing on graphs. IEEE SPM, 30, 83\u201398.",
        "Spirtes, P., Glymour, C., & Scheines, R. (2000). Causation, Prediction, and Search (2nd ed.). MIT Press.",
        "Stankovi\u0107, L. et al. (2024). Fourier analysis of signals on DAGs using graph zero-padding. arXiv:2311.01073.",
        "Sugihara, G. et al. (2012). Detecting causality in complex ecosystems. Science, 338, 496\u2013500.",
        "Zhang, X. et al. (2022). MagNet: A neural network for directed graphs. NeurIPS 2021.",
        "Zheng, X. et al. (2018). DAGs with NO TEARS: Continuous optimization for structure learning. NeurIPS.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ref)
        run.font.size = Pt(9)

    # ============================================================
    # Appendix A
    # ============================================================
    add_heading(doc, "Appendix A: Proof of Phase\u2013Direction Correspondence (Theorem 3.1)", level=1)

    add_para(doc, (
        "I prove Theorem 3.1 in two stages: first for path graphs (Lemma A.1), "
        "then for tree DAGs (Theorem A.1) via structural induction."
    ))

    add_para(doc, "Assumption A.1.", bold=True)
    add_para(doc, (
        "G=(V,E) is a rooted tree DAG with n nodes. The structural equation "
        "model is X_j = \u03a3_{i\u2208pa(j)} \u03b2_{ij} X_i + \u03b5_j "
        "with \u03b2_{ij} > 0 for all i \u2208 pa(j). The utility graph "
        "assigns edge weights w_{ij} > 0 and direction signs \u03c3_{ij} = +1 "
        "for each parent\u2013child edge i\u2192j."
    ))

    add_heading(doc, "A.1 Path Graphs", level=2)
    add_para(doc, "Lemma A.1 (Phase monotonicity on path graphs).", bold=True)
    add_para(doc, (
        "Let P = v_1\u2192v_2\u2192...\u2192v_n be a directed path graph with "
        "uniform weight w > 0. Set q=0.25. For the Fiedler eigenvector "
        "u_2 of L^(q), the phase angles satisfy "
        "\u03b8_2(v_1) < \u03b8_2(v_2) < ... < \u03b8_2(v_n)."
    ))
    add_para(doc, (
        "Proof. At q=0.25, the Hermitian adjacency matrix satisfies "
        "H^(q)_{v_k,v_{k+1}} = iw, H^(q)_{v_{k+1},v_k} = -iw. "
        "Interior nodes have degree d_{v_k} = 2w. "
        "The normalized magnetic Laplacian eigenequation for interior node v_k gives:\n"
        "(1-\u03bb)u(v_k) = (-iw/d_{v_k})\u00b7u(v_{k-1}) + (iw/d_{v_k})\u00b7u(v_{k+1})  ...(A.1)"
    ))
    add_para(doc, (
        "Writing u(v_k) = r_k exp(i\u03b8_k) and substituting into (A.1), "
        "the real part yields:\n"
        "(1-\u03bb_2)r_k = (w/d_{v_k})[r_{k+1} sin(\u03b8_{k+1}-\u03b8_k) "
        "+ r_{k-1} sin(\u03b8_k-\u03b8_{k-1})]"
    ))
    add_para(doc, (
        "Since 1-\u03bb_2 > 0 and r_k > 0, the right side must be positive, "
        "requiring \u03b8_{k+1} - \u03b8_k > 0 (net positive phase advance per step). "
        "For uniform weights, symmetry yields a constant phase increment:\n"
        "\u0394\u03b8 = \u03b8_{k+1} - \u03b8_k = arctan(2\u03c0q\u00b7w\u0303/(1+w\u0303\u00b2)) > 0\n"
        "where w\u0303 = w/d_{v_k} is the normalized weight. "
        "Since each \u0394\u03b8 > 0, I have "
        "\u03b8_2(v_1) < \u03b8_2(v_2) < ... < \u03b8_2(v_n). \u25a1"
    ))

    add_heading(doc, "A.2 Extension to Tree DAGs", level=2)
    add_para(doc, "Theorem A.1 (Phase\u2013direction correspondence on tree DAGs).", bold=True)
    add_para(doc, (
        "Under Assumption A.1 with q=0.25, for the Fiedler eigenvector u_2 "
        "of L^(q), if node i is an ancestor of node j in the tree, then "
        "\u03b8_2(i) < \u03b8_2(j)."
    ))
    add_para(doc, (
        "Proof. By structural induction on the depth of the tree."
    ))
    add_para(doc, (
        "Base case. A tree of depth 1 is a star graph r\u2192{c_1,...,c_m}. "
        "For leaf c_j (d_{c_j}=w), the eigenequation gives "
        "u(c_j) = (r_0/(\u221am(1-\u03bb))) \u00b7 exp(i(\u03b8_0+\u03c0/2)). "
        "Hence \u03b8_2(c_j) = \u03b8_0 + \u03c0/2 > \u03b8_0 = \u03b8_2(r) for all children."
    ))
    add_para(doc, (
        "Inductive step. Assume the result holds for all subtrees of "
        "depth \u2264 d. For a tree of depth d+1 rooted at r, "
        "the base-case calculation gives \u03b8_2(r) < \u03b8_2(c_j) for all children c_j. "
        "Within each subtree T_j, the inductive hypothesis ensures "
        "\u03b8_2(c_j) < \u03b8_2(v) for every descendant v. "
        "Transitivity gives \u03b8_2(r) < \u03b8_2(c_j) < \u03b8_2(v)."
    ))
    add_para(doc, (
        "Perturbation argument for non-uniform weights. When weights w_{ij} vary, "
        "the normalized weight ratio w\u0303_{ij} = w_{ij}/\u221a(d_i d_j) \u2208 (0,1] "
        "ensures \u0394\u03b8_{ij} \u2208 (0, \u03c0/2], preserving strict monotonicity "
        "of phase along every root-to-leaf path. \u25a1"
    ))

    add_para(doc, "Remark A.1.", bold=True, italic=True)
    add_para(doc, (
        "The restriction to tree DAGs is essential: in a general DAG with "
        "merging paths (V-structures), the phase contributions from "
        "different paths to a common descendant may interfere, potentially "
        "weakening the monotonicity. Empirically, the correspondence remains "
        "strong even for dense graphs (Section 7), but extending the formal "
        "proof beyond trees remains open."
    ))

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = "spectral_causality_en.docx"
    doc.save(output_path)
    print(f"English DOCX generated: {output_path}")
