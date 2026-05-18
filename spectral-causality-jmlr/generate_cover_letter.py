#!/usr/bin/env python3
"""Generate JMLR cover letter as DOCX (matches official JMLR template)."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date


def build_cover_letter():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # Sender
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Tatsuki Onishi\n")
    p.add_run(
        "Data Science and AI Innovation Research Promotion Center\n"
        "Shiga University, Japan\n"
        "bougtoir@gmail.com\n"
    )
    p.add_run(date.today().strftime("%B %d, %Y"))

    # Recipient
    p = doc.add_paragraph()
    p.add_run("Journal of Machine Learning Research")

    doc.add_paragraph()

    # Opening
    p = doc.add_paragraph()
    p.add_run("Dear Editors:")

    doc.add_paragraph()

    # Summary
    p = doc.add_paragraph()
    p.add_run(
        'I am writing to submit my manuscript '
        '\u201cSpectral Causality: Causal Direction Estimation via '
        'Magnetic Laplacians and Hodge Decomposition\u201d '
        'to the Journal of Machine Learning Research.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(
        "The manuscript introduces spectral causality, a principled framework "
        "that estimates causal directions by exploiting the spectral structure "
        "of the magnetic Laplacian. I define a Directional Predictability Index "
        "(DPI) that resolves circularity in earlier utility-based formulations, "
        "connect the framework to Hodge decomposition to separate DAG-compatible "
        "and feedback components of causal flow, and establish scale-invariance "
        "and phase-transition results. The approach is validated on synthetic DAGs, "
        "the Sachs protein signaling network, and the UCI Heart Disease dataset, "
        "with benchmarks against DirectLiNGAM, PC, GES, and NOTEARS."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(
        "This paper has not been published previously in any journal or "
        "conference proceedings. A preprint may be made available on arXiv "
        "concurrent with this submission."
    )

    doc.add_paragraph()

    # Suggested AEs and reviewers
    p = doc.add_paragraph()
    p.add_run(
        "I suggest the following action editors and referees for this submission."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Action Editors:")
    run.bold = True

    action_editors = [
        "Bryon Aragam, University of Chicago \u2014 causality, graphical models",
        "Elias Bareinboim, Columbia University \u2014 causal inference, generalizability",
        "Silvia Chiappa, DeepMind \u2014 causal inference, variational inference",
        "Kenji Fukumizu, The Institute of Statistical Mathematics, Japan "
        "\u2014 kernel methods, dimension reduction",
        "Mladen Kolar, University of Southern California "
        "\u2014 graphical models, high-dimensional statistics",
    ]
    for ae in action_editors:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(ae)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Reviewers:")
    run.bold = True

    reviewers = [
        "Joris Mooij, University of Amsterdam \u2014 causal discovery, structural causal models",
        "Dominik Janzing, Amazon T\u00fcbingen \u2014 causal inference, "
        "information-theoretic methods",
        "Biwei Huang, UC San Diego \u2014 causal discovery, latent variable models",
        "Murat Kocaoglu, Purdue University \u2014 causal inference, graphical models",
        "Frederick Eberhardt, California Institute of Technology "
        "\u2014 causal inference, Bayesian networks",
    ]
    for rv in reviewers:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rv)

    doc.add_paragraph()

    # Keywords
    p = doc.add_paragraph()
    p.add_run(
        "This submission has the following keywords: "
        "causal discovery, magnetic Laplacian, Hodge decomposition, "
        "spectral graph theory, directed graphs."
    )

    doc.add_paragraph()

    # COI confirmation
    p = doc.add_paragraph()
    p.add_run(
        "As the corresponding (and sole) author, I confirm that I have no "
        "conflict of interest with the action editors and referees suggested above."
    )

    doc.add_paragraph()

    # Disclosure
    p = doc.add_paragraph()
    run = p.add_run("Disclosure of funding and competing interests. ")
    run.bold = True
    p.add_run(
        "No external funding was received for this work. "
        "The author declares no competing interests."
    )

    doc.add_paragraph()

    # Closing
    p = doc.add_paragraph()
    p.add_run("Sincerely,")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(
        "Tatsuki Onishi\n"
        "Data Science and AI Innovation Research Promotion Center\n"
        "Shiga University, Japan"
    )

    out = "cover_letter.docx"
    doc.save(out)
    print(f"Cover letter generated: {out}")


if __name__ == "__main__":
    build_cover_letter()
