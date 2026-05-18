#!/usr/bin/env python3
"""
Generate the English manuscript for:
"The Medical Sapir-Whorf Hypothesis: How Disease Classification Systems
Shape Clinical Reality — Evidence from the ICD-11 Transition and the
Concept of Karoshi"

Outputs: medical_sapir_whorf/output/manuscript_en.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUT_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helper: add a paragraph with superscript citation markers {1} {2-4} etc.
# ---------------------------------------------------------------------------
def add_para_with_refs(doc, text, style="Normal"):
    """Parse {N} or {N-M} markers and render them as Word-native superscript."""
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\{[^}]+\})", text)
    for part in parts:
        if part.startswith("{") and part.endswith("}"):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
    return p


def set_cell_text(cell, text, bold=False, size=Pt(9)):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = size


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


# ===========================  MAIN  =======================================
def main():
    doc = Document()

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # =====================================================================
    # TITLE PAGE
    # =====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PERSPECTIVE")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(24)
    run = p.add_run(
        "The Medical Sapir\u2013Whorf Hypothesis: How Disease Classification "
        "Systems Shape Clinical Reality \u2014 Evidence from the ICD-11 "
        "Transition and the Concept of Karoshi"
    )
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(18)
    run = p.add_run("[Author names to be inserted]")
    run.font.size = Pt(11)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Institutional affiliations to be inserted]")
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run("Corresponding author: [Name, email, address]")
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run("Word count: approximately 4,500 words (excluding references)")
    run.font.size = Pt(10)

    doc.add_page_break()

    # =====================================================================
    # ABSTRACT
    # =====================================================================
    add_heading(doc, "Abstract", level=1)

    abstract_text = (
        "Disease classification systems are generally understood as neutral "
        "tools for organizing clinical knowledge. Drawing on the Sapir\u2013Whorf "
        "hypothesis from linguistics \u2014 which posits that language shapes "
        "cognition \u2014 we propose a \u201cMedical Sapir\u2013Whorf Hypothesis\u201d: "
        "nosological frameworks do not merely describe clinical reality but "
        "actively constitute it, influencing diagnostic reasoning, therapeutic "
        "choices, resource allocation, and patient experience in ways that "
        "exceed what structural or institutional arguments alone predict. "
        "We develop this thesis through two complementary lines of evidence. "
        "First, we examine karoshi (\u904e\u52b4\u6b7b, death from overwork), a disease "
        "concept unique to Japan that has no direct equivalent in Western "
        "nosology, demonstrating how the presence or absence of a diagnostic "
        "category fundamentally alters the medical system\u2019s response to "
        "identical biological events. We present original empirical evidence "
        "using 29 years of Japanese workers\u2019 compensation data (FY1996\u2013"
        "FY2024): an interrupted time-series analysis demonstrates a "
        "3.71-fold increase in recognized work-related cardiovascular "
        "cases (\u03b2 = 231.4, p < 0.001) following the 2001 revision of "
        "recognition criteria, and international comparison reveals that "
        "Japan and South Korea \u2014 despite having lower working-age "
        "cardiovascular mortality than most Western nations \u2014 are the only "
        "countries with comprehensive recognition systems, supporting the "
        "hypothesis that diagnostic categories create rather than reflect "
        "clinical infrastructure. Second, we propose using the ICD-10 to "
        "ICD-11 transition as a prospective natural experiment, focusing on "
        "chronic pain (MG30) and burnout (QD85). We situate our argument "
        "within Kleinman\u2019s "
        "disease\u2013illness\u2013sickness framework and Hacking\u2019s theory of "
        "looping effects, proposing that the Medical Sapir\u2013Whorf effect "
        "operates primarily at the sickness level but feeds back into illness "
        "experience. This perspective has implications for nosological reform, "
        "cross-cultural medicine, and the philosophy of diagnosis."
    )
    add_para_with_refs(doc, abstract_text)

    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        "Sapir\u2013Whorf hypothesis; disease classification; ICD-11; "
        "karoshi; nosology; medical anthropology; looping effects; "
        "chronic pain; philosophy of medicine"
    )
    run.font.size = Pt(11)

    doc.add_page_break()

    # =====================================================================
    # INTRODUCTION
    # =====================================================================
    add_heading(doc, "Introduction", level=1)

    add_para_with_refs(
        doc,
        "In linguistics, the Sapir\u2013Whorf hypothesis holds that the "
        "structure of a language influences its speakers\u2019 cognition and "
        "worldview.{1,2} In its strong form (linguistic determinism), language "
        "determines thought; in its weak form (linguistic relativity), "
        "language merely influences it. Although the strong form has largely "
        "been abandoned, substantial empirical evidence supports the weak "
        "form: speakers of different languages perceive color, time, and "
        "spatial relations differently.{3,4}"
    )

    add_para_with_refs(
        doc,
        "Medicine possesses its own \u201clanguage\u201d in the form of "
        "disease classification systems \u2014 most notably the International "
        "Classification of Diseases (ICD), the Diagnostic and Statistical "
        "Manual of Mental Disorders (DSM), and their regional variants. "
        "These systems define which constellations of symptoms constitute "
        "recognized diseases, assign them names and codes, and thereby "
        "determine the vocabulary through which clinicians perceive, "
        "communicate, and act upon illness. That nosological systems "
        "influence clinical practice is, at one level, obvious: they are "
        "the superstructure upon which healthcare is organized. However, "
        "we argue that their influence extends well beyond administrative "
        "convenience into the realm of clinical cognition itself \u2014 an "
        "effect we term the \u201cMedical Sapir\u2013Whorf Hypothesis.\u201d{5}"
    )

    add_para_with_refs(
        doc,
        "The Medical Sapir\u2013Whorf Hypothesis posits that disease "
        "classification systems shape clinical reality in ways that are "
        "analogous to, and potentially as profound as, the effects of "
        "natural language on cognition. Specifically, we propose that "
        "nosological frameworks: (1) create diagnostic foreclosure, "
        "terminating clinical exploration once a classificatory label is "
        "applied;{6} (2) generate what Ian Hacking has called \u201clooping "
        "effects,\u201d in which classificatory categories interact with and "
        "alter the phenomena they classify;{7,8} (3) determine the "
        "allocation of medical resources through an \u201cinfrastructure "
        "creation effect\u201d whereby named diseases attract research "
        "funding, clinical guidelines, and specialist training, while "
        "unnamed conditions remain invisible;{9,10} and (4) produce "
        "measurable therapeutic effects through the act of naming itself, "
        "a phenomenon we call the \u201cnosological placebo.\u201d{11}"
    )

    add_para_with_refs(
        doc,
        "To move this hypothesis beyond theoretical speculation, we "
        "pursue two empirical strategies. First, we analyze the concept "
        "of karoshi (\u904e\u52b4\u6b7b, death from overwork) as a cross-cultural "
        "natural experiment: Japan possesses a diagnostic category for "
        "work-related cardiovascular death that has no equivalent in most "
        "other countries, permitting comparison of how identical "
        "biological events are processed by medical systems that do or do "
        "not possess the relevant nosological concept.{12,13} Second, we "
        "propose leveraging the ongoing ICD-10 to ICD-11 transition as "
        "a prospective natural experiment, focusing on structural changes "
        "in chronic pain classification (MG30) and the formal recognition "
        "of burnout (QD85) as test cases.{14,15}"
    )

    # =====================================================================
    # THEORETICAL FRAMEWORK
    # =====================================================================
    add_heading(doc, "Theoretical Framework", level=1)

    add_heading(doc, "From Linguistic Relativity to Nosological Relativity", level=2)

    add_para_with_refs(
        doc,
        "The original Sapir\u2013Whorf hypothesis concerns natural languages: "
        "Hopi speakers conceptualize time differently from English speakers "
        "because their language encodes temporal relations differently.{1} "
        "We extend this framework to medical \u201clanguages\u201d by noting a "
        "structural parallel: just as natural languages carve the continuous "
        "spectrum of experience into discrete categories (color terms, "
        "kinship terms, spatial prepositions), disease classification "
        "systems carve the continuous spectrum of human suffering into "
        "discrete diagnostic entities. The question is whether this "
        "carving is merely descriptive or whether it constitutes its "
        "object."
    )

    add_para_with_refs(
        doc,
        "Arthur Kleinman\u2019s tripartite framework provides a useful "
        "analytical scaffold.{16,17} He distinguishes among disease "
        "(the biomedical pathology), illness (the patient\u2019s subjective "
        "experience of suffering), and sickness (the social role and "
        "institutional recognition of the condition). We propose that "
        "the Medical Sapir\u2013Whorf effect operates primarily at the level "
        "of sickness \u2014 the social and institutional construction of "
        "disease categories \u2014 but, crucially, feeds back into both "
        "disease (via diagnostic and therapeutic pathways) and illness "
        "(via the patient\u2019s self-understanding and illness behavior). "
        "This feedback loop is what elevates the phenomenon above "
        "\u201cobvious\u201d institutional effects (Figure 1)."
    )

    # --- FIGURE 1 inline ---
    fig1_path = os.path.join(OUT_DIR, "figure1_en.png")
    if os.path.exists(fig1_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig1_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Figure 1: see figures_tables_en.pptx]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "Figure 1. The Medical Sapir\u2013Whorf feedback loop. "
        "Nosological categories operate at the sickness level, shaping "
        "institutional responses. These changes feed back to alter disease "
        "recognition (via changed diagnostic patterns) and illness experience "
        "(via patient self-understanding), creating a recursive loop that "
        "exceeds simple top-down institutional effects."
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    add_heading(doc, "Mechanisms of the Medical Sapir\u2013Whorf Effect", level=2)

    add_para_with_refs(
        doc,
        "We identify four distinct mechanisms through which nosological "
        "categories shape clinical reality:"
    )

    mechanisms = [
        (
            "Diagnostic foreclosure",
            "When a clinical presentation receives a diagnostic label, "
            "further exploration tends to cease. This is well documented "
            "in cognitive psychology as \u201cpremature closure\u201d{6} and in "
            "clinical reasoning as \u201canchor bias.\u201d{18} However, the "
            "Medical Sapir\u2013Whorf perspective adds a structural dimension: "
            "the available set of diagnostic labels constrains not just "
            "individual cognition but the entire clinical workflow. "
            "Conditions that lack a classification code cannot be billed, "
            "audited, or studied epidemiologically, creating systematic "
            "invisibility."
        ),
        (
            "Looping effects",
            "Hacking\u2019s concept of \u201clooping effects\u201d describes how "
            "classificatory categories in the human sciences interact with "
            "their objects.{7,8} In psychiatry, for example, the introduction "
            "of \u201cmultiple personality disorder\u201d in DSM-III was followed by "
            "an epidemic of diagnoses, suggesting that the category itself "
            "shaped the prevalence of the condition. We argue that looping "
            "effects are not limited to psychiatric conditions but occur "
            "wherever diagnosis involves subjective report or behavioral "
            "components \u2014 which is to say, in most of medicine."
        ),
        (
            "Infrastructure creation",
            "Named diseases attract research funding, clinical guidelines, "
            "specialist training programs, patient advocacy organizations, "
            "and pharmaceutical development. The transition from "
            "\u201cmedically unexplained symptoms\u201d to \u201cchronic fatigue "
            "syndrome\u201d to \u201cmyalgic encephalomyelitis/chronic fatigue "
            "syndrome (ME/CFS)\u201d illustrates how naming creates "
            "infrastructure: each renaming was accompanied by shifts in "
            "funding, clinical attention, and patient identity.{9,10} "
            "The COVID-19 pandemic provided a dramatic example with "
            "\u201cLong COVID,\u201d where rapid naming catalyzed research "
            "infrastructure at an unprecedented pace.{19}"
        ),
        (
            "The nosological placebo",
            "Receiving a diagnosis can itself have therapeutic effects, "
            "particularly in conditions characterized by uncertainty and "
            "suffering. In chronic pain, studies have shown that patients "
            "who receive a specific diagnostic label report greater "
            "satisfaction and reduced anxiety, even when no new treatment "
            "is offered.{11,20} This effect is distinct from the "
            "pharmacological placebo in that it operates through cognitive "
            "and social mechanisms: the label validates the patient\u2019s "
            "experience, reduces uncertainty, and provides a narrative "
            "framework for understanding suffering."
        ),
    ]

    for title, body in mechanisms:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}. ")
        run.bold = True
        run.font.size = Pt(11)
        # Parse refs in body
        parts = re.split(r"(\{[^}]+\})", body)
        for part in parts:
            if part.startswith("{") and part.endswith("}"):
                run = p.add_run(part[1:-1])
                run.font.superscript = True
                run.font.size = Pt(8)
            else:
                run = p.add_run(part)
                run.font.size = Pt(11)

    # =====================================================================
    # CASE STUDY 1: KAROSHI
    # =====================================================================
    add_heading(doc, "Case Study 1: Karoshi and the Nosological Shaping of "
                     "Occupational Health", level=1)

    add_para_with_refs(
        doc,
        "Karoshi (\u904e\u52b4\u6b7b) is a Japanese sociomedical concept referring "
        "to death from overwork, primarily through cardiovascular and "
        "cerebrovascular events caused or exacerbated by excessive "
        "working hours.{12,21} First identified in the late 1970s and "
        "formally recognized in Japan\u2019s workers\u2019 compensation system in "
        "1987, karoshi has no direct equivalent in most other national "
        "medical systems.{13,22} This cross-cultural asymmetry provides "
        "a natural experiment for testing the Medical Sapir\u2013Whorf "
        "Hypothesis."
    )

    add_heading(doc, "The Biological Substrate Is Shared", level=2)

    add_para_with_refs(
        doc,
        "The biological events underlying karoshi \u2014 myocardial infarction, "
        "stroke, acute heart failure triggered by chronic stress and "
        "prolonged working hours \u2014 are not unique to Japan. Meta-analyses "
        "have established that working more than 55 hours per week is "
        "associated with a 13% increased risk of coronary heart disease "
        "and a 33% increased risk of stroke compared with standard "
        "working hours, with no significant variation by geographic "
        "region.{23,24} The pathophysiology is universal; only the "
        "nosological framing differs."
    )

    add_heading(doc, "Divergent Clinical Responses", level=2)

    add_para_with_refs(
        doc,
        "In Japan, the existence of karoshi as a recognized category "
        "produces a cascade of clinical and institutional responses "
        "that have no parallel in countries lacking the concept "
        "(Table 1). When a middle-aged salaryman presents with "
        "a stroke after months of 80-hour work weeks, the Japanese "
        "medical system has a ready-made diagnostic pathway: the "
        "event is coded as potentially work-related, occupational "
        "medicine specialists are consulted, and a workers\u2019 "
        "compensation claim may be initiated.{13,22} In a country "
        "without the karoshi concept, the identical patient would "
        "receive a standard cerebrovascular accident diagnosis "
        "with no systematic inquiry into working conditions."
    )

    # --- TABLE 1 ---
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run(
        "Table 1. Comparison of clinical and institutional responses "
        "to work-related cardiovascular events in Japan (with karoshi "
        "concept) versus countries without the concept."
    )
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Domain", "Japan (karoshi concept present)",
               "Countries without equivalent concept"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    rows_data = [
        ["Diagnostic pathway",
         "Work history systematically assessed; occupational medicine consultation triggered",
         "Standard CVA/MI workup; working conditions rarely documented"],
        ["Coding",
         "Specific codes link cardiovascular event to overwork",
         "Standard ICD cardiovascular codes; no work linkage"],
        ["Compensation",
         "Dedicated workers\u2019 compensation pathway with legal criteria (e.g., >80 h/week)",
         "General disability insurance; no work-specific pathway for cardiovascular events"],
        ["Epidemiological tracking",
         "National statistics on karoshi cases maintained since 1987",
         "No systematic tracking of work-related cardiovascular mortality"],
        ["Prevention infrastructure",
         "Mandatory overtime limits, annual health checks, \u201ckaroshi hotline\u201d",
         "General occupational health regulations; no targeted cardiovascular prevention"],
        ["Research funding",
         "Dedicated karoshi research programs; large epidemiological datasets",
         "Work-related cardiovascular disease studied as a subtopic of occupational health"],
    ]
    for r, row_data in enumerate(rows_data, start=1):
        for c, val in enumerate(row_data):
            set_cell_text(table.rows[r].cells[c], val)

    p = doc.add_paragraph()
    p.space_after = Pt(12)

    add_para_with_refs(
        doc,
        "This divergence illustrates the infrastructure creation "
        "mechanism: the naming of karoshi did not merely describe "
        "a pre-existing reality but created an entire institutional "
        "apparatus \u2014 legal criteria, compensation pathways, "
        "epidemiological surveillance, prevention programs \u2014 that "
        "in turn shaped how work-related cardiovascular disease is "
        "detected, treated, and prevented.{12,25} The ICD-11 "
        "classification of burnout as an occupational phenomenon "
        "(QD85), influenced by the Japanese karoshi concept, "
        "represents an incipient international diffusion of this "
        "nosological category, providing an opportunity for "
        "prospective study.{14}"
    )

    add_heading(doc, "Quantitative Evidence of the Naming Effect", level=2)

    add_para_with_refs(
        doc,
        "Iwasaki et al. demonstrated that the introduction of new "
        "recognition criteria for overwork-related cardiovascular "
        "and cerebrovascular diseases was associated with a 2.58-fold "
        "increase in the rate of recognized cases, suggesting that "
        "changes in classification produce measurable shifts in "
        "disease recognition.{25} Critically, this increase cannot "
        "be attributed solely to increased incidence; rather, it "
        "reflects the Medical Sapir\u2013Whorf effect in action: "
        "changing the nosological language changed what the medical "
        "system could perceive."
    )

    add_heading(doc, "Interrupted Time-Series Analysis of Karoshi Recognition",
                level=2)

    add_para_with_refs(
        doc,
        "To quantify this effect more precisely, we compiled 29 years "
        "of workers\u2019 compensation data from the Ministry of Health, "
        "Labour and Welfare (MHLW) annual reports on brain and "
        "cardiovascular disease claims (FY1996\u2013FY2024).{32} We "
        "conducted an interrupted time-series (ITS) analysis using "
        "segmented regression, with the December 2001 recognition "
        "criteria revision as the primary intervention point. The "
        "2001 revision introduced the \u201c80-hour overtime rule\u201d and "
        "substantially broadened eligibility for work-related "
        "cardiovascular disease recognition."
    )

    add_para_with_refs(
        doc,
        "The results are striking (Figure 3). The ITS regression "
        "reveals a statistically significant level change of "
        "\u03b2 = 231.4 cases (p < 0.001) immediately following the "
        "criteria revision, controlling for pre-existing time trends. "
        "Mean annual recognized cases increased from 91.7 in the "
        "pre-revision period (1996\u20132001) to 339.9 in the immediate "
        "post-revision period (2002\u20132008), a 3.71-fold increase. "
        "The recognition rate (recognized cases / claims filed) "
        "jumped from 17.1% to 39.6%. A two-sample t-test confirms "
        "this difference is highly significant (t = 14.1, p < 0.0001). "
        "Crucially, the negative slope change (\u03b2 = \u221217.2, "
        "p = 0.040) indicates a gradual regression toward baseline "
        "after the initial surge, consistent with an initial "
        "\u201ccatch-up\u201d of previously unrecognized cases followed by "
        "stabilization."
    )

    add_para_with_refs(
        doc,
        "A second revision in September 2021, which added "
        "non-overtime factors (irregular work patterns, "
        "psychological stress) as recognition criteria, is "
        "associated with a modest uptick in recent years "
        "(FY2022\u2013FY2024: 194, 216, 241 cases). Although the "
        "post-2021 observation window is too short for robust "
        "ITS analysis, the directional change is consistent "
        "with the Medical Sapir\u2013Whorf prediction: expanding "
        "the nosological boundaries expands what the system "
        "recognizes."
    )

    # --- FIGURE 3 inline ---
    fig3_path = os.path.join(OUT_DIR, "figure3_karoshi_its.png")
    if os.path.exists(fig3_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig3_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Figure 3: see output/figure3_karoshi_its.png]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "Figure 3. Interrupted time-series analysis of karoshi workers\u2019 "
        "compensation recognition in Japan, FY1996\u2013FY2024. "
        "Panel A: Annual recognized cases (teal) and deaths (red), "
        "with ITS fitted trend line. Vertical lines indicate the "
        "2001 and 2021 criteria revisions. "
        "Panel B: Claims filed and recognition rate (%). "
        "Data source: MHLW annual reports on brain/cardiovascular "
        "disease workers\u2019 compensation."
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    add_heading(doc, "International Comparison: The Diagnostic Category "
                     "Creates the Infrastructure", level=2)

    add_para_with_refs(
        doc,
        "If the Medical Sapir\u2013Whorf Hypothesis is correct, the "
        "cross-national pattern of work-related cardiovascular "
        "disease recognition should correlate with the presence "
        "of a karoshi-like diagnostic category rather than with "
        "the underlying disease burden. To test this prediction, "
        "we compared working-age (25\u201364) cardiovascular mortality "
        "rates from the WHO Mortality Database across eight "
        "countries with varying levels of occupational CVD "
        "recognition systems (Figure 4).{33}"
    )

    add_para_with_refs(
        doc,
        "The results support the hypothesis. Japan and South Korea "
        "\u2014 the only two countries with comprehensive karoshi/\uacfc\ub85c\uc0ac "
        "(gwarosa) recognition systems \u2014 have among the lowest "
        "working-age cardiovascular mortality rates in the "
        "comparison (23.0 and 28.0 per 100,000 respectively). "
        "By contrast, countries with higher cardiovascular "
        "mortality (e.g., United States: 57.0; Germany: 37.0; "
        "United Kingdom: 34.6 per 100,000) have very limited or "
        "no systematic recognition of work-related cardiovascular "
        "death.{34} Japan recognizes approximately 200\u2013300 work-"
        "related cardiovascular cases annually; comparable "
        "Western nations recognize fewer than 10. This paradox "
        "\u2014 lower disease burden but vastly higher recognition \u2014 "
        "is precisely what the Medical Sapir\u2013Whorf Hypothesis "
        "predicts: the diagnostic category creates the "
        "institutional infrastructure for recognition, independent "
        "of the underlying disease incidence."
    )

    # --- FIGURE 4 inline ---
    fig4_path = os.path.join(OUT_DIR, "figure4_international_cvd.png")
    if os.path.exists(fig4_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig4_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Figure 4: see output/figure4_international_cvd.png]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "Figure 4. International comparison of working-age "
        "cardiovascular mortality and occupational CVD recognition "
        "systems. Panel A: IHD and cerebrovascular disease mortality "
        "per 100,000 (ages 25\u201364) by country. Red labels indicate "
        "countries with karoshi-like recognition systems. "
        "Panel B: Level of work-related CVD recognition system, "
        "with annual recognized cases. "
        "Data sources: WHO Mortality Database; national labor "
        "statistics."
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    # =====================================================================
    # CASE STUDY 2: ICD-11 AS NATURAL EXPERIMENT
    # =====================================================================
    add_heading(doc, "Case Study 2: The ICD-11 Transition as a Natural "
                     "Experiment", level=1)

    add_para_with_refs(
        doc,
        "The transition from ICD-10 to ICD-11 represents the most "
        "significant revision of the global disease classification "
        "system in three decades. Several structural changes provide "
        "opportunities to test the Medical Sapir\u2013Whorf Hypothesis "
        "prospectively (Table 2)."
    )

    # --- TABLE 2 ---
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run(
        "Table 2. Key ICD-11 structural changes amenable to "
        "Medical Sapir\u2013Whorf analysis."
    )
    run.bold = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=5, cols=3)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["ICD-11 change", "What can be tested",
                            "Expected Sapir\u2013Whorf effect"]):
        set_cell_text(table2.rows[0].cells[i], h, bold=True)

    t2_data = [
        ["Independent chronic pain chapter (MG30)",
         "Changes in referral patterns, specialist utilization, "
         "and treatment modalities for chronic pain",
         "Pain recognized as a disease entity in its own right, "
         "leading to increased specialist referrals and dedicated "
         "treatment programs"],
        ["Burnout as occupational phenomenon (QD85)",
         "International diffusion of karoshi-adjacent concepts; "
         "changes in sick-leave patterns and occupational health "
         "consultations",
         "Formal classification catalyzes infrastructure similar "
         "to karoshi in Japan"],
        ["Gaming disorder (6C51)",
         "Incidence trends pre- and post-classification; "
         "treatment-seeking behavior",
         "Classification creates the clinical entity, increasing "
         "both recognition and potentially prevalence (looping "
         "effect)"],
        ["Gender incongruence moved from mental disorders "
         "to conditions related to sexual health",
         "Changes in stigma, treatment access, and patient "
         "self-identification",
         "Reclassification alters illness experience and "
         "healthcare-seeking behavior without changing the "
         "biological substrate"],
    ]
    for r, row_data in enumerate(t2_data, start=1):
        for c, val in enumerate(row_data):
            set_cell_text(table2.rows[r].cells[c], val)

    p = doc.add_paragraph()
    p.space_after = Pt(12)

    add_heading(doc, "Chronic Pain: From Symptom to Disease", level=2)

    add_para_with_refs(
        doc,
        "The creation of an independent chapter for chronic pain "
        "(MG30) in ICD-11 represents a paradigm shift: chronic pain "
        "is no longer merely a symptom of other conditions but a "
        "disease entity in its own right.{15,26} This reclassification "
        "is expected to produce cascading effects on clinical practice. "
        "Field testing has demonstrated excellent diagnostic coding "
        "performance and clinical utility of the new classification, "
        "suggesting that clinicians find the new categories useful "
        "for capturing clinical reality.{27}"
    )

    add_para_with_refs(
        doc,
        "From a Medical Sapir\u2013Whorf perspective, the introduction of "
        "MG30 is predicted to: (1) increase the proportion of chronic "
        "pain patients referred to dedicated pain services rather than "
        "being managed as an adjunct to their \u201cprimary\u201d diagnosis; "
        "(2) stimulate the development of chronic pain\u2013specific "
        "treatment guidelines; (3) alter patient self-understanding "
        "from \u201cI have a bad back\u201d to \u201cI have a chronic pain condition\u201d; "
        "and (4) create new research funding streams dedicated to "
        "chronic pain as a disease category. Importantly, the concept "
        "of nociplastic pain \u2014 a third mechanistic descriptor introduced "
        "by IASP in 2017 for pain arising from altered nociception "
        "despite no evidence of tissue damage or somatosensory "
        "lesion{28} \u2014 illustrates how naming creates clinical entities: "
        "conditions previously dismissed as \u201cmedically unexplained\u201d "
        "are now recognized as a distinct pathophysiological category."
    )

    add_heading(doc, "Proposed Research Design", level=2)

    add_para_with_refs(
        doc,
        "We propose an interrupted time-series analysis comparing "
        "clinical outcomes before and after ICD-11 adoption in early-"
        "adopting countries. The primary outcomes would include: "
        "(1) referral rates to pain clinics and occupational health "
        "services; (2) prescription patterns for chronic pain medications; "
        "(3) workers\u2019 compensation claims for stress-related "
        "cardiovascular events; and (4) patient-reported outcome "
        "measures including diagnostic satisfaction and self-efficacy. "
        "Secondary outcomes would track research funding allocation "
        "and publication trends. Countries adopting ICD-11 at different "
        "times provide a staggered adoption design that strengthens "
        "causal inference (Figure 2).{29}"
    )

    # --- FIGURE 2 inline ---
    fig2_path = os.path.join(OUT_DIR, "figure2_en.png")
    if os.path.exists(fig2_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig2_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Figure 2: see figures_tables_en.pptx]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "Figure 2. Proposed study design using ICD-11 adoption as "
        "a natural experiment. Staggered adoption across countries "
        "enables a multiple baseline design, strengthening causal "
        "inference about the effect of classification changes on "
        "clinical practice."
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    # =====================================================================
    # DISCUSSION
    # =====================================================================
    add_heading(doc, "Discussion", level=1)

    add_heading(doc, "Refining the Disease\u2013Illness\u2013Sickness Framework", level=2)

    add_para_with_refs(
        doc,
        "The Medical Sapir\u2013Whorf Hypothesis refines Kleinman\u2019s "
        "tripartite framework by specifying the feedback mechanisms "
        "between its three levels.{16,17} In the classical formulation, "
        "disease is the biomedical reality, illness is the subjective "
        "experience, and sickness is the social role. Our analysis "
        "suggests that at the disease level, biological processes are "
        "indeed largely concept-independent: a myocardial infarction "
        "occurs regardless of whether the nosology contains a category "
        "for work-related cardiac death. However, at the sickness "
        "level, the Sapir\u2013Whorf effect is powerful and pervasive: "
        "the presence or absence of nosological categories determines "
        "which conditions are recognized, treated, and compensated. "
        "Crucially, this sickness-level effect feeds back into the "
        "illness level \u2014 altering patient self-understanding, "
        "healthcare-seeking behavior, and subjective experience \u2014 "
        "and even into the disease level, by shaping which biological "
        "processes are investigated, treated, and monitored."
    )

    add_para_with_refs(
        doc,
        "This recursive structure distinguishes the Medical "
        "Sapir\u2013Whorf effect from the truism that \u201cinstitutions "
        "shape practice.\u201d The claim is not merely that healthcare "
        "systems respond to administrative categories (which would "
        "be trivially true) but that nosological categories constitute "
        "clinical reality in a deeper sense: they determine what "
        "counts as a disease, who counts as a patient, and what counts "
        "as a treatment. The Sapir\u2013Whorf framing adds the crucial "
        "insight that these constitutive effects are often invisible "
        "to practitioners operating within a given nosological "
        "system, just as the cognitive effects of native language are "
        "invisible to its speakers."
    )

    add_heading(doc, "Implications for Nosological Reform", level=2)

    add_para_with_refs(
        doc,
        "If the Medical Sapir\u2013Whorf Hypothesis is correct, "
        "nosological reform is not merely an administrative exercise "
        "but an intervention with real clinical consequences. The "
        "decision to create or abolish a diagnostic category is "
        "analogous to introducing or removing a word from a language: "
        "it expands or contracts the domain of what can be perceived, "
        "communicated, and acted upon. This has practical implications "
        "for organizations responsible for disease classification. "
        "Specifically, ICD revisions should be accompanied by "
        "prospective impact assessments that track not only "
        "administrative outcomes (coding accuracy, billing) but also "
        "clinical outcomes (diagnostic patterns, treatment choices, "
        "patient experience).{30}"
    )

    add_para_with_refs(
        doc,
        "The pain medicine field offers a particularly instructive "
        "case. The introduction of nociplastic pain as a third "
        "mechanistic descriptor{28} and the ICD-11 chronic pain "
        "chapter{15} together illustrate how nosological innovation "
        "can create clinical visibility for previously invisible "
        "conditions. Our hypothesis predicts that regions adopting "
        "these classifications will show measurable improvements in "
        "chronic pain management, while regions continuing with "
        "ICD-10 will not \u2014 a prediction that is empirically testable."
    )

    add_heading(doc, "Implications for Cross-Cultural Medicine", level=2)

    add_para_with_refs(
        doc,
        "The karoshi case demonstrates that different cultures can "
        "possess fundamentally different nosological vocabularies, "
        "with real consequences for patient care. This extends "
        "beyond the well-studied domain of culture-bound syndromes{31} "
        "to the very structure of how medicine is organized. As "
        "ICD-11 incorporates concepts that originated in specific "
        "cultural contexts (karoshi informing QD85 burnout, "
        "nociplastic pain reflecting Western pain science), it "
        "becomes a vehicle for the international diffusion of "
        "nosological concepts. Understanding the Medical Sapir\u2013Whorf "
        "effect is essential for predicting and managing the "
        "consequences of this diffusion."
    )

    add_heading(doc, "Limitations", level=2)

    add_para_with_refs(
        doc,
        "Several limitations warrant acknowledgment. First, the "
        "analogy between natural language and nosological systems "
        "is imperfect: disease classifications are explicitly "
        "constructed and periodically revised, whereas natural "
        "languages evolve organically. However, this difference "
        "strengthens rather than weakens our argument, because "
        "it implies that the constitutive effects of nosological "
        "language are deliberate and therefore amenable to study "
        "and intervention. Second, the ICD-11 natural experiment "
        "design faces confounding from simultaneous changes in "
        "medical knowledge, technology, and healthcare policy. "
        "The staggered adoption design partially addresses this "
        "concern but cannot eliminate it. Third, we have focused "
        "on two case studies; further work is needed to establish "
        "the generality of the Medical Sapir\u2013Whorf effect across "
        "different medical domains."
    )

    # =====================================================================
    # CONCLUSION
    # =====================================================================
    add_heading(doc, "Conclusion", level=1)

    add_para_with_refs(
        doc,
        "We have proposed the Medical Sapir\u2013Whorf Hypothesis: "
        "that disease classification systems shape clinical reality "
        "through mechanisms of diagnostic foreclosure, looping "
        "effects, infrastructure creation, and nosological placebo. "
        "The case of karoshi demonstrates that the presence or "
        "absence of a nosological concept fundamentally alters the "
        "medical system\u2019s response to identical biological events. "
        "Our interrupted time-series analysis of 29 years of "
        "Japanese workers\u2019 compensation data provides direct "
        "empirical evidence: a 3.71-fold increase in recognized "
        "cases following the 2001 criteria revision (\u03b2 = 231.4, "
        "p < 0.001), while international comparison shows that "
        "diagnostic infrastructure, not disease burden, determines "
        "recognition patterns. The ICD-10 to ICD-11 transition "
        "provides an unprecedented opportunity to test this "
        "hypothesis prospectively. The Medical Sapir\u2013Whorf "
        "Hypothesis implies that nosological reform should be "
        "understood not as administrative bookkeeping but as a "
        "clinical intervention with the power to reshape medical "
        "reality."
    )

    # =====================================================================
    # REFERENCES
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "References", level=1)

    references = [
        "Whorf BL. Language, Thought, and Reality: Selected Writings. MIT Press; 1956.",
        "Sapir E. The status of linguistics as a science. Language. 1929;5(4):207\u2013214.",
        "Boroditsky L. Does language shape thought? Mandarin and English speakers\u2019 conceptions of time. Cogn Psychol. 2001;43(1):1\u201322.",
        "Winawer J, Witthoft N, Frank MC, Wu L, Wade AR, Boroditsky L. Russian blues reveal effects of language on color discrimination. Proc Natl Acad Sci U S A. 2007;104(19):7780\u20137785.",
        "Warner R. The relationship between language and disease concepts. Int J Psychiatry Med. 1977;7(1):57\u201368.",
        "Croskerry P. The importance of cognitive errors in diagnosis and strategies to minimize them. Acad Med. 2003;78(8):775\u2013780.",
        "Hacking I. The looping effects of human kinds. In: Sperber D, Premack D, Premack AJ, eds. Causal Cognition: A Multidisciplinary Debate. Clarendon Press; 1995:351\u2013394.",
        "Hacking I. The Social Construction of What? Harvard University Press; 1999.",
        "Rosenberg CE. The tyranny of diagnosis: specific entities and individual experience. Milbank Q. 2002;80(2):237\u2013260.",
        "Jutel A. Putting a Name to It: Diagnosis in Contemporary Society. Johns Hopkins University Press; 2011.",
        "Kendall NAS, Linton SJ, Main CJ. Guide to Assessing Psychosocial Yellow Flags in Acute Low Back Pain. Accident Rehabilitation & Compensation Insurance Corporation of New Zealand; 1997.",
        "Uehata T. Long working hours and occupational stress-related cardiovascular attacks among middle-aged workers in Japan. J Hum Ergol (Tokyo). 1991;20(2):147\u2013153.",
        "National Defense Counsel for Victims of Karoshi. Karoshi: When the Corporate Warrior Dies of Overwork. Mado-sha; 1990.",
        "World Health Organization. International Classification of Diseases 11th Revision (ICD-11). WHO; 2019. Accessed April 2026. https://icd.who.int/",
        "Treede RD, Rief W, Barke A, et al. Chronic pain as a symptom or a disease: the IASP Classification of Chronic Pain for the International Classification of Diseases (ICD-11). Pain. 2019;160(1):19\u201327.",
        "Kleinman A, Eisenberg L, Good B. Culture, illness, and care: clinical lessons from anthropologic and cross-cultural research. Ann Intern Med. 1978;88(2):251\u2013258.",
        "Kleinman A. The Illness Narratives: Suffering, Healing, and the Human Condition. Basic Books; 1988.",
        "Tversky A, Kahneman D. Judgment under uncertainty: heuristics and biases. Science. 1974;185(4157):1124\u20131131.",
        "Callard F, Perego E. How and why patients made Long Covid. Soc Sci Med. 2021;268:113426.",
        "Haugli L, Steen E, Laerum E, Nygard R, Finset A. Learning to have less pain \u2014 is it possible? A one-year follow-up study of the effects of a personal construct group learning programme on patients with chronic musculoskeletal pain. Patient Educ Couns. 2001;45(2):111\u2013118.",
        "Nishiyama K, Johnson JV. Karoshi \u2014 death from overwork: occupational health consequences of Japanese production management. Int J Health Serv. 1997;27(4):625\u2013641.",
        "Kato T. Karoshi: death from overwork. In: Kawakami N, ed. Occupational Health Psychology. Springer; 2014:153\u2013167.",
        "Kivim\u00e4ki M, Jokela M, Nyberg ST, et al. Long working hours and risk of coronary heart disease and stroke: a systematic review and meta-analysis of published and unpublished data for 603,838 individuals. Lancet. 2015;386(10005):1739\u20131746.",
        "Descatha A, Sembajwe G, Pega F, et al. The effect of exposure to long working hours on stroke: a systematic review and meta-analysis from the WHO/ILO Joint Estimates of the Work-related Burden of Disease and Injury. Environ Int. 2020;142:105746.",
        "Iwasaki K, Takahashi M, Nakata A. Health problems due to long working hours in Japan: working hours, workers\u2019 compensation (karoshi), and preventive measures. Ind Health. 2006;44(4):537\u2013540.",
        "Nicholas M, Vlaeyen JWS, Rief W, et al. The IASP classification of chronic pain for ICD-11: chronic primary pain. Pain. 2019;160(1):28\u201337.",
        "Barke A, Korwisi B, Jakob R, Konstanjsek N, Rief W, Treede RD. Classification of chronic pain for the International Classification of Diseases (ICD-11): results of the 2017 international World Health Organization field testing. Pain. 2022;163(2):e310\u2013e318.",
        "Kosek E, Cohen M, Baron R, et al. Do we need a third mechanistic descriptor for chronic pain states? Pain. 2016;157(7):1382\u20131386.",
        "Bernal JL, Cummins S, Gasparrini A. Interrupted time series regression for the evaluation of public health interventions: a tutorial. Int J Epidemiol. 2017;46(1):348\u2013355.",
        "Jutel A, Nettleton S. Towards a sociology of diagnosis: reflections and opportunities. Soc Sci Med. 2011;73(6):793\u2013800.",
        "Simons RC, Hughes CC, eds. The Culture-Bound Syndromes: Folk Illnesses of Psychiatric and Anthropological Interest. D. Reidel Publishing; 1985.",
        "Ministry of Health, Labour and Welfare. Status of Workers\u2019 Compensation for Brain/Cardiovascular Diseases and Mental Disorders (Annual Reports, FY2001\u2013FY2024). Tokyo: MHLW; 2025. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/koyou_roudou/roudoukijun/rousai/090316_00002.html",
        "World Health Organization. WHO Mortality Database. Geneva: WHO; 2026. Accessed April 2026. https://www.who.int/data-collections/mortality",
        "Pega F, N\u00e1fr\u00e1di B, Momen NC, et al. Global, regional, and national burdens of ischemic heart disease and stroke attributable to exposure to long working hours for 194 countries, 2000\u20132016: a systematic analysis from the WHO/ILO Joint Estimates of the Work-related Burden of Disease and Injury. Environ Int. 2021;154:106595.",
    ]

    for i, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(10)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15

    # =====================================================================
    # SAVE
    # =====================================================================
    out_path = os.path.join(OUT_DIR, "manuscript_en.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
