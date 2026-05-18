#!/usr/bin/env python3
"""
Generate the Japanese manuscript for:
「医療版サピア＝ウォーフ仮説：疾患分類体系はいかにして臨床的現実を構成するか
 ―― ICD-11移行と過労死概念からの考察」

Outputs: medical_sapir_whorf/output/manuscript_ja.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUT_DIR, exist_ok=True)


def add_para_with_refs(doc, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\{[^}]+\})", text)
    for part in parts:
        if part.startswith("{") and part.endswith("}"):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
    return p


def set_cell_text(cell, text, bold=False, size=Pt(9)):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = size


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # =====================================================================
    # タイトルページ
    # =====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("展望論文")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(24)
    run = p.add_run(
        "医療版サピア＝ウォーフ仮説：疾患分類体系はいかにして臨床的現実を構成するか\n"
        "\u2015\u2015 ICD-11移行と過労死概念からの考察"
    )
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(18)
    run = p.add_run("[著者名]")
    run.font.size = Pt(11)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[所属機関]")
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run("責任著者：[氏名、メールアドレス、住所]")
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run("本文語数：約4,500語（参考文献を除く）")
    run.font.size = Pt(10)

    doc.add_page_break()

    # =====================================================================
    # 要旨
    # =====================================================================
    add_heading(doc, "要旨", level=1)

    add_para_with_refs(
        doc,
        "疾患分類体系は一般に、臨床知識を整理するための中立的な道具とみなされている。"
        "本稿では、言語が認知を形作るとするサピア＝ウォーフ仮説{1,2}を医療に援用し、"
        "「医療版サピア＝ウォーフ仮説」を提唱する。すなわち、疾患分類体系は臨床的現実を"
        "単に記述するのではなく、能動的に構成しており、診断推論・治療選択・資源配分・"
        "患者体験に対して、制度的・構造的議論だけでは説明しきれない影響を及ぼしている"
        "というものである。"
    )

    add_para_with_refs(
        doc,
        "本仮説を二つの補完的な根拠から展開する。第一に、過労死（karoshi）を"
        "取り上げる。過労死は日本に固有の疾患概念であり、西洋の疾病分類には直接の対応物"
        "がない。同一の生物学的事象に対して、当該概念の有無が医療システムの応答を根本的"
        "に変えることを示す。{12,13} さらに、29年間の日本の労災補償データ"
        "（1996〜2024年度）を用いた独自の実証分析を提示する。分断時系列分析は、"
        "2001年の認定基準改正後に認定件数が3.71倍に増加したことを示し"
        "（β = 231.4, p < 0.001）、国際比較では、日本と韓国が"
        "多くの西洋諸国より低い生産年齢心血管疾患死亡率にもかかわらず包括的な"
        "認定制度を有する唯一の国であることを明らかにし、診断カテゴリーが臨床的"
        "インフラストラクチャーを反映するのではなく創出することを支持する。"
        "第二に、ICD-10からICD-11への移行を前向きの自然実験として活用することを"
        "提案する。特に慢性疼痛（MG30）とバーンアウト（QD85）に焦点を"
        "当てる。{14,15}"
    )

    add_para_with_refs(
        doc,
        "本論はKleinmanの疾患（disease）\u2013病い（illness）\u2013病人役割（sickness）"
        "の枠組み{16,17}およびHackingのループ効果理論{7,8}に位置づけ、医療版"
        "サピア＝ウォーフ効果が主として病人役割（sickness）レベルで作動しつつ、"
        "病い（illness）体験にフィードバックする構造を論じる。この視座は疾病分類改定、"
        "異文化間医療、診断の哲学に対して含意をもつ。"
    )

    p = doc.add_paragraph()
    run = p.add_run("キーワード：")
    run.bold = True
    run = p.add_run(
        "サピア＝ウォーフ仮説、疾患分類、ICD-11、過労死、疾病分類学、"
        "医療人類学、ループ効果、慢性疼痛、医学哲学"
    )

    doc.add_page_break()

    # =====================================================================
    # 序論
    # =====================================================================
    add_heading(doc, "序論", level=1)

    add_para_with_refs(
        doc,
        "言語学におけるサピア＝ウォーフ仮説は、言語の構造がその話者の認知と世界観に"
        "影響を与えるとするものである。{1,2} 強い形態（言語決定論）では言語が思考を"
        "決定し、弱い形態（言語相対論）では言語が思考に影響を与えるとする。強い形態は"
        "概ね否定されているが、弱い形態については相当の実証的根拠がある。異なる言語の"
        "話者は色彩・時間・空間関係を異なる仕方で知覚することが示されている。{3,4}"
    )

    add_para_with_refs(
        doc,
        "医療もまた、疾患分類体系という独自の「言語」を持つ。国際疾病分類（ICD）、"
        "精神疾患の診断と統計マニュアル（DSM）、およびそれらの地域的変種は、"
        "どのような症状の集合が認定された疾患を構成するかを定義し、それらに名前と"
        "コードを付与し、それによって臨床医が病いを知覚・伝達・行動する語彙を決定する。"
        "疾病分類が臨床実践に影響するのは、一つのレベルでは自明である。医療が"
        "組織される上部構造だからである。しかし本稿では、その影響が行政的便宜を超えて"
        "臨床認知そのものの領域にまで及ぶと主張する。我々はこの効果を「医療版"
        "サピア＝ウォーフ仮説」と呼ぶ。{5}"
    )

    add_para_with_refs(
        doc,
        "医療版サピア＝ウォーフ仮説は、疾患分類体系が臨床的現実を形成する様態が、"
        "自然言語の認知への効果と類似しており、潜在的に同程度に深遠であると主張する。"
        "具体的には、疾病分類学的枠組みが以下の効果を産出すると提案する。"
        "（1）診断的閉鎖：分類上のラベルが付与されると臨床的探索が終了する傾向、{6}"
        "（2）ループ効果：Hackingが論じた、分類カテゴリーがそれが分類する現象と"
        "相互作用しそれを変容させる現象、{7,8}"
        "（3）インフラストラクチャー創出：命名された疾患が研究資金・診療ガイドライン・"
        "専門医教育を引き寄せ、未命名の状態は不可視のままに留まる効果、{9,10}"
        "（4）疾病分類学的プラセボ：命名行為自体が測定可能な治療効果を産出する"
        "現象。{11}"
    )

    add_para_with_refs(
        doc,
        "本仮説を理論的思弁を超えて展開するため、二つの実証戦略を追求する。"
        "第一に、過労死（karoshi）を異文化間自然実験として分析する。日本は労働関連"
        "心血管死に対する診断カテゴリーを有するが、大半の国にはその等価物がなく、"
        "同一の生物学的事象が疾病分類学的概念の有無によってどう処理されるかの"
        "比較が可能である。{12,13} 第二に、進行中のICD-10からICD-11への移行を"
        "前向き自然実験として活用し、慢性疼痛分類の構造変化（MG30）およびバーンアウト"
        "の公式分類（QD85）を試金石とすることを提案する。{14,15}"
    )

    # =====================================================================
    # 理論的枠組み
    # =====================================================================
    add_heading(doc, "理論的枠組み", level=1)
    add_heading(doc, "言語相対論から疾病分類相対論へ", level=2)

    add_para_with_refs(
        doc,
        "サピア＝ウォーフ仮説の原型は自然言語を対象とする。ホピ語の話者は英語話者"
        "とは異なる仕方で時間を概念化する。それはホピ語が時間関係を異なる仕方で"
        "符号化するからである。{1} 我々は構造的な並行性に注目して、この枠組みを"
        "医療「言語」に拡張する。自然言語が経験の連続体を離散的カテゴリー（色彩用語、"
        "親族名称、空間前置詞）に分割するのと同様に、疾患分類体系は人間の苦しみの"
        "連続体を離散的な診断実体に分割する。問いは、この分割が単なる記述に過ぎないか、"
        "それともその対象を構成するのかにある。"
    )

    add_para_with_refs(
        doc,
        "Arthur Kleinmanの三部構造が有用な分析的足場を提供する。{16,17} 彼は"
        "疾患（disease：生物医学的病態）、病い（illness：患者の主観的苦痛体験）、"
        "病人役割（sickness：疾患の社会的役割と制度的認知）を区別する。"
        "我々は医療版サピア＝ウォーフ効果が主として病人役割（sickness）レベルで"
        "作動すると提案する。すなわち疾患カテゴリーの社会的・制度的構成においてである。"
        "しかし決定的に重要なのは、それが疾患（disease）レベル（診断・治療経路を介して）"
        "および病い（illness）レベル（患者の自己理解と病い行動を介して）の両方に"
        "フィードバックすることである。このフィードバックループこそが、本現象を"
        "「自明な」制度的効果の域を超えて高める（図1）。"
    )

    # --- 図1 インライン ---
    fig1_path = os.path.join(OUT_DIR, "figure1_ja.png")
    if os.path.exists(fig1_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig1_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[図1：figures_tables_ja.pptx参照]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "図1. 医療版サピア＝ウォーフのフィードバックループ。疾病分類学的カテゴリーは"
        "病人役割レベルで作動し、制度的応答を形成する。この変化は疾患認識（変化した"
        "診断パターンを介して）および病い体験（患者の自己理解を介して）にフィードバック"
        "し、単純なトップダウンの制度的効果を超える再帰的ループを創出する。"
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    add_heading(doc, "医療版サピア＝ウォーフ効果のメカニズム", level=2)

    add_para_with_refs(
        doc,
        "疾病分類学的カテゴリーが臨床的現実を形成する4つの異なるメカニズムを同定する。"
    )

    mechanisms = [
        (
            "診断的閉鎖",
            "臨床的プレゼンテーションに診断名が付与されると、さらなる探索は停止する"
            "傾向がある。これは認知心理学では「早期閉鎖」{6}として、臨床推論では"
            "「アンカーバイアス」{18}として十分に文書化されている。しかし医療版"
            "サピア＝ウォーフの視点は構造的次元を加える。利用可能な診断ラベルの"
            "セットは個人の認知だけでなく臨床ワークフロー全体を制約する。分類コードを"
            "持たない状態は請求も、監査も、疫学的研究もできず、体系的な不可視性を"
            "生み出す。"
        ),
        (
            "ループ効果",
            "Hackingの「ループ効果」概念は、人間科学における分類カテゴリーがその"
            "対象と相互作用する仕方を記述する。{7,8} 精神医学では例えば、DSM-IIIに"
            "「多重人格障害」が導入された後にその診断が急増し、カテゴリー自体が疾患の"
            "有病率を形成したことが示唆された。我々はループ効果が精神疾患に限定されず、"
            "診断に主観的報告や行動的要素が関与するところ——すなわち医療の大部分——で"
            "生じると主張する。"
        ),
        (
            "インフラストラクチャー創出",
            "命名された疾患は研究資金、診療ガイドライン、専門医教育プログラム、"
            "患者団体、医薬品開発を引き寄せる。「医学的に説明できない症状」から"
            "「慢性疲労症候群」へ、さらに「筋痛性脳脊髄炎/慢性疲労症候群（ME/CFS）」"
            "への移行は、命名がインフラストラクチャーを創出する過程を例示する。"
            "各改名には資金、臨床的注目、患者アイデンティティの変動が伴った。{9,10} "
            "COVID-19パンデミックは「Long COVID」で劇的な事例を提供し、迅速な命名が"
            "前例のない速度で研究インフラストラクチャーを触媒した。{19}"
        ),
        (
            "疾病分類学的プラセボ",
            "診断を受けること自体が治療効果を持ちうる。特に不確実性と苦痛を特徴と"
            "する状態においてそうである。慢性疼痛では、特定の診断名を受けた患者は、"
            "新たな治療が提供されなくても、より高い満足度と低い不安を報告することが"
            "研究で示されている。{11,20} この効果は薬理学的プラセボとは異なり、"
            "認知的・社会的メカニズムを通じて作動する。ラベルは患者の体験を妥当化し、"
            "不確実性を低減し、苦痛を理解するための物語的枠組みを提供する。"
        ),
    ]

    for title, body in mechanisms:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}. ")
        run.bold = True
        run.font.size = Pt(11)
        parts = re.split(r"(\{[^}]+\})", body)
        for part in parts:
            if part.startswith("{") and part.endswith("}"):
                run = p.add_run(part[1:-1])
                run.font.superscript = True
                run.font.size = Pt(8)
            else:
                run = p.add_run(part)
                run.font.size = Pt(11)

    # =====================================================================
    # 事例研究1：過労死
    # =====================================================================
    add_heading(doc, "事例研究1：過労死と職業衛生の疾病分類学的形成", level=1)

    add_para_with_refs(
        doc,
        "過労死（karoshi）は、過重労働に起因する死亡、主として過度の労働時間により"
        "惹起・増悪された心血管・脳血管イベントによる死亡を指す日本の社会医学的"
        "概念である。{12,21} 1970年代後半に初めて同定され、1987年に労災補償制度で"
        "正式に認定された過労死は、大多数の国の医療制度に直接の対応物を持たない。{13,22}"
        "この異文化間非対称性が、医療版サピア＝ウォーフ仮説を検証するための自然実験"
        "を提供する。"
    )

    add_heading(doc, "生物学的基盤の共通性", level=2)

    add_para_with_refs(
        doc,
        "過労死の基盤にある生物学的事象——慢性ストレスと長時間労働により惹起された"
        "心筋梗塞、脳卒中、急性心不全——は日本に固有のものではない。メタアナリシスに"
        "より、週55時間以上の労働は標準労働時間と比較して冠動脈疾患リスクの13%増加"
        "および脳卒中リスクの33%増加と関連し、地理的地域による有意な変動はないことが"
        "確立されている。{23,24} 病態生理は普遍的であり、異なるのは疾病分類学的な"
        "枠づけのみである。"
    )

    add_heading(doc, "分岐する臨床応答", level=2)

    add_para_with_refs(
        doc,
        "日本では、認定されたカテゴリーとしての過労死の存在が、当該概念を持たない国に"
        "は並行物のない臨床的・制度的応答の連鎖を産出する（表1）。数ヶ月にわたる"
        "週80時間労働の後に脳卒中で搬送された中年サラリーマンに対して、日本の医療"
        "システムは既成の診断パスウェイを持つ。当該イベントは潜在的に業務関連として"
        "コーディングされ、産業医が紹介され、労災補償申請が開始されうる。{13,22}"
        "過労死概念を持たない国では、同一の患者は標準的な脳血管障害の診断を受け、"
        "労働条件への体系的な問い合わせは行われない。"
    )

    # --- 表1 ---
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run(
        "表1. 日本（過労死概念あり）と当該概念を持たない国における"
        "労働関連心血管イベントへの臨床的・制度的応答の比較"
    )
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["領域", "日本（過労死概念あり）", "対応概念を持たない国"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    rows_data = [
        ["診断パスウェイ",
         "労働歴が体系的に評価される；産業医コンサルテーションが発動",
         "標準的CVA/MI精査；労働条件は稀にしか記録されない"],
        ["コーディング",
         "心血管イベントを過重労働と結びつける特定コード",
         "標準的ICD心血管コード；労働との紐付けなし"],
        ["補償",
         "法的基準を伴う専用の労災補償パスウェイ（例：月80時間超の残業）",
         "一般的な障害保険；心血管イベントに対する業務特異的パスウェイなし"],
        ["疫学的追跡",
         "1987年以降の過労死症例の全国統計が維持",
         "労働関連心血管死亡の体系的追跡なし"],
        ["予防インフラ",
         "時間外労働上限規制、年次健康診断、「過労死110番」",
         "一般的な産業衛生規制；標的型心血管予防なし"],
        ["研究資金",
         "過労死専門の研究プログラム；大規模疫学データセット",
         "労働関連心血管疾患は産業衛生のサブトピックとして研究"],
    ]
    for r, row_data in enumerate(rows_data, start=1):
        for c, val in enumerate(row_data):
            set_cell_text(table.rows[r].cells[c], val)

    p = doc.add_paragraph()
    p.space_after = Pt(12)

    add_para_with_refs(
        doc,
        "この分岐はインフラストラクチャー創出メカニズムを例示する。過労死の命名は"
        "既存の現実を単に記述したのではなく、法的基準・補償パスウェイ・疫学的監視・"
        "予防プログラムという制度的装置全体を創出し、それが今度は労働関連心血管疾患の"
        "検出・治療・予防のあり方を形成した。{12,25} ICD-11におけるバーンアウトの"
        "職業的現象としての分類（QD85）は、日本の過労死概念に影響を受けた、この"
        "疾病分類学的カテゴリーの国際的拡散を表しており、前向き研究の機会を"
        "提供する。{14}"
    )

    add_heading(doc, "命名効果の量的証拠", level=2)

    add_para_with_refs(
        doc,
        "岩崎らは、過重労働関連心血管・脳血管疾患の新認定基準導入が認定症例率の"
        "2.58倍増加と関連していたことを示し、分類の変更が疾患認識に測定可能な"
        "シフトを産出することを示唆した。{25} 決定的に重要なのは、この増加が単に"
        "発生率の増加に帰せないことである。むしろ、それは作動中の医療版"
        "サピア＝ウォーフ効果を反映している。疾病分類学的言語を変えることが、"
        "医療システムが知覚しうるものを変えたのである。"
    )

    add_heading(doc, "過労死認定の分断時系列分析", level=2)

    add_para_with_refs(
        doc,
        "この効果をより精密に定量化するため、厚生労働省の脳・心臓疾患労災補償状況"
        "年次報告から29年間の労災補償データ（1996〜2024年度）を編纂した。{32} "
        "2001年12月の認定基準改正を主要介入点とする分断時系列（ITS）分析を"
        "セグメント回帰により実施した。2001年改正は「80時間ルール」を導入し、"
        "労働関連心血管疾患の認定適格性を大幅に拡大した。"
    )

    add_para_with_refs(
        doc,
        "結果は顕著である（図3）。ITS回帰は、既存の時間的傾向を統制した上で、"
        "認定基準改正直後にβ = 231.4件（p < 0.001）の統計的に有意な水準変化を"
        "示した。年間平均認定件数は改正前（1996〜2001年度）の91.7件から、"
        "改正直後（2002〜2008年度）の339.9件へと3.71倍に増加した。認定率"
        "（認定件数／請求件数）は17.1%から39.6%に急上昇した。二標本t検定は"
        "この差が高度に有意であることを確認した（t = 14.1, p < 0.0001）。"
        "注目すべきは、負の傾斜変化（β = −17.2, p = 0.040）が、初期の急増後の"
        "基線への漸進的回帰を示していることであり、これは以前認識されなかった症例の"
        "「追いつき」効果とそれに続く安定化に一致する。"
    )

    add_para_with_refs(
        doc,
        "2021年9月の第二次改正は、不規則な勤務形態や心理的ストレスなどの"
        "残業時間以外の要因を認定基準に追加したものであり、直近年度（2022〜2024"
        "年度：194、216、241件）の緩やかな上昇と関連している。2021年改正後の"
        "観察期間はロバストなITS分析には不十分であるが、方向性の変化は医療版"
        "サピア＝ウォーフの予測に一致する。すなわち、疾病分類学的境界を拡大すれば、"
        "システムが認識する範囲が拡大するのである。"
    )

    # --- 図3 インライン ---
    fig3_path = os.path.join(OUT_DIR, "figure3_karoshi_its_ja.png")
    if os.path.exists(fig3_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig3_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[図3：output/figure3_karoshi_its_ja.png参照]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "図3. 過労死の労災認定件数の分断時系列分析（1996〜2024年度）。"
        "パネルA：年間認定件数（水色）と死亡件数（赤）、ITS回帰線付き。"
        "縦線は2001年および2021年の認定基準改正を示す。"
        "パネルB：請求件数と認定率（%）。"
        "データソース：厚生労働省「脳・心臓疾患に係る労災補償状況」年次報告。"
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    add_heading(doc, "国際比較：診断カテゴリーがインフラストラクチャーを創出する",
                level=2)

    add_para_with_refs(
        doc,
        "医療版サピア＝ウォーフ仮説が正しければ、労働関連心血管疾患の国際的認定"
        "パターンは、基礎疾患負荷ではなく過労死類似の診断カテゴリーの有無と相関する"
        "はずである。この予測を検証するため、WHO死亡データベースから8カ国の"
        "生産年齢人口（25〜64歳）の心血管疾患死亡率を比較した（図4）。{33}"
    )

    add_para_with_refs(
        doc,
        "結果は仮説を支持する。日本と韓国——包括的な過労死／過労死（과로사, "
        "gwarosa）認定制度を有する唯一の二カ国——は、比較対象国の中で最も低い"
        "生産年齢心血管疾患死亡率を示す（それぞれ10万人対23.0および28.0）。"
        "対照的に、心血管疾患死亡率がより高い国（例：米国57.0、ドイツ37.0、"
        "英国34.6/10万人）では、労働関連心血管死の体系的認定制度は非常に限定的"
        "であるか存在しない。{34} 日本は年間約200〜300件の労働関連心血管疾患を"
        "認定するが、同等の西洋諸国では10件未満である。この逆説——より低い疾患負荷に"
        "もかかわらずはるかに高い認定——は、まさに医療版サピア＝ウォーフ仮説が予測"
        "するものである。診断カテゴリーが基礎疾患の発生率とは独立に、認定のための"
        "制度的インフラストラクチャーを創出するのである。"
    )

    # --- 図4 インライン ---
    fig4_path = os.path.join(OUT_DIR, "figure4_international_cvd_ja.png")
    if os.path.exists(fig4_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig4_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[図4：output/figure4_international_cvd_ja.png参照]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "図4. 生産年齢人口の心血管疾患死亡率と職業性心血管疾患認定制度の国際比較。"
        "パネルA：国別の虚血性心疾患および脳血管疾患死亡率（10万人対、25〜64歳）。"
        "赤字は過労死類似の認定制度を有する国。"
        "パネルB：労働関連心血管疾患認定制度の水準と年間認定件数。"
        "データソース：WHO死亡データベース、各国労働統計。"
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    # =====================================================================
    # 事例研究2：ICD-11
    # =====================================================================
    add_heading(doc, "事例研究2：自然実験としてのICD-11移行", level=1)

    add_para_with_refs(
        doc,
        "ICD-10からICD-11への移行は、30年ぶりの世界的疾患分類体系の最も重大な"
        "改定を表す。いくつかの構造的変更が、医療版サピア＝ウォーフ仮説を前向きに"
        "検証する機会を提供する（表2）。"
    )

    # --- 表2 ---
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run(
        "表2. 医療版サピア＝ウォーフ分析に適するICD-11の主要な構造変化"
    )
    run.bold = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=5, cols=3)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["ICD-11の変更", "検証可能な事項",
                            "予測されるサピア＝ウォーフ効果"]):
        set_cell_text(table2.rows[0].cells[i], h, bold=True)

    t2_data = [
        ["慢性疼痛独立章（MG30）",
         "慢性疼痛に対する紹介パターン・専門医利用・治療法の変化",
         "疼痛がそれ自体の疾患実体として認識され、専門医紹介の増加と専用治療プログラムの創設へ"],
        ["バーンアウトの職業的現象としての分類（QD85）",
         "過労死類似概念の国際的拡散；病気休暇パターンと産業保健相談の変化",
         "公式分類が日本における過労死と類似のインフラストラクチャーを触媒"],
        ["ゲーム障害（6C51）",
         "分類前後の発生率の推移；受療行動",
         "分類が臨床的実体を創出し、認知と潜在的に有病率の両方を増加させる（ループ効果）"],
        ["性別不合が精神障害から性の健康に関連する状態へ移動",
         "スティグマ・治療アクセス・患者の自己同定の変化",
         "再分類が生物学的基盤を変えることなく病い体験と受療行動を変容させる"],
    ]
    for r, row_data in enumerate(t2_data, start=1):
        for c, val in enumerate(row_data):
            set_cell_text(table2.rows[r].cells[c], val)

    p = doc.add_paragraph()
    p.space_after = Pt(12)

    add_heading(doc, "慢性疼痛：症状から疾患へ", level=2)

    add_para_with_refs(
        doc,
        "ICD-11における慢性疼痛独立章（MG30）の新設はパラダイムシフトを表す。"
        "慢性疼痛はもはや他の疾患の単なる症状ではなく、それ自体の疾患実体である。{15,26}"
        "この再分類は臨床実践に連鎖的効果を産出すると予測される。フィールドテストは"
        "新分類の優れた診断コーディング性能と臨床的有用性を実証しており、臨床医が"
        "臨床的現実を捕捉するために新カテゴリーが有用であると認識していることを"
        "示唆する。{27}"
    )

    add_para_with_refs(
        doc,
        "医療版サピア＝ウォーフの視点から、MG30の導入は以下を予測する。"
        "（1）慢性疼痛患者が「主」診断の付随として管理されるのではなく、専用の"
        "疼痛サービスに紹介される割合の増加、（2）慢性疼痛特異的治療ガイドラインの"
        "策定、（3）「腰が悪い」から「慢性疼痛疾患を持っている」への患者の"
        "自己理解の変容、（4）疾患カテゴリーとしての慢性疼痛に特化した新規研究"
        "資金の創出。重要なことに、侵害可塑性疼痛（nociplastic pain）の概念——"
        "2017年にIASPが導入した、組織損傷や体性感覚系の病変の証拠なしに"
        "変容した侵害受容から生じる疼痛の第三の機序記述子{28}——は、命名が臨床的"
        "実体を創出する過程を例示する。以前「医学的に説明不能」として退けられた"
        "状態が、今や独立した病態生理学的カテゴリーとして認識されている。"
    )

    add_heading(doc, "提案する研究デザイン", level=2)

    add_para_with_refs(
        doc,
        "ICD-11を早期に採用した国における採用前後の臨床アウトカムを比較する"
        "分断時系列分析を提案する。主要アウトカムは以下を含む。"
        "（1）疼痛クリニックおよび産業保健サービスへの紹介率、"
        "（2）慢性疼痛薬の処方パターン、"
        "（3）ストレス関連心血管イベントに対する労災補償申請、"
        "（4）診断満足度と自己効力感を含む患者報告アウトカム指標。"
        "副次アウトカムは研究資金配分と出版動向を追跡する。ICD-11を異なる時期に"
        "採用する国々が、因果推論を強化する段階的採用デザインを提供する"
        "（図2）。{29}"
    )

    # --- 図2 インライン ---
    fig2_path = os.path.join(OUT_DIR, "figure2_ja.png")
    if os.path.exists(fig2_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig2_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[図2：figures_tables_ja.pptx参照]")
        run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        "図2. ICD-11採用を自然実験として活用する提案研究デザイン。国家間の段階的"
        "採用が多重ベースラインデザインを可能にし、分類変更の臨床実践への効果に"
        "関する因果推論を強化する。"
    )
    run.italic = True
    run.font.size = Pt(9)
    p.space_after = Pt(18)

    # =====================================================================
    # 考察
    # =====================================================================
    add_heading(doc, "考察", level=1)

    add_heading(doc, "疾患\u2013病い\u2013病人役割の枠組みの精緻化", level=2)

    add_para_with_refs(
        doc,
        "医療版サピア＝ウォーフ仮説はKleinmanの三部構造を、その三つのレベル間の"
        "フィードバックメカニズムを特定することにより精緻化する。{16,17} 古典的定式化"
        "では、疾患（disease）は生物医学的現実、病い（illness）は主観的体験、"
        "病人役割（sickness）は社会的役割である。我々の分析は、疾患レベルでは生物学的"
        "プロセスが確かに概ね概念非依存であることを示唆する。心筋梗塞は疾病分類に"
        "労働関連心臓死のカテゴリーが含まれるか否かに関わらず発生する。しかし病人役割"
        "レベルでは、サピア＝ウォーフ効果は強力かつ広範である。疾病分類学的カテゴリーの"
        "有無がどの状態が認識・治療・補償されるかを決定する。決定的に重要なのは、"
        "この病人役割レベルの効果が病い（illness）レベルに——患者の自己理解・受療行動・"
        "主観的体験を変容させて——、さらには疾患（disease）レベルにも——どの生物学的"
        "プロセスが調査・治療・モニタリングされるかを形成して——フィードバックすること"
        "である。"
    )

    add_para_with_refs(
        doc,
        "この再帰的構造が、医療版サピア＝ウォーフ効果を「制度は実践を形成する」"
        "という自明の理から区別する。主張は医療システムが行政的カテゴリーに応答する"
        "こと（これは自明に真であろう）ではなく、疾病分類学的カテゴリーがより深い"
        "意味で臨床的現実を構成するということである。それらは何が疾患として数えられ、"
        "誰が患者として数えられ、何が治療として数えられるかを決定する。サピア＝ウォーフ"
        "の枠づけは、これらの構成的効果が所与の疾病分類体系内で作業する実践者にとって"
        "しばしば不可視であるという決定的な洞察を加える。それは母語の認知的効果がその"
        "話者にとって不可視であるのと同様である。"
    )

    add_heading(doc, "疾病分類改定への含意", level=2)

    add_para_with_refs(
        doc,
        "医療版サピア＝ウォーフ仮説が正しければ、疾病分類改定は単なる行政的作業で"
        "なく、真の臨床的帰結を伴う介入である。診断カテゴリーを創設または廃止する"
        "決定は、言語に語を導入または除去することに類似する。それは知覚・伝達・行動"
        "しうるものの領域を拡大または縮小する。これは疾患分類を担う組織にとって"
        "実践的含意を持つ。具体的には、ICD改定に際して行政的アウトカム（コーディング"
        "精度、請求）だけでなく臨床アウトカム（診断パターン、治療選択、患者体験）をも"
        "追跡する前向き影響評価が伴われるべきである。{30}"
    )

    add_heading(doc, "異文化間医療への含意", level=2)

    add_para_with_refs(
        doc,
        "過労死の事例は、異なる文化が根本的に異なる疾病分類学的語彙を持ちうることを"
        "示す。それは患者ケアに対して現実的な帰結を伴う。これは十分に研究された"
        "文化結合症候群{31}の領域を超えて、医療が組織される構造そのものにまで及ぶ。"
        "ICD-11が特定の文化的文脈から生まれた概念（過労死がQD85バーンアウトに影響し、"
        "侵害可塑性疼痛が西洋の疼痛科学を反映するなど）を取り込むにつれ、それは"
        "疾病分類学的概念の国際的拡散の媒体となる。医療版サピア＝ウォーフ効果を"
        "理解することは、この拡散の帰結を予測し管理するために不可欠である。"
    )

    add_heading(doc, "限界", level=2)

    add_para_with_refs(
        doc,
        "いくつかの限界を認める必要がある。第一に、自然言語と疾病分類体系の類推は"
        "不完全である。疾患分類は明示的に構築され定期的に改訂されるが、自然言語は"
        "有機的に進化する。しかしこの差異は我々の論拠を弱めるのではなくむしろ強化する。"
        "それは疾病分類学的言語の構成的効果が意図的であり、したがって研究と介入に"
        "服しうることを含意するからである。第二に、ICD-11の自然実験デザインは医学知識・"
        "技術・医療政策の同時的変化による交絡に直面する。段階的採用デザインはこの懸念に"
        "部分的に対処するが、排除はできない。第三に、我々は二つの事例研究に焦点を"
        "当てた。医療版サピア＝ウォーフ効果の異なる医学領域にわたる一般性を確立するには"
        "さらなる研究が必要である。"
    )

    # =====================================================================
    # 結論
    # =====================================================================
    add_heading(doc, "結論", level=1)

    add_para_with_refs(
        doc,
        "我々は医療版サピア＝ウォーフ仮説を提唱した。疾患分類体系は診断的閉鎖、"
        "ループ効果、インフラストラクチャー創出、疾病分類学的プラセボのメカニズムを"
        "通じて臨床的現実を形成する。過労死の事例は、疾病分類学的概念の有無が同一の"
        "生物学的事象に対する医療システムの応答を根本的に変えることを実証する。"
        "29年間の日本の労災補償データを用いた分断時系列分析は、2001年の認定基準"
        "改正後に認定件数が3.71倍に増加したことを示す直接的な実証的証拠を提供し"
        "（β = 231.4, p < 0.001）、国際比較は疾患負荷ではなく診断的"
        "インフラストラクチャーが認定パターンを決定することを示す。"
        "ICD-10からICD-11への移行は、この仮説を前向きに検証する前例のない機会を"
        "提供する。医療版サピア＝ウォーフ仮説は、疾病分類改定が行政的"
        "帳簿整理ではなく、医療的現実を再構成する力を持つ臨床的介入として理解される"
        "べきことを含意する。"
    )

    # =====================================================================
    # 参考文献
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "参考文献", level=1)

    references = [
        "Whorf BL. Language, Thought, and Reality: Selected Writings. MIT Press; 1956.",
        "Sapir E. The status of linguistics as a science. Language. 1929;5(4):207\u2013214.",
        "Boroditsky L. Does language shape thought? Mandarin and English speakers\u2019 conceptions of time. Cogn Psychol. 2001;43(1):1\u201322.",
        "Winawer J, Witthoft N, Frank MC, Wu L, Wade AR, Boroditsky L. Russian blues reveal effects of language on color discrimination. Proc Natl Acad Sci U S A. 2007;104(19):7780\u20137785.",
        "Warner R. The relationship between language and disease concepts. Int J Psychiatry Med. 1977;7(1):57\u201368.",
        "Croskerry P. The importance of cognitive errors in diagnosis and strategies to minimize them. Acad Med. 2003;78(8):775\u2013780.",
        "Hacking I. The looping effects of human kinds. In: Sperber D, Premack D, Premack AJ, eds. Causal Cognition: A Multidisciplinary Debate. Clarendon Press; 1995:351\u2013394.",
        "Hacking I. The Social Construction of What? Harvard University Press; 1999.",
        "Rosenberg CE. The tyranny of diagnosis: specific entities and individual experience. Milbank Q. 2002;80(2):237\u2013260.",
        "Jutel A. Putting a Name to It: Diagnosis in Contemporary Society. Johns Hopkins University Press; 2011.",
        "Kendall NAS, Linton SJ, Main CJ. Guide to Assessing Psychosocial Yellow Flags in Acute Low Back Pain. Accident Rehabilitation & Compensation Insurance Corporation of New Zealand; 1997.",
        "Uehata T. Long working hours and occupational stress-related cardiovascular attacks among middle-aged workers in Japan. J Hum Ergol (Tokyo). 1991;20(2):147\u2013153.",
        "National Defense Counsel for Victims of Karoshi. Karoshi: When the Corporate Warrior Dies of Overwork. Mado-sha; 1990.",
        "World Health Organization. International Classification of Diseases 11th Revision (ICD-11). WHO; 2019. Accessed April 2026. https://icd.who.int/",
        "Treede RD, Rief W, Barke A, et al. Chronic pain as a symptom or a disease: the IASP Classification of Chronic Pain for the International Classification of Diseases (ICD-11). Pain. 2019;160(1):19\u201327.",
        "Kleinman A, Eisenberg L, Good B. Culture, illness, and care: clinical lessons from anthropologic and cross-cultural research. Ann Intern Med. 1978;88(2):251\u2013258.",
        "Kleinman A. The Illness Narratives: Suffering, Healing, and the Human Condition. Basic Books; 1988.",
        "Tversky A, Kahneman D. Judgment under uncertainty: heuristics and biases. Science. 1974;185(4157):1124\u20131131.",
        "Callard F, Perego E. How and why patients made Long Covid. Soc Sci Med. 2021;268:113426.",
        "Haugli L, Steen E, Laerum E, Nygard R, Finset A. Learning to have less pain \u2014 is it possible? A one-year follow-up study of the effects of a personal construct group learning programme on patients with chronic musculoskeletal pain. Patient Educ Couns. 2001;45(2):111\u2013118.",
        "Nishiyama K, Johnson JV. Karoshi \u2014 death from overwork: occupational health consequences of Japanese production management. Int J Health Serv. 1997;27(4):625\u2013641.",
        "Kato T. Karoshi: death from overwork. In: Kawakami N, ed. Occupational Health Psychology. Springer; 2014:153\u2013167.",
        "Kivim\u00e4ki M, Jokela M, Nyberg ST, et al. Long working hours and risk of coronary heart disease and stroke: a systematic review and meta-analysis of published and unpublished data for 603,838 individuals. Lancet. 2015;386(10005):1739\u20131746.",
        "Descatha A, Sembajwe G, Pega F, et al. The effect of exposure to long working hours on stroke: a systematic review and meta-analysis from the WHO/ILO Joint Estimates of the Work-related Burden of Disease and Injury. Environ Int. 2020;142:105746.",
        "Iwasaki K, Takahashi M, Nakata A. Health problems due to long working hours in Japan: working hours, workers\u2019 compensation (karoshi), and preventive measures. Ind Health. 2006;44(4):537\u2013540.",
        "Nicholas M, Vlaeyen JWS, Rief W, et al. The IASP classification of chronic pain for ICD-11: chronic primary pain. Pain. 2019;160(1):28\u201337.",
        "Barke A, Korwisi B, Jakob R, Konstanjsek N, Rief W, Treede RD. Classification of chronic pain for the International Classification of Diseases (ICD-11): results of the 2017 international World Health Organization field testing. Pain. 2022;163(2):e310\u2013e318.",
        "Kosek E, Cohen M, Baron R, et al. Do we need a third mechanistic descriptor for chronic pain states? Pain. 2016;157(7):1382\u20131386.",
        "Bernal JL, Cummins S, Gasparrini A. Interrupted time series regression for the evaluation of public health interventions: a tutorial. Int J Epidemiol. 2017;46(1):348\u2013355.",
        "Jutel A, Nettleton S. Towards a sociology of diagnosis: reflections and opportunities. Soc Sci Med. 2011;73(6):793\u2013800.",
        "Simons RC, Hughes CC, eds. The Culture-Bound Syndromes: Folk Illnesses of Psychiatric and Anthropological Interest. D. Reidel Publishing; 1985.",
        "Ministry of Health, Labour and Welfare. Status of Workers\u2019 Compensation for Brain/Cardiovascular Diseases and Mental Disorders (Annual Reports, FY2001\u2013FY2024). Tokyo: MHLW; 2025. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/koyou_roudou/roudoukijun/rousai/090316_00002.html",
        "World Health Organization. WHO Mortality Database. Geneva: WHO; 2026. Accessed April 2026. https://www.who.int/data-collections/mortality",
        "Pega F, N\u00e1fr\u00e1di B, Momen NC, et al. Global, regional, and national burdens of ischemic heart disease and stroke attributable to exposure to long working hours for 194 countries, 2000\u20132016: a systematic analysis from the WHO/ILO Joint Estimates of the Work-related Burden of Disease and Injury. Environ Int. 2021;154:106595.",
    ]

    for i, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(10)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15

    # SAVE
    out_path = os.path.join(OUT_DIR, "manuscript_ja.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
