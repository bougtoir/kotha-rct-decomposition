#!/usr/bin/env python3
"""Build an ADEMP reporting checklist docx for the KOTHA simulation study."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(14 if level == 1 else 12)
    return h


def add_para(doc, text, bold=False, italic=False, align=None, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def build(out_path='ADEMP_checklist_KOTHA.docx'):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    add_para(doc, 'ADEMP checklist for the KOTHA Framework simulation study',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_para(doc, 'Prepared for: Contemporary Clinical Trials submission',
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, 'Reporting guideline: Morris, White & Crowther (2019) Stat Med 38:2074-2102.',
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_paragraph()

    # ADEMP items mapped to KOTHA manuscript
    items = [
        ('A. Aims',
         'Identify the specific aims of the simulation study.',
         'Methods → Module K: Counterfactual power simulation; Module T: Bayesian Evidence Integration',
         'Aim: quantify how enrollment-driven risk-profile shifts reduce statistical power, and evaluate Bayesian power-prior integration across discounting levels.'),
        ('A. Aims',
         'Justify why simulation is needed rather than analytical derivation.',
         'Methods → Overview',
         'No closed-form solution exists for joint distribution of power across varying control event rates and true effects; simulation is required.'),
        ('D. Data-generating mechanisms',
         'State whether resampling or parametric simulation is used.',
         'Methods → Module K; data/magnesium_ami.csv, data/statins_hf_obs.csv, data/statins_hf_rct.csv',
         'Parametric simulation using published aggregate event rates and sample sizes; scenario-specific control event rates are derived from real-world or RCT-enrolled populations.'),
        ('D. Data-generating mechanisms',
         'Describe how simple/complex the model is and its basis in real data.',
         'Methods → Module K; Results',
         'Binary outcome with fixed treatment effect (OR/HR), 1:1 randomization, and total sample size from included studies. Two real cases: magnesium in AMI and statins in HF.'),
        ('D. Data-generating mechanisms',
         'List all factors varied and their levels.',
         'Methods → Module K; Table 3',
         'Factors: enrollment scenario (S1 real-world, S2 RCT-enrolled, S3 enriched) and true effect size grid (OR 0.50–0.95, HR 0.50–1.00).'),
        ('D. Data-generating mechanisms',
         'Describe factorial design (full, partial, one-at-a-time).',
         'Methods → Module K; Results',
         'Scenarios and effect sizes crossed for every combination; results displayed as power curves.'),
        ('E. Estimands / targets',
         'Define estimands/targets of the simulation.',
         'Methods → Module K; Module T',
         'Statistical power; expected number of events; required sample size for 80% power; posterior probability of benefit; bias-adjusted treatment effect.'),
        ('M. Methods',
         'Identify all methods evaluated and justify their relevance.',
         'Methods → Module K and Module T',
         'Schoenfeld-approximation power formula; random-effects frequentist meta-analysis; hierarchical Bayesian power-prior model; GRADE/TSA assessment.'),
        ('M. Methods',
         'Report software and packages, including version where possible.',
         'Methods → Practical implementation note; requirements.txt',
         'Python 3, numpy, scipy, pandas, pymc (or equivalent), python-docx, matplotlib. Exact versions listed in requirements.txt.'),
        ('M. Methods',
         'Describe how Bayesian analyses were implemented.',
         'Methods → Module T: Hierarchical Bayesian evidence integration',
         'Power-prior discounting with fixed α grid; Markov chain Monte Carlo sampling; posterior summaries reported as median and 95% CrI.'),
        ('P. Performance measures',
         'List all performance measures and justify relevance.',
         'Methods → Module K; Module H; Results',
         'Power (%) for each scenario and true effect; required N; information fraction; cumulative Z; posterior probability P(effect < threshold).'),
        ('P. Performance measures',
         'Provide explicit formulae for non-standard performance measures.',
         'Methods → Module K',
         'Power = Φ(|log(OR)| · sqrt(D/2) − z_{α/2}); D = N(p_c + p_t)/2. Reported in Methods.'),
        ('P. Performance measures',
         'State number of simulation repetitions or sample size for Monte Carlo error.',
         'Methods → Practical implementation note',
         'Power curves computed over a dense deterministic grid of true effects; Bayesian posteriors based on Markov chains with convergence diagnostics (R̂ checked).'),
        ('Coding and execution',
         'Separate scripts for data generation and analysis of estimates.',
         'validation/run_validation.py; build_paper_cct.py; generate_cct_docx.py',
         'run_validation.py performs analyses and writes results_summary.txt; build_paper_cct.py generates manuscript, figures, and tables.'),
        ('Coding and execution',
         'Random seed and reproducibility information.',
         'validation/run_validation.py; README.md',
         'Fixed random seed for stochastic components; all outputs reproduced by make clean && make all.'),
        ('Analysis',
         'Describe graphical and tabular exploration of results.',
         'Results; Figures 2–8; Tables 3–7',
         'Power curves, forest plots, TSA plots, sensitivity tables.'),
        ('Analysis',
         'Report Monte Carlo SE or uncertainty due to finite simulation.',
         'Methods → Practical implementation note',
         'Power formula is deterministic given parameters; Bayesian CrIs capture posterior uncertainty; MCMC convergence monitored.'),
        ('Reporting',
         'Report using ADEMP structure with rationale.',
         'Methods → Module K; this checklist',
         'Aims, data-generating mechanisms, estimands, methods, and performance measures described in Methods.'),
        ('Reporting',
         'Present competing methods side-by-side where applicable.',
         'Results; Tables 3, 5, 6, 7; Figures 4, 6, 7',
         'Frequentist vs Bayesian estimates; standard vs KOTHA-enhanced GRADE assessment; scenario comparisons.'),
        ('Reporting',
         'Publish code to execute the simulation study.',
         'Data availability statement; submission_package_CCT.zip; https://github.com/bougtoir/kotha-rct-decomposition',
         'All data and code included in repository and submission zip; make cct reproduces the manuscript.'),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ['ADEMP item', 'Question / requirement', 'Manuscript location', 'KOTHA-specific response']
    for cell, text in zip(hdr_cells, headers):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    for item, question, loc, response in items:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, [item, question, loc, response]):
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    # Set column widths
    widths = [Inches(0.9), Inches(1.8), Inches(1.8), Inches(2.5)]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    doc.add_paragraph()
    add_para(doc, 'Reference', bold=True)
    add_para(doc, 'Morris TP, White IR, Crowther MJ. Using simulation studies to evaluate statistical methods. Stat Med. 2019;38(11):2074-102.', size=11)

    doc.save(out_path)
    print(f'Saved {out_path}')


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--out', default='ADEMP_checklist_KOTHA.docx')
    args = parser.parse_args()
    build(args.out)
