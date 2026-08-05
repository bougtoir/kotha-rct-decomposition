#!/usr/bin/env python3
"""Extract all tables from the CCT manuscript docx into a separate editable docx."""
from lxml import etree
from docx import Document


def build(src_docx='KOTHA_Framework_CCT.docx', out_docx='KOTHA_Framework_CCT_tables.docx'):
    src = Document(src_docx)
    dst = Document()
    for t in src.tables:
        # Deep-copy the table XML so Word-native equations (OMML) are preserved.
        tbl_copy = etree.fromstring(etree.tostring(t._tbl))
        placeholder = dst.add_table(rows=1, cols=1)
        placeholder._tbl.getparent().replace(placeholder._tbl, tbl_copy)
    dst.save(out_docx)
    print(f'Saved {out_docx}')


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--src', default='KOTHA_Framework_CCT.docx')
    parser.add_argument('--out', default='KOTHA_Framework_CCT_tables.docx')
    args = parser.parse_args()
    build(args.src, args.out)
