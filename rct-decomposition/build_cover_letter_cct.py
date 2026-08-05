#!/usr/bin/env python3
"""Generate a CCT cover letter as a standalone docx."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_paragraph(doc, text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Date and address
    add_paragraph(doc, '[Date]')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Dr. Howard Sesso, ScD, MPH')
    add_paragraph(doc, 'Editor-in-Chief, Contemporary Clinical Trials')
    add_paragraph(doc, 'Harvard Medical School, Brigham and Women\'s Hospital')
    add_paragraph(doc, 'Boston, Massachusetts, United States of America')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Dear Dr. Sesso,')
    add_paragraph(doc, '')

    # Opening
    title = (
        '"The KOTHA Framework: diagnosing structural information loss in '
        'randomized controlled trial meta-analyses to inform trial design"'
    )
    opening = (
        f'We are pleased to submit our manuscript, {title}, '
        'for consideration for publication in *Contemporary Clinical Trials* as an Original Article.'
    )
    add_paragraph(doc, opening)
    add_paragraph(doc, '')

    # Summary
    add_paragraph(doc, 'Summary', bold=True)
    summary = (
        'The Knowledge-driven Observational-Trial Harmonization Approach (KOTHA) Framework '
        'provides a reproducible, quantitative method to diagnose and address structural '
        'information loss---a systematic reduction in statistical information that occurs when '
        'clinical trial enrollment criteria, consent processes, and site selection progressively '
        'exclude higher-risk patients, diluting event rates. The framework comprises three modules: '
        'Module K (counterfactual power simulation), Module T (Bayesian evidence integration with '
        'power-prior discounting), and Module H (GRADE-compatible evidence assessment). We illustrate '
        'its implications for trial design using two canonical cases of observational-RCT divergence: '
        'intravenous magnesium in acute myocardial infarction (12 trials, 1984-1995) and statins in '
        'heart failure (5 observational studies, 2 RCTs).'
    )
    add_paragraph(doc, summary)
    add_paragraph(doc, '')

    # Relevance
    add_paragraph(doc, 'Relevance to Contemporary Clinical Trials', bold=True)
    relevance = (
        'We believe this work is well suited to the readership of *Contemporary Clinical Trials* for '
        'the following reasons. First, Module K enables trialists to simulate how enrollment-driven '
        'risk-profile shifts affect statistical power under counterfactual designs, informing decisions '
        'about prognostic enrichment, event-driven enrollment, adaptive sample-size re-estimation, and '
        'broad-eligibility pragmatic strategies. Second, the empirical applications demonstrate that '
        'the same therapeutic question can appear conclusive in observational data and inconclusive in '
        'RCTs not because the treatments differ, but because enrolled cohorts have systematically lower '
        'event rates. This reframes a class of apparently negative RCT meta-analyses as design '
        'information problems rather than true null effects. Third, the framework bridges quantitative '
        'design diagnostics and structured evidence interpretation, helping readers distinguish '
        '"evidence of no effect" from "no evidence of effect" when planning, reporting, and '
        'interpreting clinical trials.'
    )
    add_paragraph(doc, relevance)
    add_paragraph(doc, '')

    # Significance
    add_paragraph(doc, 'Significance', bold=True)
    significance = (
        'By linking enrollment representativeness, event-rate dilution, statistical power, and '
        'evidence interpretation in a single reproducible pipeline, KOTHA offers a practical tool for '
        'improving trial design and post-hoc diagnostic appraisal. All analyses are generated from '
        'publicly available aggregate data using documented Python code, supporting transparency and '
        'reproducibility. We hope that this contribution will be of interest to the journal\'s '
        'audience of clinical trial methodologists, biostatisticians, and trial designers.'
    )
    add_paragraph(doc, significance)
    add_paragraph(doc, '')

    # Exclusivity
    exclusivity = (
        'This manuscript is original, has not been published previously, and is not under '
        'consideration elsewhere. All authors have approved the manuscript and agree with its '
        'submission to *Contemporary Clinical Trials*.'
    )
    add_paragraph(doc, exclusivity)
    add_paragraph(doc, '')

    # Declarations
    add_paragraph(doc, 'Declarations', bold=True)
    add_paragraph(doc, 'Funding: [To be determined]')
    add_paragraph(doc, 'Competing interests: The authors declare no competing interests.')
    add_paragraph(doc, '')

    # Suggested reviewers
    add_paragraph(doc, 'Suggested reviewers', bold=True)
    add_paragraph(doc, '- Professor Donald A. Berry, PhD (The University of Texas MD Anderson Cancer Center) --- Bayesian adaptive clinical trial design')
    add_paragraph(doc, '- Professor David L. DeMets, PhD (University of Wisconsin-Madison) --- clinical trials methodology and monitoring')
    add_paragraph(doc, '- Professor Thomas R. Fleming, PhD (University of Washington) --- group sequential methods and clinical trial design')
    add_paragraph(doc, '')

    # Closing
    add_paragraph(doc, 'Thank you for considering our manuscript. We look forward to your response.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Sincerely,')
    add_paragraph(doc, '')
    add_paragraph(doc, '[Corresponding author name and affiliation to be determined]')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Email: [To be determined]')

    doc.save('cover_letter_CCT.docx')
    print('Saved cover_letter_CCT.docx')


if __name__ == '__main__':
    main()
