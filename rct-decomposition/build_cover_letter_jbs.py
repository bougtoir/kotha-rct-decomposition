#!/usr/bin/env python3
"""Generate a Journal of Biopharmaceutical Statistics cover letter as a standalone docx."""
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _submission_details():
    """Read word count and main-text figure/table counts from the generated manuscript."""
    md_path = Path('05_paper_jbs.md')
    word_count = '[generated after manuscript build]'
    fig_count = 6
    table_count = 2
    if md_path.exists():
        text = md_path.read_text(encoding='utf-8')
        m = re.search(r'^word_count:\s*(\d+)', text, re.MULTILINE)
        if m:
            word_count = m.group(1)
        fig_count = len(set(re.findall(r'\*\*Fig\.\s*(\d+)', text)))
        table_count = len(set(re.findall(r'\*\*Table\s*(\d+):', text)))
    return word_count, fig_count, table_count


def add_paragraph(doc, text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def main():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    add_paragraph(doc, '[Date]')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Editorial Office')
    add_paragraph(doc, 'Journal of Biopharmaceutical Statistics')
    add_paragraph(doc, 'Taylor & Francis Group')
    add_paragraph(doc, '')

    add_paragraph(doc, 'To the Editor,')
    add_paragraph(doc, '')

    title = (
        '"The KOTHA Framework: A Simulation Study of Power-Prior Integration to '
        'Correct Structural Information Loss in RCT Meta-Analyses"'
    )
    opening = (
        f'We are pleased to submit our manuscript, {title}, '
        'for consideration for publication in the *Journal of Biopharmaceutical Statistics* as an Original Article.'
    )
    add_paragraph(doc, opening)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Summary', bold=True)
    summary = (
        'The Knowledge-driven Observational-Trial Harmonization Approach (KOTHA) Framework '
        'provides a reproducible, quantitative method to diagnose and address structural '
        'information loss---a systematic reduction in statistical information that occurs when '
        'clinical trial enrollment criteria, consent processes, and site selection progressively '
        'exclude higher-risk patients, diluting event rates. The framework comprises three modules: '
        'Module K (counterfactual power simulation), Module T (Bayesian evidence integration with '
        'power-prior discounting), and Module H (GRADE-compatible evidence assessment). We illustrate '
        'its implications for trial design using two canonical cases of observational-RCT divergence: '
        'intravenous magnesium in acute myocardial infarction and statins in heart failure.'
    )
    add_paragraph(doc, summary)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Relevance to Journal of Biopharmaceutical Statistics', bold=True)
    relevance = (
        'We believe this work is well suited to the readership of the *Journal of Biopharmaceutical Statistics* '
        'for the following reasons. First, the paper centers on a statistical methodology problem---'
        'information loss in RCT meta-analyses induced by enrollment-driven risk-profile shifts---and '
        'addresses it through a formal power-prior Bayesian model and a prespecified simulation study '
        'reported under the ADEMP guidelines. Second, it compares operating characteristics (bias, RMSE, '
        '95% credible-interval coverage, and power) of the proposed KOTHA combination rule against '
        'RCT-enrolled-only, observational-only, and naive meta-analysis estimators, providing the kind of '
        'quantitative performance evaluation expected in pharmaceutical statistics. Third, the empirical '
        'illustrations (magnesium in acute myocardial infarction and statins in heart failure) show how '
        'event-rate dilution can make an intervention that appears effective in observational cohorts look '
        'inconclusive in RCTs, and how power-prior discounting can be calibrated to reduce both bias and '
        'RMSE. Finally, the entire analysis pipeline is documented and reproducible, supporting '
        'transparency in regulatory-relevant development settings.'
    )
    add_paragraph(doc, relevance)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Significance', bold=True)
    significance = (
        'By linking enrollment representativeness, event-rate dilution, statistical power, and '
        'evidence interpretation in a single reproducible framework, KOTHA offers a practical design-stage '
        'diagnostic for Phase II/III and enrichment trials. All operating-characteristics estimates are '
        'generated from a documented simulation, and all empirical numbers are derived from publicly '
        'available aggregate data using reproducible Python code. We hope that this contribution will '
        'be of interest to the journal\'s audience of biostatisticians and pharmaceutical developers.'
    )
    add_paragraph(doc, significance)
    add_paragraph(doc, '')

    exclusivity = (
        'This manuscript is original, has not been published previously, and is not under '
        'consideration elsewhere. All authors have approved the manuscript and agree with its '
        'submission to the *Journal of Biopharmaceutical Statistics*.'
    )
    add_paragraph(doc, exclusivity)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Submission details', bold=True)
    word_count, fig_count, table_count = _submission_details()
    details = (
        f'The main text is approximately {word_count} words and contains {fig_count} figures '
        f'and {table_count} tables. All figures are cited inline and are supplied as separate, '
        f'high-resolution PNG files and an editable PowerPoint file for upload. All data and '
        f'analysis code are available in the public repository '
        f'(https://github.com/bougtoir/kotha-rct-decomposition).'
    )
    add_paragraph(doc, details)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Declarations', bold=True)
    add_paragraph(doc, 'Funding: No funding was received for this study.')
    add_paragraph(doc, 'Competing interests: The authors declare no competing interests.')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Suggested reviewers', bold=True)
    add_paragraph(doc, '- Professor Donald A. Berry, PhD (The University of Texas MD Anderson Cancer Center) --- Bayesian adaptive clinical trial design')
    add_paragraph(doc, '- Professor David L. DeMets, PhD (University of Wisconsin-Madison) --- clinical trials methodology and monitoring')
    add_paragraph(doc, '- Professor Thomas R. Fleming, PhD (University of Washington) --- group sequential methods and clinical trial design')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Thank you for considering our manuscript. We look forward to your response.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Sincerely,')
    add_paragraph(doc, '')
    add_paragraph(doc, '[Corresponding author name and affiliation to be determined]')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Email: [To be determined]')

    doc.save('cover_letter_JBS.docx')
    print('Saved cover_letter_JBS.docx')


if __name__ == '__main__':
    main()
