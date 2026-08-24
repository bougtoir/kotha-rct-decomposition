#!/usr/bin/env python3
"""Generate a Clinical Trials cover letter as a standalone docx."""
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _submission_details():
    """Read word count and main-text figure/table counts from the generated manuscript."""
    md_path = Path('05_paper_clinical_trials.md')
    word_count = '[generated after manuscript build]'
    fig_count = 4
    table_count = 2
    if md_path.exists():
        text = md_path.read_text(encoding='utf-8')
        m = re.search(r'^word_count:\s*(\d+)', text, re.MULTILINE)
        if m:
            word_count = m.group(1)
        fig_count = len(set(re.findall(r'\*\*Fig\.\s*(\d+)', text)))
        table_count = len(set(re.findall(r'\*\*Table\s*(\d+):', text)))
    return word_count, fig_count, table_count


def add_paragraph(doc, text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def main():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    add_paragraph(doc, '[Date]')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Dr. Colin Begg')
    add_paragraph(doc, 'Editor-in-Chief, Clinical Trials: Journal of the Society for Clinical Trials')
    add_paragraph(doc, 'Sage Publications')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Dear Dr. Begg,')
    add_paragraph(doc, '')

    title = (
        '"The KOTHA Framework: Diagnosing Structural Information Loss in '
        'Randomized Controlled Trial Meta-Analyses to Inform Trial Design"'
    )
    opening = (
        f'We are pleased to submit our manuscript, {title}, '
        'for consideration for publication in *Clinical Trials: Journal of the Society for Clinical Trials* as an Original Research article.'
    )
    add_paragraph(doc, opening)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Summary', bold=True)
    summary = (
        'The Knowledge-driven Observational-Trial Harmonization Approach (KOTHA) Framework '
        'provides a reproducible, quantitative method to diagnose and address structural '
        'information loss---a systematic reduction in statistical information that occurs when '
        'clinical trial enrollment criteria, consent processes, and site selection progressively '
        'exclude higher-risk patients, diluting event rates. The framework comprises three modules: '
        'Module K (counterfactual power simulation), Module T (Bayesian evidence integration with '
        'power-prior discounting), and Module H (GRADE-compatible evidence assessment). We illustrate '
        'its implications for trial design using two canonical cases of observational-RCT divergence: '
        'intravenous magnesium in acute myocardial infarction and statins in heart failure.'
    )
    add_paragraph(doc, summary)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Relevance to Clinical Trials', bold=True)
    relevance = (
        'We believe this work is well suited to the readership of *Clinical Trials* for the following '
        'reasons. First, Module K enables trialists to simulate how enrollment-driven risk-profile shifts '
        'affect statistical power under counterfactual designs, informing decisions about prognostic '
        'enrichment, event-driven enrollment, adaptive sample-size re-estimation, and broad-eligibility '
        'pragmatic strategies. Second, the empirical applications demonstrate that the same therapeutic '
        'question can appear conclusive in observational data and inconclusive in RCTs not because the '
        'treatments differ, but because enrolled cohorts have systematically lower event rates. This '
        'reframes a class of apparently negative RCT meta-analyses as design-information problems '
        'rather than true null effects. Third, the framework bridges quantitative design diagnostics and '
        'structured evidence interpretation, helping readers distinguish "evidence of no effect" '
        'from "no evidence of effect" when planning, reporting, and interpreting clinical trials. The '
        'reproducible pipeline and ADEMP reporting structure further align with the journal\'s goal of '
        'increasing transparency and reducing publication bias.'
    )
    add_paragraph(doc, relevance)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Significance', bold=True)
    significance = (
        'By linking enrollment representativeness, event-rate dilution, statistical power, and '
        'evidence interpretation in a single reproducible pipeline, KOTHA offers a practical tool for '
        'improving trial design and post-hoc diagnostic appraisal. All analyses are generated from '
        'publicly available aggregate data using documented Python code, supporting transparency and '
        'reproducibility. We hope that this contribution will be of interest to the journal\'s '
        'audience of clinical trial methodologists, biostatisticians, and trial designers.'
    )
    add_paragraph(doc, significance)
    add_paragraph(doc, '')

    exclusivity = (
        'This manuscript is original, has not been published previously, and is not under '
        'consideration elsewhere. All authors have approved the manuscript and agree with its '
        'submission to *Clinical Trials: Journal of the Society for Clinical Trials*.'
    )
    add_paragraph(doc, exclusivity)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Submission details', bold=True)
    word_count, fig_count, table_count = _submission_details()
    details = (
        f'The main text is approximately {word_count} words and contains {fig_count} figures '
        f'and {table_count} tables. All figures are cited inline and are supplied as separate, '
        f'high-resolution PNG files and an editable PowerPoint file for upload. All data and '
        f'analysis code are available in the public repository '
        f'(https://github.com/bougtoir/kotha-rct-decomposition).'
    )
    add_paragraph(doc, details)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Declarations', bold=True)
    add_paragraph(doc, 'Funding: [To be determined]')
    add_paragraph(doc, 'Competing interests: [To be determined]')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Suggested reviewers', bold=True)
    add_paragraph(doc, '- Professor Donald A. Berry, PhD (The University of Texas MD Anderson Cancer Center) --- Bayesian adaptive clinical trial design')
    add_paragraph(doc, '- Professor David L. DeMets, PhD (University of Wisconsin-Madison) --- clinical trials methodology and monitoring')
    add_paragraph(doc, '- Professor Thomas R. Fleming, PhD (University of Washington) --- group sequential methods and clinical trial design')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Thank you for considering our manuscript. We look forward to your response.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Sincerely,')
    add_paragraph(doc, '')
    add_paragraph(doc, '[Corresponding author name and affiliation to be determined]')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Email: [To be determined]')

    doc.save('cover_letter_ClinicalTrials.docx')
    print('Saved cover_letter_ClinicalTrials.docx')


if __name__ == '__main__':
    main()
