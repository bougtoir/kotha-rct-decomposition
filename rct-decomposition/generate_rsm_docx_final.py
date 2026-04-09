#!/usr/bin/env python3
"""Convert corrected RSM markdown (04_paper_rsm.md) to formatted docx."""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = '/home/ubuntu/repos/wip/rct-decomposition'
MD_FILE = os.path.join(BASE, '04_paper_rsm.md')
FIG_DIR = os.path.join(BASE, 'validation', 'figures')
OUTPUT = '/home/ubuntu/KOTHA_Framework_RSM.docx'

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if align:
        p.alignment = align
    return p


def parse_inline(text):
    """Parse markdown inline formatting into (text, bold, italic) runs."""
    # Replace $...$ math with plain text (remove $)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    # Replace --- with em dash (before --)
    text = text.replace('---', '\u2014')
    # Replace -- with en dash
    text = text.replace('--', '\u2013')

    runs = []
    i = 0
    current = ''
    while i < len(text):
        # Bold+italic ***text***
        if text[i:i+3] == '***':
            if current:
                runs.append((current, False, False))
                current = ''
            end = text.find('***', i + 3)
            if end != -1:
                runs.append((text[i+3:end], True, True))
                i = end + 3
                continue
        # Bold **text**
        if text[i:i+2] == '**':
            if current:
                runs.append((current, False, False))
                current = ''
            end = text.find('**', i + 2)
            if end != -1:
                runs.append((text[i+2:end], True, False))
                i = end + 2
                continue
        # Italic *text* (not preceded/followed by *)
        if (text[i] == '*' and
            (i == 0 or text[i-1] != '*') and
            (i+1 < len(text) and text[i+1] != '*')):
            if current:
                runs.append((current, False, False))
                current = ''
            end = text.find('*', i + 1)
            if end != -1 and (end+1 >= len(text) or text[end+1] != '*'):
                runs.append((text[i+1:end], False, True))
                i = end + 1
                continue
        current += text[i]
        i += 1
    if current:
        runs.append((current, False, False))
    return runs if runs else [(text, False, False)]


def add_rich_para(text, align=None):
    """Add a paragraph with parsed inline formatting."""
    p = doc.add_paragraph()
    runs = parse_inline(text)
    for txt, bold, italic in runs:
        run = p.add_run(txt)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    if align:
        p.alignment = align
    return p


def add_table(headers, rows):
    """Add a formatted table."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header.strip())
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    # Data rows
    for i, row in enumerate(rows):
        for j in range(min(len(row), ncols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            ct = row[j].strip()
            ct = re.sub(r'\$([^$]+)\$', r'\1', ct)
            ct = ct.replace('--', '\u2013')
            run = p.add_run(ct)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    return table


def add_figure(img_path, caption, fig_num):
    """Add a figure with caption."""
    full_path = os.path.join(BASE, img_path) if not os.path.isabs(img_path) else img_path
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(full_path, width=Inches(5.5))
    else:
        print(f"  WARNING: Figure not found: {full_path}")

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap_p.add_run(f'Figure {fig_num}. ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = cap_p.add_run(caption)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def clean_math(text):
    """Convert LaTeX math to Unicode text."""
    replacements = [
        ('\\text{', ''), ('}', ''),
        ('\\sim', '~'), ('\\cdot', '\u00b7'),
        ('\\sqrt', '\u221a'), ('\\sum', '\u03a3'),
        ('\\prod', '\u220f'), ('\\alpha', '\u03b1'),
        ('\\beta', '\u03b2'), ('\\mu', '\u03bc'),
        ('\\tau', '\u03c4'), ('\\theta', '\u03b8'),
        ('\\sigma', '\u03c3'), ('\\delta', '\u03b4'),
        ('\\rho', '\u03c1'), ('\\Phi', '\u03a6'),
        ('\\phi', '\u03c6'), ('\\in', '\u2208'),
        ('\\leq', '\u2264'), ('\\geq', '\u2265'),
        ('\\neq', '\u2260'), ('\\times', '\u00d7'),
        ('\\to', '\u2192'), ('\\approx', '\u2248'),
        ('\\frac', ''), ('\\left', ''),
        ('\\right', ''), ('\\quad', '  '),
        ('\\hat', ''), ('\\log', 'log'),
        ('\\exp', 'exp'), ('\\mid', '|'),
        ('^2', '\u00b2'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r'_\{([^}]+)\}', r'_\1', text)
    text = re.sub(r'\^\{([^}]+)\}', r'^\1', text)
    return text


# ============================================================
# Parse the markdown and build the docx
# ============================================================
with open(MD_FILE, 'r') as f:
    lines = f.readlines()

print(f"Read {len(lines)} lines from {MD_FILE}")

i = 0
in_table = False
table_headers = []
table_rows = []
in_fig_section = False
current_section = ''

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue

    # Skip empty lines
    if line.strip() == '':
        i += 1
        continue

    # Title (first line starting with single #)
    if line.startswith('# ') and not line.startswith('## ') and i < 5:
        title_text = line[2:].strip()
        p = add_para(title_text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.runs[0].font.size = Pt(14)
        i += 1
        continue

    # Author/corresponding lines
    if line.startswith('**Authors**') or line.startswith('**Corresponding'):
        text = line.replace('**', '')
        add_para(text, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue

    # Section headers
    if line.startswith('#### '):
        add_heading(line[5:].strip(), level=4)
        i += 1
        continue
    if line.startswith('### '):
        add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    if line.startswith('## '):
        section_title = line[3:].strip()
        current_section = section_title

        if section_title == 'Figures':
            in_fig_section = True
            add_heading(section_title, level=2)
            i += 1
            continue
        else:
            in_fig_section = False

        add_heading(section_title, level=2)
        i += 1
        continue

    # Handle Figures section
    if in_fig_section:
        if line.startswith('**Fig.'):
            m = re.match(r'\*\*Fig\.\s*(\d+)\*\*\s*(.*)', line)
            if m:
                fig_num = int(m.group(1))
                caption = m.group(2).strip()
                # Find the image reference
                j = i + 1
                img_path = None
                while j < len(lines):
                    stripped = lines[j].strip()
                    if stripped.startswith('!['):
                        img_m = re.match(r'!\[.*?\]\((.*?)\)', stripped)
                        if img_m:
                            img_path = img_m.group(1)
                        j += 1
                        break
                    elif stripped == '':
                        j += 1
                    else:
                        break
                if img_path:
                    add_figure(img_path, caption, fig_num)
                i = j
                continue
        i += 1
        continue

    # Skip inline image references in body
    if line.strip().startswith('!['):
        i += 1
        continue

    # Flush pending table if needed
    if in_table and not line.strip().startswith('|'):
        add_table(table_headers, table_rows)
        in_table = False
        table_headers = []
        table_rows = []

    # Table detection
    if line.strip().startswith('|'):
        # Separator row
        if re.match(r'^[\|\-\s:]+$', line.strip()):
            i += 1
            continue

        cells = [c.strip() for c in line.strip().split('|')[1:-1]]

        if not in_table:
            # Check if next line is separator
            if i + 1 < len(lines) and re.match(r'^[\|\-\s:]+$', lines[i+1].strip()):
                in_table = True
                table_headers = cells
                table_rows = []
                i += 2  # skip header and separator
                continue
        else:
            table_rows.append(cells)
            # Check if next line is still table
            if i + 1 >= len(lines) or not lines[i+1].strip().startswith('|'):
                add_table(table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
            i += 1
            continue

        i += 1
        continue

    # Bold table caption
    if line.strip().startswith('**Table'):
        text = line.strip().replace('**', '')
        add_para(text, bold=True)
        i += 1
        continue

    # Math blocks ($$...$$)
    if line.strip().startswith('$$'):
        math_text = line.strip().replace('$$', '').strip()
        if not math_text:
            math_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('$$'):
                math_lines.append(lines[i].strip())
                i += 1
            math_text = ' '.join(math_lines)
            i += 1  # skip closing $$
        else:
            i += 1
        math_text = clean_math(math_text)
        add_para(math_text.strip(), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        continue

    # Bullet lists (top level)
    if line.strip().startswith('- '):
        text = line.strip()[2:]
        p = doc.add_paragraph(style='List Bullet')
        runs = parse_inline(text)
        for txt, bold, italic in runs:
            run = p.add_run(txt)
            run.bold = bold
            run.italic = italic
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        i += 1
        continue

    # Indented bullet lists
    if re.match(r'^  +- ', line):
        text = line.strip()[2:]
        p = doc.add_paragraph(style='List Bullet 2')
        runs = parse_inline(text)
        for txt, bold, italic in runs:
            run = p.add_run(txt)
            run.bold = bold
            run.italic = italic
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        i += 1
        continue

    # Numbered lists
    m = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
    if m:
        num = int(m.group(1))
        text = m.group(2)

        # Detect references (typically start with author name)
        if current_section == 'References':
            p = doc.add_paragraph()
            run = p.add_run(f'{num}. {text}')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            i += 1
            continue
        else:
            p = doc.add_paragraph(style='List Number')
            runs = parse_inline(text)
            for txt, bold, italic in runs:
                run = p.add_run(txt)
                run.bold = bold
                run.italic = italic
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            i += 1
            continue

    # Regular paragraph
    para_text = line.strip()
    if para_text:
        add_rich_para(para_text)
    i += 1

# Flush any remaining table
if in_table:
    add_table(table_headers, table_rows)

# Save
doc.save(OUTPUT)
print(f'\nSaved RSM docx to {OUTPUT}')

# Word count
word_count = 0
for p in doc.paragraphs:
    word_count += len(p.text.split())
print(f'Total word count (all content): ~{word_count}')
print(f'File size: {os.path.getsize(OUTPUT) / 1024:.0f} KB')
