#!/usr/bin/env python3
"""Build an editable docx with supplementary study-level tables."""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import build_paper as bp


def _parse_markdown_table(md):
    lines = [l.strip() for l in md.strip().split('\n') if l.strip()]
    # remove separator line
    rows = []
    for i, line in enumerate(lines):
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
        else:
            continue
        if i == 1 and all(set(c) <= set('-| ') for c in cells):
            continue
        rows.append(cells)
    return rows


def _add_table(doc, caption, md):
    rows = _parse_markdown_table(md)
    if not rows:
        return
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(rows[0]):
        hdr_cells[i].text = text
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    doc.add_paragraph()


def build(out_docx='KOTHA_Framework_CCTC_supplementary_tables.docx'):
    v = bp._compute_values()
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    run = title.add_run('Supplementary Tables')
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _add_table(doc, 'Supplementary Table S1: Study-level data for magnesium in AMI', bp._table_3(v))
    _add_table(doc, 'Supplementary Table S2: Study-level data for statins in heart failure', bp._table_4(v))

    doc.save(out_docx)
    print(f'Saved {out_docx}')


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--out', default='KOTHA_Framework_CCTC_supplementary_tables.docx')
    args = parser.parse_args()
    build(args.out)
