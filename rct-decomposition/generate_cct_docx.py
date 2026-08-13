#!/usr/bin/env python3
"""Convert CCT-format markdown to formatted docx with a title page.

Equations written in LaTeX ($...$ and $$...$$) are converted to Word-native
Office Math Markup (OMML) using pandoc.
"""
import argparse
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from lxml import etree

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# Locate a pandoc binary. The default fallback is a user-local download; if it
# is not present the PATH is searched.
_PANDOC_CANDIDATES = [
    os.path.expanduser('~/.local/pandoc/bin/pandoc'),
    shutil.which('pandoc'),
    'pandoc',
]
PANDOC = next((p for p in _PANDOC_CANDIDATES if p and os.path.isfile(p) and os.access(p, os.X_OK)), 'pandoc')


def _rewrite_docx_timestamps(docx_path):
    """Rewrite docx zip timestamps to a fixed value for reproducibility."""
    _fixed_time = (2025, 1, 1, 0, 0, 0)
    _tmp = docx_path + '.tmp'
    shutil.move(docx_path, _tmp)
    with zipfile.ZipFile(_tmp, 'r') as zin:
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                info.date_time = _fixed_time
                zout.writestr(info, data)
    os.remove(_tmp)


def build(input_md, output_docx):
    BASE = os.path.dirname(os.path.abspath(input_md))

    doc = Document()

    # Fixed document properties for reproducible docx output
    doc.core_properties.created = datetime(2025, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2025, 1, 1, 0, 0, 0, tzinfo=timezone.utc)

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- Style setup ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 2.0

    body = doc.element.body

    def sectPr():
        for child in body:
            if child.tag.endswith('sectPr'):
                return child
        return None

    def insert_element(e):
        """Insert element before the final sectPr, keeping it at the document end."""
        sp = sectPr()
        if sp is None:
            body.append(e)
        else:
            body.insert(body.index(sp), e)

    def set_run_font(run_elem, size=12):
        """Set Times New Roman and size (points) for a w:r element."""
        rPr = run_elem.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run_elem.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), str(int(size * 2)))

        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), str(int(size * 2)))

    def apply_para_format(p_elem, style='Normal', align=None, line_spacing=480, space_after=120):
        """Apply paragraph style, alignment, double line spacing and space-after."""
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)

        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            pStyle = OxmlElement('w:pStyle')
            pPr.append(pStyle)
        pStyle.set(qn('w:val'), style)

        if align is not None:
            align_map = {
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
            }
            align_val = align_map.get(align, align) if isinstance(align, int) else align
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), align_val)

        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), str(line_spacing))
        spacing.set(qn('w:lineRule'), 'auto')
        if space_after is not None:
            spacing.set(qn('w:after'), str(space_after))

    def apply_citation_superscript(p_elem):
        """Convert bracketed citation numbers (e.g. [1], [2-3]) to superscript."""
        CIT_RE = re.compile(r'(\[\d+(?:-\d+)?\])')
        for r in list(p_elem.findall(qn('w:r'))):
            # Only process text runs; skip math, drawing, or other special runs
            children = list(r)
            t_elems = [c for c in children if c.tag == qn('w:t')]
            if not t_elems:
                continue
            non_t = [c for c in children if c.tag not in (qn('w:rPr'), qn('w:t'))]
            if non_t:
                continue
            text = ''.join((t.text or '') for t in t_elems)
            if '[' not in text:
                continue
            rPr = r.find(qn('w:rPr'))
            segments = [s for s in CIT_RE.split(text) if s != '']
            p_elem.remove(r)
            for seg in segments:
                new_r = OxmlElement('w:r')
                if rPr is not None:
                    new_rPr = deepcopy(rPr)
                else:
                    new_rPr = OxmlElement('w:rPr')
                new_t = OxmlElement('w:t')
                new_t.set(qn('xml:space'), 'preserve')
                new_t.text = seg
                if CIT_RE.match(seg):
                    vert = OxmlElement('w:vertAlign')
                    vert.set(qn('w:val'), 'superscript')
                    new_rPr.append(vert)
                if len(new_rPr) > 0:
                    new_r.append(new_rPr)
                new_r.append(new_t)
                p_elem.append(new_r)

    # --- Page numbers in footer ---
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        set_run_font(run._r, size=10)

    def pandoc_docx(text):
        """Convert a markdown fragment to a temporary docx and return the Document."""
        with tempfile.TemporaryDirectory() as tmp:
            md_path = os.path.join(tmp, 'in.md')
            out_path = os.path.join(tmp, 'out.docx')
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(text + '\n')
            subprocess.run(
                [PANDOC, '-f', 'markdown', '-t', 'docx', '-o', out_path, md_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            return Document(out_path)

    def pandoc_para(text):
        """Return a w:p element rendered from markdown text by pandoc."""
        tmpdoc = pandoc_docx(text)
        if not tmpdoc.paragraphs:
            return OxmlElement('w:p')
        p = tmpdoc.paragraphs[0]._p
        return etree.fromstring(etree.tostring(p))

    def pandoc_table(md_table):
        """Return a w:tbl element rendered from a markdown table by pandoc."""
        tmpdoc = pandoc_docx(md_table)
        if not tmpdoc.tables:
            return None
        tbl = tmpdoc.tables[0]._tbl
        return etree.fromstring(etree.tostring(tbl))

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
        return h

    def add_para(text, bold=False, italic=False, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if align:
            p.alignment = align
        return p

    def add_rich_para(text, align=None, style='Normal'):
        """Add a paragraph, using pandoc so inline math becomes Word equations."""
        p_elem = pandoc_para(text)
        insert_element(p_elem)
        apply_para_format(p_elem, style=style, align=align)
        for r in p_elem.findall(qn('w:r')):
            set_run_font(r, size=12)
        apply_citation_superscript(p_elem)
        return p_elem

    def add_table(headers, rows):
        """Add a formatted table; math in headers/cells is converted to OMML."""
        ncols = len(headers)
        md = '| ' + ' | '.join(headers) + ' |\n'
        md += '|' + '|'.join(['---'] * ncols) + '|\n'
        for row in rows:
            md += '| ' + ' | '.join(str(x) for x in row[:ncols]) + ' |\n'

        tbl_elem = pandoc_table(md)
        if tbl_elem is None:
            return None

        # Table style and centering
        tblPr = tbl_elem.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl_elem.insert(0, tblPr)
        tblStyle = tblPr.find(qn('w:tblStyle'))
        if tblStyle is None:
            tblStyle = OxmlElement('w:tblStyle')
            tblPr.append(tblStyle)
        tblStyle.set(qn('w:val'), 'Table Grid')

        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')

        # Format cells: 10 pt Times New Roman, header bold, single line spacing
        for tr in tbl_elem.findall(qn('w:tr')):
            trPr = tr.find(qn('w:trPr'))
            is_header = False
            if trPr is not None:
                is_header = trPr.find(qn('w:tblHeader')) is not None
            for tc in tr.findall(qn('w:tc')):
                for p in tc.findall(qn('w:p')):
                    apply_para_format(p, style='Normal', line_spacing=240, space_after=0)
                    for r in p.findall(qn('w:r')):
                        set_run_font(r, size=10)
                        if is_header:
                            rPr = r.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                r.insert(0, rPr)
                            if rPr.find(qn('w:b')) is None:
                                rPr.append(OxmlElement('w:b'))

        insert_element(tbl_elem)
        return tbl_elem

    def clean_math(text):
        """Convert a small subset of LaTeX math to Unicode text for captions."""
        replacements = [
            ('\\text{', ''), ('}', ''),
            ('\\sim', '~'), ('\\cdot', '\u00b7'),
            ('\\sqrt', '\u221a'), ('\\sum', '\u03a3'),
            ('\\prod', '\u220f'), ('\\alpha', '\u03b1'),
            ('\\beta', '\u03b2'), ('\\mu', '\u03bc'),
            ('\\tau', '\u03c4'), ('\\theta', '\u03b8'),
            ('\\sigma', '\u03c3'), ('\\delta', '\u03b4'),
            ('\\rho', '\u03c1'), ('\\Phi', '\u03a6'),
            ('\\phi', '\u03c6'), ('\\in', '\u2208'),
            ('\\leq', '\u2264'), ('\\geq', '\u2265'),
            ('\\neq', '\u2260'), ('\\times', '\u00d7'),
            ('\\to', '\u2192'), ('\\approx', '\u2248'),
            ('\\frac', ''), ('\\left', ''),
            ('\\right', ''), ('\\quad', '  '),
            ('\\hat', ''), ('\\log', 'log'),
            ('\\exp', 'exp'), ('\\mid', '|'),
            ('^2', '\u00b2'),
        ]
        for old, new in replacements:
            text = text.replace(old, new)
        text = re.sub(r'_\{([^}]+)\}', r'_\1', text)
        text = re.sub(r'\^\{([^}]+)\}', r'^\1', text)
        return text

    def add_figure(img_path, caption, fig_num):
        """Add a figure with caption; caption math is converted to Unicode text."""
        full_path = os.path.join(BASE, img_path) if not os.path.isabs(img_path) else img_path
        caption = re.sub(r'\$([^$]+)\$', lambda m: clean_math(m.group(1)), caption)
        if os.path.exists(full_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(full_path, width=Inches(5.5))
        else:
            print(f"  WARNING: Figure not found: {full_path}")

        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cap_p.add_run(f'Figure {fig_num}. ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run = cap_p.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    # ============================================================
    # Read and parse front matter for title page
    # ============================================================
    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    metadata = {}
    start = 0
    if lines and lines[0].strip() == '---':
        for idx in range(1, len(lines)):
            if lines[idx].strip() == '---':
                start = idx + 1
                break
            m = re.match(r'^(\w+):\s*(.*)$', lines[idx].strip())
            if m:
                metadata[m.group(1).lower()] = m.group(2).strip()

    # Title page
    if metadata.get('running_head'):
        p = add_para(f"Running head: {metadata['running_head']}", bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if metadata.get('title'):
        p = add_para(metadata['title'], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in p.runs:
            run.font.size = Pt(14)
    if metadata.get('authors'):
        p = add_para(metadata['authors'], italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if metadata.get('affiliations'):
        p = add_para(f"Affiliations: {metadata['affiliations']}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if metadata.get('corresponding_author'):
        add_para(f"Corresponding author: {metadata['corresponding_author']}")
    if metadata.get('corresponding_author_address'):
        add_para(f"Corresponding author address: {metadata['corresponding_author_address']}")
    if metadata.get('word_count'):
        add_para(f"Word count: {metadata['word_count']}")
    add_para("Prepared for submission to Contemporary Clinical Trials Communications", italic=True)
    doc.add_page_break()

    # ============================================================
    # Parse the markdown body and build the docx
    # ============================================================
    print(f"Read {len(lines)} lines from {input_md}")

    i = start
    in_table = False
    table_headers = []
    table_rows = []
    current_section = ''

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Skip empty lines
        if line.strip() == '':
            i += 1
            continue

        # Title (first line starting with single #)
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            p = add_para(title_text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            p.runs[0].font.size = Pt(14)
            i += 1
            continue

        # Author/corresponding lines
        if line.startswith('**Authors**') or line.startswith('**Corresponding'):
            text = line.replace('**', '')
            add_para(text, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # Section headers
        if line.startswith('#### '):
            add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        if line.startswith('### '):
            add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('## '):
            section_title = line[3:].strip()
            current_section = section_title

            # If the document still has a Figures/Tables section at the end,
            # skip those headings because the items are already inline.
            if section_title in ('Figures', 'Tables'):
                i += 1
                continue

            add_heading(section_title, level=2)
            i += 1
            continue

        # Handle figure blocks anywhere (inline or legacy Figures section)
        if line.startswith('**Fig.'):
            m = re.match(r'\*\*Fig\.\s*(\d+)\*\*\s*(.*)', line)
            if m:
                fig_num = int(m.group(1))
                caption = m.group(2).strip()
                # Find the image reference
                j = i + 1
                img_path = None
                while j < len(lines):
                    stripped = lines[j].strip()
                    if stripped.startswith('!['):
                        img_m = re.match(r'!\[.*?\]\((.*?)\)', stripped)
                        if img_m:
                            img_path = img_m.group(1)
                        j += 1
                        break
                    elif stripped == '':
                        j += 1
                    else:
                        break
                if img_path:
                    add_figure(img_path, caption, fig_num)
                i = j
                continue
            i += 1
            continue

        # Skip bare image references that are not part of a figure block
        if line.strip().startswith('!['):
            i += 1
            continue

        # Flush pending table if needed
        if in_table and not line.strip().startswith('|'):
            add_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        # Table detection
        if line.strip().startswith('|'):
            # Separator row
            if re.match(r'^[\|\-\s:]+$', line.strip()):
                i += 1
                continue

            cells = [c.strip() for c in line.strip().split('|')[1:-1]]

            if not in_table:
                # Check if next line is separator
                if i + 1 < len(lines) and re.match(r'^[\|\-\s:]+$', lines[i+1].strip()):
                    in_table = True
                    table_headers = cells
                    table_rows = []
                    i += 2  # skip header and separator
                    continue
            else:
                table_rows.append(cells)
                # Check if next line is still table
                if i + 1 >= len(lines) or not lines[i+1].strip().startswith('|'):
                    add_table(table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                i += 1
                continue

            i += 1
            continue

        # Bold table caption
        if line.strip().startswith('**Table'):
            text = line.strip().replace('**', '')
            add_para(text, bold=True)
            i += 1
            continue

        # Math blocks ($$...$$)
        if line.strip().startswith('$$'):
            math_text = line.strip().replace('$$', '').strip()
            if not math_text:
                math_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('$$'):
                    math_lines.append(lines[i].strip())
                    i += 1
                math_text = ' '.join(math_lines)
                i += 1  # skip closing $$
            else:
                i += 1
            if math_text:
                add_rich_para('$$' + math_text + '$$', align=WD_ALIGN_PARAGRAPH.CENTER)
            continue

        # Bullet lists (top level)
        if line.strip().startswith(('- ', '* ')):
            text = line.strip()[2:]
            add_rich_para(text, style='List Bullet')
            i += 1
            continue

        # Indented bullet lists
        if re.match(r'^  +(- |\* )', line):
            text = line.strip()[2:]
            add_rich_para(text, style='List Bullet 2')
            i += 1
            continue

        # Numbered lists
        m = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if m:
            num = int(m.group(1))
            text = m.group(2)

            # References (number + plain text, no markdown formatting)
            if current_section == 'References':
                p = doc.add_paragraph()
                run = p.add_run(f'{num}. {text}')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                i += 1
                continue
            else:
                add_rich_para(text, style='List Number')
                i += 1
                continue

        # Regular paragraph
        para_text = line.strip()
        if para_text:
            add_rich_para(para_text)
        i += 1

    # Flush any remaining table
    if in_table:
        add_table(table_headers, table_rows)

    # Save
    doc.save(output_docx)
    _rewrite_docx_timestamps(output_docx)
    print(f'\nSaved CCT docx to {output_docx}')

    # Word count
    word_count = 0
    for p in doc.paragraphs:
        word_count += len(p.text.split())
    print(f'Total word count (all content): ~{word_count}')
    print(f'File size: {os.path.getsize(output_docx) / 1024:.0f} KB')

    # Generate a separate Highlights docx for the online submission system
    highlights_docx = output_docx.rsplit('.docx', 1)[0] + '_highlights.docx'
    hdoc = Document()
    hdoc.core_properties.created = datetime(2025, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
    hdoc.core_properties.modified = datetime(2025, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
    hsec = hdoc.sections[0]
    hsec.top_margin = Cm(2.54)
    hsec.bottom_margin = Cm(2.54)
    hsec.left_margin = Cm(2.54)
    hsec.right_margin = Cm(2.54)
    hstyle = hdoc.styles['Normal']
    hstyle.font.name = 'Times New Roman'
    hstyle.font.size = Pt(12)
    htitle = hdoc.add_heading('Highlights', level=1)
    for r in htitle.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = True
    with open(input_md, 'r', encoding='utf-8') as f:
        md_text = f.read()
    h_match = re.search(r'## Highlights\n\n(.*?)\n\n## ', md_text, re.S)
    highlights_text = h_match.group(1).strip() if h_match else ''
    for line in highlights_text.splitlines():
        line = line.strip()
        if line.startswith('* '):
            p = hdoc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            p = hdoc.add_paragraph(line)
        if p:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
    hdoc.save(highlights_docx)
    _rewrite_docx_timestamps(highlights_docx)
    print(f'Saved Highlights docx to {highlights_docx}')
    print(f'Highlights file size: {os.path.getsize(highlights_docx) / 1024:.0f} KB')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert CCT markdown to docx')
    parser.add_argument('input_md', nargs='?', default='05_paper_cct.md')
    parser.add_argument('output_docx', nargs='?', default='KOTHA_Framework_CCT.docx')
    args = parser.parse_args()
    build(args.input_md, args.output_docx)
