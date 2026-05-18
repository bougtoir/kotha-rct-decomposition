#!/usr/bin/env python3
"""
デスフルラン文献の批判的吟味レポート生成スクリプト（日本語版）
臨床麻酔学会発表資料
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def add_superscript_text(paragraph, text):
    """Parse text with {N} markers and add superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = paragraph.add_run(part)
    return paragraph


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def create_report():
    doc = Document()

    # --- Style setup ---
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Yu Gothic'
        h_style.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

    # === TITLE PAGE ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(72)
    run = p.add_run('デスフルラン優位性を主張する文献の\n批判的吟味')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(24)
    run = p.add_run('― セボフルランに対する真の優位性は存在するか ―')
    run.font.size = Pt(14)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(48)
    run = p.add_run('臨床麻酔学会発表用資料（ドラフト）')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run('EU規則 (EU) 2024/573 によるデスフルラン使用禁止（2026年1月1日施行）を踏まえた\n10のシチュエーションにおけるエビデンスの系統的検証')
    run.font.size = Pt(10)

    doc.add_page_break()

    # === EXECUTIVE SUMMARY ===
    doc.add_heading('要旨', level=1)
    summary_text = (
        'デスフルランは地球温暖化係数（GWP）2,540とセボフルラン（GWP ≈ 130）の約20倍であり、'
        'EU規則 (EU) 2024/573により2026年1月1日から原則使用禁止となる。'
        'デスフルラン優位性を主張する文献は発売当初から多数存在するが、その多くに方法論的問題が認められる。'
        '本稿では「デスフルランを使用すべき症例」として主張される10のシチュエーションについて、'
        '代表的な論文を特定し、方法論的問題点を検証した。'
        '加えて、各論文に対して発表された批判的レター（Letter to the Editor）およびその著者回答を調査し、'
        '「引用の非対称性」—原著のみが繰り返し引用され、批判とその回答は無視される現象—を明らかにした。'
    )
    p = doc.add_paragraph(summary_text)
    p.paragraph_format.space_after = Pt(12)

    # Key conclusion box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('結論: ')
    run.font.bold = True
    run = p.add_run(
        '10シチュエーション全てにおいて、セボフルランに対するデスフルランの臨床的優位性を支持する'
        '質の高いエビデンスは見出されなかった。「デスフルランを使用すべき症例」は事実上存在しない。'
    )

    doc.add_page_break()

    # === TABLE OF CONTENTS HEADER ===
    doc.add_heading('目次', level=1)
    toc_items = [
        '1. 肥満患者における覚醒の速さ',
        '2. 高齢者における認知機能保存',
        '3. 長時間手術における覚醒の速さ',
        '4. 脳神経外科における術後神経学的評価',
        '5. 日帰り手術における退院時間短縮',
        '6. 小児における覚醒時興奮の軽減',
        '7. 心臓手術における心筋保護効果',
        '8. 呼吸器疾患患者における安全性',
        '9. 肝・腎機能障害患者における安全性',
        '10. 筋疾患患者における安全性',
        '総括: 引用の非対称性と産業バイアス',
        '文献一覧',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =====================================================
    # SCENARIO 1: OBESITY
    # =====================================================
    doc.add_heading('1. 肥満患者における覚醒の速さ', level=1)

    doc.add_heading('主張', level=2)
    add_superscript_text(
        doc.add_paragraph(),
        'デスフルランは血液/ガス分配係数が低く（0.42 vs セボフルラン 0.65）、'
        '肥満患者でもウォッシュアウトが速いため覚醒が早い。{1,2}'
    )

    doc.add_heading('代表的論文', level=2)

    # Paper 1a
    doc.add_heading('La Colla et al. (2007)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'La Colla, 2007, "Faster wash-out and recovery for desflurane vs '
        'sevoflurane in morbidly obese patients when no premedication is used."{1}'
    )

    p = doc.add_paragraph()
    run = p.add_run('方法論的問題: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    items = [
        '非再呼吸回路（Mapleson D回路）を使用し、FGF 8 L/minのハイフロー条件下で比較。'
        '臨床現場で一般的に使用される低流量循環式回路（1-2 L/min）とは条件が大きく異なる。',
        'ハイフロー条件下ではデスフルランの血液/ガス分配係数の差が最大限に発揮されるため、'
        '臨床的に意味のある差が人工的に拡大された可能性がある。',
        '前投薬なし（no premedication）の条件であり、ミダゾラムやフェンタニル等の'
        '前投薬が一般的な実臨床とは乖離している。',
        '手術時間30分と短く、長時間手術での肥満患者における脂肪組織への蓄積の影響が反映されていない。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Paper 1b
    doc.add_heading('Kaur et al. (2013)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Kaur, 2013, "Hemodynamics and early recovery characteristics of '
        'desflurane versus sevoflurane in bariatric surgery."{2}'
    )

    p = doc.add_paragraph()
    run = p.add_run('方法論的問題: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    items = [
        'フレッシュガスフロー（FGF）が方法に記載されていない。'
        '「closed circuit breathing system」とのみ記載されており、術中維持のFGFも覚醒時のFGFも不明。',
        '覚醒時にデスフルラン群のみハイフローでウォッシュアウトした可能性を排除できない。'
        'これは薬剤の薬理学的特性ではなく、プロトコルの差異による偽の結果を生む。',
        'BIS 40-60で管理と記載されているが、揮発性麻酔薬の濃度は具体的に報告されていない。',
        'n=20/群と少数であり、多重比較の補正がなされていない。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Counter-evidence
    doc.add_heading('反証', level=3)
    p = doc.add_paragraph()
    run = p.add_run('反証: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Arain, 2005, "Choice of volatile anesthetic for the morbidly obese patient: '
        'sevoflurane or desflurane."{3}'
    )
    p = doc.add_paragraph(
        'BIS-guided titration（BIS 45-50を目標）で麻酔深度を揃えた場合、'
        'デスフルランとセボフルランの間に覚醒時間・認知機能回復に有意差はなかった。'
    )

    doc.add_heading('批判レター・エディトリアル', level=3)
    p = doc.add_paragraph()
    run = p.add_run('エディトリアル: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Eger & Shafer, 2005, "The complexity of recovery from anesthesia."{4}'
    )
    p = doc.add_paragraph(
        'Arain論文に付随するエディトリアルで、Eger & Shaferは「覚醒の速さ」が'
        '臨床的アウトカム（退院時間、合併症率、患者満足度）に直結しないことを指摘。'
        '数分の覚醒時間の差は臨床的に無意味であり、より重要なのは退院基準を満たすまでの時間であると論じた。'
    )

    doc.add_heading('FI-FA差の問題', level=3)
    p = doc.add_paragraph(
        'デスフルランは体内でほぼ代謝されない（代謝率 0.02%）ため、'
        '定常状態では吸気濃度（FI）と呼気濃度（FA）は一致するはずである。'
        'しかし臨床現場ではFI-FA差が持続的に観察される。'
        'これは肥満患者の大量の脂肪組織への分配が継続していることを意味し、'
        '「低い血液/ガス分配係数だから肥満でも速い」という主張の根本的矛盾を示す。'
        'デスフルランの脂肪/血液分配係数は27（セボフルラン: 48）であり、'
        '血液/ガス分配係数の差ほど脂肪/血液分配係数の差は大きくない。'
    )

    # Meta-analysis box
    doc.add_heading('メタアナリシス', level=3)
    add_superscript_text(
        doc.add_paragraph(),
        'Singh, 2017, "Comparison of the Recovery Profile between Desflurane and Sevoflurane '
        'in Patients Undergoing Bariatric Surgery."{5} '
        'メタアナリシスでは抜管時間に統計的有意差を認めたものの、その差は臨床的に意味のある閾値（数分）に過ぎず、'
        '退院時間やPACU滞在時間には差がないと結論。'
        '出版バイアスが検出されている（Egger検定 p=0.02）。'
    )

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=2)
    p = doc.add_paragraph()
    run = p.add_run('La Colla 2007: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】明示的な主要エンドポイントの事前指定なし。wash-out時間、抜管時間、'
        '覚醒時間等の複数アウトカムを報告しているが、いずれが主要エンドポイントかの記載がない。'
        '【多重比較検定】複数アウトカムに対する多重比較補正は実施されていない。'
    )
    p = doc.add_paragraph()
    run = p.add_run('Kaur 2013: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】明示的な主要エンドポイントの事前指定なし。血行動態と覚醒特性の両方を比較しているが、'
        '主要アウトカムの指定がない。'
        '【多重比較検定】n=20/群で複数アウトカムを比較しているが、多重比較補正は実施されていない。'
    )
    p = doc.add_paragraph()
    run = p.add_run('Singh 2017（メタアナリシス）: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】開眼時間と抜管時間を主要アウトカムとして明記。ただし退院時間・PACU滞在時間等の'
        '臨床的に重要なアウトカムは副次的扱い。'
        '【多重比較検定】5つのRCTの統合解析で複数アウトカムを比較しているが、'
        '個々のアウトカム間のBonferroni補正は実施されていない。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 2: ELDERLY POCD
    # =====================================================
    doc.add_heading('2. 高齢者における認知機能保存', level=1)

    doc.add_heading('主張', level=2)
    add_superscript_text(
        doc.add_paragraph(),
        'デスフルランは覚醒が速いため、高齢者の術後認知機能障害（POCD）が軽減される。{6}'
    )

    doc.add_heading('代表的論文', level=2)
    doc.add_heading('Tachibana et al. (2015)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Tachibana, 2015, "Recovery of postoperative cognitive function in elderly patients '
        'after a long duration of desflurane anesthesia: a pilot study."{6}'
    )

    p = doc.add_paragraph()
    run = p.add_run('方法論的問題: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    items = [
        'n=42（D群21、S群21）のパイロット研究であり、検出力が不十分。',
        '除外基準に「severe dementiaの診断がある患者」とあるが、'
        '具体的な認知機能スクリーニング閾値（MMSE ≤ 23等）が設定されていない。'
        'ベースラインで軽度〜中等度の認知機能低下がある患者が含まれた可能性がある。',
        'デスフルラン群で術後MMSEが術前より「有意に改善」と報告しているが、'
        'その変化量は約1点に過ぎず、MMSEの測定誤差（±2-3点）の範囲内。'
        '臨床的に有意義とされるMMSE変化量は2点以上。',
        'p=0.048とぎりぎりの有意水準であり、多重比較補正を行えば有意差は消失する。',
        '測定時点が術後24時間の1点のみであり、残存麻酔薬、疼痛、'
        '疼痛管理の影響を受けやすいタイミング。',
        'デスフルラン 3.5%（約0.58 MAC）とセボフルラン 1.0%（約0.50 MAC）で'
        '等MAC濃度ではなく、セボフルラン群が相対的に深い麻酔を受けた可能性。',
        'BISモニタリングによる麻酔深度の標準化が行われていない。',
        'ベースラインのMMSEスコアが低い症例（認知症の診断はないが実質的な認知機能低下がある患者）を'
        '除外して再解析すると有意差が消失することが検証されている。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('批判レターと著者回答 【重要】', level=3)

    p = doc.add_paragraph()
    run = p.add_run('批判レター: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Ding, 2015, "Desflurane anesthesia and postoperative cognitive function."{7}'
    )

    p = doc.add_paragraph('Ding et al.は以下の3点を指摘:')
    items_ding = [
        'アルコール歴、教育歴、深麻酔等のPOCDリスク因子が考慮されていない。',
        'MMSE変化量が約1点で、すべてカットオフ値24点以上であり、臨床的に無意味。',
        '術後24時間の1時点のみの測定では残存麻酔薬等の影響が大きく、認知機能の正確な評価には不十分。',
    ]
    for item in items_ding:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('著者回答: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Tachibana, 2017, "In reply: Desflurane anesthesia and cognitive function."{8}'
    )

    p = doc.add_paragraph('著者らは回答で以下を認めた（事実上の譲歩）:')
    items_reply = [
        '「教育歴、経済的状況、精神疾患等の因子を調べていない」',
        '「認知機能の結果は慎重に解釈すべきである」（"The results for cognitive function should be interpreted carefully"）',
        '「デスフルラン群とセボフルラン群のMMSEスコアの差については言及していない」'
        '（"We did not mention the difference between MMSE scores in the desflurane group and sevoflurane group"）',
        '「両群のMMSEスコアは同等であった」（"Our data showed comparable MMSE scores in the two groups"）',
        '「セボフルランも高齢者の術後MMSE保存に適した麻酔薬である」と認めた',
    ]
    for item in items_reply:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('引用の非対称性: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run(
        '原著（Tachibana 2015）は31回引用されているのに対し、'
        '批判レター（Ding 2015）は2回、著者の譲歩回答（Tachibana 2017）はほぼ引用されていない。'
        '「デスフルランは高齢者の認知機能を保存する」という主張は、著者自身が事実上撤回しているにもかかわらず、'
        '後続の文献では原著の結論のみが引用され続けている。'
    )

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Tachibana 2015: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】明示的な主要エンドポイントの事前指定なし。パイロット研究と銘打っているが、'
        'MMSE変化量、覚醒時間、抜管時間等の複数アウトカムを報告し、'
        'いずれが主要エンドポイントかの記載がない。MMSE変化量が「有意」として強調されているが、'
        '事前に主要アウトカムとして指定された形跡がない。'
        '【多重比較検定】MMSE、覚醒時間、抜管時間等の複数アウトカムに対する'
        '多重比較補正（Bonferroni等）は実施されていない。'
        'p=0.048というぎりぎりの有意水準は、多重比較補正を行えば有意差が消失するレベル。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 3: LONG DURATION SURGERY
    # =====================================================
    doc.add_heading('3. 長時間手術における覚醒の速さ', level=1)

    doc.add_heading('主張', level=2)
    add_superscript_text(
        doc.add_paragraph(),
        'デスフルランは組織溶解度が低いため、長時間手術後でもウォッシュアウトが速く、'
        '覚醒が早い。手術時間が3時間を超えると、セボフルランとの差が顕著になる。{9}'
    )

    doc.add_heading('方法論的問題', level=2)
    items = [
        '長時間手術後の覚醒時間の差（数分〜十数分）は、'
        'PACU滞在時間や退院時間には反映されないことがメタアナリシスで示されている。',
        'context-sensitive half-timeの理論的な差は、'
        '臨床的なFGF、残存オピオイド、筋弛緩薬の拮抗タイミング等の実臨床因子に埋もれる。',
        'BIS-guided titrationが行われていない研究では、'
        '麻酔深度の不均衡が覚醒時間の差を生む交絡因子となる。',
        '「覚醒の速さ」が臨床的に有益であるためには、'
        '抜管後の気道合併症、PONV、疼痛管理等の総合的な回復の質で評価すべきだが、'
        '多くの研究は抜管時間のみを主要エンドポイントとしている。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=2)
    p = doc.add_paragraph(
        '長時間手術におけるデスフルラン優位性を検証した個々のRCTの多くは、'
        '抜管時間を主要エンドポイントとして設定しているが、退院時間やPACU滞在時間は'
        '副次的エンドポイントまたは未測定であることが多い。'
        '覚醒時間（数分の差）を主要EPとすること自体が、臨床的に重要なアウトカムからの'
        '逸脱であり、結果の臨床的意義を過大評価させる研究デザインとなっている。'
        '多重比較補正については、複数の回復指標を同時に検定している研究が大半だが、'
        'Bonferroni等の補正を明示している研究は見当たらない。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 4: NEUROSURGERY
    # =====================================================
    doc.add_heading('4. 脳神経外科における術後神経学的評価', level=1)

    doc.add_heading('主張', level=2)
    add_superscript_text(
        doc.add_paragraph(),
        'デスフルランは開頭術後の速やかな覚醒により、早期の神経学的評価（wake-up test）が可能。{10}'
    )

    doc.add_heading('代表的論文', level=2)
    doc.add_heading('Dube et al. (2015)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Dube, 2015, "Comparison of intraoperative brain condition, hemodynamics and '
        'postoperative recovery between desflurane and sevoflurane in patients undergoing '
        'supratentorial craniotomy."{10}'
    )

    p = doc.add_paragraph()
    run = p.add_run('方法論的問題: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    items = [
        '抜管までの時間に有意差を認めたが（D群 8.1±2.1分 vs S群 11.3±3.4分）、'
        'GCS回復時間や退室時間に有意差なし。',
        'NACCSガイドラインによれば、この程度の覚醒時間の差は'
        '神経学的評価のタイミングに臨床的な影響を与えない。',
        'デスフルランの気道刺激性により咳嗽・バッキングのリスクがあり、'
        '開頭術後の頭蓋内圧上昇の危険性がある。',
        'NICEエビデンスレビュー（2024年）は、脳外科手術においても'
        'デスフルランの臨床的優位性を示す質の高いエビデンスはないと結論。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('メタアナリシス', level=3)
    add_superscript_text(
        doc.add_paragraph(),
        '2025年のメタアナリシス, "Comparison of desflurane and sevoflurane as maintenance '
        'inhalational anaesthetic agents for adult patients undergoing neurosurgeries."{11} '
        '抜管時間の統計的差異はあるが臨床的転帰に差はないと結論。'
    )

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Dube 2015: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】明示的な主要エンドポイントの事前指定なし。'
        '術中脳条件、血行動態、術後回復の3領域にわたる複数アウトカムを比較しているが、'
        'いずれが主要エンドポイントかの記載がない。抜管時間の差のみが強調されている。'
        '【多重比較検定】脳弛緩度、血行動態パラメータ、抜管時間、GCS回復時間、退室時間等の'
        '多数のアウトカムを比較しているが、多重比較補正は実施されていない。'
    )
    p = doc.add_paragraph()
    run = p.add_run('2025年メタアナリシス（PMID: 40046703）: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】抜管時間を主要アウトカムとしているが、臨床的に重要な転帰'
        '（ICU滞在時間、退院時間、神経学的合併症率）は副次的扱い。'
        '【多重比較検定】メタアナリシスの標準的手法に従っているが、'
        '複数アウトカム間の多重比較に対するBonferroni補正は明示されていない。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 5: AMBULATORY
    # =====================================================
    doc.add_heading('5. 日帰り手術における退院時間短縮', level=1)

    doc.add_heading('主張', level=2)
    p = doc.add_paragraph(
        'デスフルランの速やかな覚醒により日帰り手術のターンオーバーが改善し、'
        '早期退院が可能。'
    )

    doc.add_heading('代表的論文と方法論的問題', level=2)
    doc.add_heading('White et al. (2009)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'White, 2009, "Desflurane versus sevoflurane for maintenance of outpatient anesthesia: '
        'the effect on early versus late recovery and perioperative coughing."{12}'
    )

    items = [
        '早期覚醒（eye opening, following commands）はデスフルラン群で有意に速かったが、'
        '退院基準到達時間（fast-tracking criteria）には有意差なし。',
        'デスフルラン群で周術期の咳嗽（coughing）発生率が有意に高い（67% vs 33%, p<0.01）。',
        '日帰り手術における臨床的ボトルネックは覚醒ではなくPONV、疼痛管理、'
        '退院基準充足であり、数分の覚醒の差は退院ワークフローに影響しない。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('系統的レビュー', level=3)
    add_superscript_text(
        doc.add_paragraph(),
        'Gupta, 2004, "Comparison of Recovery Profile After Ambulatory Anesthesia with Propofol, '
        'Isoflurane, Sevoflurane and Desflurane: A Systematic Review."{13} '
        '退院時間に有意差なし。'
    )

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=3)
    p = doc.add_paragraph()
    run = p.add_run('White 2009: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】明示的な単一の主要エンドポイントの事前指定なし。'
        '「early vs late recovery」および「perioperative coughing」の両方を'
        'タイトルに掲げ、開眼時間、指示動作までの時間、抜管時間、fast-track基準到達時間、'
        '咳嗽発生率の少なくとも5つのアウトカムを報告。いずれが主要EPかの記載がない。'
        '【多重比較検定】5つ以上のアウトカムを比較しているが、'
        '多重比較補正（Bonferroni等）は実施されていない。'
        'fast-track基準（実際の臨床的ボトルネック）で差がなく、'
        '開眼時間等の代理エンドポイントでのみ差が出ている点は、'
        '多重比較の文脈で特に問題がある。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 6: PEDIATRIC
    # =====================================================
    doc.add_heading('6. 小児における覚醒時興奮の軽減', level=1)

    doc.add_heading('主張', level=2)
    p = doc.add_paragraph(
        'デスフルランは小児の覚醒時興奮（emergence agitation: EA）を'
        'セボフルランより軽減する可能性がある。'
    )

    doc.add_heading('メタアナリシスの結論', level=2)
    doc.add_heading('Lim et al. (2016)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Lim, 2016, "Comparison of the incidence of emergence agitation and emergence times between '
        'desflurane and sevoflurane anesthesia in children: A systematic review and meta-analysis."{14}'
    )

    items = [
        'EA発生率: デスフルラン群 25% vs セボフルラン群 25%で同等。',
        'EAの持続時間はデスフルラン群で短い傾向があるが、臨床的に意味のある差ではない。',
        'デスフルランは気道刺激性が強く小児のマスク導入には適さない。'
        '導入はセボフルランで行い維持にデスフルランを使用する二段階プロトコルが必要となり、'
        'セボフルラン単独使用に対する実用的優位性がない。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Lim 2016（メタアナリシス）: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】EA発生率と覚醒時間を主要アウトカムとして明記。'
        'メタアナリシスとしては適切にアウトカムが定義されている。'
        '【多重比較検定】メタアナリシスの標準的統計手法に従っており、'
        '異質性のI²統計量も報告。ただし臨床的結論はデスフルランの優位性なし（EA発生率同等）。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 7: CARDIAC
    # =====================================================
    doc.add_heading('7. 心臓手術における心筋保護効果', level=1)

    doc.add_heading('主張', level=2)
    p = doc.add_paragraph(
        '揮発性麻酔薬は麻酔プレコンディショニング（APC）・ポストコンディショニング効果を持ち、'
        '心臓手術における心筋保護に有益。デスフルランもこの効果を持つ。'
    )

    doc.add_heading('方法論的問題', level=2)
    doc.add_heading('Sivanna et al. (2015)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'Sivanna, 2015, "A comparative study of pharmacological myocardial protection between '
        'sevoflurane and desflurane at anaesthetic doses in patients undergoing off pump '
        'coronary artery bypass grafting surgery."{15}'
    )

    items = [
        '両群ともトロポニンT値の有意な上昇を認めたが、群間差なし。',
        '心筋保護効果はセボフルランの方が豊富なエビデンスを有する。'
        'セボフルランのプレコンディショニング効果は複数のRCTとメタアナリシスで確認されており、'
        'デスフルランが同等またはそれ以上であるとするエビデンスは不十分。',
        'JCVA 2017のエディトリアルでは、揮発性麻酔薬のクラス効果としての'
        '心筋保護はTIVA（プロポフォール）との比較では有益だが、'
        '揮発性麻酔薬間の差は明確でないと論じている。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('CABG試験における長期予後: De Hert 2009（VACMAN）', level=2)
    p = doc.add_paragraph()
    run = p.add_run('論文情報: ')
    run.font.bold = True
    add_superscript_text(
        p,
        'De Hert, 2009, "A comparison of volatile and non volatile agents for cardioprotection '
        'during on-pump coronary surgery."{16}'
    )
    p = doc.add_paragraph(
        '多施設前向き研究（VACMAN: Volatile Anesthetics and Cardioprotection '
        'Multicentre ANalysis）において、CABG患者の1年死亡率が報告された。'
    )

    # Mortality table
    mort_table = doc.add_table(rows=4, cols=2)
    mort_table.style = 'Table Grid'
    mort_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['麻酔法', '1年死亡率']
    for i, h in enumerate(headers):
        cell = mort_table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1A477A')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    mort_data = [
        ['デスフルラン', '6.9%'],
        ['セボフルラン', '3.3%'],
        ['TIVA（プロポフォール）', '12.3%'],
    ]
    for row_idx, row_data in enumerate(mort_data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = mort_table.rows[row_idx].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p = doc.add_paragraph(
        'デスフルラン群の1年死亡率（6.9%）はセボフルラン群（3.3%）の約2倍であった。'
        'TIVA群（12.3%）と比較すれば両揮発性麻酔薬とも良好だが、'
        '揮発性麻酔薬間の比較ではデスフルランの方が予後不良であり、'
        '「デスフルランが心臓手術で有利」という主張と明確に矛盾する。'
    )
    p = doc.add_paragraph(
        'この結果は、心筋保護効果が揮発性麻酔薬のクラス効果であること、'
        'さらにセボフルランが同クラス内でもデスフルランより優れている可能性を示唆する。'
    )

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Sivanna 2015: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】トロポニンT値を主要アウトカムとして比較しているが、'
        '明示的な「primary endpoint」の記載がない。血行動態指標等の複数アウトカムも同時に評価。'
        '【多重比較検定】複数の血行動態パラメータとバイオマーカーを比較しているが、'
        '多重比較補正は実施されていない。'
    )
    p = doc.add_paragraph()
    run = p.add_run('De Hert 2009: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】多施設前向きコホート研究であり、1年死亡率が主要アウトカム。'
        '事前登録された臨床試験としてのデザイン。'
        '【多重比較検定】3群比較（デスフルラン、セボフルラン、TIVA）であるが、'
        '多重比較補正の詳細は論文本文中で確認が必要。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 8: RESPIRATORY
    # =====================================================
    doc.add_heading('8. 呼吸器疾患患者における安全性', level=1)

    doc.add_heading('主張', level=2)
    p = doc.add_paragraph(
        'デスフルランの低い血液溶解度により速やかなウォッシュアウトが可能で、'
        '呼吸器疾患患者でも安全に使用できる。'
    )

    doc.add_heading('重大な反証: 気道刺激性', level=2)
    items = [
        'デスフルランはTRPA1チャネルを活性化し、喉頭C線維を刺激する。'
        'これにより咳嗽、喉頭痙攣、気管支痙攣のリスクが上昇する。',
        'COPD・喘息患者では、デスフルランの気道刺激性は特に問題となる。'
        '気管支攣縮の既往がある患者には禁忌に近い。',
        '日本のDPCデータベースを用いた大規模後ろ向きコホート研究（n=40,442）では、'
        'COPD・喘息患者における術後肺合併症に有意差はなかったが、'
        'デスフルランが安全であることを積極的に示す根拠にはならない。',
        'セボフルランは気道刺激性が低く、気管支拡張作用があるため、'
        '呼吸器疾患患者ではセボフルランが明確に優位。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('コホート研究', level=3)
    add_superscript_text(
        doc.add_paragraph(),
        'DPCデータベース研究, 2026, "Postoperative pulmonary complications of desflurane- versus '
        'sevoflurane-based general anesthesia in patients with COPD or asthma."{17} '
        'COPD 24,243例、喘息 16,199例の解析で、'
        'デスフルラン vs セボフルランの術後肺合併症に有意差なし'
        '（COPD: aRD −0.57%, 99%CI −1.8%〜+0.60%; 喘息: aRD −0.62%, 99%CI −1.8%〜+0.59%）。'
        'ただし本研究は後ろ向きコホート研究であり非劣性試験ではない。'
        '有意差なしはデスフルランの安全性を積極的に示すものではなく、'
        'またデスフルランの優位性を示すものでもない。'
        'なお同論文Introductionでは、CABG試験における1年死亡率がデスフルラン群6.9%、'
        'セボフルラン群3.3%、TIVA群12.3%であったことが引用されている。'
    )

    doc.add_heading('主要エンドポイントと多重比較検定の検証', level=3)
    p = doc.add_paragraph()
    run = p.add_run('DPC研究 2026: ')
    run.font.bold = True
    p.add_run(
        '【主要EP】術後肺合併症を主要アウトカムとして明記。'
        '後ろ向きコホート研究としては適切にアウトカムが定義されている。'
        '【多重比較検定】COPD群と喘息群の2つの解析を実施。'
        '99%信頼区間（通常の95%ではなく）を採用しており、'
        '多重比較に対する一定の配慮がなされている。'
        'ただし非劣性マージンの事前設定はなく、「有意差なし」が「安全性の証明」と同義ではない。'
    )

    doc.add_page_break()

    # =====================================================
    # SCENARIO 9: HEPATORENAL
    # =====================================================
    doc.add_heading('9. 肝・腎機能障害患者における安全性', level=1)

    doc.add_heading('主張', level=2)
    p = doc.add_paragraph(
        'デスフルランの代謝率は0.02%と極めて低く、肝・腎機能障害患者でも安全。'
        'セボフルランはCompound Aの産生やフッ化物イオンの腎毒性が懸念される。'
    )

    doc.add_heading('方法論的問題', level=2)
    items = [
        'Compound A腎毒性: ラットで認められた腎毒性はヒトでは確認されていない。'
        '20年以上の臨床使用で、セボフルランによる臨床的に意義のある腎障害の報告はない。'
        'FDAもCompound Aに関する警告を解除している。',
        'フッ化物イオン: セボフルランから産生される無機フッ化物の'
        '腎毒性閾値（50 μmol/L）を超えることは通常の臨床使用では稀。'
        'メトキシフルランとは異なり、腎毒性のメカニズムは「面積説（AUC theory）」ではなく'
        '「peak theory」では説明できないとされている。',
        'デスフルランの「低代謝率」は理論的メリットだが、'
        'セボフルランの臨床使用でも肝・腎障害の実質的リスクは極めて低い。'
        '臨床的に意味のある差をつける根拠に乏しい。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =====================================================
    # SCENARIO 10: NEUROMUSCULAR
    # =====================================================
    doc.add_heading('10. 筋疾患患者における安全性', level=1)

    doc.add_heading('主張', level=2)
    p = doc.add_paragraph(
        'デスフルランは筋疾患（筋強直性ジストロフィー等）の患者において安全に使用できる。'
    )

    doc.add_heading('方法論的問題', level=2)
    items = [
        '全ての揮発性麻酔薬は悪性高熱のトリガーとなり得る。'
        'デスフルランに特異的な安全性の優位はない。',
        '筋強直性ジストロフィーにおけるデスフルランの安全性を示す文献は症例報告レベルであり、'
        'エビデンスレベルは極めて低い。',
        'セボフルランでも同様の症例報告が存在し、揮発性麻酔薬間の差別化は不可能。',
        '筋疾患患者にはTIVA（プロポフォール+レミフェンタニル）が推奨されることが多く、'
        '揮発性麻酔薬の選択自体が議論の対象。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =====================================================
    # SUMMARY TABLE
    # =====================================================
    doc.add_heading('総括: 10シチュエーションの検証結果', level=1)

    table = doc.add_table(rows=12, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ['シチュエーション', '代表的論文', '主な方法論的問題', '結論']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1A477A')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(4.0)
        row.cells[2].width = Cm(6.0)
        row.cells[3].width = Cm(3.0)

    data = [
        ['1. 肥満', 'La Colla 2007\nKaur 2013', 'FGF未報告、ハイフロー条件、\n臨床的に無意味な差', '優位性なし'],
        ['2. 高齢者POCD', 'Tachibana 2015', 'MMSE 1点差、除外基準不備\n著者自身が譲歩', '優位性なし'],
        ['3. 長時間手術', '複数RCT', '覚醒数分の差、退院時間に\n影響なし', '優位性なし'],
        ['4. 脳外科', 'Dube 2015', '抜管時間差あるが GCS回復・\n転帰に差なし', '優位性なし'],
        ['5. 日帰り手術', 'White 2009', '退院時間に差なし、\n咳嗽増加のデメリット', '優位性なし'],
        ['6. 小児', 'Lim 2016 MA', 'EA発生率同等、気道刺激で\nマスク導入不可', '優位性なし'],
        ['7. 心臓手術', 'Sivanna 2015\nDe Hert 2009', '心筋保護はクラス効果\nCABG 1年死亡率: D6.9% S3.3%', '優位性なし'],
        ['8. 呼吸器疾患', 'DPC研究 2026', 'TRPA1活性化による\n気道刺激が禁忌的', '劣位'],
        ['9. 肝腎障害', '理論的主張', 'Compound A毒性は\nヒトで未確認', '差なし'],
        ['10. 筋疾患', '症例報告のみ', 'MHトリガーは全VA共通\nTIVA推奨', '差なし'],
        ['【総合評価】', '', '', 'デスフルランの\n優位性なし'],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
        # Color the conclusion column
        conclusion = row_data[3]
        cell = table.rows[row_idx].cells[3]
        if '劣位' in conclusion:
            set_cell_shading(cell, 'FFCCCC')
        elif '優位性なし' in conclusion:
            set_cell_shading(cell, 'FFFFCC')
        elif '差なし' in conclusion:
            set_cell_shading(cell, 'E6E6E6')

    doc.add_paragraph()  # spacing

    doc.add_page_break()

    # =====================================================
    # CITATION ASYMMETRY
    # =====================================================
    doc.add_heading('引用の非対称性と産業バイアス', level=1)

    p = doc.add_paragraph(
        'デスフルラン優位性を主張する論文群に共通するパターンとして、以下が認められた:'
    )

    items = [
        '原著論文のみが引用され、批判レターや著者の譲歩回答は引用されない'
        '（例: Tachibana 2015は31回引用、批判レターDing 2015は2回引用、著者回答はほぼゼロ）。',
        '覚醒時間（抜管時間）を主要エンドポイントとし、退院時間、PACU滞在時間、'
        '患者報告アウトカム等の臨床的に重要なエンドポイントでは差が出ないことを軽視。',
        '【主要エンドポイントの不在】本稿で検証した論文の大半（La Colla 2007、Kaur 2013、'
        'Tachibana 2015、Dube 2015、White 2009、Sivanna 2015）において、'
        '明示的な主要エンドポイントの事前指定がなされていない。'
        '複数のアウトカムを測定した上で、都合の良い結果のみを強調する'
        '「outcome reporting bias」の疑いがある。',
        '【多重比較補正の系統的欠如】上記の全論文において、複数アウトカムに対する'
        '多重比較補正（Bonferroni、FDR等）が実施されていない。'
        'ぎりぎりの有意水準（p=0.04-0.05）で「有意差あり」と報告されている結果は、'
        '多重比較補正を適用すれば有意差が消失するレベルである。',
        'FGF（フレッシュガスフロー）を報告しない、または臨床的でないハイフロー条件を使用。'
        'デスフルランの薬理学的特性が最大限に発揮される条件設計でバイアスを生む。',
        'BIS等の脳機能モニタリングによる麻酔深度の標準化が行われていない研究が多く、'
        '麻酔深度の不均衡が交絡因子となる。',
        'パイロット研究・少数例のRCTが多く、検出力不足のまま有意差を主張。',
        'デスフルラン製造元（Baxter/現Hillrom）との利益相反が開示されていない研究がある。',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =====================================================
    # REFERENCES
    # =====================================================
    doc.add_heading('文献', level=1)
    references = [
        # {1} Scenario 1: Obesity
        'La Colla L, Albertin A, La Colla G, Mangano A. Faster wash-out and recovery for desflurane vs sevoflurane in morbidly obese patients when no premedication is used. Br J Anaesth. 2007;99(3):353-358.',
        # {2} Scenario 1: Obesity
        'Kaur A, Jain AK, Sehgal R, Sood J. Hemodynamics and early recovery characteristics of desflurane versus sevoflurane in bariatric surgery. J Anaesthesiol Clin Pharmacol. 2013;29(1):36-40.',
        # {3} Scenario 1: Counter-evidence
        'Arain SR, Barth CD, Shankar H, Ebert TJ. Choice of volatile anesthetic for the morbidly obese patient: sevoflurane or desflurane. J Clin Anesth. 2005;17(6):413-419.',
        # {4} Scenario 1: Editorial
        'Eger EI II, Shafer S. The complexity of recovery from anesthesia. J Clin Anesth. 2005;17(6):411-412.',
        # {5} Scenario 1: Meta-analysis
        'Singh PM, Borle A, McGavin J, Trikha A, Sinha A. Comparison of the Recovery Profile between Desflurane and Sevoflurane in Patients Undergoing Bariatric Surgery\u2014a Meta-Analysis of Randomized Controlled Trials. Obes Surg. 2017;27(11):3031-3039.',
        # {6} Scenario 2: Elderly POCD
        'Tachibana S, Hayase T, Osuda M, Kazuma S, Yamakage M. Recovery of postoperative cognitive function in elderly patients after a long duration of desflurane anesthesia: a pilot study. J Anesth. 2015;29:627-630.',
        # {7} Scenario 2: Critical letter
        'Ding F, Zheng L, Luo T. Desflurane anesthesia and postoperative cognitive function. J Anesth. 2015. DOI: 10.1007/s00540-015-2002-3.',
        # {8} Scenario 2: Author reply
        'Tachibana S, Hayase T, Yamakage M. In reply: Desflurane anesthesia and cognitive function. J Anesth. 2017;31:637.',
        # {9} Scenario 3: Long duration
        'Eger EI II. New inhaled anesthetics. Anesthesiology. 1994;80(4):906-922.',
        # {10} Scenario 4: Neurosurgery
        'Dube SK, Pandia MP, Chaturvedi A, Bithal P, Dash HH. Comparison of intraoperative brain condition, hemodynamics and postoperative recovery between desflurane and sevoflurane in patients undergoing supratentorial craniotomy. Saudi J Anaesth. 2015;9(2):167-173.',
        # {11} Scenario 4: Neurosurgery meta-analysis
        'Comparison of desflurane and sevoflurane as maintenance inhalational anaesthetic agents for adult patients undergoing neurosurgeries: A systematic review and meta-analysis of randomised trials. Indian J Anaesth. 2025. PMID: 40046703.',
        # {12} Scenario 5: Ambulatory
        'White PF, Tang J, Wender RH, et al. Desflurane versus sevoflurane for maintenance of outpatient anesthesia: the effect on early versus late recovery and perioperative coughing. Anesth Analg. 2009;109(2):387-393.',
        # {13} Scenario 5: Systematic review
        'Gupta A, Stierer T, Zuckerman R, Sakima N, Parker SD, Fleisher LA. Comparison of Recovery Profile After Ambulatory Anesthesia with Propofol, Isoflurane, Sevoflurane and Desflurane: A Systematic Review. Anesth Analg. 2004;98:632-641.',
        # {14} Scenario 6: Pediatric meta-analysis
        'Lim BG, Lee IO, Ahn H, et al. Comparison of the incidence of emergence agitation and emergence times between desflurane and sevoflurane anesthesia in children: A systematic review and meta-analysis. Medicine. 2016;95(38):e4927.',
        # {15} Scenario 7: Cardiac
        'Sivanna U, Joshi S, Babu B, Jagadeesh AM. A comparative study of pharmacological myocardial protection between sevoflurane and desflurane at anaesthetic doses in patients undergoing off pump coronary artery bypass grafting surgery. Indian J Anaesth. 2015;59(5):282-286.',
        # {16} Scenario 7: CABG mortality
        'De Hert SG, et al. A comparison of volatile and non volatile agents for cardioprotection during on-pump coronary surgery. Anaesthesia. 2009;64(9):953-960.',
        # {17} Scenario 8: Respiratory cohort
        'Postoperative pulmonary complications of desflurane- versus sevoflurane-based general anesthesia in patients with COPD or asthma: a nationwide retrospective cohort study. J Anesth. 2026;40:59-68.',
        # {18} Regulation
        'Regulation (EU) 2024/573 of the European Parliament and of the Council of 7 February 2024 on fluorinated greenhouse gases. Official Journal of the European Union.',
        # {19} Guidance
        'NHS England. Guidance: Desflurane decommissioning and clinical use. 2024.',
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        add_superscript_text(p, f'{{{i}}} ')
        p.add_run(ref)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            if not run.font.superscript:
                run.font.size = Pt(9)

    # Save
    output_path = os.path.join(OUTPUT_DIR, 'desflurane_critical_appraisal_ja.docx')
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    create_report()
