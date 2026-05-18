#!/usr/bin/env python3
"""
JSCE投稿用カバーレター生成スクリプト
土木学会論文集特集号（水工学）向け
"""

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"

doc = Document()

# ── ページ設定（A4）──
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# ── スタイル定義 ──
style = doc.styles["Normal"]
style.font.name = "游明朝"
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = Pt(18)
style.paragraph_format.space_after = Pt(0)
rpr = style.element.find(qn("w:rPr"))
if rpr is None:
    rpr = style.element.makeelement(qn("w:rPr"), {})
    style.element.append(rpr)
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = rpr.makeelement(qn("w:rFonts"), {})
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "游明朝")


def add_para(text, bold=False, alignment=None, space_before=0, space_after=0, font_size=10.5):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "游明朝"
    run.font.size = Pt(font_size)
    run.bold = bold
    rpr = run._element.find(qn("w:rPr"))
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "游明朝")
    return p


# ── 日付 ──
add_para("2026年5月　日", alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

# ── 宛先 ──
add_para("土木学会論文集特集号（水工学）", space_after=0)
add_para("編集小委員会 御中", space_after=18)

# ── 差出人 ──
add_para("著者名：（著者名）", space_after=0)
add_para("所　属：（所属）", space_after=0)
add_para("連絡先：（メールアドレス）", space_after=0)
add_para("電　話：（電話番号）", space_after=18)

# ── タイトル ──
add_para("【投稿論文のご送付について】", bold=True, space_after=12, font_size=11)

# ── 本文 ──
add_para(
    "拝啓　時下ますますご清栄のこととお慶び申し上げます．",
    space_after=6
)

add_para(
    "このたび，土木学会論文集特集号（水工学）に下記の論文を投稿いたしたく，"
    "原稿をお送りいたします．",
    space_after=12
)

# ── 論文情報 ──
add_para("記", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_before=6, space_after=12)

add_para("1．論文題目", bold=True, space_after=0)
add_para(
    "　計画放水発電と隣接水系地下水管理を統合した新たな洪水制御フレームワークの提案"
    "　― 小田川–高梁川流域を対象とした簡易フィージビリティ分析 ―",
    space_after=12
)

add_para("2．原稿区分", bold=True, space_after=0)
add_para(
    "　論文（理論的・実証的研究成果の統合，独創性のある完結した論文）",
    space_after=12
)

add_para("3．希望査読分冊", bold=True, space_after=0)
add_para(
    "　Ｂ１分冊（水工学）：水害・氾濫，水防災，河川計画と管理，水資源",
    space_after=12
)

add_para("4．論文の概要", bold=True, space_after=0)
add_para(
    "　本論文は，計画放水時の水力発電エネルギーを利用して隣接水系の帯水層から"
    "地下水を事前に汲み上げ，洪水時の受容空間を能動的に創出する「計画放水発電–"
    "地下水管理統合型洪水制御フレームワーク」を提案するものです．"
    "2018年西日本豪雨時の小田川–高梁川流域（岡山県倉敷市真備町）を対象とした"
    "簡易シミュレーションにより，エネルギー的自立性（発電量の約9%で地下水管理が可能），"
    "帯水層貯留容量の確保可能性（浸水量の約78%を吸収），"
    "既存構造的対策との相補性（ピーク流量32%削減）を示しました．"
    "さらに，首都圏（荒川・江戸川流域）への適用可能性を検討し，"
    "人口・産業密集地での治水安全度向上への貢献可能性を議論しています．",
    space_after=12
)

add_para("5．新規性・有用性", bold=True, space_after=0)
add_para(
    "　本論文の新規性は以下の3点にあります．\n"
    "　（1）計画放水発電と帯水層管理を統合した洪水制御の概念を世界で初めて体系的に提示した点\n"
    "　（2）エネルギー収支・帯水層貯留容量・洪水ハイドログラフの3側面から定量的にフィージビリティを評価した点\n"
    "　（3）構造的対策（合流点付替え）との統合による相補的効果を定量的に示した点\n"
    "　有用性として，既存の揚水発電技術・帯水層涵養技術の延長線上にある実現可能な提案であり，"
    "流域治水の新たなメニューとして政策的検討に値すると考えます．",
    space_after=12
)

add_para("6．関連論文の開示", bold=True, space_after=0)
add_para(
    "　本論文と関連する内容の英文原稿（Commentary形式）を"
    "Nature Water誌に投稿予定です．"
    "ただし，英文原稿は日本国内事例の詳細（小田川–高梁川・荒川の定量分析）を含まず，"
    "国際的な水ガバナンス・越境水系管理の政策的議論に焦点を当てた"
    "異なる対象読者・異なる内容構成の原稿であり，"
    "本和文論文との間に実質的な内容の重複はありません．"
    "投稿要項4(3)の4)「未発表であること」および5)「二重投稿でないこと」の"
    "条件を満たしていることを申し添えます．",
    space_after=12
)

add_para("7．利益相反", bold=True, space_after=0)
add_para(
    "　本研究に関して，開示すべき利益相反はありません．",
    space_after=12
)

add_para("8．倫理規定の遵守", bold=True, space_after=0)
add_para(
    "　著者は土木学会倫理規定（土木技術者の倫理規定）および"
    "土木学会論文集に関する倫理基準を遵守し，本論文を作成いたしました．"
    "本研究で使用したデータはすべて公開情報であり，"
    "倫理審査を要する調査は含まれておりません．",
    space_after=12
)

add_para("9．提出書類一覧", bold=True, space_after=0)
add_para(
    "　（1）原稿本体（Word形式，7頁以内）\n"
    "　（2）本カバーレター\n"
    "　（3）投稿チェックリスト\n"
    "　（4）図表ファイル（PNG形式，個別ファイル4点）",
    space_after=18
)

add_para(
    "何卒ご査読のほどよろしくお願い申し上げます．",
    space_after=6
)
add_para("敬具", alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

add_para("以上", alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=18)

# ── 保存 ──
outpath = f"{OUTDIR}/jsce_cover_letter.docx"
doc.save(outpath)
print(f"カバーレターを保存しました: {outpath}")
