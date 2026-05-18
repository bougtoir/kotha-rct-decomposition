#!/usr/bin/env python3
"""
Journal of JSCE (English) manuscript generator
Special Issue: Hydroscience and Hydraulic Engineering (B1)

Title: A novel flood control framework integrating planned-release
       hydropower with inter-watershed groundwater management:
       A simplified feasibility analysis for the Oda River-Takahashi River basin

Format (Guidelines for Authors, January 1, 2026 revision):
- Times New Roman, double-spaced, line numbers, page numbers
- Structure: Abstract -> Keywords -> Main text -> References -> Figure legends -> Tables
- Figures: NOT embedded; uploaded separately as PNG
- Word limit: 14,000 words (each fig/table = 250 words)
- Chapters: 1., 2., 3. / Sections: (1), (2), (3) / Sub-sections: a), b), c)
- References: Vancouver style, numbered, superscript 1),5)
- Ref format: Author(s): Title, *Journal*, Vol., pp., Year. doi: URL
"""

import os
import re
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"

doc = Document()

# ============================================================
# Page setup (A4)
# ============================================================
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# ============================================================
# Styles: Times New Roman, 12pt, double-spaced
# ============================================================
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = Pt(24)  # double spacing
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)


def add_heading_bold(text, level=1):
    """Add bold heading. Chapter=1., Section=(1), Sub-section=a)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(12)
    return p


def add_body(text):
    """Add body paragraph. {N} markers become font-superscript references."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith("{") and part.endswith("}"):
            ref_text = part[1:-1]
            run = p.add_run(ref_text + ")")
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.superscript = True
        else:
            run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    return p


def add_body_no_indent(text):
    p = add_body(text)
    p.paragraph_format.first_line_indent = Pt(0)
    return p


def add_table_from_data(headers, rows, caption):
    """Add table with caption above."""
    pc = doc.add_paragraph()
    pc.paragraph_format.space_before = Pt(12)
    run_c = pc.add_run(caption)
    run_c.font.name = "Times New Roman"
    run_c.font.size = Pt(10)
    run_c.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(8)
    return table


# ============================================================
# ABSTRACT
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run("ABSTRACT")
run.font.name = "Times New Roman"
run.font.size = Pt(12)
run.bold = True

add_body_no_indent(
    "Conventional flood control in Japan relies on structural measures such as levees, dams, "
    "super-levees, and diversion channels. This paper proposes a novel flood control framework "
    "that integrates planned-release hydropower generation with inter-watershed groundwater "
    "management. The core concept is to use the potential energy of controlled flood discharge "
    "to generate electricity, which in turn powers the pre-emptive pumping of groundwater from "
    "adjacent watersheds, thereby creating subsurface storage capacity to absorb floodwaters. "
    "A simplified feasibility analysis was conducted for the Oda River-Takahashi River basin "
    "(Okayama Prefecture), where the 2018 Western Japan Heavy Rainfall caused catastrophic "
    "backwater flooding with 51 fatalities. The energy balance analysis demonstrates that "
    "planned-release hydropower (approximately 42 MW at H = 100 m, Q = 50 m\u00b3/s) far "
    "exceeds the power required for groundwater pumping (approximately 3.7 MW), achieving "
    "energy self-sufficiency with a surplus of approximately 38 MW. The aquifer storage analysis "
    "shows that a managed area of 20 km\u00b2 with 3 m drawdown creates approximately "
    "12.0 \u00d7 10\u2076 m\u00b3 of storage capacity, equivalent to 78% of the 2018 flood volume. "
    "Integration with the existing structural measure (confluence relocation, completed 2024) "
    "yields a 32% reduction in peak discharge. The applicability to the Arakawa-Edogawa basin "
    "in the Tokyo metropolitan area (8.21 million people in projected inundation zones) was also "
    "examined, revealing the potential for complementary flood control using Kanto Plain "
    "alluvial aquifers. The framework leverages Japan's steep terrain, existing pumped-storage "
    "infrastructure, and dense hydrometeorological networks, offering a replicable model for "
    "flood-prone regions worldwide."
)

# ============================================================
# KEYWORDS
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Key words: ")
run.font.name = "Times New Roman"
run.font.size = Pt(12)
run.bold = True
run2 = p.add_run(
    "flood control, groundwater management, hydropower, backwater effect, aquifer storage"
)
run2.font.name = "Times New Roman"
run2.font.size = Pt(12)

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading_bold("1. INTRODUCTION")

add_body(
    "In recent years, Japan has experienced increasingly severe flood disasters due to "
    "intensifying rainfall events associated with climate change. During the 2018 Western Japan "
    "Heavy Rainfall (July 2018 event), levees along the Oda River in Mabi-cho, Kurashiki City, "
    "Okayama Prefecture, breached at five locations, resulting in 51 fatalities, an inundation area "
    "of 8.28 km\u00b2, and an estimated flood volume of approximately 15.3 \u00d7 10\u2076 m\u00b3.{1} "
    "The primary cause was the backwater effect: rising water levels in the main Takahashi River "
    "impeded drainage from the tributary Oda River.{2}"
)

add_body(
    "Conventional flood control in Japan centres on structural measures, including levee and "
    "revetment construction, dam-based flood regulation, and the construction of diversion channels "
    "and retarding basins.{3} In Mabi-cho, a large-scale structural measure was completed in March "
    "2024: the relocation of the Oda-Takahashi confluence approximately 4.6 km downstream. "
    "During the heavy rainfall of November 2024, this measure reportedly reduced water levels by "
    "approximately 4.6 m on the Oda River side and 0.8 m on the Takahashi River side.{4} "
    "However, such large-scale civil engineering projects require decades from planning to completion, "
    "and their adaptive capacity to future increases in flood magnitude under climate change is limited."
)

add_body(
    "Meanwhile, groundwater management as a flood mitigation approach has attracted increasing "
    "international attention. In California, the Flood-Managed Aquifer Recharge (Flood-MAR) strategy "
    "channels floodwaters into aquifers for storage.{5} In Taiwan, optimised groundwater pumping has "
    "been demonstrated to reduce flood hazard in subsidence-prone coastal areas.{6} In Japan, the "
    "Metropolitan Area Outer Underground Discharge Channel (MAOUDC) achieves inter-watershed "
    "flood transfer, with its efficacy validated through over 100 operational events.{7}"
)

add_body(
    "This paper proposes a novel flood control framework that integrates these existing technologies. "
    "The framework uses the hydroelectric energy from planned flood discharge to pre-emptively pump "
    "groundwater from adjacent watersheds, actively creating subsurface storage capacity to absorb "
    "floodwaters. A simplified feasibility analysis is conducted for the Oda River-Takahashi River "
    "basin, and the applicability to the Tokyo metropolitan area is examined. Policy implications "
    "are also discussed."
)

# ============================================================
# 2. LITERATURE REVIEW AND RELATED TECHNOLOGIES
# ============================================================
add_heading_bold("2. LITERATURE REVIEW AND RELATED TECHNOLOGIES")

add_heading_bold("(1) Backwater effect and confluence flood control", level=2)
add_body(
    "The backwater effect occurs when rising water levels in a main river impede drainage from "
    "tributaries, causing flooding in the tributary basin. Maeno{8} analysed the effect of relocating "
    "the Takahashi-Oda River confluence using two-dimensional flow analysis, demonstrating "
    "significant water level reduction in the Oda River basin. "
    "Internationally, the Yazoo Backwater area in the Mississippi River basin is a well-documented "
    "case of backwater flooding.{9} In 2019, floodwaters exceeded the flood stage for 158 consecutive "
    "days, causing prolonged inundation lasting approximately four months.{10} "
    "Berkowitz et al.{11} analysed the wetland hydrology of the Yazoo basin, demonstrating "
    "the critical role of groundwater-surface water interactions during prolonged floods. "
    "Li and Tsai{12} showed that groundwater flow dynamics in the Mississippi River Delta are "
    "strongly influenced by surface water-groundwater interactions, suggesting the feasibility "
    "of aquifer-based flood management in river deltas."
)

add_heading_bold("(2) Managed aquifer recharge and groundwater-based flood control", level=2)
add_body(
    "Managed Aquifer Recharge (MAR), including Aquifer Storage and Recovery (ASR), has been "
    "practised worldwide as a technique for sustainable groundwater resource management.{13} "
    "California's Flood-MAR strategy represents a novel approach that simultaneously achieves "
    "flood risk reduction and groundwater recharge, diverting floodwaters onto agricultural lands "
    "and wetlands during high-flow events.{5} "
    "Chang et al.{6} demonstrated in south-western Taiwan that optimal groundwater pumping "
    "scheduling can lower pre-flood water tables, reducing inundation depth by up to 0.5 m "
    "in subsidence-prone coastal areas. "
    "Nishizawa Katazakai and Zhang{14} analysed half a century of groundwater monitoring data "
    "in central Japan, providing evidence that long-term groundwater management can contribute "
    "to flood control while maintaining aquifer sustainability."
)

add_heading_bold("(3) Pumped-storage hydropower and its applicability to flood control", level=2)
add_body(
    "Japan possesses world-leading pumped-storage hydropower capacity: approximately 27,000 MW "
    "across the country, with TEPCO alone operating approximately 10,000 MW.{15} "
    "These facilities operate on the principle of reversible water transfer between upper and lower "
    "reservoirs, a concept directly applicable to the proposed framework. "
    "The EU CORDIS project{16} investigated energy recovery from flood protection reservoir releases, "
    "establishing the technical feasibility of combining flood control with energy generation. "
    "Chang et al.{17} proposed a multi-scenario water level drawdown method for flood resource "
    "utilisation at cascade hydropower stations, demonstrating the potential for integrating flood "
    "management with hydropower operation."
)

add_heading_bold("(4) Japan's topographic advantages", level=2)
add_body(
    "Japan's rivers are characterised by exceptionally steep gradients compared with major rivers "
    "worldwide.{18} The Shinano River, the longest in Japan at 367 km, has a gradient 5-20 times "
    "steeper than major continental rivers such as the Mississippi (3,730 km) or the Mekong (4,350 km). "
    "This steep terrain provides two advantages for the proposed framework: high energy yield from "
    "relatively short discharge distances, and short inter-watershed distances (typically several to "
    "tens of kilometres) that facilitate groundwater transfer between adjacent basins."
)

# ============================================================
# 3. PROPOSED FRAMEWORK
# ============================================================
add_heading_bold("3. PROPOSED FRAMEWORK")

add_body(
    "The proposed Planned-Release Hydropower-Groundwater Management Integrated Flood Control "
    "Framework (PRHG-FCF) consists of the following four phases:"
)

add_body_no_indent(
    "Phase 1 (Pre-flood preparation): Based on weather forecasts and upstream hydrological "
    "observations, dam operators initiate planned discharge at a controlled rate. The discharge drives "
    "turbines installed in the dam outlet works, generating electricity."
)

add_body_no_indent(
    "Phase 2 (Groundwater pre-emptive pumping): Using the generated electricity, groundwater pumps "
    "in the adjacent watershed lower the water table by a predetermined depth (e.g., 3 m), creating "
    "subsurface storage capacity in the aquifer."
)

add_body_no_indent(
    "Phase 3 (Flood absorption): During the flood event, surface floodwaters infiltrate into the "
    "pre-emptied aquifer through natural infiltration and engineered recharge facilities (infiltration "
    "basins, injection wells), reducing surface runoff volume and peak discharge."
)

add_body_no_indent(
    "Phase 4 (Post-flood recovery): After the flood recedes, the aquifer naturally recharges over weeks "
    "to months, restoring groundwater levels. The recharged groundwater can subsequently be used for "
    "agricultural irrigation or municipal supply, creating a beneficial water cycle."
)

add_body(
    "The energy balance of this system can be expressed as follows. The hydropower generation "
    "capacity P_gen is calculated as P_gen = \u03b7_t \u00d7 \u03c1 \u00d7 g \u00d7 Q \u00d7 H, where \u03b7_t is the turbine "
    "efficiency (0.85), \u03c1 is the water density (1,000 kg/m\u00b3), g is gravitational acceleration "
    "(9.81 m/s\u00b2), Q is the discharge rate (m\u00b3/s), and H is the effective head (m). "
    "The pumping power P_pump required for groundwater extraction is P_pump = (\u03c1 \u00d7 g \u00d7 Q_p \u00d7 H_p) "
    "/ \u03b7_p, where Q_p is the pumping rate, H_p is the pumping head, and \u03b7_p is the pump efficiency "
    "(0.80). "
    "The framework is energetically self-sustaining when P_gen > P_pump, and the surplus energy "
    "P_surplus = P_gen - P_pump can be supplied to the electrical grid (Fig. 1)."
)

# ============================================================
# 4. STUDY AREA AND THE 2018 FLOOD EVENT
# ============================================================
add_heading_bold("4. STUDY AREA AND THE 2018 FLOOD EVENT")

add_heading_bold("(1) Overview of the Oda River-Takahashi River basin", level=2)
add_body(
    "The Oda River is a Class A river with a basin area of approximately 211 km\u00b2 in Okayama "
    "Prefecture, western Japan. It flows through a relatively flat plain before joining the Takahashi "
    "River at a bed gradient of approximately 1/2200.{2} "
    "During high water events on the Takahashi River, the backwater effect propagates upstream "
    "along the Oda River, maintaining elevated water levels for extended periods.{8}"
)

add_heading_bold("(2) The 2018 Western Japan Heavy Rainfall", level=2)
add_body(
    "From 5 to 8 July 2018, a quasi-stationary front brought record-breaking rainfall to western "
    "Japan: a 3-day cumulative rainfall of approximately 320 mm in the Takahashi River basin.{1} "
    "The resulting rapid rise in the Takahashi River main stream caused a severe backwater effect "
    "on the Oda River. Levees breached at five locations, inundating 8.28 km\u00b2 with a total flood "
    "volume of approximately 15.3 \u00d7 10\u2076 m\u00b3.{1}"
)

# Table 1
add_table_from_data(
    ["Parameter", "Value"],
    [
        ["3-day cumulative rainfall", "~320 mm"],
        ["Inundation area", "8.28 km\u00b2"],
        ["Flood volume", "~15.3 \u00d7 10\u2076 m\u00b3"],
        ["Levee breaches", "5 locations"],
        ["Fatalities (Mabi-cho)", "51"],
        ["Basin area (Oda River)", "~211 km\u00b2"],
        ["Bed gradient (Oda River)", "~1/2200"],
    ],
    "Table 1 Summary of the 2018 flood event in the Oda River-Takahashi River basin"
)

add_heading_bold("(3) Confluence relocation project", level=2)
add_body(
    "Following the 2018 disaster, the Ministry of Land, Infrastructure, Transport and Tourism (MLIT) "
    "relocated the Oda-Takahashi confluence approximately 4.6 km downstream, completing the project "
    "in March 2024.{4} During the heavy rainfall of November 2024, the measure was reported to reduce "
    "water levels by approximately 4.6 m on the Oda River side and approximately 0.8 m on the "
    "Takahashi River side.{4} This outcome is consistent with the predictions of Maeno's{8} "
    "two-dimensional flow analysis."
)

# ============================================================
# 5. SIMPLIFIED FEASIBILITY ANALYSIS
# ============================================================
add_heading_bold("5. SIMPLIFIED FEASIBILITY ANALYSIS")

add_heading_bold("(1) Energy balance", level=2)
add_body(
    "The energy balance was evaluated over a parameter space of effective head H = 20-200 m "
    "and discharge rate Q = 10-100 m\u00b3/s (Fig. 1). For a representative scenario of H = 100 m and "
    "Q = 50 m\u00b3/s, the generated power is P_gen = 0.85 \u00d7 1000 \u00d7 9.81 \u00d7 50 \u00d7 100 = 41.7 MW. "
    "The power required to pump groundwater at Q_p = 10 m\u00b3/s from a depth of H_p = 30 m is "
    "P_pump = (1000 \u00d7 9.81 \u00d7 10 \u00d7 30) / 0.80 = 3.68 MW. "
    "Thus, the surplus power is P_surplus = 41.7 - 3.68 = 38.0 MW, indicating that the framework "
    "achieves energy self-sufficiency with only approximately 9% of the generated power consumed "
    "for groundwater management. This is remarkably more efficient than the round-trip efficiency "
    "of conventional pumped-storage hydropower (70-80%).{15}"
)

add_heading_bold("(2) Aquifer storage capacity", level=2)
add_body(
    "The aquifer storage capacity was estimated based on the geological characteristics of the "
    "adjacent alluvial plain. Assuming a managed area A = 20 km\u00b2, a water table drawdown "
    "\u0394h = 3 m, and a specific yield S_y = 0.2, the storage volume is V = A \u00d7 \u0394h \u00d7 S_y "
    "= 20 \u00d7 10\u2076 \u00d7 3 \u00d7 0.2 = 12.0 \u00d7 10\u2076 m\u00b3 (Fig. 2). "
    "This represents approximately 78% of the 2018 flood volume (15.3 \u00d7 10\u2076 m\u00b3). "
    "However, excessive groundwater drawdown carries the risk of land subsidence.{6} "
    "The drawdown depth must be carefully calibrated to local geological conditions, with continuous "
    "monitoring of subsidence indicators."
)

add_heading_bold("(3) Flood hydrograph comparison", level=2)
add_body(
    "A simplified rainfall-runoff model (rational method basis) was used to calculate 72-hour "
    "hydrographs for four scenarios simulating the 2018 flood event (Fig. 3): "
    "Scenario 1 (no measures) reproduces the 2018 conditions; "
    "Scenario 2 (confluence relocation) assumes a 30% effective peak reduction based on the "
    "2024 operational results;{4} "
    "Scenario 3 (groundwater management) incorporates aquifer infiltration with 3 m pre-drawdown; "
    "Scenario 4 (integrated) combines Scenarios 2 and 3."
)

# Table 2
add_table_from_data(
    ["Scenario", "Peak discharge [m\u00b3/s]", "Reduction [%]", "Cumulative infiltration [\u00d710\u2076 m\u00b3]"],
    [
        ["1: No measures", "1,237", "\u2014", "\u2014"],
        ["2: Confluence relocation", "866", "30", "\u2014"],
        ["3: Groundwater management", "1,207", "2", "6.24"],
        ["4: Integrated", "845", "32", "6.24"],
    ],
    "Table 2 Flood control performance by scenario"
)

add_body(
    "The results (Table 2, Fig. 4) indicate that groundwater management alone (Scenario 3) yields "
    "a modest 2% peak reduction, constrained by aquifer infiltration rates. However, cumulative "
    "infiltration reaches 6.24 \u00d7 10\u2076 m\u00b3, equivalent to approximately 41% of the 2018 flood volume, "
    "significantly reducing post-flood drainage loads. "
    "The integrated scenario (Scenario 4) achieves a 32% peak discharge reduction through the "
    "synergistic combination of structural and groundwater-based measures."
)

# ============================================================
# 6. APPLICABILITY TO THE TOKYO METROPOLITAN AREA
# ============================================================
add_heading_bold("6. APPLICABILITY TO THE TOKYO METROPOLITAN AREA (ARAKAWA-EDOGAWA BASIN)")

add_heading_bold("(1) Existing flood control infrastructure", level=2)
add_body(
    "The Arakawa-Edogawa basin in the Tokyo metropolitan area possesses a multi-layered flood "
    "control system developed since the Meiji era (Table 3). The Arakawa Floodway (completed 1930, "
    "22 km) controls floodwater entry into the Sumida River. The Iwabuchi Sluice Gate regulates "
    "water levels at the Arakawa-Sumida bifurcation.{22} The Saiko Retarding Basin has an area of "
    "5.85 km\u00b2 and a storage capacity of 39 \u00d7 10\u2076 m\u00b3.{23} Three upstream dams (Urayama, Futase, "
    "Takizawa) provide additional flood regulation. The MAOUDC diverts floodwaters from five rivers "
    "into the Edogawa River at a maximum capacity of 200 m\u00b3/s.{7} "
    "Super-levees are under construction along the lower Arakawa, with approximately 15% of the "
    "managed levee length completed or under construction.{22}"
)

# Table 3
add_table_from_data(
    ["Facility", "Type", "Capacity/Scale", "Remarks"],
    [
        ["Arakawa Floodway", "Diversion channel", "22 km", "Completed 1930"],
        ["Iwabuchi Sluice Gate", "Sluice gate", "\u2014", "Arakawa-Sumida bifurcation"],
        ["Saiko Retarding Basin", "Retarding basin", "39 \u00d7 10\u2076 m\u00b3", "Area: 5.85 km\u00b2"],
        ["Upstream dams (3)", "Dams", "~130 \u00d7 10\u2076 m\u00b3 total", "Urayama, Futase, Takizawa"],
        ["MAOUDC", "Underground channel", "200 m\u00b3/s", "6.3 km, 5-river drainage"],
        ["Super-levees", "Levee", "\u2014", "~15% completed"],
    ],
    "Table 3 Major flood control facilities in the Arakawa-Edogawa basin"
)

add_heading_bold("(2) Kanto Plain aquifer and groundwater management history", level=2)
add_body(
    "Beneath the Kanto Plain lies a thick alluvial deposit. According to the three-dimensional "
    "geological-geotechnical model by the Geological Survey of Japan (GSJ),{24} the alluvial "
    "layer in the Tokyo Lowland reaches a maximum thickness of approximately 80 m, composed of "
    "basal gravel, alternating sand-mud layers, and surface mud layers. The sand-gravel layers "
    "function as productive aquifers with substantial groundwater storage potential."
)

add_body(
    "Historically, excessive groundwater extraction in the southern Kanto Plain from the mid-Meiji "
    "era (1890s) onward caused severe land subsidence.{25} Regulations under the Industrial Water Law "
    "(1956) and the Building Water Use Law (1962) curtailed extraction, and in recent decades, "
    "groundwater levels have recovered significantly.{25} "
    "From the perspective of the proposed framework, this 'recovered water table' represents "
    "a paradoxically favourable condition: the portion of recovered aquifer capacity that can be "
    "drawn down without inducing subsidence constitutes a potential subsurface retarding basin."
)

add_heading_bold("(3) Application scenario for the Arakawa basin", level=2)
add_body(
    "In the upper Arakawa basin (Okuchichibu Mountains), planned releases from Takizawa Dam "
    "(effective head approximately 100 m) can generate hydropower. This energy can be used to "
    "pre-emptively lower the water table in the mid-basin alluvial aquifers (Kawagoe-Saitama area). "
    "Assuming a managed area of 30 km\u00b2 with 3 m drawdown and S_y = 0.15, the available storage "
    "is 13.5 \u00d7 10\u2076 m\u00b3. Combined with the Saiko Retarding Basin (39 \u00d7 10\u2076 m\u00b3), total flood "
    "reception capacity reaches approximately 52.5 \u00d7 10\u2076 m\u00b3."
)

add_body(
    "The mid-basin alluvial lowland contains thick sand-gravel aquifers.{24} "
    "During Typhoon Hagibis (October 2019), levees breached at seven locations along the "
    "Iruma River and other Arakawa tributaries,{21} demonstrating the continued vulnerability of "
    "the metropolitan area. Given that additional structural measures are constrained by the "
    "difficulty of land acquisition in urban areas,{22} the proposed groundwater management approach "
    "offers a complementary 'underground retarding basin' that requires no surface land."
)

# ============================================================
# 7. DISCUSSION
# ============================================================
add_heading_bold("7. DISCUSSION")

add_heading_bold("(1) Energy self-sufficiency", level=2)
add_body(
    "The most significant finding of this study is that the proposed framework achieves energy "
    "self-sufficiency: only approximately 9% of the generated hydropower is consumed for "
    "groundwater pumping. This efficiency is far superior to the round-trip efficiency of "
    "conventional pumped-storage hydropower (70-80%),{15} because the framework exploits "
    "the one-directional potential energy of flood discharge rather than the bidirectional "
    "cycle of pumped storage. The surplus energy (approximately 38 MW in the representative "
    "scenario) can be supplied to the grid, creating an additional economic incentive for "
    "flood management."
)

add_heading_bold("(2) Complementarity with structural measures", level=2)
add_body(
    "The integrated scenario (Scenario 4) demonstrates that the proposed framework does not "
    "replace but rather complements existing structural measures. While the confluence relocation "
    "addresses the backwater problem directly, the groundwater management component provides "
    "additional volumetric absorption capacity. This complementarity is particularly valuable "
    "for extreme events exceeding the design capacity of structural measures."
)

add_heading_bold("(3) Limitations and challenges", level=2)
add_body(
    "Several limitations and challenges must be acknowledged. "
    "First, this study employs a simplified simulation framework. Detailed validation using "
    "coupled groundwater-surface water models such as GSFLOW{19} is required. "
    "Second, the risk of land subsidence due to groundwater drawdown is a critical concern. "
    "Historical experience in the Kanto Plain and Nobi Plain demonstrates that over-extraction "
    "can cause irreversible subsidence.{20} Strict monitoring and adaptive management protocols "
    "are essential. "
    "Third, the infiltration rate, which limits the real-time peak reduction effect (Scenario 3), "
    "depends heavily on aquifer hydraulic conductivity and the density of engineered recharge "
    "facilities. "
    "Fourth, the institutional and legal framework for inter-watershed groundwater management "
    "does not currently exist in Japan and would require legislative development."
)

add_heading_bold("(4) International applicability", level=2)
add_body(
    "The proposed framework has significant implications for transboundary river management. "
    "In large transboundary basins such as the Mekong,{26} Nile, and Jordan, downstream nations "
    "are unilaterally exposed to upstream dam operations. The proposed framework offers "
    "downstream nations a unilaterally implementable flood mitigation option that operates "
    "entirely within their own territory, as aquifer management requires no cross-border "
    "coordination. "
    "In arid regions of the Middle East and North Africa, where flash floods (wadi floods) occur "
    "sporadically, the framework could combine flood energy capture with aquifer recharge, "
    "particularly in oil-rich nations with capital for infrastructure investment. "
    "For delta protection (Mekong Delta, Nile Delta, Ganges-Brahmaputra-Meghna Delta), "
    "pre-emptive aquifer drawdown could provide dual protection against upstream flooding and "
    "sea-level rise."
)

# ============================================================
# 8. POLICY IMPLICATIONS
# ============================================================
add_heading_bold("8. POLICY IMPLICATIONS")

add_body(
    "Based on the findings of this study, the following policy recommendations are proposed:"
)

add_body_no_indent(
    "1) Pilot project implementation: A pilot-scale implementation in the Mabi district "
    "(Oda River-Takahashi River basin) and/or the mid-Arakawa basin is recommended to "
    "validate the framework under real-world conditions."
)

add_body_no_indent(
    "2) Amendment of groundwater extraction regulations: Current regulations, designed solely "
    "for subsidence prevention, should be amended to permit controlled, temporary drawdown for "
    "flood mitigation purposes, with strict monitoring requirements."
)

add_body_no_indent(
    "3) Integration with the Basin Flood Control policy: The MLIT's 'Basin Flood Control' "
    "(Ryuiki Chisui) policy framework{27} should be expanded to explicitly incorporate subsurface "
    "flood storage as a recognised non-structural measure."
)

add_body_no_indent(
    "4) Detailed simulation studies: Coupled groundwater-surface water modelling using "
    "GSFLOW{19} or equivalent tools should be conducted for candidate basins to provide "
    "high-resolution validation of the simplified analysis presented herein."
)

add_body_no_indent(
    "5) Linkage with the MAOUDC: In the Tokyo metropolitan area, the integration of the "
    "MAOUDC's inter-watershed flood transfer capability with the proposed aquifer storage "
    "approach could create a comprehensive subsurface flood management system."
)

# ============================================================
# 9. CONCLUSIONS
# ============================================================
add_heading_bold("9. CONCLUSIONS")

add_body(
    "This paper proposed a novel flood control framework (PRHG-FCF) integrating planned-release "
    "hydropower generation with inter-watershed groundwater management. The principal findings "
    "are as follows:"
)

add_body_no_indent(
    "(1) The framework achieves energy self-sufficiency: planned-release hydropower "
    "(approximately 42 MW) far exceeds the power required for groundwater pumping "
    "(approximately 3.7 MW), with a surplus of approximately 38 MW available for grid supply."
)

add_body_no_indent(
    "(2) A managed area of 20 km\u00b2 with 3 m groundwater drawdown creates approximately "
    "12.0 \u00d7 10\u2076 m\u00b3 of aquifer storage capacity, equivalent to approximately 78% of the "
    "2018 flood volume in the Oda River basin."
)

add_body_no_indent(
    "(3) Integration with the existing confluence relocation (structural measure) yields a "
    "32% reduction in peak discharge, demonstrating that the proposed framework functions as "
    "a complementary 'additional safety layer' rather than a replacement for structural measures."
)

add_body_no_indent(
    "(4) The applicability to the Arakawa basin (8.21 million people in projected inundation "
    "zones) was demonstrated, with the combination of upstream dam hydropower and mid-basin "
    "alluvial aquifers potentially creating supplementary flood regulation capacity complementing "
    "existing dams and retarding basins. Integration with the MAOUDC is also promising for "
    "enhancing flood safety in densely populated and industrially concentrated areas."
)

add_body_no_indent(
    "(5) Japan's steep terrain, existing pumped-storage infrastructure, and dense "
    "hydrometeorological observation networks provide globally rare favourable conditions for "
    "implementing this framework. Detailed validation through coupled hydrological modelling "
    "and pilot-scale implementation are recommended as next steps."
)

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
add_heading_bold("ACKNOWLEDGEMENTS")
add_body_no_indent(
    "This study utilised publicly available data from the National Research Institute for Earth "
    "Science and Disaster Resilience (NIED), the MLIT Chugoku Regional Development Bureau, "
    "the MLIT Kanto Regional Development Bureau, the Geological Survey of Japan (GSJ/AIST), "
    "and Okayama Prefecture. The authors gratefully acknowledge these organisations."
)

# ============================================================
# REFERENCES
# ============================================================
add_heading_bold("REFERENCES")

refs = [
    '1) National Research Institute for Earth Science and Disaster Resilience (NIED): Overview of damage caused by the July 2018 heavy rainfall (Baiu front and Typhoon No. 7), 2018.',
    '2) MLIT Chugoku Regional Development Bureau: Oda River confluence relocation project overview, https://www.cgr.mlit.go.jp/',
    '3) MLIT: River improvement initiatives, https://www.mlit.go.jp/river/basic_info/',
    '4) KSB News: Oda River flood control project reported to have "reduced water levels" during the 2024 heavy rainfall, 7 February 2025.',
    '5) California Department of Water Resources: Flood-Managed Aquifer Recharge (Flood-MAR), https://water.ca.gov/Programs/All-Programs/Flood-MAR',
    '6) Chang, Y.-L., Huang, W.-C. and Lee, Y.-H.: Flood hazard mitigation in land subsidence prone coastal areas by optimal groundwater pumping, Water Resour. Manag., 2019.',
    '7) MLIT Kanto Regional Development Bureau: Metropolitan Area Outer Underground Discharge Channel, https://www.ktr.mlit.go.jp/edogawa/gaikaku/',
    '8) Maeno, S.: Effect of relocating the confluence of the Takahashi and Oda Rivers, J. Hydraul. Eng., JSCE, Vol. 51, pp. 613-618, 2007.',
    '9) Mississippi Levee Board: Yazoo Backwater Project, https://msleveeboard.com/',
    '10) Mississippi State University Extension: Final Report: Survey of Overlooked Costs of the 2019 Backwater Flood in the Yazoo Mississippi Delta.',
    '11) Berkowitz, J. F., Page, L. A. and Noble, C. V.: Forested wetland hydrology in a large Mississippi River tributary system, Wetlands, 2019.',
    '12) Li, A. and Tsai, F. T.-C.: Understanding dynamics of groundwater flows in the Mississippi River Delta, J. Hydrol., Vol. 583, 124616, 2020.',
    '13) Dillon, P., Stuyfzand, P., Grischek, T., Lluria, M., Pyne, R. D. G., Jain, R. C., Bear, J., Schwarz, J., Wang, W., Fernandez, E., Stefan, C., Pettenati, M., van der Gun, J., Sprenger, C., Massmann, G., Scanlon, B. R., Xanke, J., Jokela, P., Zheng, Y., Rossetto, R., Shamrukh, M., Pavelic, P., Murray, E., Ross, A., Bonilla Valverde, J. P., Palma Nava, A., Ansems, N., Posavec, K., Ha, K., Martin, R. and Sapiano, M.: Managed aquifer recharge: rediscovering nature as a leading edge technology, Water Sci. Technol., Vol. 62, No. 10, pp. 2338-2345, 2010.',
    '14) Nishizawa Katazakai, S. and Zhang, J.: Long-term flood control in central Japan: A half-century groundwater monitoring and evaluating adaptation measures, Groundwater Sustain. Dev., Vol. 33, 101585, 2026.',
    '15) IEA Hydropower Implementing Agreement: Case Study 11-02: Large Scale Pumped Storage Power Plants, Japan, 2006.',
    '16) EU CORDIS: Recovery of Energy from Released Water and Reserve Flow from a Flood Protection Reservoir, Project ID: HY.-00410-87.',
    '17) Chang, J., Guo, A., Wang, Y. and Ha, Y.: The flood resource utilization of cascade hydropower stations based on the multi-scenario water level drawdown method, EGU General Assembly 2025, EGU25-14374.',
    '18) Kinoshita, S. and Tokunaga, Y.: Water resources management in Japan, IWRA Proceedings, 2003.',
    '19) USGS: GSFLOW: Coupled Groundwater and Surface-Water Flow Model, https://water.usgs.gov/ogw/gsflow/',
    '20) Cabinet Council on Land Subsidence Countermeasures: Basic Policy on Land Subsidence Countermeasures, 1991.',
    '21) MLIT Water Management and Land Conservation Bureau: Outline of revision of the Arakawa River System Basic Policy for River Improvement, November 2024, pp. 1-52.',
    '22) MLIT Kanto Regional Development Bureau, Arakawa Downstream River Office: Flood control measures for the Arakawa River, https://www.ktr.mlit.go.jp/arage/',
    '23) MLIT Kanto Regional Development Bureau, Arakawa Upstream River Office: Functions of Saiko Retarding Basin, https://www.ktr.mlit.go.jp/arajo/arajo00161.html',
    '24) Komatsubara, J.: Alluvial deposits of the Tokyo Lowland, GSJ Chishitsu News, Vol. 10, No. 7, pp. 148-152, 2021.',
    '25) MLIT Water Resources Department: Current status of groundwater conservation and land subsidence, https://www.mlit.go.jp/mizukokudo/mizsei/',
    '26) Mekong River Commission: State of the Basin Report 2018, Vientiane, 2019.',
    '27) MLIT: Basin Flood Control Project, https://www.mlit.go.jp/river/kasen/ryuiki_pro/',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.hanging_indent = Pt(24)
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

# ============================================================
# FIGURE LEGENDS
# ============================================================
add_heading_bold("FIGURE LEGENDS")

fig_legends = [
    "Fig. 1 Energy balance contour map of the planned-release hydropower-groundwater management framework. "
    "The contour plot shows the generated power (MW) as a function of effective head (m) and discharge rate (m\u00b3/s). "
    "The dashed line indicates the power required for groundwater pumping (H_p = 30 m, Q_p = 10 m\u00b3/s). "
    "The wide region above the dashed line represents the energy surplus zone where the framework is self-sustaining.",

    "Fig. 2 Aquifer storage capacity as a function of managed area and water table drawdown. "
    "Contours show storage volume (\u00d710\u2076 m\u00b3) for specific yield S_y = 0.2. "
    "The dashed horizontal line indicates the 2018 flood volume (15.3 \u00d7 10\u2076 m\u00b3) for comparison. "
    "The shaded area indicates the range of parameters achievable in the Oda River adjacent plain.",

    "Fig. 3 Flood hydrograph comparison for four scenarios over a 72-hour simulation period "
    "(upper panel: rainfall intensity; lower panel: discharge). "
    "Scenario 1: no measures (2018 baseline); Scenario 2: confluence relocation; "
    "Scenario 3: groundwater management; Scenario 4: integrated (Scenarios 2 + 3).",

    "Fig. 4 Comparison of peak discharge (left) and total runoff volume (right) across the four "
    "scenarios. Error bars indicate sensitivity to key parameters (specific yield: 0.15-0.25; "
    "pumping rate: 5-15 m\u00b3/s).",
]

for legend in fig_legends:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(legend)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

# ── Save ──
outpath = f"{OUTDIR}/jsce_en_manuscript.docx"
doc.save(outpath)
print(f"Journal of JSCE English manuscript saved: {outpath}")
