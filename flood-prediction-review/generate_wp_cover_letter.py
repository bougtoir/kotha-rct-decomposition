#!/usr/bin/env python3
"""
Generate Water Policy (IWA Publishing) cover letter in docx format.
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"

doc = Document()

# ─── Page setup (A4) ───
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# ─── Base style ───
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
style.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def add_para_red(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.italic = True
    return p


# ── Date ──
today = date.today().strftime("%d %B %Y")
add_para(today, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ── Addressee ──
add_para("The Editor-in-Chief")
add_para("Water Policy")
add_para("IWA Publishing", space_after=Pt(12))

# ── Subject ──
add_para("Dear Editor,", space_after=Pt(12))

# ── Body ──
add_para(
    "I am writing to submit the enclosed manuscript entitled "
    "\"Unilateral flood control through subsurface storage management: "
    "a hydro-political analysis of planned-release hydropower "
    "and inter-watershed groundwater transfer\" "
    "for consideration as a Research Paper in Water Policy."
)

add_para(
    "This paper proposes a novel flood control framework that integrates "
    "planned-release hydropower generation with inter-watershed groundwater "
    "management, enabling downstream nations or sub-national jurisdictions to "
    "implement flood mitigation unilaterally, without requiring upstream "
    "cooperation. The paper situates this technical proposal within the "
    "hydro-hegemony literature, demonstrating its relevance to transboundary "
    "flood governance in major river basins including the Mekong, Nile, "
    "Ganges-Brahmaputra-Meghna, Indus, and Tigris-Euphrates."
)

add_para(
    "The manuscript addresses a critical gap at the intersection of flood "
    "control technology and water governance: while the hydro-hegemony "
    "framework has illuminated the power dynamics of transboundary water "
    "allocation, it has given comparatively less attention to flood risk as "
    "a dimension of water power asymmetry. By showing that subsurface storage "
    "management can provide downstream states with autonomous flood control "
    "capacity, this paper opens a new dimension in debates on water security, "
    "climate adaptation, and the politics of flood infrastructure."
)

add_para(
    "A simplified feasibility analysis of the Oda River\u2013Takahashi River "
    "basin in western Japan, where catastrophic backwater flooding killed 51 "
    "people in 2018, demonstrates that such a system can be energy "
    "self-sufficient (consuming only 9% of generated power for groundwater "
    "pumping) and can create subsurface storage equivalent to 78% of observed "
    "flood volumes. The framework's applicability to the Tokyo metropolitan "
    "area (8.21 million people in projected inundation zones) and to deltaic "
    "and arid-region contexts is also discussed."
)

add_para(
    "I believe this manuscript is well suited for Water Policy because it "
    "bridges the gap between technical flood management and water governance "
    "analysis, directly engaging with the hydro-hegemony framework and "
    "international water law while grounding its arguments in quantitative "
    "feasibility analysis."
)

add_para(
    "This manuscript has not been published previously and is not under "
    "consideration by any other journal. All data and simulation code are "
    "publicly available. The author declares no conflicts of interest and "
    "no external funding was received for this work."
)

add_para(
    "Thank you for considering this manuscript. I look forward to your "
    "response.",
    space_after=Pt(12),
)

# ── Closing ──
add_para("Sincerely,", space_after=Pt(24))

add_para("Tatsuki Onishi")
add_para_red("(Position, Affiliation)")
add_para_red("(Address)")
add_para_red("(E-mail)")

# ── Save ──
outpath = os.path.join(OUTDIR, "wp_cover_letter.docx")
doc.save(outpath)
print(f"Water Policy cover letter saved: {outpath}")
