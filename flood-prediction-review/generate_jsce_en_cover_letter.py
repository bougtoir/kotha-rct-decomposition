#!/usr/bin/env python3
"""
Cover Letter for Journal of JSCE (English)
Special Issue: Hydroscience and Hydraulic Engineering (B1)
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"

doc = Document()

# Page setup (A4)
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# Style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.line_spacing = Pt(22)
style.paragraph_format.space_after = Pt(0)


def add_para(text, bold=False, alignment=None, space_before=0, space_after=0,
             font_size=11, color=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_placeholder(text):
    return add_para(text, color=RGBColor(255, 0, 0))


# ── Date ──
add_para("May __, 2026", alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

# ── Addressee ──
add_para("Editorial Committee", space_after=0)
add_para("Journal of JSCE, Special Issue: Hydroscience and Hydraulic Engineering", space_after=18)

# ── Sender ──
add_placeholder("Author name: (Author name)")
add_placeholder("Affiliation: (Affiliation)")
add_placeholder("E-mail: (E-mail address)")
add_placeholder("Tel: (Phone number)")
add_para("", space_after=12)

# ── Title ──
add_para("Re: Manuscript Submission", bold=True, space_after=12, font_size=12)

# ── Body ──
add_para("Dear Editor,", space_after=12)

add_para(
    "We are pleased to submit the following manuscript for consideration "
    "for publication in the Journal of JSCE, Special Issue on Hydroscience "
    "and Hydraulic Engineering (B1 Section).",
    space_after=18
)

# ── Manuscript information ──
add_para("1. Manuscript title", bold=True, space_after=0)
add_para(
    "   A novel flood control framework integrating planned-release hydropower "
    "with inter-watershed groundwater management: A simplified feasibility analysis "
    "for the Oda River-Takahashi River basin",
    space_after=12
)

add_para("2. Article type", bold=True, space_after=0)
add_para(
    "   Academic paper (integrated theoretical and empirical research with original contributions)",
    space_after=12
)

add_para("3. Preferred review section", bold=True, space_after=0)
add_para(
    "   B1 Section (Hydroscience and Hydraulic Engineering): "
    "Flood damage, flood control, river management and planning, water resources",
    space_after=12
)

add_para("4. Summary", bold=True, space_after=0)
add_para(
    "   This paper proposes a Planned-Release Hydropower-Groundwater Management "
    "Integrated Flood Control Framework (PRHG-FCF) that uses hydroelectric energy "
    "from controlled flood discharge to pre-emptively pump groundwater from adjacent "
    "watersheds, thereby creating subsurface storage capacity to absorb floodwaters. "
    "A simplified feasibility analysis for the Oda River-Takahashi River basin "
    "(Mabi-cho, Kurashiki City, Okayama Prefecture), the site of catastrophic backwater "
    "flooding during the 2018 Western Japan Heavy Rainfall, demonstrates: "
    "(1) energy self-sufficiency (only ~9% of generated power consumed for groundwater "
    "management); (2) aquifer storage capacity equivalent to ~78% of the 2018 flood volume; "
    "and (3) a 32% peak discharge reduction when integrated with the existing confluence "
    "relocation measure. The applicability to the Tokyo metropolitan area "
    "(Arakawa-Edogawa basin, 8.21 million people in projected inundation zones) "
    "is also examined.",
    space_after=12
)

add_para("5. Novelty and significance", bold=True, space_after=0)
add_para(
    "   The novelty of this study lies in three aspects:\n"
    "   (1) The first systematic presentation of a flood control concept integrating "
    "planned-release hydropower with aquifer management.\n"
    "   (2) Quantitative feasibility assessment from three perspectives: energy balance, "
    "aquifer storage capacity, and flood hydrograph modification.\n"
    "   (3) Quantitative demonstration of the complementary effect when combined with "
    "structural measures (confluence relocation).\n"
    "   The proposed framework builds upon existing technologies (pumped-storage hydropower, "
    "managed aquifer recharge) and offers a practical new component for Japan's "
    "Basin Flood Control (Ryuiki Chisui) policy.",
    space_after=12
)

add_para("6. Disclosure of related publications", bold=True, space_after=0)
add_para(
    "   A related English-language manuscript (Commentary format) is being prepared "
    "for submission to Nature Water. The Nature Water manuscript focuses on international "
    "water governance and transboundary river management policy implications, "
    "and does not contain the detailed quantitative analysis of the Oda River-Takahashi River "
    "and Arakawa River basins presented in this paper. "
    "The two manuscripts target different readerships and have distinct content structures "
    "with no substantive overlap. "
    "This submission satisfies the requirements for originality and non-duplicate submission.",
    space_after=12
)

add_para("7. Conflicts of interest", bold=True, space_after=0)
add_para(
    "   The authors declare that there are no conflicts of interest.",
    space_after=12
)

add_para("8. Ethics compliance", bold=True, space_after=0)
add_para(
    "   The authors have complied with the JSCE Code of Ethics and the ethical standards "
    "for manuscripts submitted to the Journal of JSCE. All data used in this study "
    "are publicly available, and no ethical review was required.",
    space_after=12
)

add_para("9. Submitted documents", bold=True, space_after=0)
add_para(
    "   (1) Title Page (Word format)\n"
    "   (2) Main Manuscript (Word format, double-spaced)\n"
    "   (3) This Cover Letter\n"
    "   (4) Figure files (PNG format, 4 individual files)\n"
    "   (5) Submission checklist",
    space_after=18
)

add_para(
    "We respectfully request your consideration and review of the enclosed manuscript.",
    space_after=12
)

add_para("Sincerely,", space_after=12)
add_placeholder("(Author name)")
add_placeholder("(Affiliation)")

# ── Save ──
outpath = f"{OUTDIR}/jsce_en_cover_letter.docx"
doc.save(outpath)
print(f"Cover letter saved: {outpath}")
