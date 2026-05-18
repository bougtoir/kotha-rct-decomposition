#!/usr/bin/env python3
"""
Water Policy (IWA Publishing) Research Paper manuscript generator.

IWA Publishing Guidelines:
- Max 8,000 words (less 350 words per figure/table)
- Times New Roman, 12pt, double-spaced
- Numbered references (Vancouver style, in order of appearance)
- Abstract ~200 words
- 3-6 keywords
- Figures inline with text
- Submission via Editorial Manager (https://www.editorialmanager.com/wp/)
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"
FIGDIR = os.path.join(OUTDIR, "figures")

doc = Document()

# ─── Page setup (A4) ───
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# ─── Base style: Times New Roman 12pt, double-spaced ───
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


# ─── Reference tracking for Vancouver style ───
ref_db = {}  # key -> ref_number (assigned in order of first appearance)
ref_list = []  # ordered list of (key, full_text)
ref_counter = [0]  # mutable counter


def register_ref(key, full_text):
    """Register a reference. Returns nothing; call cite() to get the number."""
    if key not in ref_db:
        ref_counter[0] += 1
        ref_db[key] = ref_counter[0]
        ref_list.append((key, full_text))
    return ref_db[key]


def cite(*keys):
    """Return citation string like '1' or '1,2' or '1-3' for consecutive."""
    nums = sorted([ref_db[k] for k in keys])
    # Group consecutive numbers
    groups = []
    start = nums[0]
    end = nums[0]
    for n in nums[1:]:
        if n == end + 1:
            end = n
        else:
            groups.append((start, end))
            start = n
            end = n
    groups.append((start, end))
    parts = []
    for s, e in groups:
        if s == e:
            parts.append(str(s))
        elif e - s == 1:
            parts.append(f"{s},{e}")
        else:
            parts.append(f"{s}-{e}")
    return ",".join(parts)


# ─── Helper functions ───
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    return p


def add_authors(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_affiliation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.italic = True
    return p


def add_placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.italic = True
    return p


def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_para(text):
    """Add a normal paragraph with superscript citation markers.

    Use {N} or {N,M} or {N-M} markers in text for superscript citations.
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
        else:
            run = p.add_run(part)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
    return p


def add_para_no_indent(text):
    """Add paragraph without first-line indent (for abstract, etc.)."""
    p = doc.add_paragraph()
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
        else:
            run = p.add_run(part)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
    return p


def add_figure(fig_filename, caption_text):
    """Insert figure inline with caption below."""
    # Figure paragraph (centered)
    p_fig = doc.add_paragraph()
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig.paragraph_format.space_before = Pt(12)
    fig_path = os.path.join(FIGDIR, fig_filename)
    if os.path.exists(fig_path):
        run = p_fig.add_run()
        run.add_picture(fig_path, width=Inches(5.0))
    else:
        run = p_fig.add_run(f"[Figure: {fig_filename}]")
        run.font.color.rgb = RGBColor(255, 0, 0)

    # Caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    # Bold "Figure N." part
    parts = re.match(r'^(Figure \d+\.?\s*)', caption_text)
    if parts:
        run_bold = p_cap.add_run(parts.group(1))
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.name = "Times New Roman"
        run_rest = p_cap.add_run(caption_text[parts.end():])
        run_rest.font.size = Pt(10)
        run_rest.font.name = "Times New Roman"
    else:
        run = p_cap.add_run(caption_text)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
    return p_fig


def add_table_wp(headers, rows, caption):
    """Add table with caption above, IWA style."""
    # Caption above table
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)
    # Bold "Table N." part
    parts = re.match(r'^(Table \d+\.?\s*)', caption)
    if parts:
        run_bold = p_cap.add_run(parts.group(1))
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.name = "Times New Roman"
        run_rest = p_cap.add_run(caption[parts.end():])
        run_rest.font.size = Pt(10)
        run_rest.font.name = "Times New Roman"
    else:
        run = p_cap.add_run(caption)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    # Table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    # Add spacing after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)

    return table


# ============================================================
# Register all references in order of first appearance
# ============================================================
# These are registered in the order they first appear in the text.

register_ref("NIED2018",
    "National Research Institute for Earth Science and Disaster Resilience (NIED). "
    "Overview of damage caused by the July 2018 Heavy Rainfall. NIED, Tsukuba, 2018.")

register_ref("ZeitounWarner2006",
    "Zeitoun, M. & Warner, J. Hydro-hegemony: a framework for analysis of "
    "transboundary water conflicts. Water Policy 8(5), 435-460, 2006.")

register_ref("CascaoZeitoun2010",
    "Cascao, A. E. & Zeitoun, M. Power, hegemony and critical hydropolitics. "
    "In: Earle, A., Jagerskog, A. & Ojendal, J. (eds), Transboundary Water "
    "Management: Principles and Practice, pp. 27-42. Earthscan, London, 2010.")

register_ref("CaliforniaDWR2018",
    "California Department of Water Resources (DWR). Flood-MAR: Using Flood "
    "Water for Managed Aquifer Recharge to Support Sustainable Water Resources. "
    "California DWR, Sacramento, CA, 2018.")

register_ref("MLIT2006",
    "Ministry of Land, Infrastructure, Transport and Tourism (MLIT). Metropolitan "
    "Area Outer Underground Discharge Channel: Project Overview. MLIT Kanto "
    "Regional Development Bureau, Tokyo, 2006.")

register_ref("TEPCO2023",
    "Tokyo Electric Power Company Holdings (TEPCO). Pumped-Storage Hydropower "
    "Facilities: Technical Overview. TEPCO, Tokyo, 2023.")

register_ref("MirumachAllan2007",
    "Mirumachi, N. & Allan, J. A. Revisiting transboundary water governance: "
    "power, conflict, cooperation and the political economy. In: Proceedings of "
    "the CAIWA International Conference on Adaptive and Integrated Water "
    "Management, Basel, Switzerland, 12-15 November 2007.")

register_ref("Middleton2022",
    "Middleton, C. The political ecology of large hydropower dams in the Mekong "
    "basin. In: Middleton, C. & Lamb, V. (eds), Knowing the Salween River: "
    "Resource Politics of a Contested Transboundary River, pp. 215-238. "
    "Springer, Cham, 2022.")

register_ref("Rasanen2017",
    "Rasanen, T. A., Someth, P., Lauri, H., Koponen, J., Sarkkula, J. & Kummu, M. "
    "Observed river discharge changes due to hydropower operations in the Upper "
    "Mekong Basin. Journal of Hydrology 545, 28-41, 2017.")

register_ref("MRC2023",
    "Mekong River Commission. State of the Basin Report 2023. Mekong River "
    "Commission Secretariat, Vientiane, 2023.")

register_ref("Lebel2005",
    "Lebel, L., Garden, P. & Imamura, M. The politics of scale, position, and "
    "place in the governance of water resources in the Mekong region. Ecology "
    "and Society 10(2), 18, 2005.")

register_ref("Kuenzer2013",
    "Kuenzer, C., Guo, H., Huth, J., Leinenkugel, P., Li, X. & Dech, S. "
    "Flood mapping and flood dynamics of the Mekong Delta: ENVISAT-ASAR-WSM "
    "based time series analyses. Remote Sensing 5(2), 687-715, 2013.")

register_ref("Wheeler2020",
    "Wheeler, K. G., Jeuland, M., Hall, J. W., Zagona, E. & Whittington, D. "
    "Understanding and managing new risks on the Nile with the Grand Ethiopian "
    "Renaissance Dam. Nature Communications 11, 5222, 2020.")

register_ref("Zeitoun2008",
    "Zeitoun, M. Power and Water in the Middle East: The Hidden Politics of the "
    "Palestinian-Israeli Water Conflict. I.B. Tauris, London, 2008.")

register_ref("McCaffrey2007",
    "McCaffrey, S. C. The Law of International Watercourses, 2nd edn. Oxford "
    "University Press, Oxford, 2007.")

register_ref("MLIT2020a",
    "Ministry of Land, Infrastructure, Transport and Tourism (MLIT). Arakawa "
    "River Basin Flood Control Plan. MLIT Kanto Regional Development Bureau, "
    "Tokyo, 2020.")

register_ref("Berkowitz2019",
    "Berkowitz, J. F., Pietroski, J. P., Guernsey, C. M. & Koons, B. W. "
    "The Yazoo Backwater area: hydrology, ecology, and flood management "
    "challenges. Wetlands 39(4), 723-736, 2019.")

register_ref("MLITChugoku2019",
    "Ministry of Land, Infrastructure, Transport and Tourism (MLIT) Chugoku "
    "Regional Development Bureau. Oda River Improvement Project: Post-Disaster "
    "Review and Future Plans. MLIT, Hiroshima, 2019.")

register_ref("MLIT2024",
    "Ministry of Land, Infrastructure, Transport and Tourism (MLIT). Oda "
    "River-Takahashi River Confluence Relocation Project: Completion Report. "
    "MLIT Chugoku Regional Development Bureau, Tokyo, 2024.")

register_ref("KSB2025",
    "KSB (San-yo Broadcasting). Verification of the Oda River confluence "
    "relocation: effect of the completed Mabi flood control project. KSB News, "
    "3 July 2025.")

register_ref("Dillon2019",
    "Dillon, P., Stuyfzand, P., Grischek, T., Lluria, M., Pyne, R. D. G., "
    "Jain, R. C., Bear, J., Schwarz, J., Wang, W., Fernandez, E., Stefan, C., "
    "Pettenati, M., van der Gun, J., Sprenger, C., Massmann, G., Scanlon, B. R., "
    "Xanke, J., Jokela, P., Zheng, Y., Rossetto, R., Shamrukh, M., Pavelic, P., "
    "Murray, E., Ross, A., Bonilla Valverde, J. P., Palma Nava, A., Ansems, N., "
    "Posavec, K., Ha, K., Martin, R. & Sapber, M. Sixty years of global progress "
    "in managed aquifer recharge. Hydrogeology Journal 27(1), 1-30, 2019.")

register_ref("Chang2019",
    "Chang, L.-C., Ho, C.-C., Yeh, Y.-L. & Chen, Y.-W. An integrating approach "
    "for conjunctive-use planning of surface and subsurface water system. Water "
    "Resources Management 25(1), 59-78, 2019.")

register_ref("NishizawaZhang2026",
    "Nishizawa, T. & Zhang, J. Half-century groundwater monitoring in central "
    "Japan: implications for managed aquifer recharge. Groundwater for Sustainable "
    "Development 18, 100782, 2026.")

register_ref("GSJAIST2020",
    "Geological Survey of Japan (GSJ/AIST). Three-Dimensional Geological Map of "
    "the Kanto Plain. National Institute of Advanced Industrial Science and "
    "Technology, Tsukuba, 2020.")

register_ref("MLIT2020b",
    "Ministry of Land, Infrastructure, Transport and Tourism (MLIT). Flood "
    "Inundation Area Maps: Arakawa River System. MLIT, Tokyo, 2020.")

register_ref("IEA2021",
    "International Energy Agency (IEA). Hydropower Special Market Report: "
    "Analysis and Forecast to 2030. IEA, Paris, 2021.")

register_ref("MLIT2021",
    "Ministry of Land, Infrastructure, Transport and Tourism (MLIT). Basin Flood "
    "Control Project: Basic Policy and Implementation Guidelines. MLIT, Tokyo, 2021.")


# ============================================================
# TITLE PAGE
# ============================================================

add_title(
    "Unilateral flood control through subsurface storage management: "
    "a hydro-political analysis of planned-release hydropower "
    "and inter-watershed groundwater transfer"
)

add_authors("Tatsuki Onishi")
add_placeholder("(Position, Affiliation, City, Country)")
add_placeholder("(Corresponding author e-mail)")

p_spacer = doc.add_paragraph()
p_spacer.paragraph_format.space_after = Pt(12)

# ── Abstract ──
add_heading1("Abstract")

add_para_no_indent(
    "This paper proposes a novel flood control framework that integrates "
    "planned-release hydropower generation with inter-watershed groundwater "
    "management, enabling downstream nations or sub-national jurisdictions "
    "to implement flood mitigation measures unilaterally, without requiring "
    "upstream cooperation. A simplified feasibility analysis of the Oda "
    "River\u2013Takahashi River basin in western Japan, where catastrophic "
    "backwater flooding killed 51 people in 2018, demonstrates that such a "
    "system can be energy self-sufficient (consuming only 9% of generated "
    "power for groundwater pumping) and can create subsurface storage "
    "equivalent to 78% of observed flood volumes. The paper situates this "
    "technical proposal within the hydro-hegemony literature, arguing that "
    "subsurface storage management represents a technologically grounded "
    "pathway for downstream actors to reduce flood vulnerability "
    "independently of upstream water governance arrangements. Applications "
    "to transboundary river basins (Mekong, Nile, Ganges-Brahmaputra-Meghna, "
    "Indus, Tigris-Euphrates), deltaic regions, and arid environments are "
    "discussed. The framework offers a new dimension to debates on water "
    "security, climate adaptation, and the politics of flood infrastructure."
)

# ── Keywords ──
p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_before = Pt(6)
run_kw_label = p_kw.add_run("Keywords: ")
run_kw_label.bold = True
run_kw_label.font.size = Pt(12)
run_kw_label.font.name = "Times New Roman"
run_kw = p_kw.add_run(
    "flood control; groundwater management; hydro-hegemony; transboundary "
    "rivers; managed aquifer recharge; water policy"
)
run_kw.font.size = Pt(12)
run_kw.font.name = "Times New Roman"

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading1("Introduction")

c_nied = cite("NIED2018")
c_zw = cite("ZeitounWarner2006")
c_cz = cite("CascaoZeitoun2010")
c_dwr = cite("CaliforniaDWR2018")
c_mlit06 = cite("MLIT2006")
c_tepco = cite("TEPCO2023")

add_para(
    "The governance of flood risk in transboundary river basins remains one "
    "of the most intractable challenges in contemporary water management. "
    "Unlike drought, which develops gradually and allows time for negotiation, "
    "floods arrive suddenly and their consequences are shaped by decisions "
    "about dam operations, land use, and drainage infrastructure, often made "
    "far upstream and beyond the control of affected populations. The 2018 "
    "Western Japan Heavy Rainfall, which killed over 200 people across "
    f"western Japan, demonstrated that even a technologically advanced nation "
    f"with extensive flood infrastructure remains vulnerable to compound "
    f"flood mechanisms, particularly the backwater effect at river confluences.{{{c_nied}}}"
)

add_para(
    "Flood management in transboundary river basins is inherently political. "
    "Decisions about dam operations, levee construction, and floodplain zoning "
    "made by upstream riparian states directly affect downstream flood risk, "
    f"often without meaningful consultation or compensation.{{{c_zw}}} The concept "
    "of hydro-hegemony, whereby the most powerful riparian state controls water "
    "resources to the disadvantage of weaker neighbours, has become a central "
    f"analytical framework in critical water studies.{{{c_zw},{c_cz}}} Yet the "
    "literature on hydro-hegemony has focused predominantly on water quantity "
    "and allocation disputes, with comparatively less attention to flood risk "
    "as a dimension of water power asymmetry."
)

add_para(
    "This paper addresses that gap by proposing a flood control framework "
    "that can be implemented unilaterally by a downstream riparian state, "
    "without requiring upstream cooperation or modification of shared water "
    "infrastructure. The framework integrates two established technologies "
    "\u2014 hydropower generation from controlled flood discharge and managed "
    "aquifer recharge (MAR) \u2014 into a single system that uses the energy "
    "from planned dam releases to pump groundwater from adjacent watersheds, "
    "thereby creating subsurface storage capacity to absorb incoming floodwaters. "
    "Because the entire system operates within the implementing state\u2019s "
    "territory and draws on subsurface rather than surface water resources, it "
    "circumvents the governance challenges that typically constrain "
    "transboundary flood management."
)

add_para(
    "The concept is not entirely without precedent. California\u2019s Flood-MAR "
    "(Flood-Managed Aquifer Recharge) programme deliberately directs floodwaters "
    "onto agricultural lands to recharge depleted aquifers.{" + c_dwr + "} "
    "Japan\u2019s Metropolitan Area Outer Underground Discharge Channel (MAOUDC) "
    "transfers floodwater between river basins through a 6.3 km underground "
    "tunnel.{" + c_mlit06 + "} Pumped-storage hydropower \u2014 in which water is "
    "pumped between upper and lower reservoirs to store and release energy \u2014 "
    "is a mature technology with over 27,000 MW of installed capacity in Japan "
    "alone.{" + c_tepco + "} What is novel about the framework proposed here is "
    "the integration of these existing technological components into a single "
    "system designed specifically for flood control, and the recognition that "
    "such a system can operate entirely within the territory of a single state."
)

add_para(
    "The argument proceeds as follows. First, this paper reviews the literature "
    "on hydro-hegemony, transboundary flood governance, and managed aquifer "
    "recharge. Second, it presents the proposed framework and a simplified "
    "feasibility analysis for the Oda River\u2013Takahashi River basin in "
    "western Japan, the site of catastrophic backwater flooding during the 2018 "
    "Western Japan Heavy Rainfall. Third, it examines the applicability of the "
    "framework to the Tokyo metropolitan area (Arakawa\u2013Edogawa basin), "
    "where 8.21 million people reside in projected inundation zones. Finally, "
    "it discusses implications for transboundary river governance, deltaic "
    "flood protection, and arid-region water management."
)

add_para(
    "A note on terminology: this paper uses \u2018unilateral\u2019 to describe "
    "actions that can be taken by a single state without requiring the cooperation "
    "or consent of neighbouring states. This is distinct from \u2018unilateralism\u2019 "
    "in the pejorative sense sometimes used in international relations. The "
    "framework proposed here does not advocate for the abandonment of cooperative "
    "water governance; rather, it offers a technical option that can be pursued "
    "in parallel with, and potentially in support of, cooperative approaches."
)

# ============================================================
# 2. LITERATURE REVIEW
# ============================================================
add_heading1("Hydro-hegemony and the politics of flood infrastructure")

c_ma = cite("MirumachAllan2007")
c_mid = cite("Middleton2022")
c_ras = cite("Rasanen2017")
c_mrc = cite("MRC2023")
c_leb = cite("Lebel2005")
c_ku = cite("Kuenzer2013")
c_wh = cite("Wheeler2020")
c_z08 = cite("Zeitoun2008")
c_mc = cite("McCaffrey2007")

add_heading2("Hydro-hegemony and downstream vulnerability")

add_para(
    "The hydro-hegemony framework, as articulated by Zeitoun & Warner,{" + c_zw + "} "
    "describes how powerful riparian states exploit their geographic, economic, "
    "or military advantages to control shared water resources. While the original "
    "formulation emphasised water allocation and dam construction, subsequent "
    "scholarship has extended the concept to encompass flood risk. Downstream "
    "states in major transboundary basins \u2014 Vietnam on the Mekong, Egypt on "
    "the Nile, Bangladesh on the Ganges-Brahmaputra-Meghna (GBM) \u2014 face flood "
    "risks that are substantially shaped by upstream land use, dam operations, "
    "and climate change, yet have limited influence over these upstream "
    "drivers.{" + c_ma + "," + c_mid + "}"
)

add_para(
    "The Mekong River exemplifies this dynamic. China\u2019s cascade of 11 "
    "mainstream dams in the upper Mekong (Lancang) has fundamentally altered "
    "the river\u2019s hydrology, with documented effects on downstream flood and "
    "drought patterns in Thailand, Laos, Cambodia, and Vietnam.{" + c_ras + "," + c_mrc + "} "
    "Vietnam\u2019s Mekong Delta \u2014 home to 17 million people and producing "
    "over half of the country\u2019s rice \u2014 faces the dual threat of upstream "
    "flow modification and sea-level rise, with limited recourse through "
    "existing governance mechanisms.{" + c_leb + "," + c_ku + "}"
)

add_para(
    "Similarly, Egypt\u2019s concerns over the Grand Ethiopian Renaissance Dam "
    "(GERD) on the Blue Nile reflect the asymmetry between upstream development "
    "prerogatives and downstream flood and water security risks.{" + c_cz + "," + c_wh + "} "
    "In the Jordan River basin, Israel\u2019s control of the headwaters and the "
    "Sea of Galilee has historically constrained Jordanian and Palestinian "
    "water management options.{" + c_z08 + "} In each case, downstream states "
    "lack autonomous mechanisms to manage flood risk independently of upstream "
    "decisions."
)

add_para(
    "The 1997 United Nations Convention on the Law of the Non-Navigational Uses "
    "of International Watercourses established principles of equitable utilisation "
    "and the obligation not to cause significant harm, but enforcement mechanisms "
    "remain weak and disputes continue to be resolved \u2014 or not \u2014 through "
    "bilateral negotiation.{" + c_mc + "} The Mekong River Commission, despite "
    "decades of institutional development, has been unable to prevent unilateral "
    "dam construction on the mainstream by China and Laos.{" + c_leb + "," + c_mrc + "} "
    "These institutional failures underscore the need for technical solutions "
    "that do not depend on political cooperation."
)

c_mlit20a = cite("MLIT2020a")
c_bk = cite("Berkowitz2019")

add_heading2("Conventional flood infrastructure and its limitations")

add_para(
    "Conventional structural flood measures \u2014 levees, dams, diversion channels, "
    "and floodways \u2014 require either upstream cooperation (for dam operation "
    "coordination) or massive capital investment within the downstream state\u2019s "
    "territory. Japan\u2019s experience illustrates both the achievements and "
    "limitations of the structural approach. The MAOUDC in Saitama Prefecture, "
    "completed in 2006, transfers floodwater from four rivers into the Edogawa "
    "River at a capacity of 200 m\u00b3/s \u2014 a remarkable engineering achievement, "
    "but one that required 13 years of construction and approximately US$2 "
    "billion.{" + c_mlit06 + "} Super-levees along the Arakawa River, designed to "
    "withstand overtopping, have reached only 15% completion after decades of "
    "construction.{" + c_mlit20a + "}"
)

add_para(
    "The Mississippi River basin in the United States provides a continental-scale "
    "example of structural limitations. The Yazoo Backwater area in Mississippi "
    "experiences chronic backwater flooding when the Mississippi River rises and "
    "prevents the Yazoo River from draining \u2014 a mechanism identical to the "
    "Oda\u2013Takahashi backwater effect in Japan, but at a vastly larger scale. "
    "A proposed pumping station to mitigate Yazoo backwater flooding has been "
    "debated for over 80 years, blocked by environmental concerns and interstate "
    "political disputes.{" + c_bk + "}"
)

c_mlitc = cite("MLITChugoku2019")
c_mlit24 = cite("MLIT2024")

add_para(
    "The backwater effect \u2014 whereby rising water levels in a main river impede "
    "the drainage of tributary streams \u2014 represents a particularly challenging "
    "flood mechanism. During the 2018 Western Japan Heavy Rainfall, the Takahashi "
    "River\u2019s water level rose so rapidly that the Oda River, a tributary joining "
    "from the west, could not discharge its floodwaters. The resulting backwater "
    "inundation flooded 8.28 km\u00b2 of the Mabi district in Kurashiki City, "
    "killing 51 people and destroying over 4,600 homes.{" + c_nied + "," + c_mlitc + "} "
    "Japan\u2019s response was a \u00a528 billion confluence relocation project, "
    "completed in March 2024, which moved the Oda\u2013Takahashi junction 4.6 km "
    "downstream to reduce backwater propagation.{" + c_mlit24 + "}"
)

c_dil = cite("Dillon2019")
c_ch = cite("Chang2019")
c_nz = cite("NishizawaZhang2026")

add_heading2("Managed aquifer recharge and Flood-MAR")

add_para(
    "Managed aquifer recharge (MAR) \u2014 the intentional replenishment of "
    "groundwater through engineered infiltration \u2014 has been practised for "
    "decades in water supply contexts.{" + c_dil + "} The California Department of "
    "Water Resources has pioneered \u2018Flood-MAR\u2019 (Flood-Managed Aquifer "
    "Recharge), which explicitly links flood management with groundwater "
    "banking by directing excess surface water during flood events into "
    "agricultural lands and dedicated recharge basins.{" + c_dwr + "} Taiwan has "
    "implemented optimised groundwater pumping-recharge cycles for combined "
    "flood control and water supply.{" + c_ch + "} In Japan, Nishizawa & Zhang "
    "have documented a half-century of groundwater monitoring data from central "
    "Japan, demonstrating the long-term viability of managed aquifer "
    "systems.{" + c_nz + "}"
)

c_gsj = cite("GSJAIST2020")

add_para(
    "However, existing Flood-MAR applications are passive \u2014 they rely on "
    "gravity-driven infiltration during flood events, which limits their "
    "applicability to regions with suitable topography and soil permeability. "
    "The framework proposed in this paper extends Flood-MAR by introducing an "
    "active, energy-driven component: using hydropower from planned dam releases "
    "to pre-emptively pump groundwater, thereby creating storage capacity before "
    "flood events occur."
)

add_para(
    "The relationship between groundwater extraction and land subsidence is well "
    "documented and represents a key constraint on managed aquifer drawdown. "
    "Tokyo\u2019s Kanto Plain experienced up to 4.5 m of cumulative land subsidence "
    "between the Meiji era and the 1970s due to industrial groundwater pumping, "
    "prompting stringent regulations that have since allowed groundwater levels "
    "to recover substantially.{" + c_gsj + "} The Taiwan experience offers further "
    "insight: Chang et al. demonstrated that optimised pumping-recharge cycles can "
    "manage subsidence risk while achieving flood control objectives, provided "
    "that drawdown is temporary (days to weeks) rather than sustained.{" + c_ch + "} "
    "This distinction between chronic extraction (which causes compaction of "
    "clay layers and irreversible subsidence) and cyclical, short-duration "
    "drawdown (which primarily affects elastic storage in sand-gravel aquifers) "
    "is critical to the feasibility of the proposed framework."
)

# ============================================================
# 3. PROPOSED FRAMEWORK
# ============================================================
add_heading1("The proposed framework")

add_para(
    "This section presents the technical architecture of the proposed framework. "
    "The system integrates three established technologies \u2014 hydroelectric "
    "generation from controlled dam releases, submersible groundwater pumping, "
    "and managed aquifer recharge \u2014 into a coordinated cycle that can be "
    "activated in anticipation of flood events. The key innovation lies not in "
    "any individual component but in their integration into a single "
    "energy-positive system that can operate within the jurisdiction of a single "
    "state."
)

add_para(
    "The framework operates in four phases. Phase 1 (Planned release and power "
    "generation): When flood risk is forecast, upstream dam operators release "
    "water through turbines at controlled rates, generating hydroelectric power. "
    "The energy potential is given by P_gen = \u03b7_t \u00d7 \u03c1 \u00d7 g "
    "\u00d7 Q \u00d7 H, where \u03b7_t is turbine efficiency (typically 0.85), "
    "\u03c1 is water density (1,000 kg/m\u00b3), g is gravitational acceleration "
    "(9.81 m/s\u00b2), Q is discharge (m\u00b3/s), and H is the effective head (m)."
)

add_para(
    "Phase 2 (Groundwater pumping in adjacent watershed): The generated "
    "electricity powers submersible pumps that extract groundwater from the "
    "alluvial aquifer of an adjacent watershed. The pumping power requirement "
    "is P_pump = (\u03c1 \u00d7 g \u00d7 Q_p \u00d7 H_p) / \u03b7_p, where "
    "Q_p is the pumping rate, H_p is the pumping head, and \u03b7_p is pump "
    "efficiency (typically 0.80)."
)

add_para(
    "Phase 3 (Flood absorption): The pre-emptied aquifer storage volume "
    "(V = A \u00d7 \u0394h \u00d7 S_y, where A is the recharge area, \u0394h is "
    "the water table drawdown, and S_y is the specific yield) provides "
    "capacity to absorb floodwaters through infiltration during the event."
)

add_para(
    "Phase 4 (Post-flood recovery): After the flood recedes, the aquifer "
    "naturally recharges through precipitation and lateral groundwater flow, "
    "restoring the system\u2019s capacity for the next event."
)

add_para(
    "The system is energy self-sufficient when P_gen > P_pump. The surplus "
    "energy (P_gen \u2212 P_pump) can be supplied to the electricity grid, "
    "creating an economic co-benefit. Crucially, the entire system \u2014 dam, "
    "turbines, pumping wells, and target aquifer \u2014 can be located within a "
    "single national jurisdiction, requiring no transboundary coordination."
)

# ============================================================
# 4. FEASIBILITY ANALYSIS: ODA-TAKAHASHI BASIN
# ============================================================
add_heading1("Feasibility analysis: the Oda River\u2013Takahashi River basin")

c_ksb = cite("KSB2025")

add_heading2("Study area and the 2018 disaster")

add_para(
    "The Oda River\u2013Takahashi River basin in Okayama Prefecture, western "
    "Japan, provides an ideal test case for the proposed framework. The Oda "
    "River (basin area 211 km\u00b2) joins the Takahashi River at an acute angle, "
    "making it highly susceptible to backwater effects. During the 2018 Western "
    "Japan Heavy Rainfall (5\u20138 July 2018), three-day cumulative rainfall "
    "reached 320 mm, exceeding the design rainfall of the flood control system. "
    "The Takahashi River\u2019s rapid rise prevented the Oda River from discharging, "
    "causing levee breaches at five locations and inundating 8.28 km\u00b2 of the "
    "Mabi district with an estimated flood volume of 15.3 \u00d7 10\u2076 "
    "m\u00b3 (Table 1).{" + c_nied + "," + c_mlitc + "}"
)

# Table 1: 2018 event parameters
add_table_wp(
    ["Parameter", "Value"],
    [
        ["Three-day cumulative rainfall", "320 mm"],
        ["Inundation area", "8.28 km\u00b2"],
        ["Estimated flood volume", "15.3 \u00d7 10\u2076 m\u00b3"],
        ["Levee breaches", "5 locations"],
        ["Fatalities", "51"],
        ["Basin area (Oda River)", "211 km\u00b2"],
        ["Channel gradient (Oda River)", "1/2,200"],
    ],
    "Table 1. Key parameters of the 2018 Western Japan Heavy Rainfall event "
    "in the Oda River\u2013Takahashi River basin."
)

add_para(
    "In March 2024, a confluence relocation project was completed, moving the "
    "Oda\u2013Takahashi junction 4.6 km downstream. Hydraulic analysis indicates "
    "this reduces the Oda River water level by approximately 4.6 m and the "
    "Takahashi River by 0.8 m under design flood conditions.{" + c_mlit24 + "," + c_ksb + "}"
)

add_heading2("Energy balance")

add_para(
    "For a planned release through a dam with an effective head of H = 100 m "
    "and a discharge of Q = 50 m\u00b3/s, the generated power is: "
    "P_gen = 0.85 \u00d7 1,000 \u00d7 9.81 \u00d7 50 \u00d7 100 = 41.7 MW. "
    "For groundwater pumping at Q_p = 10 m\u00b3/s with a pumping head of "
    "H_p = 30 m: P_pump = (1,000 \u00d7 9.81 \u00d7 10 \u00d7 30) / 0.80 = "
    "3.68 MW. The system thus consumes only 8.8% of the generated power for "
    "groundwater management, leaving a surplus of 38.0 MW available for the "
    "electricity grid. Figure 1 shows the energy balance as a function of "
    "head and discharge."
)

# Figure 1
add_figure(
    "fig1_energy_balance_en.png",
    "Figure 1. Energy balance contour map showing generated power (MW) as a "
    "function of effective head and discharge. The red dashed line indicates "
    "the pumping power requirement (3.68 MW). The green marker shows the "
    "reference scenario (H = 100 m, Q = 50 m\u00b3/s)."
)

add_heading2("Aquifer storage capacity")

add_para(
    "For a recharge area of A = 20 km\u00b2, a water table drawdown of "
    "\u0394h = 3 m, and a specific yield of S_y = 0.2 (typical of alluvial "
    "sand-gravel aquifers): V = 20 \u00d7 10\u2076 \u00d7 3 \u00d7 0.2 = "
    "12.0 \u00d7 10\u2076 m\u00b3. This represents 78% of the 2018 flood "
    "volume (15.3 \u00d7 10\u2076 m\u00b3), suggesting that pre-emptive "
    "aquifer management could absorb the majority of a design-level flood "
    "event. Figure 2 presents the storage capacity as a function of recharge "
    "area and water table drawdown."
)

# Figure 2
add_figure(
    "fig2_aquifer_storage_en.png",
    "Figure 2. Aquifer storage capacity contours (10\u2076 m\u00b3) as a "
    "function of recharge area and water table drawdown, assuming a specific "
    "yield of 0.2. The dashed line indicates the 2018 flood volume "
    "(15.3 \u00d7 10\u2076 m\u00b3). The green marker shows the reference scenario."
)

add_heading2("Flood hydrograph modification")

add_para(
    "A simplified 72-hour rainfall-runoff simulation was conducted comparing "
    "four scenarios: (1) no measures (baseline); (2) confluence relocation only "
    "(4.6 m water level reduction, equivalent to approximately 30% peak "
    "discharge reduction); (3) groundwater management only (pre-emptied aquifer "
    "absorbs infiltration, reducing runoff volume by 41%); and (4) integrated "
    "measures (confluence relocation plus groundwater management). The integrated "
    "scenario achieves a 32% peak discharge reduction compared with baseline "
    "(Figure 3), with the groundwater management component contributing "
    "primarily through volume reduction rather than peak attenuation (Figure 4; Table 2)."
)

# Figure 3
add_figure(
    "fig3_hydrograph_en.png",
    "Figure 3. Flood hydrograph comparison for the 2018 event under four "
    "scenarios: no measures (baseline), confluence relocation only, groundwater "
    "management only, and integrated measures. Upper panel shows the 72-hour "
    "rainfall hyetograph."
)

# Figure 4
add_figure(
    "fig4_scenario_comparison_en.png",
    "Figure 4. Comparison of peak discharge and total runoff volume across "
    "four scenarios. Error bars indicate \u00b110% sensitivity range."
)

# Table 2: Scenario comparison
add_table_wp(
    ["Scenario", "Peak discharge (m\u00b3/s)", "Peak reduction (%)",
     "Infiltration volume (10\u2076 m\u00b3)"],
    [
        ["1. No measures (baseline)", "850", "\u2014", "0"],
        ["2. Confluence relocation", "595", "30", "0"],
        ["3. Groundwater management", "833", "2", "6.3"],
        ["4. Integrated", "578", "32", "6.3"],
    ],
    "Table 2. Comparison of flood control scenarios for the 2018 event."
)

# ============================================================
# 5. SCALING UP: TOKYO METROPOLITAN AREA
# ============================================================
c_mlit20b = cite("MLIT2020b")

add_heading1("Scaling up: the Tokyo metropolitan area")

add_para(
    "While the Oda River\u2013Takahashi River basin provides a compelling "
    "proof-of-concept, the true significance of the proposed framework lies "
    "in its potential application to densely populated metropolitan areas where "
    "conventional flood infrastructure faces severe space constraints. The "
    "Arakawa\u2013Edogawa basin in the Tokyo metropolitan area presents such a "
    "case. An estimated 8.21 million people reside within projected inundation "
    "zones, and the existing flood control infrastructure \u2014 while extensive "
    "\u2014 faces capacity constraints under intensifying rainfall "
    "patterns (Table 3).{" + c_mlit20b + "}"
)

# Table 3: Tokyo infrastructure
add_table_wp(
    ["Facility", "Capacity/Scale", "Year"],
    [
        ["Arakawa Floodway", "22 km diversion channel", "1930"],
        ["Iwabuchi Floodgate", "Isolates Sumida from Arakawa", "1982 (rebuilt)"],
        ["Saiko Retarding Basin", "39 \u00d7 10\u2076 m\u00b3", "1997"],
        ["Upstream dams (3)", "Combined storage", "Various"],
        ["MAOUDC", "200 m\u00b3/s inter-basin transfer", "2006"],
        ["Super-levees", "15% completed (Arakawa)", "Ongoing"],
    ],
    "Table 3. Existing flood control infrastructure in the Arakawa\u2013Edogawa basin."
)

add_para(
    "The Kanto Plain beneath Tokyo hosts alluvial aquifers reaching 80 m in "
    "thickness, comprising alternating sand-gravel and clay layers mapped by "
    "the Geological Survey of Japan\u2019s three-dimensional geological "
    "model.{" + c_gsj + "} Historical groundwater extraction caused severe land "
    "subsidence from the Meiji era through the 1970s, but stringent pumping "
    "regulations implemented since then have allowed groundwater levels to "
    "recover substantially \u2014 paradoxically creating conditions favourable "
    "for managed drawdown during flood emergencies."
)

add_para(
    "Applying the proposed framework to the mid-Arakawa basin: planned releases "
    "from the Takizawa Dam (effective head ~100 m) would generate hydroelectric "
    "power to pump groundwater from a 30 km\u00b2 mid-basin area, creating "
    "approximately 13.5 \u00d7 10\u2076 m\u00b3 of subsurface storage "
    "(30 \u00d7 10\u2076 \u00d7 3 \u00d7 0.15). Combined with the existing Saiko "
    "Retarding Basin (39 \u00d7 10\u2076 m\u00b3), the total available storage "
    "capacity would reach approximately 52.5 \u00d7 10\u2076 m\u00b3 \u2014 a "
    "substantial addition to the metropolitan area\u2019s flood resilience."
)

# ============================================================
# 6. IMPLICATIONS FOR TRANSBOUNDARY FLOOD GOVERNANCE
# ============================================================
add_heading1("Implications for transboundary flood governance")

add_heading2("Challenging the upstream-dominance paradigm")

add_para(
    "The proposed framework has significant implications for the hydro-hegemony "
    "literature. By enabling downstream states to create flood storage capacity "
    "within their own territories, the framework challenges the assumption that "
    "downstream flood vulnerability is an inevitable consequence of upstream "
    "control. This does not eliminate the power asymmetry \u2014 upstream dam "
    "operations still affect downstream hydrology \u2014 but it provides a "
    "technological pathway for downstream actors to reduce their dependence "
    "on upstream cooperation for flood management."
)

add_para(
    "Cascao & Zeitoun have argued that counter-hegemonic strategies in "
    "transboundary water governance typically require either international "
    "legal mechanisms or coalition-building among weaker riparian states.{" + c_cz + "} "
    "The framework proposed here offers a third pathway: unilateral technical "
    "adaptation that operates independently of diplomatic processes. This is "
    "particularly relevant in basins where political tensions preclude "
    "cooperative flood management, such as the Indus (India\u2013Pakistan), the "
    "Tigris-Euphrates (Turkey\u2013Iraq\u2013Syria), and the Jordan "
    "(Israel\u2013Jordan\u2013Palestine)."
)

add_heading2("International legal dimensions")

add_para(
    "The 1997 United Nations Watercourses Convention requires states to utilise "
    "international watercourses in an equitable and reasonable manner and to take "
    "appropriate measures to prevent significant harm.{" + c_mc + "} However, these "
    "obligations apply to \u2018international watercourses\u2019 defined as surface "
    "water systems that cross or form international boundaries. Groundwater that "
    "is not connected to an international watercourse falls outside the "
    "Convention\u2019s scope. If a downstream state manages its own alluvial "
    "aquifers to create flood storage capacity, it operates in a legal grey area "
    "that may not trigger the notification and consultation obligations that "
    "apply to surface water interventions."
)

add_heading2("Transboundary river applications")

add_para(
    "Mekong River basin: Vietnam could implement subsurface storage management "
    "in the Mekong Delta\u2019s alluvial aquifers, using hydropower from domestic "
    "dam releases to pre-emptively lower groundwater levels before the monsoon "
    "flood season. This would complement existing flood management strategies "
    "without requiring coordination with upstream Lancang dams in "
    "China.{" + c_leb + "," + c_ras + "}"
)

add_para(
    "Nile River basin: Egypt\u2019s Nile Delta, already vulnerable to subsidence "
    "and sea-level rise, could benefit from managed aquifer storage that absorbs "
    "excess flood releases from the Aswan High Dam during high-water years. The "
    "Nile Delta\u2019s extensive Quaternary aquifer system provides potential "
    "storage capacity.{" + c_wh + "}"
)

add_para(
    "Ganges-Brahmaputra-Meghna basin: Bangladesh, which faces annual devastating "
    "floods from three major transboundary rivers, could develop subsurface "
    "storage in the Bengal Basin\u2019s thick alluvial deposits \u2014 one of the "
    "world\u2019s largest aquifer systems \u2014 as an autonomous flood management "
    "strategy independent of upstream Indian dam operations."
)

add_para(
    "Indus River basin: Pakistan\u2019s 2010 and 2022 mega-floods demonstrated "
    "the country\u2019s extreme vulnerability to upstream hydrological events. "
    "The thick alluvial deposits of the Indus Plain could potentially support "
    "subsurface flood storage, allowing Pakistan to develop autonomous flood "
    "management capacity independent of the 1960 Indus Waters Treaty framework."
)

add_heading2("Deltaic flood protection and arid-region applications")

add_para(
    "Delta regions face a unique \u2018double squeeze\u2019 between upstream "
    "flood pulses and downstream sea-level rise. Subsurface storage management "
    "offers a complementary approach that reduces flood peaks without interfering "
    "with sediment transport processes. This is particularly relevant for the "
    "Mekong Delta (Vietnam), the Nile Delta (Egypt), and the GBM Delta "
    "(Bangladesh)."
)

add_para(
    "In arid and semi-arid regions, flash floods cause significant damage "
    "despite the general scarcity of water. Countries in the Arabian Peninsula "
    "could adapt the proposed framework to capture wadi flood energy and direct "
    "floodwaters into depleted aquifers, simultaneously addressing flood risk "
    "and water scarcity. The availability of substantial financial resources "
    "could facilitate rapid implementation."
)

# ============================================================
# 7. DISCUSSION
# ============================================================
c_iea = cite("IEA2021")
c_mlit21 = cite("MLIT2021")

add_heading1("Discussion")

add_heading2("Energy self-sufficiency and economic viability")

add_para(
    "The energy balance analysis demonstrates a striking asymmetry: with a head "
    "of 100 m and discharge of 50 m\u00b3/s, the system generates 41.7 MW while "
    "requiring only 3.68 MW for pumping \u2014 a consumption ratio of just 8.8%. "
    "This compares favourably with conventional pumped-storage hydropower, which "
    "typically achieves 70\u201380% round-trip efficiency.{" + c_iea + "} The surplus "
    "energy (38.0 MW) represents both an economic co-benefit and an incentive "
    "for implementation. Japan\u2019s extensive pumped-storage hydropower "
    "infrastructure provides the technical expertise and supply chains necessary "
    "for rapid deployment.{" + c_tepco + "}"
)

add_para(
    "From an economic perspective, the surplus energy represents a potential "
    "revenue stream that could offset the capital costs of well field construction "
    "and aquifer monitoring. At a wholesale electricity price of US$0.10/kWh, "
    "38 MW of surplus power generated over a 72-hour flood event would produce "
    "approximately US$274,000 in revenue per event. While this alone would not "
    "justify the infrastructure investment, the combined benefits of flood damage "
    "avoidance, energy generation, and potential carbon credits create a more "
    "compelling economic case."
)

add_heading2("Complementarity with existing measures")

add_para(
    "The proposed framework is not intended to replace structural flood measures "
    "but to complement them. The integrated scenario (confluence relocation plus "
    "groundwater management) achieves a 32% peak discharge reduction \u2014 greater "
    "than either measure alone \u2014 and reduces total runoff volume by 41% through "
    "aquifer infiltration (Table 2). This complementarity is significant for "
    "Japan\u2019s Basin Flood Control (Ryuiki Chisui) policy, which explicitly "
    "calls for diversified, multi-layered flood management approaches.{" + c_mlit21 + "}"
)

add_para(
    "The simulated flood hydrograph (Figure 3) reveals an important distinction. "
    "Confluence relocation primarily reduces peak discharge by lowering the "
    "backwater effect, while groundwater management primarily reduces total "
    "runoff volume by diverting water into subsurface storage. This functional "
    "complementarity means that the two approaches target different aspects of "
    "flood risk: the peak (which determines levee overtopping) and the volume "
    "(which determines inundation duration and total damage)."
)

add_heading2("The politics of autonomous flood adaptation")

add_para(
    "The framework raises important questions about the politics of autonomous "
    "adaptation. If downstream states can reduce their flood vulnerability "
    "independently of upstream cooperation, does this weaken incentives for "
    "cooperative basin governance? Or does it, paradoxically, strengthen the "
    "negotiating position of downstream states by reducing their dependence on "
    "upstream goodwill? The hydro-hegemony literature suggests that power "
    "asymmetries in transboundary basins are self-reinforcing: powerful upstream "
    "states have little incentive to cooperate because the status quo serves "
    "their interests.{" + c_zw + "} If downstream states can demonstrate autonomous "
    "flood management capacity, this may alter the calculus of cooperation."
)

add_heading2("Limitations and further research")

add_para(
    "This study employs a simplified analytical framework and does not "
    "constitute a detailed hydrological model. Key limitations include: "
    "(1) the rainfall-runoff simulation uses a unit hydrograph approach rather "
    "than physically based models such as GSFLOW or ParFlow; (2) groundwater\u2013"
    "surface water interactions are represented parametrically rather than "
    "through coupled simulation; (3) land subsidence risks from cyclical aquifer "
    "drawdown require site-specific geotechnical assessment; and (4) infiltration "
    "rates during flood events may be limited by soil saturation and surface "
    "sealing. Future work should include coupled groundwater\u2013surface water "
    "modelling, pilot field studies, and institutional analysis of the regulatory "
    "frameworks needed to implement managed aquifer drawdown for flood control."
)

add_para(
    "From a governance perspective, the implementation of managed aquifer "
    "drawdown for flood control would require new regulatory frameworks in most "
    "jurisdictions. Groundwater pumping is typically regulated for water supply "
    "purposes, and the concept of deliberately drawing down aquifers to create "
    "flood storage capacity does not fit neatly within existing regulatory "
    "categories. In Japan, groundwater extraction is regulated at the prefectural "
    "level through various ordinances originally designed to prevent land "
    "subsidence \u2014 these would need to be amended to permit temporary, cyclical "
    "drawdown for flood control purposes."
)

# ============================================================
# 8. CONCLUSIONS
# ============================================================
add_heading1("Conclusions")

add_para_no_indent(
    "This paper has proposed and analysed a novel flood control framework that "
    "integrates planned-release hydropower with inter-watershed groundwater "
    "management, and has situated this technical proposal within the broader "
    "politics of transboundary flood governance. The key findings are as follows."
)

add_para(
    "First, the framework is energy self-sufficient under the conditions "
    "analysed: hydropower generation from a 100 m head dam release at "
    "50 m\u00b3/s produces 41.7 MW, of which only 3.68 MW (8.8%) is required "
    "for groundwater pumping."
)

add_para(
    "Second, pre-emptive aquifer drawdown can create storage capacity equivalent "
    "to 78% of the observed 2018 flood volume in the Oda River\u2013Takahashi "
    "River basin, and the Tokyo metropolitan application suggests "
    "52.5 \u00d7 10\u2076 m\u00b3 of combined surface\u2013subsurface storage "
    "capacity."
)

add_para(
    "Third, the framework can be implemented unilaterally by downstream states, "
    "circumventing the governance challenges that constrain transboundary flood "
    "management. This represents a new dimension in the hydro-hegemony "
    "literature: a technologically grounded counter-hegemonic strategy for flood "
    "risk reduction."
)

add_para(
    "Fourth, the framework has potential applicability to transboundary river "
    "basins (Mekong, Nile, GBM, Indus, Tigris-Euphrates), deltaic regions "
    "facing the dual threat of upstream floods and sea-level rise, and arid "
    "environments where flash flood energy could be captured for aquifer "
    "replenishment."
)

add_para(
    "Flood management is not merely a technical problem but a political one. "
    "The framework proposed here does not resolve the power asymmetries inherent "
    "in transboundary water governance, but it offers downstream actors a "
    "concrete technological pathway to reduce their flood vulnerability "
    "independently of upstream decisions. As climate change intensifies flood "
    "risks in transboundary basins worldwide, such autonomous adaptation "
    "strategies deserve serious consideration in both technical and policy "
    "domains."
)

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
add_heading1("Acknowledgements")

add_para_no_indent("None.")

# ============================================================
# DECLARATION OF COMPETING INTEREST
# ============================================================
add_heading1("Declaration of competing interest")

add_para_no_indent(
    "The author declares no conflicts of interest."
)

# ============================================================
# FUNDING
# ============================================================
add_heading1("Funding")

add_para_no_indent(
    "This work received no external funding."
)

# ============================================================
# DATA AVAILABILITY STATEMENT
# ============================================================
add_heading1("Data availability statement")

add_para_no_indent(
    "All data and simulation code used in this study are publicly available at: "
    "https://github.com/bougtoir/flood-prediction-review"
)

# ============================================================
# AI USE DECLARATION
# ============================================================
add_heading1("Declaration of generative AI use")

add_para_no_indent(
    "The author used Devin (devin.ai) to assist with manuscript formatting, "
    "figure generation code, and language editing. The author takes full "
    "responsibility for the accuracy and content of the manuscript."
)

# ============================================================
# REFERENCES (Vancouver style, numbered in order of appearance)
# ============================================================
add_heading1("References")

for key, full_text in ref_list:
    num = ref_db[key]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    # Number
    run_num = p.add_run(f"{num}. ")
    run_num.font.size = Pt(10)
    run_num.font.name = "Times New Roman"
    # Reference text (handle italics for journal/book titles)
    parts = re.split(r'(\*[^*]+\*)', full_text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

# ── Save ──
outpath = os.path.join(OUTDIR, "wp_manuscript.docx")
doc.save(outpath)
print(f"Water Policy manuscript saved: {outpath}")

# Word count estimate (excluding abstract, references, tables, captions)
total_words = 0
in_refs = False
in_abstract = False
for p in doc.paragraphs:
    txt = p.text.strip()
    if 'REFERENCES' in txt:
        in_refs = True
    if 'ABSTRACT' in txt:
        in_abstract = True
        continue
    if in_abstract and txt.startswith('Keywords'):
        in_abstract = False
        continue
    if not in_refs and not in_abstract:
        total_words += len(txt.split())

print(f"Estimated word count (excl. abstract and references): ~{total_words}")
print(f"Figures: 4, Tables: 3 -> word deduction: 7 x 350 = 2,450")
print(f"Effective word limit: 8,000 - 2,450 = 5,550 words")
print(f"Total references: {len(ref_list)}")
