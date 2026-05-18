#!/usr/bin/env python3
"""
Reporting Guidelines Checklist for Water Alternatives Research Article — as .docx
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"

doc = Document()

# Page setup (A4)
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(20)
section.right_margin = Mm(20)

# Style
style = doc.styles["Normal"]
style.font.name = "Verdana"
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(4)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Verdana"
    run.font.size = Pt(14)
    run.bold = True
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = "Verdana"
    run.font.size = Pt(10)
    run.italic = True
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Verdana"
    if level == 1:
        run.font.size = Pt(11)
        run.bold = True
    else:
        run.font.size = Pt(10)
        run.bold = True


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Verdana"
    run.font.size = Pt(9)
    return p


def add_checklist_table(items):
    """items: list of (number, item_text, section, status)"""
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.style = "Table Grid"

    # Header
    headers = ["#", "Item", "Section", "✓"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Verdana"
        run.font.size = Pt(8)
        run.bold = True
        # Grey background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "D9E2F3"
        })
        shading.append(shading_elm)

    # Rows
    for row_idx, (num, item_text, section_text, status) in enumerate(items):
        row = table.rows[row_idx + 1]
        data = [str(num), item_text, section_text, status]
        for col_idx, val in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Verdana"
            run.font.size = Pt(8)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(9.0)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(1.5)

    return table


# ── Title ──
add_title("Reporting Guidelines Checklist")
add_subtitle(
    "Water Alternatives Research Article\n"
    "Unilateral flood control through subsurface storage management"
)

# ── Introductory note ──
add_note(
    "There is no established reporting guideline equivalent to STROBE/CONSORT for "
    "water governance/hydro-political research articles. This checklist is adapted "
    "from best practices for interdisciplinary water research, incorporating "
    "elements from: FAIR Data Principles (data transparency), CRediT (Contributor "
    "Roles Taxonomy), hydrological model reporting (Clark et al., 2015), and WaA "
    "Guidelines for Authors (journal-specific requirements)."
)

# ── A. Study Design and Framing ──
add_heading("A. Study Design and Framing")
add_checklist_table([
    (1, "Research question clearly stated", "Introduction", "☐"),
    (2, "Theoretical framework identified (hydro-hegemony)", "Lit. Review", "☐"),
    (3, "Study type described (feasibility + governance analysis)", "Introduction", "☐"),
    (4, "Scope and limitations acknowledged", "Discussion", "☐"),
    (5, "Policy relevance articulated", "Throughout", "☐"),
])

# ── B. Literature Review ──
add_heading("B. Literature Review")
add_checklist_table([
    (6, "Key theories reviewed (hydro-hegemony, Flood-MAR)", "Section 2", "☐"),
    (7, "Transboundary water governance literature engaged", "Section 2", "☐"),
    (8, "Technical precedents cited (California DWR, ASR, pumped-storage)", "Section 2", "☐"),
    (9, "Research gap identified", "Introduction", "☐"),
    (10, "Prior work on study area cited", "Section 4", "☐"),
])

# ── C. Data and Methods ──
add_heading("C. Data and Methods")
add_checklist_table([
    (11, "Data sources identified", "Throughout", "☐"),
    (12, "Data publicly available or access described", "Acknowledgements", "☐"),
    (13, "Analytical approach described (energy balance, storage, hydrograph)", "Section 3", "☐"),
    (14, "Key parameters listed with values and sources", "Section 4, tables 1–3", "☐"),
    (15, "Equations presented with variable definitions", "Section 3", "☐"),
    (16, "Assumptions stated explicitly", "Discussion (Limitations)", "☐"),
    (17, "Model type described (simplified analytical vs. physically-based)", "Discussion", "☐"),
    (18, "Sensitivity analysis conducted (±10% error bars)", "figure 4, Discussion", "☐"),
])

# ── D. Results Presentation ──
add_heading("D. Results Presentation")
add_checklist_table([
    (19, "Results presented with quantitative values", "Section 4", "☐"),
    (20, "Scenarios compared systematically", "table 2, figure 4", "☐"),
    (21, "Figures and tables cited in order", "Verified", "☐"),
    (22, "All figures/tables have captions", "Verified", "☐"),
    (23, "Units consistent (SI/metric) throughout", "", "☐"),
    (24, "Uncertainty quantified where applicable", "Error bars, ±10%", "☐"),
])

# ── E. Discussion and Governance Analysis ──
add_heading("E. Discussion and Governance Analysis")
add_checklist_table([
    (25, "Results interpreted within theoretical framework", "Discussion", "☐"),
    (26, "Policy implications discussed", "Discussion", "☐"),
    (27, "International applicability assessed", "Section 6", "☐"),
    (28, "Multiple transboundary basins considered", "Section 6", "☐"),
    (29, "Legal dimensions addressed (UN Watercourses Convention)", "Section 6", "☐"),
    (30, "Limitations explicitly stated", "Discussion", "☐"),
    (31, "Future research directions identified", "Discussion", "☐"),
    (32, "Alternative explanations/counterarguments considered", "Discussion", "☐"),
])

# ── F. Ethics and Transparency ──
add_heading("F. Ethics and Transparency")
add_checklist_table([
    (33, "AI use declaration present", "Cover Page", "☐"),
    (34, "Acknowledgements present", "Cover Page", "☐"),
    (35, "Conflicts of interest: none declared", "(not required by WaA)", "☐"),
    (36, "Related submissions disclosed", "Cover letter", "☐"),
    (37, "No author-identifying info in main text (double blind)", "Verified", "☐"),
    (38, "CRediT author contributions", "To be added by author", "☐"),
])

# ── G. Reference Quality ──
add_heading("G. Reference Quality")
add_checklist_table([
    (39, "All citations appear in reference list", "Verified (27/27)", "☐"),
    (40, "No orphan references", "Verified", "☐"),
    (41, "Mix of peer-reviewed and grey literature appropriate", "", "☐"),
    (42, "Recent literature included (within 5 years)", "", "☐"),
    (43, "Non-English sources cited where relevant", "", "☐"),
    (44, "Government/institutional sources properly cited", "", "☐"),
])

# ── Final note ──
add_note("")
add_note(
    "Note: This checklist is provided as a quality assurance tool. Water Alternatives "
    "does not require a formal reporting guideline checklist for submission, but "
    "systematic verification of these items helps ensure manuscript quality and "
    "completeness."
)

# ── Save ──
outpath = f"{OUTDIR}/waa_reporting_checklist.docx"
doc.save(outpath)
print(f"WaA reporting checklist saved: {outpath}")
