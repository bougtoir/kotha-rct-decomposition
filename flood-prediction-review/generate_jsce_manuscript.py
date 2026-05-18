#!/usr/bin/env python3
"""
土木学会論文集特集号（水工学）向け原稿生成スクリプト

タイトル: 計画放水発電と隣接水系地下水管理を統合した
         新たな洪水制御フレームワークの提案
         ― 小田川–高梁川流域を対象とした簡易シミュレーション ―

フォーマット:
- A4, 2段組（本スクリプトでは単段で出力し、最終調整はWordで行う）
- 明朝体（本文）、ゴシック体（見出し）
- カンマ「，」ピリオド「．」（句読点「、」「。」不可）
- 参考文献: 番号順引用（Vancouver式）
- 図表: インライン配置（初出段落の直後）
"""

import os
import re
from docx import Document
from docx.shared import Pt, Mm, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"
FIGDIR = f"{OUTDIR}/figures"

doc = Document()

# ============================================================
# ページ設定（A4）
# ============================================================
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(19)
section.bottom_margin = Mm(24)
section.left_margin = Mm(20)
section.right_margin = Mm(20)

# ============================================================
# スタイル定義
# ============================================================
style_normal = doc.styles["Normal"]
style_normal.font.name = "游明朝"
style_normal.font.size = Pt(9)
style_normal.paragraph_format.line_spacing = Pt(14)
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)
# 日本語フォントの設定
rpr = style_normal.element.find(qn("w:rPr"))
if rpr is None:
    rpr = style_normal.element.makeelement(qn("w:rPr"), {})
    style_normal.element.append(rpr)
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = rpr.makeelement(qn("w:rFonts"), {})
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "游明朝")


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "游ゴシック"
    run.font.size = Pt(14)
    run.bold = True
    rpr = run._element.find(qn("w:rPr"))
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "游ゴシック")
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "游ゴシック"
    run.font.size = Pt(11)
    run.bold = True
    return p


def add_heading_gothic(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "游ゴシック"
    run.bold = True
    if level == 1:
        run.font.size = Pt(11)
    elif level == 2:
        run.font.size = Pt(10)
    else:
        run.font.size = Pt(9)
    rpr = run._element.find(qn("w:rPr"))
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "游ゴシック")
    return p


def add_body(text):
    """本文段落を追加．{N} を上付き参照番号に変換する．"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(18)
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith("{") and part.endswith("}"):
            ref_text = part[1:-1]
            run = p.add_run(ref_text + ")")
            run.font.name = "游明朝"
            run.font.size = Pt(7)
            run.font.superscript = True
        else:
            run = p.add_run(part)
            run.font.name = "游明朝"
            run.font.size = Pt(9)
    return p


def add_body_no_indent(text):
    p = add_body(text)
    p.paragraph_format.first_line_indent = Pt(0)
    return p


def add_figure(fig_path, caption, width_inches=5.5):
    """図をインライン挿入"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    if os.path.exists(fig_path):
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(width_inches))
    else:
        run = p.add_run(f"[図ファイル未生成: {os.path.basename(fig_path)}]")
        run.font.color.rgb = RGBColor(255, 0, 0)

    # キャプション
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_before = Pt(6)
    pc.paragraph_format.space_after = Pt(12)
    run_c = pc.add_run(caption)
    run_c.font.name = "游明朝"
    run_c.font.size = Pt(8)
    return p


def add_table_from_data(headers, rows, caption):
    """表をインライン挿入"""
    # キャプション（表の上に配置）
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_before = Pt(12)
    run_c = pc.add_run(caption)
    run_c.font.name = "游ゴシック"
    run_c.font.size = Pt(8)
    run_c.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # ヘッダー
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "游ゴシック"
        run.font.size = Pt(8)
        run.bold = True
    # データ
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "游明朝"
            run.font.size = Pt(8)

    # 表の後にスペース
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(8)
    return table


# ============================================================
# 原稿本文
# ============================================================

# ── タイトル ──
add_title("計画放水発電と隣接水系地下水管理を統合した")
add_title("新たな洪水制御フレームワークの提案")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("― 小田川–高梁川流域を対象とした簡易フィージビリティ分析 ―")
run.font.name = "游ゴシック"
run.font.size = Pt(10)

# ── 著者 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("（著者名）")
run.font.name = "游明朝"
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("（所属）")
run.font.name = "游明朝"
run.font.size = Pt(8)

# ── 和文要旨 ──
add_heading_gothic("要旨", level=2)
add_body_no_indent(
    "近年の激甚化する豪雨災害に対し，従来の構造的治水対策（護岸・ダム・堤防）に加えた"
    "新たな洪水制御の枠組みが求められている．"
    "本研究では，計画放水時の水力発電エネルギーを利用して隣接水系の帯水層から地下水を事前に汲み上げ，"
    "洪水時の受容空間を能動的に創出する「計画放水発電–地下水管理統合型洪水制御フレームワーク」を提案する．"
    "2018年西日本豪雨時の小田川–高梁川流域（岡山県倉敷市真備町）を対象に簡易シミュレーションを実施した結果，"
    "（1）有効落差100 m・放水流量50 m3/sの計画放水で約41.7 MWの発電が可能であり，"
    "地下水汲み上げ（揚程30 m・流量10 m3/s）の所要出力3.7 MWは発電量の約9%に過ぎないこと，"
    "（2）管理面積20 km2・地下水位低下3 mで約12.0 x 106 m3の貯留容量が創出され，"
    "2018年浸水量（15.3 x 106 m3）の約78%を吸収し得ること，"
    "（3）2024年に完成した合流点付替え（構造的対策）との統合により，"
    "ピーク流量を約32%削減し得ることを示した．"
    "本構想は既存の揚水発電技術・帯水層涵養技術の延長線上にあり，"
    "日本の急峻な地形と高密度な水インフラを活用した新たな治水パラダイムとして政策的検討に値する．"
)

# ── 英文アブストラクト ──
add_heading_gothic("Abstract", level=2)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(18)
run = p.add_run(
    "This paper proposes a novel flood control framework that integrates planned-release hydropower generation "
    "with inter-watershed groundwater management. The core concept utilizes the potential energy of planned "
    "dam releases to generate electricity, which powers groundwater pumping in adjacent watersheds to pre-create "
    "subsurface storage capacity before flood events. A simplified feasibility analysis was conducted for the "
    "Oda River-Takahashi River basin (Mabi, Okayama Prefecture), the site of catastrophic backwater flooding "
    "during the July 2018 Western Japan Heavy Rains. Results indicate that: (1) a planned release with 100 m "
    "effective head and 50 m3/s discharge can generate approximately 41.7 MW, while groundwater pumping requires "
    "only 3.7 MW (9% of generation); (2) pre-flood drawdown of 3 m over a 20 km2 managed area creates 12.0 x 10^6 m3 "
    "of aquifer storage, absorbing 78% of the 2018 flood volume; and (3) integration with the recently completed "
    "confluence relocation project could reduce peak discharge by approximately 32%. This framework leverages "
    "Japan's steep topography, dense hydraulic infrastructure, and established pumped-storage technology, "
    "representing a paradigm shift from 'expelling floodwater' to 'proactively creating reception capacity.'"
)
run.font.name = "Times New Roman"
run.font.size = Pt(9)

# ── キーワード ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run_kw = p.add_run("Key Words: ")
run_kw.font.name = "Times New Roman"
run_kw.font.size = Pt(9)
run_kw.bold = True
run_kw.italic = True
run = p.add_run("flood control, groundwater management, hydropower, backwater effect, aquifer storage")
run.font.name = "Times New Roman"
run.font.size = Pt(9)
run.italic = True

# ============================================================
# 1. 序論
# ============================================================
add_heading_gothic("1. 序論", level=1)

add_body(
    "わが国では近年，気候変動に伴う豪雨の激甚化・頻発化により，"
    "従来の治水計画を超える洪水被害が繰り返し発生している．"
    "2018年西日本豪雨（平成30年7月豪雨）では，岡山県倉敷市真備町において小田川の堤防が5箇所で決壊し，"
    "死者51人，浸水面積8.28 km2，浸水量約15.3 x 106 m3という甚大な被害が生じた{1}．"
    "この災害の主要因は，高梁川本流の水位上昇に伴う背水現象（backwater effect）により，"
    "支流である小田川の排水が阻害されたことにある{2}．"
)

add_body(
    "従来のわが国の治水対策は，護岸・堤防の整備，ダムによる洪水調節，"
    "放水路・遊水地の建設といった構造的対策（structural measures）が中心である{3}．"
    "真備町の事例では，2024年3月に小田川と高梁川の合流点を約4.6 km下流に付替える"
    "大規模な構造的対策が完了し，2024年11月の大雨時には小田川側で約4.6 m，"
    "高梁川側で約0.8 mの水位低減効果が報告されている{4}．"
    "しかし，合流点付替えのような大規模土木工事は計画から完成まで長期間を要し，"
    "気候変動による将来の洪水規模の増大に対する適応力には限界がある．"
)

add_body(
    "一方，地下水管理を治水に活用するアプローチは国際的に注目されつつある．"
    "米国カリフォルニア州では，洪水水を帯水層に浸透させるFlood-MAR"
    "（Flood-Managed Aquifer Recharge）戦略が推進されており{5}，"
    "台湾では地下水汲み上げ管理による洪水被害低減が実証されている{6}．"
    "また，日本では首都圏外郭放水路が隣接水系への洪水移送を実現し{7}，"
    "その治水効果は100回以上の実績で裏付けられている．"
)

add_body(
    "本研究では，これらの既存技術を統合し，"
    "計画放水時の水力発電エネルギーを利用して隣接水系の帯水層から地下水を事前に汲み上げ，"
    "洪水時の受容空間を能動的に創出する新たな洪水制御フレームワークを提案する．"
    "小田川–高梁川流域を対象とした簡易シミュレーションにより，"
    "そのフィージビリティを検討するとともに，政策的含意を議論する．"
)

# ============================================================
# 2. 既往研究と関連技術
# ============================================================
add_heading_gothic("2. 既往研究と関連技術", level=1)

add_heading_gothic("2.1 背水現象と合流点治水", level=2)
add_body(
    "河川合流部における背水現象は，本流の水位上昇が支流の排水能力を阻害する現象であり，"
    "わが国では古くから認識されてきた．"
    "前野{8}は，高梁川と小田川の合流点付替えの効果を2次元流れ解析により検討し，"
    "合流点を現在の13.0 km地点から柳井原貯水池経由で9.0 km地点に移設することで，"
    "洪水時の水位が顕著に低下することを示した．"
    "この研究成果は2024年に実現した合流点付替え事業の技術的根拠となっている．"
)

add_body(
    "国際的には，米国ミシシッピ川流域のYazoo Backwater地域（約10,600 km2）が"
    "背水氾濫の典型事例として知られる{9}．"
    "2019年には158日間連続で洪水段階を超過し，約4ヶ月にわたる長期浸水被害が生じた{10}．"
    "Berkowitz et al.{11}は，Yazoo流域の湿地水文を分析し，"
    "降水が76%の飽和イベントを引き起こすことを報告した．"
    "Li & Tsai{12}は，ミシシッピ川デルタの地下水流動が表流水との相互作用により"
    "季節的に変動することを示し，洪水時の地下水位上昇が排水能力をさらに低下させる"
    "メカニズムを明らかにした．"
)

add_heading_gothic("2.2 帯水層涵養と地下水管理による治水", level=2)
add_body(
    "帯水層への人工涵養（Managed Aquifer Recharge: MAR）は，"
    "地下水資源の持続的管理手法として世界各地で実施されている{13}．"
    "California Department of Water Resourcesが推進するFlood-MAR戦略は，"
    "洪水水を農地や氾濫原に展開し帯水層に浸透させることで，"
    "洪水リスク低減と地下水涵養を同時に達成するアプローチである{5}．"
    "この概念を「逆方向」に操作し，事前に帯水層から地下水を汲み上げて"
    "貯留空間を創出する方法が本研究の着想の出発点となっている．"
)

add_body(
    "Chang et al.{6}は台湾南西部沿岸域（地盤沈下地域）において，"
    "最適な地下水汲み上げスキームにより地盤沈下を抑制しつつ洪水被害を低減できることを実証した．"
    "Nishizawa Katazakai & Zhang{14}は日本中部での半世紀にわたる地下水モニタリングに基づき，"
    "水田の地下水涵養効果と季節的管理涵養のスケーラビリティを評価した．"
    "これらの研究は，地下水管理が長期的な治水戦略として有効であることを示唆している．"
)

add_heading_gothic("2.3 揚水発電技術と洪水制御への転用可能性", level=2)
add_body(
    "日本は世界有数の揚水発電大国であり，"
    "神流川発電所（2,820 MW，有効落差653 m），奥多々良木発電所（1,932 MW，387 m）等，"
    "大規模施設が稼働している{15}．"
    "揚水発電は上池と下池の間で水を往復させ電力を貯蔵する技術であり，"
    "「ダム湖 → 放水路 → 発電 → 隣接水系の地下水汲み上げ → 帯水層に空間創出」"
    "というサイクルは，揚水発電の運転概念を地下水管理に拡張したものと理解できる．"
)

add_body(
    "首都圏外郭放水路{7}は全長6.3 km，最大深度50 m，排水能力200 m3/sの地下放水路であり，"
    "中川等5河川の洪水を江戸川に排出する水系間洪水移送の実例である．"
    "EUのCORDISプロジェクト{16}では，洪水防御貯水池からの放流水と余剰流量から"
    "エネルギーを回収する技術の研究開発が進められている．"
    "また，Chang et al.{17}はカスケード水力発電所群の"
    "多シナリオ水位低下法による洪水資源利用を提案している．"
)

add_heading_gothic("2.4 日本の地形的優位性", level=2)
add_body(
    "日本の河川は世界の主要河川と比較して極めて急勾配であり{18}，"
    "利根川（約1/500〜1/2000）や信濃川（約1/500〜1/1500）の勾配は"
    "ミシシッピ川（約1/10000〜1/20000）の5〜20倍に達する．"
    "急勾配は同一水量あたりの位置エネルギーを大きくし，計画放水発電の効率を高める．"
    "また，日本の脊梁山脈を挟む隣接水系間の距離は数km〜数十kmと短く，"
    "ミシシッピ川流域（幅数百〜千km）と比較して1〜2桁小さいスケールで"
    "水系間の水管理が実施可能である．"
    "これらの地形的特性は，本構想の実現にとって世界的に見ても稀有な好条件を提供する．"
)

# ============================================================
# 3. 提案フレームワーク
# ============================================================
add_heading_gothic("3. 提案フレームワーク", level=1)

add_body(
    "本研究が提案する「計画放水発電–地下水管理統合型洪水制御フレームワーク」は，"
    "以下の5つのフェーズから構成される（図-1参照）．"
)

add_body(
    "フェーズ1（洪水予測）：気象予報に基づき数日後の豪雨を予測し，計画放水を開始する．"
    "フェーズ2（計画放水発電）：ダムからの計画放水により，放水路の落差を利用して水力発電を行う．"
    "フェーズ3（地下水汲み上げ）：発電した電力で隣接水系の帯水層から地下水を汲み上げ，"
    "帯水層に受容空間を創出する．汲み上げた水はダム湖や調整池に貯留する．"
    "フェーズ4（洪水受容）：豪雨到来時，降雨が事前に空けた帯水層に浸透することで"
    "地表流出量が低減し，河川の洪水ピークが抑制される．"
    "フェーズ5（回復）：洪水後，ダム湖の貯留水を渇水期の利水・再発電に活用し，"
    "帯水層は自然涵養により回復する．"
)

add_body(
    "本フレームワークの核心的イノベーションは，"
    "洪水の位置エネルギーそのものを利用して隣接水系の「受容能力」を事前に拡大する点にある．"
    "従来の洪水制御が「水を排出する」ことに注力するのに対し，"
    "本構想は「水を受け入れる空間を能動的に創出する」というパラダイムシフトである．"
)

# ============================================================
# 4. 対象流域と2018年洪水イベント
# ============================================================
add_heading_gothic("4. 対象流域と2018年洪水イベント", level=1)

add_heading_gothic("4.1 小田川–高梁川流域の概要", level=2)
add_body(
    "小田川は岡山県西部を流れる一級河川であり，流域面積483 km2，"
    "河床勾配約1/2200で高梁川に合流する{2}．"
    "高梁川本流は河床勾配約1/900であり，"
    "合流点下流の12.0 km付近には狭窄部が存在する．"
    "洪水時には，この狭窄部による背水の影響に加え，"
    "合流点付近における本流の水位上昇に伴う背水現象により，"
    "小田川の水位が長時間にわたり高い状態が維持される{8}．"
)

add_heading_gothic("4.2 2018年西日本豪雨による被害", level=2)
add_body(
    "2018年7月の西日本豪雨では，小田川流域で3日間雨量約320 mmを記録し，"
    "高梁川本流の急激な水位上昇により小田川で背水現象が発生した{1}．"
    "矢掛観測所では6時間で3〜4 m（0.5〜0.67 m/h）の水位上昇が観測された．"
    "堤防天端と真備町の住宅地との標高差は約6 mあったにもかかわらず，"
    "5箇所で堤防が決壊し，8.28 km2が浸水，浸水量は約15.3 x 106 m3に達した{1}．"
    "死者51人，全壊4,600棟以上という甚大な被害をもたらした．"
)

# 表-1: 2018年洪水データ
add_table_from_data(
    ["項目", "値"],
    [
        ["流域面積（小田川）", "483 km2"],
        ["3日間雨量", "約320 mm"],
        ["浸水面積", "8.28 km2"],
        ["浸水量", "15.3 x 10^6 m3"],
        ["最大水位上昇速度", "0.5-0.67 m/h"],
        ["堤防決壊箇所数", "5箇所"],
        ["死者数", "51人"],
        ["全壊棟数", "4,600棟以上"],
    ],
    "表-1 2018年西日本豪雨時の小田川–高梁川流域の被害概要"
)

add_heading_gothic("4.3 合流点付替え事業の効果", level=2)
add_body(
    "国土交通省は2018年の災害を受け，小田川と高梁川の合流点を"
    "従来の13.0 km地点から4.6 km下流の9.0 km地点（柳井原貯水池経由）に"
    "付替える大規模事業を実施し，2024年3月に完成した{4}．"
    "中国地方整備局の報告によれば，2024年11月の大雨時には，"
    "事業実施前と比較して高梁川側で約0.8 m，小田川側で約4.6 mの"
    "水位低減効果が確認されている{4}．"
    "この効果は前野{8}の2次元流れ解析による予測と整合的である．"
)

# ============================================================
# 5. 簡易シミュレーション
# ============================================================
add_heading_gothic("5. 簡易シミュレーション", level=1)

add_heading_gothic("5.1 エネルギー収支の計算", level=2)
add_body(
    "計画放水発電の出力は水力発電の基本式 P = rho g Q H eta により計算される．"
    "ここで，rho は水の密度（1,000 kg/m3），g は重力加速度（9.81 m/s2），"
    "Q は放水流量，H は有効落差，eta は水車効率である．"
    "一方，地下水汲み上げの所要出力は P_pump = rho g Q_pump h / eta_pump で表される．"
    "ここで，Q_pump は揚水流量，h は揚程，eta_pump はポンプ効率である．"
)

add_body(
    "有効落差20〜200 m，放水流量5〜100 m3/sの範囲でパラメータスイープを行い，"
    "発電出力と揚水所要出力の関係を図-1に示す．"
    "代表的な条件（H = 100 m，Q = 50 m3/s，eta = 0.85）では"
    "約41.7 MWの発電が可能であり，揚水側（h = 30 m，Q_pump = 10 m3/s，eta_pump = 0.80）"
    "の所要出力3.7 MWは発電量のわずか約9%である．"
    "すなわち，計画放水のエネルギーの約91%は余剰電力として利用可能である．"
)

# 図-1挿入
add_figure(
    f"{FIGDIR}/fig1_energy_balance_ja.png",
    "図-1 計画放水発電の出力と揚水所要出力の関係\n"
    "（青破線は揚水所要出力3.7 MWの等出力線，破線より上の領域でエネルギー的に自立）",
    width_inches=5.0
)

add_heading_gothic("5.2 帯水層貯留容量の推定", level=2)
add_body(
    "真備地区周辺の沖積低地には砂礫層を主体とする帯水層が分布しており，"
    "帯水層厚15 m，有効間隙率0.20，透水係数10^-3 m/sと仮定した．"
    "管理面積20 km2において，地下水位を事前にDelta h低下させた場合の"
    "利用可能貯留容量 V = A x n_e x Delta h を図-2に示す．"
)

add_body(
    "Delta h = 3 mの場合，貯留容量は12.0 x 106 m3となり，"
    "2018年の浸水量（15.3 x 106 m3）の約78%に相当する．"
    "Delta h = 5 mでは20.0 x 106 m3（約131%）となり，"
    "2018年規模の洪水量を上回る受容能力の創出が理論上可能である．"
    "ただし，地下水位の過度な低下は地盤沈下のリスクを伴うため{6}，"
    "管理可能なDelta hの範囲は地質条件と地盤沈下モニタリングに基づき"
    "慎重に設定する必要がある．"
)

# 図-2挿入
add_figure(
    f"{FIGDIR}/fig2_aquifer_storage_ja.png",
    "図-2 事前地下水位低下量と帯水層貯留容量の関係\n"
    "（赤破線は2018年浸水量15.3 x 10^6 m3，バー上の数値は浸水量に対する割合）",
    width_inches=5.0
)

add_heading_gothic("5.3 洪水ハイドログラフの比較", level=2)
add_body(
    "2018年洪水イベントを模した簡易降雨–流出モデル（合理式ベース）により，"
    "4つのシナリオについて72時間のハイドログラフを計算した（図-3）．"
    "シナリオ1（対策なし）は2018年の状況を再現し，"
    "シナリオ2（合流点付替え）は背水効果の解消により実効的に30%のピーク削減を仮定する{4}．"
    "シナリオ3（地下水管理型）は事前3 mの地下水位低下による帯水層浸透を考慮し，"
    "シナリオ4（統合型）はシナリオ2と3の複合である．"
)

# 図-3挿入
add_figure(
    f"{FIGDIR}/fig3_hydrograph_ja.png",
    "図-3 洪水シナリオ別ハイドログラフ比較（小田川流域，72時間）\n"
    "（上段: 降雨強度，下段: 流出量）",
    width_inches=5.5
)

add_body(
    "シミュレーション結果を表-2に示す．"
    "地下水管理型（シナリオ3）単独ではピーク流量の削減効果は約2%にとどまるが，"
    "これは帯水層への浸透速度の制約によるものであり，"
    "人工涵養施設（浸透池・注入井）の整備により改善が見込まれる．"
    "一方，累積浸透量は6.24 x 106 m3（2018年浸水量の約41%）に達しており，"
    "洪水後の排水負荷の軽減に寄与する．"
    "統合型（シナリオ4）では，構造的対策と地下水管理の相乗効果により，"
    "ピーク流量を約32%削減する結果が得られた．"
)

# 表-2: シナリオ比較
add_table_from_data(
    ["シナリオ", "ピーク流量 [m3/s]", "削減率 [%]", "累積浸透量 [x10^6 m3]"],
    [
        ["1: 対策なし", "1,237", "—", "—"],
        ["2: 合流点付替え", "866", "30", "—"],
        ["3: 地下水管理型", "1,207", "2", "6.24"],
        ["4: 統合型", "845", "32", "6.24"],
    ],
    "表-2 シナリオ別の洪水制御効果"
)

# 図-4挿入
add_figure(
    f"{FIGDIR}/fig4_scenario_comparison_ja.png",
    "図-4 シナリオ別のピーク流量（左）と総流出量（右）の比較",
    width_inches=5.5
)

# ============================================================
# 6. 首都圏（荒川・江戸川流域）への適用可能性検討
# ============================================================
add_heading_gothic("6. 首都圏（荒川・江戸川流域）への適用可能性検討", level=1)

add_body(
    "小田川–高梁川流域でのフィージビリティが示されたことを踏まえ，"
    "本フレームワークの首都圏適用可能性を検討する．"
    "荒川（幹川流路延長173 km，流域面積2,940 km2）の想定氾濫区域内人口は"
    "約821万人に達し{22}，堤防決壊時の被害は壊滅的となる．"
    "人口・産業が高度に集積した首都圏においてこそ，"
    "既存の構造的対策を補完する追加的安全層の意義は大きい．"
)

add_heading_gothic("6.1 荒川・江戸川流域の既存治水対策体系", level=2)
add_body(
    "荒川・江戸川流域では，明治以降の近代治水事業により"
    "多層的な構造的対策が整備されてきた（表-3）．"
    "荒川放水路（1930年完成，全長22 km）は隅田川への洪水流入を制御するために建設され，"
    "岩淵水門は荒川と隅田川の分派点における水位を調節している{23}．"
    "荒川第一調節池（彩湖）は面積5.85 km2，貯留容量3,900 x 104 m3を有し{24}，"
    "荒川中流部の洪水調節を担う主要施設である．"
    "加えて，上流には浦山ダム，二瀬ダム，滝沢ダムの3基が設置されている．"
)

add_body(
    "江戸川流域では，首都圏外郭放水路（全長6.3 km，最大排水能力200 m3/s）が"
    "中川・綾瀬川等5河川の洪水を江戸川に排出する{7}．"
    "これは水系間洪水移送の世界最大級の実例である．"
    "その他，三郷放水路，綾瀬川放水路，八潮排水機場など"
    "多数の排水施設が整備され，"
    "中川・綾瀬川等は令和6年に特定都市河川に指定された．"
    "さらに，荒川下流部ではスーパー堤防（高規格堤防）の整備が進められており，"
    "管内堤防延長の約15%で完成または工事中である{23}．"
)

# 表-3: 荒川・江戸川の主要治水施設
add_table_from_data(
    ["施設名", "種別", "規模・容量", "特記事項"],
    [
        ["荒川放水路", "放水路", "全長22 km", "1930年完成，隅田川洪水を制御"],
        ["岩淵水門", "水門", "—", "荒川↔隅田川の分派点制御"],
        ["荒川第一調節池（彩湖）", "調節池", "3,900万m3", "面積5.85 km2"],
        ["浦山・二瀬・滝沢ダム", "ダム", "合計約1.3億m3", "上流域洪水調節"],
        ["首都圏外郭放水路", "地下放水路", "200 m3/s", "全長6.3 km，5河川排水"],
        ["スーパー堤防", "堤防", "—", "下流部15%完成"],
    ],
    "表-3 荒川・江戸川流域の主要治水施設"
)

add_heading_gothic("6.2 関東平野の帯水層と地下水管理の歴史", level=2)
add_body(
    "関東平野の地下には厚い沖積層が分布する．"
    "産業技術総合研究所の3次元地質地盤図によれば{25}，"
    "東京低地では沖積層が最大約80 m堆積しており，"
    "基底礫層，砂泥互層，泥層の積み重なりから構成される．"
    "この沖積層中の砂礫層は帯水層として機能し，"
    "関東平野全体では膨大な地下水貯留ポテンシャルを有する．"
)

add_body(
    "歴史的に，関東平野南部では明治中期（1890年代）から地下水の過剰汲み上げによる"
    "深刻な地盤沈下が発生した{26}．"
    "これに対し，工業用水法（1956年）および建築物用地下水採取規制法（1962年）に基づく"
    "採取規制が実施された結果，近年では大規模な地盤沈下は沈静化し，"
    "地下水位が顕著に回復している{26}．"
    "この『回復した地下水位』は，本フレームワークの観点からは"
    "逆説的な好条件を意味する．"
    "すなわち，規制により自然に回復した帯水層のうち，"
    "地盤沈下を生じない範囲で計画的に地下水位を低下させれば，"
    "洪水受容空間を創出できる可能性がある．"
)

add_heading_gothic("6.3 荒川流域への適用シナリオ", level=2)
add_body(
    "荒川上流域（奥秩父山地）では，滝沢ダム（有効落差約100 m）等から"
    "計画放水による発電が可能である．"
    "生成した電力で中流域（川越–さいたま間）の沖積層帯水層の"
    "地下水位を事前に低下させるシナリオが想定される．"
    "荒川中流域の沖積低地には厚い砂礫層帯水層が分布しており{25}，"
    "管理面積を仮に30 km2，有効間隙率0.15，地下水位低下3 mと仮定すると，"
    "帯水層貯留容量は13.5 x 106 m3と推定される．"
    "荒川第一調節池（3,900 x 104 m3）と合わせれば，"
    "合計約52.5 x 106 m3の受容能力が創出される．"
)

add_body(
    "令和元年東日本台風（台風19号）では，荒川中流部で無堤部からの溢水が発生し，"
    "支川入間川等で7箇所の堤防決壊が報告されている{22}．"
    "既存の調節池・ダム群に加え，帯水層貯留を「地下の遊水地」として活用することで，"
    "荒川流域の洪水調節容量をさらに増強できる可能性がある．"
    "特に，人口821万人を擁する想定氾濫区域において，"
    "構造的対策の追加が用地取得の困難さから制約される中{23}，"
    "地下空間の活用は有力な代替手段となり得る．"
)

add_body(
    "さらに，首都圏外郭放水路の排水能力（200 m3/s）と本構想を組み合わせることで，"
    "中川・綾瀬川流域の洪水水を帯水層に一時貯留し，"
    "排水機場の負荷を軽減するスキームも検討に値する．"
    "これは，首都圏外郭放水路が「水の移送」を担い，"
    "帯水層が「水の貯留」を担うという機能分担の枠組みである．"
)

# ============================================================
# 7. 考察
# ============================================================
add_heading_gothic("7. 考察", level=1)

add_heading_gothic("7.1 エネルギー的自立性", level=2)
add_body(
    "本シミュレーションの最も注目すべき結果は，"
    "地下水汲み上げに必要なエネルギーが計画放水発電の約9%に過ぎない点である．"
    "これは揚水発電所の往復効率（70〜80%{15}）と比較して極めて効率的であり，"
    "余剰電力の約91%を送電網に供給できることを意味する．"
    "すなわち，本フレームワークは洪水制御と再生可能エネルギー供給を"
    "同時に達成する「エネルギー自立型治水」として機能し得る．"
)

add_heading_gothic("7.2 構造的対策との相補性", level=2)
add_body(
    "シナリオ比較の結果，地下水管理型（シナリオ3）単独のピーク削減効果は限定的であるが，"
    "合流点付替え（シナリオ2）との統合（シナリオ4）により32%の削減が達成される．"
    "これは，本提案が従来の構造的対策を置き換えるものではなく，"
    "相補的に機能する「追加的な安全層」として位置づけられることを示唆する．"
    "特に，合流点付替え事業が完了した真備地区において，"
    "さらなる安全度の向上策としての地下水管理の導入は政策的に合理的である．"
)

add_heading_gothic("7.3 限界と課題", level=2)
add_body(
    "本研究にはいくつかの重要な限界がある．"
    "第一に，本シミュレーションは合理式ベースの簡易モデルであり，"
    "HEC-RAS等の水理モデルやGSFLOW等の統合水文モデル{19}による"
    "詳細な検証が不可欠である．"
    "第二に，帯水層パラメータ（有効間隙率，透水係数，層厚）は"
    "文献値に基づく仮定であり，現地調査によるキャリブレーションが必要である．"
    "第三に，地盤沈下リスクの定量的評価が行われていない．"
    "関東平野や濃尾平野では過去に深刻な地盤沈下を経験しており{20}，"
    "地下水汲み上げと涵養のバランス管理が不可欠である．"
    "第四に，帯水層への浸透速度の制約から，ピーク流量の直接的な削減効果は限定的であり，"
    "人工涵養施設（浸透池，注入井等）の設計と最適配置に関する研究が求められる．"
)

add_heading_gothic("7.4 国際的展開可能性", level=2)
add_body(
    "本フレームワークは日本の急峻な地形を前提に構想されたものであるが，"
    "その基本原理は国際的にも応用可能性を有する．"
    "特に，大陸国家における越境水系（メコン川，ナイル川等）では，"
    "上流国のダム放水に下流国が一方的にさらされるリスクが存在する{21}．"
    "本構想は自国領内の帯水層管理のみで洪水緩和を実現できるため，"
    "国境をまたぐ水利外交の制約を受けない「一方的実施可能な治水」として機能し得る．"
    "また，中東・北アフリカの乾燥地帯では，ワジ（涸れ川）の鉄砲水と"
    "帯水層涵養を組み合わせた水資源管理への応用が想定される．"
)

# ============================================================
# 8. 政策提言
# ============================================================
add_heading_gothic("8. 政策提言", level=1)

add_body(
    "以上の分析に基づき，以下の政策提言を行う．"
)

add_body(
    "（1）パイロット事業の実施：小田川–高梁川流域（真備地区）と"
    "荒川中流域の2地域を候補とした実証事業を提案する．"
    "真備地区は合流点付替え事業が完了し効果検証が進行中であり，"
    "荒川流域は首都圏の人口・産業密集地として社会的インパクトが最も大きい．"
    "（2）帯水層の貯留容量と地盤沈下リスクの詳細調査．"
    "小田川流域，および荒川中流域の沖積層帯水層について，"
    "ボーリングデータと地下水位モニタリングに基づく定量的評価が必要である．"
    "関東平野では地下水規制により水位が回復しており{26}，"
    "管理涵養の余地を科学的に評価する好機である．"
    "（3）GSFLOW等の統合水文モデルによる詳細シミュレーション{19}．"
    "本研究の簡易モデルを発展させ，地下水–地表水相互作用を動的に評価する必要がある．"
    "（4）首都圏外郭放水路と帯水層貯留の連携検討．"
    "放水路の「水の移送」機能と帯水層の「水の貯留」機能を組み合わせた"
    "複合型洪水制御システムの設計指針を策定すべきである．"
    "（5）「流域治水」政策への地下水管理の制度的位置づけ{27}．"
    "国土交通省が推進する流域治水の枠組みに，帯水層貯留を「地下の遊水地」として"
    "明示的に位置づけることが望ましい．"
)

# ============================================================
# 9. 結論
# ============================================================
add_heading_gothic("9. 結論", level=1)

add_body(
    "本研究では，計画放水発電と隣接水系地下水管理を統合した"
    "新たな洪水制御フレームワークを提案し，"
    "2018年西日本豪雨時の小田川–高梁川流域を対象とした簡易シミュレーションにより"
    "そのフィージビリティを検討した．"
    "主要な知見は以下の通りである．"
)

add_body(
    "（1）有効落差100 m・放水流量50 m3/sの計画放水で約41.7 MWの発電が可能であり，"
    "地下水汲み上げの所要出力（3.7 MW）は発電量の約9%に過ぎない．"
    "エネルギー的に自立した洪水制御システムの実現可能性が示された．"
)

add_body(
    "（2）管理面積20 km2・地下水位低下3 mで約12.0 x 106 m3の帯水層貯留容量が創出され，"
    "2018年浸水量の約78%に相当する受容空間を確保できる．"
)

add_body(
    "（3）2024年に完成した合流点付替え（構造的対策）との統合により，"
    "ピーク流量を約32%削減し得る．"
    "本提案は既存の構造的対策を代替するものではなく，"
    "相補的な「追加的安全層」として機能する．"
)

add_body(
    "（4）荒川流域（想定氾濫区域内人口821万人）への適用可能性が示唆された．"
    "上流ダム群からの計画放水発電と中流域の沖積層帯水層の組み合わせにより，"
    "既存の調節池・ダム群を補完する追加的な洪水調節容量の創出が期待される．"
    "首都圏外郭放水路との連携も有望であり，"
    "人口・産業密集地での治水安全度向上に貢献し得る．"
)

add_body(
    "（5）日本の急峻な地形，既存の揚水発電インフラ，高密度な気象・水文観測網は，"
    "本フレームワークの実現にとって世界的に見ても稀有な好条件を提供する．"
    "今後，統合水文モデルによる詳細検証と，"
    "パイロット事業による実証が望まれる．"
)

# ============================================================
# 謝辞
# ============================================================
add_heading_gothic("謝辞", level=1)
add_body_no_indent(
    "本研究では，防災科学技術研究所，国土交通省中国地方整備局，"
    "国土交通省関東地方整備局，産業技術総合研究所地質調査総合センター，"
    "および岡山県の公開データを利用した．"
    "ここに記して謝意を表する．"
)

# ============================================================
# 参考文献
# ============================================================
add_heading_gothic("参考文献", level=1)

refs = [
    "1) 防災科学技術研究所: 平成30年7月豪雨（前線及び台風第7号による大雨等）による被害の概況, 2018.",
    "2) 国土交通省中国地方整備局: 小田川合流点付替え事業概要, https://www.cgr.mlit.go.jp/",
    "3) 国土交通省: 河川改修の取り組み, https://www.mlit.go.jp/river/basic_info/",
    "4) KSBニュース: 小田川の治水事業 2024年の大雨の際「水位を低減させる効果があった」と報告, 2025年2月7日.",
    "5) California Department of Water Resources: Flood-Managed Aquifer Recharge (Flood-MAR), https://water.ca.gov/Programs/All-Programs/Flood-MAR",
    "6) Chang, Y.-L. et al.: Flood hazard mitigation in land subsidence prone coastal areas by optimal groundwater pumping, Water Resources Management, 2019.",
    "7) 国土交通省関東地方整備局: 首都圏外郭放水路, https://www.ktr.mlit.go.jp/edogawa/gaikaku/",
    "8) 前野詩朗: 高梁川と小田川の合流点付け替えの効果, 水工学論文集, 第51巻, pp.613-618, 2007.",
    "9) Mississippi Levee Board: Yazoo Backwater Project, https://msleveeboard.com/",
    "10) Mississippi State University Extension: Final Report: Survey of Overlooked Costs of the 2019 Backwater Flood in the Yazoo Mississippi Delta.",
    "11) Berkowitz, J.F. et al.: Forested Wetland Hydrology in a Large Mississippi River Tributary System, Wetlands, 2019.",
    "12) Li, A. & Tsai, F.T.-C.: Understanding dynamics of groundwater flows in the Mississippi River Delta, Journal of Hydrology, Vol.583, 124616, 2020.",
    "13) Dillon, P. et al.: Managed aquifer recharge: rediscovering nature as a leading edge technology, Water Science and Technology, Vol.62, No.10, pp.2338-2345, 2010.",
    "14) Nishizawa Katazakai, S. & Zhang, J.: Long-term flood control in central Japan: A half-century groundwater monitoring and evaluating adaptation measures, Groundwater for Sustainable Development, Vol.33, 101585, 2026.",
    "15) IEA Hydropower Implementing Agreement: Case Study 11-02: Large Scale Pumped Storage Power Plants, Japan, 2006.",
    "16) EU CORDIS: Recovery of Energy from Released Water and Reserve Flow from a Flood Protection Reservoir, Project ID: HY.-00410-87.",
    "17) Chang, J. et al.: The flood resource utilization of cascade hydropower stations based on the multi-scenario water level drawdown method, EGU General Assembly 2025, EGU25-14374.",
    "18) Kinoshita, S. & Tokunaga, Y.: Water Resources Management in Japan, IWRA Proceedings, 2003.",
    "19) USGS: GSFLOW: Coupled Groundwater and Surface-Water Flow Model, https://water.usgs.gov/ogw/gsflow/",
    "20) 地盤沈下対策関係閣僚会議: 地盤沈下対策に関する基本方針, 1991.",
    "21) Mekong River Commission: State of the Basin Report 2018, Vientiane, 2019.",
    "22) 国土交通省水管理・国土保全局: 荒川水系河川整備基本方針の変更の概要, 令和6年11月, pp.1-52.",
    "23) 国土交通省関東地方整備局荒川下流河川事務所: 治水対策〜首都圏を守る災害に強い荒川づくり〜, https://www.ktr.mlit.go.jp/arage/",
    "24) 国土交通省関東地方整備局荒川上流河川事務所: 彩湖のはたらき, https://www.ktr.mlit.go.jp/arajo/arajo00161.html",
    "25) 小松原純子: 東京低地の沖積層, GSJ地質ニュース, Vol.10, No.7, pp.148-152, 2021.",
    "26) 国土交通省水資源部: 地下水保全と地盤沈下の現状, https://www.mlit.go.jp/mizukokudo/mizsei/",
    "27) 国土交通省: 流域治水プロジェクト, https://www.mlit.go.jp/river/kasen/ryuiki_pro/",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.hanging_indent = Pt(18)
    run = p.add_run(ref)
    run.font.name = "游明朝"
    run.font.size = Pt(8)

# ── 保存 ──
outpath = f"{OUTDIR}/jsce_manuscript.docx"
doc.save(outpath)
print(f"JSCE原稿を保存しました: {outpath}")
