#!/usr/bin/env python3
"""
JSCE投稿用タイトルページ生成スクリプト
Guidelines for Authors (2026年1月1日改訂) のサンプルファイルに準拠

提出書類:
  - Title Page（本スクリプト）: Article Type, 著者情報, COI, Funding, Author Contributions
  - Main Text（別ファイル）: Abstract, Keywords, 本文, References, Figure legends, Tables
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"

# サンプルファイルをベースにして作成
doc = Document()

# ── ページ設定（A4）──
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# ── スタイル ──
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.line_spacing = Pt(22)
style.paragraph_format.space_after = Pt(0)


def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = True
    return p


def add_normal(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_placeholder(text):
    return add_normal(text, color=RGBColor(255, 0, 0))


# ============================================================
# Article Type
# ============================================================
add_heading("Article Type")
add_normal("Academic paper")

# ============================================================
# Title
# ============================================================
add_heading("TITLE")
add_normal(
    "A novel flood control framework integrating planned-release hydropower "
    "with inter-watershed groundwater management: "
    "A simplified feasibility analysis for the Oda River–Takahashi River basin",
    bold=True
)
p = doc.add_paragraph()
run = p.add_run(
    "（和文題目）計画放水発電と隣接水系地下水管理を統合した"
    "新たな洪水制御フレームワークの提案"
    "　― 小田川–高梁川流域を対象とした簡易フィージビリティ分析 ―"
)
run.font.name = "Times New Roman"
run.font.size = Pt(10)
run.italic = True

# ============================================================
# Authors
# ============================================================
add_heading("AUTHORS")
add_placeholder(
    "（著者名）1"
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
add_placeholder(
    "1（JSCE会員種別）, （職位）, （所属）\n"
    "  （住所）\n"
    "  E-mail: （メールアドレス） (Corresponding Author)"
)

# ============================================================
# Conflicts of Interest
# ============================================================
add_heading("CONFLICTS OF INTEREST")
add_normal("The authors declare that there are no conflicts of interest.")

# ============================================================
# Sources of Funding
# ============================================================
add_heading("SOURCES OF FUNDING")
add_normal(
    "This research received no specific grant from any funding agency "
    "in the public, commercial or not-for-profit sectors."
)

# ============================================================
# Author Contributions
# ============================================================
add_heading("AUTHOR CONTRIBUTIONS")
add_placeholder(
    "（著者名） was responsible for the overall study concept, "
    "simulation design, data analysis, and manuscript preparation."
)

# ============================================================
# Related Publications Disclosure
# ============================================================
add_heading("DISCLOSURE OF RELATED PUBLICATIONS")
add_normal(
    "A related English-language manuscript (Commentary format) is being prepared "
    "for submission to Nature Water. The English manuscript focuses on international "
    "water governance and transboundary river management policy implications, "
    "and does not contain the detailed quantitative analysis of the Oda River–Takahashi River "
    "and Arakawa River basins presented in this paper. "
    "The two manuscripts target different readerships and have distinct content structures "
    "with no substantive overlap."
)

# ── 保存 ──
outpath = f"{OUTDIR}/jsce_title_page.docx"
doc.save(outpath)
print(f"タイトルページを保存しました: {outpath}")
