#!/usr/bin/env python3
"""
Nature Water Commentary 原稿生成スクリプト

Title: Harnessing flood energy to create subsurface storage:
       A hydropower–groundwater framework for integrated flood control

Format: Nature Water Commentary / Perspective
- ~2,000–3,000 words body text
- Up to 4 figures/tables
- Up to 30 references
- Figures submitted as separate files
"""

import os
import re
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"
FIGDIR = f"{OUTDIR}/figures"

doc = Document()

# ── Page setup (A4) ──
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# ── Styles ──
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = Pt(24)  # double spacing
style_normal.paragraph_format.space_after = Pt(0)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def add_body(text):
    """Body paragraph with superscript refs {N}."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(36)
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith("{") and part.endswith("}"):
            ref_text = part[1:-1]
            run = p.add_run(ref_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.superscript = True
        else:
            run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    return p


def add_body_no_indent(text):
    p = add_body(text)
    p.paragraph_format.first_line_indent = Pt(0)
    return p


# ============================================================
# Title Page
# ============================================================
add_title("Harnessing flood energy to create subsurface storage:")
add_title("A hydropower-groundwater framework for integrated flood control")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
run = p.add_run("[Author Names]")
run.font.size = Pt(12)
run.bold = True

p = doc.add_paragraph()
run = p.add_run("[Affiliations]")
run.font.size = Pt(11)
run.italic = True

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Correspondence: [email]")
run.font.size = Pt(11)

# ── Abstract ──
add_heading("Abstract", level=1)
add_body_no_indent(
    "Conventional flood control relies on structural measures -- levees, dams, and floodways -- "
    "that expel excess water. We propose a paradigm shift: using the potential energy of planned "
    "dam releases to generate electricity that powers groundwater pumping in adjacent watersheds, "
    "proactively creating subsurface reception capacity before flood events. A feasibility analysis "
    "of the Oda-Takahashi River basin in Japan, site of catastrophic backwater flooding in 2018, "
    "demonstrates that hydropower from a 100 m head release (41.7 MW) can power groundwater "
    "pumping (3.7 MW, 9% of generation) to create aquifer storage absorbing 78% of observed "
    "flood volume. This framework exploits Japan's steep topography and established pumped-storage "
    "infrastructure, but its principles extend to transboundary river basins where unilateral "
    "flood mitigation within national borders is geopolitically valuable, and to arid regions "
    "where coupling flash-flood energy with aquifer recharge addresses both flood risk and "
    "water scarcity."
)

# ============================================================
# Main Text
# ============================================================
add_heading("The backwater problem", level=1)

add_body(
    "When a mainstem river swells, it can block tributary discharge -- a phenomenon known as the "
    "backwater effect. This mechanism caused catastrophic flooding in Mabi, Okayama Prefecture, Japan, "
    "during the July 2018 Western Japan Heavy Rains: the Takahashi River's water level prevented "
    "the Oda River from draining, leading to levee breaches at five locations, inundating 8.28 km2 "
    "and claiming 51 lives{1}. Half a world away, the Yazoo River basin in Mississippi, USA, "
    "experienced 158 consecutive days above flood stage in 2019 as the Mississippi River's high "
    "water prevented tributary outflow{2}. These events share a common mechanism: the failure of "
    "one watershed to discharge into another."
)

add_body(
    "Current responses to backwater flooding are predominantly structural. Japan relocated the "
    "Oda-Takahashi confluence 4.6 km downstream (completed March 2024), achieving a 4.6 m water "
    "level reduction on the Oda River side{3}. The US has debated a pump station for the Yazoo "
    "Backwater area since 1941, stalled by environmental concerns{2}. Both approaches -- river "
    "engineering and mechanical pumping -- focus on moving water away. We propose an alternative: "
    "creating space to receive it."
)

add_heading("A new framework: flood energy for subsurface storage", level=1)

add_body(
    "Our proposed framework operates in five phases. First, weather forecasts trigger planned "
    "dam releases days before a flood event. Second, the release generates hydropower using "
    "the elevation drop -- in Japan's steep rivers, effective heads of 50-200 m are common, "
    "yielding 5-170 MW at typical flow rates{4}. Third, this electricity powers groundwater "
    "pumping in adjacent watersheds, drawing down the water table to create void space in "
    "aquifers. Fourth, when heavy rainfall arrives, a portion infiltrates into the pre-emptied "
    "aquifer rather than becoming surface runoff, reducing flood peaks. Fifth, after the event, "
    "stored water serves irrigation or re-generation during dry seasons."
)

add_body(
    "The energy arithmetic is compelling. For a representative planned release (100 m head, "
    "50 m3/s, 85% turbine efficiency), hydropower output is 41.7 MW. Groundwater pumping from "
    "30 m depth at 10 m3/s with 80% pump efficiency requires only 3.7 MW -- roughly 9% of "
    "generation (Fig. 1). The remaining 91% is surplus electricity for the grid. Unlike "
    "pumped-storage hydropower, which cycles water between two reservoirs with 70-80% round-trip "
    "efficiency{4}, this framework converts flood energy into a flood-mitigation asset with "
    "minimal energy cost."
)

add_heading("Feasibility: the Oda-Takahashi case", level=1)

add_body(
    "We tested this concept against the 2018 Mabi flooding. The Oda River basin (483 km2) "
    "received approximately 320 mm of rainfall over three days, generating an estimated "
    "15.3 x 10^6 m3 of floodwater in the inundation zone{1}. The surrounding alluvial plain "
    "contains sand-gravel aquifers (estimated thickness 15 m, effective porosity 0.20). "
    "Pre-emptying 3 m of water table over a 20 km2 managed area would create 12.0 x 10^6 m3 "
    "of subsurface storage -- 78% of the observed flood volume (Fig. 2). At 5 m drawdown, "
    "storage exceeds the flood volume entirely."
)

add_body(
    "A simplified rainfall-runoff simulation comparing four scenarios (Fig. 3) reveals that "
    "groundwater management alone achieves modest peak reduction (2%), limited by infiltration "
    "rates. However, combined with the completed confluence relocation, the integrated approach "
    "reduces peak discharge by 32% and absorbs 41% of total runoff volume into the aquifer "
    "(Fig. 4). The groundwater component thus functions not as a replacement for structural "
    "measures but as a complementary safety layer."
)

add_heading("Scaling up: Tokyo's Arakawa-Edogawa basin", level=1)

add_body(
    "If the Oda-Takahashi case demonstrates feasibility in a mid-sized basin, the true policy "
    "significance lies in densely populated metropolitan areas. The Arakawa River basin "
    "(drainage area 2,940 km2, main channel length 173 km) protects approximately 8.2 million "
    "people within its projected inundation zone{13}. A levee breach would be catastrophic "
    "for Japan's political and economic core."
)

add_body(
    "Tokyo's flood defence already involves a remarkable multi-layered system: three upstream "
    "dams (Urayama, Futase, Takizawa), the Arakawa First Retarding Basin (Saiko, capacity "
    "39 x 10^6 m3){14}, the Iwabuchi Sluice Gate controlling flow between the Arakawa and "
    "Sumida rivers, super levees (15% of lower reach completed), and the Metropolitan Area "
    "Outer Underground Discharge Channel (6.3 km, 200 m3/s) serving the adjacent Edogawa "
    "basin{15}. Despite this infrastructure, Typhoon Hagibis (2019) caused overbank flooding "
    "and seven levee breaches on Arakawa tributaries{13}, exposing residual vulnerability."
)

add_body(
    "The Kanto Plain beneath Tokyo holds alluvial deposits up to 80 m thick{16}, with "
    "sand-gravel aquifers of substantial storage potential. Critically, decades of strict "
    "groundwater extraction regulations (Industrial Water Law 1956, Building Water Law 1962) "
    "have allowed water tables to recover significantly from mid-20th-century over-extraction "
    "that caused severe land subsidence{17}. This recovered water table, paradoxically, "
    "represents an opportunity: carefully managed drawdown within subsidence-safe limits could "
    "create 'underground retarding basins' complementing the existing surface facilities. "
    "A conservative estimate -- 30 km2 managed area, effective porosity 0.15, 3 m drawdown -- "
    "yields 13.5 x 10^6 m3 of additional storage, augmenting the Saiko's 39 x 10^6 m3 "
    "by over one-third."
)

add_body(
    "Moreover, coupling aquifer storage with the Metropolitan Outer Discharge Channel -- "
    "where the channel provides 'water transfer' and the aquifer provides 'water detention' -- "
    "offers a functional division that could substantially reduce pump station loads during "
    "extreme events. In land-constrained megacities where surface expansion of retarding "
    "basins is prohibitively expensive, subsurface storage exploits an otherwise unused "
    "dimension of urban infrastructure."
)

add_heading("Why Japan -- and why beyond Japan", level=1)

add_body(
    "Japan's geography uniquely suits this framework. Its rivers have gradients 5-20 times "
    "steeper than the Mississippi{5}, maximizing hydropower per unit flow. Adjacent watersheds "
    "are separated by only 5-30 km (versus hundreds of kilometres in continental basins), "
    "making inter-basin water management feasible at practical scales. Japan also operates "
    "the world's largest pumped-storage fleet, including the 2,820 MW Kannagawa facility{4}, "
    "providing both technical precedent and operational expertise directly transferable to the "
    "proposed framework."
)

add_body(
    "Yet the principles are not Japan-specific. Three international contexts merit attention. "
    "First, transboundary rivers: in basins like the Mekong, Nile, and Jordan, downstream "
    "nations face flood risk from upstream dam operations over which they have limited control{6}. "
    "A groundwater management approach allows flood mitigation using only domestic aquifers, "
    "circumventing the geopolitical constraints of water-sharing agreements. This represents "
    "a form of 'unilaterally implementable flood control' -- a concept with significant "
    "diplomatic value."
)

add_body(
    "Second, river deltas: the Mekong Delta, Nile Delta, and Ganges-Brahmaputra-Meghna Delta "
    "face compounding threats from upstream flooding and sea-level rise{7}. Pre-flood aquifer "
    "drawdown could address both by creating storage for upstream flood pulses while managing "
    "the freshwater-saltwater interface."
)

add_body(
    "Third, arid regions: in the Middle East and North Africa, flash floods in wadis "
    "(ephemeral streams) represent both a hazard and an untapped water resource. Countries "
    "with sovereign wealth (Saudi Arabia's NEOM project, UAE's food security programmes) "
    "could deploy this framework to capture flash-flood energy for aquifer recharge, "
    "simultaneously mitigating flood damage and augmenting scarce groundwater supplies. "
    "California's Flood-MAR programme{8} already demonstrates the reverse direction "
    "(surface flood to aquifer), validating the hydrogeological feasibility."
)

add_body(
    "Where rivers form international borders, the framework offers additional value. "
    "The nation experiencing backwater flooding can manage its own alluvial aquifer "
    "without requiring upstream cooperation, diplomatic negotiation, or joint infrastructure "
    "investment -- a critical advantage in politically fragile basins."
)

add_heading("Challenges and next steps", level=1)

add_body(
    "Several challenges require resolution. Land subsidence from excessive drawdown is a "
    "well-documented risk{9}; Japan's Kanto Plain experienced severe subsidence from "
    "industrial groundwater extraction in the mid-20th century, leading to strict regulations "
    "that any new framework must navigate. Infiltration rates during intense rainfall events "
    "may limit real-time peak reduction, necessitating engineered recharge facilities "
    "(infiltration basins, injection wells) rather than relying on natural percolation alone."
)

add_body(
    "Detailed validation using coupled surface-subsurface models (GSFLOW{10}, ParFlow{11}) "
    "is essential to replace the simplified analysis presented here. Two pilot basins are "
    "proposed: Oda-Takahashi (Mabi district, where confluence relocation is complete and "
    "effect verification is ongoing) and the Arakawa middle reach (where metropolitan-scale "
    "social impact maximizes the value of incremental safety improvements). "
    "Institutional frameworks for 'subsurface flood detention' -- the underground "
    "equivalent of surface retarding basins -- need development within Japan's 'Watershed-based "
    "Flood Management' policy{12} and analogous frameworks internationally."
)

add_body(
    "The integration of flood control with renewable energy generation, groundwater management, "
    "and drought resilience represents a systems-level innovation. By viewing floods not merely "
    "as hazards to be repelled but as energy resources to be harvested and storage opportunities "
    "to be created, we open a new design space for adaptive water management in an era of "
    "climate uncertainty."
)

# ============================================================
# Competing Interests
# ============================================================
add_heading("Competing Interests", level=2)
add_body_no_indent(
    "The authors declare no competing interests. A related Japanese-language manuscript "
    "on the same conceptual framework, focusing on the Oda-Takahashi case study, has been "
    "submitted to the Journal of JSCE (Special Issue: Hydraulic Engineering)."
)

# ============================================================
# References
# ============================================================
add_heading("References", level=1)

refs = [
    "1. National Research Institute for Earth Science and Disaster Resilience (NIED). Report on July 2018 Western Japan Heavy Rains. (2018).",
    "2. Mississippi State University Extension. Final Report: Survey of Overlooked Costs of the 2019 Backwater Flood in the Yazoo Mississippi Delta. (2020).",
    "3. KSB News. Oda River flood control project effectiveness confirmed during November 2024 heavy rain. 7 February (2025).",
    "4. IEA Hydropower Implementing Agreement. Case Study 11-02: Large Scale Pumped Storage Power Plants, Japan. (2006).",
    "5. Kinoshita, S. & Tokunaga, Y. Water Resources Management in Japan. IWRA Proceedings (2003).",
    "6. Mekong River Commission. State of the Basin Report 2018. (Vientiane, 2019).",
    "7. Minderhoud, P. S. J. et al. Mekong delta much lower than previously assumed in sea-level rise impact assessments. Nat. Commun. 10, 3847 (2019).",
    "8. California Department of Water Resources. Flood-Managed Aquifer Recharge (Flood-MAR). https://water.ca.gov/Programs/All-Programs/Flood-MAR",
    "9. Chang, Y.-L. et al. Flood hazard mitigation in land subsidence prone coastal areas by optimal groundwater pumping. Water Resour. Manag. (2019).",
    "10. USGS. GSFLOW: Coupled Groundwater and Surface-Water Flow Model. https://water.usgs.gov/ogw/gsflow/",
    "11. Maxwell, R. M. et al. A high-resolution simulation of groundwater and surface water over most of the continental US with ParFlow v3. Geosci. Model Dev. 8, 923-937 (2015).",
    "12. Ministry of Land, Infrastructure, Transport and Tourism (MLIT). Watershed-based Flood Management Project. https://www.mlit.go.jp/river/kasen/ryuiki_pro/ (in Japanese).",
    "13. MLIT Water Management and National Land Conservation Bureau. Overview of the Revision of the Arakawa River Basin Flood Control Basic Policy. (November 2024).",
    "14. MLIT Kanto Regional Development Bureau. Saiko (Arakawa First Retarding Basin). https://www.ktr.mlit.go.jp/arajo/arajo00161.html",
    "15. MLIT Kanto Regional Development Bureau. Metropolitan Area Outer Underground Discharge Channel. https://www.ktr.mlit.go.jp/edogawa/gaikaku/",
    "16. Komatsubara, J. Alluvial deposits of the Tokyo Lowland. GSJ Geol. News 10(7), 148-152 (2021).",
    "17. MLIT Water Resources Department. Groundwater Conservation and Land Subsidence Status. https://www.mlit.go.jp/mizukokudo/mizsei/",
    "18. Maeno, S. Effect of relocation of confluence site of the Takahashi River and the Oda River. J. Hydraul. Eng. JSCE 51, 613-618 (2007).",
    "19. Berkowitz, J. F. et al. Forested wetland hydrology in a large Mississippi River tributary system. Wetlands (2019).",
    "20. Li, A. & Tsai, F.T.-C. Understanding dynamics of groundwater flows in the Mississippi River Delta. J. Hydrol. 583, 124616 (2020).",
    "21. Nishizawa Katazakai, S. & Zhang, J. Long-term flood control in central Japan. Groundw. Sustain. Dev. 33, 101585 (2026).",
    "22. Cosgrove, B. et al. NOAA's National Water Model. JAWRA 60, 247-272 (2024).",
    "23. Chang, J. et al. Flood resource utilization of cascade hydropower stations. EGU General Assembly 2025, EGU25-14374.",
    "24. Dillon, P. et al. Managed aquifer recharge: rediscovering nature as a leading edge technology. Water Sci. Technol. 62, 2338-2345 (2010).",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(36)
    p.paragraph_format.hanging_indent = Pt(36)
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

# ── Figure Legends ──
add_heading("Figure Legends", level=1)

legends = [
    ("Fig. 1", "Hydropower output from planned dam release as a function of effective head (H) "
     "and discharge (Q). Blue dashed line indicates the pumping power requirement (3.7 MW) for "
     "groundwater extraction at 30 m lift, 10 m3/s flow. The area above this line represents "
     "conditions where the system is energy-positive."),
    ("Fig. 2", "Aquifer storage capacity created by pre-flood groundwater drawdown over a 20 km2 "
     "managed area (effective porosity 0.20). Red dashed line: observed flood volume during the "
     "2018 event (15.3 x 10^6 m3). Percentages indicate the fraction of 2018 flood volume "
     "that could be absorbed."),
    ("Fig. 3", "Comparison of flood hydrographs under four scenarios for the Oda River basin "
     "(72-hour event modelled on 2018 rainfall). Upper panel: rainfall intensity; lower panel: "
     "discharge. Scenarios: (1) no action, (2) confluence relocation, (3) groundwater management, "
     "(4) integrated approach."),
    ("Fig. 4", "Flood control effectiveness by scenario. (a) Peak discharge comparison; "
     "(b) Total runoff volume comparison. Percentages indicate reduction relative to no-action baseline."),
]

for fig_id, legend in legends:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run_id = p.add_run(f"{fig_id}  ")
    run_id.font.name = "Times New Roman"
    run_id.font.size = Pt(11)
    run_id.bold = True
    run_text = p.add_run(legend)
    run_text.font.name = "Times New Roman"
    run_text.font.size = Pt(11)

# ── Save ──
outpath = f"{OUTDIR}/nature_water_commentary.docx"
doc.save(outpath)
print(f"Nature Water Commentary saved: {outpath}")
