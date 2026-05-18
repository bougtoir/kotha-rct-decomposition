#!/usr/bin/env python3
"""
Cover Letter (Email Draft) for Water Alternatives — Research Article
WaA submission is by email; this serves as the email body.
British English, Verdana to match WaA style.

Note: Declarations (COI, Funding, AI use, Author contributions) go in the
manuscript Cover Page (Section 1), NOT in this email/cover letter.
Per WaA guidelines, the Cover Page of the manuscript contains all author
information, acknowledgements, and declarations. This email simply introduces
the submission.
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"

doc = Document()

# Page setup (A4)
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# Style
style = doc.styles["Normal"]
style.font.name = "Verdana"
style.font.size = Pt(10)
style.paragraph_format.line_spacing = Pt(18)
style.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, alignment=None, space_before=0,
             space_after=6, font_size=10, color=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Verdana"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_placeholder(text):
    return add_para(text, color=RGBColor(255, 0, 0))


# ── Header ──
add_para("Cover Letter / Submission Email Draft", bold=True, font_size=12,
         space_after=12)

add_para("To: managing_editor@water-alternatives.org", bold=True)

add_para(
    "Subject: Submission of Research Article — "
    "\"Unilateral flood control through subsurface storage management\"",
    bold=True, space_after=12
)

# ── Salutation ──
add_para("Dear Managing Editor,")

# ── Body paragraph 1: Submission ──
add_para("")
add_para(
    "I am pleased to submit the attached manuscript, entitled "
    "\"Unilateral flood control through subsurface storage management: "
    "A hydro-political analysis of planned-release hydropower and "
    "inter-watershed groundwater transfer\", for consideration as a "
    "Research Article in Water Alternatives."
)

# ── Body paragraph 2: Brief summary ──
add_para(
    "The paper proposes a novel flood control framework that enables "
    "downstream nations to implement flood mitigation measures unilaterally "
    "\u2014 without requiring upstream cooperation \u2014 by integrating "
    "planned-release hydropower generation with inter-watershed groundwater "
    "management. The concept is analysed through feasibility studies of "
    "two Japanese basins (Oda River\u2013Takahashi River and "
    "Arakawa\u2013Edogawa) and situated within the hydro-hegemony "
    "literature, with applications discussed for the Mekong, Nile, "
    "Ganges-Brahmaputra-Meghna, Indus, Tigris-Euphrates, and Jordan basins."
)

# ── Body paragraph 3: Fit to WaA ──
add_para(
    "I believe the manuscript is well suited to Water Alternatives because "
    "it engages directly with hydro-hegemony theory, bridges technical "
    "feasibility analysis with water governance scholarship, and has "
    "implications for transboundary flood management across multiple "
    "global river systems."
)

# ── Body paragraph 4: Related submissions ──
add_para(
    "Disclosure: A related manuscript focusing on quantitative engineering "
    "analysis is being submitted to the Journal of JSCE (Special Issue: "
    "Hydroscience and Hydraulic Engineering, B1). The two manuscripts "
    "differ substantially in scope and audience \u2014 the JSCE paper is a "
    "technical engineering paper, while the Water Alternatives manuscript "
    "focuses on water governance and hydro-political dimensions."
)

# ── Body paragraph 5: Confirmations ──
add_para(
    "The manuscript has not been published previously and is not under "
    "consideration elsewhere in its current form. Author information, "
    "declarations (conflicts of interest, funding, AI use, and author "
    "contributions), and acknowledgements are provided on the cover page "
    "(first page) of the attached manuscript, as per WaA guidelines."
)

# ── Body paragraph 6: Attachment note ──
add_para(
    "Attached: Manuscript file (.docx), prepared using the WaA template.",
    italic=True
)

# ── Closing ──
add_para("")
add_para("Kind regards,")

add_para("")
add_para("Tatsuki Onishi")
add_placeholder("(Affiliation)")
add_placeholder("(Email)")

# ── Save ──
outpath = f"{OUTDIR}/waa_cover_letter.docx"
doc.save(outpath)
print(f"WaA cover letter (email draft) saved: {outpath}")
