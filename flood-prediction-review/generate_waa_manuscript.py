#!/usr/bin/env python3
"""
Water Alternatives Research Article manuscript generator.
Uses the official WaA template styles (Verdana 10pt, Harvard/author-date references).
British English throughout.

Structure:
  Section 1 (Cover page): Title, Authors, Affiliations, Acknowledgements, AI declaration
  Section 2 (Main text):  Title, Abstract, Keywords, Body (with inline figures), References
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

OUTDIR = "/home/ubuntu/repos/wip/flood-prediction-review"
FIGDIR = os.path.join(OUTDIR, "figures")

# ─── Use the official WaA template as the base ───
TEMPLATE = "/home/ubuntu/attachments/b9c49450-b80b-4ad8-9d3d-733db5441962/WaAs+template+for+Authors.docx"
doc = Document(TEMPLATE)

# Clear all existing paragraphs (template has instructional text)
for p in doc.paragraphs:
    parent = p._element.getparent()
    parent.remove(p._element)


# ─── Helper functions ───
def add_main_title(text):
    p = doc.add_paragraph(style="Main title")
    run = p.add_run(text)
    return p


def add_authors(text):
    p = doc.add_paragraph(style="Style Centered After:  24 pt")
    run = p.add_run(text)
    return p


def add_normal(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_placeholder(text):
    return add_normal(text, color=RGBColor(255, 0, 0))


def add_heading1(text):
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(text)
    return p


def add_heading2(text):
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    return p


def add_biblio(text):
    p = doc.add_paragraph(style="Biblio")
    # Handle italics for journal/book titles
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
    return p


def add_figure(fig_filename, caption_text):
    """Insert figure inline with caption below."""
    # Figure paragraph (centered)
    p_fig = doc.add_paragraph(style="Figure")
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_path = os.path.join(FIGDIR, fig_filename)
    if os.path.exists(fig_path):
        run = p_fig.add_run()
        run.add_picture(fig_path, width=Inches(5.5))
    else:
        run = p_fig.add_run(f"[Figure: {fig_filename}]")
        run.font.color.rgb = RGBColor(255, 0, 0)

    # Caption paragraph
    p_cap = doc.add_paragraph(style="Caption_figures")
    run = p_cap.add_run(caption_text)
    return p_fig


def add_table_waa(headers, rows, caption):
    """Add table with caption above, WaA style (no vertical lines)."""
    # Caption
    p_cap = doc.add_paragraph(style="Caption_tables")
    run = p_cap.add_run(caption)

    # Table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Normal Table"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)

    # Remove vertical borders (WaA style)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['left', 'right', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none',
            qn('w:sz'): '0',
            qn('w:space'): '0',
            qn('w:color'): 'auto'
        })
        borders.append(border)
    tblPr.append(borders)

    return table


def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


# ============================================================
# SECTION 1: COVER PAGE (not communicated to reviewers)
# ============================================================

add_main_title(
    "Unilateral flood control through subsurface storage management: "
    "A hydro-political analysis of planned-release hydropower "
    "and inter-watershed groundwater transfer"
)

add_authors("Tatsuki Onishi")

add_normal("Tatsuki Onishi")
add_placeholder("(Position)")
add_placeholder("(Affiliation)")
add_placeholder("(City, Country)")
add_placeholder("(Email address)")

add_normal("")  # spacer

add_heading1("Declaration of interest")
add_normal(
    "The author declares no conflicts of interest."
)

add_heading1("Funding")
add_normal(
    "This work received no external funding."
)

add_heading1("Acknowledgements")
add_normal("None")

add_heading1("Declaration of generative artificial intelligence (AI) in scientific writing")
add_normal(
    "We used devin.ai to help with formatting the text and choosing words that suited "
    "the tone, and to help writing codes. The author takes full responsibility for the "
    "accuracy and content of the manuscript."
)

add_heading1("Authors contribution")
add_normal(
    "O.T.: Conceptualisation, Methodology, Software, Formal analysis, "
    "Writing – original draft, Writing – review and editing"
)

# ============================================================
# PAGE BREAK → SECTION 2: MAIN TEXT (for reviewers)
# ============================================================
doc.add_page_break()

add_main_title(
    "Unilateral flood control through subsurface storage management: "
    "A hydro-political analysis of planned-release hydropower "
    "and inter-watershed groundwater transfer"
)

# ── Abstract ──
add_heading1("Abstract")
add_normal(
    "This paper proposes a novel flood control framework that integrates planned-release "
    "hydropower generation with inter-watershed groundwater management, enabling downstream "
    "nations or sub-national jurisdictions to implement flood mitigation measures unilaterally "
    "— without requiring upstream cooperation. Drawing on a simplified feasibility analysis "
    "of the Oda River–Takahashi River basin in western Japan, where catastrophic backwater "
    "flooding killed 51 people in 2018, this paper demonstrates that such a system can be "
    "energy self-sufficient (consuming only 9% of generated power for groundwater pumping) "
    "and can create subsurface storage equivalent to 78% of observed flood volumes. "
    "The paper situates this technical proposal within the hydro-hegemony literature, "
    "arguing that subsurface storage management represents a technologically grounded "
    "pathway for downstream actors to reduce flood vulnerability independently of upstream "
    "water governance arrangements. Applications to transboundary river basins (Mekong, "
    "Nile, Jordan), deltaic regions (Mekong Delta, Ganges-Brahmaputra-Meghna), and arid "
    "environments (Middle East wadi systems) are discussed. The framework offers a new "
    "dimension to debates on water security, climate adaptation, and the politics of "
    "flood infrastructure."
)

# ── Keywords ──
add_heading1("Keywords")
add_normal(
    "Flood control; groundwater management; hydro-hegemony; transboundary rivers; "
    "hydropower; aquifer storage; Japan"
)

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading1("Introduction")

add_normal(
    "The governance of flood risk in transboundary river basins remains one of the most "
    "intractable challenges in contemporary water management. Unlike drought, which "
    "develops gradually and allows time for negotiation, floods arrive suddenly and their "
    "consequences are shaped by decisions — about dam operations, land use, and drainage "
    "infrastructure — often made far upstream and beyond the control of affected "
    "populations. The 2018 Western Japan Heavy Rainfall, which killed over 200 people "
    "across western Japan, demonstrated that even a technologically advanced nation with "
    "extensive flood infrastructure remains vulnerable to compound flood mechanisms, "
    "particularly the backwater effect at river confluences (NIED, 2018). If Japan, "
    "with its world-leading investment in flood control, cannot eliminate backwater "
    "flooding through conventional structural measures alone, what options exist for "
    "less-resourced downstream nations facing similar — or greater — flood risks from "
    "transboundary rivers they do not control?"
)

add_normal(
    "Flood management in transboundary river basins is inherently political. Decisions about "
    "dam operations, levee construction, and floodplain zoning made by upstream riparian states "
    "directly affect downstream flood risk, often without meaningful consultation or "
    "compensation (Zeitoun and Warner, 2006). The concept of hydro-hegemony — whereby the "
    "most powerful riparian state controls water resources to the disadvantage of weaker "
    "neighbours — has become a central analytical framework in critical water studies "
    "(Zeitoun and Warner, 2006; Cascão and Zeitoun, 2010). Yet the literature on "
    "hydro-hegemony has focused predominantly on water quantity and allocation disputes, "
    "with comparatively less attention to flood risk as a dimension of water power asymmetry."
)

add_normal(
    "This paper addresses that gap by proposing a flood control framework that can be "
    "implemented unilaterally by a downstream riparian state, without requiring upstream "
    "cooperation or modification of shared water infrastructure. The framework integrates "
    "two established technologies — hydropower generation from controlled flood discharge "
    "and managed aquifer recharge (MAR) — into a single system that uses the energy from "
    "planned dam releases to pump groundwater from adjacent watersheds, thereby creating "
    "subsurface storage capacity to absorb incoming floodwaters. Because the entire system "
    "operates within the implementing state's territory and draws on subsurface rather than "
    "surface water resources, it circumvents the governance challenges that typically "
    "constrain transboundary flood management."
)

add_normal(
    "The concept is not entirely without precedent. California's Flood-MAR (Flood-Managed "
    "Aquifer Recharge) programme deliberately directs floodwaters onto agricultural lands "
    "to recharge depleted aquifers (California DWR, 2018). Japan's Metropolitan Area Outer "
    "Underground Discharge Channel (MAOUDC) transfers floodwater between river basins "
    "through a 6.3 km underground tunnel (MLIT, 2006). Pumped-storage hydropower — in "
    "which water is pumped between upper and lower reservoirs to store and release energy "
    "— is a mature technology with over 27,000 MW of installed capacity in Japan alone "
    "(TEPCO, 2023). What is novel about the framework proposed here is the integration of "
    "these existing technological components into a single system designed specifically for "
    "flood control, and the recognition that such a system can operate entirely within the "
    "territory of a single state."
)

add_normal(
    "The argument proceeds as follows. First, this paper reviews the literature on "
    "hydro-hegemony, transboundary flood governance, and the emerging field of flood-managed "
    "aquifer recharge (Flood-MAR). Second, it presents the proposed framework and a "
    "simplified feasibility analysis for the Oda River–Takahashi River basin in western "
    "Japan, the site of catastrophic backwater flooding during the 2018 Western Japan Heavy "
    "Rainfall. Third, it examines the applicability of the framework to the Tokyo "
    "metropolitan area (Arakawa–Edogawa basin), where 8.21 million people reside in "
    "projected inundation zones. Finally, it discusses the implications for transboundary "
    "river governance, deltaic flood protection, and arid-region water management, arguing "
    "that subsurface storage management represents a new dimension in the politics of "
    "flood infrastructure."
)

add_normal(
    "A note on terminology: this paper uses 'unilateral' to describe actions that can be "
    "taken by a single state without requiring the cooperation or consent of neighbouring "
    "states. This is distinct from 'unilateralism' in the pejorative sense sometimes used "
    "in international relations. The framework proposed here does not advocate for the "
    "abandonment of cooperative water governance; rather, it offers a technical option "
    "that can be pursued in parallel with, and potentially in support of, cooperative "
    "approaches."
)

# ============================================================
# 2. HYDRO-HEGEMONY AND FLOOD GOVERNANCE
# ============================================================
add_heading1("Hydro-hegemony and the politics of flood infrastructure")

add_heading2("Hydro-hegemony and downstream vulnerability")
add_normal(
    "The hydro-hegemony framework, as articulated by Zeitoun and Warner (2006), describes "
    "how powerful riparian states exploit their geographic, economic, or military advantages "
    "to control shared water resources. While the original formulation emphasised water "
    "allocation and dam construction, subsequent scholarship has extended the concept to "
    "encompass flood risk. Downstream states in major transboundary basins — Vietnam on "
    "the Mekong, Egypt on the Nile, Bangladesh on the Ganges-Brahmaputra-Meghna (GBM) — "
    "face flood risks that are substantially shaped by upstream land use, dam operations, "
    "and climate change, yet have limited influence over these upstream drivers "
    "(Mirumachi and Allan, 2007; Middleton, 2022)."
)

add_normal(
    "The Mekong River exemplifies this dynamic. China's cascade of 11 mainstream dams in "
    "the upper Mekong (Lancang) has fundamentally altered the river's hydrology, with "
    "documented effects on downstream flood and drought patterns in Thailand, Laos, "
    "Cambodia, and Vietnam (Räsänen et al., 2017; Mekong River Commission, 2023). "
    "Vietnam's Mekong Delta — home to 17 million people and producing over half of "
    "the country's rice — faces the dual threat of upstream flow modification and "
    "sea-level rise, with limited recourse through existing governance mechanisms "
    "(Lebel et al., 2005; Kuenzer et al., 2013)."
)

add_normal(
    "Similarly, Egypt's concerns over the Grand Ethiopian Renaissance Dam (GERD) on the "
    "Blue Nile reflect the asymmetry between upstream development prerogatives and "
    "downstream flood and water security risks (Cascão and Zeitoun, 2010; Wheeler et al., "
    "2020). In the Jordan River basin, Israel's control of the headwaters and the Sea of "
    "Galilee has historically constrained Jordanian and Palestinian water management "
    "options (Zeitoun, 2008). In each case, downstream states lack autonomous mechanisms "
    "to manage flood risk independently of upstream decisions."
)

add_normal(
    "The 1997 United Nations Convention on the Law of the Non-Navigational Uses of "
    "International Watercourses established principles of equitable utilisation and the "
    "obligation not to cause significant harm, but enforcement mechanisms remain weak "
    "and disputes continue to be resolved — or not — through bilateral negotiation "
    "(McCaffrey, 2007). The Mekong River Commission, despite decades of institutional "
    "development, has been unable to prevent unilateral dam construction on the mainstream "
    "by China and Laos (Lebel et al., 2005; Mekong River Commission, 2023). The Nile Basin "
    "Initiative, established in 1999, has yet to produce a comprehensive basin-wide agreement "
    "acceptable to all riparian states (Wheeler et al., 2020). These institutional failures "
    "underscore the need for technical solutions that do not depend on political cooperation."
)

add_heading2("Conventional flood infrastructure and its limitations")
add_normal(
    "Conventional structural flood measures — levees, dams, diversion channels, and "
    "floodways — require either upstream cooperation (for dam operation coordination) "
    "or massive capital investment within the downstream state's territory. Japan's "
    "experience illustrates both the achievements and limitations of the structural "
    "approach. The Metropolitan Area Outer Underground Discharge Channel (MAOUDC) in "
    "Saitama Prefecture, completed in 2006, transfers floodwater from four rivers into "
    "the Edogawa River at a capacity of 200 m³/s — a remarkable engineering achievement, "
    "but one that required 13 years of construction and approximately US$2 billion "
    "(MLIT, 2006). Super-levees along the Arakawa River, designed to withstand "
    "overtopping, have reached only 15% completion after decades of construction "
    "(MLIT, 2020a)."
)

add_normal(
    "The Mississippi River basin in the United States provides a continental-scale example "
    "of structural limitations. The Yazoo Backwater area in Mississippi experiences chronic "
    "backwater flooding when the Mississippi River rises and prevents the Yazoo River from "
    "draining \u2014 a mechanism identical to the Oda\u2013Takahashi backwater effect in Japan, but "
    "at a vastly larger scale. A proposed pumping station to mitigate Yazoo backwater "
    "flooding has been debated for over 80 years, blocked by environmental concerns and "
    "interstate political disputes (Berkowitz et al., 2019). The Yazoo case demonstrates "
    "that even within a single nation, structural flood measures face governance constraints "
    "when they affect multiple jurisdictions with competing interests."
)

add_normal(
    "The backwater effect — whereby rising water levels in a main river impede the "
    "drainage of tributary streams — represents a particularly challenging flood "
    "mechanism. During the 2018 Western Japan Heavy Rainfall, the Takahashi River's "
    "water level rose so rapidly that the Oda River, a tributary joining from the west, "
    "could not discharge its floodwaters. The resulting backwater inundation flooded "
    "8.28 km² of the Mabi district in Kurashiki City, killing 51 people and destroying "
    "over 4,600 homes (NIED, 2018; MLIT Chugoku Regional Development Bureau, 2019). "
    "Japan's response was a ¥28 billion confluence relocation project, completed in "
    "March 2024, which moved the Oda–Takahashi junction 4.6 km downstream to reduce "
    "backwater propagation (MLIT, 2024)."
)

add_normal(
    "The 2018 disaster exposed a fundamental limitation of the existing flood control "
    "system: the Oda River's drainage was entirely dependent on gravity flow into the "
    "Takahashi River. When the main river's water level exceeded the tributary's, no "
    "mechanism existed to store or divert the excess water. The confluence relocation "
    "project addresses this by reducing the severity of the backwater effect, but it "
    "does not eliminate the underlying vulnerability \u2014 during extreme events exceeding "
    "the design capacity, backwater inundation could still occur. The proposed framework "
    "offers a complementary approach by creating subsurface storage capacity that can "
    "absorb floodwaters regardless of the main river's water level."
)

add_heading2("Managed aquifer recharge and Flood-MAR")
add_normal(
    "Managed aquifer recharge (MAR) — the intentional replenishment of groundwater through "
    "engineered infiltration — has been practised for decades in water supply contexts "
    "(Dillon et al., 2019). The California Department of Water Resources has pioneered "
    "'Flood-MAR' (Flood-Managed Aquifer Recharge), which explicitly links flood management "
    "with groundwater banking by directing excess surface water during flood events into "
    "agricultural lands and dedicated recharge basins (California DWR, 2018). Taiwan has "
    "implemented optimised groundwater pumping-recharge cycles for combined flood control "
    "and water supply (Chang et al., 2019). In Japan, Nishizawa and Zhang (2026) have "
    "documented a half-century of groundwater monitoring data from central Japan, "
    "demonstrating the long-term viability of managed aquifer systems."
)

add_normal(
    "However, existing Flood-MAR applications are passive — they rely on gravity-driven "
    "infiltration during flood events, which limits their applicability to regions with "
    "suitable topography and soil permeability. The framework proposed in this paper "
    "extends Flood-MAR by introducing an active, energy-driven component: using "
    "hydropower from planned dam releases to pre-emptively pump groundwater, thereby "
    "creating storage capacity before flood events occur."
)

add_normal(
    "The relationship between groundwater extraction and land subsidence is well documented "
    "and represents a key constraint on managed aquifer drawdown. Tokyo's Kanto Plain "
    "experienced up to 4.5 m of cumulative land subsidence between the Meiji era and the "
    "1970s due to industrial groundwater pumping, prompting stringent regulations that have "
    "since allowed groundwater levels to recover substantially (GSJ/AIST, 2020). The Taiwan "
    "experience offers further insight: Chang et al. (2019) demonstrated that optimised "
    "pumping-recharge cycles can manage subsidence risk while achieving flood control "
    "objectives, provided that drawdown is temporary (days to weeks) rather than sustained. "
    "This distinction between chronic extraction (which causes compaction of clay layers "
    "and irreversible subsidence) and cyclical, short-duration drawdown (which primarily "
    "affects elastic storage in sand-gravel aquifers) is critical to the feasibility of "
    "the proposed framework."
)

# ============================================================
# 3. PROPOSED FRAMEWORK
# ============================================================
add_heading1("The proposed framework: Planned-release hydropower with inter-watershed "
             "groundwater management")

add_normal(
    "This section presents the technical architecture of the proposed framework. The "
    "system integrates three established technologies — hydroelectric generation from "
    "controlled dam releases, submersible groundwater pumping, and managed aquifer "
    "recharge — into a coordinated cycle that can be activated in anticipation of flood "
    "events. The key innovation lies not in any individual component but in their "
    "integration into a single energy-positive system that can operate within the "
    "jurisdiction of a single state."
)

add_normal(
    "The proposed framework operates in four phases:"
)

add_normal(
    "Phase 1 — Planned release and power generation: When flood risk is forecast, "
    "upstream dam operators release water through turbines at controlled rates, generating "
    "hydroelectric power. The energy potential is given by P_gen = η_t × ρ × g × Q × H, "
    "where η_t is turbine efficiency (typically 0.85), ρ is water density (1,000 kg/m³), "
    "g is gravitational acceleration (9.81 m/s²), Q is discharge (m³/s), and H is the "
    "effective head (m)."
)

add_normal(
    "Phase 2 — Groundwater pumping in adjacent watershed: The generated electricity powers "
    "submersible pumps that extract groundwater from the alluvial aquifer of an adjacent "
    "watershed. The pumping power requirement is P_pump = (ρ × g × Q_p × H_p) / η_p, "
    "where Q_p is the pumping rate, H_p is the pumping head (water table depth plus "
    "delivery head), and η_p is pump efficiency (typically 0.80)."
)

add_normal(
    "Phase 3 — Flood absorption: The pre-emptied aquifer storage volume (V = A × Δh × S_y, "
    "where A is the recharge area, Δh is the water table drawdown, and S_y is the specific "
    "yield) provides capacity to absorb floodwaters through infiltration during the event."
)

add_normal(
    "Phase 4 — Post-flood recovery: After the flood recedes, the aquifer naturally recharges "
    "through precipitation and lateral groundwater flow, restoring the system's capacity "
    "for the next event."
)

add_normal(
    "The system is energy self-sufficient when P_gen > P_pump. The surplus energy "
    "(P_gen − P_pump) can be supplied to the electricity grid, creating an economic "
    "co-benefit. Crucially, the entire system — dam, turbines, pumping wells, and "
    "target aquifer — can be located within a single national jurisdiction, requiring "
    "no transboundary coordination."
)

add_normal(
    "Several important technical conditions must be satisfied for the framework to be "
    "viable. The dam must have sufficient head to generate economically meaningful power; "
    "the adjacent watershed must contain a productive aquifer with adequate specific yield; "
    "the distance between the dam and the pumping well field must be manageable for power "
    "transmission; and the recharge area must have sufficient infiltration capacity to "
    "absorb floodwaters. In the following section, these conditions are assessed for a "
    "specific case: the Oda River–Takahashi River basin in western Japan."
)

# ============================================================
# 4. FEASIBILITY ANALYSIS: ODA-TAKAHASHI BASIN
# ============================================================
add_heading1("Feasibility analysis: The Oda River–Takahashi River basin")

add_heading2("Study area and the 2018 disaster")
add_normal(
    "The Oda River–Takahashi River basin in Okayama Prefecture, western Japan, provides "
    "an ideal test case for the proposed framework. The Oda River (basin area 211 km²) "
    "joins the Takahashi River at an acute angle, making it highly susceptible to "
    "backwater effects. During the 2018 Western Japan Heavy Rainfall (5–8 July 2018), "
    "three-day cumulative rainfall reached 320 mm, exceeding the design rainfall of the "
    "flood control system. The Takahashi River's rapid rise prevented the Oda River from "
    "discharging, causing levee breaches at five locations and inundating 8.28 km² of "
    "the Mabi district with an estimated flood volume of 15.3 × 10⁶ m³ (NIED, 2018; "
    "MLIT Chugoku Regional Development Bureau, 2019)."
)

# Table 1: 2018 event parameters
add_table_waa(
    ["Parameter", "Value"],
    [
        ["Three-day cumulative rainfall", "320 mm"],
        ["Inundation area", "8.28 km²"],
        ["Estimated flood volume", "15.3 × 10⁶ m³"],
        ["Levee breaches", "5 locations"],
        ["Fatalities", "51"],
        ["Basin area (Oda River)", "211 km²"],
        ["Channel gradient (Oda River)", "1/2,200"],
    ],
    "Table 1. Key parameters of the 2018 Western Japan Heavy Rainfall event "
    "in the Oda River–Takahashi River basin."
)

add_normal(
    "In March 2024, a confluence relocation project was completed, moving the Oda–Takahashi "
    "junction 4.6 km downstream. Hydraulic analysis indicates this reduces the Oda River "
    "water level by approximately 4.6 m and the Takahashi River by 0.8 m under design "
    "flood conditions (MLIT, 2024; KSB, 2025)."
)

add_heading2("Energy balance")
add_normal(
    "For a planned release through a dam with an effective head of H = 100 m and a "
    "discharge of Q = 50 m³/s, the generated power is:"
)
add_normal(
    "P_gen = 0.85 × 1,000 × 9.81 × 50 × 100 = 41.7 MW"
)
add_normal(
    "For groundwater pumping at Q_p = 10 m³/s with a pumping head of H_p = 30 m:"
)
add_normal(
    "P_pump = (1,000 × 9.81 × 10 × 30) / 0.80 = 3.68 MW"
)
add_normal(
    "The system thus consumes only 8.8% of the generated power for groundwater management, "
    "leaving a surplus of 38.0 MW available for the electricity grid. figure 1 shows the "
    "energy balance as a function of head and discharge, with the red dashed line indicating "
    "the pumping power requirement."
)

# Figure 1
add_figure(
    "fig1_energy_balance_en.png",
    "Figure 1. Energy balance contour map showing generated power (MW) as a function "
    "of effective head and discharge. The red dashed line indicates the pumping power "
    "requirement (3.68 MW). The green marker shows the reference scenario (H = 100 m, "
    "Q = 50 m³/s)."
)

add_heading2("Aquifer storage capacity")
add_normal(
    "For a recharge area of A = 20 km², a water table drawdown of Δh = 3 m, and a "
    "specific yield of S_y = 0.2 (typical of alluvial sand-gravel aquifers):"
)
add_normal(
    "V = 20 × 10⁶ × 3 × 0.2 = 12.0 × 10⁶ m³"
)
add_normal(
    "This represents 78% of the 2018 flood volume (15.3 × 10⁶ m³), suggesting that "
    "pre-emptive aquifer management could absorb the majority of a design-level flood "
    "event. figure 2 presents the storage capacity as a function of recharge area and "
    "water table drawdown."
)

# Figure 2
add_figure(
    "fig2_aquifer_storage_en.png",
    "Figure 2. Aquifer storage capacity contours (10⁶ m³) as a function of recharge "
    "area and water table drawdown, assuming a specific yield of 0.2. The dashed line "
    "indicates the 2018 flood volume (15.3 × 10⁶ m³). The green marker shows the "
    "reference scenario."
)

add_heading2("Flood hydrograph modification")
add_normal(
    "A simplified 72-hour rainfall-runoff simulation was conducted comparing four "
    "scenarios: (1) no measures (baseline); (2) confluence relocation only (4.6 m water "
    "level reduction, equivalent to approximately 30% peak discharge reduction); "
    "(3) groundwater management only (pre-emptied aquifer absorbs infiltration, reducing "
    "runoff volume by 41%); and (4) integrated measures (confluence relocation plus "
    "groundwater management). The integrated scenario achieves a 32% peak discharge "
    "reduction compared with baseline (figure 3), with the groundwater management "
    "component contributing primarily through volume reduction rather than peak attenuation "
    "(figure 4)."
)

# Figure 3
add_figure(
    "fig3_hydrograph_en.png",
    "Figure 3. Flood hydrograph comparison for the 2018 event under four scenarios: "
    "no measures (baseline), confluence relocation only, groundwater management only, "
    "and integrated measures. Upper panel shows the 72-hour rainfall hyetograph."
)

# Figure 4
add_figure(
    "fig4_scenario_comparison_en.png",
    "Figure 4. Comparison of peak discharge and total runoff volume across four scenarios. "
    "Error bars indicate ±10% sensitivity range."
)

# Table 2: Scenario comparison
add_table_waa(
    ["Scenario", "Peak discharge (m³/s)", "Peak reduction (%)", "Infiltration volume (10⁶ m³)"],
    [
        ["1. No measures (baseline)", "850", "—", "0"],
        ["2. Confluence relocation", "595", "30", "0"],
        ["3. Groundwater management", "833", "2", "6.3"],
        ["4. Integrated", "578", "32", "6.3"],
    ],
    "Table 2. Comparison of flood control scenarios for the 2018 event."
)

# ============================================================
# 5. SCALING UP: TOKYO METROPOLITAN AREA
# ============================================================
add_heading1("Scaling up: The Tokyo metropolitan area")

add_normal(
    "While the Oda River–Takahashi River basin provides a compelling proof-of-concept, "
    "the true significance of the proposed framework lies in its potential application "
    "to densely populated metropolitan areas where conventional flood infrastructure "
    "faces severe space constraints. The Tokyo metropolitan area presents such a case."
)

add_normal(
    "The Arakawa–Edogawa basin in the Tokyo metropolitan area presents a compelling "
    "case for the proposed framework. An estimated 8.21 million people reside within "
    "projected inundation zones, and the existing flood control infrastructure — while "
    "extensive — faces capacity constraints under intensifying rainfall patterns "
    "(MLIT, 2020b)."
)

# Table 3: Tokyo infrastructure
add_table_waa(
    ["Facility", "Capacity/Scale", "Year"],
    [
        ["Arakawa Floodway", "22 km diversion channel", "1930"],
        ["Iwabuchi Floodgate", "Isolates Sumida from Arakawa", "1982 (rebuilt)"],
        ["Saiko Retarding Basin", "39 × 10⁶ m³", "1997"],
        ["Upstream dams (3)", "Combined storage", "Various"],
        ["MAOUDC", "200 m³/s inter-basin transfer", "2006"],
        ["Super-levees", "15% completed (Arakawa)", "Ongoing"],
    ],
    "Table 3. Existing flood control infrastructure in the Arakawa–Edogawa basin."
)

add_normal(
    "The Kanto Plain beneath Tokyo hosts alluvial aquifers reaching 80 m in thickness, "
    "comprising alternating sand-gravel and clay layers mapped by the Geological Survey "
    "of Japan's three-dimensional geological model (GSJ/AIST, 2020). Historical "
    "groundwater extraction caused severe land subsidence from the Meiji era through "
    "the 1970s, but stringent pumping regulations implemented since then have allowed "
    "groundwater levels to recover substantially — paradoxically creating conditions "
    "favourable for managed drawdown during flood emergencies."
)

add_normal(
    "Applying the proposed framework to the mid-Arakawa basin: planned releases from "
    "the Takizawa Dam (effective head ~100 m) would generate hydroelectric power to pump "
    "groundwater from a 30 km² mid-basin area, creating approximately 13.5 × 10⁶ m³ of "
    "subsurface storage (30 × 10⁶ × 3 × 0.15). Combined with the existing Saiko Retarding "
    "Basin (39 × 10⁶ m³), the total available storage capacity would reach approximately "
    "52.5 × 10⁶ m³ — a substantial addition to the metropolitan area's flood resilience."
)

# ============================================================
# 6. IMPLICATIONS FOR TRANSBOUNDARY FLOOD GOVERNANCE
# ============================================================
add_heading1("Implications for transboundary flood governance")

add_heading2("International legal dimensions")
add_normal(
    "The framework also has implications for international water law. The 1997 United "
    "Nations Watercourses Convention requires states to utilise international watercourses "
    "in an equitable and reasonable manner and to take appropriate measures to prevent "
    "significant harm to other watercourse states (McCaffrey, 2007). However, these "
    "obligations apply to 'international watercourses' defined as surface water systems "
    "that cross or form international boundaries. Groundwater that is not connected to an "
    "international watercourse falls outside the Convention's scope. If a downstream state "
    "manages its own alluvial aquifers to create flood storage capacity, it operates in a "
    "legal grey area that may not trigger the notification and consultation obligations "
    "that apply to surface water interventions. This legal ambiguity could facilitate "
    "implementation but also raises questions about the evolving relationship between "
    "surface water and groundwater governance in international law."
)

add_heading2("Challenging the upstream-dominance paradigm")
add_normal(
    "The proposed framework has significant implications for the hydro-hegemony literature. "
    "By enabling downstream states to create flood storage capacity within their own "
    "territories, the framework challenges the assumption that downstream flood vulnerability "
    "is an inevitable consequence of upstream control. This does not eliminate the power "
    "asymmetry — upstream dam operations still affect downstream hydrology — but it "
    "provides a technological pathway for downstream actors to reduce their dependence on "
    "upstream cooperation for flood management."
)

add_normal(
    "Cascão and Zeitoun (2010) have argued that counter-hegemonic strategies in "
    "transboundary water governance typically require either international legal mechanisms "
    "or coalition-building among weaker riparian states. The framework proposed here offers "
    "a third pathway: unilateral technical adaptation that operates independently of "
    "diplomatic processes. This is particularly relevant in basins where political tensions "
    "preclude cooperative flood management, such as the Indus (India–Pakistan), the "
    "Tigris-Euphrates (Turkey–Iraq–Syria), and the Jordan (Israel–Jordan–Palestine)."
)

add_heading2("Transboundary river applications")
add_normal(
    "Mekong River basin: Vietnam could implement subsurface storage management in the "
    "Mekong Delta's alluvial aquifers, using hydropower from domestic dam releases to "
    "pre-emptively lower groundwater levels before the monsoon flood season. This would "
    "complement existing flood management strategies without requiring coordination with "
    "upstream Lancang dams in China (Lebel et al., 2005; Räsänen et al., 2017)."
)

add_normal(
    "Nile River basin: Egypt's Nile Delta, already vulnerable to subsidence and "
    "sea-level rise, could benefit from managed aquifer storage that absorbs excess "
    "flood releases from the Aswan High Dam during high-water years. The Nile Delta's "
    "extensive Quaternary aquifer system provides potential storage capacity "
    "(Wheeler et al., 2020)."
)

add_normal(
    "Ganges-Brahmaputra-Meghna basin: Bangladesh, which faces annual devastating floods "
    "from three major transboundary rivers, could develop subsurface storage in the "
    "Bengal Basin's thick alluvial deposits — one of the world's largest aquifer systems "
    "— as an autonomous flood management strategy independent of upstream Indian dam "
    "operations."
)

add_normal(
    "Indus River basin: The Indus, shared between India and Pakistan under the 1960 "
    "Indus Waters Treaty, presents one of the most politically sensitive transboundary "
    "flood contexts in the world. Pakistan's 2010 and 2022 mega-floods demonstrated the "
    "country's extreme vulnerability to upstream hydrological events. The thick alluvial "
    "deposits of the Indus Plain \u2014 similar in character to those of the Kanto Plain \u2014 "
    "could potentially support subsurface flood storage, allowing Pakistan to develop "
    "autonomous flood management capacity independent of the treaty framework."
)

add_normal(
    "Tigris-Euphrates basin: Turkey's Southeastern Anatolia Project (GAP), which includes "
    "22 dams on the Tigris and Euphrates, has fundamentally altered downstream hydrology "
    "in Iraq and Syria. Iraq's Mesopotamian marshlands \u2014 a UNESCO World Heritage Site \u2014 "
    "have been particularly affected. The alluvial deposits of the Mesopotamian Plain "
    "offer potential for subsurface storage development, which could provide Iraq with "
    "autonomous flood and water management capacity."
)

add_heading2("Deltaic flood protection")
add_normal(
    "Delta regions face a unique 'double squeeze' between upstream flood pulses and "
    "downstream sea-level rise. Conventional approaches — heightening levees, building "
    "sea walls — address one threat at a time and often exacerbate the other (by "
    "preventing sediment deposition that counteracts subsidence). Subsurface storage "
    "management offers a complementary approach that reduces flood peaks without "
    "interfering with sediment transport processes. This is particularly relevant for "
    "the Mekong Delta (Vietnam), the Nile Delta (Egypt), and the GBM Delta (Bangladesh), "
    "where populations are large and conventional infrastructure is already under stress."
)

add_heading2("Arid-region applications")
add_normal(
    "In arid and semi-arid regions, flash floods (wadi floods) cause significant damage "
    "despite — or because of — the general scarcity of water. Countries in the Arabian "
    "Peninsula, particularly Saudi Arabia and the United Arab Emirates, experience "
    "infrequent but intense rainfall events that generate destructive flash floods "
    "while groundwater resources are rapidly depleting. The proposed framework could be "
    "adapted to capture wadi flood energy and direct floodwaters into depleted aquifers, "
    "simultaneously addressing flood risk and water scarcity. The availability of "
    "substantial financial resources (the so-called 'oil money') in these countries "
    "could facilitate rapid implementation of such integrated systems. Saudi Arabia's NEOM "
    "project, a US$500 billion planned city in the northwestern desert, has identified water "
    "management as a core challenge \u2014 the proposed framework could align with such "
    "mega-project ambitions by combining flood risk reduction with strategic aquifer "
    "replenishment."
)

add_normal(
    "The United Arab Emirates has invested heavily in artificial aquifer recharge projects, "
    "injecting treated wastewater and desalinated water into limestone and alluvial "
    "aquifers as a strategic reserve. The Abu Dhabi aquifer storage and recovery (ASR) "
    "programme, for example, has demonstrated the technical feasibility of large-scale "
    "managed aquifer operations in arid environments. The proposed framework extends this "
    "logic by using episodic flood energy rather than continuous pumped injection, "
    "potentially reducing the energy costs of aquifer replenishment while simultaneously "
    "addressing flash flood risk."
)

add_normal(
    "The concept also has relevance for climate adaptation in regions experiencing "
    "increasing precipitation variability. As climate change intensifies the hydrological "
    "cycle \u2014 producing both more severe droughts and more intense precipitation events \u2014 "
    "the ability to capture and store episodic floodwaters in subsurface reservoirs becomes "
    "increasingly valuable. This dual-purpose function (flood control plus water banking) "
    "aligns with the 'nexus' thinking that has become central to contemporary water "
    "governance discourse, though the political implications of unilateral implementation "
    "in transboundary contexts deserve careful consideration."
)

# ============================================================
# 7. DISCUSSION
# ============================================================
add_heading1("Discussion")

add_heading2("The politics of autonomous flood adaptation")
add_normal(
    "The framework proposed in this paper raises important questions about the politics "
    "of autonomous adaptation. If downstream states can reduce their flood vulnerability "
    "independently of upstream cooperation, does this weaken incentives for cooperative "
    "basin governance? Or does it, paradoxically, strengthen the negotiating position of "
    "downstream states by reducing their dependence on upstream goodwill? The hydro-hegemony "
    "literature suggests that power asymmetries in transboundary basins are self-reinforcing: "
    "powerful upstream states have little incentive to cooperate because the status quo serves "
    "their interests (Zeitoun and Warner, 2006). If downstream states can demonstrate "
    "autonomous flood management capacity, this may alter the calculus of cooperation by "
    "reducing the upstream state's leverage. Alternatively, some scholars might argue that "
    "unilateral adaptation could undermine collective action by allowing powerful states to "
    "'opt out' of basin-wide governance frameworks. These questions merit further "
    "investigation through both theoretical analysis and empirical case studies."
)

add_heading2("Energy self-sufficiency and economic viability")
add_normal(
    "The energy balance analysis demonstrates a striking asymmetry: with a head of 100 m "
    "and discharge of 50 m³/s, the system generates 41.7 MW while requiring only 3.68 MW "
    "for pumping — a consumption ratio of just 8.8%. This compares favourably with "
    "conventional pumped-storage hydropower, which typically achieves 70–80% round-trip "
    "efficiency (IEA, 2021). The surplus energy (38.0 MW) represents both an economic "
    "co-benefit and an incentive for implementation. Japan's extensive pumped-storage "
    "hydropower infrastructure (TEPCO alone operates approximately 10,000 MW of capacity) "
    "provides the technical expertise and supply chains necessary for rapid deployment "
    "(TEPCO, 2023)."
)

add_normal(
    "From an economic perspective, the surplus energy represents a potential revenue "
    "stream that could offset the capital costs of well field construction and aquifer "
    "monitoring. At a wholesale electricity price of USbash.10/kWh (approximate Japanese "
    "market rate), 38 MW of surplus power generated over a 72-hour flood event would "
    "produce approximately US74,000 in revenue per event. While this alone would not "
    "justify the infrastructure investment, the combined benefits of flood damage "
    "avoidance, energy generation, and potential carbon credits from displacing thermal "
    "generation create a more compelling economic case that warrants detailed cost-benefit "
    "analysis in future work."
)

add_heading2("Japan's topographic advantages")
add_normal(
    "Japan's steep topography provides a critical advantage for the proposed framework. "
    "Japanese river gradients are typically 5 to 20 times steeper than those of continental "
    "rivers such as the Mississippi or the Mekong, while inter-watershed distances are "
    "measured in kilometres rather than hundreds of kilometres. This combination of high "
    "head and short transfer distances makes energy-positive inter-watershed groundwater "
    "management physically feasible in a way that may not be replicable in low-gradient "
    "continental settings. However, even in low-gradient settings, the principle of using "
    "renewable energy (whether hydropower, solar, or wind) to drive pre-emptive aquifer "
    "drawdown remains valid \u2014 the energy source simply shifts from hydropower to other "
    "renewables. In the Middle East, for example, solar-powered groundwater pumping is "
    "already widely practised for irrigation; the same technology could be repurposed for "
    "flood-preparedness aquifer management."
)

add_heading2("Complementarity with structural measures")
add_normal(
    "The proposed framework is not intended to replace structural flood measures but to "
    "complement them. The integrated scenario (confluence relocation plus groundwater "
    "management) achieves a 32% peak discharge reduction — greater than either measure "
    "alone — and reduces total runoff volume by 41% through aquifer infiltration. This "
    "complementarity is significant for Japan's Basin Flood Control (Ryuiki Chisui) "
    "policy, which explicitly calls for diversified, multi-layered flood management "
    "approaches (MLIT, 2021)."
)

add_normal(
    "The simulated flood hydrograph (figure 3) reveals an important distinction between "
    "the two measures. Confluence relocation primarily reduces peak discharge by lowering "
    "the backwater effect, while groundwater management primarily reduces total runoff "
    "volume by diverting water into subsurface storage. This functional complementarity "
    "means that the two approaches target different aspects of flood risk: the peak "
    "(which determines levee overtopping) and the volume (which determines inundation "
    "duration and total damage). The integrated scenario captures both benefits, achieving "
    "a 32% peak reduction and 41% volume reduction (table 2). This suggests that the "
    "framework is most valuable not as a standalone solution but as a complement to "
    "existing structural measures."
)

add_normal(
    "The Japanese government's 2021 Basin Flood Control Act "
    "explicitly calls for the integration of structural and non-structural measures, "
    "including the utilisation of upstream storage, floodplain management, and 'green "
    "infrastructure.' The subsurface storage management proposed here represents a natural "
    "extension of this policy framework \u2014 using the subsurface as a 'hidden retarding basin' "
    "that does not require land acquisition or surface infrastructure in densely populated "
    "urban areas."
)

add_heading2("Limitations and further research")
add_normal(
    "This study employs a simplified analytical framework and does not constitute a "
    "detailed hydrological model. Key limitations include: (1) the rainfall-runoff "
    "simulation uses a unit hydrograph approach rather than physically based models "
    "such as GSFLOW or ParFlow; (2) groundwater-surface water interactions are "
    "represented parametrically rather than through coupled simulation; (3) land "
    "subsidence risks from cyclical aquifer drawdown require site-specific geotechnical "
    "assessment; and (4) infiltration rates during flood events may be limited by soil "
    "saturation and surface sealing. Future work should include coupled groundwater-surface "
    "water modelling (e.g. using GSFLOW), pilot field studies, and institutional analysis "
    "of the regulatory frameworks needed to implement managed aquifer drawdown for flood "
    "control purposes."
)

add_normal(
    "The international applicability discussion is necessarily preliminary. Each "
    "transboundary basin presents unique hydrogeological, political, and institutional "
    "conditions that would require detailed case-specific analysis. The framework's "
    "applicability to low-gradient settings (e.g. Bangladesh, the Netherlands) and "
    "arid regions (e.g. Saudi Arabia) requires further feasibility assessment."
)

add_normal(
    "From a governance perspective, the implementation of managed aquifer drawdown for "
    "flood control would require new regulatory frameworks in most jurisdictions. "
    "Groundwater pumping is typically regulated for water supply purposes, and the concept "
    "of deliberately drawing down aquifers to create flood storage capacity does not fit "
    "neatly within existing regulatory categories. In Japan, groundwater extraction is "
    "regulated at the prefectural level through various ordinances originally designed to "
    "prevent land subsidence \u2014 these would need to be amended to permit temporary, cyclical "
    "drawdown for flood control purposes. Similar regulatory challenges would arise in "
    "other jurisdictions."
)

# ============================================================
# 8. CONCLUSIONS
# ============================================================
add_heading1("Conclusions")

add_normal(
    "This paper has proposed and analysed a novel flood control framework that integrates planned-release "
    "hydropower with inter-watershed groundwater management, and has situated this technical "
    "proposal within the broader politics of transboundary flood governance. The key findings "
    "are as follows:"
)

add_normal(
    "First, the framework is energy self-sufficient under the conditions analysed: "
    "hydropower generation from a 100 m head dam release at 50 m³/s produces 41.7 MW, "
    "of which only 3.68 MW (8.8%) is required for groundwater pumping."
)

add_normal(
    "Second, pre-emptive aquifer drawdown can create storage capacity equivalent to 78% "
    "of the observed 2018 flood volume in the Oda River–Takahashi River basin, and the "
    "Tokyo metropolitan application suggests 52.5 × 10⁶ m³ of combined surface-subsurface "
    "storage capacity."
)

add_normal(
    "Third, the framework can be implemented unilaterally by downstream states, "
    "circumventing the governance challenges that constrain transboundary flood management. "
    "This represents a new dimension in the hydro-hegemony literature: a technologically "
    "grounded counter-hegemonic strategy for flood risk reduction."
)

add_normal(
    "Fourth, the framework has potential applicability to transboundary river basins "
    "(Mekong, Nile, GBM, Indus, Tigris-Euphrates, Jordan), deltaic regions facing "
    "the dual threat of upstream floods and sea-level rise, and arid environments where "
    "flash flood energy could be captured for aquifer replenishment."
)

add_normal(
    "Flood management is not merely a technical problem but a political one. The "
    "framework proposed here does not resolve the power asymmetries inherent in "
    "transboundary water governance, but it offers downstream actors a concrete "
    "technological pathway to reduce their flood vulnerability independently of "
    "upstream decisions. As climate change intensifies flood risks in transboundary "
    "basins worldwide, such autonomous adaptation strategies deserve serious "
    "consideration in both technical and policy domains."
)

add_normal(
    "Finally, while this paper has focused on the technical and political dimensions of "
    "the proposed framework, successful implementation would also require engagement with "
    "affected communities, particularly regarding groundwater drawdown and potential "
    "(temporary) effects on well water levels and land stability. The history of groundwater "
    "management in Japan — from the uncontrolled extraction that caused severe subsidence to "
    "the stringent regulations that enabled recovery — offers both cautionary lessons and a "
    "model for adaptive governance. The challenge for the next generation of water "
    "management is to harness these hard-won lessons in service of climate adaptation, "
    "while respecting the legitimate concerns of those who live above the aquifers we "
    "propose to manage."
)

# ============================================================
# REFERENCES (Harvard/author-date style, WaA format)
# ============================================================
add_heading1("References")

refs = [
    "Berkowitz, J.F.; Pietroski, J.P.; Guernsey, C.M. and Koons, B.W. 2019. The Yazoo Backwater area: Hydrology, ecology, and flood management challenges. *Wetlands* 39(4): 723-736.",
    "California DWR (Department of Water Resources). 2018. *Flood-MAR: Using flood water for managed aquifer recharge to support sustainable water resources*. Sacramento, CA: California Department of Water Resources.",
    "Cascão, A.E. and Zeitoun, M. 2010. Power, hegemony and critical hydropolitics. In Earle, A.; Jägerskog, A. and Öjendal, J. (Eds), *Transboundary water management: Principles and practice*, pp. 27-42. London: Earthscan.",
    "Chang, L.-C.; Ho, C.-C.; Yeh, Y.-L. and Chen, Y.-W. 2019. An integrating approach for conjunctive-use planning of surface and subsurface water system. *Water Resources Management* 25(1): 59-78.",
    "Dillon, P.; Stuyfzand, P.; Grischek, T.; Lluria, M.; Pyne, R.D.G.; Jain, R.C.; Bear, J.; Schwarz, J.; Wang, W.; Fernandez, E.; Stefan, C.; Pettenati, M.; van der Gun, J.; Sprenger, C.; Massmann, G.; Scanlon, B.R.; Xanke, J.; Jokela, P.; Zheng, Y.; Rossetto, R.; Shamrukh, M.; Pavelic, P.; Murray, E.; Ross, A.; Bonilla Valverde, J.P.; Palma Nava, A.; Ansems, N.; Posavec, K.; Ha, K.; Martin, R. and Sapber, M. 2019. Sixty years of global progress in managed aquifer recharge. *Hydrogeology Journal* 27(1): 1-30.",
    "GSJ/AIST (Geological Survey of Japan). 2020. *Three-dimensional geological map of the Kanto Plain*. Tsukuba: National Institute of Advanced Industrial Science and Technology.",
    "IEA (International Energy Agency). 2021. *Hydropower special market report: Analysis and forecast to 2030*. Paris: International Energy Agency.",
    "KSB (San-yo Broadcasting). 2025. Verification of the Oda River confluence relocation: Effect of the completed Mabi flood control project. KSB News, 3 July 2025.",
    "Kuenzer, C.; Guo, H.; Huth, J.; Leinenkugel, P.; Li, X. and Dech, S. 2013. Flood mapping and flood dynamics of the Mekong Delta: ENVISAT-ASAR-WSM based time series analyses. *Remote Sensing* 5(2): 687-715.",
    "McCaffrey, S.C. 2007. *The law of international watercourses*. 2nd ed. Oxford: Oxford University Press.",
    "Lebel, L.; Garden, P. and Imamura, M. 2005. The politics of scale, position, and place in the governance of water resources in the Mekong region. *Ecology and Society* 10(2): 18, www.ecologyandsociety.org/vol10/iss2/art18/",
    "Mekong River Commission. 2023. *State of the basin report 2023*. Vientiane: Mekong River Commission Secretariat.",
    "Middleton, C. 2022. The political ecology of large hydropower dams in the Mekong basin. In Middleton, C. and Lamb, V. (Eds), *Knowing the Salween River: Resource politics of a contested transboundary river*, pp. 215-238. Cham: Springer.",
    "Mirumachi, N. and Allan, J.A. 2007. Revisiting transboundary water governance: Power, conflict, cooperation and the political economy. In Proceedings of the CAIWA International Conference on Adaptive and Integrated Water Management, Basel, Switzerland, 12-15 November 2007.",
    "MLIT (Ministry of Land, Infrastructure, Transport and Tourism). 2006. *Metropolitan Area Outer Underground Discharge Channel: Project overview*. Tokyo: MLIT Kanto Regional Development Bureau.",
    "MLIT. 2020a. *Arakawa River basin flood control plan*. Tokyo: MLIT Kanto Regional Development Bureau.",
    "MLIT. 2020b. *Flood inundation area maps: Arakawa River system*. Tokyo: MLIT.",
    "MLIT. 2021. *Basin Flood Control Project: Basic policy and implementation guidelines*. Tokyo: MLIT.",
    "MLIT. 2024. *Oda River–Takahashi River confluence relocation project: Completion report*. Tokyo: MLIT Chugoku Regional Development Bureau.",
    "MLIT Chugoku Regional Development Bureau. 2019. *Oda River improvement project: Post-disaster review and future plans*. Hiroshima: MLIT.",
    "NIED (National Research Institute for Earth Science and Disaster Resilience). 2018. *Overview of damage caused by the July 2018 Heavy Rainfall*. Tsukuba: NIED.",
    "Nishizawa, T. and Zhang, J. 2026. Half-century groundwater monitoring in central Japan: Implications for managed aquifer recharge. *Groundwater for Sustainable Development* 18: 100782.",
    "Räsänen, T.A.; Someth, P.; Lauri, H.; Koponen, J.; Sarkkula, J. and Kummu, M. 2017. Observed river discharge changes due to hydropower operations in the Upper Mekong Basin. *Journal of Hydrology* 545: 28-41.",
    "TEPCO (Tokyo Electric Power Company Holdings). 2023. *Pumped-storage hydropower facilities: Technical overview*. Tokyo: TEPCO.",
    "Wheeler, K.G.; Jeuland, M.; Hall, J.W.; Zagona, E. and Whittington, D. 2020. Understanding and managing new risks on the Nile with the Grand Ethiopian Renaissance Dam. *Nature Communications* 11: 5222.",
    "Zeitoun, M. 2008. *Power and water in the Middle East: The hidden politics of the Palestinian-Israeli water conflict*. London: I.B. Tauris.",
    "Zeitoun, M. and Warner, J. 2006. Hydro-hegemony: A framework for analysis of transboundary water conflicts. *Water Policy* 8(5): 435-460.",
]

for ref in refs:
    add_biblio(ref)


# ── Save ──
outpath = os.path.join(OUTDIR, "waa_manuscript.docx")
doc.save(outpath)
print(f"WaA manuscript saved: {outpath}")

# Word count estimate
total_words = 0
in_refs = False
in_abstract = False
for p in doc.paragraphs:
    txt = p.text.strip()
    if 'References' in txt and p.style and 'Heading' in p.style.name:
        in_refs = True
    if 'Abstract' in txt and p.style and 'Heading' in p.style.name:
        in_abstract = True
        continue
    if in_abstract and p.style and 'Heading' in p.style.name and 'Abstract' not in txt:
        in_abstract = False
    if not in_refs and not in_abstract:
        total_words += len(txt.split())

print(f"Estimated word count (excl. abstract and references): ~{total_words}")
