#!/usr/bin/env python3
"""Generate DOCX concept document for Clinical Noise Inverse Problem with inline figures."""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_superscript_text(paragraph, text):
    """Parse text with {ref} markers and create runs with superscript."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
        else:
            paragraph.add_run(part)


def create_docx():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Title
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        'Clinical Noise Inverse Problem:\n'
        'A Framework for Decomposing Treatment Effects and Confounders in Medical Research'
    )
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    # Abstract
    doc.add_heading('Abstract', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Clinical research fundamentally seeks to separate the true effect of an intervention '
        'from confounders, measurement error, and random biological variation. We propose the '
        'Clinical Noise Inverse Problem (CNIP) framework, which reconceptualizes confounders '
        'and random variation as structured, modelable "noise" (N) that can be explicitly '
        'reconstructed and subtracted from observed outcomes, leaving a cleaner estimate of '
        'the treatment signal (S). The framework consists of a four-phase pipeline: '
        '(1) noise forward model construction from auxiliary clinical data, '
        '(2) Bayesian noise parameter estimation, '
        '(3) residual generation for clean signal extraction, and '
        '(4) hypothesis-free discovery in residuals. '
        'We describe applications to both randomized controlled trials (RCTs) — where noise '
        'modeling can effectively increase sample size by a factor of 1/(1−ρ²) — and '
        'retrospective observational studies, where explicit noise modeling complements '
        'existing causal inference methods. This concept document outlines the theoretical '
        'foundation, practical pipeline, and potential impact of the CNIP approach.'
    )

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Every observed clinical outcome is a composite of the true treatment effect and '
        'numerous sources of variation — confounders, measurement error, inter-individual '
        'biological variability, and chance fluctuation. Traditional statistical methods '
        '(regression adjustment, propensity score matching, randomization) address these '
        'sources separately. We propose a unified conceptual framework — the '
    )
    run = p.add_run('Clinical Noise Inverse Problem (CNIP)')
    run.bold = True
    p.add_run(
        ' — that treats confounders and random variation as structured noise that can be '
        'explicitly modeled and subtracted from observed data.'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run('The key insight is: if we can model and reproduce the noise, we can remove it.')
    run.bold = True
    p.add_run(
        ' Rather than relying solely on averaging over large samples to dilute random '
        'variation, or on post hoc adjustment for known confounders, the noise inverse '
        'approach builds an explicit forward model of how noise generates observed data, '
        'then solves the inverse problem to estimate and subtract the noise component.'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('This framework has the potential to:')

    benefits = [
        'Reduce required sample sizes in RCTs by improving signal-to-noise ratio',
        'Strengthen causal inference in retrospective studies by explicitly modeling confounding structure',
        'Enable hypothesis-free discovery by examining residuals after noise removal',
        'Unify existing methods (propensity scores, mixed models, instrumental variables) '
        'under a single information-theoretic framework',
    ]
    for b in benefits:
        bp = doc.add_paragraph(b, style='List Bullet')

    # 2. Fundamental Formulation
    doc.add_heading('2. Fundamental Formulation', level=1)

    doc.add_heading('2.1 The Observation Equation', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'For any clinical study, the observed outcome for patient '
    )
    run = p.add_run('i')
    run.italic = True
    p.add_run(' can be written:')

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq.add_run('Y_obs(i) = S(i) + N(i)')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('where ')
    run = p.add_run('Y_obs(i)')
    run.bold = True
    p.add_run(' is the observed outcome (primary endpoint), ')
    run = p.add_run('S(i)')
    run.bold = True
    p.add_run(' is the true treatment effect (signal), and ')
    run = p.add_run('N(i)')
    run.bold = True
    p.add_run(' is the aggregate noise from all non-treatment sources of variation.')

    doc.add_heading('2.2 Noise Decomposition', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('The noise term N(i) can be further decomposed into four components:')

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq.add_run('N(i) = N_conf(i) + N_bio(i) + N_meas(i) + N_rand(i)')
    run.bold = True
    run.font.size = Pt(12)

    # --- Fig 1 inline ---
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Figure 1 illustrates the conceptual framework, showing the decomposition of '
        'observed outcomes into signal and noise components, with the four noise subcategories '
        'and their respective modelability.'
    )

    # Insert Fig 1
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig1_conceptual_framework.png'):
        doc.add_picture('fig1_conceptual_framework.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('Figure 1. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        'Clinical Noise Inverse Problem — Conceptual Framework. '
        'The observed outcome Y_obs(i) is decomposed into signal S(i) (treatment effect) '
        'and noise N(i) (confounders, biological variability, measurement error, '
        'and irreducible randomness).'
    )
    run.font.size = Pt(10)

    # Table 1: Noise components
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('Table 1 summarizes the four noise components, their characteristics, and modelability.')

    p_tcap = doc.add_paragraph()
    p_tcap.paragraph_format.space_before = Pt(14)
    run = p_tcap.add_run('Table 1. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_tcap.add_run('Noise component decomposition in clinical research.')
    run.font.size = Pt(10)

    table = doc.add_table(rows=5, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Component', 'Description', 'Examples', 'Modelability']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    rows_data = [
        ['N_conf', 'Measured confounders', 'Age, sex, BMI, comorbidities, medications', 'High'],
        ['N_bio', 'Biological variability', 'Genetic polymorphisms, circadian variation, disease stage', 'Medium'],
        ['N_meas', 'Measurement error', 'Instrument precision, inter-observer variability', 'Medium'],
        ['N_rand', 'Irreducible randomness', 'Stochastic biological processes', 'Low'],
    ]
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)

    doc.add_heading('2.3 The Inverse Problem', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'The forward problem is: given noise parameters θ, predict the noise contribution '
        'N(i; θ). The inverse problem is: given observed data and a treatment model, '
        'estimate the noise parameters:'
    )

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq.add_run('θ̂ = argmin_θ  D( Y_obs,  S_model + F(θ) )')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'The residual after noise subtraction yields a cleaner estimate of the signal: '
        'Ŝ(i) = Y_obs(i) − F(θ̂; X_i), where X_i represents auxiliary data for patient i.'
    )

    # 3. The Four-Phase Pipeline
    doc.add_heading('3. The Four-Phase Clinical Noise Inverse Pipeline', level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Figure 2 illustrates the complete four-phase pipeline. The pipeline takes clinical '
        'outcome data, auxiliary patient data, and domain knowledge as inputs, and produces '
        'a clean treatment effect estimate with uncertainty quantification.'
    )

    # Insert Fig 2
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig2_pipeline.png'):
        doc.add_picture('fig2_pipeline.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('Figure 2. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        'Four-Phase Clinical Noise Inverse Pipeline. '
        'Phase 1: noise forward model construction; Phase 2: Bayesian parameter estimation; '
        'Phase 3: residual generation; Phase 4: hypothesis-free discovery. '
        'Iterative refinement between Phase 3 and Phase 2 improves noise estimates.'
    )
    run.font.size = Pt(10)

    # Phase descriptions
    phases = [
        ('Phase 1: Noise Forward Model Construction',
         'The first phase builds an explicit model F(θ; X) that predicts the noise '
         'contribution from measurable sources. Inputs include patient-level covariates '
         '(demographics, comorbidities, medications), temporal auxiliary data (serial lab '
         'values, vital signs), and institutional variables (site, operator, device). '
         'Methods range from parametric (GLM, structural equation models) to non-parametric '
         '(gradient boosting, neural networks with interpretability constraints). Variable '
         'selection is guided by causal directed acyclic graphs (DAGs) to ensure only '
         'non-treatment sources are modeled.'),
        ('Phase 2: Bayesian Noise Parameter Estimation',
         'The second phase estimates noise parameters θ from observed data. Bayesian inference '
         '— via MCMC, variational inference, or amortized neural network approaches — provides '
         'full posterior distributions p(θ | Y_obs), enabling uncertainty quantification. '
         'Cross-validation and calibration checks ensure the estimates are reliable.'),
        ('Phase 3: Residual Generation (Clean Signal Extraction)',
         'The third phase subtracts estimated noise to obtain cleaner treatment effect estimates: '
         'Ŝ(i) = Y_obs(i) − F(θ̂; X_i). For continuous outcomes, direct subtraction suffices. '
         'For count or event data, probabilistic thinning removes noise-attributed events. '
         'Doubly-robust estimation combines the noise model with an outcome model for '
         'robustness to partial misspecification. Quality is assessed by residual variance '
         'reduction, independence from measured confounders, and signal preservation tests.'),
        ('Phase 4: Hypothesis-Free Discovery in Residuals',
         'After removing known noise, the fourth phase examines residuals for unexpected patterns: '
         'treatment effect heterogeneity, novel subgroup effects, or biomarker signals. '
         'Methods include anomaly detection, data-driven subgroup discovery, and temporal pattern '
         'analysis. Safeguards include FDR control, pre-registration, and replication requirements.'),
    ]

    for title, desc in phases:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.add_run(desc)

    # 4. Application to RCTs
    doc.add_heading('4. Application to Randomized Controlled Trials', level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Randomization eliminates systematic confounding on average, but individual-level '
        'variation remains substantial. In a typical RCT, between-patient variability often '
        'dominates the treatment effect, and prognostic covariates (baseline disease severity, '
        'biomarkers) create substantial noise that standard analysis does not fully exploit.'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'The CNIP approach exploits auxiliary data routinely collected but underutilized in '
        'standard RCT analysis: baseline covariates, serial biomarkers, concomitant medications, '
        'and site/operator variables. If the noise model explains a fraction ρ'
    )
    run = p.add_run('2')
    run.font.superscript = True
    p.add_run(' of outcome variance, the effective sample size increases by a factor of 1/(1 − ρ')
    run = p.add_run('2')
    run.font.superscript = True
    p.add_run(').')

    # Insert Fig 3
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Figure 3 compares the CNIP application in RCT and retrospective settings, '
        'highlighting the different auxiliary channels and expected benefits in each context.'
    )

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig3_rct_vs_retrospective.png'):
        doc.add_picture('fig3_rct_vs_retrospective.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('Figure 3. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        'CNIP application comparison. Left: RCT setting with randomization-balanced confounders '
        'and underutilized auxiliary channels. Right: retrospective/observational setting with '
        'rich EHR-derived auxiliary data and entangled confounders.'
    )
    run.font.size = Pt(10)

    # Sample size section
    doc.add_heading('4.1 Sample Size Implications', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('If baseline covariates and serial biomarkers explain 40% of outcome variance (ρ')
    run = p.add_run('2')
    run.font.superscript = True
    p.add_run(' = 0.4), the effective sample size becomes n_eff = n_actual / 0.6 = 1.67 × n_actual. '
              'A trial of 600 patients with CNIP achieves the statistical power of a 1000-patient '
              'conventional trial (Figure 4).')

    # Insert Fig 4
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig4_sample_size_reduction.png'):
        doc.add_picture('fig4_sample_size_reduction.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('Figure 4. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        'Sample size implications of noise modeling. Left: effective sample size multiplier '
        'as a function of noise model accuracy ρ². Right: illustrative power curves showing '
        'earlier achievement of 80% power with CNIP versus conventional analysis.'
    )
    run.font.size = Pt(10)

    # 5. Application to Retrospective Studies
    doc.add_heading('5. Application to Retrospective Studies', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'In retrospective studies, treatment assignment is non-random and confounders are '
        'entangled with the treatment decision. The CNIP approach explicitly models the noise '
        'structure using the rich auxiliary data available in electronic health records (EHRs): '
        'serial lab values, continuous vital signs, timestamped medication records, and '
        'NLP-extracted features from clinical notes.'
    )

    # Table 2: Comparison with existing methods
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Table 2 compares CNIP with existing causal inference methods for observational studies.'
    )

    p_tcap = doc.add_paragraph()
    p_tcap.paragraph_format.space_before = Pt(14)
    run = p_tcap.add_run('Table 2. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_tcap.add_run('Comparison of CNIP with existing causal inference methods.')
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=6, cols=4, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ['Method', 'Noise Model', 'Strengths', 'Limitations']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    rows2 = [
        ['Propensity score', 'P(T|X)', 'Intuitive, well-validated',
         'Ignores outcome model'],
        ['Instrumental variables', 'Exogenous variation', 'Addresses unmeasured confounding',
         'Requires valid instrument'],
        ['Difference-in-differences', 'Time trends', 'Controls time-invariant confounders',
         'Parallel trends assumption'],
        ['Regression discontinuity', 'Threshold rules', 'Quasi-experimental rigor',
         'Local estimate only'],
        ['CNIP', 'Explicit F(θ; X)', 'Unified framework; residual discovery; flexible',
         'Requires rich auxiliary data'],
    ]
    for r_idx, row_data in enumerate(rows2):
        for c_idx, cell_text in enumerate(row_data):
            cell = table2.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)
            if r_idx == 4:
                run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'The CNIP framework does not replace these methods but provides a complementary lens. '
        'In particular, CNIP can be combined with propensity scores (doubly-robust estimation) '
        'or used in settings where instrumental variables are unavailable.'
    )

    # 6. Methodological Considerations
    doc.add_heading('6. Methodological Considerations', level=1)

    doc.add_heading('6.1 Identifiability', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'The decomposition Y = S + N is not automatically identifiable. Three assumptions '
        'are required: (1) conditional independence — given auxiliary data X, noise N is '
        'independent of treatment assignment T (satisfied by design in RCTs; requires careful '
        'modeling in observational studies); (2) model separability — the noise forward model '
        'does not absorb the treatment signal; (3) auxiliary sufficiency — the auxiliary data '
        'is rich enough to capture the dominant noise structure.'
    )

    doc.add_heading('6.2 Model Validation Protocol', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('A rigorous validation protocol includes five steps:')

    validation_steps = [
        'Noise model accuracy: cross-validated R² or AUC significantly above null model',
        'Residual balance: |r| < 0.05 between residuals and all measured confounders',
        'Signal preservation: known treatment effects recovered within 95% CI',
        'Placebo test: no spurious signal detected in placebo arm or pre-treatment period',
        'Sensitivity analysis: conclusions robust to reasonable model perturbations',
    ]
    for v in validation_steps:
        doc.add_paragraph(v, style='List Number')

    doc.add_heading('6.3 Ethical and Regulatory Considerations', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'The noise model and its assumptions must be fully documented and reproducible. '
        'For confirmatory analyses, noise model specifications should be pre-registered. '
        'Novel statistical frameworks require validation against established methods before '
        'regulatory acceptance. Rich auxiliary data raises privacy concerns that may require '
        'federated learning or differential privacy approaches.'
    )

    # 7. Potential Impact
    doc.add_heading('7. Potential Impact and Future Directions', level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run('Near-term applications')
    run.bold = True
    p.add_run(
        ' include integration into adaptive trial designs for sample size re-estimation, '
        'application to EHR-based comparative effectiveness research, and meta-analysis '
        'enhancement through pooling of residuals across studies.'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run('Longer-term directions')
    run.bold = True
    p.add_run(
        ' include personalized noise models from wearable sensors and longitudinal EHR data, '
        'real-time clinical decision support, and potential for regulatory acceptance of '
        'CNIP-based trials with smaller sample size requirements.'
    )

    # 8. Summary
    doc.add_heading('8. Summary', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'The Clinical Noise Inverse Problem framework reconceptualizes confounders and random '
        'variation as structured noise that can be explicitly modeled and subtracted. By building '
        'forward models of noise from auxiliary clinical data and solving the inverse problem, '
        'we can extract cleaner treatment signals — potentially enabling faster conclusions in '
        'RCTs, stronger causal inference in retrospective studies, and hypothesis-free discovery '
        'of treatment effect heterogeneity. The framework provides a practical four-phase pipeline '
        'applicable to a wide range of clinical research settings.'
    )

    out = 'clinical_noise_inverse_problem.docx'
    doc.save(out)
    print(f'Saved: {out}')
    return out


if __name__ == '__main__':
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    create_docx()
