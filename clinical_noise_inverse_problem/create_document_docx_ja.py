#!/usr/bin/env python3
"""Generate Japanese DOCX concept document for Clinical Noise Inverse Problem with inline figures."""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_superscript_text(paragraph, text):
    """Parse text with {ref} markers and create runs with superscript."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
        else:
            paragraph.add_run(part)


def create_docx_ja():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Title
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        '臨床ノイズ逆問題（CNIP）:\n'
        '医療研究における治療効果と交絡因子の分解フレームワーク'
    )
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    # Abstract
    doc.add_heading('要旨', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '臨床研究の根本的な課題は、介入の真の効果を交絡因子、測定誤差、ランダムな生物学的変動'
        'から分離することである。本概念文書では、臨床ノイズ逆問題（Clinical Noise Inverse '
        'Problem; CNIP）フレームワークを提案する。CNIPは交絡因子と偶然変動を構造化された'
        'モデル化可能な「ノイズ」（N）として再定義し、観察データから明示的にノイズを再構成・'
        '差引することで、より高精度な治療シグナル（S）の推定を可能にする。フレームワークは'
        '4フェーズのパイプラインで構成される: '
        '(1) 補助臨床データからのノイズ前方モデル構築、'
        '(2) ベイズノイズパラメータ推定、'
        '(3) 残差生成によるクリーンシグナル抽出、'
        '(4) 残差中の仮説フリー発見。'
        'ランダム化比較試験（RCT）ではノイズモデリングにより有効サンプルサイズを1/(1−ρ')
    run = p.add_run('2')
    run.font.superscript = True
    p.add_run(
        ')倍に増加させ得ること、また後ろ向き観察研究では既存の因果推論手法を補完する'
        '明示的ノイズモデリングが可能であることを示す。'
    )

    # 1. はじめに
    doc.add_heading('1. はじめに', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '臨床研究で観察されるすべてのアウトカムは、真の治療効果と多様な変動源 — 交絡因子、'
        '測定誤差、個人間の生物学的変動、偶然変動 — の合成物である。従来の統計手法（回帰調整、'
        '傾向スコアマッチング、ランダム化）はこれらの変動源を個別に扱ってきた。本文書では、'
        '統一的な概念フレームワーク — '
    )
    run = p.add_run('臨床ノイズ逆問題（CNIP）')
    run.bold = True
    p.add_run(
        ' — を提案する。CNIPは交絡因子と偶然変動を構造化されたノイズとして扱い、'
        '観察データから明示的にモデル化・差引することを目指す。'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run('核心的洞察: ノイズをモデル化・再現できれば、それを除去できる。')
    run.bold = True
    p.add_run(
        '大標本平均でランダム変動を希釈したり、事後的に既知の交絡因子を調整するだけでなく、'
        'ノイズ逆問題アプローチは、ノイズがどのように観察データを生成するかの明示的な前方モデルを'
        '構築し、逆問題を解くことでノイズ成分を推定・差引する。'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('本フレームワークの潜在的効果:')

    benefits = [
        'RCTにおいてシグナル対ノイズ比の改善により必要サンプルサイズを縮小',
        '後ろ向き研究において交絡構造の明示的モデリングにより因果推論を強化',
        'ノイズ除去後の残差検証による仮説フリーの発見を実現',
        '既存手法（傾向スコア、混合モデル、操作変数法）を情報理論的フレームワークで統合',
    ]
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    # 2. 基本定式化
    doc.add_heading('2. 基本定式化', level=1)

    doc.add_heading('2.1 観測方程式', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('臨床研究において、患者')
    run = p.add_run('i')
    run.italic = True
    p.add_run('の観察されたアウトカムは以下のように表現できる:')

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq.add_run('Y_obs(i) = S(i) + N(i)')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('ここで、')
    run = p.add_run('Y_obs(i)')
    run.bold = True
    p.add_run('は観察されたアウトカム（主エンドポイント）、')
    run = p.add_run('S(i)')
    run.bold = True
    p.add_run('は真の治療効果（シグナル）、')
    run = p.add_run('N(i)')
    run.bold = True
    p.add_run('は治療以外のすべての変動源から成る総合ノイズである。')

    doc.add_heading('2.2 ノイズの分解', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('ノイズ項 N(i) は4つの成分に分解できる:')

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq.add_run('N(i) = N_conf(i) + N_bio(i) + N_meas(i) + N_rand(i)')
    run.bold = True
    run.font.size = Pt(12)

    # --- 図1 インライン ---
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '図1に概念フレームワークを示す。観察アウトカムのシグナルとノイズ成分への分解、'
        '4つのノイズサブカテゴリーとそれぞれのモデル化可能性を図示している。'
    )

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig1_conceptual_framework_ja.png'):
        doc.add_picture('fig1_conceptual_framework_ja.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('図1. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        '臨床ノイズ逆問題 — 概念フレームワーク。'
        '観察アウトカム Y_obs(i) をシグナル S(i)（治療効果）とノイズ N(i)'
        '（交絡因子、生物学的変動、測定誤差、不可避のランダムネス）に分解する。'
    )
    run.font.size = Pt(10)

    # 表1: ノイズ成分
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('表1に4つのノイズ成分の特性とモデル化可能性をまとめる。')

    p_tcap = doc.add_paragraph()
    p_tcap.paragraph_format.space_before = Pt(14)
    run = p_tcap.add_run('表1. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_tcap.add_run('臨床研究におけるノイズ成分の分解。')
    run.font.size = Pt(10)

    table = doc.add_table(rows=5, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['成分', '説明', '具体例', 'モデル化可能性']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    rows_data = [
        ['N_conf', '測定可能な交絡因子', '年齢、性別、BMI、合併症、投薬', '高'],
        ['N_bio', '生物学的変動', '遺伝的多型、概日リズム、疾患ステージ', '中'],
        ['N_meas', '測定誤差', '機器精度、観察者間変動', '中'],
        ['N_rand', '不可避のランダムネス', '確率的生物学的プロセス', '低'],
    ]
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)

    doc.add_heading('2.3 逆問題', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '順問題はノイズパラメータθが与えられたときにノイズ寄与N(i; θ)を予測すること、'
        '逆問題は観察データと治療モデルからノイズパラメータを推定することである:'
    )

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq.add_run('θ̂ = argmin_θ  D( Y_obs,  S_model + F(θ) )')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'ノイズ差引後の残差がシグナルのよりクリーンな推定値となる: '
        'Ŝ(i) = Y_obs(i) − F(θ̂; X_i)。ここでX_iは患者iの補助データを表す。'
    )

    # 3. 4フェーズパイプライン
    doc.add_heading('3. 4フェーズ臨床ノイズ逆問題パイプライン', level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '図2に完全な4フェーズパイプラインを示す。パイプラインは臨床アウトカムデータ、'
        '補助患者データ、ドメイン知識を入力とし、不確実性の定量化を伴うクリーンな'
        '治療効果推定値を出力する。'
    )

    # 図2 インライン
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig2_pipeline_ja.png'):
        doc.add_picture('fig2_pipeline_ja.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('図2. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        '4フェーズ臨床ノイズ逆問題パイプライン。'
        'Phase 1: ノイズ前方モデル構築; Phase 2: ベイズパラメータ推定; '
        'Phase 3: 残差生成; Phase 4: 仮説フリー発見。'
        'Phase 3とPhase 2間の反復精緻化によりノイズ推定が改善される。'
    )
    run.font.size = Pt(10)

    # フェーズの説明
    phases = [
        ('Phase 1: ノイズ前方モデル構築',
         '第1フェーズでは、測定可能なソースからノイズ寄与を予測する明示的モデル F(θ; X) を'
         '構築する。入力には患者レベルの共変量（人口統計、合併症、投薬）、時系列補助データ'
         '（経時的検査値、バイタルサイン）、施設変数（施設、術者、機器）が含まれる。'
         '手法はパラメトリック（GLM、構造方程式モデル）からノンパラメトリック'
         '（勾配ブースティング、解釈可能性制約付きニューラルネットワーク）まで多岐にわたる。'
         '変数選択は因果DAG（有向非巡回グラフ）にガイドされ、治療効果以外のソースのみが'
         'モデル化されることを保証する。'),
        ('Phase 2: ベイズノイズパラメータ推定',
         '第2フェーズでは、観察データからノイズパラメータθを推定する。ベイズ推論 — '
         'MCMC、変分推論、または償却型ニューラルネットワークアプローチ — により、'
         '完全な事後分布 p(θ | Y_obs) を得て不確実性の定量化を可能にする。'
         '交差検証とキャリブレーションチェックにより推定値の信頼性を確保する。'),
        ('Phase 3: 残差生成（クリーンシグナル抽出）',
         '第3フェーズでは、推定ノイズを差し引いてよりクリーンな治療効果推定値を得る: '
         'Ŝ(i) = Y_obs(i) − F(θ̂; X_i)。連続アウトカムには直接差引で十分であり、'
         'カウント/イベントデータには確率的間引きでノイズ起因のイベントを除去する。'
         '二重頑健推定はノイズモデルとアウトカムモデルを組み合わせ、部分的な'
         'モデル誤特定に対する頑健性を提供する。品質は残差分散減少、測定交絡因子'
         'からの独立性、シグナル保存テストで評価される。'),
        ('Phase 4: 残差中の仮説フリー発見',
         '既知のノイズを除去した後、第4フェーズでは残差中の予期しないパターン — '
         '治療効果の異質性、新規サブグループ効果、バイオマーカーシグナル — を探索する。'
         '手法には異常検知、データ駆動型サブグループ発見、時間的パターン分析が含まれる。'
         'セーフガードとしてFDR制御、事前登録、独立コホートでの再現が必要である。'),
    ]

    for title_text, desc in phases:
        doc.add_heading(title_text, level=2)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.add_run(desc)

    # 4. RCTへの応用
    doc.add_heading('4. ランダム化比較試験（RCT）への応用', level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'ランダム化は平均的に系統的交絡を排除するが、個人レベルの変動は依然として大きい。'
        '典型的なRCTでは、患者間変動が治療効果を上回ることが多く、予後共変量'
        '（ベースラインの疾患重症度、バイオマーカー）が実質的なノイズを生む。'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'CNIPアプローチは、標準的なRCT解析では活用不足の補助データを利用する: '
        'ベースライン共変量、経時的バイオマーカー、併用薬、施設/術者変数。'
        'ノイズモデルがアウトカム分散のρ'
    )
    run = p.add_run('2')
    run.font.superscript = True
    p.add_run('の割合を説明する場合、有効サンプルサイズは1/(1 − ρ')
    run = p.add_run('2')
    run.font.superscript = True
    p.add_run(')倍に増加する。')

    # 図3 インライン
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '図3にRCTと後ろ向き研究におけるCNIP適用の比較を示す。'
        '各セッティングにおける補助チャンネルと期待される利点を対比する。'
    )

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig3_rct_vs_retrospective_ja.png'):
        doc.add_picture('fig3_rct_vs_retrospective_ja.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('図3. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        'CNIP適用比較。左: ランダム化により交絡が均衡したRCTセッティング。'
        '右: 交絡が治療と絡み合う後ろ向き/観察研究セッティング。'
    )
    run.font.size = Pt(10)

    # サンプルサイズ
    doc.add_heading('4.1 サンプルサイズへの含意', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('ベースライン共変量と経時的バイオマーカーがアウトカム分散の40%を説明する場合（ρ')
    run = p.add_run('2')
    run.font.superscript = True
    p.add_run(' = 0.4）、有効サンプルサイズはn_eff = n_actual / 0.6 = 1.67 × n_actualとなる。'
              'CNIPを用いた600例の試験は、従来法の1000例の試験と同等の統計的検出力を達成する（図4）。')

    # 図4 インライン
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    if os.path.exists('fig4_sample_size_reduction_ja.png'):
        doc.add_picture('fig4_sample_size_reduction_ja.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(14)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run('図4. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cap.add_run(
        'ノイズモデリングによるサンプルサイズへの含意。左: ノイズモデル精度ρ²に対する'
        '有効サンプルサイズ倍率。右: CNIPにより従来法と比較して80%検出力をより早く'
        '達成できることを示す概念的な検出力曲線。'
    )
    run.font.size = Pt(10)

    # 4.2 具体例
    doc.add_heading('4.2 適用例: 周術期アウトカム試験', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '新しい麻酔プロトコルの術後合併症に対するRCTを考える。主エンドポイント（S）は'
        '30日複合合併症率であり、ノイズ源（N）にはASA-PS分類、手術複雑度、年齢、BMI'
        '（N_conf）、術前炎症マーカー（CRP、IL-6）や遺伝的変異（CYP多型）'
        '（N_bio）、施設間の合併症判定のばらつき（N_meas）、確率的周術期イベント'
        '（N_rand）が含まれる。'
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('CNIPパイプラインは以下のように適用される:')

    steps = [
        'Phase 1: 治療以外の全変数から合併症リスクを予測する前方モデルを構築',
        'Phase 2: 帰無仮説下での個人レベルの予測合併症率を推定',
        'Phase 3: 残差 = 観察値 − 予測値 → プロトコル効果のクリーンな推定',
        'Phase 4: 残差が予期しない変数（手術時間帯、特定の手術サブタイプなど）でクラスタリングしていないか探索',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    # 5. 後ろ向き研究への応用
    doc.add_heading('5. 後ろ向き（観察）研究への応用', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '後ろ向き研究では治療割り当てはランダムでなく、交絡因子が治療決定と絡み合っている。'
        'CNIPアプローチは電子カルテ（EHR）に含まれる豊富な補助データを活用して'
        'ノイズ構造を明示的にモデル化する: 経時的検査値、連続バイタルサイン、'
        'タイムスタンプ付き投薬記録、臨床記録からのNLP抽出特徴量。'
    )

    # 表2: 既存手法との比較
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('表2にCNIPと既存の因果推論手法を比較する。')

    p_tcap = doc.add_paragraph()
    p_tcap.paragraph_format.space_before = Pt(14)
    run = p_tcap.add_run('表2. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_tcap.add_run('CNIPと既存の因果推論手法の比較。')
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=6, cols=4, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ['手法', 'ノイズモデル', '強み', '限界']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    rows2 = [
        ['傾向スコア', 'P(T|X)', '直感的、十分に検証済み', 'アウトカムモデルを無視'],
        ['操作変数法', '外生的変動', '未測定交絡に対処可能', '有効な操作変数が必要'],
        ['差分の差分法', '時間トレンド', '時間不変交絡を制御', '平行トレンド仮定'],
        ['回帰不連続デザイン', '閾値ルール', '準実験的厳密性', '局所推定のみ'],
        ['CNIP', '明示的F(θ; X)', '統一フレームワーク; 残差発見; 柔軟', '豊富な補助データが必要'],
    ]
    for r_idx, row_data in enumerate(rows2):
        for c_idx, cell_text in enumerate(row_data):
            cell = table2.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)
            if r_idx == 4:
                run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'CNIPフレームワークはこれらの手法を置き換えるものではなく、補完的な視点を提供する。'
        '特にCNIPは傾向スコアと組み合わせた二重頑健推定や、操作変数が利用できない'
        'セッティングで活用できる。'
    )

    # 6. 方法論的考察
    doc.add_heading('6. 方法論的考察', level=1)

    doc.add_heading('6.1 識別可能性', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'Y = S + N の分解は自動的に識別可能ではない。3つの仮定が必要である: '
        '(1) 条件付き独立性 — 補助データXが与えられたとき、ノイズNは治療割り当てTと独立'
        '（RCTでは設計上満足; 観察研究では慎重なモデリングが必要）; '
        '(2) モデル分離可能性 — ノイズ前方モデルが治療シグナルを吸収しない; '
        '(3) 補助十分性 — 補助データが支配的なノイズ構造を捕捉するに十分豊富である。'
    )

    doc.add_heading('6.2 モデル検証プロトコル', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run('厳密な検証プロトコルは5つのステップで構成される:')

    validation_steps = [
        'ノイズモデル精度: 交差検証によるR²またはAUCがヌルモデルを有意に上回る',
        '残差のバランス: 残差と全測定交絡因子の相関 |r| < 0.05',
        'シグナル保存: 既知の治療効果を持つデータに適用し、回復された効果が真の効果の95%信頼区間内',
        'プラセボテスト: プラセボ群または治療前期間で偽シグナルが検出されない',
        '感度分析: ノイズモデルの仮定に対する合理的な摂動にロバスト',
    ]
    for v in validation_steps:
        doc.add_paragraph(v, style='List Number')

    doc.add_heading('6.3 倫理的・規制的考察', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        'ノイズモデルとその仮定は完全に文書化され再現可能でなければならない。'
        '確証的解析ではノイズモデルの仕様を事前登録すべきである。'
        '新規統計フレームワークは規制上の受容に先立ち、確立された手法との検証が必要である。'
        '豊富な補助データはプライバシーの懸念を生じ、連合学習や差分プライバシーが'
        '必要となる場合がある。'
    )

    # 7. 潜在的影響と今後の方向性
    doc.add_heading('7. 潜在的影響と今後の方向性', level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run('近期的応用')
    run.bold = True
    p.add_run(
        'として、サンプルサイズ再推定のための適応的試験デザインへの統合、'
        'EHRベースの比較有効性研究への適用、研究間の残差プーリングによるメタ解析の強化がある。'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run('長期的方向性')
    run.bold = True
    p.add_run(
        'として、ウェアラブルセンサーや縦断的EHRデータからの個別化ノイズモデル、'
        'リアルタイム臨床意思決定支援、CNIPベースの試験におけるより小さな'
        'サンプルサイズでの規制承認の可能性がある。'
    )

    # 8. まとめ
    doc.add_heading('8. まとめ', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(
        '臨床ノイズ逆問題フレームワークは、交絡因子と偶然変動を構造化されたノイズとして'
        '再定義し、明示的にモデル化・差引できるものとして捉え直す。補助臨床データから'
        'ノイズの前方モデルを構築し逆問題を解くことで、よりクリーンな治療シグナルの抽出が'
        '可能となる — RCTにおけるより早い結論、後ろ向き研究におけるより強い因果推論、'
        'そして治療効果の異質性に関する仮説フリーの発見を潜在的に実現する。'
        '本フレームワークは幅広い臨床研究セッティングに適用可能な実践的な'
        '4フェーズパイプラインを提供する。'
    )

    out = 'clinical_noise_inverse_problem_ja.docx'
    doc.save(out)
    print(f'Saved: {out}')
    return out


if __name__ == '__main__':
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    create_docx_ja()
