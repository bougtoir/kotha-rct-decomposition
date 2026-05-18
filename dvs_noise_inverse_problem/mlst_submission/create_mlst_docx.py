#!/usr/bin/env python3
"""Generate MLST-formatted English DOCX with inline figures."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.dirname(BASE_DIR)

def add_caption(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_r = p.add_run(text)
        run_r.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_figure(doc, img_path, caption, fig_num):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        add_caption(doc, caption, f"Figure {fig_num}. ")
    else:
        add_caption(doc, f"[Figure {fig_num} not found: {img_path}]")

def make_table(doc, headers, rows, caption=None, tab_num=None):
    if caption and tab_num:
        add_caption(doc, caption, f"Table {tab_num}. ")
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacing

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(
        'Noise as Signal: A Systematic Review and Physics-Informed Framework '
        'for Solving the Noise Inverse Problem in Dynamic Vision Sensors '
        'Applied to Astronomical Observation'
    )
    run.bold = True
    run.font.size = Pt(16)

    # Authors placeholder
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.add_run('[Authors and Affiliations]').font.size = Pt(11)

    doc.add_paragraph()

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'Dynamic Vision Sensors (DVS) offer microsecond temporal resolution and high dynamic range '
        '(>120 dB), making them promising for astronomical observation. However, background activity '
        '(BA) noise dominates under low-light conditions, severely limiting sensitivity. '
        'We present a systematic review of 25 publications across four domains\u2014DVS noise physics, '
        'DVS denoising, DVS astronomical applications, and noise inverse problem methodologies from '
        'non-DVS fields (notably LIGO gravitational-wave detection)\u2014and identify five unexplored '
        'gaps (G1\u2013G5) at their intersections. '
        'Building on this analysis, we propose Physics-Informed DeepClean for DVS (PI-DC-DVS), '
        'a unified algorithm that integrates circuit-level physical models with auxiliary-channel '
        'regression and Bayesian inference to reconstruct and subtract noise with high fidelity. '
        'We further introduce a six-tier calibration framework (Cal-1 through Cal-6), where Cal-6 '
        'repurposes satellite light trails\u2014conventionally regarded as light pollution\u2014as natural '
        'calibration sources exploiting DVS-specific advantages. '
        'A proof-of-concept demonstration on the EBSSA space-observation dataset achieves 90.3% '
        'noise event removal while preserving satellite trajectories. '
        'The theoretical signal-to-noise ratio (SNR) improvement scales as 1/(1\u2212\u03b1), where '
        '\u03b1 is the noise model accuracy, offering 10\u00d7 improvement at \u03b1=0.9 and '
        '100\u00d7 at \u03b1=0.99. '
        'This framework has the potential to extend DVS detection limits by 2\u20134 magnitudes and '
        'enable template-free discovery of faint, fast-moving objects such as sub-50 m near-Earth objects.'
    )

    kw = doc.add_paragraph()
    r = kw.add_run('Keywords: ')
    r.bold = True
    kw.add_run('dynamic vision sensor, event camera, noise inverse problem, gravitational-wave '
               'detector, astronomical observation, physics-informed neural network, calibration')

    # ================================================================
    # 1. Introduction
    # ================================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Conventional noise reduction in imaging is formulated as the signal inverse problem: '
        'given an observation y = h * s + n, recover the signal s. '
        'This work explores the complementary perspective\u2014the noise inverse problem\u2014where '
        'the noise generation mechanism is modelled as a physical forward process and solved '
        'inversely to reconstruct and subtract noise, leaving only the signal in the residual.'
    )
    doc.add_paragraph(
        'Dynamic Vision Sensors (DVS) are neuromorphic sensors in which each pixel independently '
        'and asynchronously emits an event when the logarithmic intensity change exceeds a '
        'threshold [1,4]. Inspired by change-detection neurons in insect compound eyes, DVS offer '
        'microsecond temporal resolution, >120 dB dynamic range, and sparse output, making them '
        'attractive for space situational awareness (SSA) [12,14,15] and fast optical astronomy [16]. '
        'Under low-light conditions, however, shot-noise-induced background activity (BA) becomes '
        'the dominant source of events, overwhelming faint astronomical signals [1,2].'
    )
    doc.add_paragraph(
        'The core insight of this work is: Solve the noise inverse problem with high precision '
        '\u2192 reconstruct and subtract noise \u2192 only signal remains \u2192 structural SNR improvement.'
    )
    doc.add_paragraph(
        'This paradigm has been spectacularly successful in gravitational-wave (GW) astronomy, '
        'where DeepClean [17] and related methods [18,19,20] use auxiliary witness channels to '
        'model and subtract non-stationary instrumental noise from the strain signal. '
        'The iDQ framework [24] further assigns per-event noise probabilities in real time. '
        'Yet no equivalent pipeline exists for DVS astronomical observation.'
    )
    doc.add_paragraph(
        'In this paper we: (i) systematically survey 25 publications across four domains '
        '(section 2); (ii) identify five research gaps G1\u2013G5 at their intersections (section 3); '
        '(iii) propose the PI-DC-DVS algorithm integrating physics-informed modelling with '
        'auxiliary-channel regression (section 4); (iv) design a six-tier calibration framework '
        'including a novel satellite-trail calibration tier (section 5); and (v) present a '
        'proof-of-concept demonstration on real space-observation data (section 6).'
    )

    # ================================================================
    # 2. Systematic Literature Survey
    # ================================================================
    doc.add_heading('2. Systematic Literature Survey', level=1)
    doc.add_paragraph(
        'We classify 25 relevant publications into four domains (Figure 1).'
    )

    # Figure 1
    add_figure(doc,
               os.path.join(PARENT_DIR, 'fig1_gap_map_en.png'),
               'Gap map showing the four surveyed domains (A: DVS noise physics, B: DVS denoising, '
               'C: DVS astronomical applications, D: noise inverse problem methods) and the five '
               'identified research gaps (G1\u2013G5) at their intersections.',
               1)

    # Domain A
    doc.add_heading('2.1 Domain A: DVS Noise Physics (5 publications)', level=2)

    make_table(doc,
        ['#', 'Reference', 'Key Contribution'],
        [
            ['A1', 'Graca, Delbruck (2023) [1]', 'Proved 2\u00d7 photon shot noise as fundamental lower bound; bias optimisation guidelines'],
            ['A2', 'McReynolds et al. (2023) [2]', 'ON/OFF alternation statistics for shot noise identification'],
            ['A3', 'Graca et al. (2024) [3]', 'SciDVS: 1.7% temporal contrast sensitivity at 0.7 lux (scientific-grade DVS)'],
            ['A4', 'Delbruck et al. (2021) [4]', 'Feedback control framework; noise dependence on bias settings'],
            ['A5', 'Graca, Delbruck (2025) [5]', 'Large-signal ODE DVS pixel model with first-passage-time stochastic event generation; >1000\u00d7 speedup; usable as noise forward model'],
        ],
        'Domain A: DVS noise physics publications.', 1)

    # Domain B
    doc.add_heading('2.2 Domain B: DVS Denoising Methods (7 publications)', level=2)

    make_table(doc,
        ['#', 'Reference', 'Category', 'Key Contribution'],
        [
            ['B1', 'Delbruck (2008) [6]', 'Empirical', 'First spatio-temporal neighbourhood filter'],
            ['B2', 'Liu, Delbruck (2008) [7]', 'Empirical', 'Optical-flow-based filtering'],
            ['B3', 'Baldwin et al. (2020) [8]', 'Probabilistic+DL', 'Event Probability Mask; DVSNOISE20 dataset'],
            ['B4', 'McReynolds et al. (2023) [2]', 'Physics-statistical', 'ON/OFF alternation filtering (=A2)'],
            ['B5', 'Fang et al. (2024) [9]', 'DL', 'WedNet: real-time window-based denoising'],
            ['B6', 'Wu et al. (2024) [10]', 'DL', 'ASTEDNet: direct event stream processing'],
            ['B7', 'Shiba et al. (2025) [11]', 'Joint estimation', 'Simultaneous motion+noise estimation (closest prior work)'],
        ],
        'Domain B: DVS denoising publications.', 2)

    # Domain C
    doc.add_heading('2.3 Domain C: DVS Astronomical Applications (5 publications)', level=2)

    make_table(doc,
        ['#', 'Reference', 'Application', 'Key Contribution'],
        [
            ['C1', 'Afshar et al. (2019) [12]', 'SSA', 'First event-based space observation dataset (EBSSA)'],
            ['C2', 'Chin et al. (2019) [13]', 'Attitude', 'Event-camera star tracking'],
            ['C3', 'Joubert et al. (2022) [14]', 'SSA', 'FIESTA: unsupervised real-time detection/tracking'],
            ['C4', 'Gedek et al. (2019) [15]', 'SSA', 'Observational evaluation of event cameras for SSA'],
            ['C5', 'Hoang (2023) [16]', 'HE astro', 'Neuromorphic cameras for Cherenkov telescopes'],
        ],
        'Domain C: DVS astronomical application publications.', 3)

    # Domain D
    doc.add_heading('2.4 Domain D: Noise Inverse Problem Methods in Non-DVS Fields (7 publications)', level=2)

    make_table(doc,
        ['#', 'Reference', 'Field', 'Key Contribution'],
        [
            ['D1', 'Vajente et al. (2020) [17]', 'GW (LIGO)', 'DeepClean: aux-channel noise regression; most successful noise inverse problem implementation'],
            ['D2', 'Dooney et al. (2025) [18]', 'GW', 'DeepExtractor: signal/glitch reconstruction via noise distribution modelling'],
            ['D3', 'Wang et al. (2024) [19]', 'GW', 'WaveFormer: transformer-based GW denoising'],
            ['D4', 'Chatterjee, Jani (2025) [20]', 'GW', 'Robust signal reconstruction under noise artefacts'],
            ['D5', 'Cao et al. (2024) [21]', 'DVS+imaging', 'Noise2Image: static scene recovery from DVS noise (paradigm shift)'],
            ['D6', 'Essick et al. (2021) [24]', 'GW (LIGO)', 'iDQ: per-event noise probability estimation'],
            ['D7', 'Selig, Ensslin (2015) [25]', 'Astronomy', 'D3PO: simultaneous Bayesian signal+noise reconstruction for photon data'],
        ],
        'Domain D: Noise inverse problem publications from non-DVS fields.', 4)

    # ================================================================
    # 3. Gap Analysis
    # ================================================================
    doc.add_heading('3. Gap Analysis: Five Unexplored Research Directions', level=1)
    doc.add_paragraph(
        'Cross-domain analysis reveals five gaps (Table 5).'
    )

    make_table(doc,
        ['Gap', 'Description', 'Domains'],
        [
            ['G1', 'Rigorous Bayesian inverse-problem formulation from DVS circuit physics', 'A\u2192D'],
            ['G2', 'LIGO DeepClean-style auxiliary-channel DVS noise prediction', 'D\u2192A'],
            ['G3', 'Astronomy-specific DVS noise inverse problem pipeline', 'A+B+D\u2192C'],
            ['G4', 'Template-free object detection via noise residual analysis', 'C+D'],
            ['G5', 'Integrated demonstration: SciDVS + large telescope + noise inverse problem', 'A+C+D'],
        ],
        'Five identified research gaps at the intersection of domains A\u2013D.', 5)

    doc.add_heading('3.1 G1: Bayesian Inverse Problem from DVS Circuit Physics', level=2)
    doc.add_paragraph(
        'The forward model of [5] generates noise events as e_noise = F(\u03b8_pixel, I_bg, T, bias). '
        'The Bayesian inverse problem seeks the posterior: '
        'p(\u03b8|e_obs) \u221d p(e_obs|\u03b8) \u00b7 p(\u03b8), '
        'where p(e_obs|\u03b8) is the inhomogeneous Poisson process likelihood and p(\u03b8) derives '
        'from the physical parameter ranges of [5]. '
        'Following D3PO [25], joint estimation of signal s and noise parameters \u03b8 yields '
        'p(\u03b8, s|e_obs) \u221d p(e_obs|\u03b8, s) \u00b7 p(\u03b8) \u00b7 p(s).'
    )

    doc.add_heading('3.2 G2: Auxiliary-Channel Noise Prediction', level=2)
    doc.add_paragraph(
        'By analogy with DeepClean [17], DVS auxiliary channels\u2014temperature sensor, '
        'accelerometer, and background illuminance monitor\u2014can independently measure '
        'environmental drivers of non-stationary noise. '
        'Dark current depends exponentially on temperature: I_dark \u221d exp(\u2212E_g/2k_BT); '
        'a 1\u00b0C rise approximately doubles the dark current.'
    )

    doc.add_heading('3.3 G3: Astronomy-Specific DVS Noise Inverse Problem Pipeline', level=2)
    doc.add_paragraph(
        'G3 integrates G1 and G2 into a four-stage pipeline (Figure 2): '
        '(1) noise forward model construction, (2) Bayesian inverse problem solution, '
        '(3) residual event stream generation, (4) astronomical calibration and verification.'
    )

    # Figure 2
    add_figure(doc,
               os.path.join(PARENT_DIR, 'fig2_g3_pipeline_en.png'),
               'System architecture of the G3 astronomy-specific DVS noise inverse problem pipeline.',
               2)

    make_table(doc,
        ['Method', 'Noise model', 'Aux. channels', 'Astro. use', 'Bayesian'],
        [
            ['Shiba et al. [11]', 'Data-driven', '\u2717', '\u2717', '\u2717'],
            ['Noise2Image [21]', 'Illuminance-dep.', '\u2717', '\u2717', '\u2717'],
            ['FIESTA [14]', 'Threshold', '\u2717', '\u2713', '\u2717'],
            ['DeepClean [17]', 'Physics+ML', '\u2713', '\u2717 (LIGO)', '\u2717'],
            ['D3PO [25]', 'Bayesian', '\u2717', '\u2713 (\u03b3-ray)', '\u2713'],
            ['G3 (proposed)', 'Physics(A5)+ML', '\u2713', '\u2713', '\u2713'],
        ],
        'Comparison of G3 with prior denoising approaches.', 6)

    doc.add_heading('3.4 G4: Template-Free Object Detection via Noise Residual', level=2)
    doc.add_paragraph(
        'Traditional astronomical detection relies on signal templates. '
        'G4 inverts this: if the noise is precisely subtracted, any structure remaining in the '
        'residual is signal\u2014regardless of its morphology. '
        'Detection sensitivity depends solely on noise model accuracy, not on signal assumptions. '
        'This mirrors unmodelled burst searches in GW astronomy.'
    )

    make_table(doc,
        ['Aspect', 'Conventional approach', 'G4 approach'],
        [
            ['Detection principle', 'Signal template matching', 'Deviation from noise model'],
            ['Prior knowledge', 'Signal morphology known', 'Noise generation known'],
            ['Detectable targets', 'Known object types only', 'Unknown types included'],
            ['Sensitivity factor', 'Template accuracy', 'Noise model accuracy'],
            ['LIGO analogy', 'Matched filtering', 'Burst search (unmodelled)'],
        ],
        'Paradigm shift: template-dependent vs. template-free detection.', 7)

    doc.add_heading('3.5 G5: Integrated Demonstration Roadmap', level=2)

    make_table(doc,
        ['Phase', 'Configuration', 'Aperture', 'Objective'],
        [
            ['1', 'SciDVS + small telescope', '0.3\u20130.5 m', 'Pipeline validation, known-source recovery'],
            ['2', 'SciDVS + medium telescope', '1\u20132 m', 'Faint-object detection limit measurement'],
            ['3', 'SciDVS + large telescope', '4 m class', 'Template-free detection demonstration'],
        ],
        'Phased roadmap for integrated SciDVS + telescope demonstration.', 8)

    # ================================================================
    # 4. Proposed Algorithm: PI-DC-DVS
    # ================================================================
    doc.add_heading('4. Proposed Algorithm: Physics-Informed DeepClean for DVS (PI-DC-DVS)', level=1)

    doc.add_heading('4.1 Fundamental Principle', level=2)
    doc.add_paragraph(
        'Let \u03b1 denote the noise model accuracy: '
        '\u03b1 = 1 \u2212 ||\u0115_noise \u2212 e_noise,true|| / ||e_noise,true||. '
        'The residual noise level after subtraction is \u03c3_residual = (1\u2212\u03b1) \u00b7 \u03c3_original, '
        'yielding an SNR improvement ratio: SNR_after / SNR_before = 1/(1\u2212\u03b1). '
        'At \u03b1=0.9 this gives 10\u00d7 improvement; at \u03b1=0.99, 100\u00d7 (Figure 3).'
    )

    doc.add_heading('4.2 Algorithm Overview', level=2)
    doc.add_paragraph(
        'PI-DC-DVS operates in four phases:'
    )
    doc.add_paragraph(
        'Phase 1 (Offline Calibration): Record dark events e_dark (lens cap), flat-field events '
        'e_flat (integrating sphere), and thermal sweep events e_thermal (\u0394T = \u00b15\u00b0C). '
        'Fit the A5 forward model [5] to obtain per-pixel parameter maps via MAP estimation: '
        '\u03b8\u0302 = argmax_\u03b8 p(e_cal|\u03b8) \u00b7 p(\u03b8|\u03b8_prior).'
    )
    doc.add_paragraph(
        'Phase 2 (Online Inference): A Physics-Informed Neural Network (PI-DC-Net) predicts '
        'per-pixel noise rates. The network has three layers: '
        '(a) Physics model layer (fixed weights, A5 model) providing baseline noise rate; '
        '(b) Auxiliary-channel coupling layer (MLP [64-32-1]) learning non-stationary variations; '
        '(c) Spatio-temporal correlation layer (Conv2D, 3\u00d73 kernel) learning inter-pixel correlations. '
        'Output: \u03bb\u0302_noise(x,y,t) = \u03bb_physics(x,y,t) \u00b7 (1 + \u0394\u03bb_aux(t)) + \u0394\u03bb_corr(x,y,t). '
        'Per-event noise probability: P_noise(e_i) = \u03bb\u0302_noise / (\u03bb\u0302_noise + \u03bb\u0302_signal) '
        'following iDQ [24].'
    )
    doc.add_paragraph(
        'Phase 3 (Residual Generation): '
        'Soft subtraction (recommended): assign weight w_i = 1 \u2212 P_noise(e_i), '
        'retain events with w_i > threshold. '
        'Hard subtraction (fast): retain events with P_noise(e_i) < \u03c4.'
    )
    doc.add_paragraph(
        'Phase 4 (Adaptive Updates): Monitor residual Poisson statistics; trigger online weight '
        'updates when deviations exceed thresholds. '
        'Apply Kalman-filter-like drift correction to \u03b8_pixel_map.'
    )

    # Figure 3 — SNR improvement analysis (first cited in section 4.1)
    add_figure(doc,
               os.path.join(PARENT_DIR, 'fig4_sn_improvement.png'),
               'SNR improvement analysis. '
               '(a) Fano factor spatial map showing noise-dominated (Fano \u2248 1) vs. '
               'signal-containing (Fano \u226b 1) pixels; '
               '(b) Temporal dynamics of noise and signal event rates; '
               '(c) Per-pixel SNR distribution before and after noise subtraction.',
               3)

    doc.add_heading('4.3 Expected Performance', level=2)

    make_table(doc,
        ['Condition', 'Threshold-based', 'PI-DC-DVS', 'Rationale'],
        [
            ['Dark current (<0.1 lux)', '~50% removal', '~90% removal', 'Physics model precisely predicts dark current'],
            ['Temperature drift (\u00b15\u00b0C/h)', 'Not addressed', 'Real-time tracking', 'Auxiliary channel + adaptive update'],
            ['Pixel mismatch', 'Uniform threshold', 'Per-pixel optimised', 'Calibration provides \u03b8_pixel_map'],
            ['Mechanical vibration', 'Not addressed', 'Predicted & removed', 'Accelerometer (proven by LIGO DeepClean)'],
        ],
        'Expected noise removal performance: PI-DC-DVS vs. conventional threshold-based denoising.', 9)

    # ================================================================
    # 5. Calibration Framework
    # ================================================================
    doc.add_heading('5. Proposed Calibration Framework', level=1)

    doc.add_heading('5.1 Six-Tier Calibration Dataset', level=2)

    make_table(doc,
        ['Tier', 'Condition', 'Purpose', 'Pass criterion'],
        [
            ['Cal-1', 'Dark (lens cap)', 'Pure noise reference', '\u03c7\u00b2/dof < 1.5'],
            ['Cal-2', 'Thermal sweep (\u00b15\u00b0C)', 'Temperature dependence', 'Residual < 10% all T'],
            ['Cal-3', 'Flat-field (integrating sphere)', 'Shot noise statistics', '\u03b1_flat > 0.9'],
            ['Cal-4', 'Dynamic patterns', 'Injection-recovery', 'AUC > 0.95'],
            ['Cal-5', 'Simulated astronomy', 'End-to-end pipeline', '\u0394m_lim > 2 mag'],
            ['Cal-6', 'Satellite trails', 'In-operation verification', 'Detection rate > 95%'],
        ],
        'Six-tier calibration framework for DVS noise model validation.', 10)

    doc.add_heading('5.2 Cal-6: Satellite Trail Calibration \u2014 Repurposing Light Pollution', level=2)
    doc.add_paragraph(
        'We propose repurposing satellite light trails\u2014conventionally regarded as light pollution\u2014'
        'as natural calibration sources for the noise model.'
    )
    doc.add_paragraph(
        'Rationale: Artificial satellites have precisely predictable trajectories: TLE orbital data '
        'provide position, velocity, and transit time to microsecond precision. '
        'Reflected sunlight magnitude can be estimated from the reflective cross-section, solar angle, '
        'and attitude model. This makes satellite transits a natural injection-recovery test.'
    )
    doc.add_paragraph(
        'DVS-Specific Advantages: '
        '(1) Abundant calibration opportunities (Starlink alone: several thousand satellites, '
        'dozens of transits per night); '
        '(2) Real observing conditions (unlike laboratory injection); '
        '(3) No saturation (DVS >120 dB dynamic range vs. CCD saturation); '
        '(4) Temporal compatibility (DVS \u03bcs resolution matches LEO angular speed); '
        '(5) No additional hardware required.'
    )
    doc.add_paragraph(
        'Verification metrics: Detection rate P_detect(m) vs. predicted magnitude; '
        'magnitude agreement |m_detect \u2212 m_pred| < 0.5 mag; trajectory residuals. '
        'Pass criterion: detection rate > 95% for satellites brighter than m_lim \u2212 1.'
    )

    doc.add_heading('5.3 Calibration Pipeline', level=2)
    doc.add_paragraph(
        'The pipeline consists of pre-observation calibration (Cal-1 through Cal-5, all must pass) '
        'followed by in-operation calibration (Cal-6, runs continuously). '
        'Cal-6 failures trigger the Phase 4 adaptive update mechanism.'
    )

    doc.add_heading('5.4 Real-Time Monitoring Metrics', level=2)

    make_table(doc,
        ['Metric', 'Definition', 'Threshold', 'Remediation'],
        [
            ['Residual Poissonicity', 'KS test p-value of residual ISI', 'p > 0.05', 'Trigger Phase 4 update'],
            ['Residual rate stability', 'CV of non-signal residual rate', 'CV < 0.3', 'Temperature drift correction'],
            ['Aux. channel consistency', 'NN vs. physics model divergence', '|\u0394\u03bb/\u03bb| < 0.5', 'Aux. channel anomaly alert'],
            ['Pixel anomaly rate', 'Fraction with \u03b8 outside 3\u03c3', '< 1%', 'Hot-pixel mask update'],
        ],
        'Real-time monitoring metrics during observation.', 11)

    # ================================================================
    # 6. Proof-of-Concept Demonstration
    # ================================================================
    doc.add_heading('6. Proof-of-Concept Demonstration', level=1)

    doc.add_heading('6.1 Dataset', level=2)
    doc.add_paragraph(
        'We use the Event-Based Space Situational Awareness (EBSSA) dataset [12]: '
        '236 recordings from a DAVIS240C sensor observing satellites and stars, with 572 labelled '
        'space objects.'
    )

    doc.add_heading('6.2 Simplified Noise Model', level=2)
    doc.add_paragraph(
        'As EBSSA lacks auxiliary channel data, we implement a simplified version of PI-DC-DVS '
        'using the Fano factor (variance-to-mean ratio of event rates across time bins). '
        'Pixels with Fano factor > 2 are classified as signal candidates. '
        'Per-event noise probability: P_noise(e_i) = \u03bb_noise(x_i,y_i) / \u03bb_total(x_i,y_i), '
        'where \u03bb_noise is estimated from noise-dominated pixels (Fano \u2264 2).'
    )

    doc.add_heading('6.3 Results', level=2)
    doc.add_paragraph(
        'From 1,800,674 input events, probabilistic thinning with threshold \u03c4 = 0.5 yields '
        '175,261 residual events\u2014a 90.3% noise removal rate. '
        'The residual event stream clearly reveals satellite trajectories that are buried in '
        'noise in the raw stream (Figure 4). '
        'Signal candidate pixels (2,294 out of 76,800) concentrate along known satellite tracks.'
    )

    # Figure 4
    add_figure(doc,
               os.path.join(PARENT_DIR, 'fig3_noise_inverse_demo.png'),
               'Proof-of-concept demonstration on EBSSA data. '
               '(a) Raw event accumulation (1,800,674 events); '
               '(b) Estimated noise rate map; '
               '(c) Per-event noise probability distribution; '
               '(d) Residual events after 90.3% noise removal (175,261 events).',
               4)

    doc.add_paragraph(
        'The Fano factor spatial map (Figure 3) shows clear separation between noise-dominated '
        'pixels (Fano \u2248 1, Poisson) and signal-containing pixels (Fano \u226b 1). '
        'Temporal analysis confirms that the noise rate is approximately stationary, while signal '
        'events cluster around satellite transit times.'
    )

    doc.add_heading('6.4 Limitations', level=2)
    doc.add_paragraph(
        'This demonstration uses a simplified noise model without auxiliary channels. '
        'The full PI-DC-DVS algorithm with physics-informed neural network, auxiliary-channel '
        'regression, and online adaptation is expected to achieve substantially higher noise model '
        'accuracy (\u03b1 > 0.95) and correspondingly greater SNR improvement.'
    )

    # ================================================================
    # 7. Discussion
    # ================================================================
    doc.add_heading('7. Discussion', level=1)
    doc.add_paragraph(
        'The proposed framework synthesises ideas from multiple fields. '
        'The key advance is the integration of DVS-specific circuit physics [5] as a structural '
        'prior in the neural network, combined with auxiliary-channel regression proven in '
        'GW astronomy [17,24]. This physics-informed approach offers interpretability, data '
        'efficiency, and robustness to distribution shift compared to purely data-driven denoising.'
    )
    doc.add_paragraph(
        'Cal-6 exemplifies "turning a bug into a feature": the proliferation of satellite '
        'constellations, widely lamented as degrading ground-based astronomy, here provides a '
        'continuous, cost-free, and statistically abundant source of calibration signals. '
        'While CCD sensors saturate on bright satellite trails, DVS sensors record them '
        'quantitatively across their full dynamic range.'
    )
    doc.add_paragraph(
        'The noise inverse problem framework is not specific to astronomy. '
        'Potential applications include neural calcium imaging, industrial inspection '
        '(faint defect detection), and autonomous driving (adverse lighting conditions).'
    )

    # ================================================================
    # 8. Conclusion
    # ================================================================
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph(
        'We have presented a systematic review of DVS noise modelling, denoising, astronomical '
        'application, and noise inverse problem methodologies, identifying five unexplored research '
        'gaps (G1\u2013G5). The proposed PI-DC-DVS algorithm integrates circuit-level physics, '
        'auxiliary-channel regression, and Bayesian inference into a unified framework for '
        'high-precision noise reconstruction and subtraction. The six-tier calibration framework '
        '(Cal-1 through Cal-6) provides quantitative validation, with Cal-6 innovatively repurposing '
        'satellite light trails as natural calibration sources. A proof-of-concept demonstration '
        'on real space-observation data achieves 90.3% noise removal, and the theoretical framework '
        'predicts SNR improvements of 10\u2013100\u00d7 at noise model accuracies of '
        '\u03b1 = 0.9\u20130.99. All necessary methodological components exist in adjacent fields; '
        'the contribution of this work is their systematic identification and proposed integration '
        'for DVS astronomical observation.'
    )

    # ================================================================
    # References
    # ================================================================
    doc.add_heading('References', level=1)
    refs = [
        '[1] Graca R, Delbruck T 2023 Optimal biasing and physical limits of DVS event noise. arXiv:2304.04019',
        '[2] McReynolds B, Graca R, Delbruck T 2023 Exploiting alternating DVS shot noise event pair statistics. arXiv:2304.03494',
        '[3] Graca R, Zhou S, McReynolds B, Delbruck T 2024 SciDVS: a scientific event camera with 1.7% temporal contrast sensitivity at 0.7 lux. ESSERC 2024',
        '[4] Delbruck T, Graca R, Paluch M 2021 Feedback control of event cameras. CVPRW 2021',
        '[5] Graca R, Delbruck T 2025 Towards a physically realistic computationally efficient DVS pixel model. arXiv:2505.07386',
        '[6] Delbruck T 2008 Frame-free dynamic digital vision. Proc. Intl. Symp. Secure-Life Electronics',
        '[7] Liu S-C, Delbruck T 2008 Adaptive time-slice block-matching optical flow. BMVC',
        '[8] Baldwin RW, Almatrafi M, Asari V, Hirakawa K 2020 Event probability mask (EPM) and EDnCNN. CVPR 2020',
        '[9] Fang H, Wu J, Hou Q, Dong W, Shi G 2024 Fast window-based event denoising. IEEE TPAMI',
        '[10] Wu W, Yao H, Zhai C, Dai Z, Zhu X 2024 Event camera denoising using ASTEDNet. ISPRS Archives XLVIII-4-2024',
        '[11] Shiba S, Aoki Y, Gallego G 2025 Simultaneous motion and noise estimation with event cameras. ICCV 2025',
        '[12] Afshar S, Nicholson AP, van Schaik A, Cohen G 2019 Event-based object detection and tracking for SSA. arXiv:1911.08730',
        '[13] Chin T-J, Bagchi S, Eriksson A, van Schaik A 2019 Star tracking using an event camera. CVPRW 2019',
        '[14] Joubert D, Afshar S et al 2022 FIESTA: real-time event-based feature consolidation for SSA. Front. Neurosci. 16 821157',
        '[15] Gedek M, Zolnowski M, Delbruck T et al 2019 Observational evaluation of event cameras for SSA. EESA',
        '[16] Hoang J 2023 Neuromorphic cameras for atmospheric Cherenkov telescopes. arXiv:2310.16321',
        '[17] Vajente G, Huang Y, Isi M et al 2020 Machine-learning nonstationary noise out of gravitational-wave detectors. Phys. Rev. D 101 042003',
        '[18] Dooney T, Narola H et al 2025 DeepExtractor: time-domain reconstruction of signals and glitches in GW data. arXiv:2501.18423',
        '[19] Wang H, Zhou Y, Cao Z et al 2024 WaveFormer: transformer-based denoising for GW data. Mach. Learn.: Sci. Technol. 5 015046',
        '[20] Chatterjee C, Jani K 2025 No glitch in the matrix: robust GW signal reconstruction. ApJ',
        '[21] Cao R, Galor D, Kohli A, Yates JL, Waller L 2024 Noise2Image: noise-enabled static scene recovery for event cameras. Optica',
        '[22] Gallego G et al 2020 Event-based vision: a survey. IEEE TPAMI 42(1) 154-180',
        '[23] Stetzler S, Juric M et al 2025 An efficient shift-and-stack algorithm applied to detection catalogs. AJ 170 352',
        '[24] Essick R, Godwin P, Hanna C et al 2021 iDQ: statistical inference of non-astrophysical noise transients in GW detectors. Mach. Learn.: Sci. Technol. 2 015004',
        '[25] Selig M, Ensslin TA 2015 D3PO \u2013 denoising, deconvolving, and decomposing photon observations. A&A 574 A74',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(10)

    # Save
    out_path = os.path.join(BASE_DIR, 'dvs_noise_inverse_mlst.docx')
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == '__main__':
    main()
