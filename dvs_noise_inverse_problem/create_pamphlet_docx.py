#!/usr/bin/env python3
"""
市民向けA4パンフレット（1-2枚）: DVSノイズ逆問題の研究紹介
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_styled_paragraph(doc, text, font_size=10.5, bold=False, color=None,
                         alignment=None, space_before=0, space_after=4,
                         font_name='Yu Gothic'):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): font_name})
    rPr.append(rFonts)
    return p


def add_heading_box(doc, title, subtitle=None):
    """Add a colored heading box using a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1  # center
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, '1B3A5C')  # dark navy

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Yu Gothic'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'Yu Gothic'})
    rPr.append(rFonts)

    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(6)
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(200, 220, 255)
        run2.font.name = 'Yu Gothic'
        rPr2 = run2._element.get_or_add_rPr()
        rFonts2 = rPr2.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'Yu Gothic'})
        rPr2.append(rFonts2)


def add_section_heading(doc, number, title, color_hex=(27, 58, 92)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    # Number badge
    run_num = p.add_run(f' {number} ')
    run_num.font.size = Pt(12)
    run_num.font.bold = True
    run_num.font.color.rgb = RGBColor(255, 255, 255)
    run_num.font.name = 'Yu Gothic'
    # Add shading to the number run
    rPr = run_num._element.get_or_add_rPr()
    shd = rPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:fill'): '{:02X}{:02X}{:02X}'.format(*color_hex),
    })
    rPr.append(shd)

    # Title text
    run_title = p.add_run(f'  {title}')
    run_title.font.size = Pt(13)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(*color_hex)
    run_title.font.name = 'Yu Gothic'
    rPr2 = run_title._element.get_or_add_rPr()
    rFonts = rPr2.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'Yu Gothic'})
    rPr2.append(rFonts)


def add_highlight_box(doc, text, bg_color='E8F0FE', border_color='4285F4'):
    """Add a highlighted text box using a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Yu Gothic'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'Yu Gothic'})
    rPr.append(rFonts)


def main():
    doc = Document()

    # Page setup: A4, narrow margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ========== PAGE 1 ==========

    # Title banner
    add_heading_box(
        doc,
        '「ノイズ」で宇宙を見る',
        '次世代カメラが切り拓く、微弱天体発見の新手法'
    )

    # --- Section 1: What is this research? ---
    add_section_heading(doc, '1', 'この研究は何をするの？')

    add_styled_paragraph(
        doc,
        '夜空の写真を撮ると、星以外にも「ノイズ」と呼ばれる砂嵐のような'
        'ちらつきが写り込みます。従来の天文学では、このノイズは「邪魔者」'
        'として捨てられてきました。',
        font_size=10.5
    )

    add_highlight_box(
        doc,
        'この研究の発想：ノイズを捨てるのではなく、ノイズの「仕組み」を解き明かして、'
        '精密に再現・除去する。すると、ノイズに隠れていた微かな天体の光が浮かび上がる！',
        bg_color='FFF3E0', border_color='FF9800'
    )

    add_styled_paragraph(
        doc,
        'これは医療でいえば、レントゲン写真から骨だけでなく「影」の原因も突き止めて、'
        'よりクリアな画像を得るようなものです。',
        font_size=10
    )

    # --- Section 2: DVS Camera ---
    add_section_heading(doc, '2', '「昆虫の目」を持つ次世代カメラ — DVS')

    add_styled_paragraph(
        doc,
        '本研究では、DVS（ダイナミック・ビジョン・センサー）という特殊なカメラを使います。'
        '従来のカメラが「写真」のようにシャッターで一枚ずつ記録するのに対し、'
        'DVSは昆虫の複眼のように、各ピクセルが独立して明るさの「変化」だけを記録します。',
        font_size=10.5
    )

    # Comparison table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['', '従来のカメラ（CCD）', 'DVS（イベントカメラ）']
    comparisons = [
        ['時間分解能', '数十ミリ秒', '1マイクロ秒（10万倍）'],
        ['ダイナミックレンジ', '60-70 dB', '120 dB以上'],
        ['データ出力', 'フレーム全体', '変化した点だけ（省エネ）'],
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1B3A5C')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.name = 'Yu Gothic'
    for row_idx, row_data in enumerate(comparisons):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Yu Gothic'
            if col_idx == 2:
                set_cell_shading(cell, 'E8F0FE')

    # --- Section 3: Key idea ---
    add_section_heading(doc, '3', '核心アイデア：ノイズの「逆問題」')

    add_styled_paragraph(
        doc,
        '従来のノイズ除去は、「ノイズっぽいものを機械的にカットする」方法でした。'
        'これだと、微かな天体の信号まで一緒に消えてしまいます。',
        font_size=10.5
    )

    add_styled_paragraph(
        doc,
        '本研究は逆の発想です：',
        font_size=10.5, bold=True
    )

    # Step-by-step
    steps = [
        ('Step 1', 'ノイズの物理法則をモデル化', 'DVSチップの回路特性・温度・背景光から、各ピクセルのノイズ発生率を物理方程式で予測'),
        ('Step 2', 'ノイズを精密に「再現」', 'AIが物理モデルを学習し、観測中にリアルタイムでノイズを再構成'),
        ('Step 3', '観測データからノイズを差し引く', '残ったものが「本物の天体信号」→ 従来見えなかった微弱天体が出現！'),
    ]
    for step_name, title, desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(14)
        run_step = p.add_run(f'{step_name}: ')
        run_step.font.size = Pt(10)
        run_step.font.bold = True
        run_step.font.color.rgb = RGBColor(27, 58, 92)
        run_step.font.name = 'Yu Gothic'
        run_title = p.add_run(f'{title}\n')
        run_title.font.size = Pt(10)
        run_title.font.bold = True
        run_title.font.name = 'Yu Gothic'
        run_desc = p.add_run(f'  → {desc}')
        run_desc.font.size = Pt(9.5)
        run_desc.font.name = 'Yu Gothic'
        run_desc.font.color.rgb = RGBColor(80, 80, 80)

    # --- Section 4: Demo result ---
    add_section_heading(doc, '4', '実証結果：90%のノイズ除去に成功')

    add_styled_paragraph(
        doc,
        '公開データセット（EBSSA: 衛星観測データ）を使った概念実証では、'
        '約180万個のイベントから90.3%のノイズイベントを除去し、'
        '衛星の軌跡を鮮明に浮かび上がらせることに成功しました。',
        font_size=10.5
    )

    # Add demo figure if available
    demo_fig = os.path.join(BASE_DIR, 'fig3_noise_inverse_demo.png')
    if os.path.exists(demo_fig):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(demo_fig, width=Inches(5.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(4)
        run_cap = cap.add_run(
            '図：ノイズ除去の実証結果。左が元データ（ノイズだらけ）、右がノイズ除去後（衛星軌跡が鮮明に）'
        )
        run_cap.font.size = Pt(8.5)
        run_cap.font.italic = True
        run_cap.font.name = 'Yu Gothic'
        run_cap.font.color.rgb = RGBColor(100, 100, 100)

    add_highlight_box(
        doc,
        'S/N比（信号対ノイズ比）の理論的改善：ノイズ精度90%で10倍、99%で100倍の改善が可能。'
        'これは検出限界を2〜4等級改善することに相当し、これまで見えなかった暗い天体の発見につながります。',
        bg_color='E8F5E9', border_color='4CAF50'
    )

    # --- Section 5: Satellite calibration ---
    add_section_heading(doc, '5', '人工衛星の「光害」を味方に変える')

    add_styled_paragraph(
        doc,
        '天文学では、人工衛星の光跡は「光害」として問題視されています。'
        'しかし本研究では、この光跡を逆に「校正信号」として活用する新手法（Cal-6）を提案します。',
        font_size=10.5
    )

    add_styled_paragraph(
        doc,
        '人工衛星は軌道が正確にわかっているため、「いつ・どこに・どの明るさで」光るかが予測できます。'
        'この予測と実際の観測を照合することで、ノイズモデルの精度を自動的・継続的に検証できます。'
        'スターリンク衛星だけでも数千基が周回しており、毎晩数十回の校正機会が得られます。',
        font_size=10
    )

    # --- Section 6: Future impact ---
    add_section_heading(doc, '6', '将来の展望：何が見つかるか')

    impacts = [
        '小惑星の早期発見 — 直径10〜50mの地球近傍小天体（NEO）の検出。地球防衛に貢献',
        '高速天体現象の観測 — マイクロ秒単位の天体フラッシュ（ガンマ線バースト残光など）',
        '宇宙状況認識（SSA） — 人工衛星やスペースデブリのリアルタイム監視の高精度化',
        '他分野への応用 — 神経科学（脳イメージング）、工業検査、自動運転の暗所認識',
    ]
    for impact in impacts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Pt(14)
        # Split at first —
        parts = impact.split(' — ', 1)
        run_title = p.add_run(f'● {parts[0]}')
        run_title.font.size = Pt(10)
        run_title.font.bold = True
        run_title.font.name = 'Yu Gothic'
        if len(parts) > 1:
            run_desc = p.add_run(f' — {parts[1]}')
            run_desc.font.size = Pt(9.5)
            run_desc.font.name = 'Yu Gothic'
            run_desc.font.color.rgb = RGBColor(80, 80, 80)

    # Footer
    doc.add_paragraph()  # spacer
    footer_table = doc.add_table(rows=1, cols=1)
    cell = footer_table.rows[0].cells[0]
    set_cell_shading(cell, 'F5F5F5')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'この研究は、重力波検出で実証された「ノイズの物理モデリング＋AI」技術を\n'
        '天文観測用の次世代カメラに初めて適用するものです。\n'
        'コード・データはオープンソースで公開しています。'
    )
    run.font.size = Pt(9)
    run.font.name = 'Yu Gothic'
    run.font.color.rgb = RGBColor(100, 100, 100)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'Yu Gothic'})
    rPr.append(rFonts)

    # Save
    out_path = os.path.join(BASE_DIR, 'pamphlet_dvs_noise_inverse.docx')
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    main()
