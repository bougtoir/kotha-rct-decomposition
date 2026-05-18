"""
Generate STROBE checklist for cross-sectional study (EJA submission).
European Journal of Anaesthesiology requires STROBE for observational studies.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def create_strobe_checklist():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(29.7)  # A4 landscape
        section.page_height = Cm(21.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'STROBE Statement\u2014Checklist of items that should be included in reports of '
        'cross-sectional studies')
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Impact of the EU desflurane regulation on secondary market prices of '
        'anaesthetic vaporisers')
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    # STROBE items for cross-sectional study
    items = [
        ('Title and abstract', '1',
         '(a) Indicate the study\u2019s design with a commonly used term in the title or abstract',
         'Title: \u201ccross-sectional time-series analysis\u201d'),
        ('', '',
         '(b) Provide in the abstract an informative and balanced summary of what was done and found',
         'Abstract: structured abstract with BACKGROUND, OBJECTIVE, DESIGN, SETTING, '
         'MAIN OUTCOME MEASURES, RESULTS, CONCLUSION'),
        ('Introduction', '', '', ''),
        ('Background/rationale', '2',
         'Explain the scientific background and rationale for the investigation being reported',
         'Introduction, paragraphs 1\u20134'),
        ('Objectives', '3',
         'State specific objectives, including any prespecified hypotheses',
         'Introduction, final paragraph: hypothesis that desflurane prices would decline '
         'while sevoflurane/isoflurane remain stable'),
        ('Methods', '', '', ''),
        ('Study design', '4',
         'Present key elements of study design early in the paper',
         'Methods, \u201cStudy design and data source\u201d: cross-sectional time-series analysis'),
        ('Setting', '5',
         'Describe the setting, locations, and relevant dates, including periods of recruitment, '
         'exposure, follow-up, and data collection',
         'Methods: eBay/Terapeak, data collected March 2026, covering 28 March 2023\u201324 March 2026'),
        ('Participants', '6',
         'Give the eligibility criteria, and the sources and methods of selection of participants',
         'Methods, \u201cEligibility criteria\u201d: inclusion/exclusion criteria detailed'),
        ('Variables', '7',
         'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers',
         'Methods, \u201cVariables\u201d: primary outcome (sale price), exposure (regulatory period), '
         'agent type'),
        ('Data sources/measurement', '8',
         'For each variable of interest, give sources of data and details of methods of assessment',
         'Methods: Terapeak product research tool; automated data retrieval with manual '
         'screening against eligibility criteria'),
        ('Bias', '9',
         'Describe any efforts to address potential sources of bias',
         'Methods/Discussion: single marketplace to avoid cross-listing duplicates; '
         'natural comparator groups (sevoflurane, isoflurane)'),
        ('Study size', '10',
         'Explain how the study size was arrived at',
         'Methods, \u201cStatistical analysis\u201d: all available transactions within Terapeak '
         'window; no a priori sample size calculation'),
        ('Quantitative variables', '11',
         'Explain how quantitative variables were handled in the analyses',
         'Methods: prices in USD; regulatory period classified by milestones; '
         'trend analyses used continuous (days) and ordinal (phase 1\u20135) variables'),
        ('Statistical methods', '12',
         '(a) Describe all statistical methods',
         'Methods, \u201cStatistical analysis\u201d: Mann\u2013Whitney U, Welch\u2019s t-test, '
         'Spearman \u03c1, Kendall \u03c4, Kruskal\u2013Wallis, Cohen\u2019s d, LOWESS'),
        ('', '',
         '(b) Describe any methods used to examine subgroups and interactions',
         'Analyses performed separately for each agent type; supplementary asking-price analysis'),
        ('', '',
         '(c) Explain how missing data were addressed',
         'Methods, \u201cEligibility criteria\u201d: listings with missing price data excluded'),
        ('', '',
         '(d) If applicable, describe analytical methods taking account of sampling strategy',
         'Not applicable (complete enumeration of available sales)'),
        ('', '',
         '(e) Describe any sensitivity analyses',
         'Welch\u2019s t-test as sensitivity analysis; quarterly aggregation for trend robustness'),
        ('Results', '', '', ''),
        ('Participants', '13',
         '(a) Report numbers of individuals at each stage of study',
         'Results, paragraph 1: 1,033 total sales after exclusion criteria applied'),
        ('', '',
         '(b) Give reasons for non-participation at each stage',
         'Methods, \u201cEligibility criteria\u201d: exclusion criteria listed'),
        ('', '',
         '(c) Consider use of a flow diagram',
         'Not applicable (marketplace data, not participant recruitment)'),
        ('Descriptive data', '14',
         '(a) Give characteristics of study participants and information on exposures',
         'Results, Table 1: n, mean, SD, median, IQR, range by agent and period'),
        ('', '',
         '(b) Indicate number of participants with missing data for each variable of interest',
         'No missing data after eligibility screening'),
        ('Outcome data', '15',
         'Report numbers of outcome events or summary measures',
         'Results, Tables 1\u20132: all summary statistics and trend results by agent type'),
        ('Main results', '16',
         '(a) Give unadjusted estimates and, if applicable, confounder-adjusted estimates',
         'Results: Spearman \u03c1, Kendall \u03c4, Mann\u2013Whitney U P values, '
         'Cohen\u2019s d reported for each agent'),
        ('', '',
         '(b) Report category boundaries when continuous variables were categorized',
         'Methods, \u201cVariables\u201d: regulatory period cutpoints defined '
         '(April 2022, October 2023, February 2024, January 2026)'),
        ('', '',
         '(c) If relevant, consider translating estimates into meaningful clinical measures',
         'Results: percentage price decline (31%) and dollar values reported'),
        ('Other analyses', '17',
         'Report other analyses done\u2014e.g. subgroup, interaction, or sensitivity analyses',
         'Results: supplementary asking-price cross-sectional analysis; quarterly trend analysis'),
        ('Discussion', '', '', ''),
        ('Key results', '18',
         'Summarise key results with reference to study objectives',
         'Discussion, paragraph 1'),
        ('Limitations', '19',
         'Discuss limitations, potential bias, imprecision, and multiplicity of analyses',
         'Discussion, \u201cLimitations\u201d paragraph: single marketplace, no condition data, '
         'small post-ban n, global marketplace, no pre-proposal baseline'),
        ('Interpretation', '20',
         'Give a cautious overall interpretation considering objectives, limitations, '
         'multiplicity of analyses, similar studies, and other relevant evidence',
         'Discussion, paragraphs 2\u20134: comparison with literature, regulatory obsolescence'),
        ('Generalisability', '21',
         'Discuss the generalisability (external validity) of the study results',
         'Discussion, final paragraphs: implications for facilities in other jurisdictions'),
        ('Other information', '', '', ''),
        ('Funding', '22',
         'Give the source of funding and the role of the funders',
         'Acknowledgements: no external funding'),
    ]

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ['Section/Topic', 'Item No.', 'Recommendation', 'Reported on page/section']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for section, item_no, recommendation, location in items:
        row = table.add_row()
        cells = row.cells

        # Section heading rows
        if recommendation == '' and location == '' and item_no == '':
            for cell in cells:
                set_cell_shading(cell, 'E8F0FE')
            p = cells[0].paragraphs[0]
            run = p.add_run(section)
            run.bold = True
            run.font.size = Pt(8)
            continue

        for i, (text, align) in enumerate([
            (section, WD_ALIGN_PARAGRAPH.LEFT),
            (item_no, WD_ALIGN_PARAGRAPH.CENTER),
            (recommendation, WD_ALIGN_PARAGRAPH.LEFT),
            (location, WD_ALIGN_PARAGRAPH.LEFT),
        ]):
            p = cells[i].paragraphs[0]
            p.alignment = align
            run = p.add_run(text)
            run.font.size = Pt(8)

    path = os.path.join(out_dir, 'eja_strobe_checklist.docx')
    doc.save(path)
    print(f"EJA STROBE checklist saved: {path}")
    return path


if __name__ == '__main__':
    create_strobe_checklist()
