"""
Generate EJA (European Journal of Anaesthesiology) cover letter as editable .docx file.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def add_para(doc, text, size=Pt(11), bold=False, italic=False,
             alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p


def write_eja_cover_letter():
    doc = setup_doc()

    # Date
    add_para(doc, '[Date]', space_after=Pt(12))

    # Addressee
    add_para(doc, 'The Editor-in-Chief', bold=True, space_after=Pt(0))
    add_para(doc, 'European Journal of Anaesthesiology', italic=True, space_after=Pt(0))
    add_para(doc, '', space_after=Pt(12))

    # Subject line
    p = doc.add_paragraph()
    run = p.add_run('Re: ')
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(
        'Submission of original article \u2013 '
        '\u201cEconomic consequences of the European Union desflurane ban for '
        'anaesthetic equipment: a cross-sectional time-series analysis of '
        'secondary market vaporiser prices\u201d')
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)

    # Salutation
    add_para(doc, 'Dear Editor,', space_after=Pt(12))

    # Body paragraphs
    add_para(doc,
        'We are pleased to submit the above manuscript for consideration for publication '
        'in the European Journal of Anaesthesiology as an original article.',
        space_after=Pt(8))

    add_para(doc,
        'The European Union\u2019s prohibition of desflurane for routine anaesthesia '
        '(Regulation (EU) 2024/573), effective 1 January 2026, represents the first mandatory '
        'governmental restriction on a specific anaesthetic agent based on environmental grounds. '
        'While the clinical and environmental rationale for this regulation has been extensively '
        'debated, the downstream economic consequences for existing equipment remain unexplored.',
        space_after=Pt(8))

    add_para(doc,
        'In this study, we analysed 1,033 completed eBay sales of anaesthetic vaporisers '
        '(desflurane, sevoflurane and isoflurane) over three years, spanning the full legislative '
        'trajectory from the EC proposal through to post-ban implementation. Using complementary '
        'statistical approaches (Spearman rank correlation, Kendall \u03c4 trend test, '
        'Mann\u2013Whitney U and between-agent effect size comparison), we provide the first '
        'empirical evidence that the EU desflurane regulation was associated with a progressive, '
        'agent-specific decline in secondary market vaporiser prices. The desflurane effect size '
        'was significantly larger than that of sevoflurane (P=0.043), and sevoflurane and isoflurane '
        'vaporiser prices remained stable throughout, serving as natural controls.',
        space_after=Pt(8))

    add_para(doc,
        'Critically, a substantial proportion of the price depreciation occurred during the '
        'legislative process itself, before the ban took effect. This anticipatory market '
        'response suggests that the regulation was perceived as credible and well signalled, '
        'generating predictable economic consequences. Our data indicate that early compliance '
        'during the legislative phase\u2014rather than delayed action at the point of '
        'prohibition\u2014was associated with better cost recovery, freeing capital for '
        'reinvestment in alternative equipment and thereby supporting broader anaesthetic '
        'choice and patient care.',
        space_after=Pt(8))

    add_para(doc,
        'We believe this work is particularly well suited for EJA because: (1) it directly '
        'addresses the consequences of European anaesthesia policy, specifically the EU F-gas '
        'Regulation which is of primary relevance to EJA\u2019s readership; (2) it provides '
        'empirical evidence that well-designed environmental regulation produces predictable, '
        'measurable economic outcomes\u2014a finding of interest to both clinicians and '
        'policymakers; (3) it demonstrates that early compliance yields economic advantages '
        'that can support equipment reinvestment and patient care; and (4) it offers a '
        'generalisable framework for anticipating equipment depreciation under future '
        'environmental regulations (e.g. nitrous oxide).',
        space_after=Pt(8))

    add_para(doc,
        'This manuscript has not been published elsewhere and is not under consideration by '
        'another journal. All authors have read and approved the final manuscript and meet '
        'ICMJE authorship criteria. There are no conflicts of interest to declare. '
        'No ethical approval was required as the study analysed publicly available, anonymised '
        'marketplace data with no human participant involvement.',
        space_after=Pt(8))

    add_para(doc,
        'The study is reported in accordance with the STROBE guidelines for cross-sectional '
        'studies. The completed STROBE checklist is provided as supplementary material.',
        space_after=Pt(8))

    add_para(doc,
        'Thank you for considering this manuscript. We look forward to your response.',
        space_after=Pt(12))

    # Closing
    add_para(doc, 'Yours sincerely,', space_after=Pt(24))

    add_para(doc, '[Corresponding author name]', bold=True, space_after=Pt(0))
    add_para(doc, '[Department]', space_after=Pt(0))
    add_para(doc, '[Institution]', space_after=Pt(0))
    add_para(doc, '[Address]', space_after=Pt(0))
    add_para(doc, '[Email]', space_after=Pt(0))

    path = os.path.join(out_dir, 'eja_cover_letter.docx')
    doc.save(path)
    print(f"EJA cover letter saved: {path}")
    return path


if __name__ == '__main__':
    write_eja_cover_letter()
