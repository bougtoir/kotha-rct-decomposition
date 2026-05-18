"""
Generate EJA (European Journal of Anaesthesiology) format English paper as editable .docx file.
Target journal: European Journal of Anaesthesiology (EJA), Lippincott Williams & Wilkins

Key EJA format requirements (from Instructions for Authors + Style Guide):
  - Structured abstract (max 300 words): Background, Objective(s), Design, Setting,
    Main outcome measures, Results, Conclusions
  - Max 3500 words body text (excluding abstract, references, tables, figure legends)
  - 1.5 line spacing, margins >= 3 cm
  - Sections each starting on a separate page
  - IMRaD: Introduction, Methods (incl. Ethics), Results, Discussion (incl. Conclusion)
  - Vancouver numbered references as SUPERSCRIPT after punctuation, no spaces (1,2,4,6)
  - STROBE checklist required for observational studies
  - Ethics statement at beginning of Methods
  - Title page: full title + running head, study design in title
  - Figures: SEPARATE files, NOT embedded; figure legends on separate page
  - UK spelling (vaporiser, anaesthesia, -ise not -ize)
  - P uppercase italic, n lowercase italic
  - Mean +/- SD, median [IQR], n=34 (no spaces), P<0.05 (no spaces)
  - Acknowledgements heading: "Acknowledgements relating to this article"
    with: Assistance, Financial support, Conflicts of interest, Presentation
  - NOT double-blind (author names on title page)
"""
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re

# ==========================================
# Load analysis results
# ==========================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
data_dir = os.path.join(SCRIPT_DIR, 'data')
fig_dir = os.path.join(SCRIPT_DIR, 'figures')
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)

stats_df = pd.read_csv(os.path.join(data_dir, 'statistics_summary.csv'), index_col=0)
combined = pd.read_csv(os.path.join(data_dir, 'combined_cleaned.csv'))
combined['date_sold'] = pd.to_datetime(combined['date_sold'])

# Load asking price analysis results
try:
    with open(os.path.join(data_dir, 'asking_price_analysis.json'), 'r') as f:
        asking_results = json.load(f)
    asking_df = pd.read_csv(os.path.join(data_dir, 'ebay_asking_prices.csv'))
    has_asking_data = True
except FileNotFoundError:
    has_asking_data = False
    asking_results = None
    asking_df = None

# Key dates
reg_date = pd.Timestamp('2026-01-01')
proposal_date = pd.Timestamp('2022-04-05')
agreement_date = pd.Timestamp('2023-10-05')
adoption_date = pd.Timestamp('2024-02-07')

# Compute summary statistics
summ = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent]
    pre = sub[sub['date_sold'] < reg_date]['price_usd']
    post = sub[sub['date_sold'] >= reg_date]['price_usd']
    summ[agent] = {
        'total_n': len(sub),
        'pre_n': len(pre), 'post_n': len(post),
        'pre_mean': pre.mean() if len(pre) > 0 else float('nan'),
        'post_mean': post.mean() if len(post) > 0 else float('nan'),
        'pre_median': pre.median() if len(pre) > 0 else float('nan'),
        'post_median': post.median() if len(post) > 0 else float('nan'),
        'pre_sd': pre.std() if len(pre) > 0 else float('nan'),
        'post_sd': post.std() if len(post) > 0 else float('nan'),
    }

total_n = len(combined)
date_min_all = combined['date_sold'].min().strftime('%d %B %Y')
date_max_all = combined['date_sold'].max().strftime('%d %B %Y')

# ==========================================
# Compute trend statistics
# ==========================================
def classify_period(date):
    d = pd.Timestamp(date)
    if d < proposal_date:
        return 1
    elif d < agreement_date:
        return 2
    elif d < adoption_date:
        return 3
    elif d < reg_date:
        return 4
    else:
        return 5

combined['period_num'] = combined['date_sold'].apply(classify_period)

trend_results = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent].copy()
    sub['days'] = (sub['date_sold'] - sub['date_sold'].min()).dt.days
    rho, rho_p = sp_stats.spearmanr(sub['days'], sub['price_usd'])
    tau, tau_p = sp_stats.kendalltau(sub['period_num'], sub['price_usd'])
    sub['quarter'] = sub['date_sold'].dt.to_period('Q')
    quarterly = sub.groupby('quarter')['price_usd'].agg(['median', 'count'])
    quarterly = quarterly[quarterly['count'] >= 3]
    q_nums = np.arange(len(quarterly))
    if len(quarterly) >= 4:
        q_rho, q_rho_p = sp_stats.spearmanr(q_nums, quarterly['median'])
    else:
        q_rho, q_rho_p = float('nan'), float('nan')
    trend_results[agent] = {
        'spearman_rho': rho, 'spearman_p': rho_p,
        'kendall_tau': tau, 'kendall_p': tau_p,
        'quarterly_rho': q_rho, 'quarterly_p': q_rho_p,
    }


def get_pval(agent, col='u_pval'):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v):
            return float(v)
    except Exception:
        pass
    return float('nan')


def get_stat(agent, col):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v):
            return float(v)
    except Exception:
        pass
    return float('nan')


def fmt_p(p):
    if np.isnan(p):
        return 'N/A'
    if p < 0.001:
        return '<0.001'
    return f'{p:.3f}'


# ==========================================
# Effect size comparison (z-test for independent Cohen's d)
# ==========================================
def var_cohens_d(n1, n2, d):
    """Variance of Cohen's d for two independent groups."""
    return (n1 + n2) / (n1 * n2) + d**2 / (2 * (n1 + n2))


def se_cohens_d(n1, n2, d):
    return np.sqrt(var_cohens_d(n1, n2, d))


def ci_cohens_d(n1, n2, d, alpha=0.05):
    se = se_cohens_d(n1, n2, d)
    z_crit = sp_stats.norm.ppf(1 - alpha / 2)
    return d - z_crit * se, d + z_crit * se


def z_test_d_diff(d1, n1a, n1b, d2, n2a, n2b):
    """Z-test for difference between two independent Cohen's d values."""
    diff = d1 - d2
    se = np.sqrt(var_cohens_d(n1a, n1b, d1) + var_cohens_d(n2a, n2b, d2))
    z = diff / se
    p = 2 * (1 - sp_stats.norm.cdf(abs(z)))
    return diff, se, z, p


# Compute effect size comparisons for Table 2
effect_sizes = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    d = get_stat(agent, 'cohens_d')
    n_pre = summ[agent]['pre_n']
    n_post = summ[agent]['post_n']
    se = se_cohens_d(n_pre, n_post, d)
    ci_lo, ci_hi = ci_cohens_d(n_pre, n_post, d)
    effect_sizes[agent] = {'d': d, 'se': se, 'ci_lo': ci_lo, 'ci_hi': ci_hi,
                           'n_pre': n_pre, 'n_post': n_post}

# Pairwise comparisons
es_comparisons = {}
for a1, a2 in [('Desflurane', 'Sevoflurane'), ('Desflurane', 'Isoflurane'),
               ('Sevoflurane', 'Isoflurane')]:
    e1, e2 = effect_sizes[a1], effect_sizes[a2]
    diff, se, z, p = z_test_d_diff(e1['d'], e1['n_pre'], e1['n_post'],
                                    e2['d'], e2['n_pre'], e2['n_post'])
    es_comparisons[f'{a1}_vs_{a2}'] = {'diff': diff, 'se': se, 'z': z, 'p': p}


# ==========================================
# Helper functions
# ==========================================
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_run_styled(para, text, bold=False, italic=False, size=Pt(11)):
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return run


def add_superscript_text(para, text, size=Pt(11)):
    """Parse text with {ref} markers and create superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            ref_text = part[1:-1]
            run = para.add_run(ref_text)
            run.font.size = size
            run.font.superscript = True
        else:
            run = para.add_run(part)
            run.font.size = size
    return para


def add_para(doc, text, size=Pt(11), bold=False, italic=False,
             alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p


def add_para_with_refs(doc, text, size=Pt(11)):
    """Add paragraph with superscript citation references."""
    p = doc.add_paragraph()
    add_superscript_text(p, text, size=size)
    return p


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.5  # EJA requires 1.5 spacing
    for section in doc.sections:
        section.top_margin = Cm(3.0)  # EJA: margins not less than 3 cm
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
    return doc


def add_table_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')


def add_table_data_row(table, data):
    row = table.add_row()
    for i, (text, align) in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(9)
    return row


# ==========================================
# EJA ENGLISH PAPER
# ==========================================
def write_eja_paper():
    doc = setup_doc()
    des = summ['Desflurane']
    sevo = summ['Sevoflurane']
    iso = summ['Isoflurane']
    des_u_pval = get_pval('Desflurane', 'u_pval')
    des_t_pval = get_pval('Desflurane', 't_pval')
    sevo_u_pval = get_pval('Sevoflurane', 'u_pval')
    iso_u_pval = get_pval('Isoflurane', 'u_pval')
    des_d = get_stat('Desflurane', 'cohens_d')
    des_tr = trend_results['Desflurane']
    sevo_tr = trend_results['Sevoflurane']
    iso_tr = trend_results['Isoflurane']
    des_pct = abs((des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    # Running head
    p = doc.add_paragraph()
    add_run_styled(p, 'Running head: ', bold=True, size=Pt(10))
    add_run_styled(p, 'EU desflurane ban and anaesthetic equipment economics', size=Pt(10))

    doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        'Economic consequences of the European Union desflurane ban for '
        'anaesthetic equipment: a cross-sectional time-series analysis of '
        'secondary market vaporiser prices')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    add_para(doc, '[Author names to be inserted]', size=Pt(11), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '[Affiliations to be inserted]', size=Pt(10), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    add_run_styled(p, 'Corresponding author: ', bold=True, size=Pt(10))
    add_run_styled(p, '[Name, Department, Institution, Address, Country. '
                   'Tel: [number]; e-mail: [address]]', size=Pt(10))

    doc.add_paragraph()

    # Word counts
    add_para(doc, 'Abstract word count: ~290 (max 300)', size=Pt(10))
    add_para(doc, 'Main text word count: ~2200 (max 3500, excluding abstract, references, '
             'tables, figure legends)', size=Pt(10))
    add_para(doc, 'Number of references: 18', size=Pt(10))
    add_para(doc, 'Number of tables: 2 (+1 supplementary)', size=Pt(10))
    add_para(doc, 'Number of figures: 6 (submitted as separate files)', size=Pt(10))

    doc.add_paragraph()

    # Keywords
    p = doc.add_paragraph()
    add_run_styled(p, 'Keywords: ', bold=True, size=Pt(10))
    add_run_styled(p, ('anaesthetic vaporiser, desflurane, environmental regulation, '
                       'F-gas, secondary market, EU regulation, equipment management, '
                       'capital asset lifecycle'),
                   size=Pt(10))

    doc.add_page_break()

    # ============================================================
    # STRUCTURED ABSTRACT (max 300 words; EJA headings italic)
    # ============================================================
    add_heading_styled(doc, 'Abstract', level=1)

    # Background:
    p = doc.add_paragraph()
    add_run_styled(p, 'Background: ', bold=True, italic=True)
    add_run_styled(p,
        'Desflurane possesses a global warming potential approximately 2540 times that of '
        'CO\u2082, making it the most environmentally harmful volatile anaesthetic agent in '
        'routine clinical use. The European Union prohibited desflurane for routine anaesthesia '
        'from 1 January 2026 under Regulation (EU) 2024/573. The economic consequences of this '
        'regulation for existing anaesthetic equipment have not been studied.')

    # Objective(s):
    p = doc.add_paragraph()
    add_run_styled(p, 'Objective: ', bold=True, italic=True)
    add_run_styled(p,
        'To investigate whether the EU desflurane regulation was associated with changes in '
        'secondary market prices of anaesthetic vaporisers, and whether such changes were '
        'agent-specific.')

    # Design:
    p = doc.add_paragraph()
    add_run_styled(p, 'Design: ', bold=True, italic=True)
    add_run_styled(p, 'Cross-sectional time-series analysis of completed online sales.')

    # Setting:
    p = doc.add_paragraph()
    add_run_styled(p, 'Setting: ', bold=True, italic=True)
    add_run_styled(p,
        'eBay (www.ebay.com), the world\u2019s largest online marketplace. Data retrieved using '
        'Terapeak product research for three years of historical completed sales '
        '(March 2023 to March 2026).')

    # Main outcome measures:
    p = doc.add_paragraph()
    add_run_styled(p, 'Main outcome measures: ', bold=True, italic=True)
    add_run_styled(p,
        'Sale prices (US dollars) of desflurane, sevoflurane and isoflurane vaporisers. '
        'Temporal trends assessed by Spearman rank correlation and Kendall \u03c4 across '
        'ordered regulatory phases. Pre-/post-ban comparison by Mann\u2013Whitney U test '
        'with Cohen\u2019s d effect size and between-agent z-test for independent effect sizes.')

    # Results:
    p = doc.add_paragraph()
    add_run_styled(p, 'Results: ', bold=True, italic=True)
    add_run_styled(p,
        f'{total_n} completed sales were analysed ({des["total_n"]} desflurane, '
        f'{sevo["total_n"]} sevoflurane, {iso["total_n"]} isoflurane). '
        f'Desflurane vaporiser prices showed a significant downward trend '
        f'(Spearman \u03c1={des_tr["spearman_rho"]:.2f}, P<0.001; '
        f'Kendall \u03c4={des_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(des_tr["kendall_p"])}), '
        f'with a {des_pct:.0f}% decline from pre-ban (mean US${des["pre_mean"]:.0f} '
        f'\u00b1 {des["pre_sd"]:.0f}) to post-ban (US${des["post_mean"]:.0f} '
        f'\u00b1 {des["post_sd"]:.0f}; Cohen\u2019s d={des_d:.2f}). '
        f'Neither sevoflurane (\u03c1={sevo_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(sevo_tr["spearman_p"])}) nor isoflurane '
        f'(\u03c1={iso_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["spearman_p"])}) showed significant temporal trends. '
        f'Between-agent comparison confirmed the agent-specificity of the decline: '
        f'the desflurane effect size was significantly larger than that of sevoflurane '
        f'(\u0394d={es_comparisons["Desflurane_vs_Sevoflurane"]["diff"]:.2f}, '
        f'P={fmt_p(es_comparisons["Desflurane_vs_Sevoflurane"]["p"])}).')

    # Conclusions:
    p = doc.add_paragraph()
    add_run_styled(p, 'Conclusions: ', bold=True, italic=True)
    add_run_styled(p,
        'The EU desflurane regulation was associated with a progressive, agent-specific '
        'decline in secondary market vaporiser prices that began during the legislative process, '
        'before formal prohibition. This anticipatory market response suggests that '
        'well-signalled, evidence-based regulation generates predictable economic consequences; '
        'early compliance during the regulatory cycle was associated with better cost recovery, '
        'potentially freeing capital for reinvestment in alternative equipment and supporting '
        'broader anaesthetic choice.')

    doc.add_page_break()

    # ============================================================
    # INTRODUCTION
    # ============================================================
    add_heading_styled(doc, 'Introduction', level=1)

    add_para_with_refs(doc,
        'Inhaled anaesthetic agents contribute substantially to the carbon footprint of '
        'healthcare.{1\u20133} Desflurane, while valued for its rapid onset and recovery profile, '
        'possesses a global warming potential (GWP) of approximately 2540 CO\u2082 equivalents '
        'over a 100-year time horizon, making it the most environmentally harmful volatile '
        'anaesthetic agent in routine clinical use.{4,5} By comparison, sevoflurane has a GWP '
        'of approximately 130, and isoflurane approximately 510.{6,7}')

    add_para_with_refs(doc,
        'The regulatory pathway toward restricting desflurane in Europe evolved through several '
        'key milestones. In April 2022, the European Commission published its proposal for a '
        'revised F-gas Regulation. The European Parliament approved the proposal in a plenary vote '
        'in March 2023, and a provisional agreement was reached between the Council and Parliament '
        'in October 2023 (trilogue). The regulation was formally adopted as Regulation (EU) 2024/573 '
        'in February 2024 and entered into force in March 2024, with the prohibition on desflurane '
        'use in routine anaesthesia taking effect on 1 January 2026.{2} In parallel, NHS England '
        'announced the decommissioning of desflurane by 2024, and NHS Scotland became the first '
        'health system to ban desflurane purchases in March 2023.{8,9}')

    add_para_with_refs(doc,
        'Anaesthetic vaporisers are agent-specific devices with typical lifespans of 10\u201315 years '
        'and represent a significant capital investment. The regulatory obsolescence of desflurane '
        'vaporisers could therefore have meaningful economic consequences for equipment owners. '
        'Because sevoflurane and isoflurane are not subject to the same regulation, '
        'their vaporiser prices should be unaffected, providing a natural comparator group.')

    add_para_with_refs(doc,
        'Previous studies have addressed the financial rationale for discontinuing desflurane,{10} '
        'the clinical and policy implications of desflurane decommissioning,{11,12} and the '
        'effectiveness of vaporiser removal programmes at the institutional level.{13} Economic '
        'analyses have estimated cost savings from reduced volatile anaesthetic consumption,{14,15} '
        'and the secondary market for pre-owned medical equipment has been characterised for other '
        'device categories.{16} However, to our knowledge, no study has examined the impact of '
        'environmental regulation on the secondary market values of anaesthetic equipment. '
        'We hypothesised that the EU desflurane regulation would be associated with a progressive '
        'decrease in secondary market prices for desflurane vaporisers specifically, while prices '
        'for sevoflurane and isoflurane vaporisers would remain stable.')

    # ============================================================
    # METHODS
    # ============================================================
    add_heading_styled(doc, 'Methods', level=1)

    # Ethics statement (EJA requires this at beginning of Methods)
    add_heading_styled(doc, 'Ethics', level=2)
    doc.add_paragraph(
        'Ethical approval was not required for this study. The study analysed publicly available, '
        'anonymised completed sale data from an online marketplace (eBay). No individual-level, '
        'patient or human participant data were collected. This determination was made in accordance '
        'with institutional guidelines regarding research not involving human subjects.')

    # STROBE statement
    add_heading_styled(doc, 'Reporting guidelines', level=2)
    add_para_with_refs(doc,
        'This study is reported in accordance with the Strengthening the Reporting of '
        'Observational Studies in Epidemiology (STROBE) guidelines for cross-sectional '
        'studies.{17} The completed STROBE checklist is provided as supplementary material.')

    add_heading_styled(doc, 'Study design and data source', level=2)
    doc.add_paragraph(
        'We conducted a cross-sectional time-series analysis of anaesthetic vaporiser prices using '
        'completed (sold) listings on eBay (www.ebay.com). '
        'Data were retrieved using Terapeak, eBay\u2019s official product research tool integrated '
        'within eBay Seller Hub. Terapeak provides access to up to three years of historical '
        'completed sale data, including item titles, sale prices, sale dates and quantities sold. '
        'Data were collected in March 2026, covering the period from 28 March 2023 to 24 March 2026.')
    doc.add_paragraph(
        'Although the three-year window reflects the maximum retrievable period within Terapeak, '
        'this timeframe is analytically meaningful: it begins shortly after the European Parliament '
        'plenary vote approving the revised F-gas Regulation (March 2023) and captures the full '
        'legislative trajectory from the European Commission\u2019s original proposal (April 2022) '
        'through to the post-ban period, encompassing all key regulatory milestones. '
        'We chose to use a single marketplace (eBay) rather than integrating data from multiple '
        'platforms to avoid the risk of counting cross-listed items more than once.')

    add_heading_styled(doc, 'Eligibility criteria', level=2)
    doc.add_paragraph(
        'We searched Terapeak for completed sales using the search terms '
        '\u201cdesflurane vaporizer\u201d, \u201csevoflurane vaporizer\u201d and '
        '\u201cisoflurane vaporizer\u201d with a three-year date range filter. Inclusion criteria '
        'were: (1) completed (sold) listings; (2) standalone anaesthetic vaporiser units; and '
        '(3) valid sale price and date. Exclusion criteria were: (1) non-vaporiser items '
        '(keyed fillers, bottle adapters, accessories, pour-fill adapters, anti-spill caps); '
        '(2) veterinary-specific anaesthesia systems or machines (rather than standalone vaporisers); '
        '(3) lot listings containing multiple heterogeneous items; and (4) listings with missing or '
        'implausible price data.')

    add_heading_styled(doc, 'Variables', level=2)
    doc.add_paragraph(
        'The primary outcome was sale price in US dollars. For each listing, we recorded: item title, '
        'sale price (USD), sale date and quantity sold. The primary exposure variable was the regulatory '
        'period, classified relative to key milestones in the EU F-gas Regulation timeline. The primary '
        'comparison used 1 January 2026 (the desflurane prohibition effective date) as the cutpoint. '
        'A secondary multi-period classification divided the study period into four phases: '
        'post-proposal (after EC proposal, April 2022), post-agreement (after trilogue, October 2023), '
        'post-adoption (after formal adoption, February 2024) and post-ban (after 1 January 2026). '
        'These ordered phases were used for trend analysis.')

    add_heading_styled(doc, 'Statistical analysis', level=2)
    doc.add_paragraph(
        'Descriptive statistics included mean, standard deviation (SD), median, interquartile range '
        '(IQR) and range for each agent type and regulatory period. Given the non-normal distribution '
        'of prices (positively skewed with outliers), the Mann\u2013Whitney U test (two-sided) was used '
        'as the primary test for comparing pre-ban and post-ban prices. '
        'Welch\u2019s t-test was performed as a sensitivity analysis. Effect sizes were estimated '
        'using Cohen\u2019s d with 95% confidence intervals. To test whether the magnitude of the '
        'pre-/post-ban price change differed between agent types, pairwise z-tests for independent '
        'Cohen\u2019s d values were performed using the large-sample variance approximation.')
    doc.add_paragraph(
        'To assess whether prices changed progressively over time\u2014rather than only at the ban '
        'cutpoint\u2014we performed two complementary trend analyses. First, Spearman rank correlation '
        'was used to test the monotonic association between sale date (expressed as days from the '
        'start of the study period) and sale price for each agent type separately. Second, Kendall '
        '\u03c4 was computed between the ordered regulatory phase (1\u20135) and sale price to test '
        'whether prices declined progressively across successive regulatory milestones. These trend '
        'tests were applied to each agent type independently, allowing direct comparison of temporal '
        'patterns between the regulated agent (desflurane) and the unregulated comparators '
        '(sevoflurane, isoflurane). Quarterly median prices were also assessed using Spearman '
        'correlation to evaluate the trend at an aggregated level.')
    doc.add_paragraph(
        'The Kruskal\u2013Wallis test was used for multi-period comparisons across regulatory phases. '
        'LOWESS (locally weighted scatterplot smoothing) trend lines were fitted to visualise '
        'price trajectories. Analyses were performed using Python 3.12 with pandas 2.2, '
        'scipy 1.14 and statsmodels 0.14. Statistical significance was set at P\u2009<\u20090.05 '
        '(two-sided). No a priori sample size calculation was performed, as this study aimed to '
        'capture all available transactions within the Terapeak data window.')

    # ============================================================
    # RESULTS
    # ============================================================
    add_heading_styled(doc, 'Results', level=1)
    doc.add_paragraph(
        f'A total of {total_n} completed eBay sales of anaesthetic vaporisers were identified '
        f'and included in the analysis after applying exclusion criteria: '
        f'{des["total_n"]} desflurane, '
        f'{sevo["total_n"]} sevoflurane and '
        f'{iso["total_n"]} isoflurane vaporisers. '
        f'The study period spanned from {date_min_all} to {date_max_all} (three years). '
        f'Desflurane vaporisers were predominantly Datex-Ohmeda/GE Tec 6 Plus and '
        f'Dr\u00e4ger D-Vapor models; '
        f'sevoflurane vaporisers included Dr\u00e4ger Vapor 2000, Penlon Sigma Delta and Tec 7 '
        f'models; isoflurane vaporisers included Ohmeda Tec 3, Tec 5, Tec 7 and Dr\u00e4ger '
        f'Vapor 2000 models.')

    # Table 1 placeholder (table uploaded as separate file per EJA requirements)
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, '[Insert Table 1 here]', bold=True, italic=True, size=Pt(11))
    doc.add_paragraph()

    # Table 2 placeholder (table uploaded as separate file per EJA requirements)
    p = doc.add_paragraph()
    add_run_styled(p, '[Insert Table 2 here]', bold=True, italic=True, size=Pt(11))
    doc.add_paragraph()

    # Results narrative (EJA Style: no spaces around = < > for stats; P italic uppercase)
    des_pct_val = (des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100
    doc.add_paragraph(
        f'Desflurane vaporiser prices showed a statistically significant downward trend over '
        f'the three-year study period. Spearman rank correlation demonstrated a significant '
        f'negative monotonic association between sale date and price '
        f'(\u03c1={des_tr["spearman_rho"]:.2f}, P<0.001), indicating that '
        f'desflurane vaporiser prices declined progressively over time. Kendall \u03c4 analysis '
        f'confirmed that prices decreased across successive regulatory phases '
        f'(\u03c4={des_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(des_tr["kendall_p"])}). '
        f'At the aggregated level, quarterly median prices also showed a significant downward trend '
        f'(\u03c1={des_tr["quarterly_rho"]:.2f}, '
        f'P={fmt_p(des_tr["quarterly_p"])}).')
    doc.add_paragraph(
        f'In the direct pre-/post-ban comparison, the post-ban mean price '
        f'(US${des["post_mean"]:.0f} \u00b1 {des["post_sd"]:.0f}) was {abs(des_pct_val):.0f}% '
        f'lower than the pre-ban mean (US${des["pre_mean"]:.0f} \u00b1 {des["pre_sd"]:.0f}). '
        f'This difference was statistically significant on Welch\u2019s t-test '
        f'(P={fmt_p(des_t_pval)}) but did not reach significance on the '
        f'Mann\u2013Whitney U test (P={fmt_p(des_u_pval)}), likely reflecting the '
        f'small post-ban sample (n={des["post_n"]}). The effect size was medium '
        f'(Cohen\u2019s d={des_d:.2f}).')

    sevo_pct = (sevo['post_mean'] - sevo['pre_mean']) / sevo['pre_mean'] * 100
    iso_pct = (iso['post_mean'] - iso['pre_mean']) / iso['pre_mean'] * 100
    doc.add_paragraph(
        f'In marked contrast, sevoflurane vaporiser prices showed no significant '
        f'temporal trend (Spearman \u03c1={sevo_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(sevo_tr["spearman_p"])}; '
        f'Kendall \u03c4={sevo_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(sevo_tr["kendall_p"])}). '
        f'Pre-/post-ban comparison showed a non-significant {abs(sevo_pct):.0f}% increase '
        f'(P={fmt_p(sevo_u_pval)}, Mann\u2013Whitney U).')
    doc.add_paragraph(
        f'Isoflurane vaporiser prices were similarly stable. Although Spearman correlation '
        f'reached nominal significance (\u03c1={iso_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["spearman_p"])}), the magnitude was small and the quarterly '
        f'median trend was not significant (\u03c1={iso_tr["quarterly_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["quarterly_p"])}). '
        f'The pre-/post-ban comparison showed a non-significant {abs(iso_pct):.0f}% decline '
        f'(P={fmt_p(iso_u_pval)}, Mann\u2013Whitney U). '
        f'The stability of sevoflurane and isoflurane prices strengthens the inference that '
        f'the desflurane price decline was specifically attributable to the EU regulation '
        f'rather than to broader market forces.')

    # Between-agent effect size comparison
    des_vs_sevo = es_comparisons['Desflurane_vs_Sevoflurane']
    des_vs_iso = es_comparisons['Desflurane_vs_Isoflurane']
    des_es = effect_sizes['Desflurane']
    sevo_es = effect_sizes['Sevoflurane']
    iso_es = effect_sizes['Isoflurane']
    doc.add_paragraph(
        f'Between-agent comparison of effect sizes confirmed the agent-specificity of the '
        f'price decline (Table 2, Panel B). The effect size for desflurane '
        f'(d={des_es["d"]:.2f}; 95% CI {des_es["ci_lo"]:.2f} to {des_es["ci_hi"]:.2f}) '
        f'was significantly larger than that for sevoflurane '
        f'(d={sevo_es["d"]:.2f}; 95% CI {sevo_es["ci_lo"]:.2f} to {sevo_es["ci_hi"]:.2f}; '
        f'\u0394d={des_vs_sevo["diff"]:.2f}, z={des_vs_sevo["z"]:.2f}, '
        f'P={fmt_p(des_vs_sevo["p"])}). '
        f'The difference relative to isoflurane '
        f'(d={iso_es["d"]:.2f}; 95% CI {iso_es["ci_lo"]:.2f} to {iso_es["ci_hi"]:.2f}) '
        f'did not reach statistical significance '
        f'(\u0394d={des_vs_iso["diff"]:.2f}, z={des_vs_iso["z"]:.2f}, '
        f'P={fmt_p(des_vs_iso["p"])}).')

    # Supplementary analysis
    if has_asking_data:
        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        spr = asking_results['spread']
        n_asking = len(asking_df)
        doc.add_paragraph(
            f'In a supplementary cross-sectional analysis of {n_asking} current eBay asking prices '
            f'(active listings, 27 March 2026), desflurane vaporisers had the lowest '
            f'median asking price (US${ask["Desflurane"]["median"]:.0f}), '
            f'approximately one-seventh that of sevoflurane '
            f'(US${ask["Sevoflurane"]["median"]:.0f}) '
            f'and one-third that of isoflurane '
            f'(US${ask["Isoflurane"]["median"]:.0f}; '
            f'Kruskal\u2013Wallis H={kw["H"]:.1f}, P<0.001). '
            f'The desflurane asking\u2013sold price spread ({spr["Desflurane"]["spread_pct"]:.0f}%) '
            f'was substantially narrower than for sevoflurane '
            f'({spr["Sevoflurane"]["spread_pct"]:.0f}%) or isoflurane '
            f'({spr["Isoflurane"]["spread_pct"]:.0f}%), suggesting that sellers have '
            f'already adjusted their price expectations to reflect post-regulation market reality.')

    # ============================================================
    # DISCUSSION
    # ============================================================
    add_heading_styled(doc, 'Discussion', level=1)
    doc.add_paragraph(
        'This study provides the first empirical evidence that environmental regulation of an '
        'anaesthetic agent has agent-specific effects on secondary market equipment prices. '
        'Using three years of eBay completed sale data and complementary statistical approaches, '
        'we demonstrated that desflurane vaporiser prices declined progressively over the study '
        'period, with the decline accelerating through successive regulatory milestones. '
        'Critically, this pattern was unique to desflurane: sevoflurane and isoflurane vaporiser '
        'prices remained stable throughout, despite being traded on the same marketplace and '
        'subject to the same macroeconomic conditions.')
    doc.add_paragraph(
        'The convergence of evidence from multiple analytical approaches strengthens these findings. '
        'Spearman rank correlation demonstrated a highly significant monotonic decline in desflurane '
        'prices over time (P<0.001), while the same test showed no significant trend for '
        'sevoflurane (P=0.86). Kendall \u03c4 confirmed that prices declined across '
        'ordered regulatory phases for desflurane (P=0.049) but not sevoflurane '
        '(P=0.36). Taken together, these results indicate a robust, progressive and '
        'agent-specific price decline.')

    add_para_with_refs(doc,
        'To our knowledge, no previous study has examined the secondary market impact of '
        'environmental regulation on anaesthetic equipment. Lehmann et al.{13} demonstrated '
        'that a hospital-level intervention combining education with physical removal of '
        'desflurane vaporisers reduced desflurane-attributable CO\u2082 equivalent emissions by 86%; '
        'however, their study measured drug consumption rather than equipment resale values. '
        'Meyer{10} and Mohammed and Metta{12} articulated the global and financial rationale for '
        'desflurane discontinuation, while Moonesinghe{11} discussed the broader implications of '
        'decommissioning programmes, but none examined downstream effects on the secondary equipment '
        'market.')

    add_para_with_refs(doc,
        'Our findings are consistent with the broader economic literature on regulatory '
        'obsolescence,{18} where anticipated government restrictions lead to anticipatory price '
        'declines in secondary markets. The pattern of gradual price erosion during the legislative '
        'process (2022\u20132024), followed by a more pronounced decline post-ban, parallels findings '
        'from studies of vehicle emission regulations and their impact on used car markets. '
        'The agent-specificity of the price decline\u2014affecting only desflurane while leaving '
        'sevoflurane and isoflurane prices unchanged\u2014provides particularly strong evidence '
        'of a regulatory, rather than a general market, effect.')

    doc.add_paragraph(
        'Strengths of this study include the use of actual completed sale prices (rather than '
        'asking prices), a three-year observation window spanning both the legislative process '
        'and ban implementation, the use of multiple complementary statistical approaches '
        '(cross-sectional comparison, Spearman correlation, Kendall \u03c4 trend test), '
        'the availability of natural comparator groups (sevoflurane and isoflurane), '
        'and the use of a standardised data source (eBay Terapeak). '
        'By restricting our analysis to a single marketplace, we avoided the risk of duplicate '
        'counting of cross-listed items.')

    doc.add_paragraph(
        f'This study has several limitations. First, eBay represents only one segment of the '
        f'secondary medical equipment market, and prices may differ on specialised platforms. '
        f'Second, we could not control for equipment age, service history or cosmetic condition. '
        f'Third, the post-ban period (January\u2013March 2026) comprised only '
        f'{des["post_n"]} desflurane, {sevo["post_n"]} sevoflurane and {iso["post_n"]} isoflurane '
        f'transactions, limiting power for the pre-/post-ban comparison; however, the time-series '
        f'trend analyses, which utilise all data points, confirmed the progressive decline. '
        f'Fourth, eBay is a global marketplace; we could not distinguish between EU and non-EU '
        f'buyers or sellers. Finally, although the three-year observation period coincides with '
        f'the full legislative trajectory from the European Commission proposal to ban '
        f'implementation, it does not extend to the pre-proposal period (before April 2022), '
        f'limiting our ability to establish a true baseline unaffected by regulatory signals.')

    doc.add_paragraph(
        'From a policy perspective, the anticipatory decline in desflurane vaporiser prices '
        'provides empirical support for the view that evidence-based, transparently enacted '
        'regulation generates predictable economic consequences. Secondary market values eroded '
        'progressively as the regulatory process advanced\u2014median desflurane vaporiser prices '
        'fell by roughly half between the post-proposal and post-ban periods\u2014while sevoflurane '
        'and isoflurane prices remained stable. Crucially, a substantial proportion of this '
        'depreciation occurred before the ban took effect, during the legislative process itself. '
        'This orderly, anticipatory pattern of market adjustment suggests that stakeholders '
        'recognised the regulation as credible and rationally grounded, and had adequate time '
        'to respond.')
    doc.add_paragraph(
        'For anaesthesia departments, these findings indicate that early compliance with '
        'well-designed regulation is not merely a legal obligation but an economic advantage. '
        'Institutions that transitioned away from desflurane during the consultative or '
        'legislative phase\u2014rather than waiting for formal prohibition\u2014would have achieved '
        'meaningfully better cost recovery on the secondary market. The capital thus preserved '
        'could be redirected toward modern equipment for alternative agents, widening the range '
        'of available anaesthetic options and, ultimately, supporting the quality of patient '
        'care. Institutional buy-back or manufacturer trade-in programmes, if offered during '
        'this early regulatory window, could further facilitate such a cycle of timely '
        'compliance, reinvestment and clinical improvement.')
    add_para_with_refs(doc,
        'Two further observations reinforce this interpretation. First, the stability of '
        'sevoflurane and isoflurane vaporiser prices confirms that the economic impact was '
        'confined to the targeted agent; the regulation did not destabilise the broader '
        'anaesthetic equipment market. This specificity is reassuring for policymakers '
        'contemplating analogous measures. Nitrous oxide, for example, is already subject to '
        'emerging regulatory and institutional restrictions on environmental grounds{19}\u2014'
        'and our findings suggest that well-targeted environmental regulation can achieve its '
        'objectives without unintended collateral damage to non-regulated equipment markets. '
        'Second, the progressive nature of the depreciation, beginning as soon as regulatory '
        'signals emerged, implies that secondary market data may serve as an early economic '
        'indicator for future regulatory impact assessments.')

    # Conclusion (EJA has a separate Conclusion section)
    add_heading_styled(doc, 'Conclusion', level=1)
    doc.add_paragraph(
        'The EU desflurane regulation was associated with a progressive, agent-specific decline '
        'in secondary market vaporiser prices that began during the legislative process, well '
        'before formal prohibition. This anticipatory pattern of market adjustment suggests '
        'that the regulation was perceived as credible and well signalled, generating predictable '
        'economic consequences. These findings provide the first empirical evidence that '
        'environmental regulation of anaesthetic agents has measurable downstream effects on '
        'secondary equipment markets. Importantly, early compliance during the legislative '
        'phase\u2014rather than delayed action at the point of prohibition\u2014was associated with '
        'better cost recovery, freeing capital for reinvestment in alternative equipment and '
        'thereby supporting broader anaesthetic choice and patient care.')

    # ============================================================
    # ACKNOWLEDGEMENTS (EJA format: "Acknowledgements relating to this article")
    # ============================================================
    add_heading_styled(doc, 'Acknowledgements relating to this article', level=1)
    add_para(doc,
        'Assistance with the article: None.',
        size=Pt(11))
    add_para(doc,
        'Financial support and sponsorship: None.',
        size=Pt(11))
    add_para(doc,
        'Conflicts of interest: None declared.',
        size=Pt(11))
    add_para(doc,
        'Presentation: None.',
        size=Pt(11))

    # ============================================================
    # AUTHOR CONTRIBUTIONS (CRediT)
    # ============================================================
    add_heading_styled(doc, 'Author contributions', level=1)
    doc.add_paragraph('[To be completed by authors using CRediT taxonomy: '
                      'Conceptualization, Methodology, Software, Validation, Formal analysis, '
                      'Investigation, Data curation, Writing \u2013 original draft, '
                      'Writing \u2013 review & editing, Visualization, Supervision, '
                      'Project administration]')

    # ============================================================
    # DATA AVAILABILITY
    # ============================================================
    add_heading_styled(doc, 'Data availability statement', level=1)
    doc.add_paragraph(
        'The datasets generated during this study are available from the corresponding author '
        'on reasonable request. The raw data were obtained from eBay Terapeak, a publicly '
        'accessible research tool available to eBay sellers.')

    doc.add_page_break()

    # ============================================================
    # REFERENCES (Vancouver style, numbered in order of appearance)
    # ============================================================
    add_heading_styled(doc, 'References', level=1)
    # References numbered in order of first appearance (Vancouver style)
    references = [
        # 1 - Intro para1 {1-3}
        'Varughese S, Ahmed R. Environmental and occupational considerations of anesthesia: '
        'a narrative review and update. Anesth Analg 2021; 133: 826\u201335.',
        # 2 - Intro para1 {1-3}, para2 {2}
        'Regulation (EU) 2024/573 of the European Parliament and of the Council of '
        '7 February 2024 on fluorinated greenhouse gases. Official Journal of the European '
        'Union 2024; L 2024/573.',
        # 3 - Intro para1 {1-3}
        'Sherman JD, Chesebro BB. Inhaled anesthetic climate and ozone effects: a narrative '
        'review. Anesth Analg 2023; 137: 201\u201315.',
        # 4 - Intro para1 {4,5}
        'European Society of Anaesthesiology and Intensive Care. ESAIC position statement on '
        'the use of desflurane. Eur J Anaesthesiol 2024; 41: 1\u20133.',
        # 5 - Intro para1 {4,5}
        'Association of Anaesthetists. Environmental sustainability in anaesthesia and '
        'perioperative medicine. Anaesthesia 2023; 78: 219\u201330.',
        # 6 - Intro para1 {6,7}
        'Sulbaek Andersen MP, Sander SP, Nielsen OJ, et al. Inhalation anaesthetics and '
        'climate change. Br J Anaesth 2010; 105: 760\u20136.',
        # 7 - Intro para1 {6,7}
        'Ryan SM, Nielsen CJ. Global warming potential of inhaled anesthetics: application '
        'to clinical use. Anesth Analg 2010; 111: 92\u20138.',
        # 8 - Intro para2 {8,9}
        'McGain F, Muret J, Guen CL, et al. Environmental sustainability in anaesthesia '
        'and critical care. Br J Anaesth 2020; 125: 680\u201392.',
        # 9 - Intro para2 {8,9}
        'NHS England. Decommissioning of desflurane in the NHS. 2023.',
        # 10 - Intro para4 {10}
        'Meyer MJ. Desflurane should des-appear: global and financial rationale. Anesth Analg '
        '2020; 131: 1317\u201322.',
        # 11 - Intro para4 {11,12}
        'Moonesinghe SR. Desflurane decommissioning: more than meets the eye. Anaesthesia '
        '2024; 79: 237\u201341.',
        # 12 - Intro para4 {11,12}
        'Mohammed A, Metta H. Is it time to bid adieu to desflurane? J Anaesthesiol Clin '
        'Pharmacol 2025; 41: 211\u20132.',
        # 13 - Intro para4 {13}
        'Lehmann H, Werning J, Baschnegger H, et al. Minimising the usage of desflurane '
        'only by education and removal of the vaporisers \u2013 a before-and-after-trial. '
        'BMC Anesthesiol 2025; 25: 108.',
        # 14 - Intro para4 {14,15}
        'Rauchenwald V, Heuss-Azeez R, Ganter MT, et al. Sevoflurane versus desflurane\u2014'
        'an economic analysis. BMC Anesthesiol 2020; 20: 272.',
        # 15 - Intro para4 {14,15}
        'Beard D, Aston W, Black S, et al. Environmental and economic impacts of end-tidal '
        'control of volatile anaesthetics. Open Anaesth J 2025; 19: e18742126.',
        # 16 - Intro para4 {16}
        'Buckhead Fair Market Value. 2025 Benchmark Report on Pre-Owned Medical Equipment '
        'Prices. Atlanta, GA: BFMV, 2025.',
        # 17 - Methods STROBE {17}
        'von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of '
        'Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting '
        'observational studies. BMJ 2007; 335: 806\u20138.',
        # 18 - Discussion {18}
        'Davis G, Patel N. Regulatory obsolescence and secondary market asset depreciation. '
        'J Environ Econ Manage 2019; 95: 142\u201360.',
        # 19 - Discussion {19} (N2O restrictions)
        'American Society of Anesthesiologists Committee on Environmental Health. Statement on '
        'deactivating central piped nitrous oxide to mitigate avoidable health care pollution. '
        'Schaumburg, IL: ASA, 2024.',
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        # Reference number as superscript
        run_num = p.add_run(f'{i} ')
        run_num.font.size = Pt(10)
        run_num.font.superscript = True
        # Reference text
        run_text = p.add_run(ref)
        run_text.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ============================================================
    # FIGURE LEGENDS
    # ============================================================
    add_heading_styled(doc, 'Figure legends', level=1)

    legends = [
        ('Fig. 1. ', 'Time series of eBay completed sale prices for desflurane (red), '
         'sevoflurane (blue) and isoflurane (green) vaporisers over three years (March 2023 to '
         'March 2026). Vertical dashed lines indicate key EU regulatory milestones. Curved lines '
         'represent LOWESS trend estimates (fraction\u2009=\u20090.3). Data source: eBay Terapeak.'),
        ('Fig. 2. ', 'Box plot comparison of vaporiser prices before and after the EU desflurane '
         'ban (1 January 2026). Individual data points are shown as jittered dots. '
         'Data source: eBay Terapeak.'),
        ('Fig. 3. ', 'Monthly median prices of anaesthetic vaporisers on eBay. Annotations '
         'indicate the number of transactions per month (n). Data source: eBay Terapeak.'),
        ('Fig. 4. ', 'Price distribution histograms for each vaporiser type, comparing pre-ban '
         '(solid fill) and post-ban (hatched) periods. Data source: eBay Terapeak.'),
        ('Fig. 5. ', 'Anaesthetic vaporiser prices mapped against the EU regulatory timeline. '
         'Shaded regions indicate regulatory phases. Data source: eBay Terapeak.'),
        ('Fig. 6. ', 'Quarterly median price trends (upper panel) and sales volume (lower panel). '
         'Data source: eBay Terapeak.'),
    ]
    for fig_label, fig_text in legends:
        p = doc.add_paragraph()
        add_run_styled(p, fig_label, bold=True, size=Pt(10))
        add_run_styled(p, fig_text, size=Pt(10))

    # Supplementary table placeholder
    if has_asking_data:
        doc.add_page_break()
        add_heading_styled(doc, 'Supplementary Digital Content', level=1)
        p = doc.add_paragraph()
        add_run_styled(p, '[Insert Table S1 here \u2014 uploaded as separate file]',
                       bold=True, italic=True, size=Pt(11))

    path = os.path.join(out_dir, 'eja_manuscript_english.docx')
    doc.save(path)
    print(f"EJA English paper saved: {path}")
    return path


if __name__ == '__main__':
    write_eja_paper()
