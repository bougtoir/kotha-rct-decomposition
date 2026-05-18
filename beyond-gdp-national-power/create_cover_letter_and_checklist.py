"""
Generate cover letter and STROBE reporting guideline checklist for Cliometrica submission.

Outputs:
  manuscript/cover_letter.docx              — Cover letter to Managing Editor
  manuscript/strobe_checklist.docx           — STROBE checklist for cross-sectional studies
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = os.path.join(os.path.dirname(__file__), "manuscript")
os.makedirs(OUT, exist_ok=True)


def create_cover_letter():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("[Date]")

    doc.add_paragraph()

    # Addressee
    p = doc.add_paragraph()
    p.add_run("Professor Claude Diebolt\nManaging Editor, ").bold = False
    run = p.add_run("Cliometrica")
    run.italic = True
    p.add_run("\nBETA/CNRS, University of Strasbourg\n"
              "61 Avenue de la Forêt Noire\n"
              "67085 Strasbourg Cedex, France")

    doc.add_paragraph()

    # Subject
    p = doc.add_paragraph()
    run = p.add_run("Re: Submission of original manuscript — ")
    run.bold = True
    run = p.add_run('"Network Exclusion and State Collapse: From Maritime Isolation to '
                    'Technological Access Denial in the Long Run of History"')
    run.bold = True

    doc.add_paragraph()

    # Salutation
    doc.add_paragraph("Dear Professor Diebolt,")

    doc.add_paragraph()

    # Body
    doc.add_paragraph(
        "We are pleased to submit the above-titled manuscript for consideration for "
        "publication in Cliometrica. This manuscript has not been published elsewhere and "
        "is not under consideration by any other journal. All authors have approved the "
        "manuscript and agree with its submission to Cliometrica."
    )

    doc.add_paragraph(
        "The manuscript addresses a fundamental question in economic history: why does "
        "exclusion from international exchange networks so reliably precede state collapse? "
        "Using a comparative dataset of 96 historical polities spanning antiquity to the "
        "present, we distinguish between deliberate closure (policy-based trade bans, "
        "sakoku) and what we term 'technical network exclusion'—involuntary disconnection "
        "from the dominant exchange network of an era due to geographic or technological "
        "constraints. We argue that the critical channel is technology flow: technically "
        "excluded polities were severed not just from trade but from the diffusion of "
        "military techniques, institutional innovations, and frontier knowledge."
    )

    doc.add_paragraph(
        "Reclassifying seven technically excluded polities—including the Han Dynasty, "
        "the Khmer Empire, and Kievan Rus'—transforms a non-significant baseline "
        "association between closure and conquest (Fisher's exact test p = 0.187) into a "
        "significant one (p = 0.020). All seven were eventually conquered. A dose–response "
        "gradient emerges across closure types: technical exclusion (zero technology "
        "transfer) shows a 100% conquest rate, policy bans (reduced but non-zero transfer) "
        "show rates below 80%, and open polities show the lowest rates. The core stock–flow "
        "odds ratio (OR = 1.774) remains invariant across all reclassification scenarios."
    )

    doc.add_paragraph(
        "We believe this manuscript offers four features that may interest Cliometrica\u2019s "
        "readership. First, it provides a quantitative restatement of the long-recognized "
        "pattern that contact between civilizations at different technological levels tends "
        "to end unfavorably for the less advanced party (Diamond 1997)\u2014but grounded in a "
        "systematic, cross-historical dataset rather than selective case studies. Second, "
        "we examine polities that pursued conditional closure\u2014restricting broad engagement "
        "while deliberately maintaining selective technology transfer channels (e.g., "
        "Tokugawa Japan\u2019s rangaku, Qing China\u2019s Canton system)\u2014and find mixed outcomes, "
        "suggesting that the policy question is not binary (open or closed) but conditional "
        "on which channels are preserved. Third, the technology-flow mechanism we propose "
        "generalizes beyond geographic isolation: as the dominant network shifts from sea lanes "
        "to semiconductors, AI infrastructure, and advanced robotics, the same logic of "
        "cumulative divergence through exclusion may apply to contemporary states. Fourth, "
        "the sensitivity analysis framework (3 reclassification levels \u00d7 2 outcome "
        "assignments = 6 scenarios) provides a replicable template for other researchers "
        "working with categorical historical data."
    )

    doc.add_paragraph(
        "The manuscript comprises approximately 9,000 words of body text, 4 tables, "
        "4 figures, and one supplementary table (Table S1) listing all 96 polities with "
        "modern-country equivalents, periods of existence, and specific turning-point "
        "events. All data and code are available upon request."
    )

    doc.add_paragraph(
        "We confirm that this work is original, has not been published before, and is not "
        "being considered for publication elsewhere. There are no competing interests to "
        "declare. No external funding was received for this research."
    )

    doc.add_paragraph(
        "We suggest the following researchers as potential reviewers, given their expertise "
        "in quantitative economic history and cliometrics:"
    )

    # Reviewer suggestions
    reviewers = [
        ("Jörg Baten", "University of Tübingen", "quantitative historical data analysis"),
        ("Stephen Broadberry", "University of Oxford", "long-run comparative economic history"),
        ("Peter Turchin", "University of Connecticut / Complexity Science Hub Vienna",
         "quantitative historical dynamics (cliodynamics)"),
    ]
    for name, affil, expertise in reviewers:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}")
        run.bold = True
        p.add_run(f" ({affil}) — {expertise}")

    doc.add_paragraph()

    doc.add_paragraph(
        "Thank you for considering our submission. We look forward to your response."
    )

    doc.add_paragraph()

    # Closing
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[Author Name]\n[Affiliation]\n[Email]\n[ORCID]")

    path = os.path.join(OUT, "cover_letter.docx")
    doc.save(path)
    print(f"  Cover letter saved to {path}")


def create_strobe_checklist():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(2)

    doc.add_heading(
        "STROBE Statement — Checklist of Items for Reports of Cross-Sectional Studies",
        level=1,
    )
    p = doc.add_paragraph(
        'Applied to: "Network Exclusion and State Collapse: From Maritime Isolation to '
        'Technological Access Denial in the Long Run of History"'
    )
    p.runs[0].italic = True

    doc.add_paragraph()

    # STROBE items for cross-sectional studies
    items = [
        # (Item No, Section, Recommendation, Reported on page/section, Comment)
        ("1(a)", "Title and abstract",
         "Indicate the study's design with a commonly used term in the title or the abstract",
         "Title page; Abstract",
         'Title indicates scope ("in the Long Run of History"); Abstract describes comparative historical design (N=96 polities, six eras)'),

        ("1(b)", "Title and abstract",
         "Provide in the abstract an informative and balanced summary of what was done and what was found",
         "Abstract",
         "Abstract reports methods (Fisher test, logistic regression, bootstrap), key results (p=0.020, OR=1.774), and interpretation"),

        ("2", "Background/rationale",
         "Explain the scientific background and rationale for the investigation being reported",
         "Section 1 (Introduction)",
         "Reviews literature on trade openness and state survival; identifies gap in treatment of involuntary isolation"),

        ("3", "Objectives",
         "State specific objectives, including any prespecified hypotheses",
         "Section 1 (Introduction), ¶3",
         "Three contributions stated: (1) construct comparative dataset, (2) introduce technical network exclusion concept, (3) conduct sensitivity analysis"),

        ("4", "Study design",
         "Present key elements of study design early in the paper",
         "Section 2 (Data and Classification)",
         "Cross-sectional comparative design with 96 polities; variables defined; classification framework presented"),

        ("5", "Setting",
         "Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection",
         "Section 2.1",
         "Six eras from antiquity to contemporary; each polity assigned period of existence; full details in Table S1"),

        ("6(a)", "Participants",
         "Give the eligibility criteria, and the sources and methods of selection of participants",
         "Section 2.1",
         "Polities selected from standard reference works (Findlay & O'Rourke 2007; Kennedy 1987; Turchin 2009); inclusion criteria: identifiable dominant strategy and outcome"),

        ("7", "Variables",
         "Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers",
         "Section 2.1, 2.2",
         "Outcome: overtaken/disrupted/survived; Exposure: closure type (5 categories); Covariates: era, geographic barrier, external threat, institutional quality, tech position, regime duration, external patron"),

        ("8", "Data sources/measurement",
         "For each variable of interest, give sources of data and details of methods of assessment",
         "Section 2.1; Table S1",
         "Coding based on historical reference works; each polity's turning-point event documented in Table S1"),

        ("9", "Bias",
         "Describe any efforts to address potential sources of bias",
         "Section 3; Section 6",
         "Sensitivity analysis with 6 scenarios (3 reclassification × 2 disrupted assignment) to test coding robustness; Discussion section addresses selection bias and survivor bias"),

        ("10", "Study size",
         "Explain how the study size was arrived at",
         "Section 2.1",
         "N=96 polities from comprehensive survey of historical polities across six eras; not sample-based but census of major historical entities"),

        ("11", "Quantitative variables",
         "Explain how quantitative variables were handled in the analyses",
         "Section 3",
         "Binary outcome (overtaken vs. survived) constructed from 3-class outcome; two approaches for 'disrupted' category (as conquered / as survived); continuous covariates described"),

        ("12(a)", "Statistical methods",
         "Describe all statistical methods, including those used to control for confounding",
         "Section 3",
         "Fisher's exact test (one-sided); multivariate logistic regression with 7 covariates; bootstrap validation (5,000 iterations)"),

        ("12(b)", "Statistical methods",
         "Describe any methods used to examine subgroups and interactions",
         "Section 5",
         "Subgroup analysis by closure type (5 categories); cross-tabulation of conquest rates"),

        ("12(c)", "Statistical methods",
         "Explain how missing data were addressed",
         "Section 2.1",
         "Complete-case dataset; no missing values by design (all variables coded for all 96 polities)"),

        ("12(d)", "Statistical methods",
         "If applicable, describe analytical methods taking account of sampling strategy",
         "N/A",
         "Census of historical polities, not probability sample; acknowledged in Discussion as limitation"),

        ("12(e)", "Statistical methods",
         "Describe any sensitivity analyses",
         "Section 5",
         "Six-scenario sensitivity analysis: 3 reclassification levels (baseline, +5 strong, +7 all) × 2 disrupted-outcome assignments (as conquered, as survived)"),

        ("13", "Participants",
         "Report numbers of individuals at each stage of study",
         "Section 2.1; Section 4",
         "N=96 polities; breakdown by outcome (overtaken/disrupted/survived), by dominant strategy (stock/flow), by closure type"),

        ("14", "Descriptive data",
         "Give characteristics of study participants and information on exposures and potential confounders",
         "Section 2; Table 1; Table S1",
         "Table 1 lists 7 reclassification candidates with rationale; Table S1 provides full dataset with all variables for all 96 polities"),

        ("15", "Outcome data",
         "Report numbers of outcome events or summary measures",
         "Section 4; Table 2; Fig. 1",
         "Conquest rates by closure status across all 6 scenarios; Fisher p-values; relative risks"),

        ("16(a)", "Main results",
         "Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision",
         "Section 4; Section 5; Table 2; Table 3",
         "Unadjusted: Fisher p-values and conquest rate comparisons; Adjusted: multivariate logistic regression ORs with 95% CIs"),

        ("16(b)", "Main results",
         "Report category boundaries when continuous variables were categorized",
         "Section 2.1",
         "Era categories defined (ancient/medieval/early modern/modern/20th century/contemporary); closure types defined"),

        ("16(c)", "Main results",
         "If relevant, consider translating estimates of relative risk into absolute risk for a meaningful time period",
         "Section 4",
         "Absolute conquest rates reported for each group (e.g., 86.4% network closure vs. 60.8% no closure)"),

        ("17", "Other analyses",
         "Report other analyses done—e.g., analyses of subgroups and interactions, and sensitivity analyses",
         "Section 5; Fig. 3; Fig. 4",
         "Full sensitivity analysis across 6 scenarios; subgroup analysis by closure type; bootstrap OR validation"),

        ("18", "Key results",
         "Summarise key results with reference to study objectives",
         "Section 6.1–6.2 (Discussion)",
         "Technology flow disruption as mechanism (6.1); dose–response gradient; first contact and civilizational divergence (6.2); OR stability (6.4)"),

        ("19", "Limitations",
         "Discuss limitations of the study, taking into account sources of potential bias or imprecision",
         "Section 6.5",
         "Small-sample (N=96); retrospective coding subjectivity; non-independence of sequential polities; disrupted-category ambiguity; speculative modern extension"),

        ("20", "Interpretation",
         "Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, and results from similar studies",
         "Section 6.1–6.3; Section 7",
         "Technology flow disruption mechanism; extension to modern technological access exclusion (AI, semiconductors); cautious framing of forward-looking implications"),

        ("21", "Generalisability",
         "Discuss the generalisability (external validity) of the study results",
         "Section 6.3; Section 6.5",
         "Generalisation to modern technological exclusion discussed (6.3); limitations of extrapolation acknowledged (6.5); limited to major historical polities"),

        ("22", "Funding",
         "Give the source of funding and the role of the funders",
         "Statements and Declarations",
         "No external funding received"),
    ]

    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Item No.", "Section/Topic", "STROBE Recommendation",
               "Reported in", "Comment / How Addressed"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
        # Shade header
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elem)

    for item_no, section, recommendation, reported, comment in items:
        row_cells = table.add_row().cells
        vals = [item_no, section, recommendation, reported, comment]
        for i, val in enumerate(vals):
            row_cells[i].text = val
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    # Set column widths
    widths = [Cm(1.3), Cm(2.5), Cm(6.5), Cm(3.0), Cm(5.2)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    doc.add_paragraph()
    p = doc.add_paragraph(
        "Reference: von Elm E, Altman DG, Egger M, Pocock SJ, Gøtzsche PC, Vandenbroucke JP; "
        "STROBE Initiative. The Strengthening the Reporting of Observational Studies in Epidemiology "
        "(STROBE) statement: guidelines for reporting observational studies. "
    )
    run = p.add_run("Lancet")
    run.italic = True
    p.add_run(". 2007;370(9596):1453–1457.")
    for run in p.runs:
        run.font.size = Pt(8)

    path = os.path.join(OUT, "strobe_checklist.docx")
    doc.save(path)
    print(f"  STROBE checklist saved to {path}")


if __name__ == "__main__":
    print("Creating cover letter...")
    create_cover_letter()
    print("Creating STROBE checklist...")
    create_strobe_checklist()
    print("\nDone.")
