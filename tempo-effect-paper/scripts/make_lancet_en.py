"""Create Lancet Correspondence — English."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIG_DIR = '/home/ubuntu/figures'

def add_para(doc, text, bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_figure(doc, path, caption, width=6.0):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.italic = True
        cap.paragraph_format.space_after = Pt(12)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Title
add_para(doc,
    "The Forgotten Tempo Effect: How Delayed Childbearing Shapes "
    "Simultaneously Living Population Size Across OECD Countries",
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(doc, "Tatsuki Onishi", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "[Affiliation]", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc, "Correspondence", bold=True, size=10, space_after=6)
add_para(doc, "Word count: ~300 words (excluding references and figure legend)",
         italic=True, size=9, space_after=12)

body = (
"Population decline projections typically focus on total fertility rate (TFR), "
"yet the timing of births\u2014the tempo effect\u2014exerts an independent and "
"underappreciated influence on simultaneously living population size. "
"Bongaarts and Feeney formalised tempo-adjusted fertility in 1998,\u00b9 "
"and Goldstein, Lutz, and Scherbov demonstrated that delayed childbearing "
"reduces the number of generations alive at any moment for EU-15 countries.\u00b2 "
"However, this mechanism remains largely absent from contemporary policy "
"discourse on demographic decline.\n\n"
"We constructed an endogenous renewal model coupling age-specific fertility "
"(normal schedule centred on mean age at childbearing [MAC]) with Gompertz "
"survival, calibrated to life expectancy at birth. Using UN World Population "
"Prospects 2024 data,\u00b3 we validated the model across 38 OECD member states "
"plus China and the Democratic Republic of the Congo (DRC) over 1970\u20132023. "
"A dynamic variant updating TFR, life expectancy, and MAC every decade "
"achieved a median absolute percentage error of 4\u00b76% (mean 6\u00b77%) against "
"observed population trajectories (figure). France (0\u00b74%), Japan (1\u00b76%), "
"and Italy (1\u00b74%) showed excellent fit. Even the parsimonious static model "
"(fixed 2000 parameters) yielded median error of 4\u00b77% over 23 years.\n\n"
"The policy implication is direct: because MAC determines generational overlap, "
"a 5-year shift in mean childbearing age (e.g., from 25 to 30) reduces "
"simultaneously living population by approximately one-sixth, independent "
"of TFR.\u2074 This suggests that pronatalist policies addressing only quantum "
"(number of children) while ignoring tempo (timing) will systematically "
"overestimate their population impact. Conversely, policies that modestly "
"lower AFB\u2014through housing support, childcare infrastructure, or educational "
"reform\u2014could buffer population decline even without raising TFR.\n\n"
"Our 40-country validation demonstrates that this simple demographic identity "
"reproduces observed population dynamics with surprising accuracy. "
"The tempo effect deserves renewed attention as a policy lever for managing "
"the pace of demographic transition.\u2075"
)
add_para(doc, body, size=12, space_after=12)

add_para(doc, "References", bold=True, size=11, space_after=4)
refs = [
    "1. Bongaarts J, Feeney G. On the quantum and tempo of fertility. Popul Dev Rev 1998; 24: 271\u201391.",
    "2. Goldstein JR, Lutz W, Scherbov S. Long-term population decline in Europe: the relative importance of tempo effects and generational length. Popul Dev Rev 2003; 29: 699\u2013707.",
    "3. United Nations, Department of Economic and Social Affairs, Population Division. World Population Prospects 2024. https://population.un.org/wpp/",
    "4. Bongaarts J, Sobotka T. A demographic explanation for the recent rise in European fertility. Popul Dev Rev 2012; 38: 83\u2013120.",
    "5. Lutz W, Skirbekk V, Testa MR. The low-fertility trap hypothesis: forces that may lead to further postponement and fewer births in Europe. Vienna Yearb Popul Res 2006; 4: 167\u201392.",
]
for r in refs:
    add_para(doc, r, size=10, space_after=2)

doc.add_page_break()
add_para(doc, "Figure", bold=True, size=11, space_after=4)
add_figure(doc, f'{FIG_DIR}/fig1_showcase.png',
    "Figure. Endogenous renewal model with Gompertz survival versus UN WPP 2024 observed population, "
    "1970\u20132023. Dynamic model (blue dashed) updates parameters every 10 years; static model (red dotted) "
    "uses fixed 1970 parameters. Black line = UN WPP 2024 estimates. Six representative countries shown: "
    "Japan, China, United States, Republic of Korea, Germany, and DRC.",
    width=6.5)

doc.save('/home/ubuntu/Lancet_Correspondence_EN.docx')
print("OK: Lancet_Correspondence_EN.docx")
