# -*- coding: utf-8 -*-
"""Create cover letters for PDR Research Note (EN only, per user request)."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def make_cover_letter(path, paragraphs, font_name='Times New Roman', font_size=12, line_spacing=1.5):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = line_spacing

    for text, kwargs in paragraphs:
        p = doc.add_paragraph()
        bold = kwargs.get('bold', False)
        italic = kwargs.get('italic', False)
        size = kwargs.get('size', font_size)
        align = kwargs.get('align', None)
        space_after = kwargs.get('space_after', 6)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
    doc.save(path)
    print(f"OK: {path}")


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)

today = "April 14, 2026"
R = WD_ALIGN_PARAGRAPH.RIGHT

# =====================================================================
# PDR Research Note -- English Cover Letter
# =====================================================================
pdr_en = [
    (today, {'align': R, 'space_after': 18}),
    ("The Editors\nPopulation and Development Review\nPopulation Council\n"
     "One Dag Hammarskjold Plaza\nNew York, NY 10017, USA",
     {'space_after': 18}),
    ("Dear Editors,", {'space_after': 12}),
    ('Re: Submission of Notes and Commentary \u2014 "The Forgotten Tempo Effect: Delayed '
     'Childbearing, Simultaneously Living Population, and the Pace of Social Adaptation '
     'Across OECD Countries"',
     {'bold': True, 'space_after': 12}),
    ("We wish to submit the enclosed manuscript for consideration as a Notes and Commentary "
     "article in Population and Development Review. This paper revisits the tempo effect \u2014 "
     "the independent influence of birth timing on population size \u2014 which, despite seminal "
     "contributions by Bongaarts and Feeney (1998) and the foundational work by Goldstein, "
     "Lutz, and Scherbov (2003) published in this journal, has largely disappeared from "
     "contemporary policy discourse on demographic decline.",
     {'space_after': 8}),
    ("We believe PDR is the natural venue for this work for three reasons:\n\n"
     "1. Intellectual lineage: Goldstein, Lutz, and Scherbov\u2019s (2003) demonstration that "
     "generational length changes reduce population size was published in PDR. Our paper "
     "extends their EU-15 analysis to 40 countries with 20 additional years of data, using "
     "a complementary modelling approach.\n\n"
     "2. Empirical contribution: Using a parsimonious four-parameter endogenous renewal model "
     "validated against UN WPP 2024 data for 38 OECD member states plus China and the DRC "
     "over 1970\u20132023, we achieve a median absolute percentage error of 4.6% \u2014 demonstrating "
     "that quantum, tempo, and survival alone explain the vast majority of population dynamics "
     "across diverse demographic contexts.\n\n"
     "3. Policy reframing: We introduce the concept that tempo-sensitive policies control not "
     "merely population size but the speed at which societies must adapt their institutions "
     "to demographic change. This \u201cpace of adaptation\u201d framing \u2014 showing that a 5-year MAC "
     "increase reduces simultaneously living population by ~1/6 independent of TFR \u2014 offers "
     "a novel perspective for the PDR readership.",
     {'space_after': 8}),
    ("The manuscript is approximately 6,500 words with 7 figures, 4 tables, and 3 appendices "
     "(including a comparative table of national population projection methodologies for 15 "
     "countries/agencies, and a natural experiments analysis examining countries with major "
     "political/border changes as exogenous shocks). It follows GATHER reporting guidelines. "
     "The manuscript is formatted for double-anonymised review; author-identifying information "
     "appears only in this cover letter.",
     {'space_after': 8}),
    ("The work is original and is not under consideration for publication as a full article "
     "elsewhere. All authors have approved the manuscript and declare no competing interests.",
     {'space_after': 12}),
    ("Yours sincerely,", {'space_after': 18}),
    ("Tatsuki Onishi\n[Affiliation]\n[Email]\n[ORCID]",
     {'space_after': 6}),
]
make_cover_letter(os.path.join(OUT_DIR, 'CoverLetter_PDR_EN.docx'), pdr_en)

# =====================================================================
# Demographic Research -- English Cover Letter
# =====================================================================
demres_en = [
    (today, {'align': R, 'space_after': 18}),
    ("The Editors\nDemographic Research\nMax Planck Institute for Demographic Research\n"
     "Konrad-Zuse-Str. 1\n18057 Rostock, Germany",
     {'space_after': 18}),
    ("Dear Editors,", {'space_after': 12}),
    ('Re: Submission of Research Article \u2014 "The Forgotten Tempo Effect: Delayed '
     'Childbearing, Simultaneously Living Population, and the Pace of Social Adaptation '
     'Across OECD Countries"',
     {'bold': True, 'space_after': 12}),
    ("We wish to submit the enclosed manuscript for consideration as a Research Article "
     "in Demographic Research. This paper revisits the tempo effect \u2014 the independent "
     "influence of birth timing on population size \u2014 which, despite seminal contributions "
     "by Bongaarts and Feeney (1998) and foundational work by Goldstein, Lutz, and "
     "Scherbov (2003), has largely disappeared from contemporary policy discourse on "
     "demographic decline.",
     {'space_after': 8}),
    ("We believe Demographic Research is the ideal venue for this work for three reasons:\n\n"
     "1. Intellectual home: Demographic Research has been central to the development of "
     "tempo-adjustment methodology, publishing key contributions including Kohler and "
     "Ortega (2002) on tempo-adjusted parity progression measures and Frejka and Sobotka "
     "(2008) on diverse European fertility patterns. Our paper extends this tradition by "
     "providing the first systematic 40-country empirical validation of the tempo effect on "
     "simultaneously living population.\n\n"
     "2. Empirical contribution: Using a parsimonious four-parameter endogenous renewal model "
     "validated against UN WPP 2024 data for 38 OECD member states plus China and the DRC "
     "over 1970\u20132023, we achieve a median absolute percentage error of 4.6% \u2014 demonstrating "
     "that quantum, tempo, and survival alone explain the vast majority of population dynamics "
     "across diverse demographic contexts.\n\n"
     "3. Policy reframing: We introduce the concept that tempo-sensitive policies control not "
     "merely population size but the speed at which societies must adapt their institutions "
     "to demographic change. This \u2018pace of adaptation\u2019 framing \u2014 showing that a 5-year MAC "
     "increase reduces simultaneously living population by approximately one-sixth, independent "
     "of TFR \u2014 offers a novel perspective that aligns with the journal\u2019s commitment to "
     "publishing research with clear policy relevance.",
     {'space_after': 8}),
    ("The manuscript is approximately 6,500 words with 7 figures, 4 tables, and 3 appendices "
     "(including a comparative table of national population projection methodologies for 15 "
     "countries/agencies, and a natural experiments analysis examining countries with major "
     "political/border changes as exogenous shocks). It follows GATHER reporting guidelines "
     "and the Demographic Research author guidelines. All analytical code and data sources "
     "are documented for full reproducibility. The manuscript is formatted for double-anonymised "
     "review; author-identifying information appears only in this cover letter.",
     {'space_after': 8}),
    ("The work is original and is not under consideration for publication elsewhere. "
     "All authors have approved the manuscript and declare no competing interests.",
     {'space_after': 12}),
    ("Yours sincerely,", {'space_after': 18}),
    ("Tatsuki Onishi\n[Affiliation]\n[Email]\n[ORCID]",
     {'space_after': 6}),
]
make_cover_letter(os.path.join(OUT_DIR, 'CoverLetter_DemRes_EN.docx'), demres_en)

print("Cover letters created.")
