# -*- coding: utf-8 -*-
"""Create PDR Research Note - Japanese, with country projection methods appendix."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIG_DIR = '/home/ubuntu/figures'

def add_para(doc, text, bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_figure(doc, path, caption, width=6.0):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.italic = True
        cap.paragraph_format.space_after = Pt(12)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
style = doc.styles['Normal']
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# Title page
add_para(doc, "NOTES AND COMMENTARY", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc, "忘れられたテンポ効果：出産の遅延、同時在生人口、\nそしてOECD諸国における社会適応速度の制御",
         bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para(doc, "[匿名査読のため著者名を削除]", italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "語数：約5,500語", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

# Abstract
add_heading_styled(doc, "要旨", level=2)
add_para(doc,
    "人口予測と少子化対策は、出生率のカンタム（出生数）に圧倒的に焦点を当て、テンポ効果——"
    "出産タイミングが同時に生存する世代数に与える独立した影響——を見落としてきた。本稿では、"
    "年齢別出生率（平均出産年齢MACを中心とする正規分布）とGompertzパラメトリック生存関数"
    "（平均寿命にキャリブレーション）を結合した簡素な内生更新モデルにより、この過小評価された"
    "メカニズムを再考する。国連世界人口推計2024のデータを用い、OECD加盟38カ国＋中国＋コンゴ"
    "民主共和国の計40カ国で1970–2023年の検証を行った結果、動的モデル（10年ごとにパラメータ更新）"
    "は観測された人口軌跡に対し絶対百分率誤差の中央値4.6%を達成した。MACの5年上昇は、TFRとは"
    "独立に同時在生人口を約6分の1減少させることを示す。この知見は、出生数のみに対処する少子化"
    "対策がその人口学的影響を体系的に過大評価すること、そしてテンポに敏感な介入——住宅、保育、"
    "教育改革——が人口転換のペースと社会構造的適応の速度を管理する未活用のレバーとなりうる"
    "ことを意味する。", size=12, space_after=18)

add_para(doc, "キーワード：テンポ効果、同時在生人口、第一子出産年齢、平均出産年齢、"
         "Gompertz生存、OECD、人口転換、人口予測", italic=True, size=10, space_after=18)
doc.add_page_break()

# Section 1
add_heading_styled(doc, "1. 忘れられたテンポ効果", level=1)
add_para(doc,
    "人口減少をめぐるグローバルな議論は、単一の指標——合計特殊出生率（TFR）——に支配されている。"
    "TFRが置換水準（女性1人当たり約2.1人）を下回ると、警鐘が鳴る。しかしこのフレーミングは、"
    "ある瞬間に同時に生存する人数を決定する第二の独立した人口学的力——出産のタイミング——を"
    "体系的に見落としている。\n\n"
    "BongaartsとFeeney（1998）は、出生率のカンタム（出生数）とテンポ（出産タイミング）の区別を"
    "定式化し、女性が出産を先送りすると——たとえコーホート完結出生率が変わらなくても——"
    "期間TFRが機械的に押し下げられることを示した。Goldstein, Lutz, Scherbov（2003）はこの知見を"
    "さらに進め、EU15カ国において出産の遅延がある時点で同時に生存する世代数を減少させ、"
    "女性1人当たりの出生数とは独立に人口減少をもたらすことを実証した。彼らの要因分解は、"
    "世代長の変化が予測される人口減少のかなりの部分を説明することを示した。\n\n"
    "この基礎的研究にもかかわらず、テンポの次元は現代の政策議論からほぼ姿を消している。"
    "韓国の記録的な47兆ウォンの少子化対策から日本の歴代「少子化社会対策大綱」に至るまで、"
    "OECD諸国の最近の少子化対策パッケージを概観すると、出生数の増加にほぼ排他的に焦点を"
    "当てていることが分かる。同時在生人口（SLP）——ある暦上の時点で生存する人の総数——は"
    "カンタムとテンポの両方により形成されるが、体系的な政策的関心を受けるのはカンタムだけである。\n\n"
    "本稿は、単純だが実証的に根拠のある人口モデルを用いてテンポ効果を再考し、3つの問いに答える。"
    "（1）テンポ・カンタム・生存の三要素で構築された簡素なモデルは観測された人口軌跡を再現できるか？"
    "（2）多様な人口学的文脈において、テンポ効果はカンタムに対してどの程度大きいか？"
    "（3）社会が人口変動に合わせて制度を適応させるスピードに対する含意は何か？",
    size=12, space_after=12)

# Section 2
add_heading_styled(doc, "2. モデルとデータ", level=1)
add_heading_styled(doc, "2.1 内生更新モデル", level=2)
add_para(doc,
    "離散時間・単一性別の人口モデルを構築する。人口ベクトルP(t)は毎年以下のように更新される。\n\n"
    "(a) 生存：年齢xの個人はGompertzハザード関数 h(x) = a·exp(b·x) から導かれる生存確率で"
    "年齢x+1まで生存する。生存関数は S(x) = exp[−(a/b)(exp(bx)−1)]。パラメータaは出生時平均"
    "余命e₀が観測値と一致するようキャリブレーション。bは0.085に固定。\n\n"
    "(b) 出生：出生は内生的に生成。年齢別出生率（ASFR）は平均出産年齢（MAC）を中心とし"
    "標準偏差σの正規密度をTFRにスケーリング。時刻tの出生数は Σ P_x(t)·f·ASFR(x)。\n\n"
    "この最小限のパラメータ化には期間あたり4つの入力値のみが必要：TFR、平均寿命（e₀）、MAC、σ。"
    "モデルは移民を意図的に省略し、カンタム・テンポ・生存の純粋な人口学的メカニズムを分離する。",
    size=12, space_after=12)

add_heading_styled(doc, "2.2 データ", level=2)
add_para(doc,
    "全入力パラメータと検証データは国連世界人口推計2024（UN DESA 2024）から取得。"
    "分析対象は40カ国：OECD全加盟38カ国＋中国＋コンゴ民主共和国（DRC）。"
    "初期人口年齢構造（5歳階級）は1歳刻みに内挿。TFR、e₀、MACは1950–2023年の各暦年について抽出。\n\n"
    "GATHER報告ガイドライン（Stevens et al. 2016）に従い：入力データは国連人口部から公開入手可能。"
    "全モデルコードとパラメータは文書化済み。分析アプローチは完全に再現可能。",
    size=12, space_after=12)

add_heading_styled(doc, "2.3 モデルのバリアント", level=2)
add_para(doc,
    "2つのバリアントを実装する。\n\n"
    "静的モデル：パラメータ（TFR、e₀、MAC）を基準年の値に固定。"
    "4つの基準年（1970、1980、1990、2000）で前方投影し160の国×基準年の組み合わせを得る。\n\n"
    "動的モデル：10年ごとに観測値を用いてパラメータを更新し、全40カ国で1970年から2023年まで実行。",
    size=12, space_after=12)

# Section 3
add_heading_styled(doc, "3. OECD全体の検証結果", level=1)
add_heading_styled(doc, "3.1 全体的な適合度", level=2)
add_para(doc,
    "表1にモデル性能を要約する。動的モデルは53年間でMAPE中央値4.6%（平均6.7%）を達成し、"
    "最終人口比の平均は0.999（SD=0.189）——体系的バイアスが無視できることを示す。"
    "静的モデルの適合度は予測期間に伴い劣化する：中央値4.7%（基準年2000）から7.3%（基準年1970）。\n\n"
    "40カ国のうち30カ国で動的MAPEが10%未満、20カ国で5%未満、6カ国で2%未満。"
    "最適合国——フランス（0.4%）、コスタリカ（0.9%）、フィンランド（0.8%）、チェコ（1.3%）、"
    "スロベニア（1.4%）、イタリア（1.4%）——は緩やかな人口転換を経験し大規模な移民ショックがない国々。",
    size=12, space_after=6)

add_para(doc, "表1. 40カ国のモデル適合度の要約（モデルバリアント・基準年別）",
         bold=True, italic=True, size=10, space_after=4)
table = doc.add_table(rows=6, cols=6)
table.style = 'Light Shading Accent 1'
headers = ['モデル', '期間(年)', 'N', 'MAPE平均(%)', 'MAPE中央値(%)', '最終比率(mean±SD)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data_rows = [
    ['静的 (1970)', '50', '40', '12.4', '7.3', '1.272 ± 0.481'],
    ['静的 (1980)', '43', '40', '9.6', '7.7', '1.023 ± 0.288'],
    ['静的 (1990)', '33', '40', '7.8', '6.5', '0.953 ± 0.198'],
    ['静的 (2000)', '23', '40', '5.1', '4.7', '0.914 ± 0.101'],
    ['動的 (10年)', '53', '40', '6.7', '4.6', '0.999 ± 0.189'],
]
for i, rd in enumerate(data_rows):
    for j, val in enumerate(rd):
        table.rows[i+1].cells[j].text = val
add_para(doc, "", size=6, space_after=6)

add_heading_styled(doc, "3.2 不適合の原因", level=2)
add_para(doc,
    "MAPEが10%を超える国には共通の特徴がある。移民主導の人口増加はオーストラリア（13.5%）、"
    "カナダ（12.2%）、ルクセンブルク（21.5%）、イスラエル（13.9%）の不適合を説明する。"
    "本モデルはカンタム・テンポ・生存メカニズムを分離するため意図的に移民を除外しており、"
    "残差的不適合は移民成分を定量化する。\n\n"
    "急速な出生率転換がメキシコ（23.3%）、トルコ（17.0%）、中国（15.6%）、コロンビア（13.1%）を説明。"
    "韓国（11.9%）は極端な出生率低下と近年の移民の両効果を併せ持つ。"
    "リトアニア（12.0%）とラトビア（8.5%）はEU加盟後の移民流出を反映。",
    size=12, space_after=12)

# Figures
add_figure(doc, f'{FIG_DIR}/fig1_showcase.png',
    "図1. 代表6カ国のモデル vs 観測人口軌跡（1970–2023）。動的モデル（青破線）は10年ごとに"
    "パラメータ更新。静的モデル（赤点線）は1970年基準。黒実線＝UN WPP 2024。", width=6.0)
add_figure(doc, f'{FIG_DIR}/fig2_all_countries.png',
    "図2. 全40カ国のモデル検証。各パネル右上に動的モデルのMAPEを表示。", width=6.5)
add_figure(doc, f'{FIG_DIR}/fig3_heatmap.png',
    "図3. 静的モデルMAPE（%）：国×基準年。緑＝良好、赤＝不良。スケール上限30%。", width=5.0)
add_figure(doc, f'{FIG_DIR}/fig4_comparison.png',
    "図4. 静的 vs 動的モデル比較。左：国別MAPE。右：最終人口比率（2023年）。", width=6.0)
add_figure(doc, f'{FIG_DIR}/fig5_bias.png',
    "図5. モデルバイアス分析（基準年2000）。(A) vs TFR、(B) vs 平均寿命、(C) vs MAC。"
    "体系的関係は観察されず、モデル性能が人口学的文脈に対し頑健であることを示す。", width=6.0)

# Section 4
doc.add_page_break()
add_heading_styled(doc, "4. 政策レバーとしてのテンポ効果：社会適応速度の制御", level=1)
add_para(doc,
    "検証結果は、4つのパラメータだけのモデルが観測された人口軌跡を誤差中央値5%未満で再現"
    "できることを確認する。この簡素さにより、カンタム（TFR）、生存（e₀）、テンポ（MAC）の"
    "人口規模へのそれぞれの寄与が透明になる。\n\n"
    "テンポの経路は世代の重なりを通じて作用する。平均出産年齢が25歳の場合、おおよそ4世代"
    "（0, 25, 50, 75歳）が同時に生存する。MACが30歳に上昇すると世代間隔が広がり約3.3の"
    "重なり世代となり、同時在生人口が約6分の1減少する——各女性の出生数が同じであっても。\n\n"
    "このメカニズムには人口規模を超えて人口変動のペースに及ぶ決定的な政策的含意がある。"
    "TFR=1.5で同一だがMAC=25 vs MAC=33の2カ国を考える。MACが高い国は重なる世代が少ない"
    "ため暦年あたりの実効的な人口減少が速い。この加速は年金制度改革、医療インフラ拡充、"
    "労働市場再構築のために利用可能な時間を圧縮する。\n\n"
    "テンポに敏感な政策は単に何人が存在するかだけでなく、社会が人口変動に合わせてその社会"
    "構造を適応させなければならない速度を制御する。若い家庭向けの手頃な住宅、普遍的な保育、"
    "早期の出産を不利にしない教育課程の再構築を通じてAFBを緩やかに引き下げる政策は、"
    "TFRを引き上げなくとも人口減少のペースを緩め制度的調整のための時間を稼ぐことができる。\n\n"
    "問いは「どうすれば出生数を増やせるか」だけでなく「人口転換の速度をどう管理するか」でもある。"
    "テンポ効果は後者に対するメカニズムを提供し、それは現代の政策設計において見落とされてきた。\n\n"
    "40カ国の検証は、これが理論的な珍事ではなく、多様な人口学的文脈——転換後の日本"
    "（MAC=31.4、TFR=1.2）から転換前のDRC（MAC=24.8、TFR=6.1）まで——で作用する"
    "量的に有意な力であることを実証する。", size=12, space_after=12)

# Section 5
add_heading_styled(doc, "5. 限界", level=1)
add_para(doc,
    "いくつかの限界を認める。第一に、モデルは移民を除外しており、これはオーストラリア、カナダ、"
    "ルクセンブルク等の主要な不適合原因。除外は自然な人口学的メカニズムを分離するための"
    "意図的なものだが、大規模な純移民のある国への直接的な適用を制限する。第二に、正規分布の"
    "出生スケジュールは簡略化であり実際のASFRは歪みや二峰性を示す場合がある。第三に、10年ごとの"
    "パラメータ更新は10年内の急速な転換を見逃す可能性がある。第四に、Gompertz生存関数は成人"
    "死亡率を良好に適合するが乳児・小児死亡率の近似精度はやや低い。これらの限界にもかかわらず、"
    "モデルの簡素な構造は政策コミュニケーションにおける利点である。", size=12, space_after=12)

# Section 6
add_heading_styled(doc, "6. 結論", level=1)
add_para(doc,
    "同時在生人口に対するテンポ効果は人口学理論において確立されているが政策実践においては忘れ"
    "去られている。UN WPP 2024データに対し40カ国で検証した簡素なモデルを用いて、出産のタイミング"
    "が人口規模に量的に実質的な影響を及ぼすこと——それが出生率のカンタムとは独立に作用すること——"
    "を示した。政策的含意は人口規模を超えて人口変動のペースに及ぶ：テンポに敏感な介入は社会が"
    "制度を適応させなければならない速度を制御でき、従来の少子化対策を補完するレバーを提供する。"
    "テンポ効果を人口学的影響評価と人口政策設計に体系的に組み込むことを提言する。",
    size=12, space_after=18)

# References
add_heading_styled(doc, "参考文献", level=1)
refs = [
    'Bongaarts, J. and G. Feeney. 1998. "On the quantum and tempo of fertility," Population and Development Review 24(2): 271-291.',
    'Bongaarts, J. and T. Sobotka. 2012. "A demographic explanation for the recent rise in European fertility," Population and Development Review 38(1): 83-120.',
    'Goldstein, J.R., W. Lutz, and S. Scherbov. 2003. "Long-term population decline in Europe," Population and Development Review 29(4): 699-707.',
    'Lutz, W., V. Skirbekk, and M.R. Testa. 2006. "The low-fertility trap hypothesis," Vienna Yearbook of Population Research 4: 167-192.',
    'Stevens, G.A., L. Alkema, R.E. Black, et al. 2016. "The GATHER statement," The Lancet 388(10062): e19-e23.',
    'United Nations. 2024. World Population Prospects 2024. https://population.un.org/wpp/',
    'Gonand, F. 2005. "Assessing the Robustness of Demographic Projections in OECD Countries," OECD Working Papers No. 464.',
]
for r in refs:
    add_para(doc, r, size=11, space_after=4)

# Appendix A: GATHER
doc.add_page_break()
add_heading_styled(doc, "付録A：GATHER準拠声明", level=2)
add_para(doc,
    "本研究は人口推計を報告するものであり、GATHER（Stevens et al. 2016）に準拠する。\n\n"
    "・項目1-3（目的、方法、対象集団）：第1-2節に記述。\n"
    "・項目4-7（データ入力）：全入力データはUN WPP 2024から取得、公開入手可能。\n"
    "・項目8-10（データ調整）：初期人口年齢構造は5歳階級から1歳刻みに内挿。\n"
    "・項目11-13（モデリング）：Gompertz生存、正規出生スケジュール、内生更新を第2.1節に記述。\n"
    "・項目14-16（不確実性、結果）：MAPEと最終比率を適合指標として報告。モデルは決定論的。\n"
    "・項目17-18（解釈、再現性）：コードとデータソースを文書化。分析コードは要請に応じ提供。",
    size=11, space_after=12)

# Appendix B: Country Projection Methods
doc.add_page_break()
add_heading_styled(doc, "付録B：OECD諸国の公式人口予測手法と仮定の比較", level=2)
add_para(doc,
    "本付録は、分析対象40カ国の公式人口予測手法と主要仮定を要約する。全ての国がコーホート"
    "要因法（cohort-component method）の変種を基礎としているが、出生タイミングの扱い、死亡率"
    "改善モデル、移民仮定、シナリオ構造において大きく異なる。これらの差異は、本モデルの意図的な"
    "4パラメータへの簡素化を文脈化する。", size=11, space_after=12)

add_para(doc, "表B1. 国・機関別の公式人口予測手法の概要", bold=True, italic=True, size=10, space_after=4)

tbl = doc.add_table(rows=16, cols=5)
tbl.style = 'Light Shading Accent 1'
hdr = ['国・機関', '手法', '出生率仮定', '死亡率仮定', '移民の扱い']
for i, h in enumerate(hdr):
    tbl.rows[0].cells[i].text = h

rows_data = [
    ['国連WPP 2024\n（全対象国）', 'コーホート要因法\nベイズ確率的予測',
     'ベイズ階層モデル\nTFR軌跡＋不確実性', 'Lee-Carter変種\n国別ドリフト', '純移民を仮定\n長期平均に収束'],
    ['日本（社人研）', 'コーホート要因法\n出生3×死亡3バリアント',
     'コーホート出生モデル\n中位TFR=1.20(2070)\nMAC=32.8', 'Lee-Carterモデル\ne₀=85.9(M)/91.8(F)',
     '年齢・性別純移民\n約16.3万人/年'],
    ['米国（Census Bureau）', 'コーホート要因法\n主＋移民3バリアント',
     '人種別ASFR\nTFR→約1.75(2060)', '死因別モデル\ne₀≈83.9(2100)', '主変数：4シナリオ\n約110万人/年（主）'],
    ['ドイツ（Destatis）', 'コーホート要因法\n27バリアント(3×3×3)',
     'TFR 1.29-1.65\nMAC≈31.7-32.1', 'e₀ 82.6-86.4(M)\n85.9-89.3(F)', '純移民3水準\n15万/25万/35万人'],
    ['英国（ONS）', 'コーホート要因法\n主＋9バリアント',
     'ASFR; 主TFR≈1.59\n高/低バリアント', '死亡率改善モデル\ne₀≈83.9(M)/86.3(F)',
     '長期純移民≈31.5万\nバリアント：12.6-51.5万'],
    ['フランス（INSEE）', 'コーホート要因法\n中央＋成分別3バリアント',
     'TFR≈1.80中央\n高2.10/低1.60', 'トレンド外挿\ne₀≈87.5(M)/90.0(F)', '純移民+7万人/年'],
    ['韓国（KOSTAT）', 'コーホート要因法\n3シナリオ',
     'コーホートモデル\n中位TFR→1.08(2040)', 'Lee-Carter\ne₀=88.0(M)/91.4(F)',
     '国籍別純移民\n約6-10万人/年'],
    ['イタリア（ISTAT）', 'コーホート要因法\n中位＋4シナリオ',
     'TFR≈1.40中位\n範囲1.20-1.60', 'Lee-Carter\ne₀≈85.8(M)/89.2(F)', '純移民≈+15-23万人/年'],
    ['豪州（ABS）', 'コーホート要因法\n3系列(A/B/C)',
     'TFR 1.55-1.85\n系列B: 1.62', '死亡率改善率外挿\ne₀≈87(M)/89(F)',
     'NOM依存度高\n系列B:約23.5万/年'],
    ['カナダ（StatCan）', 'コーホート要因法\n＋マイクロシミュレーション',
     'TFR 1.40-1.60\n中位1.49', 'Lee-Carter変種\ne₀≈86(M)/89(F)',
     '純移民約40-50万/年\n主要成長ドライバー'],
    ['Eurostat\n（EU加盟国）', 'コーホート要因法\n収束モデル',
     '加盟国間TFR\n部分収束', '死亡率改善率\n収束モデル', '長期純移民に\n収束する国別経路'],
    ['中国（NBS）', 'コーホート要因法\n（定期公表なし）',
     'TFR=1.0-1.2(2022-23)\nUN WPPは回復仮定', 'モデル生命表\ne₀≈78.6',
     '低い国際移民\n国内移動は全国\n予測に反映なし'],
    ['DRC（国家予測なし）', 'UN WPPに依存\n独立予測なし',
     'TFR≈6.1(2023)\nUN:漸減仮定', 'モデル生命表\ne₀≈60.7', '低い純移民\n難民流は非体系的'],
    ['メキシコ（CONAPO）', 'コーホート要因法\n3バリアント',
     'TFR→約1.7(2050)', 'トレンド外挿\ne₀≈79(M)/83(F)', '純移出→ほぼゼロ\n≈5万人(2050)'],
    ['トルコ（TurkStat）', 'コーホート要因法\n3シナリオ',
     'TFR 1.51→約1.60\n長期', '改善モデル\ne₀≈80(M)/84(F)', '純移入≈20-30万/年\n難民含む'],
]
for i, rd in enumerate(rows_data):
    for j, val in enumerate(rd):
        tbl.rows[i+1].cells[j].text = val

add_para(doc, "", size=6, space_after=8)
add_para(doc,
    "出典：UN DESA(2024), 社人研(2023), US Census Bureau(2023), Destatis(2025), ONS(2025), "
    "INSEE(2021), KOSTAT(2023), ISTAT(2023), ABS(2018), Statistics Canada(2024), "
    "Eurostat(2024), CONAPO(2018), TurkStat(2023)。",
    italic=True, size=9, space_after=12)

add_heading_styled(doc, "B.1 共通点と主要な差異", level=3)
add_para(doc,
    "全ての国家予測システムがコーホート要因法を基礎構造として共有し、年齢別人口を出生・死亡・"
    "移民の仮定を用いて反復的に加齢させる。本モデルとの関連で重要な差異は以下の通り：\n\n"
    "・出生タイミングの扱い：多くの国家予測はMACとσでパラメータ化するのではなく完全なASFR"
    "スケジュールを指定する。日本（社人研）と韓国はタイミングシフトを明示的に追跡するコーホート"
    "出生モデルを使用。本モデルの正規分布簡略化は中心傾向を捕捉するがスケジュール形状は捉えない。\n\n"
    "・死亡率モデル：国家機関は通常Lee-Carterまたはその拡張を使用。本モデルのGompertz生存"
    "（パラメータaのみキャリブレーション）はより簡素だが年齢別死亡率パターンの柔軟性は低い。\n\n"
    "・移民：最も変動が大きい成分であり本モデルが意図的に除外するもの。移民依存国（豪州、"
    "カナダ、ルクセンブルク、イスラエル）では移民仮定が長期予測を支配する。\n\n"
    "・シナリオ構造：韓国の3バリアントからドイツの27バリアントまで幅広い。UN WPPはベイズ"
    "確率的予測で完全な不確実性分布を提供。本モデルの決定論的単一軌跡は不確実性定量化を"
    "テンポ・カンタム分解の透明性と引き換えにしている。\n\n"
    "・テンポの扱い：注目すべきことに、いずれの国家予測システムもカンタムとテンポ成分への"
    "明示的分解を行っていない。出生タイミングはASFRを通じて暗黙的に入るが、MACの同時在生人口"
    "への独立した寄与は分離されない。このギャップが本研究の動機である。",
    size=11, space_after=12)

add_heading_styled(doc, "B.2 モデル比較への含意", level=3)
add_para(doc,
    "本モデルは国家予測システムの代替ではなく、テンポ・カンタム・生存分解を明示することによる"
    "補完を目的とする。上表は、最も精緻な国家システムでさえ同じ基本構造（コーホート要因法）を"
    "共有し、パラメータ推定法とシナリオ構造において主に異なり、明示的テンポ分解を一様に欠くこと"
    "を示す。4パラメータモデルが同じ人口に対し動的MAPE中央値4.6%を達成する性能は、移民を含む"
    "完全パラメータ化国家モデルの精度には及ばないものの、テンポ経路の量的重要性を確立するには"
    "十分である。", size=11, space_after=12)

doc.save('/home/ubuntu/PDR_Research_Note_JP.docx')
print("OK: PDR_Research_Note_JP.docx")
