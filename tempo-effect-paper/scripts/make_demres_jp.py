# -*- coding: utf-8 -*-
"""Create Demographic Research submission — Japanese manuscript."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'figures')
OUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)


def add_para(doc, text, bold=False, italic=False, size=12, align=None,
             space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_figure(doc, path, title, note=None, width=6.0):
    """Insert figure with title above and optional note below."""
    t = doc.add_paragraph()
    r = t.add_run(title)
    r.font.size = Pt(11)
    r.bold = True
    t.paragraph_format.space_after = Pt(4)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        p.paragraph_format.space_after = Pt(2)
    if note:
        cap = doc.add_paragraph()
        r2 = cap.add_run(note)
        r2.font.size = Pt(9)
        r2.italic = True
        cap.paragraph_format.space_after = Pt(12)


# ==============================================================
# Document setup
# ==============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
style = doc.styles['Normal']
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# ==============================================================
# Title page
# ==============================================================
add_para(doc, "Research Article", bold=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc,
    "忘れられたテンポ効果：出産の遅延、同時在生人口、\n"
    "そしてOECD諸国における社会適応速度の制御",
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para(doc, "[匿名査読のため著者名を削除]",
         italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "語数（注釈・参考文献除く）：約6,500語",
         italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

# ==============================================================
# Structured Abstract (DR format)
# ==============================================================
add_heading_styled(doc, "要旨", level=2)

add_para(doc, "BACKGROUND（背景）", bold=True, size=12, space_after=2)
add_para(doc,
    "人口予測と少子化対策は出生率のカンタム（出生数）に圧倒的に焦点を当て、テンポ効果——"
    "出産タイミングが同時に生存する世代数に与える独立した影響——を見落としてきた。",
    size=12, space_after=6)

add_para(doc, "OBJECTIVE（目的）", bold=True, size=12, space_after=2)
add_para(doc,
    "3つの問いに答える：(1) テンポ・カンタム・生存の三要素で構築された簡素なモデルは観測された"
    "人口軌跡を再現できるか？ (2) 多様な人口学的文脈においてテンポ効果はカンタムに対してどの"
    "程度大きいか？ (3) 社会が制度を適応させるスピードに対する含意は何か？",
    size=12, space_after=6)

add_para(doc, "METHODS（方法）", bold=True, size=12, space_after=2)
add_para(doc,
    "年齢別出生率（平均出産年齢MACを中心とする正規分布）とGompertzパラメトリック生存関数"
    "（平均寿命にキャリブレーション）を結合した離散時間内生更新モデルを構築。国連世界人口推計"
    "2024データを用い、40カ国（OECD38カ国＋中国＋DRC）で1970–2023年の検証を実施。",
    size=12, space_after=6)

add_para(doc, "RESULTS（結果）", bold=True, size=12, space_after=2)
add_para(doc,
    "動的モデル（10年ごとにパラメータ更新）は絶対百分率誤差の中央値（MAPE）4.6%を達成。"
    "MACの5年上昇はTFRとは独立に同時在生人口を約6分の1減少させる。",
    size=12, space_after=6)

add_para(doc, "CONCLUSIONS（結論）", bold=True, size=12, space_after=2)
add_para(doc,
    "テンポに敏感な介入——住宅、保育、教育改革——は人口転換のペースと社会構造的適応の"
    "速度を管理する未活用のレバーとなりうる。",
    size=12, space_after=6)

add_para(doc, "CONTRIBUTION（貢献）", bold=True, size=12, space_after=2)
add_para(doc,
    "同時在生人口に対するテンポ効果の初の体系的40カ国実証検証を提供し、多様な人口学的文脈に"
    "おけるその量的有意性を実証するとともに、制度的適応の速度を制御する政策レバーとしての"
    "テンポの概念を導入する。",
    size=12, space_after=12)

add_para(doc,
    "キーワード：テンポ効果、同時在生人口、平均出産年齢、Gompertz生存、内生更新モデル、"
    "OECD、人口転換、人口予測、出産の先送り、社会適応",
    italic=True, size=10, space_after=18)
doc.add_page_break()

# ==============================================================
# 1. Introduction
# ==============================================================
add_heading_styled(doc, "1. 忘れられたテンポ効果", level=1)
add_para(doc,
    "人口減少をめぐるグローバルな議論は、単一の指標——合計特殊出生率（TFR）——に支配されている。"
    "TFRが置換水準（女性1人当たり約2.1人）を下回ると、警鐘が鳴る。しかしこのフレーミングは、"
    "ある瞬間に同時に生存する人数を決定する第二の独立した人口学的力——出産のタイミング——を"
    "体系的に見落としている。",
    size=12, space_after=12)

add_para(doc,
    "人口学的翻訳——バイタルイベントのタイミングの変化が基礎的なカンタムとは独立に期間指標を"
    "変動させるという概念——はRyder（1964）により導入された。BongaartsとFeeney（1998）は"
    "出生率のカンタム（出生数）とテンポ（出産タイミング）の区別を定式化し、女性が出産を先送り"
    "すると——たとえコーホート完結出生率が変わらなくても——期間TFRが機械的に押し下げられること"
    "を示した。彼らのテンポ調整済みTFRは多くの国で観測された期間出生率が出産の基礎的傾向より"
    "かなり低いことを明らかにし、この知見はKohlerとOrtega（2002）のテンポ調整済みパリティ"
    "進行測度によって確認・拡張された。Sobotka（2004）は、南欧・東欧に出現した超低出生率"
    "（TFR≦1.3）（Kohler, Billari, and Ortega 2002）の多くが、希望子ども数の恒久的減少では"
    "なく出産の先送りに帰因しうることを実証した。",
    size=12, space_after=12)

add_para(doc,
    "Goldstein, Lutz, Scherbov（2003）はこの知見をさらに進め、EU15カ国において出産の遅延が"
    "ある時点で同時に生存する世代数を減少させ、女性1人当たりの出生数とは独立に人口減少を"
    "もたらすことを実証した。彼らの要因分解は、世代長の変化が予測される人口減少のかなりの部分を"
    "説明することを示した。このメカニズムは同時在生人口（SLP）——ある暦上の時点で生存する人の"
    "総数——を通じて作用し、SLPはカンタムとテンポの両方により形成される。",
    size=12, space_after=12)

add_para(doc,
    "この基礎的研究にもかかわらず、テンポの次元は現代の政策議論からほぼ姿を消している。"
    "これらの出産先送り傾向は、遅延する家族形成、価値観の変化、持続的な置換水準以下の出生率を"
    "特徴とする広範な第二の人口転換の一部である（Lesthaeghe 2010）。韓国の記録的な47兆ウォンの"
    "少子化対策から日本の歴代「少子化社会対策大綱」に至るまで、OECD諸国の最近の少子化対策"
    "パッケージを概観すると、出生数の増加にほぼ排他的に焦点を当てていることが分かる（Thévenon "
    "2011）。カンタムだけが体系的な政策的関心を受け、人口規模の形成におけるテンポの独立した"
    "役割は無視されている。",
    size=12, space_after=12)

add_para(doc,
    "本稿は、単純だが実証的に根拠のある人口モデルを用いてテンポ効果を再考し、3つの問いに答える："
    "(1) テンポ・カンタム・生存の三要素で構築された簡素なモデルは観測された人口軌跡を再現できるか？"
    "(2) 多様な人口学的文脈において、テンポ効果はカンタムに対してどの程度大きいか？"
    "(3) 社会が人口変動に合わせて制度を適応させるスピードに対する含意は何か？",
    size=12, space_after=12)

# ==============================================================
# 2. Model and Data
# ==============================================================
add_heading_styled(doc, "2. モデルとデータ", level=1)

add_heading_styled(doc, "2.1 内生更新モデル", level=2)
add_para(doc,
    "離散時間・単一性別の人口モデルを構築する。人口ベクトルP(t) = [P\u2080(t), P\u2081(t), ..., "
    "P\u2081\u2080\u2080(t)]は毎年更新される。本モデルは、事実上すべての国家・国際人口予測の基礎を"
    "なすコーホート要因法のファミリーに属する（Preston, Heuveline, and Guillot 2001）。"
    "各タイムステップで：",
    size=12, space_after=8)

add_para(doc,
    "(a) 生存：年齢xの個人はGompertzハザード関数 h(x) = a\u00b7exp(b\u00b7x) から導かれる生存確率で"
    "年齢x+1まで生存する。Gompertz（1825）により最初に提案され、成人死亡率の最も広く使用される"
    "パラメトリックモデルである。生存関数は S(x) = exp[\u2212(a/b)(exp(bx)\u22121)]。パラメータaは"
    "出生時平均余命 e\u2080 = \u222b\u2080^\u221e S(x)dx が観測値と一致するようキャリブレーション。"
    "bは0.085に固定。",
    size=12, space_after=8)

add_para(doc,
    "(b) 出生：出生は内生的に生成。年齢別出生率（ASFR）は平均出産年齢（MAC）を中心とし"
    "標準偏差\u03c3の正規密度をTFRにスケーリング。時刻tの出生数は "
    "\u03a3_{x=15}^{49} P_x(t)\u00b7f\u00b7ASFR(x)（fは女性人口比率）。",
    size=12, space_after=8)

add_para(doc,
    "この最小限のパラメータ化には期間あたり4つの入力値のみが必要：TFR、平均寿命（e\u2080）、"
    "MAC、\u03c3。モデルは移民を意図的に省略し、カンタム・テンポ・生存の純粋な人口学的メカニズムを"
    "分離する。人口モメンタム——出生率が置換水準に低下した後も若い年齢構造により人口が増加し"
    "続ける傾向（Keyfitz 1971）——は年齢構造化ダイナミクスを通じて内生的に捕捉される。",
    size=12, space_after=12)

add_heading_styled(doc, "2.2 データ", level=2)
add_para(doc,
    "全入力パラメータと検証データは国連世界人口推計2024（United Nations 2024）から取得。"
    "分析対象は40カ国：OECD全加盟38カ国（2024年時点）＋中国＋コンゴ民主共和国（DRC）。"
    "人口転換の全段階を網羅するよう選定。初期人口年齢構造（5歳階級・両性）は1歳刻みに内挿。"
    "TFR、e\u2080、MACは1950–2023年の各暦年について抽出。",
    size=12, space_after=8)

add_para(doc,
    "GATHER報告ガイドライン（Stevens et al. 2016）に従い：入力データは国連人口部から"
    "公開入手可能。全モデルコードとパラメータは文書化済み。分析アプローチは完全に再現可能。",
    size=12, space_after=12)

add_heading_styled(doc, "2.3 モデルのバリアント", level=2)
add_para(doc,
    "2つのバリアントを実装する：",
    size=12, space_after=8)

add_para(doc,
    "静的モデル：パラメータ（TFR、e\u2080、MAC）を基準年の値に固定し予測期間中一定に保持。"
    "4つの基準年（1970、1980、1990、2000）で前方投影し、160の国\u00d7基準年の組み合わせを得る。",
    size=12, space_after=8)

add_para(doc,
    "動的モデル：10年ごとに観測されたUN WPP値を用いてパラメータを更新（例：1970年パラメータを"
    "1970–1979年に、1980年パラメータを1980–1989年に使用）し、全40カ国で1970年から2023年まで"
    "実行。このバリアントは定期的な再キャリブレーションが適合度を実質的に改善するかどうか、"
    "ひいてはモデルの構造的仮定が妥当かどうかを検証する。",
    size=12, space_after=12)

# ==============================================================
# 3. Results
# ==============================================================
add_heading_styled(doc, "3. OECD全体の検証結果", level=1)

add_heading_styled(doc, "3.1 全体的な適合度", level=2)
add_para(doc,
    "表1にモデル性能を要約する。動的モデルは53年間の期間でMAPE中央値4.6%（平均6.7%）を達成し、"
    "最終人口比の平均は0.999（SD=0.189）——体系的バイアスが無視できることを示す。静的モデルの"
    "適合度は予測期間に伴い劣化する：中央値4.7%（基準年2000、23年間）から7.3%（基準年1970、"
    "50年間）。急速な人口変動期にパラメータを固定した場合の予想通りの結果である。",
    size=12, space_after=6)

add_para(doc,
    "40カ国のうち30カ国で動的MAPEが10%未満、20カ国で5%未満、6カ国で2%未満。"
    "最適合国——フランス（0.4%）、コスタリカ（0.9%）、フィンランド（0.8%）、チェコ（1.3%）、"
    "スロベニア（1.4%）、イタリア（1.4%）——は比較的緩やかな人口転換を経験し大規模な移民"
    "ショックがない国々である。",
    size=12, space_after=6)

# Table 1
add_para(doc, "表1：40カ国のモデル適合度の要約（モデルバリアント・基準年別）",
         bold=True, size=11, space_after=4)

table = doc.add_table(rows=6, cols=6)
table.style = 'Light Shading Accent 1'
headers = ['モデル', '期間（年）', 'N', 'MAPE平均（%）',
           'MAPE中央値（%）', '最終比率（mean\u00b1SD）']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data_rows = [
    ['静的（1970）', '50', '40', '12.4', '7.3', '1.272 \u00b1 0.481'],
    ['静的（1980）', '43', '40', '9.6', '7.7', '1.023 \u00b1 0.288'],
    ['静的（1990）', '33', '40', '7.8', '6.5', '0.953 \u00b1 0.198'],
    ['静的（2000）', '23', '40', '5.1', '4.7', '0.914 \u00b1 0.101'],
    ['動的（10年）', '53', '40', '6.7', '4.6', '0.999 \u00b1 0.189'],
]
for i, row_data in enumerate(data_rows):
    for j, val in enumerate(row_data):
        table.rows[i + 1].cells[j].text = val
add_para(doc, "", size=6, space_after=6)

add_heading_styled(doc, "3.2 不適合の原因", level=2)
add_para(doc,
    "MAPEが10%を超える国には共通の特徴がある。移民主導の人口増加はオーストラリア（13.5%）、"
    "カナダ（12.2%）、スイス（7.2%）、ルクセンブルク（21.5%）、イスラエル（13.9%）の不適合を"
    "説明する——すべて純移民が自然増加を超えて人口を実質的に増加させた国である。本モデルは"
    "カンタム・テンポ・生存メカニズムを分離するため意図的に移民を除外しており、残差的不適合は"
    "移民成分を直接定量化する。",
    size=12, space_after=8)

add_para(doc,
    "急速な出生率転換が残りの外れ値を説明する：メキシコ（23.3%）、トルコ（17.0%）、"
    "中国（15.6%）、コロンビア（13.1%）は研究期間中にTFRが3～5低下。動的モデルは10年ごとの"
    "更新によりこれを部分的に捕捉するが、10年内の変化は未考慮のまま残る。",
    size=12, space_after=8)

add_para(doc,
    "韓国（11.9%）は両効果を併せ持つ：極端な出生率低下（TFR 4.5→0.7）と近年の移民。"
    "リトアニア（12.0%）とラトビア（8.5%）はEU加盟後の移民流出による人口減少を反映する。",
    size=12, space_after=12)

# Figures
add_figure(doc, os.path.join(FIG_DIR, 'fig1_showcase.png'),
    "図1：代表6カ国のモデル vs 観測人口軌跡（1970–2023）",
    note="注：動的モデル（青破線）は10年ごとにパラメータ更新。静的モデル（赤点線）は1970年"
    "基準年パラメータを使用。黒実線＝UN WPP 2024観測値。",
    width=6.0)

add_figure(doc, os.path.join(FIG_DIR, 'fig2_all_countries.png'),
    "図2：全40カ国のモデル検証",
    note="注：各パネル右上に動的モデルのMAPEを表示。国名アルファベット順。",
    width=6.5)

add_figure(doc, os.path.join(FIG_DIR, 'fig3_heatmap.png'),
    "図3：静的モデルMAPE（%）：国×基準年",
    note="注：緑＝良好な適合、赤＝不良な適合。スケール上限30%。",
    width=5.0)

add_figure(doc, os.path.join(FIG_DIR, 'fig4_comparison.png'),
    "図4：静的 vs 動的モデル比較",
    note="注：左パネル＝国別MAPE。右パネル＝最終人口比率（2023年のモデル/観測値）。"
    "動的モデルが一貫して静的バリアントを上回る。",
    width=6.0)

add_figure(doc, os.path.join(FIG_DIR, 'fig5_bias.png'),
    "図5：モデルバイアス分析（基準年2000）",
    note="注：(a) vs TFR、(b) vs 平均寿命、(c) vs MAC。体系的関係は観察されず、"
    "モデル性能が人口学的文脈に対し頑健であることを示す。",
    width=6.0)

# ==============================================================
# 4. Discussion
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "4. 政策レバーとしてのテンポ効果：社会適応速度の制御", level=1)
add_para(doc,
    "検証結果は、4つのパラメータだけ——TFR、平均寿命、MAC、出生スケジュールの幅——のモデルが"
    "観測された人口軌跡を誤差中央値5%未満で再現できることを確認する。この簡素さにより、"
    "カンタム（TFR）、生存（e\u2080）、テンポ（MAC）の人口規模へのそれぞれの寄与が透明になる。",
    size=12, space_after=12)

add_para(doc,
    "テンポの経路は世代の重なりを通じて作用する。平均出産年齢が25歳の場合、おおよそ4世代"
    "（0, 25, 50, 75歳）が同時に生存する。MACが30歳に上昇すると世代間隔が広がり約3.3の"
    "重なり世代（0, 30, 60, 90歳）となり、同時在生人口が約6分の1減少する——各女性の出生数が"
    "同じであっても。BongaartsとFeeney（2006）は元のカンタム・テンポの枠組みを拡張し、"
    "テンポの歪みが出生率指標だけでなくすべてのライフサイクルイベントに影響することを示し、"
    "タイミングの変化が期間指標の調整をはるかに超える人口レベルの帰結をもたらすことを強調した。",
    size=12, space_after=12)

add_para(doc,
    "このメカニズムには人口規模を超えて人口変動のペースに及ぶ決定的な政策的含意がある。"
    "TFR=1.5で同一だがMAC=25 vs MAC=33の2カ国を考える。MACが高い国は重なる世代が少ない"
    "ため暦年あたりの実効的な人口減少が速い。この加速は年金制度改革、医療インフラ拡充、"
    "労働市場再構築のために利用可能な時間を圧縮する。Lutz, Sanderson, Scherbov（2001）は"
    "今世紀中の世界人口増加の終焉を予測した。本分析はテンポがその終点にどれだけ速く到達するか"
    "を決定することを示す。",
    size=12, space_after=12)

add_para(doc,
    "この観点からすると、テンポに敏感な政策は単に何人が存在するかだけでなく、社会が人口変動に"
    "合わせてその社会構造を適応させなければならない速度を制御する。若い家庭向けの手頃な住宅、"
    "普遍的な保育、または早期の出産を不利にしない教育課程の再構築を通じて第一子出産年齢を緩やか"
    "に引き下げる政策は、TFRを引き上げなくとも人口減少のペースを緩め制度的調整のための時間を"
    "稼ぐことができる。各国比較のエビデンスは、家族政策が出産タイミングに影響しうることを示唆する"
    "が、完結出生率への効果はより穏やかである（Gauthier 2007; Thévenon 2011）。McDonald"
    "（2000）は制度における男女平等が持続的な出生率の前提条件であると論じた。本フレームワークは、"
    "そうした平等がキャリアのペナルティなしに早期の出産を可能にすることで、テンポ経路を通じても"
    "作用することを示唆する。",
    size=12, space_after=12)

add_para(doc,
    "この視点は政策の問題設定を再構成する。問いは「どうすれば出生数を増やせるか」だけでなく"
    "「人口転換の速度をどう管理するか」でもある。テンポ効果は後者に対するメカニズムを提供し、"
    "それは現代の政策設計において見落とされてきた。Lutz, Skirbekk, Testa（2006）は低出生率の"
    "罠——人口規範の低下と構造的制約が出産の先送りを強化する——について警告した。Bongaartsと"
    "Sobotka（2012）はその後、一部の欧州諸国が先送り傾向の逆転を開始しテンポの歪みが縮小"
    "するにつれて期間TFRが回復したことを示した。本モデルは、先送り傾向の部分的な逆転でさえ"
    "同時在生人口に実質的な効果をもたらしうることを示唆する。",
    size=12, space_after=12)

add_para(doc,
    "OECD諸国、中国、DRCにわたる40カ国の検証は、これが理論的な珍事ではなく、多様な人口学的"
    "文脈——転換後の日本（MAC=31.4、TFR=1.2）から転換前のDRC（MAC=24.8、TFR=6.1）まで——"
    "で作用する量的に有意な力であることを実証する。Myrskylä, Kohler, Billari（2009）は先進的"
    "発展が出生率低下を逆転しうることを示した。本フレームワークはこれを補完し、カンタムが低い"
    "ままであってもテンポの調整が人口軌跡を変えることを示す。移民パラメータなしで観測された人口"
    "軌跡を再現するモデルの能力は、同時在生人口の決定におけるカンタム・テンポ・生存の三要素の"
    "優位性をさらに強調する。",
    size=12, space_after=12)

# ==============================================================
# 5. Limitations
# ==============================================================
add_heading_styled(doc, "5. 限界", level=1)
add_para(doc,
    "いくつかの限界を認める。第一に、モデルは移民を除外しており、これはオーストラリア、カナダ、"
    "ルクセンブルク等の主要な不適合原因である。除外は自然な人口学的メカニズムを分離するための"
    "意図的なものだが、大規模な純移民のある国への直接的な適用を制限する。第二に、正規分布の"
    "出生スケジュールは簡略化であり、実際のASFRは歪みや二峰性を示す場合がある（Frejkaと"
    "Sobotka 2008は欧州各国間の出生スケジュール形状のかなりの差異を記録している）。第三に、"
    "10年ごとのパラメータ更新は10年内の急速な転換を見逃す可能性がある（例：2010年代の韓国の"
    "出生率崩壊）。第四に、Gompertz生存関数は成人死亡率を良好に適合するが乳児・小児死亡率の"
    "近似精度はやや低い。国家予測は通常Lee–Carter（Lee and Carter 1992）のようなより柔軟な"
    "死亡率モデルを使用する。これらの限界にもかかわらず、モデルの簡素な構造は政策コミュニケー"
    "ションにおける利点である：テンポ・カンタム・生存の分解を透明で解釈可能にする。",
    size=12, space_after=12)

# ==============================================================
# 6. Conclusion
# ==============================================================
add_heading_styled(doc, "6. 結論", level=1)
add_para(doc,
    "同時在生人口に対するテンポ効果は人口学理論において確立されているが政策実践においては忘れ"
    "去られている。UN WPP 2024データに対し40カ国で検証した簡素なモデルを用いて、出産のタイミング"
    "が人口規模に量的に実質的な影響を及ぼすこと——それが出生率のカンタムとは独立に作用する"
    "こと——を示した。政策的含意は人口規模を超えて人口変動のペースに及ぶ：テンポに敏感な介入は"
    "社会が制度を適応させなければならない速度を制御でき、従来の少子化対策を補完するレバーを"
    "提供する。テンポ効果を人口学的影響評価と人口政策設計に体系的に組み込むことを提言する。",
    size=12, space_after=18)

# ==============================================================
# Data availability statement
# ==============================================================
add_heading_styled(doc, "データ利用可能性声明", level=2)
add_para(doc,
    "全入力データは国連世界人口推計2024から取得しており、https://population.un.org/wpp/ で"
    "公開されている。全結果と図を再現するための分析コードは著者に問い合わせることで入手可能"
    "であり、受理後に公開リポジトリに登録する。",
    size=12, space_after=18)

# ==============================================================
# References (alphabetical, author-date style per DR house style)
# ==============================================================
add_heading_styled(doc, "参考文献", level=1)

refs = [
    'Bongaarts, J. and Feeney, G. (1998). On the quantum and tempo of fertility. '
    'Population and Development Review 24(2): 271\u2013291. '
    'doi:10.2307/2807974.',

    'Bongaarts, J. and Feeney, G. (2006). The quantum and tempo of life-cycle events. '
    'Vienna Yearbook of Population Research 4: 115\u2013151. '
    'doi:10.1553/populationyearbook2006s115.',

    'Bongaarts, J. and Sobotka, T. (2012). A demographic explanation for the recent rise '
    'in European fertility. Population and Development Review 38(1): 83\u2013120. '
    'doi:10.1111/j.1728-4457.2012.00473.x.',

    'Frejka, T. and Sobotka, T. (2008). Fertility in Europe: Diverse, delayed and below '
    'replacement. Demographic Research 19(3): 15\u201346. '
    'doi:10.4054/DemRes.2008.19.3.',

    'Gauthier, A.H. (2007). The impact of family policies on fertility in industrialized '
    'countries: A review of the literature. Population Research and Policy Review '
    '26(3): 323\u2013346. doi:10.1007/s11113-007-9033-x.',

    'Goldstein, J.R. and Kreyenfeld, M. (2011). Has East Germany overtaken West Germany? '
    'Recent trends in order-specific fertility. Population and Development Review '
    '37(3): 453\u2013472. doi:10.1111/j.1728-4457.2011.00430.x.',

    'Goldstein, J.R., Lutz, W., and Scherbov, S. (2003). Long-term population decline in '
    'Europe: The relative importance of tempo effects and generational length. Population '
    'and Development Review 29(4): 699\u2013707. doi:10.1111/j.1728-4457.2003.00699.x.',

    'Gompertz, B. (1825). On the nature of the function expressive of the law of human '
    'mortality, and on a new mode of determining the value of life contingencies. '
    'Philosophical Transactions of the Royal Society of London 115: 513\u2013583. '
    'doi:10.1098/rstl.1825.0026.',

    'Gonand, F. (2005). Assessing the robustness of demographic projections in OECD '
    'countries. OECD Economics Department Working Papers No. 464. Paris: OECD Publishing. '
    'doi:10.1787/010523837414.',

    'Keyfitz, N. (1971). On the momentum of population growth. Demography 8(1): 71\u201380. '
    'doi:10.2307/2060339.',

    'Kohler, H.-P., Billari, F.C., and Ortega, J.A. (2002). The emergence of lowest-low '
    'fertility in Europe during the 1990s. Population and Development Review 28(4): '
    '641\u2013680. doi:10.1111/j.1728-4457.2002.00641.x.',

    'Kohler, H.-P. and Ortega, J.A. (2002). Tempo-adjusted period parity progression '
    'measures, fertility postponement and completed cohort fertility. Demographic Research '
    '6(6): 91\u2013144. doi:10.4054/DemRes.2002.6.6.',

    'Lee, R.D. and Carter, L.R. (1992). Modeling and forecasting U.S. mortality. Journal '
    'of the American Statistical Association 87(419): 659\u2013671. '
    'doi:10.1080/01621459.1992.10475265.',

    'Lesthaeghe, R. (2010). The unfolding story of the Second Demographic Transition. '
    'Population and Development Review 36(2): 211\u2013251. '
    'doi:10.1111/j.1728-4457.2010.00328.x.',

    'Lutz, W., Sanderson, W., and Scherbov, S. (2001). The end of world population growth. '
    'Nature 412: 543\u2013545. doi:10.1038/35087589.',

    'Lutz, W., Skirbekk, V., and Testa, M.R. (2006). The low-fertility trap hypothesis: '
    'Forces that may lead to further postponement and fewer births in Europe. Vienna '
    'Yearbook of Population Research 4: 167\u2013192. '
    'doi:10.1553/populationyearbook2006s167.',

    'McDonald, P. (2000). Gender equity in theories of fertility transition. Population '
    'and Development Review 26(3): 427\u2013439. doi:10.1111/j.1728-4457.2000.00427.x.',

    'Myrskyl\u00e4, M., Kohler, H.-P., and Billari, F.C. (2009). Advances in development '
    'reverse fertility declines. Nature 460: 741\u2013743. doi:10.1038/nature08230.',

    'Preston, S.H., Heuveline, P., and Guillot, M. (2001). Demography: Measuring and '
    'modeling population processes. Oxford: Blackwell.',

    'Ryder, N.B. (1964). The process of demographic translation. Demography 1(1): 74\u201382. '
    'doi:10.2307/2060033.',

    'Sobotka, T. (2004). Is lowest-low fertility in Europe explained by the postponement '
    'of childbearing? Population and Development Review 30(2): 195\u2013220. '
    'doi:10.1111/j.1728-4457.2004.010_1.x.',

    'Stevens, G.A., Alkema, L., Black, R.E., Boerma, J.T., Collins, G.S., Ezzati, M., '
    'Grove, J.T., Hogan, D.R., Hosseinpoor, A.R., Lawn, J.E., Marusi\u0107, A., Mathers, '
    'C.D., Murray, C.J.L., Stanton, C., Tell, G.S., and Wardlaw, T.M. (2016). Guidelines '
    'for Accurate and Transparent Health Estimates Reporting: The GATHER statement. The '
    'Lancet 388(10062): e19\u2013e23. doi:10.1016/S0140-6736(16)30388-9.',

    'Th\u00e9venon, O. (2011). Family policies in OECD countries: A comparative analysis. '
    'Population and Development Review 37(1): 57\u201387. '
    'doi:10.1111/j.1728-4457.2011.00390.x.',

    'United Nations, Department of Economic and Social Affairs, Population Division (2024). '
    'World Population Prospects 2024. New York: United Nations. '
    'https://population.un.org/wpp/.',

    'Witte, J.C. and Wagner, G.G. (1995). Declining fertility in East Germany after '
    'unification: A demographic response to socioeconomic change. Population and '
    'Development Review 21(2): 387\u2013397. doi:10.2307/2137500.',
]
for r in refs:
    add_para(doc, r, size=11, space_after=6)

# ==============================================================
# Appendix A: GATHER
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "付録A：GATHER準拠声明", level=2)
add_para(doc,
    "本研究は人口推計を報告するものであり、GATHER（Stevens et al. 2016）に準拠する。主要項目：",
    size=11, space_after=6)
add_para(doc,
    "\u2022 項目1\u20133（目的、方法、対象集団）：第1\u20132節に記述。\n"
    "\u2022 項目4\u20137（データ入力）：全入力データはUN WPP 2024から取得、population.un.org/wppで"
    "公開入手可能。一次データ収集なし。\n"
    "\u2022 項目8\u201310（データ調整）：初期人口年齢構造は5歳階級から1歳刻みに均等分布で内挿。\n"
    "\u2022 項目11\u201313（モデリング）：Gompertz生存、正規出生スケジュール、内生更新を第2.1節に記述。"
    "期間あたり4パラメータ（TFR、e\u2080、MAC、\u03c3）。\n"
    "\u2022 項目14\u201316（不確実性、結果）：MAPEと最終比率を適合指標として報告。正式な不確実性区間なし。"
    "モデルは決定論的。\n"
    "\u2022 項目17\u201318（解釈、再現性）：コードとデータソースを文書化。分析コードは要請に応じ提供。",
    size=11, space_after=12)

# ==============================================================
# Appendix B: National Projection Methods
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "付録B：OECD諸国の公式人口予測手法と仮定の比較", level=2)

add_para(doc,
    "本付録は、分析対象のOECD諸国および追加2カ国（中国、DRC）の公式人口予測手法と主要仮定を"
    "要約する。全ての国がコーホート要因法（Preston, Heuveline, and Guillot 2001; Gonand 2005）の"
    "変種を基礎としているが、出生タイミングの扱い、死亡率改善モデル、移民仮定、シナリオ構造に"
    "おいて大きく異なる。これらの差異は、本モデルの意図的な4パラメータへの簡素化を文脈化する。",
    size=11, space_after=12)

# Table B1
add_para(doc, "表B1：国・機関別の公式人口予測手法の概要",
         bold=True, size=11, space_after=4)

tbl = doc.add_table(rows=16, cols=5)
tbl.style = 'Light Shading Accent 1'
hdr = ['国・機関', '手法', '出生率仮定', '死亡率仮定', '移民の扱い']
for i, h in enumerate(hdr):
    tbl.rows[0].cells[i].text = h

rows_data = [
    ['国連WPP 2024\n（全対象国）', 'コーホート要因法\nベイズ確率的予測',
     'ベイズ階層モデル\nTFR軌跡＋不確実性', 'Lee\u2013Carter変種\n国別ドリフト',
     '純移民を仮定\n長期平均に収束'],
    ['日本（社人研）', 'コーホート要因法\n出生3\u00d7死亡3バリアント',
     'コーホート出生モデル\n中位TFR=1.20(2070)\nMAC=32.8',
     'Lee\u2013Carterモデル\ne\u2080=85.9(M)/91.8(F)', '年齢・性別純移民\n約16.3万人/年'],
    ['米国（Census Bureau）', 'コーホート要因法\n主＋移民3バリアント',
     '人種別ASFR\nTFR\u21921.75(2060)', '死因別モデル\ne\u2080\u224883.9(2100)',
     '主変数：4シナリオ\n約110万人/年（主）'],
    ['ドイツ（Destatis）', 'コーホート要因法\n27バリアント(3\u00d73\u00d73)',
     'TFR 1.29\u20131.65\nMAC\u224831.7\u201332.1', 'e\u2080 82.6\u201386.4(M)\n85.9\u201389.3(F)',
     '純移民3水準\n15万/25万/35万人'],
    ['英国（ONS）', 'コーホート要因法\n主＋9バリアント',
     'ASFR; 主TFR\u22481.59\n高/低バリアント', '死亡率改善モデル\ne\u2080\u224883.9(M)/86.3(F)',
     '長期純移民\u2248315千人\nバリアント：126k\u2013515k'],
    ['フランス（INSEE）', 'コーホート要因法\n中央＋成分別3バリアント',
     'TFR\u22481.80中央\n高2.10/低1.60', 'トレンド外挿\ne\u2080\u224887.5(M)/90.0(F)',
     '純移民+7万人/年'],
    ['韓国（KOSTAT）', 'コーホート要因法\n3シナリオ',
     'コーホートモデル\n中位TFR\u21921.08(2040)', 'Lee\u2013Carter\ne\u2080=88.0(M)/91.4(F)',
     '国籍別純移民\n約6\u201310万人/年'],
    ['イタリア（ISTAT）', 'コーホート要因法\n中位＋4シナリオ',
     'TFR\u22481.40中位\n範囲1.20\u20131.60', 'Lee\u2013Carter\ne\u2080\u224885.8(M)/89.2(F)',
     '純移民\u2248+15\u201323万人/年'],
    ['豪州（ABS）', 'コーホート要因法\n3系列(A/B/C)',
     'TFR 1.55\u20131.85\n系列B: 1.62', '死亡率改善率外挿\ne\u2080\u224887(M)/89(F)',
     'NOM依存度高\n系列B:約23.5万/年'],
    ['カナダ（StatCan）', 'コーホート要因法\n＋マイクロシミュレーション',
     'TFR 1.40\u20131.60\n中位1.49', 'Lee\u2013Carter変種\ne\u2080\u224886(M)/89(F)',
     '純移民約40\u201350万/年\n主要成長ドライバー'],
    ['Eurostat\n（EU加盟国）', 'コーホート要因法\n収束モデル',
     '加盟国間TFR\n部分収束', '死亡率改善率\n収束モデル',
     '長期純移民に\n収束する国別経路'],
    ['中国（NBS）', 'コーホート要因法\n（定期公表なし）',
     'TFR=1.0\u20131.2(2022\u201323)\nUN WPPは回復仮定', 'モデル生命表\ne\u2080\u224878.6',
     '低い国際移民\n国内移動は全国\n予測に反映なし'],
    ['DRC（国家予測なし）', 'UN WPPに依存\n独立予測なし',
     'TFR\u22486.1(2023)\nUN:漸減仮定', 'モデル生命表\ne\u2080\u224860.7',
     '低い純移民\n難民流は非体系的'],
    ['メキシコ（CONAPO）', 'コーホート要因法\n3バリアント',
     'TFR\u2192約1.7(2050)', 'トレンド外挿\ne\u2080\u224879(M)/83(F)',
     '純移出\u2192ほぼゼロ\n\u22485万人(2050)'],
    ['トルコ（TurkStat）', 'コーホート要因法\n3シナリオ',
     'TFR 1.51\u2192約1.60\n長期', '改善モデル\ne\u2080\u224880(M)/84(F)',
     '純移入\u224820\u201330万/年\n難民含む'],
]
for i, rd in enumerate(rows_data):
    for j, val in enumerate(rd):
        tbl.rows[i + 1].cells[j].text = val

add_para(doc, "", size=6, space_after=8)
add_para(doc,
    "出典：UN DESA(2024), 社人研(2023), US Census Bureau(2023), Destatis(2025), ONS(2025), "
    "INSEE(2021), KOSTAT(2023), ISTAT(2023), ABS(2018), Statistics Canada(2024), "
    "Eurostat(2024), CONAPO(2018), TurkStat(2023)。",
    italic=True, size=9, space_after=12)

add_heading_styled(doc, "B.1 共通点と主要な差異", level=3)
add_para(doc,
    "全ての国家予測システムがコーホート要因法を基礎構造として共有し、年齢別人口を出生・死亡・"
    "移民の仮定を用いて反復的に加齢させる。本モデルとの関連で重要な差異は以下の通り：",
    size=11, space_after=8)
add_para(doc,
    "\u2022 出生タイミングの扱い：多くの国家予測はMACと\u03c3でパラメータ化するのではなく完全なASFR"
    "スケジュールを指定する。日本（社人研）と韓国はタイミングシフトを明示的に追跡するコーホート"
    "出生モデルを使用。本モデルの正規分布簡略化は中心傾向を捕捉するがスケジュール形状は捉えない。\n\n"
    "\u2022 死亡率モデル：国家機関は通常Lee\u2013Carter（Lee and Carter 1992）またはその拡張を使用。"
    "本モデルのGompertz生存（パラメータaのみキャリブレーション）はより簡素だが年齢別死亡率"
    "パターンの柔軟性は低い。\n\n"
    "\u2022 移民：最も変動が大きい成分であり本モデルが意図的に除外するもの。移民依存国（豪州、"
    "カナダ、ルクセンブルク、イスラエル）では移民仮定が長期予測を支配する。本モデルの不適合"
    "（MAPE 12\u201322%）は省略された移民成分を直接反映する。\n\n"
    "\u2022 シナリオ構造：韓国の3バリアントからドイツの27バリアントまで幅広い。UN WPPはベイズ"
    "確率的予測で完全な不確実性分布を提供。本モデルの決定論的単一軌跡は不確実性定量化を"
    "テンポ・カンタム分解の透明性と引き換えにしている。\n\n"
    "\u2022 テンポの扱い：注目すべきことに、いずれの国家予測システムもカンタムとテンポ成分への"
    "明示的分解を行っていない。出生タイミングはASFRを通じて暗黙的に入るが、MACの同時在生人口"
    "への独立した寄与は分離されない。このギャップが本研究の動機である。",
    size=11, space_after=12)

add_heading_styled(doc, "B.2 モデル比較への含意", level=3)
add_para(doc,
    "本モデルは国家予測システムの代替ではなく、テンポ・カンタム・生存分解を明示することによる"
    "補完を目的とする。上表は、最も精緻な国家システムでさえ同じ基本構造（コーホート要因法）を"
    "共有し、パラメータ推定法とシナリオ構造において主に異なり、明示的テンポ分解を一様に欠くこと"
    "を示す。4パラメータモデルが同じ人口に対し動的MAPE中央値4.6%を達成する性能は、移民を含む"
    "完全パラメータ化国家モデルの精度には及ばないものの、テンポ経路の量的重要性を確立するには"
    "十分である。",
    size=11, space_after=12)

# ==============================================================
# Appendix C: Natural Experiments
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "付録C：自然実験 — 政治的・国境変更による外生的ショック", level=2)

add_para(doc,
    "本モデルは移民を意図的に除外している。この設計選択はカンタム・テンポ・生存の純粋なメカニズム"
    "の分離を可能にするが、重要な問いを提起する：政治的事象の結果として大規模な人口再配分が"
    "生じた場合、モデルはどのように機能するのか？1970年から2023年の間に大規模な国境変更や"
    "国家解体を経験した国々は、外生的な移民ショックが人口に事実上課された自然実験を提供する。"
    "本付録では5つのケースを分析し、内生更新フレームワークの頑健性と限界を評価する。",
    size=11, space_after=12)

# --- C.1 Germany ---
add_heading_styled(doc, "C.1 ドイツ：移民ショックとしての再統一（1990年）", level=3)
add_para(doc,
    "1990年10月3日のドイツ再統一は、41年間にわたり著しく異なる人口レジームの下で発展した"
    "2つの人口を統合した。東ドイツ（ドイツ民主共和国）は1990年時点で西側より低い平均寿命"
    "（約74.5年 vs 約76.0年）、より若い出産年齢（MAC\u224825.1 vs 28.3）、そしてより高いが"
    "低下中のTFR（1.52 vs 1.45）を有していた。再統一の直後には大規模な東から西への移住"
    "（1989\u20131992年に約190万人）と東部での劇的な出生率崩壊（TFRは1994年に0.77まで低下；"
    "Witte and Wagner 1995）が生じた。GoldsteinとKreyenfeld（2011）はその後、出産の遅延化を"
    "含む東ドイツの出生行動の西側パターンへの漸進的収束を記録した。",
    size=11, space_after=6)

add_para(doc,
    "東西ドイツをそれぞれ固有の人口パラメータでモデル化し、両者を1970年から前方に実行し、"
    "2つのモデル人口を合計して合成軌跡を構築する。この合成軌跡は反事実を表す：2つの地域が"
    "人口学的に分離されたままであった場合（すなわち国境を越える移民がない場合）、ドイツの"
    "人口はどのようになっていたか？これを観測された統一ドイツと比較することで、再統一に"
    "関連する移民と統合効果の累積的影響が明らかになる。",
    size=11, space_after=6)

add_para(doc, "表C1：ドイツ再統一：合成東西ドイツ vs 観測された統一軌跡",
         bold=True, size=11, space_after=4)

tbl_c1 = doc.add_table(rows=6, cols=4)
tbl_c1.style = 'Light Shading Accent 1'
for i, h in enumerate(['年', '合成E+W（百万人）', '観測値（百万人）', '乖離率（%）']):
    tbl_c1.rows[0].cells[i].text = h
c1_data = [
    ['1990', '76.5', '79.4', '\u22123.6'],
    ['2000', '75.3', '82.2', '\u22128.4'],
    ['2010', '73.7', '81.8', '\u22129.9'],
    ['2020', '71.0', '83.2', '\u221214.6'],
    ['2023', '70.1', '83.3', '\u221215.8'],
]
for i, rd in enumerate(c1_data):
    for j, val in enumerate(rd):
        tbl_c1.rows[i + 1].cells[j].text = val
add_para(doc, "", size=6, space_after=6)

add_para(doc,
    "合成軌跡は再統一時点（1990年）で観測人口を3.6%下回り、2023年までに15.8%に拡大する。"
    "この拡大するギャップは閉鎖モデルにない3つの複合プロセスを反映する：(a)統一ドイツへの"
    "年間約30万\u201340万人の純移民、(b)地域人口動態を変容させた東から西への国内移動、"
    "(c)本モデルの分離レジームが1990年以降捕捉しない東ドイツの出生率・死亡率の西側水準への"
    "収束。合成東西モデルの全体MAPE 6.4%は、主要分析（表1）におけるドイツの比較的低い"
    "適合度がモデルの構造的失敗ではなく、主として再統一の移民効果に帰因することを確認する。",
    size=11, space_after=6)

add_figure(doc, os.path.join(FIG_DIR, 'fig_germany_reunification.png'),
    "図A-1：ドイツ再統一分析",
    note="注：(a)東西ドイツの個別モデル人口と合成軌跡 vs 観測された統一ドイツ。"
    "(b)合成・統一モデルの観測軌跡からの乖離率。1990年以降の拡大するギャップは再統一の"
    "累積的な移民・統合効果を定量化する。",
    width=6.5)

# --- C.2 Czechoslovakia ---
add_heading_styled(doc, "C.2 チェコスロバキア：ビロード離婚（1993年）", level=3)
add_para(doc,
    "1993年1月1日のチェコスロバキアの平和的解体はチェコとスロバキアの2つの独立国家を生んだ。"
    "国境を越える移民は比較的限定的であり、最小限の人口再配分を伴う国家分割という比較的"
    "クリーンな自然実験を構成する。本モデルは1970\u20132023年の全期間でチェコのMAPE 6.3%、"
    "スロバキアのMAPE 9.9%を達成する。スロバキアのより大きな誤差はEU加盟（2004年）後の"
    "チェコおよび西欧への移民流出を反映しており、閉鎖モデルはこれを捕捉できない。",
    size=11, space_after=12)

# --- C.3 Yugoslavia ---
add_heading_styled(doc, "C.3 ユーゴスラビア：解体と紛争（1991\u20132001年）", level=3)
add_para(doc,
    "ユーゴスラビアの解体は武力紛争、民族浄化、大規模な難民流出を伴い、全ての後継国家に"
    "影響を及ぼした。これは本サンプルにおける最も極端な移民ショックを表す。モデル性能は"
    "後継国間で大きく異なる：クロアチア（MAPE 4.1%）と北マケドニア（6.4%）は合理的な"
    "適合度を示し、比較的安定した紛争後の人口動態を反映する。ボスニア・ヘルツェゴビナ"
    "（8.1%）とスロベニア（12.2%）はより大きな誤差を示す——ボスニアは戦争関連の人口損失"
    "と避難、スロベニアは小規模なEU加盟国としての移民流入による成長が原因。セルビア"
    "（7.1%）とモンテネグロ（8.1%）はその中間に位置する。",
    size=11, space_after=12)

# --- C.4 Baltic ---
add_heading_styled(doc, "C.4 バルト三国：ソ連解体と移民流出（1991年）", level=3)
add_para(doc,
    "エストニア、ラトビア、リトアニアは1991年にソビエト連邦から独立し、その後ロシア系住民の"
    "著しい流出と西欧への移民（特に2004年のEU加盟後）が続いた。モデルMAPEは4.8%"
    "（エストニア）から7.1%（リトアニア）の範囲であり、閉鎖モデルが考慮しない持続的な"
    "移民流出を反映する。これらのケースは、穏やかだが持続的な純移出（年間人口の約0.5\u20131.0%）"
    "であっても30年間に蓄積されればモデルと観測の実質的な乖離を生むことを示す。",
    size=11, space_after=12)

# --- C.5 Ethiopia/Eritrea ---
add_heading_styled(doc, "C.5 エチオピアとエリトリア：分離（1993年）", level=3)
add_para(doc,
    "エリトリア独立（1993年）は高出生率の人口転換の最中にある2つの人口を分離した。モデル"
    "性能は顕著に異なる：エチオピア（MAPE 16.5%）は実質的な過大推計を示し、10年ごとの"
    "パラメータ更新が捕捉する以上に速い10年内の出生率低下を反映する。エリトリア"
    "（MAPE 37.8%）は自然実験サンプルで最大の誤差を示し、長期の徴兵制度、紛争関連の"
    "移民流出、そして非常に不確実なベースライン人口データが原因。これらのケースは、"
    "急速な人口変動が大規模な人口移動と同時に生じる紛争影響下・データ希少な環境における"
    "モデルの限界を浮き彫りにする。",
    size=11, space_after=6)

add_figure(doc, os.path.join(FIG_DIR, 'fig_natural_experiments_summary.png'),
    "図A-2：自然実験の概要",
    note="注：大規模な政治的・国境変更を経験した5ケースの人口軌跡。"
    "緑破線は後継国モデルの合成合計、黒実線は観測された（分裂前の）軌跡、"
    "色付き線は個別の後継国を示す。赤い垂直線は政治的変化の年を示す。",
    width=6.5)

# --- C.6 Synthesis ---
add_heading_styled(doc, "C.6 総合：自然実験が明らかにすること", level=3)

add_para(doc, "表C2：全自然実験ケースのモデル性能の要約",
         bold=True, size=11, space_after=4)

tbl_c2 = doc.add_table(rows=15, cols=4)
tbl_c2.style = 'Light Shading Accent 1'
for i, h in enumerate(['国', '事象（年）', 'MAPE（%）', '不適合の主因']):
    tbl_c2.rows[0].cells[i].text = h
c2_data = [
    ['ドイツ（合成東西）', '再統一（1990）', '6.4', '移民＋国内移動'],
    ['チェコ', 'ビロード離婚（1993）', '6.3', 'EU後の移民流出（穏やか）'],
    ['スロバキア', 'ビロード離婚（1993）', '9.9', 'EU・西欧への移民流出'],
    ['クロアチア', 'ユーゴ解体（1991）', '4.1', '紛争後の安定化'],
    ['スロベニア', 'ユーゴ解体（1991）', '12.2', '移民流入（EU加盟国）'],
    ['ボスニア', 'ユーゴ解体（1991）', '8.1', '戦争関連の避難'],
    ['セルビア', 'ユーゴ解体（1991）', '7.1', '難民流、移民流出'],
    ['北マケドニア', 'ユーゴ解体（1991）', '6.4', '穏やかな移民効果'],
    ['モンテネグロ', 'ユーゴ解体（1991）', '8.1', '小国、変動的なフロー'],
    ['エストニア', 'ソ連解体（1991）', '4.8', 'ロシア系住民の流出'],
    ['ラトビア', 'ソ連解体（1991）', '6.6', '移民流出（民族＋EU）'],
    ['リトアニア', 'ソ連解体（1991）', '7.1', '持続的な移民流出'],
    ['エチオピア', 'エリトリア独立（1993）', '16.5', '急速な出生率低下'],
    ['エリトリア', 'エリトリア独立（1993）', '37.8', '紛争、徴兵、移民流出'],
]
for i, rd in enumerate(c2_data):
    for j, val in enumerate(rd):
        tbl_c2.rows[i + 1].cells[j].text = val
add_para(doc, "", size=6, space_after=6)

add_para(doc,
    "これらの自然実験は内生更新モデルに対して3つの重要な知見をもたらす。",
    size=11, space_after=4)
add_para(doc,
    "第一に、モデルは大規模な政治的激変を経験した国々に対しても、事象後の移民が穏やかで"
    "人口データが信頼できる場合にはMAPE 8%未満と合理的に良好に機能する。クロアチア（4.1%）、"
    "エストニア（4.8%）、チェコ（6.3%）はいずれもこの範囲に収まり、カンタム・テンポ・生存"
    "メカニズムが政治的不連続性の文脈でも人口動態の大部分を捕捉することを示す。\n\n"
    "第二に、モデルと観測の乖離の大きさは移民成分の直接的な推定を提供する。ドイツの2023年"
    "までの15.8%の合成・観測ギャップは、移民が自然増加だけでは生じなかったであろう約1,350万人"
    "——総人口の約6分の1に相当——を追加したことを意味する。この定量化は閉鎖モデルが自然な"
    "人口動態を分離するからこそ可能である。\n\n"
    "第三に、モデルの限界は紛争影響下・データ希少な環境（エリトリア：37.8%）および持続的な"
    "移民流出が人口の大きな割合を除去する場合（スロベニア：12.2%、スロバキア：9.9%）において"
    "最も顕著である。これらのケースは内生更新フレームワークの境界条件を特定し、持続的な"
    "人口フローを経験する国々に対する移民モデリングの重要性を補強する。",
    size=11, space_after=12)

add_para(doc,
    "ドイツのケースは主要分析結果の解釈にとって特に示唆的である。40カ国検証における"
    "ドイツの動的MAPEはカンタム・テンポ・生存フレームワークの失敗ではなく、大規模な"
    "外生的移民事象としての再統一の量的足跡を反映する。モデルの「誤差」は実際にはシグナル"
    "である：3年間で190万人の国内移民を吸収し、その後30年間の持続的な国際移民が続くという"
    "政治的変革の人口学的影響を測定している。この簡素なモデルがこの移民シグナルを基盤となる"
    "人口動態から分離できることは、分解アプローチの妥当性を裏付ける。",
    size=11, space_after=12)

outpath = os.path.join(OUT_DIR, 'DemRes_Research_Article_JP.docx')
doc.save(outpath)
print(f'OK: {outpath}')
