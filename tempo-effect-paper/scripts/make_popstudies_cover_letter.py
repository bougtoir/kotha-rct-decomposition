"""Create Population Studies cover letter — English (.docx).

v3: Removed rejection history. Aligned with 'three levers' policy framework.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)


def add_para(doc, text, bold=False, italic=False, size=12, align=None,
             space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Date
add_para(doc, "[Date]", size=12, space_after=12)

# Addressee
add_para(doc, "The Editors", bold=True, size=12, space_after=2)
add_para(doc, "Population Studies", italic=True, size=12, space_after=2)
add_para(doc, "Population Investigation Committee", size=12, space_after=2)
add_para(doc, "London School of Economics and Political Science",
         size=12, space_after=18)

add_para(doc, "Dear Editors,", size=12, space_after=12)

# Opening — what the paper is
add_para(doc,
    "We are pleased to submit our manuscript entitled \u201cQuantifying the Tempo Effect on "
    "Simultaneously Living Population: Evidence from 40 Countries, 1970\u20132023\u201d for "
    "consideration as a research article in Population Studies.",
    size=12, space_after=12)

# The problem and why it matters
add_para(doc,
    "When discussing population decline, scholarship and policy almost exclusively focus "
    "on two demographic forces: how many children are born (fertility quantum) and how "
    "long people live (survival). Immigration is then invoked as the principal corrective "
    "mechanism. Yet demographic theory has long recognised a third, independent force: the "
    "timing of births. When childbearing is postponed, generational spacing widens, fewer "
    "generations overlap at any given moment, and the simultaneously living population "
    "(SLP) shrinks\u2014even if the number of children per woman remains unchanged. This tempo "
    "channel, identified by Goldstein et al. (2003), has received remarkably "
    "little empirical follow-up.",
    size=12, space_after=12)

# What the paper contributes
add_para(doc,
    "Our paper provides five contributions. First, we demonstrate that the tempo effect is "
    "invisible within standard projection frameworks. Comparing a tempo-invariant model "
    "(fertility level and survival updated decadally, but mean age at childbearing [MAC] "
    "held fixed) with a tempo-responsive model (all parameters updated), we find nearly "
    "identical fit (median absolute percentage error 4.5% versus 4.6%). This near-equivalence "
    "is itself the key insight: the standard period total fertility rate (TFR) mechanically "
    "absorbs the depression caused by rising MAC, conflating quantum decline with "
    "postponement. Policymakers who rely on period TFR cannot distinguish between fewer "
    "births and delayed births.",
    size=12, space_after=12)

add_para(doc,
    "Second, we show that explicit tempo decomposition restores visibility. Using the "
    "Bongaarts\u2013Feeney tempo-adjusted TFR (TFR*), which strips the timing distortion from "
    "period TFR, we achieve the best overall fit (median absolute percentage error 4.3%), "
    "with the largest improvements in countries experiencing strong postponement: the "
    "Republic of Korea, China, and Colombia. Where TFR* does not improve fit (e.g., France, "
    "Japan), the implication is that the TFR decline was predominantly quantum\u2014itself a "
    "policy-relevant distinction.",
    size=12, space_after=12)

add_para(doc,
    "Third, using a parsimonious endogenous renewal model validated against United Nations "
    "World Population Prospects 2024 data for 40 countries (38 Organisation for Economic "
    "Co-operation and Development members, China, the Democratic Republic of the Congo) over "
    "1970\u20132023, we find that the observed 4\u20136 year rise in MAC independently reduced the "
    "simultaneously living population (SLP) by 8\u201317%\u2014equivalent to 15\u201340 years of "
    "below-replacement fertility. Fourth, we decompose population change into quantum, tempo, "
    "and survival components, showing that tempo is typically the second-largest contributor "
    "in post-transitional countries.",
    size=12, space_after=12)

add_para(doc,
    "Fifth, we demonstrate that higher MAC accelerates the annual pace of population "
    "decline. This has a direct policy implication: the demographic response to population "
    "decline should not be framed as a binary choice between \u2018more births\u2019 and \u2018more "
    "immigration.\u2019 Birth timing constitutes a third, independent, and actionable policy "
    "lever. Tempo-sensitive interventions\u2014affordable housing for young families, universal "
    "childcare, educational pathways that do not penalise early parenthood\u2014could slow "
    "the pace of population decline and expand the window for institutional adaptation, "
    "without requiring increases in either fertility quantum or immigration. A bibliometric "
    "analysis of PubMed and a review of 15 national projection systems confirm that this "
    "tempo channel is currently absent from both the scholarly conversation and official "
    "demographic assessments.",
    size=12, space_after=12)

# Why Population Studies
add_para(doc,
    "We believe Population Studies is the ideal venue for this work. The journal\u2019s "
    "long tradition of publishing foundational contributions on fertility tempo and "
    "formal demography\u2014including the Bongaarts\u2013Feeney framework and related work on "
    "demographic translation\u2014means that the readership is uniquely positioned to "
    "evaluate and build upon our findings. Our argument that population policy requires "
    "a third lever, alongside quantum and survival, speaks directly to the journal\u2019s "
    "core audience of researchers and policymakers engaged with population dynamics.",
    size=12, space_after=12)

# Formalities
add_para(doc,
    "The manuscript has not been published elsewhere and is not under consideration by "
    "any other journal. All authors have approved the manuscript and agree to its "
    "submission. The authors declare no conflicts of interest.",
    size=12, space_after=12)

add_para(doc,
    "We look forward to your consideration.",
    size=12, space_after=18)

add_para(doc, "Yours sincerely,", size=12, space_after=6)
add_para(doc, "[Author name(s)]", size=12, space_after=2)
add_para(doc, "[Institutional affiliation]", size=12, space_after=2)
add_para(doc, "[Email address]", size=12, space_after=2)

outpath = os.path.join(OUT_DIR, 'CoverLetter_PopStudies_EN.docx')
doc.save(outpath)
print(f'OK: {outpath}')
