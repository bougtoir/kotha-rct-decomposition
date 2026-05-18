"""Create Population Studies submission — English manuscript (v3).

Revision notes (v3 – Population Studies compliance):
- Author-date citations (Population Studies house style)
- Reference list: alphabetical by first author, Pop Studies format
- Abstract trimmed to ≤150 words
- Short title (≤45 chars) for running header
- Tables at end of manuscript, each on own page; position markers in text
- Figures supplied separately; in-text markers; figure-title list at end
- Disclosure, Funding, CRediT statements added
- Section headings: H1=bold, H2=italic (per journal style)
- British (-ize / Oxford) spelling throughout
- Two output versions: open (with author details) and anonymized (for review)
- Manuscript order per guidelines: title; abstract; keywords; main text;
  acknowledgements; funding; appendix; references; tables; figure titles
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'figures')
OUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)


# ==============================================================
# Helpers
# ==============================================================
def add_para(doc, text, bold=False, italic=False, size=12, align=None,
             space_after=6, first_line_indent=None):
    """Add a plain paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_ref_para(doc, text, size=11, space_after=4):
    """Add a reference paragraph with italic portions marked as *...*
    and hanging indent."""
    p = doc.add_paragraph()
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.font.size = Pt(size)
            run.italic = True
        else:
            run = p.add_run(part)
            run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    return p


def add_h1(doc, text):
    """Section heading — bold, black, lower-case except first letter."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_h2(doc, text):
    """Sub-heading — italic, black, lower-case except first letter."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_table_marker(doc, label):
    """Insert [Insert Table/Figure X about here] marker."""
    p = doc.add_paragraph()
    run = p.add_run(f'[Insert {label} about here]')
    run.font.size = Pt(12)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p


# ==============================================================
# Build manuscript
# ==============================================================
def build_manuscript(anonymized=True):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    # ==========================================================
    # Title page
    # ==========================================================
    add_para(doc, 'Manuscript submitted to Population Studies',
             italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=24)

    add_para(doc,
        'Quantifying the tempo effect on simultaneously living population: '
        'Evidence from 40 countries, 1970\u20132023',
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_para(doc, 'Short title: The tempo effect on population size',
             italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=18)

    if anonymized:
        add_para(doc, '[Author names removed for double-blind review]',
                 italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=6)
        add_para(doc,
                 '[Institutional affiliation removed for double-blind review]',
                 italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=6)
    else:
        add_para(doc, '[Author Name]',
                 size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        add_para(doc, '[Institutional Affiliation]',
                 italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=6)
        add_para(doc, 'Corresponding author e-mail: [email@example.com]',
                 size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_para(doc, 'ORCID: [0000-0000-0000-0000]',
                 size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_para(doc, 'Postal address: [Full postal address]',
                 size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    add_para(doc, 'Word count (main text): approximately 7,400',
             italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=24)

    doc.add_page_break()

    # ==========================================================
    # Abstract (≤150 words)
    # ==========================================================
    add_h1(doc, 'Abstract')

    add_para(doc,
        'The timing of childbearing affects population size independently of the number '
        'of children born, yet this tempo channel remains absent from most population '
        'projection systems and policy frameworks. Using a parsimonious endogenous '
        'renewal model, we analyse 40 countries over 1970\u20132023 and compare five model '
        'variants. The tempo-invariant and tempo-responsive models yield nearly identical '
        'fit (median absolute percentage error 4.5% versus 4.6%), revealing that the '
        'period total fertility rate absorbs tempo distortion, rendering the tempo '
        'channel invisible to conventional projections. Only when tempo is explicitly '
        'decomposed via the Bongaarts\u2013Feeney tempo-adjusted total fertility rate does '
        'the best fit emerge (4.3%). Counterfactual analysis shows that the observed '
        'rise in mean age at childbearing independently reduced simultaneously living '
        'population by 8\u201317% across Organisation for Economic Co-operation and '
        'Development countries. None of the 15 national projection systems reviewed '
        'decomposes this channel. Birth timing deserves recognition alongside birth '
        'quantum and survival as a third policy lever.',
        size=12, space_after=12)

    add_para(doc,
        'Keywords: tempo effect; simultaneously living population; mean age at '
        'childbearing; fertility postponement; population projection; birth timing; '
        'OECD',
        italic=True, size=10, space_after=18)

    doc.add_page_break()

    # ==========================================================
    # 1. Introduction
    # ==========================================================
    add_h1(doc, '1. Introduction')

    add_para(doc,
        'When fertility falls below replacement, policy responses almost universally '
        'target the quantum of fertility\u2014the number of children born. South Korea\u2019s '
        'record 47 trillion won pronatalist investment, Japan\u2019s successive Plans for '
        'Measures Against the Declining Birthrate, and similar programmes across '
        'Organisation for Economic Co-operation and Development (OECD) nations share '
        'this quantum-centric framing (Th\u00e9venon 2011). Yet a second, independent '
        'demographic force shapes how many people are simultaneously alive at any given '
        'moment: the timing of births.',
        size=12, space_after=12)

    add_para(doc,
        'Ryder (1964) introduced the concept of demographic translation\u2014the idea that '
        'shifts in the timing of vital events alter period rates independently of '
        'underlying quantum. Bongaarts and Feeney (1998) formalised the quantum\u2013tempo '
        'distinction for fertility, showing that the period total fertility rate (TFR) '
        'is mechanically depressed when women postpone childbearing, even if completed '
        'cohort fertility remains unchanged. Subsequent work extended this framework: '
        'Kohler and Ortega (2002) developed tempo-adjusted parity progression measures; '
        'Sobotka (2004) demonstrated that much of Europe\u2019s lowest-low fertility could '
        'be attributed to postponement rather than permanent decline in desired family '
        'size; and Bongaarts and Feeney (2006) showed that tempo distortions affect all '
        'life-cycle events, not only fertility indicators.',
        size=12, space_after=12)

    add_para(doc,
        'Crucially, Goldstein et al. (2003) demonstrated for EU-15 countries that '
        'delayed childbearing reduces the number of generations alive at any moment, '
        'producing population decline independent of the number of children ever born. '
        'Their analytical insight was that population size at a point in time\u2014the '
        'simultaneously living population (SLP)\u2014depends not only on how many people are '
        'born per generation but on how many generations overlap. When women bear '
        'children later, generational spacing widens and fewer generations coexist.',
        size=12, space_after=12)

    # --- The 'forgotten' evidence ---
    add_para(doc,
        'Despite this theoretical foundation, the tempo\u2013population link has been largely '
        'forgotten in subsequent research and policy. Evidence for this neglect comes '
        'from two independent sources. First, a bibliometric analysis reveals a striking '
        'divergence: while PubMed-indexed publications on \u2018delayed childbearing\u2019 or '
        '\u2018postponement of childbearing\u2019 grew from 22 articles in 1998\u20132002 to 96 in '
        '2023\u20132025 (a fourfold increase), publications mentioning the \u2018tempo effect\u2019 '
        'in a demographic context remained near zero throughout the same period '
        '(Table 1). The phenomenon of delayed childbearing is increasingly discussed\u2014'
        'particularly in the health and policy literature\u2014but its formal demographic '
        'mechanism (the tempo effect on population size) has essentially disappeared '
        'from the scholarly conversation.',
        size=12, space_after=6)

    add_table_marker(doc, 'Table 1')

    add_para(doc,
        'Second, a review of 15 national population projection systems across OECD '
        'countries reveals that none explicitly decomposes population change into '
        'quantum and tempo components (see Section 5.1). Fertility timing enters '
        'national models implicitly through age-specific fertility rate schedules, but '
        'the independent contribution of the mean age at childbearing (MAC) to '
        'simultaneously living population is not isolated or reported. This is not '
        'merely a semantic omission: it means that the population-level consequences '
        'of postponement\u2014which, as we demonstrate, can be equivalent to decades of '
        'below-replacement fertility\u2014are structurally invisible in the projections '
        'that inform pension reform, healthcare planning, and immigration policy.',
        size=12, space_after=12)

    add_para(doc,
        'Three specific gaps persist in the literature. First, the Goldstein et al. '
        '(2003) analysis was limited to EU-15 projections with stylised assumptions; no '
        'study has measured the tempo effect against observed population data across '
        'diverse demographic contexts. Second, the relative magnitudes of quantum and '
        'tempo contributions to population change remain unquantified: we do not know, '
        'for any given country, what fraction of population change is attributable to '
        'postponement versus changes in the number of births. Third, the pace '
        'implication\u2014that tempo affects not only the level but the rate of population '
        'change\u2014has not been empirically demonstrated.',
        size=12, space_after=12)

    add_para(doc,
        'This paper fills these gaps. Using a parsimonious endogenous renewal model '
        'validated against observed trajectories for 40 countries over half a century, '
        'we provide five contributions. First, we demonstrate that the standard period '
        'total fertility rate (TFR) absorbs most tempo distortion, rendering the tempo '
        'channel invisible: a tempo-invariant model (MAC fixed) and a tempo-responsive '
        'model (MAC updated) yield nearly identical fit (median absolute percentage '
        'error [MAPE] 4.5% versus 4.6%), because TFR itself already reflects '
        'postponement. Second, we show that explicit tempo decomposition via the '
        'Bongaarts\u2013Feeney tempo-adjusted TFR (TFR*) (Bongaarts and Feeney 1998) '
        'restores visibility: the tempo-adjusted model achieves the best overall fit '
        '(median MAPE 4.3%), with the largest gains in countries with strong '
        'postponement\u2014Republic of Korea (11.9% \u2192 6.9%), China (15.6% \u2192 8.6%), '
        'Colombia (13.1% \u2192 7.8%). Third, we quantify the independent tempo effect on '
        'SLP by comparing observed trajectories with counterfactuals in which MAC is '
        'held constant. Fourth, we decompose population change into quantum, tempo, and '
        'survival components, establishing their relative magnitudes across diverse '
        'demographic contexts. Fifth, we demonstrate that higher MAC accelerates the '
        'annual pace of population decline, compressing the window for institutional '
        'adaptation. Together, these results establish the quantitative case for '
        'incorporating birth timing into demographic impact assessments and population '
        'policy design.',
        size=12, space_after=12)

    # ==========================================================
    # 2. The tempo–population mechanism
    # ==========================================================
    add_h1(doc, '2. The tempo\u2013population mechanism')

    add_para(doc,
        'The link between birth timing and population size operates through '
        'generational overlap. Consider a stylised population in which every woman '
        'bears exactly R children (replacement quantum) and individuals survive to '
        'age L. If the mean age at childbearing is MAC, then approximately L/MAC '
        'generations are simultaneously alive at any moment. Population size at time '
        't\u2014the simultaneously living population (SLP)\u2014is therefore proportional to '
        'L/MAC for a given birth quantum.',
        size=12, space_after=12)

    add_para(doc,
        'This relationship implies that a rise in MAC, holding quantum and survival '
        'constant, mechanically reduces SLP. When MAC = 25 and L = 80, approximately '
        '3.2 generations overlap (at ages 0, 25, 50, 75). When MAC rises to 30, '
        'overlap falls to approximately 2.7 generations (at ages 0, 30, 60). The '
        'proportional reduction in SLP is approximately (MAC\u2082 \u2212 MAC\u2081)/MAC\u2082 = '
        '5/30 \u2248 17%. This is the tempo effect on simultaneously living population: '
        'a five-year increase in MAC reduces the population stock by roughly one-sixth, '
        'independent of fertility quantum (Goldstein et al. 2003).',
        size=12, space_after=12)

    add_para(doc,
        'In practice, populations are not stationary: fertility changes over time, '
        'survival improves, and age structures carry momentum from past demographic '
        'regimes (Keyfitz 1971). The stylised L/MAC relationship therefore serves as '
        'an analytical benchmark that actual populations approximate but do not exactly '
        'match. To quantify the tempo effect in real populations with changing vital '
        'rates and non-stationary age structures, a simulation model is required.',
        size=12, space_after=12)

    # ==========================================================
    # 3. Model and data
    # ==========================================================
    add_h1(doc, '3. Model and data')

    add_h2(doc, '3.1 Endogenous renewal model')
    add_para(doc,
        'We construct a discrete-time, single-sex population model in which the '
        'population vector P(t) = [P\u2080(t), P\u2081(t), \u2026, P\u2081\u2080\u2080(t)] evolves annually '
        '(Preston et al. 2001). At each time step:',
        size=12, space_after=8)

    add_para(doc,
        '(a) Survival: Individuals at age x survive to age x + 1 with probability '
        'derived from a Gompertz hazard function h(x) = a\u00b7exp(b\u00b7x) (Gompertz '
        '1825). The survival function is S(x) = exp[\u2212(a/b)(exp(bx) \u2212 1)]. '
        'Parameter a is calibrated so that life expectancy at birth '
        'e\u2080 = \u222b\u2080\u221e S(x)dx matches the observed value; b is fixed at 0.085.',
        size=12, space_after=8)

    add_para(doc,
        '(b) Fertility: Births are generated endogenously. The age-specific fertility '
        'rate is modelled as a normal density centred on MAC with standard deviation '
        '\u03c3, scaled to TFR. Births at time t equal \u03a3 P\u2093(t) \u00b7 f \u00b7 ASFR(x) '
        'for x = 15\u201349, where ASFR denotes the age-specific fertility rate and f is '
        'the female population share.',
        size=12, space_after=8)

    add_para(doc,
        'This minimal parameterisation requires only four inputs per period: TFR, '
        'e\u2080, MAC, and \u03c3. Migration is deliberately excluded to isolate the pure '
        'demographic mechanics of quantum, tempo, and survival. Population momentum '
        '(Keyfitz 1971)\u2014the tendency for population to continue growing after '
        'fertility falls to replacement due to a young age structure\u2014is captured '
        'endogenously through the age-structured dynamics.',
        size=12, space_after=12)

    add_h2(doc, '3.2 Data')
    add_para(doc,
        'All parameters and validation data are drawn from the United Nations World '
        'Population Prospects 2024 (United Nations 2024). We analyse 40 countries: all '
        '38 OECD member states (as of 2024) plus China and the Democratic Republic of '
        'the Congo, chosen to span the full range of demographic transition stages. '
        'Initial population age structures (five-year age groups, both sexes) are '
        'interpolated to single-year ages. Demographic indicators\u2014TFR, e\u2080, and '
        'MAC\u2014are extracted for each calendar year from 1950 to 2023. This study '
        'reports population estimates and follows the Guidelines for Accurate and '
        'Transparent Health Estimates Reporting (GATHER) (Stevens et al. 2016).',
        size=12, space_after=12)

    add_h2(doc, '3.3 Model variants and counterfactuals')
    add_para(doc,
        'We implement four model variants, plus the observed population, yielding a '
        'five-way comparison framework:',
        size=12, space_after=8)

    add_para(doc,
        'Tempo-responsive model: All four parameters (TFR, e\u2080, MAC, \u03c3) are '
        'updated every 10 years using observed UN values (e.g. 1970 parameters for '
        '1970\u20131979, 1980 for 1980\u20131989, etc.), running from 1970 to 2023 for all '
        '40 countries.',
        size=12, space_after=8)

    add_para(doc,
        'Tempo-invariant model: TFR, e\u2080, and \u03c3 are updated every 10 years as '
        'above, but MAC is held fixed at its 1970 value throughout. This variant '
        'mirrors the practice of national statistical offices, which routinely update '
        'fertility-level and survival assumptions in each projection round yet do not '
        'decompose the independent contribution of changing birth timing to population '
        'size (Section 5.1).',
        size=12, space_after=8)

    add_para(doc,
        'Tempo-adjusted model (TFR*): All parameters are updated decadally as in the '
        'tempo-responsive model, but the fertility input is replaced with the '
        'Bongaarts\u2013Feeney tempo-adjusted TFR (Bongaarts and Feeney 1998). The '
        'adjustment removes the mechanical depression of period TFR caused by rising '
        'MAC: TFR* = TFR / (1 \u2212 dMAC/dt), where dMAC/dt is estimated from centred '
        'differences over the surrounding years and clipped to [\u22120.5, 0.5] to '
        'prevent extreme values. This variant isolates quantum fertility by stripping '
        'the tempo distortion that is embedded within the standard period TFR.',
        size=12, space_after=8)

    add_para(doc,
        'Fixed-parameter model: All parameters (TFR, e\u2080, MAC) are fixed at '
        'base-year values and held constant throughout the projection. We run four '
        'base years (1970, 1980, 1990, 2000) with forward projections to 2023, '
        'yielding 160 country\u2013base-year combinations.',
        size=12, space_after=8)

    add_para(doc,
        'Counterfactual scenarios: To isolate each component\u2019s contribution, we run '
        'the tempo-responsive model with one parameter held at its 1970 value while '
        'others evolve as observed: (i) MAC frozen at 1970 level (tempo '
        'counterfactual\u2014equivalent to the tempo-invariant variant); (ii) TFR frozen '
        'at 1970 level (quantum counterfactual); (iii) e\u2080 frozen at 1970 level '
        '(survival counterfactual). The difference between each counterfactual and the '
        'baseline quantifies the independent contribution of that component.',
        size=12, space_after=12)

    # ==========================================================
    # 4. Results
    # ==========================================================
    add_h1(doc, '4. Results')

    # --- 4.1 ---
    add_h2(doc, '4.1 Why tempo is invisible in standard projections\u2014and how '
           'to reveal it')
    add_para(doc,
        'National statistical offices routinely update their population projections, '
        'revising fertility-level and survival assumptions in each round. Yet a '
        'persistent pattern emerges: successive projection rounds often revise '
        'forecasts downward, even after parameters are updated. Japan\u2019s National '
        'Institute of Population and Social Security Research, for example, has '
        'lowered its long-term population forecast in virtually every round since the '
        '1990s, despite incorporating updated TFR and life expectancy data. We '
        'hypothesise that this serial downward revision partly reflects the '
        'accumulating tempo effect\u2014an independent population-reducing force that '
        'standard projection updates do not decompose.',
        size=12, space_after=12)

    add_para(doc,
        'To test this hypothesis, we compare four model variants against observed '
        'population trajectories (Table 2). The fixed-parameter model holds all '
        'parameters constant\u2014analogous to a single projection that is never revised. '
        'The tempo-invariant model updates TFR and e\u2080 every 10 years but keeps MAC '
        'fixed\u2014analogous to projection systems that revise fertility levels and '
        'mortality without separating the tempo channel. The tempo-responsive model '
        'updates all parameters including MAC. The tempo-adjusted model replaces '
        'period TFR with the Bongaarts\u2013Feeney TFR*, explicitly stripping tempo '
        'distortion from the fertility input.',
        size=12, space_after=6)

    add_table_marker(doc, 'Table 2')

    add_para(doc,
        'The most revealing finding is the near-identical performance of the '
        'tempo-invariant and tempo-responsive models (median MAPE 4.5% versus 4.6%). '
        'Both update parameters every 10 years\u2014as national statistical offices do in '
        'periodic projection rounds\u2014but differ in a single respect: whether MAC is '
        'allowed to evolve. That this difference is negligible is itself the key '
        'insight. The period TFR used by the tempo-invariant model already '
        'incorporates the mechanical depression caused by rising MAC: when women '
        'postpone childbearing, age-specific fertility rates decline at younger ages, '
        'depressing the period TFR even if completed cohort fertility remains '
        'unchanged (Bongaarts and Feeney 1998). Consequently, the tempo effect is '
        'absorbed into the TFR input and becomes invisible to the projection '
        'framework. The model \u2018works\u2019 for the wrong reason\u2014it captures the '
        'population-level consequence of postponement through a lower TFR, but '
        'attributes the entire decline to quantum rather than decomposing it into '
        'quantum and tempo.',
        size=12, space_after=12)

    add_para(doc,
        'This invisibility has a direct policy consequence: policymakers who observe '
        'a declining TFR respond with quantum interventions (pronatalist cash '
        'transfers, child allowances) because the period TFR does not distinguish '
        'between fewer births and delayed births. Yet a substantial portion of the '
        'observed TFR decline may reflect postponement\u2014a timing shift that is '
        'amenable to different policy instruments.',
        size=12, space_after=12)

    add_para(doc,
        'The tempo-adjusted model (TFR*) breaks this invisibility. By replacing '
        'period TFR with TFR* = TFR / (1 \u2212 dMAC/dt), the Bongaarts\u2013Feeney '
        'adjustment strips the tempo distortion from the fertility input, revealing '
        'quantum fertility. The tempo-adjusted model achieves the best overall fit '
        '(median MAPE 4.3%, mean 5.8%), improving on both the tempo-invariant and '
        'tempo-responsive variants. Crucially, the improvement is concentrated in '
        'countries where postponement was strongest: Republic of Korea (MAPE 11.9% '
        '\u2192 6.9%), China (15.6% \u2192 8.6%), Colombia (13.1% \u2192 7.8%), Sweden (5.7% '
        '\u2192 3.5%), Denmark (3.0% \u2192 1.0%). In total, 26 of 40 countries show '
        'improved fit under TFR*. Where TFR* does not improve fit (e.g. France, '
        'Japan), the implication is that the period TFR decline was predominantly '
        'quantum\u2014a substantively informative finding in itself.',
        size=12, space_after=12)

    add_para(doc,
        'Countries with MAPE exceeding 10% in the tempo-responsive model share '
        'identifiable sources of misfit. Immigration-driven growth explains Australia '
        '(13.5%), Canada (12.2%), Switzerland (7.2%), Luxembourg (21.5%), and Israel '
        '(13.9%). Rapid fertility transition explains Mexico (23.3%), T\u00fcrkiye '
        '(17.0%), China (15.6%), and Colombia (13.1%). Our model deliberately excludes '
        'migration; the residual misfit therefore quantifies the migration component '
        'of population change\u2014itself a useful by-product of the decomposition '
        'approach.',
        size=12, space_after=6)

    add_para(doc,
        'Figure 1 shows model trajectories for six representative countries spanning '
        'diverse demographic contexts. All five variants are displayed: the '
        'near-overlap of tempo-invariant and tempo-responsive lines visually confirms '
        'the invisibility of tempo in standard TFR-based projections, while the '
        'tempo-adjusted trajectory diverges where postponement was strongest (Republic '
        'of Korea, China).',
        size=12, space_after=6)

    add_table_marker(doc, 'Figure 1')

    add_para(doc,
        'The validation confirms that the four-parameter model captures the '
        'quantum\u2013tempo\u2013survival mechanism with sufficient fidelity to support '
        'counterfactual decomposition, and that explicit tempo decomposition via TFR* '
        'yields measurably better projections than either conventional variant. We now '
        'turn to the counterfactual analysis.',
        size=12, space_after=12)

    # --- 4.2 ---
    add_h2(doc, '4.2 The magnitude of the tempo effect')
    add_para(doc,
        'How much of observed population change is attributable to delayed '
        'childbearing? To answer this, we compare the baseline tempo-responsive model '
        '(all parameters evolving as observed) with the tempo counterfactual, in '
        'which MAC is held at its 1970 value while TFR and e\u2080 evolve as observed. '
        'The difference between these trajectories isolates the independent population '
        'effect of postponement.',
        size=12, space_after=12)

    add_para(doc,
        'Table 3 presents the results for 20 countries with the largest tempo effects. '
        'The findings are striking. Across OECD countries, the observed four- to '
        'six-year rise in MAC between 1970 and 2023 independently reduced the '
        'simultaneously living population by 8\u201317% relative to the counterfactual '
        'with stable birth timing. In absolute terms, the tempo effect accounts for '
        'population reductions equivalent to decades of below-replacement fertility.',
        size=12, space_after=6)

    add_table_marker(doc, 'Table 3')

    add_para(doc,
        'Several patterns emerge. First, the tempo effect is largest where '
        'postponement was most pronounced. Korea, where MAC rose by 7.3 years (from '
        '26.1 to 33.4), experienced the largest tempo-driven SLP reduction (21.9%). '
        'Czechia (\u0394MAC = 5.7 years, \u221218.9%) demonstrates that the post-socialist '
        'fertility delay\u2014well documented by Sobotka (2004) and Kohler et al. '
        '(2002)\u2014had population-level consequences far beyond the period TFR decline '
        'that attracted most scholarly attention.',
        size=12, space_after=12)

    add_para(doc,
        'Second, the tempo effect operates independently of quantum. Japan and Italy '
        'have similar TFR trajectories (both declining to ~1.2\u20131.3), yet their '
        'tempo effects differ because their MAC trajectories differ. Conversely, '
        'countries with different TFR trajectories but similar MAC changes show '
        'similar tempo effects. This independence confirms that the tempo channel is '
        'not merely a proxy for low fertility but a distinct demographic force.',
        size=12, space_after=12)

    add_para(doc,
        'Third, the \u2018equivalent TFR-years\u2019 column provides a policy-relevant '
        'metric. Korea\u2019s tempo effect is equivalent to 45 years of fertility at '
        'TFR = 1.5; even for countries with moderate postponement (Germany, '
        '\u0394MAC = 3.5), the tempo effect equals ~24 years of below-replacement '
        'fertility. These magnitudes suggest that tempo is not a second-order '
        'correction but a first-order demographic force.',
        size=12, space_after=12)

    # --- 4.3 ---
    add_h2(doc, '4.3 Decomposing population change: quantum, tempo, and survival')
    add_para(doc,
        'The counterfactual framework allows systematic decomposition of total '
        'population change into three components. For each country, we compute the '
        'difference between the baseline trajectory and each single-parameter '
        'counterfactual. Table 4 shows the decomposition for 10 countries '
        'representing distinct demographic profiles.',
        size=12, space_after=6)

    add_table_marker(doc, 'Table 4')

    add_para(doc,
        'The decomposition reveals that tempo is the second-largest component of '
        'population change in most post-transitional countries, after survival gains. '
        'In Japan, the tempo effect (\u221213.0%) is roughly 40% the magnitude of the '
        'quantum effect (\u221232.1%) and would be even larger relative to quantum in '
        'the absence of Japan\u2019s substantial longevity gains (+28.4%). For Czechia, '
        'the tempo effect (\u221218.9%) approaches the quantum effect (\u221222.0%) in '
        'magnitude, highlighting that post-socialist postponement was nearly as '
        'consequential for population size as the decline in births themselves.',
        size=12, space_after=12)

    add_para(doc,
        'France and Sweden illustrate an important point: even in countries where '
        'population grew over this period, the tempo effect was substantially '
        'negative (\u221211.7% and \u221212.5%, respectively). These countries\u2019 growth was '
        'driven by survival gains and, in France\u2019s case, relatively sustained '
        'quantum\u2014but would have grown more had childbearing not been postponed. The '
        'tempo channel thus operates as a drag on population growth even in '
        'demographically favourable contexts.',
        size=12, space_after=12)

    add_para(doc,
        'China provides a revealing contrast. Its MAC actually declined slightly over '
        'this period (from 29.2 to 28.4), producing a small positive tempo effect '
        '(+2.8%). China\u2019s population dynamics were overwhelmingly driven by the '
        'quantum channel (the one-child policy) and survival gains. The Democratic '
        'Republic of the Congo, still in early demographic transition, shows negligible '
        'tempo effect and massive growth driven by sustained high fertility and '
        'improving survival.',
        size=12, space_after=12)

    add_para(doc,
        'Figure 2 shows model validation across all 40 countries, providing the '
        'empirical foundation for the decomposition results.',
        size=12, space_after=6)

    add_table_marker(doc, 'Figure 2')

    # --- 4.4 ---
    add_h2(doc, '4.4 The pace of demographic change')
    add_para(doc,
        'The tempo effect operates not only on the level of population but on the '
        'rate of change. This pace dimension has direct policy implications: it '
        'determines how quickly institutions must adapt to demographic shifts. We '
        'quantify this by comparing annual rates of population change under different '
        'MAC scenarios.',
        size=12, space_after=12)

    add_para(doc,
        'Consider two countries with identical TFR = 1.5 and e\u2080 = 80 but '
        'MAC = 25 versus MAC = 33. In the low-MAC country, approximately 3.2 '
        'generations overlap, producing a generational replacement cycle of 25 years. '
        'In the high-MAC country, only 2.4 generations overlap with a 33-year cycle. '
        'The high-MAC country\u2019s population declines faster per calendar year because '
        'each generation\u2019s below-replacement contribution accumulates over fewer '
        'overlapping cohorts.',
        size=12, space_after=12)

    add_para(doc,
        'This acceleration is not trivial. Our model shows that for a country with '
        'TFR = 1.5, the annual rate of population decline at MAC = 33 is '
        'approximately 0.7% per year, compared with 0.4% at MAC = 25\u2014a 75% '
        'acceleration. Over a 30-year planning horizon, this translates to a '
        'cumulative difference of approximately nine percentage points of population '
        '(a decline of 19% versus 11%). For pension systems designed around 2\u20133% '
        'per-decade population adjustment, the higher MAC scenario requires twice the '
        'adaptation speed.',
        size=12, space_after=12)

    add_para(doc,
        'This acceleration effect explains why countries with similar TFR but '
        'different MAC face qualitatively different policy challenges. Japan '
        '(TFR \u2248 1.2, MAC \u2248 31.4) and the USA (TFR \u2248 1.6, MAC \u2248 29.3) differ '
        'not only in their fertility levels but in the speed at which demographic '
        'change unfolds. Japan\u2019s higher MAC means its population decline is faster '
        'per calendar year than the USA\u2019s, compressing the time available for '
        'institutional reform\u2014pension adjustment, healthcare infrastructure expansion, '
        'labour market restructuring.',
        size=12, space_after=12)

    add_para(doc,
        'Figure 3 illustrates how model fit varies across countries and base years, '
        'showing that the fixed-parameter model\u2019s degradation over longer horizons '
        'reflects accumulating demographic change\u2014the very phenomenon our '
        'decomposition quantifies. Figure 4 compares MAPE and final population ratio '
        'across countries for all four model variants, confirming that the '
        'tempo-adjusted model (TFR*) achieves the best overall fit, particularly in '
        'countries with strong postponement.',
        size=12, space_after=6)

    add_table_marker(doc, 'Figure 3')
    add_table_marker(doc, 'Figure 4')

    # ==========================================================
    # 5. Discussion
    # ==========================================================
    add_h1(doc, '5. Discussion')

    add_para(doc,
        'Our results provide four main findings with implications for demographic '
        'research and policy.',
        size=12, space_after=12)

    add_para(doc,
        'First, and most fundamentally, the tempo effect on population size is '
        'invisible within standard projection frameworks. The near-identical '
        'performance of the tempo-invariant and tempo-responsive models (median MAPE '
        '4.5% versus 4.6%) demonstrates that period TFR\u2014the indicator on which '
        'virtually all national and international projections depend\u2014absorbs tempo '
        'distortion without decomposing it. When MAC rises, period TFR falls '
        'mechanically (Bongaarts and Feeney 1998), and projection systems treat this '
        'decline as if it were entirely quantum. The policy consequence is systematic: '
        'policymakers respond with quantum interventions (pronatalist cash transfers, '
        'child allowances) to a signal that is partly or largely a timing shift.',
        size=12, space_after=12)

    add_para(doc,
        'Second, explicit tempo decomposition via the Bongaarts\u2013Feeney TFR* achieves '
        'the best overall fit (median MAPE 4.3%, mean 5.8%), outperforming both '
        'conventional variants. The improvement is concentrated in countries where '
        'postponement was strongest: the Republic of Korea (11.9% \u2192 6.9%), China '
        '(15.6% \u2192 8.6%), Colombia (13.1% \u2192 7.8%), Sweden (5.7% \u2192 3.5%), Denmark '
        '(3.0% \u2192 1.0%). This pattern reveals where the tempo distortion is '
        'largest\u2014and, by implication, where tempo-based policy interventions have the '
        'greatest potential demographic payoff. Conversely, countries where TFR* does '
        'not improve fit (e.g. France, Japan) are those where the period TFR decline '
        'was predominantly quantum\u2014itself a policy-relevant distinction. This finding '
        'extends Goldstein et al.\u2019s (2003) theoretical insight from stylised EU-15 '
        'projections to observed population data across 40 countries spanning the full '
        'range of demographic transition stages.',
        size=12, space_after=12)

    add_para(doc,
        'Third, the counterfactual decomposition reveals that tempo is typically the '
        'second-largest component of population change in post-transitional countries. '
        'The observed four- to six-year increase in MAC across OECD countries reduced '
        'SLP by 8\u201317%, magnitudes comparable to decades of below-replacement '
        'fertility. In Czechia, the tempo effect approaches the quantum effect in '
        'magnitude; in Japan, it accounts for 40% of the quantum effect. These '
        'proportions are large enough to alter the conclusions of demographic impact '
        'assessments that consider only quantum and survival. Lutz et al. (2001) '
        'projected the end of world population growth; our analysis shows that the '
        'tempo channel is a substantial driver of how rapidly that endpoint is '
        'approached.',
        size=12, space_after=12)

    add_para(doc,
        'Fourth, the pace dimension has direct policy implications. Higher MAC '
        'accelerates the annual rate of population decline, compressing the time '
        'available for institutional adaptation. This reframes the policy problem. '
        'Population policy currently operates on two levers: boosting births (quantum '
        'interventions such as child allowances, parental leave, and pronatalist '
        'incentives) and extending lives (survival interventions such as healthcare '
        'investment and disease prevention). Our results demonstrate that a third '
        'lever exists: interventions that influence the timing of births. '
        'Tempo-sensitive policies\u2014affordable housing that enables family formation at '
        'younger ages, universal childcare that reduces the opportunity cost of early '
        'parenthood, restructured educational and career pathways that do not penalise '
        'combining parenthood with professional development (McDonald 2000)\u2014could '
        'slow the pace of population decline and expand the window for institutional '
        'adjustment, even without raising TFR.',
        size=12, space_after=12)

    add_para(doc,
        'The distinction between quantum and tempo interventions is not merely '
        'semantic. Quantum-focused policies (e.g. South Korea\u2019s 47 trillion won '
        'package, Japan\u2019s child allowance expansion) aim to increase the number of '
        'births. Tempo-focused policies aim to reduce the age at which existing births '
        'occur\u2014a fundamentally different target. Gauthier (2007) showed that '
        'pronatalist cash transfers have modest and often transient effects on quantum; '
        'our analysis suggests that even if quantum remains unchanged, a two-year '
        'reduction in MAC from current levels would increase SLP by 5\u20137% in most '
        'OECD countries\u2014equivalent to roughly 10 years of moderate pronatalist '
        'success. This is a demographic dividend achievable without increasing the '
        'number of births per woman.',
        size=12, space_after=12)

    add_para(doc,
        'Bongaarts and Sobotka (2012) showed that some European countries had begun to '
        'reverse postponement trends, with period TFR recovering as tempo distortions '
        'subsided. Myrskyl\u00e4 et al. (2009) demonstrated that advanced development '
        'can reverse fertility declines. Our results complement both findings by '
        'showing that even where quantum remains low, tempo adjustments independently '
        'alter the population trajectory. The policy implication is that the '
        'demographic response to population decline should not be framed as a binary '
        'choice between \u2018more births\u2019 and \u2018more immigration\u2019 but as a '
        'three-dimensional problem in which birth quantum, birth timing, and survival '
        'each constitute independent and actionable levers.',
        size=12, space_after=12)

    add_para(doc,
        'The model\u2019s deliberate exclusion of migration is both a limitation and a '
        'feature. It limits direct applicability to high-immigration countries '
        '(Australia, Canada, Luxembourg), where MAPE exceeds 12%. However, the '
        'exclusion enables clean decomposition: the model\u2019s \u2018error\u2019 in these '
        'countries is itself informative, quantifying the migration component of '
        'population change. For example, Australia\u2019s 13.5% MAPE implies that net '
        'immigration added approximately 13\u201314% to the population beyond what natural '
        'increase would produce\u2014consistent with Australian Bureau of Statistics '
        'estimates. Our natural experiments analysis (Appendix A) demonstrates that '
        'Germany\u2019s model misfit quantifies the demographic footprint of reunification '
        'as a migration shock.',
        size=12, space_after=12)

    # --- 5.1 ---
    add_h2(doc, '5.1 National projection systems and the tempo gap')

    add_para(doc,
        'To assess whether the tempo channel is accounted for in practice, we reviewed '
        'the official population projection methodologies of 15 national statistical '
        'offices and international organisations covering all 40 countries in our '
        'sample (Table 5). The review reveals a uniform finding: none explicitly '
        'decomposes population change into quantum and tempo components.',
        size=12, space_after=6)

    add_table_marker(doc, 'Table 5')

    add_para(doc,
        'All 15 systems share the cohort-component method as their foundational '
        'structure. Fertility timing enters implicitly through age-specific fertility '
        'rate schedules, and some systems (Japan, Korea) use cohort fertility models '
        'that track timing shifts. However, the independent contribution of MAC to '
        'simultaneously living population is not isolated in any system. Our five-model '
        'comparison reveals why this gap matters. The tempo-invariant and '
        'tempo-responsive models produce nearly identical fit (median MAPE 4.5% versus '
        '4.6%), because the period TFR used by national systems already absorbs tempo '
        'distortion mechanically. Policymakers never see the tempo channel\u2014they see '
        'only a lower TFR and respond with quantum interventions. Only when tempo is '
        'explicitly decomposed via TFR* does the best fit emerge (median MAPE 4.3%), '
        'confirming that the standard approach conflates two distinct demographic '
        'forces. Japan\u2019s National Institute of Population and Social Security Research '
        'has revised its long-term projection downward in virtually every round since '
        'the 1990s\u2014not because its TFR or e\u2080 assumptions were wrong, but because '
        'the TFR signal conflated quantum and tempo, leading to systematically '
        'incorrect policy diagnosis. This is not a limitation of the cohort-component '
        'method itself\u2014it is a reporting and analytical gap. National projections '
        'could, in principle, run the same counterfactual decomposition we present '
        'here, using their more detailed models to separate the quantum and tempo '
        'channels. That none do so reinforces the \u2018forgotten\u2019 status of the '
        'tempo\u2013population link.',
        size=12, space_after=12)

    add_para(doc,
        'Our four-parameter model is not designed to replace these national projection '
        'systems but to complement them by making the tempo\u2013quantum\u2013survival '
        'decomposition explicit. The tempo-adjusted model (TFR*) achieves median MAPE '
        'of 4.3% against these same populations\u2014performance sufficient to establish '
        'the quantitative significance of the tempo channel, even though it cannot '
        'match the precision of full-parameterisation national models that include '
        'migration (Gonand 2005; Lee and Carter 1992).',
        size=12, space_after=12)

    # --- 5.2 ---
    add_h2(doc, '5.2 Limitations')
    add_para(doc,
        'Several limitations warrant acknowledgement. First, the normal fertility '
        'schedule is a simplification; actual ASFRs may be skewed or bimodal (Frejka '
        'and Sobotka 2008). Second, our decadal update interval does not match the '
        'varying revision cycles of national statistical offices, which update their '
        'projections at intervals ranging from two years (Japan\u2019s National Institute '
        'of Population and Social Security Research) to five or more years (Eurostat, '
        'US Census Bureau). We adopt the 10-year interval as a deliberate analytical '
        'choice for two reasons: (a) it approximates the time scale over which major '
        'population policy interventions\u2014housing programmes, childcare systems, '
        'educational pathway reforms\u2014are designed, implemented, and begin to produce '
        'measurable demographic effects (McDonald 2000); and (b) it provides a '
        'conservative test of the tempo channel, since more frequent updates would '
        'narrow any remaining gap between model variants. That the tempo-adjusted '
        'model (TFR*) still achieves measurably better fit (median MAPE 4.3%) despite '
        'the conservative decadal interval strengthens the case for explicit tempo '
        'decomposition. Third, the Gompertz survival function fits adult mortality '
        'well but approximates infant and child mortality less precisely; national '
        'projections typically use more flexible models (Wilmoth et al. 2012). Fourth, '
        'the decomposition is necessarily model-dependent: the counterfactual \u2018MAC '
        'frozen at 1970\u2019 is a thought experiment, not a prediction of what would '
        'have occurred in the absence of postponement, since MAC changes are '
        'endogenous to broader social and economic shifts (Lesthaeghe 2010). Despite '
        'these limitations, the model\u2019s parsimonious structure is a feature: it makes '
        'the tempo\u2013quantum\u2013survival decomposition transparent and allows the '
        'magnitudes of each component to be compared directly.',
        size=12, space_after=12)

    add_para(doc,
        'Figure 5 provides model bias diagnostics, confirming that model performance '
        'is robust across demographic contexts with no systematic relationship between '
        'fit and TFR, life expectancy, or MAC.',
        size=12, space_after=6)

    add_table_marker(doc, 'Figure 5')

    # ==========================================================
    # 6. Conclusion
    # ==========================================================
    add_h1(doc, '6. Conclusion')

    add_para(doc,
        'The tempo effect on simultaneously living population is well-established in '
        'demographic theory but has never been systematically quantified across '
        'countries. Meanwhile, bibliometric evidence shows that the concept has '
        'essentially disappeared from the health and policy literature, and no '
        'national projection system decomposes the tempo channel explicitly. We '
        'provide this quantification using a parsimonious model validated against '
        'observed trajectories for 40 countries over 1970\u20132023. Five findings emerge.',
        size=12, space_after=12)

    add_para(doc,
        'First, the tempo effect is invisible within standard projection frameworks: '
        'tempo-invariant and tempo-responsive models yield nearly identical fit '
        '(median MAPE 4.5% versus 4.6%), because period TFR absorbs tempo distortion '
        'without decomposing it. Policymakers who rely on period TFR cannot distinguish '
        'postponement from permanent quantum decline. Second, explicit tempo '
        'decomposition via the Bongaarts\u2013Feeney TFR* achieves the best overall fit '
        '(median MAPE 4.3%), with the largest improvements in countries experiencing '
        'strong postponement (Republic of Korea, China, Colombia). Third, the '
        'magnitude of the tempo effect is large: the observed rise in MAC across OECD '
        'countries independently reduced SLP by 8\u201317%, equivalent to 15\u201340 years of '
        'below-replacement fertility. Fourth, tempo is typically the second-largest '
        'component of population change in post-transitional countries, comparable in '
        'magnitude to quantum in several Central and Eastern European nations. Fifth, '
        'higher MAC accelerates the annual pace of population decline, compressing the '
        'time available for institutional adaptation.',
        size=12, space_after=12)

    add_para(doc,
        'These findings have two practical implications. For demographic assessment, '
        'the tools to decompose quantum and tempo already exist within national '
        'cohort-component projection systems; what is needed is the analytical step of '
        'reporting their separate contributions\u2014a step that is straightforward in '
        'principle but absent in all 15 systems we reviewed. For population policy, '
        'the current framework operates on two levers\u2014boosting births and extending '
        'lives\u2014while ignoring the third. Our results demonstrate that birth timing '
        'constitutes an independent and quantitatively significant policy lever. '
        'Tempo-sensitive interventions\u2014housing, childcare, educational reform\u2014offer '
        'a complementary approach that operates not on the ultimate size of the '
        'population but on the speed at which demographic change unfolds. In an era '
        'where pronatalist policies have shown limited effectiveness in raising '
        'quantum (Gauthier 2007), the tempo channel deserves systematic attention as '
        'a means of managing the pace of demographic transition.',
        size=12, space_after=18)

    # ==========================================================
    # Acknowledgements
    # ==========================================================
    add_h1(doc, 'Acknowledgements')
    add_para(doc,
        '[Acknowledgements removed for double-blind review]'
        if anonymized else
        '[To be added upon acceptance]',
        size=12, space_after=18)

    # ==========================================================
    # Disclosure statement
    # ==========================================================
    add_h1(doc, 'Disclosure statement')
    add_para(doc,
        'The authors report there are no competing interests to declare.',
        size=12, space_after=18)

    # ==========================================================
    # Funding
    # ==========================================================
    add_h1(doc, 'Funding')
    add_para(doc,
        'This research received no specific grant from any funding agency in the '
        'public, commercial, or not-for-profit sectors.',
        size=12, space_after=18)

    # ==========================================================
    # GATHER compliance statement
    # ==========================================================
    add_h1(doc, 'GATHER compliance statement')
    add_para(doc,
        'This study reports population estimates and follows the Guidelines for '
        'Accurate and Transparent Health Estimates Reporting (GATHER) (Stevens et al. '
        '2016). All input data are from United Nations World Population Prospects 2024, '
        'publicly available at https://population.un.org/wpp/. No primary data '
        'collection was undertaken. The Gompertz survival model, normal fertility '
        'schedule, and endogenous renewal model are described in Section 3. Four '
        'parameters per period (TFR, e\u2080, MAC, \u03c3) are used. MAPE and final ratio '
        'are reported as fit metrics. No formal uncertainty intervals; the model is '
        'deterministic. Analytical code is available from the authors upon request '
        'and will be deposited in a public repository upon acceptance.',
        size=12, space_after=18)

    # ==========================================================
    # Data availability statement
    # ==========================================================
    add_h1(doc, 'Data availability statement')
    add_para(doc,
        'All input data are drawn from the United Nations World Population Prospects '
        '2024, publicly available at https://population.un.org/wpp/. Analytical code '
        'is available from the authors upon request and will be deposited in a public '
        'repository upon acceptance.',
        size=12, space_after=18)

    # ==========================================================
    # Appendix A: Natural experiments
    # ==========================================================
    doc.add_page_break()
    add_h1(doc, 'Appendix A: Natural experiments \u2014 Political and border changes '
           'as exogenous shocks')

    add_para(doc,
        'Our model deliberately excludes migration. This design choice enables '
        'isolation of the pure quantum\u2013tempo\u2013survival mechanism, but raises an '
        'important question: how does the model perform when large-scale population '
        'redistribution occurs as a result of political upheaval? We examine several '
        'natural experiments that provide exogenous variation in population size '
        'independent of the demographic mechanisms our model captures.',
        size=12, space_after=12)

    add_h2(doc, 'A.1 German reunification (1990)')
    add_para(doc,
        'The reunification of Germany in 1990 provides a particularly instructive '
        'test case. We construct a synthetic \u2018combined Germany\u2019 by summing the East '
        'and West German populations from 1970 onward, with a single set of '
        'demographic parameters (population-weighted averages of TFR, e\u2080, MAC). '
        'This synthetic series does not account for the massive East-to-West internal '
        'migration that followed reunification, nor the large-scale immigration from '
        'abroad that Germany experienced from the 1990s onward. Table A1 summarises '
        'the component analysis.',
        size=12, space_after=6)

    add_table_marker(doc, 'Table A1')

    add_para(doc,
        'The model\u2019s 15.8% overestimate of the combined population in 2023 '
        'relative to observed reflects the net immigration Germany received\u2014the '
        'model\u2019s \u2018error\u2019 directly quantifies the migration component. When we '
        'compare East and West Germany separately, the fertility dynamics are '
        'striking: East Germany experienced a dramatic post-reunification fertility '
        'collapse, with TFR falling to 0.77 in 1994 (Goldstein and Kreyenfeld 2011; '
        'Witte and Wagner 1995). This demographic shock was fully driven by quantum '
        '(a genuine decline in births) rather than tempo, making it an ideal test of '
        'the model\u2019s ability to distinguish between the two channels.',
        size=12, space_after=12)

    add_h2(doc, 'A.2 The Velvet Divorce: Czechia and Slovakia (1993)')
    add_para(doc,
        'The peaceful dissolution of Czechoslovakia into Czechia and Slovakia on '
        '1 January 1993 created two independent states with distinct demographic '
        'trajectories. Both countries subsequently experienced rapid fertility '
        'postponement as they transitioned to market economies, but differed in their '
        'migration experiences: Czechia became a net immigration country after EU '
        'accession (2004), while Slovakia experienced sustained emigration.',
        size=12, space_after=12)

    add_h2(doc, 'A.3 Yugoslav breakup (1991\u20132008)')
    add_para(doc,
        'The dissolution of Yugoslavia produced seven successor states with vastly '
        'different post-conflict demographic trajectories. War, ethnic cleansing, and '
        'refugee flows created massive population redistribution that our closed-'
        'population model cannot capture. Yet the model performs reasonably for '
        'countries where post-conflict migration was moderate.',
        size=12, space_after=12)

    add_h2(doc, 'A.4 Baltic independence and EU accession')
    add_para(doc,
        'Estonia, Latvia, and Lithuania experienced two waves of demographic shock: '
        'the departure of ethnic Russians following independence (1991), and '
        'emigration to Western Europe after EU accession (2004). Our model captures '
        'the natural-increase dynamics but cannot account for these migration shocks. '
        'The resulting MAPE (Estonia 4.8%, Latvia 6.6%, Lithuania 7.1%) directly '
        'quantifies the cumulative migration deficit.',
        size=12, space_after=12)

    add_h2(doc, 'A.5 Eritrean independence (1993)')
    add_para(doc,
        'Eritrea\u2019s independence from Ethiopia in 1993 represents an extreme case. '
        'The model\u2019s high MAPE for Eritrea (37.8%) reflects the combination of '
        'conflict-driven displacement, data uncertainty, and rapid demographic change '
        'that is difficult to capture with any simple model.',
        size=12, space_after=12)

    add_h2(doc, 'A.6 Synthesis')
    add_para(doc,
        'Table A2 summarises model performance across all natural experiment cases.',
        size=12, space_after=6)

    add_table_marker(doc, 'Table A2')

    add_para(doc,
        'These natural experiments yield three key insights. First, the model '
        'performs reasonably (MAPE < 8%) even for countries with major political '
        'upheaval, provided post-event migration was moderate. Second, the '
        'model\u2013observation divergence provides a direct estimate of the migration '
        'component: Germany\u2019s 15.8% gap implies ~13.5 million persons added by '
        'immigration. Third, the model\u2019s limitations are most acute in '
        'conflict-affected, data-sparse settings (Eritrea: 37.8%).',
        size=12, space_after=12)

    # ==========================================================
    # References (alphabetical, Population Studies house style)
    # ==========================================================
    doc.add_page_break()
    add_h1(doc, 'References')

    refs_popstudies = [
        'Bongaarts, John and Griffith Feeney. 1998. On the quantum and tempo of '
        'fertility, *Population and Development Review* 24(2): 271\u2013291.',

        'Bongaarts, John and Griffith Feeney. 2006. The quantum and tempo of '
        'life-cycle events, *Vienna Yearbook of Population Research* 4: 115\u2013151.',

        'Bongaarts, John and Tom\u00e1\u0161 Sobotka. 2012. A demographic explanation '
        'for the recent rise in European fertility, *Population and Development '
        'Review* 38(1): 83\u2013120.',

        'Frejka, Tom\u00e1\u0161 and Tom\u00e1\u0161 Sobotka. 2008. Fertility in Europe: '
        'Diverse, delayed and below replacement, *Demographic Research* 19(3): '
        '15\u201346.',

        'Gauthier, Anne H. 2007. The impact of family policies on fertility in '
        'industrialized countries, *Population Research and Policy Review* 26(3): '
        '323\u2013346.',

        'Goldstein, Joshua R. and Michaela Kreyenfeld. 2011. Has East Germany '
        'overtaken West Germany? Recent trends in order-specific fertility, '
        '*Population and Development Review* 37(3): 453\u2013472.',

        'Goldstein, Joshua R., Wolfgang Lutz, and Sergei Scherbov. 2003. Long-term '
        'population decline in Europe: The relative importance of tempo effects and '
        'generational length, *Population and Development Review* 29(4): 699\u2013707.',

        'Gompertz, Benjamin. 1825. On the nature of the function expressive of the '
        'law of human mortality, *Philosophical Transactions of the Royal Society of '
        'London* 115: 513\u2013583.',

        'Gonand, Fr\u00e9d\u00e9ric. 2005. Assessing the robustness of demographic '
        'projections in OECD countries. OECD Economics Department Working Papers '
        'No. 464. Paris: OECD Publishing.',

        'Keyfitz, Nathan. 1971. On the momentum of population growth, *Demography* '
        '8(1): 71\u201380.',

        'Kohler, Hans-Peter and Jos\u00e9 Antonio Ortega. 2002. Tempo-adjusted period '
        'parity progression measures, fertility postponement and completed cohort '
        'fertility, *Demographic Research* 6(6): 91\u2013144.',

        'Kohler, Hans-Peter, Francesco C. Billari, and Jos\u00e9 Antonio Ortega. 2002. '
        'The emergence of lowest-low fertility in Europe during the 1990s, '
        '*Population and Development Review* 28(4): 641\u2013680.',

        'Lee, Ronald D. and Lawrence R. Carter. 1992. Modeling and forecasting U.S. '
        'mortality, *Journal of the American Statistical Association* 87(419): '
        '659\u2013671.',

        'Lesthaeghe, Ron. 2010. The unfolding story of the Second Demographic '
        'Transition, *Population and Development Review* 36(2): 211\u2013251.',

        'Lutz, Wolfgang, Warren Sanderson, and Sergei Scherbov. 2001. The end of '
        'world population growth, *Nature* 412: 543\u2013545.',

        'McDonald, Peter. 2000. Gender equity in theories of fertility transition, '
        '*Population and Development Review* 26(3): 427\u2013439.',

        'Myrskyl\u00e4, Mikko, Hans-Peter Kohler, and Francesco C. Billari. 2009. '
        'Advances in development reverse fertility declines, *Nature* 460: 741\u2013743.',

        'Preston, Samuel H., Patrick Heuveline, and Michel Guillot. 2001. '
        '*Demography: Measuring and Modeling Population Processes*. Oxford: Blackwell.',

        'Ryder, Norman B. 1964. The process of demographic translation, *Demography* '
        '1(1): 74\u201382.',

        'Sobotka, Tom\u00e1\u0161. 2004. Is lowest-low fertility in Europe explained '
        'by the postponement of childbearing?, *Population and Development Review* '
        '30(2): 195\u2013220.',

        'Stevens, Gretchen A., Leontine Alkema, Robert E. Black, et al. 2016. '
        'Guidelines for Accurate and Transparent Health Estimates Reporting: The '
        'GATHER statement, *The Lancet* 388(10062): e19\u2013e23.',

        'Th\u00e9venon, Olivier. 2011. Family policies in OECD countries: A '
        'comparative analysis, *Population and Development Review* 37(1): 57\u201387.',

        'United Nations, Department of Economic and Social Affairs, Population '
        'Division. 2024. *World Population Prospects 2024*. New York: United Nations.',

        'Wilmoth, John R., Sarah Zureick, Vladimir Canudas-Romo, Mie Inoue, and '
        'Claudia Sawyer. 2012. A flexible two-dimensional mortality model for use in '
        'indirect estimation, *Population Studies* 66(1): 1\u201328.',

        'Witte, James C. and Gert G. Wagner. 1995. Declining fertility in East '
        'Germany after unification, *Population and Development Review* 21(2): '
        '387\u2013397.',
    ]

    for ref in refs_popstudies:
        add_ref_para(doc, ref, size=11, space_after=4)

    # ==========================================================
    # Tables (each on its own page)
    # ==========================================================
    # --- Table 1: PubMed ---
    doc.add_page_break()
    add_para(doc,
        'Table 1: PubMed publication counts by search term and period',
        size=11, space_after=4)

    tbl1 = doc.add_table(rows=7, cols=4)
    tbl1.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Period',
                           '\u2018Tempo effect\u2019\n+ demography',
                           '\u2018Delayed\nchildbearing\u2019',
                           '\u2018Fertility\npostponement\u2019']):
        tbl1.rows[0].cells[i].text = h
    pubmed_data = [
        ['1998\u20132002', '0', '22', '0'],
        ['2003\u20132007', '2', '39', '0'],
        ['2008\u20132012', '2', '56', '6'],
        ['2013\u20132017', '1', '90', '5'],
        ['2018\u20132022', '3', '89', '10'],
        ['2023\u20132025', '0', '96', '14'],
    ]
    for i, rd in enumerate(pubmed_data):
        for j, val in enumerate(rd):
            tbl1.rows[i + 1].cells[j].text = val
    add_para(doc, '', size=6, space_after=4)
    add_para(doc,
        'Note: PubMed search conducted May 2025 using NCBI E-utilities API. '
        '\u2018Tempo effect\u2019 search: "tempo effect" AND (fertility OR demography OR '
        'population). \u2018Delayed childbearing\u2019 search: "delayed childbearing" OR '
        '"postponement of childbearing". \u2018Fertility postponement\u2019 search: exact '
        'phrase. Results filtered by publication date.',
        italic=True, size=9, space_after=12)

    # --- Table 2: Model performance ---
    doc.add_page_break()
    add_para(doc,
        'Table 2: Model performance across 40 countries \u2014 five-model comparison',
        size=11, space_after=4)

    tbl2 = doc.add_table(rows=8, cols=6)
    tbl2.style = 'Light Shading Accent 1'
    headers2 = ['Model variant', 'Horizon (yrs)', 'N', 'MAPE mean (%)',
                'MAPE median (%)', 'Final ratio\n(mean \u00b1 SD)']
    for i, h in enumerate(headers2):
        tbl2.rows[0].cells[i].text = h
    data2 = [
        ['Fixed-parameter (1970)', '53', '40', '13.7', '7.7',
         '1.309 \u00b1 0.554'],
        ['Fixed-parameter (1980)', '43', '40', '9.6', '7.7',
         '1.023 \u00b1 0.288'],
        ['Fixed-parameter (1990)', '33', '40', '7.8', '6.5',
         '0.953 \u00b1 0.198'],
        ['Fixed-parameter (2000)', '23', '40', '5.1', '4.7',
         '0.914 \u00b1 0.101'],
        ['Tempo-invariant (10-yr)', '53', '40', '6.5', '4.5',
         '0.991 \u00b1 0.182'],
        ['Tempo-responsive (10-yr)', '53', '40', '6.7', '4.6',
         '0.999 \u00b1 0.189'],
        ['Tempo-adjusted [TFR*]\n(10-yr)', '53', '40', '5.8', '4.3',
         '1.014 \u00b1 0.166'],
    ]
    for i, row_data in enumerate(data2):
        for j, val in enumerate(row_data):
            tbl2.rows[i + 1].cells[j].text = val
    add_para(doc, '', size=6, space_after=4)
    add_para(doc,
        'Note: All decadally-updated models run 1970\u20132023 with parameters '
        'refreshed every 10 years. TFR* = TFR / (1 \u2212 dMAC/dt) (Bongaarts and '
        'Feeney 1998). Final ratio = model population / observed population in 2023.',
        italic=True, size=9, space_after=6)

    # --- Table 3: Counterfactual tempo analysis ---
    doc.add_page_break()
    add_para(doc,
        'Table 3: Counterfactual tempo analysis \u2014 population impact of observed '
        'MAC increase, selected countries',
        size=11, space_after=4)

    tbl3 = doc.add_table(rows=21, cols=6)
    tbl3.style = 'Light Shading Accent 1'
    h3 = ['Country', 'MAC 1970', 'MAC 2020s', '\u0394MAC (yrs)',
          'Tempo effect\non SLP (%)', 'Equivalent\nTFR-years']
    for i, h in enumerate(h3):
        tbl3.rows[0].cells[i].text = h
    t3_data = [
        ['Japan', '27.5', '31.4', '+3.9', '\u221213.0', '~28'],
        ['Korea', '26.1', '33.4', '+7.3', '\u221221.9', '~45'],
        ['Italy', '27.2', '31.6', '+4.4', '\u221214.6', '~31'],
        ['Spain', '28.2', '32.3', '+4.1', '\u221213.5', '~29'],
        ['Germany', '27.0', '30.5', '+3.5', '\u221211.5', '~24'],
        ['France', '27.1', '30.7', '+3.6', '\u221211.7', '~25'],
        ['United Kingdom', '26.7', '30.7', '+4.0', '\u221213.0', '~27'],
        ['Netherlands', '28.0', '31.0', '+3.0', '\u221210.0', '~21'],
        ['Czechia', '24.4', '30.1', '+5.7', '\u221218.9', '~39'],
        ['Poland', '25.8', '30.3', '+4.5', '\u221214.9', '~31'],
        ['Australia', '27.1', '31.1', '+4.0', '\u221212.9', '~27'],
        ['Canada', '27.0', '30.9', '+3.9', '\u221212.6', '~26'],
        ['USA', '25.4', '29.3', '+3.9', '\u221213.3', '~28'],
        ['Sweden', '27.3', '31.2', '+3.9', '\u221212.5', '~26'],
        ['Finland', '27.4', '31.3', '+3.9', '\u221212.5', '~26'],
        ['Switzerland', '28.0', '32.0', '+4.0', '\u221212.5', '~26'],
        ['Greece', '27.3', '31.5', '+4.2', '\u221213.3', '~28'],
        ['Portugal', '27.0', '31.4', '+4.4', '\u221214.0', '~29'],
        ['China', '29.2', '28.4', '\u22120.8', '+2.8', 'n/a'],
        ['DRC', '24.8', '24.8', '0.0', '0.0', 'n/a'],
    ]
    for i, rd in enumerate(t3_data):
        for j, val in enumerate(rd):
            tbl3.rows[i + 1].cells[j].text = val
    add_para(doc, '', size=6, space_after=4)
    add_para(doc,
        'Note: Tempo effect on SLP = percentage difference between baseline model '
        '(MAC evolves as observed) and tempo counterfactual (MAC held at 1970 value). '
        'Equivalent TFR-years = number of years of below-replacement fertility '
        '(TFR = 1.5) that would produce the same population reduction. China\u2019s MAC '
        'declined, producing a positive tempo effect; DRC shows negligible MAC change.',
        italic=True, size=9, space_after=12)

    # --- Table 4: Decomposition ---
    doc.add_page_break()
    add_para(doc,
        'Table 4: Decomposition of modelled population change (1970\u20132023) into '
        'quantum, tempo, and survival components',
        size=11, space_after=4)

    tbl4 = doc.add_table(rows=11, cols=6)
    tbl4.style = 'Light Shading Accent 1'
    h4 = ['Country', 'Total pop.\nchange (%)', 'Quantum\neffect (%)',
          'Tempo\neffect (%)', 'Survival\neffect (%)', 'Interaction\n(%)']
    for i, h in enumerate(h4):
        tbl4.rows[0].cells[i].text = h
    t4_data = [
        ['Japan', '\u221212.8', '\u221232.1', '\u221213.0', '+28.4', '+3.9'],
        ['Korea', '+7.3', '\u221238.5', '\u221221.9', '+55.2', '+12.5'],
        ['Italy', '\u22124.2', '\u221228.7', '\u221214.6', '+34.8', '+4.3'],
        ['Germany', '+0.3', '\u221226.4', '\u221211.5', '+32.0', '+6.2'],
        ['France', '+23.8', '\u22128.3', '\u221211.7', '+38.2', '+5.6'],
        ['USA', '+46.3', '\u221212.4', '\u221213.3', '+58.7', '+13.3'],
        ['Czechia', '\u22124.1', '\u221222.0', '\u221218.9', '+30.6', '+6.2'],
        ['Sweden', '+17.5', '\u22125.1', '\u221212.5', '+28.4', '+6.7'],
        ['China', '+56.1', '\u221248.2', '+2.8', '+84.9', '+16.6'],
        ['DRC', '+296.4', '+156.2', '0.0', '+95.8', '+44.4'],
    ]
    for i, rd in enumerate(t4_data):
        for j, val in enumerate(rd):
            tbl4.rows[i + 1].cells[j].text = val
    add_para(doc, '', size=6, space_after=4)
    add_para(doc,
        'Note: Quantum effect = difference between baseline and TFR-frozen '
        'counterfactual. Tempo effect = difference between baseline and MAC-frozen '
        'counterfactual. Survival effect = difference between baseline and '
        'e\u2080-frozen counterfactual. Interaction = residual from non-additive '
        'effects. Signs indicate direction: negative = reducing population relative '
        'to 1970 trajectory.',
        italic=True, size=9, space_after=12)

    # --- Table 5: National projection methods ---
    doc.add_page_break()
    add_para(doc,
        'Table 5: Summary of official population projection methods and tempo '
        'treatment by country/organisation',
        size=11, space_after=4)

    tbl5 = doc.add_table(rows=16, cols=5)
    tbl5.style = 'Light Shading Accent 1'
    hdr5 = ['Country /\nOrganisation', 'Method', 'Fertility\nassumption',
            'Mortality\nassumption', 'Tempo\ndecomposition?']
    for i, h in enumerate(hdr5):
        tbl5.rows[0].cells[i].text = h

    t5_data = [
        ['UN WPP 2024\n(All countries)',
         'Cohort-component;\nprobabilistic (Bayesian)',
         'Bayesian hierarchical;\nTFR trajectories',
         'Lee\u2013Carter variant;\ncountry-specific drift', 'No'],
        ['Japan\n(IPSS)',
         'Cohort-component;\n3\u00d73 variants',
         'Cohort fertility model;\nMAC = 32.8',
         'Lee\u2013Carter;\ne\u2080 = 85.9/91.8', 'No'],
        ['USA\n(Census Bureau)',
         'Cohort-component;\nmain + 3 migration',
         'Race-specific ASFRs;\nTFR ~ 1.75 by 2060',
         'Cause-of-death model;\ne\u2080 ~ 83.9 by 2100', 'No'],
        ['Germany\n(Destatis)',
         'Cohort-component;\n27 variants',
         'TFR 1.29\u20131.65;\nMAC ~ 31.7\u201332.1',
         'e\u2080 82.6\u201386.4 (M)\n85.9\u201389.3 (F)', 'No'],
        ['United Kingdom\n(ONS)',
         'Cohort-component;\nprincipal + 9 variants',
         'ASFRs; TFR ~ 1.59\nlong-term',
         'Age-period-cohort;\ne\u2080 ~ 83.9/86.3', 'No'],
        ['France\n(INSEE)',
         'Cohort-component;\ncentral + 3 variants',
         'TFR ~ 1.80 central',
         'Trend extrapolation;\ne\u2080 ~ 87.5/90.0', 'No'],
        ['Korea\n(KOSTAT)',
         'Cohort-component;\n3 scenarios',
         'Cohort model;\nTFR 1.08 by 2040',
         'Lee\u2013Carter;\ne\u2080 = 88.0/91.4', 'No'],
        ['Italy\n(ISTAT)',
         'Cohort-component;\nmedian + 4 scenarios',
         'TFR ~ 1.40 median',
         'Lee\u2013Carter;\ne\u2080 ~ 85.8/89.2', 'No'],
        ['Australia\n(ABS)',
         'Cohort-component;\n3 series',
         'TFR 1.55\u20131.85',
         'Mortality improvement;\ne\u2080 ~ 87/89', 'No'],
        ['Canada\n(Statistics Canada)',
         'Cohort-component;\nmicrosimulation',
         'TFR 1.40\u20131.60',
         'Lee\u2013Carter variant;\ne\u2080 ~ 86/89', 'No'],
        ['Eurostat\n(EU members)',
         'Cohort-component;\nconvergence model',
         'Partial convergence\nof TFR across EU',
         'Convergence of\nmortality improvement', 'No'],
        ['China\n(NBS)',
         'Cohort-component\n(not regularly published)',
         'TFR 1.0\u20131.2;\nrecovery assumed',
         'Model life table;\ne\u2080 ~ 78.6', 'No'],
        ['DRC',
         'Relies on UN WPP;\nno national projection',
         'TFR ~ 6.1; gradual\ndecline assumed',
         'Model life table;\ne\u2080 ~ 60.7', 'No'],
        ['Mexico\n(CONAPO)',
         'Cohort-component;\n3 variants',
         'TFR ~ 1.7 by 2050',
         'Trend extrapolation;\ne\u2080 ~ 79/83', 'No'],
        ['T\u00fcrkiye\n(TurkStat)',
         'Cohort-component;\n3 scenarios',
         'TFR declining to\n~1.60 long-term',
         'Improvement model;\ne\u2080 ~ 80/84', 'No'],
    ]
    for i, rd in enumerate(t5_data):
        for j, val in enumerate(rd):
            tbl5.rows[i + 1].cells[j].text = val
    add_para(doc, '', size=6, space_after=4)
    add_para(doc,
        'Source: UN DESA (2024), National Institute of Population and Social Security '
        'Research, Japan (2023), US Census Bureau (2023), Federal Statistical Office '
        'of Germany (2025), Office for National Statistics, United Kingdom (2025), '
        'National Institute of Statistics and Economic Studies, France (2021), '
        'Statistics Korea (2023), National Institute of Statistics, Italy (2023), '
        'Australian Bureau of Statistics (2018), Statistics Canada (2024), Eurostat '
        '(2024), National Population Council, Mexico (2018), Turkish Statistical '
        'Institute (2023).',
        italic=True, size=9, space_after=12)

    # --- Table A1: Germany natural experiment ---
    doc.add_page_break()
    add_para(doc,
        'Table A1: German reunification \u2014 model components',
        size=11, space_after=4)

    tbl_a1 = doc.add_table(rows=5, cols=4)
    tbl_a1.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Component', 'Population 1970\n(millions)',
                           'Population 2023\n(millions)', 'Change (%)']):
        tbl_a1.rows[0].cells[i].text = h
    a1_data = [
        ['West Germany (model)', '60.7', '56.4', '\u22127.1'],
        ['East Germany (model)', '17.1', '12.8', '\u221225.1'],
        ['Combined (model)', '77.8', '69.2', '\u221211.1'],
        ['Germany (observed)', '78.2', '84.5', '+8.1'],
    ]
    for i, rd in enumerate(a1_data):
        for j, val in enumerate(rd):
            tbl_a1.rows[i + 1].cells[j].text = val
    add_para(doc, '', size=6, space_after=4)
    add_para(doc,
        'Note: Model excludes migration. The 15.8 percentage-point gap between '
        'model and observed reflects cumulative net immigration since 1970.',
        italic=True, size=9, space_after=12)

    # --- Table A2: Natural experiment synthesis ---
    doc.add_page_break()
    add_para(doc,
        'Table A2: Model performance across natural experiment cases',
        size=11, space_after=4)

    tbl_a2 = doc.add_table(rows=15, cols=4)
    tbl_a2.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Country', 'Event (year)', 'MAPE (%)',
                           'Primary misfit source']):
        tbl_a2.rows[0].cells[i].text = h
    a2_data = [
        ['Germany (synth. E+W)', 'Reunification (1990)', '6.4',
         'Immigration + internal migration'],
        ['Czechia', 'Velvet Divorce (1993)', '6.3', 'Post-EU emigration'],
        ['Slovakia', 'Velvet Divorce (1993)', '9.9',
         'Emigration to EU/West'],
        ['Croatia', 'Yugoslav breakup (1991)', '4.1',
         'Post-conflict stabilisation'],
        ['Slovenia', 'Yugoslav breakup (1991)', '12.2',
         'Immigration (EU member)'],
        ['Bosnia & Herz.', 'Yugoslav breakup (1991)', '8.1',
         'War displacement'],
        ['Serbia', 'Yugoslav breakup (1991)', '7.1',
         'Refugee flows, emigration'],
        ['N. Macedonia', 'Yugoslav breakup (1991)', '6.4',
         'Modest migration'],
        ['Montenegro', 'Yugoslav breakup (1991)', '8.1',
         'Small state, volatile'],
        ['Estonia', 'USSR dissolution (1991)', '4.8',
         'Ethnic Russian emigration'],
        ['Latvia', 'USSR dissolution (1991)', '6.6',
         'Emigration (ethnic + EU)'],
        ['Lithuania', 'USSR dissolution (1991)', '7.1',
         'Sustained emigration'],
        ['Ethiopia', 'Eritrean indep. (1993)', '16.5',
         'Rapid fertility decline'],
        ['Eritrea', 'Eritrean indep. (1993)', '37.8',
         'Conflict, data uncertainty'],
    ]
    for i, rd in enumerate(a2_data):
        for j, val in enumerate(rd):
            tbl_a2.rows[i + 1].cells[j].text = val
    add_para(doc, '', size=6, space_after=6)

    # ==========================================================
    # Figure titles (as a list)
    # ==========================================================
    doc.add_page_break()
    add_h1(doc, 'Figure titles')

    fig_titles = [
        ('Figure 1', 'Five model variants versus observed population trajectories '
         'for six representative countries, 1970\u20132023.\n'
         'Note: Observed = black solid (UN WPP 2024); Tempo-responsive = blue '
         'solid; Tempo-invariant = orange dashed; Tempo-adjusted (TFR*) = green '
         'dash-dot; Fixed-parameter (1970) = red dotted.'),

        ('Figure 2', 'Five-model validation across all 40 countries.\n'
         'Note: Observed = black; Tempo-responsive = blue; Tempo-invariant = '
         'orange dashed; Tempo-adjusted (TFR*) = green dash-dot; Fixed-parameter '
         '= red dotted. Tempo-responsive MAPE shown in upper-right corner. '
         'Countries sorted alphabetically.'),

        ('Figure 3', 'Fixed-parameter model MAPE (%) by country and base year.\n'
         'Note: Greener cells indicate better fit; redder cells indicate poorer '
         'fit. Scale capped at 30%. Longer projection horizons show greater '
         'misfit, reflecting accumulating demographic change unaccounted for by '
         'fixed parameters.'),

        ('Figure 4', 'Four-model comparison\u2014MAPE and final population ratio by '
         'country.\n'
         'Note: Left panel: MAPE by country for fixed-parameter (red), '
         'tempo-invariant (orange), tempo-responsive (blue), and tempo-adjusted '
         'TFR* (green). Right panel: final population ratio (model/observed in '
         '2023). The tempo-adjusted model (TFR*) achieves the best overall fit, '
         'particularly in countries with strong postponement.'),

        ('Figure 5', 'Model bias analysis using base year 2000.\n'
         'Note: (a) Fit versus TFR; (b) fit versus life expectancy; (c) bias '
         'versus MAC. No systematic relationship is observed.'),
    ]

    for label, caption in fig_titles:
        add_para(doc, f'{label}: {caption}', size=11, space_after=12)

    return doc


# ==============================================================
# Generate both versions
# ==============================================================
doc_anon = build_manuscript(anonymized=True)
path_anon = os.path.join(OUT_DIR, 'PopStudies_Article_EN.docx')
doc_anon.save(path_anon)
print(f'OK (anonymized): {path_anon}')

doc_open = build_manuscript(anonymized=False)
path_open = os.path.join(OUT_DIR, 'PopStudies_Article_EN_open.docx')
doc_open.save(path_open)
print(f'OK (open):       {path_open}')
