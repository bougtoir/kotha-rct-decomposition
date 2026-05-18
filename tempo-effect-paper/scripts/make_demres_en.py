"""Create Demographic Research submission — English manuscript."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'figures')
OUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)


def add_para(doc, text, bold=False, italic=False, size=12, align=None,
             space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_figure(doc, path, title, note=None, width=6.0):
    """Insert figure with title above and optional note below, per DR guidelines."""
    # Title above figure (regular text, not part of graphic)
    t = doc.add_paragraph()
    r = t.add_run(title)
    r.font.size = Pt(11)
    r.bold = True
    t.paragraph_format.space_after = Pt(4)
    # Image
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        p.paragraph_format.space_after = Pt(2)
    # Note/Source below figure
    if note:
        cap = doc.add_paragraph()
        r2 = cap.add_run(note)
        r2.font.size = Pt(9)
        r2.italic = True
        cap.paragraph_format.space_after = Pt(12)


# ==============================================================
# Document setup — 12pt, double-spaced, 1-inch margins, no page numbers
# ==============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# ==============================================================
# Title page
# ==============================================================
add_para(doc, "Research Article", bold=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc,
    "The Forgotten Tempo Effect: Delayed Childbearing, Simultaneously Living "
    "Population, and the Pace of Social Adaptation Across OECD Countries",
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para(doc, "[Author names removed for review]",
         italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "Word count (excluding notes and references): approximately 6,500",
         italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

# ==============================================================
# Structured Abstract (max 250 words per DR guidelines)
# ==============================================================
add_heading_styled(doc, "Abstract", level=2)

add_para(doc, "BACKGROUND", bold=True, size=12, space_after=2)
add_para(doc,
    "Population projections and pronatalist policies overwhelmingly emphasise the quantum of "
    "fertility\u2014how many children are born\u2014while neglecting the tempo effect: the independent "
    "influence of birth timing on the number of generations simultaneously alive.",
    size=12, space_after=6)

add_para(doc, "OBJECTIVE", bold=True, size=12, space_after=2)
add_para(doc,
    "We ask three questions: (1) Can a parsimonious model built on the tempo\u2013quantum\u2013survival "
    "triad reproduce observed population trajectories? (2) How large is the tempo effect relative "
    "to quantum across diverse demographic contexts? (3) What are the implications for the pace "
    "at which societies must adapt their institutions?",
    size=12, space_after=6)

add_para(doc, "METHODS", bold=True, size=12, space_after=2)
add_para(doc,
    "We construct a discrete-time endogenous renewal model coupling normal-distributed "
    "age-specific fertility (centred on mean age at childbearing, MAC) with Gompertz parametric "
    "survival calibrated to life expectancy. The model is validated against United Nations "
    "World Population Prospects 2024 data for 40 countries (38 OECD member states plus China "
    "and the DRC) over 1970\u20132023.",
    size=12, space_after=6)

add_para(doc, "RESULTS", bold=True, size=12, space_after=2)
add_para(doc,
    "The dynamic model (parameters updated decadally) achieves a median absolute percentage "
    "error (MAPE) of 4.6%. A 5-year increase in MAC reduces simultaneously living population "
    "by approximately one-sixth, independent of TFR.",
    size=12, space_after=6)

add_para(doc, "CONCLUSIONS", bold=True, size=12, space_after=2)
add_para(doc,
    "Tempo-sensitive interventions\u2014housing, childcare, educational reform\u2014offer an "
    "underutilised lever for managing the pace of demographic transition and social "
    "structural adaptation.",
    size=12, space_after=6)

add_para(doc, "CONTRIBUTION", bold=True, size=12, space_after=2)
add_para(doc,
    "We provide the first systematic 40-country empirical validation of the tempo effect on "
    "simultaneously living population, demonstrating its quantitative significance across "
    "diverse demographic contexts and introducing the concept of tempo as a policy lever "
    "controlling the speed of institutional adaptation.",
    size=12, space_after=12)

add_para(doc,
    "Keywords: tempo effect, simultaneously living population, mean age at childbearing, "
    "Gompertz survival, endogenous renewal model, OECD, demographic transition, "
    "population projection, fertility postponement, social adaptation",
    italic=True, size=10, space_after=18)
doc.add_page_break()

# ==============================================================
# 1. Introduction
# ==============================================================
add_heading_styled(doc, "1. The forgotten tempo effect", level=1)
add_para(doc,
    "The global conversation about population decline is dominated by a single indicator: the "
    "total fertility rate (TFR). When TFR falls below replacement level (approximately 2.1 "
    "children per woman), alarm bells sound. Yet this framing systematically neglects a second, "
    "independent demographic force that shapes how many people are simultaneously alive at any "
    "given moment: the timing of births.",
    size=12, space_after=12)

add_para(doc,
    "The concept of demographic translation\u2014the idea that shifts in the timing of vital events "
    "alter period rates independently of underlying quantum\u2014was introduced by Ryder (1964). "
    "Bongaarts and Feeney (1998) formalised the distinction between fertility quantum (number "
    "of births) and tempo (timing of births), showing that period TFR is mechanically depressed "
    "when women postpone childbearing\u2014even if completed cohort fertility remains unchanged. "
    "Their tempo-adjusted TFR revealed that observed period fertility in many countries was "
    "substantially lower than the underlying propensity to bear children, a finding confirmed "
    "and extended by Kohler and Ortega (2002) using tempo-adjusted parity progression measures. "
    "Sobotka (2004) demonstrated that much of Europe\u2019s lowest-low fertility (TFR \u2264 1.3), which "
    "had emerged across southern and eastern Europe (Kohler, Billari, and Ortega 2002), could "
    "be attributed to the postponement of childbearing rather than to a permanent decline in "
    "desired family size.",
    size=12, space_after=12)

add_para(doc,
    "Goldstein, Lutz, and Scherbov (2003) took this insight further, demonstrating for EU-15 "
    "countries that delayed childbearing reduces the number of generations alive at any moment, "
    "producing population decline independent of the number of children ever born per woman. "
    "Their decomposition showed that generational length changes accounted for a substantial "
    "fraction of projected population decline. This mechanism operates through the simultaneously "
    "living population (SLP)\u2014the stock of persons alive at a given calendar moment\u2014which is "
    "shaped by both quantum and tempo.",
    size=12, space_after=12)

add_para(doc,
    "Despite this foundational work, the tempo dimension has largely disappeared from "
    "contemporary policy discourse. These postponement trends are part of the broader Second "
    "Demographic Transition characterised by delayed family formation, value shifts, and "
    "sustained below-replacement fertility (Lesthaeghe 2010). A review of recent pronatalist "
    "policy packages across OECD countries\u2014from South Korea\u2019s record-setting 47 trillion won "
    "commitment to Japan\u2019s successive \u2018Plans for Measures Against the Declining Birthrate\u2019"
    "\u2014reveals an almost exclusive focus on increasing the number of births (Th\u00e9venon 2011). "
    "Only quantum receives systematic policy attention; the independent role of tempo in "
    "shaping population size is ignored.",
    size=12, space_after=12)

add_para(doc,
    "This paper revisits the tempo effect through a simple but empirically grounded demographic "
    "model and asks three questions: (1) Can a parsimonious model built on the tempo\u2013quantum\u2013"
    "survival triad reproduce observed population trajectories? (2) How large is the tempo "
    "effect relative to quantum across diverse demographic contexts? (3) What are the "
    "implications for the pace at which societies must adapt their institutions to demographic "
    "change?",
    size=12, space_after=12)

# ==============================================================
# 2. Model and Data
# ==============================================================
add_heading_styled(doc, "2. Model and data", level=1)

add_heading_styled(doc, "2.1 Endogenous renewal model", level=2)
add_para(doc,
    "We construct a discrete-time, single-sex population model in which the population vector "
    "P(t) = [P\u2080(t), P\u2081(t), ..., P\u2081\u2080\u2080(t)] evolves annually. The model belongs to the family "
    "of cohort-component methods that form the basis of virtually all national and international "
    "population projections (Preston, Heuveline, and Guillot 2001). At each time step:",
    size=12, space_after=8)

add_para(doc,
    "(a) Survival: Individuals at age x survive to age x+1 with probability s(x) derived from "
    "a Gompertz hazard function h(x) = a\u00b7exp(b\u00b7x), first proposed by Gompertz (1825) and "
    "remaining the most widely used parametric model for adult mortality. The survival function "
    "is S(x) = exp[\u2212(a/b)(exp(bx)\u22121)]. The parameter a is calibrated so that life expectancy "
    "at birth e\u2080 = \u222b\u2080^\u221e S(x)dx matches the observed value, with b fixed at 0.085.",
    size=12, space_after=8)

add_para(doc,
    "(b) Fertility: Births are generated endogenously. The age-specific fertility rate (ASFR) "
    "is modelled as a normal density centred on the mean age at childbearing (MAC) with standard "
    "deviation \u03c3, scaled to the total fertility rate (TFR). Births at time t equal "
    "\u03a3_{x=15}^{49} P_x(t) \u00b7 f \u00b7 ASFR(x), where f is the female population share.",
    size=12, space_after=8)

add_para(doc,
    "This minimal parameterisation requires only four inputs per period: TFR, life expectancy "
    "(e\u2080), MAC, and \u03c3. The model deliberately omits migration, which allows us to isolate the "
    "pure demographic mechanics of quantum, tempo, and survival. Population momentum\u2014the "
    "tendency for population to continue growing after fertility falls to replacement due to "
    "a young age structure (Keyfitz 1971)\u2014is captured endogenously through the age-structured "
    "dynamics.",
    size=12, space_after=12)

add_heading_styled(doc, "2.2 Data", level=2)
add_para(doc,
    "All input parameters and validation data are drawn from the United Nations World Population "
    "Prospects 2024 (United Nations 2024). We analyse 40 countries: all 38 OECD member states "
    "(as of 2024) plus China and the Democratic Republic of the Congo (DRC), chosen to span the "
    "full range of demographic transition stages. Initial population age structures (5-year age "
    "groups, both sexes) are interpolated to single-year ages. Demographic indicators\u2014TFR, "
    "e\u2080, and MAC\u2014are extracted for each calendar year from 1950 to 2023.",
    size=12, space_after=8)

add_para(doc,
    "Following the GATHER reporting guidelines (Stevens et al. 2016), we note: input data are "
    "publicly available from the UN Population Division; all model code and parameters are "
    "documented; the analytical approach is fully reproducible.",
    size=12, space_after=12)

add_heading_styled(doc, "2.3 Model variants", level=2)
add_para(doc,
    "We implement two variants:",
    size=12, space_after=8)

add_para(doc,
    "Static model: Parameters (TFR, e\u2080, MAC) are fixed at their base-year values and held "
    "constant throughout the projection horizon. We run four base years (1970, 1980, 1990, 2000) "
    "with forward projections to 2020\u20132023, yielding 160 country\u2013base-year combinations.",
    size=12, space_after=8)

add_para(doc,
    "Dynamic model: Parameters are updated every 10 years using observed UN WPP values (e.g., "
    "1970 parameters for 1970\u20131979, 1980 parameters for 1980\u20131989, etc.), running from 1970 "
    "to 2023 for all 40 countries. This variant tests whether periodic recalibration "
    "substantially improves fit and, by extension, whether the model\u2019s structural assumptions "
    "are sound.",
    size=12, space_after=12)

# ==============================================================
# 3. Results
# ==============================================================
add_heading_styled(doc, "3. OECD-wide validation results", level=1)

add_heading_styled(doc, "3.1 Overall fit", level=2)
add_para(doc,
    "Table 1 summarises model performance. The dynamic model achieves a median MAPE of 4.6% "
    "(mean 6.7%) over a 53-year horizon, with a mean final population ratio of 0.999 "
    "(SD = 0.189)\u2014indicating negligible systematic bias. The static model\u2019s fit degrades with "
    "projection horizon: from median 4.7% (base year 2000, 23-year horizon) to 7.3% "
    "(base year 1970, 50-year horizon), as expected when parameters are held fixed during "
    "periods of rapid demographic change.",
    size=12, space_after=6)

add_para(doc,
    "Across the 40 countries, 30 achieve dynamic MAPE below 10%, 20 below 5%, and 6 below 2%. "
    "The best-fitting countries\u2014France (0.4%), Costa Rica (0.9%), Finland (0.8%), Czechia (1.3%), "
    "Slovenia (1.4%), and Italy (1.4%)\u2014are those with relatively smooth demographic transitions "
    "and limited immigration shocks.",
    size=12, space_after=6)

# Table 1
add_para(doc, "Table 1: Summary of model fit across 40 countries, by model variant "
         "and base year",
         bold=True, size=11, space_after=4)

table = doc.add_table(rows=6, cols=6)
table.style = 'Light Shading Accent 1'
headers = ['Model variant', 'Horizon (yrs)', 'N', 'MAPE mean (%)',
           'MAPE median (%)', 'Final ratio (mean\u00b1SD)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data_rows = [
    ['Static (1970)', '50', '40', '12.4', '7.3', '1.272 \u00b1 0.481'],
    ['Static (1980)', '43', '40', '9.6', '7.7', '1.023 \u00b1 0.288'],
    ['Static (1990)', '33', '40', '7.8', '6.5', '0.953 \u00b1 0.198'],
    ['Static (2000)', '23', '40', '5.1', '4.7', '0.914 \u00b1 0.101'],
    ['Dynamic (10-yr)', '53', '40', '6.7', '4.6', '0.999 \u00b1 0.189'],
]
for i, row_data in enumerate(data_rows):
    for j, val in enumerate(row_data):
        table.rows[i + 1].cells[j].text = val
add_para(doc, "", size=6, space_after=6)

add_heading_styled(doc, "3.2 Sources of misfit", level=2)
add_para(doc,
    "Countries with MAPE exceeding 10% share common characteristics. Immigration-driven growth "
    "explains misfit in Australia (13.5%), Canada (12.2%), Switzerland (7.2%), Luxembourg (21.5%), "
    "and Israel (13.9%)\u2014all countries where net migration substantially augmented population "
    "beyond what natural increase alone would produce. Our model deliberately excludes migration "
    "to isolate the quantum\u2013tempo\u2013survival mechanism; the residual misfit thus quantifies the "
    "migration component.",
    size=12, space_after=8)

add_para(doc,
    "Rapid fertility transition explains the remaining outliers: Mexico (23.3%), T\u00fcrkiye (17.0%), "
    "China (15.6%), and Colombia (13.1%) experienced TFR declines of 3\u20135 children per woman over "
    "the study period. The dynamic model partially captures this through decadal updates, but "
    "within-decade changes remain unaccounted for.",
    size=12, space_after=8)

add_para(doc,
    "Republic of Korea (11.9%) combines both effects: extreme fertility decline (from TFR 4.5 "
    "in 1970 to 0.7 in 2023) plus recent immigration. Lithuania (12.0%) and Latvia (8.5%) "
    "reflect emigration-driven population loss following EU accession.",
    size=12, space_after=12)

# Figures 1\u20135
add_figure(doc, os.path.join(FIG_DIR, 'fig1_showcase.png'),
    "Figure 1: Model versus observed population trajectories for six representative "
    "countries, 1970\u20132023",
    note="Note: Dynamic model (blue dashed) updates parameters decadally; static model "
    "(red dotted) uses 1970 base-year parameters; black solid = UN WPP 2024 observed.",
    width=6.0)

add_figure(doc, os.path.join(FIG_DIR, 'fig2_all_countries.png'),
    "Figure 2: Model validation across all 40 countries",
    note="Note: Dynamic model MAPE shown in upper-right corner of each panel. "
    "Countries sorted alphabetically.",
    width=6.5)

add_figure(doc, os.path.join(FIG_DIR, 'fig3_heatmap.png'),
    "Figure 3: Static model MAPE (%) by country and base year",
    note="Note: Greener cells indicate better fit; redder cells indicate poorer fit. "
    "Scale capped at 30%.",
    width=5.0)

add_figure(doc, os.path.join(FIG_DIR, 'fig4_comparison.png'),
    "Figure 4: Static versus dynamic model comparison",
    note="Note: Left panel: MAPE by country. Right panel: final population ratio "
    "(model/observed in 2023). The dynamic model consistently outperforms the static variant.",
    width=6.0)

add_figure(doc, os.path.join(FIG_DIR, 'fig5_bias.png'),
    "Figure 5: Model bias analysis using base year 2000",
    note="Note: (a) Fit versus TFR; (b) fit versus life expectancy; (c) bias versus MAC. "
    "No systematic relationship is observed, suggesting model performance is robust "
    "across demographic contexts.",
    width=6.0)

# ==============================================================
# 4. Discussion
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "4. The tempo effect as a policy lever: Controlling the speed "
                   "of social adaptation", level=1)
add_para(doc,
    "The validation results confirm that a model with just four parameters\u2014TFR, life expectancy, "
    "MAC, and fertility schedule width\u2014can reproduce observed population trajectories with "
    "median error under 5%. This parsimony makes transparent the distinct contributions of "
    "quantum (TFR), survival (e\u2080), and tempo (MAC) to population size.",
    size=12, space_after=12)

add_para(doc,
    "The tempo channel operates through generational overlap. When the mean age at childbearing "
    "is 25, approximately four generations (0, 25, 50, 75) are simultaneously alive. When MAC "
    "rises to 30, generational spacing widens to approximately 3.3 overlapping generations "
    "(0, 30, 60, 90), reducing the simultaneously living population by roughly one-sixth\u2014even "
    "if each woman bears exactly the same number of children. Bongaarts and Feeney (2006) "
    "extended their original quantum\u2013tempo framework to show that tempo distortions affect "
    "not only fertility indicators but all life-cycle events, reinforcing that timing shifts "
    "have population-level consequences far beyond period rate adjustments.",
    size=12, space_after=12)

add_para(doc,
    "This mechanism has a crucial policy implication that extends beyond population size to the "
    "pace of demographic change. Consider two countries with identical TFR = 1.5 but MAC = 25 "
    "versus MAC = 33. The country with higher MAC experiences faster effective population "
    "decline per calendar year because fewer generations overlap. This acceleration compresses "
    "the time available for institutional adaptation\u2014pension system reform, healthcare "
    "infrastructure expansion, labour market restructuring. Lutz, Sanderson, and Scherbov "
    "(2001) projected the end of world population growth within this century; our analysis "
    "shows that tempo determines how rapidly that endpoint is approached.",
    size=12, space_after=12)

add_para(doc,
    "Viewed from this angle, tempo-sensitive policies do not merely affect how many people "
    "exist; they control the speed at which societies must adapt their social structures to "
    "demographic change. Policies that modestly reduce the age at first birth\u2014through "
    "affordable housing for young families, universal childcare, or restructured educational "
    "pathways that do not penalise early parenthood\u2014could slow the pace of population decline "
    "and buy time for institutional adjustment, even without raising TFR. Cross-national "
    "evidence suggests that family policies can influence fertility timing, though their "
    "effects on completed fertility are more modest (Gauthier 2007; Th\u00e9venon 2011). "
    "McDonald (2000) argues that gender equity in institutions is a prerequisite for "
    "sustained fertility; our framework suggests that such equity also operates through "
    "the tempo channel by enabling earlier childbearing without career penalties.",
    size=12, space_after=12)

add_para(doc,
    "This perspective reframes the policy problem. The question is not only \u2018how do we increase "
    "births?\u2019 but also \u2018how do we manage the speed of demographic transition?\u2019 The tempo effect "
    "provides a mechanism for the latter that has been largely overlooked in contemporary "
    "policy design. Lutz, Skirbekk, and Testa (2006) warned of a low-fertility trap in which "
    "declining population norms and structural constraints reinforce postponement; Bongaarts "
    "and Sobotka (2012) later showed that some European countries had begun to reverse "
    "postponement trends, with period TFR recovering as tempo distortions subsided. Our "
    "model suggests that even partial reversal of postponement trends can have substantial "
    "effects on simultaneously living population.",
    size=12, space_after=12)

add_para(doc,
    "Our 40-country validation across OECD countries, China, and the DRC demonstrates that "
    "this is not a theoretical curiosity but a quantitatively significant force that operates "
    "across diverse demographic contexts\u2014from post-transition Japan (MAC = 31.4, TFR = 1.2) "
    "to pre-transition DRC (MAC = 24.8, TFR = 6.1). Myrskyl\u00e4, Kohler, and Billari (2009) "
    "showed that advanced development can reverse fertility declines; our framework "
    "complements this by showing that even where quantum remains low, tempo adjustments "
    "alter the population trajectory. The model\u2019s ability to reproduce observed population "
    "trajectories without migration parameters further underscores the primacy of the "
    "quantum\u2013tempo\u2013survival triad in determining simultaneously living population.",
    size=12, space_after=12)

# ==============================================================
# 5. Limitations
# ==============================================================
add_heading_styled(doc, "5. Limitations", level=1)
add_para(doc,
    "Several limitations warrant acknowledgement. First, the model excludes migration, which "
    "is the dominant source of misfit for countries such as Australia, Canada, and Luxembourg. "
    "The exclusion is deliberate\u2014to isolate the natural demographic mechanism\u2014but limits "
    "direct applicability to countries with large net migration. Second, the normal fertility "
    "schedule is a simplification; actual ASFRs may be skewed or bimodal (Frejka and Sobotka "
    "2008 document considerable cross-national variation in fertility schedule shapes across "
    "Europe). Third, decadal parameter updates may miss rapid within-decade transitions (e.g., "
    "Korea\u2019s fertility collapse in the 2010s). Fourth, the Gompertz survival function fits "
    "adult mortality well but approximates infant and child mortality less precisely; national "
    "projections typically use more flexible mortality models such as the Lee\u2013Carter framework "
    "(Lee and Carter 1992). Despite these limitations, the model\u2019s parsimonious structure is "
    "a feature for policy communication: it makes the tempo\u2013quantum\u2013survival decomposition "
    "transparent and interpretable.",
    size=12, space_after=12)

# ==============================================================
# 6. Conclusion
# ==============================================================
add_heading_styled(doc, "6. Conclusion", level=1)
add_para(doc,
    "The tempo effect on simultaneously living population is well-established in demographic "
    "theory but forgotten in policy practice. Using a parsimonious model validated across 40 "
    "countries against UN WPP 2024 data, we show that birth timing exerts a quantitatively "
    "substantial influence on population size\u2014one that operates independently of fertility "
    "quantum. The policy implication extends beyond population size to the pace of demographic "
    "change: tempo-sensitive interventions can control the speed at which societies must adapt "
    "their institutions, offering a complementary lever to conventional pronatalist measures. "
    "We call for the systematic incorporation of tempo effects into demographic impact "
    "assessments and population policy design.",
    size=12, space_after=18)

# ==============================================================
# Data availability statement (required by DR)
# ==============================================================
add_heading_styled(doc, "Data availability statement", level=2)
add_para(doc,
    "All input data are drawn from the United Nations World Population Prospects 2024, which "
    "are publicly available at https://population.un.org/wpp/. The analytical code required "
    "to reproduce all results and figures is available from the authors upon request and will "
    "be deposited in a public repository upon acceptance.",
    size=12, space_after=18)

# ==============================================================
# References (alphabetical, author-date style per DR house style)
# ==============================================================
add_heading_styled(doc, "References", level=1)

refs = [
    'Bongaarts, J. and Feeney, G. (1998). On the quantum and tempo of fertility. '
    'Population and Development Review 24(2): 271\u2013291. '
    'doi:10.2307/2807974.',

    'Bongaarts, J. and Feeney, G. (2006). The quantum and tempo of life-cycle events. '
    'Vienna Yearbook of Population Research 4: 115\u2013151. '
    'doi:10.1553/populationyearbook2006s115.',

    'Bongaarts, J. and Sobotka, T. (2012). A demographic explanation for the recent rise '
    'in European fertility. Population and Development Review 38(1): 83\u2013120. '
    'doi:10.1111/j.1728-4457.2012.00473.x.',

    'Frejka, T. and Sobotka, T. (2008). Fertility in Europe: Diverse, delayed and below '
    'replacement. Demographic Research 19(3): 15\u201346. '
    'doi:10.4054/DemRes.2008.19.3.',

    'Gauthier, A.H. (2007). The impact of family policies on fertility in industrialized '
    'countries: A review of the literature. Population Research and Policy Review '
    '26(3): 323\u2013346. doi:10.1007/s11113-007-9033-x.',

    'Goldstein, J.R. and Kreyenfeld, M. (2011). Has East Germany overtaken West Germany? '
    'Recent trends in order-specific fertility. Population and Development Review '
    '37(3): 453\u2013472. doi:10.1111/j.1728-4457.2011.00430.x.',

    'Goldstein, J.R., Lutz, W., and Scherbov, S. (2003). Long-term population decline in '
    'Europe: The relative importance of tempo effects and generational length. Population '
    'and Development Review 29(4): 699\u2013707. doi:10.1111/j.1728-4457.2003.00699.x.',

    'Gompertz, B. (1825). On the nature of the function expressive of the law of human '
    'mortality, and on a new mode of determining the value of life contingencies. '
    'Philosophical Transactions of the Royal Society of London 115: 513\u2013583. '
    'doi:10.1098/rstl.1825.0026.',

    'Gonand, F. (2005). Assessing the robustness of demographic projections in OECD '
    'countries. OECD Economics Department Working Papers No. 464. Paris: OECD Publishing. '
    'doi:10.1787/010523837414.',

    'Keyfitz, N. (1971). On the momentum of population growth. Demography 8(1): 71\u201380. '
    'doi:10.2307/2060339.',

    'Kohler, H.-P., Billari, F.C., and Ortega, J.A. (2002). The emergence of lowest-low '
    'fertility in Europe during the 1990s. Population and Development Review 28(4): '
    '641\u2013680. doi:10.1111/j.1728-4457.2002.00641.x.',

    'Kohler, H.-P. and Ortega, J.A. (2002). Tempo-adjusted period parity progression '
    'measures, fertility postponement and completed cohort fertility. Demographic Research '
    '6(6): 91\u2013144. doi:10.4054/DemRes.2002.6.6.',

    'Lee, R.D. and Carter, L.R. (1992). Modeling and forecasting U.S. mortality. Journal '
    'of the American Statistical Association 87(419): 659\u2013671. '
    'doi:10.1080/01621459.1992.10475265.',

    'Lesthaeghe, R. (2010). The unfolding story of the Second Demographic Transition. '
    'Population and Development Review 36(2): 211\u2013251. '
    'doi:10.1111/j.1728-4457.2010.00328.x.',

    'Lutz, W., Sanderson, W., and Scherbov, S. (2001). The end of world population growth. '
    'Nature 412: 543\u2013545. doi:10.1038/35087589.',

    'Lutz, W., Skirbekk, V., and Testa, M.R. (2006). The low-fertility trap hypothesis: '
    'Forces that may lead to further postponement and fewer births in Europe. Vienna '
    'Yearbook of Population Research 4: 167\u2013192. '
    'doi:10.1553/populationyearbook2006s167.',

    'McDonald, P. (2000). Gender equity in theories of fertility transition. Population '
    'and Development Review 26(3): 427\u2013439. doi:10.1111/j.1728-4457.2000.00427.x.',

    'Myrskyl\u00e4, M., Kohler, H.-P., and Billari, F.C. (2009). Advances in development '
    'reverse fertility declines. Nature 460: 741\u2013743. doi:10.1038/nature08230.',

    'Preston, S.H., Heuveline, P., and Guillot, M. (2001). Demography: Measuring and '
    'modeling population processes. Oxford: Blackwell.',

    'Ryder, N.B. (1964). The process of demographic translation. Demography 1(1): 74\u201382. '
    'doi:10.2307/2060033.',

    'Sobotka, T. (2004). Is lowest-low fertility in Europe explained by the postponement '
    'of childbearing? Population and Development Review 30(2): 195\u2013220. '
    'doi:10.1111/j.1728-4457.2004.010_1.x.',

    'Stevens, G.A., Alkema, L., Black, R.E., Boerma, J.T., Collins, G.S., Ezzati, M., '
    'Grove, J.T., Hogan, D.R., Hosseinpoor, A.R., Lawn, J.E., Marusi\u0107, A., Mathers, '
    'C.D., Murray, C.J.L., Stanton, C., Tell, G.S., and Wardlaw, T.M. (2016). Guidelines '
    'for Accurate and Transparent Health Estimates Reporting: The GATHER statement. The '
    'Lancet 388(10062): e19\u2013e23. doi:10.1016/S0140-6736(16)30388-9.',

    'Th\u00e9venon, O. (2011). Family policies in OECD countries: A comparative analysis. '
    'Population and Development Review 37(1): 57\u201387. '
    'doi:10.1111/j.1728-4457.2011.00390.x.',

    'United Nations, Department of Economic and Social Affairs, Population Division (2024). '
    'World Population Prospects 2024. New York: United Nations. '
    'https://population.un.org/wpp/.',

    'Witte, J.C. and Wagner, G.G. (1995). Declining fertility in East Germany after '
    'unification: A demographic response to socioeconomic change. Population and '
    'Development Review 21(2): 387\u2013397. doi:10.2307/2137500.',
]
for r in refs:
    add_para(doc, r, size=11, space_after=6)

# ==============================================================
# Appendix A: GATHER Compliance Statement
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "Appendix A: GATHER compliance statement", level=2)
add_para(doc,
    "This study reports population estimates and follows the Guidelines for Accurate and "
    "Transparent Health Estimates Reporting (GATHER; Stevens et al. 2016). Key items:",
    size=11, space_after=6)
add_para(doc,
    "\u2022 Items 1\u20133 (Objectives, methods, populations): Described in Sections 1\u20132.\n"
    "\u2022 Items 4\u20137 (Data inputs): All input data from UN WPP 2024, publicly available at "
    "population.un.org/wpp. No primary data collection.\n"
    "\u2022 Items 8\u201310 (Data adjustments): Initial population age structures interpolated from "
    "5-year to single-year groups by uniform distribution within each group.\n"
    "\u2022 Items 11\u201313 (Modelling): Gompertz survival, normal fertility schedule, endogenous "
    "renewal described in Section 2.1. Four parameters per period (TFR, e\u2080, MAC, \u03c3).\n"
    "\u2022 Items 14\u201316 (Uncertainty, results): MAPE and final ratio reported as fit metrics. "
    "No formal uncertainty intervals; model is deterministic.\n"
    "\u2022 Items 17\u201318 (Interpretation, reproducibility): Code and data sources documented. "
    "Analytical code available upon request.",
    size=11, space_after=12)

# ==============================================================
# Appendix B: National Projection Methods Comparison
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "Appendix B: National population projection methods and "
                   "assumptions across OECD countries", level=2)

add_para(doc,
    "This appendix summarises the official population projection methodologies and key "
    "assumptions used by national statistical offices in OECD countries and the two additional "
    "countries (China, DRC) included in our analysis. All countries employ variants of the "
    "cohort-component method (Preston, Heuveline, and Guillot 2001; Gonand 2005), but differ "
    "substantially in their treatment of fertility timing, mortality improvement models, "
    "migration assumptions, and scenario structures. These differences contextualise our "
    "model\u2019s deliberate simplification to four parameters.",
    size=11, space_after=12)

# Table B1
add_para(doc, "Table B1: Summary of official population projection methods by "
         "country/organisation",
         bold=True, size=11, space_after=4)

tbl = doc.add_table(rows=16, cols=5)
tbl.style = 'Light Shading Accent 1'
hdr = ['Country / Organisation', 'Method', 'Fertility assumption',
       'Mortality assumption', 'Migration treatment']
for i, h in enumerate(hdr):
    tbl.rows[0].cells[i].text = h

rows_data = [
    ['UN WPP 2024\n(All countries)', 'Cohort-component;\nprobabilistic (Bayesian)',
     'Bayesian hierarchical model;\nTFR trajectories with uncertainty',
     'Lee\u2013Carter variant with\ncountry-specific drift',
     'Net migration assumed;\nconverges to long-run average'],
    ['Japan (IPSS)', 'Cohort-component;\ndeterministic with\n3 fertility \u00d7 3 mortality variants',
     'Cohort fertility model;\nMedium TFR = 1.20 (2070);\nMAC = 32.8',
     'Lee\u2013Carter model;\ne\u2080 = 85.9 (M) / 91.8 (F) by 2070',
     'Net migration by age/sex;\nbased on recent trends;\n~163k/year'],
    ['USA (Census Bureau)', 'Cohort-component;\nmain + 3 migration\nvariants',
     'Race/ethnicity-specific ASFRs;\nTFR converges ~1.75\nby 2060',
     'Cause-of-death model;\nLee\u2013Carter for residual;\ne\u2080 \u2248 83.9 by 2100',
     'Main variable;\n4 scenarios from zero\nto high immigration;\n~1.1M/yr main'],
    ['Germany (Destatis)', 'Cohort-component;\n27 variants (3\u00d73\u00d73)',
     '3 variants: TFR 1.29\u20131.65;\nMAC \u2248 31.7\u201332.1 by 2070',
     '3 variants: e\u2080 82.6\u201386.4 (M)\n85.9\u201389.3 (F) by 2070',
     '3 net migration levels:\n150k / 250k / 350k per year'],
    ['UK (ONS)', 'Cohort-component;\nprincipal + 9 variants',
     'ASFRs by age; principal\nTFR \u2248 1.59 long-term;\nhigh/low variants',
     'Mortality improvement\nmodel (age-period-cohort);\ne\u2080 \u2248 83.9 (M) / 86.3 (F)',
     'Long-term net migration\n\u2248 315k principal;\nvariants: 126k\u2013515k'],
    ['France (INSEE)', 'Cohort-component;\ncentral + 3 variants\nper component',
     'TFR \u2248 1.80 central;\nhigh 2.10, low 1.60',
     'Trend extrapolation;\ne\u2080 \u2248 87.5 (M) / 90.0 (F)\nby 2070',
     'Net migration +70k/yr\ncentral; variants \u00b1'],
    ['Korea (KOSTAT)', 'Cohort-component;\n3 scenarios (low/medium/high)',
     'Cohort model; medium\nTFR rises to 1.08 by 2040;\nMAC continues rising',
     'Lee\u2013Carter;\ne\u2080 = 88.0 (M) / 91.4 (F)\nby 2072',
     'Net migration by\nnationality; \u224860\u2013100k/yr\nnet inflow'],
    ['Italy (ISTAT)', 'Cohort-component;\nmedian + 4 scenarios',
     'TFR \u2248 1.40 median;\nrange 1.20\u20131.60',
     'Lee\u2013Carter;\ne\u2080 \u2248 85.8 (M) / 89.2 (F)',
     'Net migration \u2248 +150\u2013230k/yr;\nage-sex specific'],
    ['Australia (ABS)', 'Cohort-component;\n3 series (A/B/C)',
     'TFR 1.55\u20131.85;\nSeries B: 1.62',
     'Mortality improvement\nrates extrapolated;\ne\u2080 \u2248 87 (M) / 89 (F)',
     'High reliance on NOM;\nSeries B: ~235k/yr;\nkey driver of growth'],
    ['Canada (StatCan)', 'Cohort-component;\nmicrosimulation\n(Demosim) + cohort',
     'TFR 1.40\u20131.60;\nmedium 1.49',
     'Lee\u2013Carter variant;\ne\u2080 \u2248 86 (M) / 89 (F)',
     'Net migration\n~400\u2013500k/yr;\nprimary growth driver'],
    ['Eurostat\n(EU Member States)', 'Cohort-component;\nconvergence model\nacross EU',
     'Partial convergence of\nTFR across member states;\ncountry-specific ASFRs',
     'Convergence of mortality\nimprovement rates;\nage-period model',
     'Convergence toward\nlong-run net migration;\ncountry-specific paths'],
    ['China (NBS)', 'Cohort-component\n(not regularly\npublished officially)',
     'TFR officially reported\nat 1.0\u20131.2 (2022\u201323);\nUN WPP assumes recovery',
     'Model life table;\ne\u2080 \u2248 78.6 (2023)',
     'Low international\nmigration; internal\nmigration not in\nnational projection'],
    ['DRC (No national\noffice projection)', 'Relies on UN WPP;\nno independent\nnational projection',
     'TFR \u2248 6.1 (2023);\ngradual decline assumed\nin UN model',
     'Model life table;\ne\u2080 \u2248 60.7 (2023)',
     'Low net migration;\nrefugee flows\nnot systematically\nmodelled'],
    ['Mexico (CONAPO)', 'Cohort-component;\n3 variants',
     'TFR declining toward\n1.7 by 2050;\nmedium variant',
     'Trend extrapolation;\ne\u2080 \u2248 79 (M) / 83 (F)',
     'Net emigration\nturning to near-zero;\n\u224850k net by 2050'],
    ['T\u00fcrkiye (TurkStat)', 'Cohort-component;\n3 scenarios',
     'TFR declining from\n1.51 (2023) to ~1.60\nlong-term',
     'Improvement model;\ne\u2080 \u2248 80 (M) / 84 (F)',
     'Net immigration\n~200\u2013300k/yr;\nrefugee component'],
]
for i, rd in enumerate(rows_data):
    for j, val in enumerate(rd):
        tbl.rows[i + 1].cells[j].text = val

add_para(doc, "", size=6, space_after=8)

add_para(doc,
    "Source: UN DESA (2024), IPSS Japan (2023), US Census Bureau (2023), Destatis Germany "
    "(2025), ONS UK (2025), INSEE France (2021), KOSTAT Korea (2023), ISTAT Italy (2023), "
    "ABS Australia (2018), Statistics Canada (2024), Eurostat (2024), CONAPO Mexico (2018), "
    "TurkStat (2023). National Bureau of Statistics of China does not publish regular "
    "subnational-to-national projection documents in the same format as OECD members; values "
    "drawn from UN WPP 2024 and published census analyses. DRC has no independent national "
    "statistical projection.",
    italic=True, size=9, space_after=12)

add_heading_styled(doc, "B.1 Common features and key differences", level=3)
add_para(doc,
    "All national projection systems share the cohort-component method as their foundational "
    "structure, iteratively ageing a population by single or five-year age groups using "
    "fertility, mortality, and migration assumptions. Key differences relevant to our model "
    "include:",
    size=11, space_after=8)
add_para(doc,
    "\u2022 Treatment of fertility timing: Most national projections specify full age-specific "
    "fertility rate schedules (ASFRs) rather than parameterising fertility by MAC and \u03c3 as we "
    "do. Japan (IPSS) and Korea (KOSTAT) use cohort fertility models that explicitly track "
    "timing shifts. Our normal-distribution simplification captures the central tendency but "
    "not schedule shape.\n\n"
    "\u2022 Mortality models: National offices typically use Lee\u2013Carter (Lee and Carter 1992) or "
    "its extensions (Li\u2013Lee for coherent multi-population forecasting, used by Eurostat). Our "
    "Gompertz survival with a single calibrated parameter (a, with b fixed) is more "
    "parsimonious but less flexible for age-specific mortality patterns, particularly at "
    "young ages.\n\n"
    "\u2022 Migration: This is the component most variable across countries and the one our model "
    "deliberately excludes. For immigration-dependent countries (Australia, Canada, Luxembourg, "
    "Israel), migration assumptions dominate long-term projections. Our model\u2019s misfit in these "
    "countries (MAPE 12\u201322%) directly reflects the omitted migration component.\n\n"
    "\u2022 Scenario structure: Countries range from 3 variants (Korea, T\u00fcrkiye) to 27 (Germany). "
    "The UN WPP uses Bayesian probabilistic projections providing full uncertainty "
    "distributions. Our deterministic model offers a single trajectory per parameter set, "
    "trading uncertainty quantification for transparency of the tempo\u2013quantum decomposition.\n\n"
    "\u2022 Tempo treatment: Notably, none of the national projection systems explicitly decomposes "
    "population change into quantum and tempo components. Fertility timing enters implicitly "
    "through ASFRs, but the independent contribution of MAC to simultaneously living population "
    "is not isolated. This gap motivates our study.",
    size=11, space_after=12)

add_heading_styled(doc, "B.2 Implications for model comparison", level=3)
add_para(doc,
    "Our model is not designed to replace national projection systems but to complement them "
    "by making the tempo\u2013quantum\u2013survival decomposition explicit. The table above demonstrates "
    "that even the most sophisticated national systems share the same fundamental structure "
    "(cohort-component), differ primarily in their parameter estimation methods and scenario "
    "structures, and uniformly lack explicit tempo decomposition. Our 4-parameter model "
    "achieves median MAPE of 4.6% (dynamic) to 4.7% (static, base 2000) against these same "
    "populations\u2014performance that is sufficient to establish the quantitative significance of "
    "the tempo channel, even though it cannot match the precision of full-parameterisation "
    "national models that include migration.",
    size=11, space_after=12)

# ==============================================================
# Appendix C: Natural Experiments
# ==============================================================
doc.add_page_break()
add_heading_styled(doc, "Appendix C: Natural experiments \u2014 Political and border "
                   "changes as exogenous shocks", level=2)

add_para(doc,
    "Our model deliberately excludes migration. This design choice enables isolation of the "
    "pure quantum\u2013tempo\u2013survival mechanism, but raises an important question: how does the "
    "model perform when large-scale population redistribution occurs as a result of political "
    "events? Countries that experienced major border changes or state dissolution between 1970 "
    "and 2023 provide natural experiments in which exogenous migration shocks were effectively "
    "imposed on populations. In this appendix, we analyse five such cases to assess the "
    "robustness and limitations of the endogenous renewal framework.",
    size=11, space_after=12)

# --- C.1 Germany ---
add_heading_styled(doc, "C.1 Germany: Reunification as a migration shock (1990)", level=3)
add_para(doc,
    "German reunification on 3 October 1990 merged two populations that had evolved under "
    "sharply different demographic regimes for 41 years. East Germany (German Democratic "
    "Republic, GDR) had lower life expectancy (~74.5 years versus ~76.0 in the West in 1990), "
    "earlier childbearing (MAC \u2248 25.1 vs. 28.3), and higher but declining TFR (1.52 vs. "
    "1.45). The immediate aftermath of reunification saw massive East-to-West migration "
    "(~1.9 million between 1989 and 1992) and a dramatic fertility collapse in the East "
    "(TFR fell to 0.77 in 1994; Witte and Wagner 1995). Goldstein and Kreyenfeld (2011) "
    "later documented the gradual convergence of East German fertility behaviour toward "
    "Western patterns, including the adoption of later childbearing.",
    size=11, space_after=6)

add_para(doc,
    "We model East and West Germany separately using their distinct demographic parameters, "
    "run both forward from 1970, and construct a synthetic combined trajectory by summing "
    "the two modelled populations. This synthetic trajectory represents the counterfactual: "
    "what would Germany\u2019s population have looked like had the two regions remained "
    "demographically separate (i.e., with no cross-border migration)? Comparing this with "
    "observed unified Germany reveals the cumulative impact of reunification-related migration "
    "and integration effects.",
    size=11, space_after=6)

add_para(doc, "Table C1: Germany reunification: synthetic East+West versus observed "
         "unified trajectory",
         bold=True, size=11, space_after=4)

tbl_c1 = doc.add_table(rows=6, cols=4)
tbl_c1.style = 'Light Shading Accent 1'
for i, h in enumerate(['Year', 'Synthetic E+W (M)', 'Observed (M)', 'Deviation (%)']):
    tbl_c1.rows[0].cells[i].text = h
c1_data = [
    ['1990', '76.5', '79.4', '\u22123.6'],
    ['2000', '75.3', '82.2', '\u22128.4'],
    ['2010', '73.7', '81.8', '\u22129.9'],
    ['2020', '71.0', '83.2', '\u221214.6'],
    ['2023', '70.1', '83.3', '\u221215.8'],
]
for i, rd in enumerate(c1_data):
    for j, val in enumerate(rd):
        tbl_c1.rows[i + 1].cells[j].text = val
add_para(doc, "", size=6, space_after=6)

add_para(doc,
    "The synthetic trajectory underestimates observed population by 3.6% at reunification "
    "(1990), growing to 15.8% by 2023. This widening gap reflects three compounding processes "
    "absent from the closed model: (a) net immigration to unified Germany averaging "
    "~300,000\u2013400,000 persons per year; (b) internal East-to-West migration altering regional "
    "demographic dynamics; and (c) convergence of East German fertility and mortality toward "
    "Western levels, which our separate-regime model does not capture post-1990. The overall "
    "MAPE of 6.4% for the synthetic East+West model confirms that Germany\u2019s relatively poor "
    "fit in the main analysis (Table 1) is attributable primarily to reunification\u2019s migration "
    "effects rather than to structural model failure.",
    size=11, space_after=6)

add_figure(doc, os.path.join(FIG_DIR, 'fig_germany_reunification.png'),
    "Figure A-1: Germany reunification analysis",
    note="Note: (a) Modelled East and West Germany populations with synthetic combined "
    "trajectory versus observed unified Germany. (b) Percentage deviation of synthetic "
    "and unified models from observed trajectory. The widening gap after 1990 quantifies "
    "the cumulative migration and integration effects of reunification.",
    width=6.5)

# --- C.2 Czechoslovakia ---
add_heading_styled(doc, "C.2 Czechoslovakia: The Velvet Divorce (1993)", level=3)
add_para(doc,
    "The peaceful dissolution of Czechoslovakia on 1 January 1993 created two independent "
    "states\u2014Czechia and Slovakia\u2014with relatively limited cross-border migration. This case "
    "represents a relatively clean natural experiment: a state partition with minimal population "
    "redistribution. Our model achieves MAPE of 6.3% for Czechia and 9.9% for Slovakia over "
    "the full 1970\u20132023 period. Slovakia\u2019s higher error reflects emigration to Czechia and "
    "Western Europe following EU accession (2004), which the closed model cannot capture.",
    size=11, space_after=12)

# --- C.3 Yugoslavia ---
add_heading_styled(doc, "C.3 Yugoslavia: Dissolution and conflict (1991\u20132001)", level=3)
add_para(doc,
    "The breakup of Yugoslavia involved armed conflict, ethnic cleansing, and massive refugee "
    "flows affecting all successor states. This represents the most extreme migration shock in "
    "our sample. Model performance varies substantially across successors: Croatia (MAPE 4.1%) "
    "and North Macedonia (6.4%) show reasonable fit, reflecting relatively stable post-conflict "
    "demographics. Bosnia and Herzegovina (8.1%) and Slovenia (12.2%) show larger errors\u2014"
    "Bosnia due to war-related population loss and displacement, Slovenia due to "
    "immigration-driven growth as a small EU member state. Serbia (7.1%) and Montenegro "
    "(8.1%) fall in between. The range of model performance across Yugoslav successor states "
    "illustrates how conflict-driven migration creates heterogeneous deviations from the "
    "endogenous renewal baseline.",
    size=11, space_after=12)

# --- C.4 Baltic States ---
add_heading_styled(doc, "C.4 Baltic States: USSR dissolution and emigration (1991)",
                   level=3)
add_para(doc,
    "Estonia, Latvia, and Lithuania gained independence in 1991 from the Soviet Union, followed "
    "by significant emigration of ethnic Russians and outmigration to Western Europe (especially "
    "after EU accession in 2004). Model MAPE ranges from 4.8% (Estonia) to 7.1% (Lithuania), "
    "reflecting the persistent emigration that the closed model does not account for. These "
    "cases demonstrate that even moderate but sustained net emigration (\u22480.5\u20131.0% of "
    "population annually) accumulates over three decades to produce substantial "
    "model-observation divergence.",
    size=11, space_after=12)

# --- C.5 Ethiopia/Eritrea ---
add_heading_styled(doc, "C.5 Ethiopia and Eritrea: Separation (1993)", level=3)
add_para(doc,
    "Eritrean independence (1993) separated two populations in the midst of high-fertility "
    "demographic transition. Model performance differs markedly: Ethiopia (MAPE 16.5%) shows "
    "substantial overprojection, reflecting within-decade fertility decline faster than "
    "captured by decadal parameter updates. Eritrea (MAPE 37.8%) shows the largest error in "
    "our natural experiments sample, driven by prolonged military conscription, "
    "conflict-related emigration, and highly uncertain baseline demographic data. These "
    "cases highlight the model\u2019s limitations in conflict-affected, data-sparse settings "
    "where rapid demographic change co-occurs with large-scale displacement.",
    size=11, space_after=6)

add_figure(doc, os.path.join(FIG_DIR, 'fig_natural_experiments_summary.png'),
    "Figure A-2: Natural experiments summary",
    note="Note: Population trajectories for five cases of major political/border change. "
    "Dashed green lines show synthetic sums of successor-state models; solid black lines "
    "show observed (pre-split) trajectories; coloured lines show individual successor "
    "states. Red vertical lines mark the year of political change.",
    width=6.5)

# --- C.6 Synthesis ---
add_heading_styled(doc, "C.6 Synthesis: What natural experiments reveal", level=3)

add_para(doc, "Table C2: Model performance across all natural experiment cases",
         bold=True, size=11, space_after=4)

tbl_c2 = doc.add_table(rows=15, cols=4)
tbl_c2.style = 'Light Shading Accent 1'
for i, h in enumerate(['Country', 'Event (year)', 'MAPE (%)',
                        'Primary source of misfit']):
    tbl_c2.rows[0].cells[i].text = h
c2_data = [
    ['Germany (synthetic E+W)', 'Reunification (1990)', '6.4',
     'Immigration + internal migration'],
    ['Czechia', 'Velvet Divorce (1993)', '6.3',
     'Post-EU emigration (moderate)'],
    ['Slovakia', 'Velvet Divorce (1993)', '9.9',
     'Emigration to EU/West'],
    ['Croatia', 'Yugoslav breakup (1991)', '4.1',
     'Post-conflict stabilisation'],
    ['Slovenia', 'Yugoslav breakup (1991)', '12.2',
     'Immigration (EU member)'],
    ['Bosnia & Herz.', 'Yugoslav breakup (1991)', '8.1',
     'War-related displacement'],
    ['Serbia', 'Yugoslav breakup (1991)', '7.1',
     'Refugee flows, emigration'],
    ['N. Macedonia', 'Yugoslav breakup (1991)', '6.4',
     'Modest migration effects'],
    ['Montenegro', 'Yugoslav breakup (1991)', '8.1',
     'Small state, volatile flows'],
    ['Estonia', 'USSR dissolution (1991)', '4.8',
     'Ethnic Russian emigration'],
    ['Latvia', 'USSR dissolution (1991)', '6.6',
     'Emigration (ethnic + EU)'],
    ['Lithuania', 'USSR dissolution (1991)', '7.1',
     'Sustained emigration'],
    ['Ethiopia', 'Eritrean indep. (1993)', '16.5',
     'Rapid fertility decline'],
    ['Eritrea', 'Eritrean indep. (1993)', '37.8',
     'Conflict, conscription, emigration'],
]
for i, rd in enumerate(c2_data):
    for j, val in enumerate(rd):
        tbl_c2.rows[i + 1].cells[j].text = val
add_para(doc, "", size=6, space_after=6)

add_para(doc,
    "These natural experiments yield three key insights for the endogenous renewal model:",
    size=11, space_after=4)
add_para(doc,
    "First, the model performs reasonably well (MAPE < 8%) even for countries that experienced "
    "major political upheaval, provided that post-event migration was moderate and demographic "
    "data are reliable. Croatia (4.1%), Estonia (4.8%), and Czechia (6.3%) all fall within "
    "this range, demonstrating that the quantum\u2013tempo\u2013survival mechanism captures the bulk "
    "of population dynamics even in contexts of political discontinuity.\n\n"
    "Second, the magnitude of model-observation divergence provides a direct estimate of the "
    "migration component. Germany\u2019s 15.8% synthetic-observed gap by 2023 implies that "
    "immigration added approximately 13.5 million persons\u2014equivalent to one-sixth of the "
    "total population\u2014beyond what natural increase alone would have produced. This "
    "quantification is only possible because the closed model isolates natural demographic "
    "dynamics.\n\n"
    "Third, the model\u2019s limitations are most acute in conflict-affected, data-sparse settings "
    "(Eritrea: 37.8%) and where sustained emigration removes a large fraction of the "
    "population (Slovenia: 12.2%, Slovakia: 9.9%). These cases identify the boundary "
    "conditions for the endogenous renewal framework and reinforce the importance of "
    "migration modelling for countries experiencing sustained population flows.",
    size=11, space_after=12)

add_para(doc,
    "The Germany case is particularly instructive for interpreting the main analysis results. "
    "Germany\u2019s dynamic MAPE in the 40-country validation reflects not a failure of the "
    "quantum\u2013tempo\u2013survival framework but rather the quantitative footprint of reunification "
    "as a massive, exogenous migration event. The model\u2019s \u2018error\u2019 is, in fact, the signal: "
    "it measures the demographic impact of a political transformation that functioned as the "
    "equivalent of absorbing 1.9 million internal migrants in three years, followed by three "
    "decades of sustained international immigration. That our parsimonious model can separate "
    "this migration signal from the underlying demographic dynamics validates the "
    "decomposition approach.",
    size=11, space_after=12)

outpath = os.path.join(OUT_DIR, 'DemRes_Research_Article_EN.docx')
doc.save(outpath)
print(f'OK: {outpath}')
