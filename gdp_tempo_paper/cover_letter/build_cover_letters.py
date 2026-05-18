"""Build docx versions of the EN/JA cover letters from the markdown sources.

Single column, business-letter style: 12pt body for EN, MS Mincho for JA, 1.5
line spacing, 25mm margins. Italic and bold inline runs (Markdown ** and *) are
rendered. Plain hyphen lists ("- item") are rendered as bullet paragraphs.
"""
from __future__ import annotations

import os
import re

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_LINE_SPACING


HERE = os.path.dirname(os.path.abspath(__file__))


def _add_inline(paragraph, text: str) -> None:
    """Render a single paragraph's text with **bold** and *italic* runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        chunk = m.group(0)
        if chunk.startswith("**"):
            r = paragraph.add_run(chunk[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(chunk[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build(md_path: str, docx_path: str, font_name: str, font_size_pt: float) -> None:
    with open(md_path, encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f.readlines()]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    # East-Asian font hint so Japanese characters render with the same family
    rpr = style.element.rPr
    if rpr is not None:
        rfonts = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
        if rfonts is None:
            from docx.oxml.ns import qn
            rfonts = style.element.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        from docx.oxml.ns import qn
        rfonts.set(qn("w:eastAsia"), font_name)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)

    for raw in lines:
        if raw.strip() == "":
            p = doc.add_paragraph("")
        elif raw.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, raw[2:])
        else:
            p = doc.add_paragraph()
            _add_inline(p, raw)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    doc.save(docx_path)
    print(f"wrote {docx_path}")


def main() -> None:
    build(
        os.path.join(HERE, "cover_letter_en.md"),
        os.path.join(HERE, "cover_letter_en.docx"),
        font_name="Times New Roman",
        font_size_pt=12,
    )
    build(
        os.path.join(HERE, "cover_letter_ja.md"),
        os.path.join(HERE, "cover_letter_ja.docx"),
        font_name="MS Mincho",
        font_size_pt=11,
    )


if __name__ == "__main__":
    main()
